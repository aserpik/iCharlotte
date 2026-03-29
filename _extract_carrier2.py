"""Extract text from Carrier004.docx and write to _carrier_text.txt, then signal completion."""
from docx import Document
import os, sys

doc_path = r'Z:\Shared\Current Clients\3100 - PARSAC\094 - Haire v. City of Wildomar\STATUS\[Draft] Carrier004.docx'
out_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), '_carrier_text.txt')
done_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), '_extract_done.txt')

try:
    doc = Document(doc_path)
    lines = []
    for p in doc.paragraphs:
        if p.text.strip():
            lines.append(p.text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                lines.append(' | '.join(cells))

    with open(out_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(lines))

    with open(done_path, 'w') as f:
        f.write(f"OK: {len(lines)} lines")
except Exception as e:
    with open(done_path, 'w') as f:
        f.write(f"ERROR: {e}")
