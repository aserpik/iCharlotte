# -*- coding: utf-8 -*-
"""
Rebuild Defendant's Confidential Mediation Brief, weaving in facts from the
Rimkus "Report of Findings" (Amor Camatcho, P.E., 12/22/2023).

RULE: Every change/addition made to the original brief is rendered in BOLD.
Original brief text is rendered in regular weight.
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = "Times New Roman"
SIZE = 12

doc = Document()

# ---- base style ----
normal = doc.styles["Normal"]
normal.font.name = normal.font.name
normal.font.name = FONT
normal.font.size = Pt(SIZE)
rpr = normal.element.get_or_add_rPr()
rfonts = rpr.get_or_add_rFonts()
rfonts.set(qn("w:ascii"), FONT)
rfonts.set(qn("w:hAnsi"), FONT)

# margins
for s in doc.sections:
    s.top_margin = Inches(1)
    s.bottom_margin = Inches(1)
    s.left_margin = Inches(1)
    s.right_margin = Inches(1)


def _set_runs(p, runs):
    """runs: list of (text, bold) tuples."""
    for text, bold in runs:
        r = p.add_run(text)
        r.bold = bool(bold)
        r.font.name = FONT
        r.font.size = Pt(SIZE)
        rb = r._element.get_or_add_rPr().get_or_add_rFonts()
        rb.set(qn("w:ascii"), FONT)
        rb.set(qn("w:hAnsi"), FONT)
    return p


def body(runs, indent=True, space_after=0):
    """A double-spaced body paragraph. `runs` is a string or list of (text,bold)."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.space_after = Pt(space_after)
    if indent:
        pf.first_line_indent = Inches(0.5)
    if isinstance(runs, str):
        runs = [(runs, False)]
    _set_runs(p, runs)
    return p


def single(runs, indent=False, align=None, space_after=0, bold_all=False):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_after = Pt(space_after)
    if indent:
        pf.first_line_indent = Inches(0.5)
    if align is not None:
        p.alignment = align
    if isinstance(runs, str):
        runs = [(runs, bold_all)]
    _set_runs(p, runs)
    return p


def h1(numeral, title):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.space_before = Pt(6)
    tabs = pf.tab_stops
    tabs.add_tab_stop(Inches(0.5))
    r = p.add_run(numeral + "\t")
    r.bold = True
    r2 = p.add_run(title)
    r2.bold = True
    r2.underline = True
    for rr in (r, r2):
        rr.font.name = FONT
        rr.font.size = Pt(SIZE)
    return p


def h2(letter, title):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.left_indent = Inches(0.5)
    pf.first_line_indent = Inches(-0.0)
    tabs = pf.tab_stops
    tabs.add_tab_stop(Inches(1.0))
    r = p.add_run(letter + "\t")
    r.bold = True
    r2 = p.add_run(title)
    r2.bold = True
    r2.underline = True
    for rr in (r, r2):
        rr.font.name = FONT
        rr.font.size = Pt(SIZE)
    return p


def subhead(text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    r = p.add_run(text)
    r.bold = True
    r.underline = True
    r.font.name = FONT
    r.font.size = Pt(SIZE)
    return p


# ============================== CAPTION ==============================
single("BORDIN SEMMER LLP")
single("Joshua Bordin-Wosk, State Bar No. 241077")
single("jbordinwosk@bordinsemmer.com")
single("Andrei V. Serpik, State Bar No. 301260")
single("aserpik@bordinsemmer.com")
single("101 Continental Blvd., Suite 700")
single("El Segundo, CA 90245")
single("Telephone:\t(323) 457-2110")
single("Facsimile:\t(323) 457-2120")
single("Attorneys for Defendant, MICHELLE MATTHEWS", space_after=12)

single("SUPERIOR COURT OF THE STATE OF CALIFORNIA",
       align=WD_ALIGN_PARAGRAPH.CENTER, bold_all=True)
single("COUNTY OF LOS ANGELES", align=WD_ALIGN_PARAGRAPH.CENTER,
       bold_all=True, space_after=6)

# caption table
tbl = doc.add_table(rows=1, cols=2)
tbl.style = "Table Grid"
left = tbl.cell(0, 0)
right = tbl.cell(0, 1)
left.width = Inches(3.2)
right.width = Inches(3.0)

lp = left.paragraphs[0]
_set_runs(lp, [("RON DUDASH and MIMI DUDASH\n", False)])
_set_runs(left.add_paragraph(), [("\t\tPlaintiff,\n", False)])
_set_runs(left.add_paragraph(), [("\tv.\n", False)])
_set_runs(left.add_paragraph(), [("MICHELLE MATTHEWS; and DOES 1-10, inclusive,\n", False)])
_set_runs(left.add_paragraph(), [("\t\tDefendants.", False)])

rp = right.paragraphs[0]
_set_runs(rp, [("Case No.:  24TRCV00960", False)])
_set_runs(right.add_paragraph(), [("(Assigned for all purposes to Hon. Patricia A. Young, Dept. B)", False)])
right.add_paragraph()
_set_runs(right.add_paragraph(), [("DEFENDANT'S CONFIDENTIAL MEDIATION BRIEF", True)])
right.add_paragraph()
_set_runs(right.add_paragraph(), [("Complaint Filed:\tMarch 19, 2024", False)])
_set_runs(right.add_paragraph(), [("Trial Date:\t\tSeptember 21, 2026", False)])

doc.add_paragraph()

# ============================== I. INTRODUCTION ==============================
h1("I.", "INTRODUCTION")

body([
    ("This action arises from a property dispute between Plaintiffs Ron Dudash and "
     "Mimi Dudash (\"Plaintiffs\") and Defendant Michelle Matthews (\"Defendant\") "
     "concerning adjacent commercial properties in Gardena, California. Plaintiffs own "
     "the commercial building located at 2224 W. Rosecrans Avenue (the \"Dudash "
     "Property\"). Defendant owns the adjacent property located at 2222 W. Rosecrans "
     "Avenue (the \"Matthews Property\"). Plaintiffs allege that an unpaved section of "
     "the parking lot on the Matthews Property and the installation of a steel tubing "
     "bumper on the exterior wall of the Dudash Property caused rainwater to flood their "
     "building between January 18 and January 20, 2023. Based on these allegations, "
     "Plaintiffs assert claims for continuing nuisance and continuing trespass.", False),
])

body([
    ("Defendant disputes liability. As an initial matter, Plaintiffs' claims are barred "
     "by the three-year statute of limitations for injury to real property and there is a "
     "pending motion for summary judgment on this issue scheduled to be heard on "
     "_______. Furthermore, as to the steel bumper, Plaintiffs impliedly consented to "
     "the alleged trespass by actively maintaining, painting, and sealing the steel "
     "tubing bumper on an annual basis for twenty years without ever requesting its "
     "removal. Second, Plaintiffs will not be able to establish causation. Physical "
     "evidence and independent engineering analysis demonstrate that Defendant's "
     "property did not cause the water intrusion. ", False),
    ("Defendant's retained forensic engineer, Amor Camatcho, P.E. (California Engineer "
     "No. 40561) of Rimkus, inspected both properties on November 30, 2023 and, in his "
     "December 22, 2023 Report of Findings, concluded that the water intrusion and "
     "suspected fungal growth were caused by the lack of maintenance of the flashing or "
     "sealant along the interface of the steel tubing bumper and the building's exterior "
     "wall—and were not caused by inadequate slopes or surface water drainage from "
     "the Matthews Property parking lot. ", True),
    ("The parking lot on the Matthews Property slopes away from the Dudash Property, and "
     "the City of Gardena investigated the premises in response to Plaintiffs’ "
     "complaints and found no unpermitted grading on Defendant’s property. The water "
     "damage was instead self-inflicted by Plaintiffs' negligent application of "
     "do-it-yourself cement and sealants over their own building's engineered weep "
     "screeds, which trapped moisture inside the wall cavity and caused the resulting "
     "mold and rot.", False),
])

body([
    ("Defendant also challenges the nature and scope of Plaintiffs' claimed injuries and "
     "damages. Plaintiffs' demand for over $750,000 in lost wages and loss of future "
     "earning capacity is entirely speculative, as it relies on Plaintiffs’ "
     "voluntary retirement and lacks any supporting medical or psychological evidence of "
     "incapacitation. Plaintiffs' property damage claims are similarly unsupported by "
     "documentation and are contradicted by their own prior sworn statements. In a "
     "previous small claims lawsuit, Plaintiffs argued to the court that the exact same "
     "mold damage was caused by a former commercial tenant's clothing inventory, a "
     "shifting narrative that Plaintiff Ron Dudash admitted to during his deposition. "
     "Additionally, the objective tenancy timeline refutes Plaintiffs' claim for three "
     "months of lost rent, as a new commercial tenant occupied the space within "
     "approximately one month of the prior tenant vacating. Finally, as absentee "
     "commercial landlords who did not physically occupy the affected units, Plaintiffs "
     "are legally precluded from recovering general damages for emotional distress.", False),
])

body([
    ("Given the insurmountable legal and evidentiary hurdles facing both liability and "
     "causation, Defendant approaches this mediation in good faith to explore a "
     "reasonable resolution, but Plaintiffs must recognize the significant structural "
     "problems in their case.", False),
])

# ============================== II. STATEMENT OF FACTS ==============================
h1("II.", "STATEMENT OF FACTS")

subhead("Property History")
body([
    ("Both the Dudash Property and the Matthews Property were originally constructed in "
     "1959. Sometime between 1989 and 2001, a square steel tubing bumper was installed "
     "along the exterior east wall of the Dudash Property. ", False),
    ("The defense engineering inspection later measured this bumper as a nominal 5-inch "
     "square steel tube, 56 feet 9 inches long, bolted to the face of the wall and "
     "positioned approximately 25½ to 28 inches above the adjacent grade.", True),
])
body([
    ("In 2004, Plaintiffs purchased the Dudash Property. The steel bumper was already "
     "attached to the exterior wall at the time of their purchase. Over the subsequent "
     "two decades, Plaintiffs maintained the bumper themselves, regularly wire brushing, "
     "painting, and applying sealant behind the structure's fasteners.", False),
])
body([
    ("In March 2009, Defendant's mother, who owned the Matthews Property at the time, "
     "removed an approximate 20-by-20-foot section of asphalt in her parking lot.", False),
])
body([
    ("Defendant acquired the Matthews Property on or about ______. ", False),
    ("Ms. Matthews has stated that the unpaved area at the southwest corner of the "
     "parking lot had been unpaved for over 25 years and was intended to eventually "
     "become a garden, and that she did not install the steel tubing bumper on the "
     "Dudash Property and does not know who installed it or when.", True),
])

subhead("Defense Engineering Inspection")
body([
    ("On November 30, 2023, Defendant's retained forensic engineer, Amor Camatcho, "
     "P.E.—a registered professional engineer with more than 25 years of structural "
     "engineering experience who specializes in water intrusion investigations—"
     "conducted an on-site inspection of both properties, interviewed the City of "
     "Gardena building code enforcement officer, and reviewed the relevant documents and "
     "codes. His findings are set forth in a December 22, 2023 Report of Findings that "
     "was peer-reviewed by Paul D. Colman, P.E. These findings are discussed in the "
     "Liability section below.", True),
])

subhead("Alleged Water Intrusion Events")
body([
    ("In January 2012, Plaintiffs allegedly experienced their first water intrusion "
     "event through the east wall of their building.", False),
])
body([
    ("Between 2014 and 2023, Plaintiffs allegedly experienced three to five additional "
     "water intrusion events. In response to these recurring leaks, Plaintiffs applied "
     "do-it-yourself patches to the exterior wall using cement and black sealants. This "
     "application covered and effectively sealed the building's engineered weep screeds "
     "along the exterior wall. ", False),
    ("The defense engineer independently observed these conditions during his inspection, "
     "noting that portions of the weep screed had been sealed with a black sealant "
     "material or covered with a light-gauge metal material and then painted.", True),
])
body([
    ("Between 2019 and 2020, a 40-unit townhome complex was constructed directly behind "
     "both properties. As part of this new development, a pre-existing chain-link fence "
     "was replaced with a solid block wall, altering the local surface drainage patterns "
     "in the immediate vicinity.", False),
])
body([
    ("Between January 18 and January 20, 2023, heavy rainstorms allegedly caused "
     "flooding inside Plaintiff’s property.", False),
])
body([("[first demand]", False)])
body([
    ("Plaintiffs filed their initial Complaint in this action on March 19, 2024, "
     "followed by a First Amended Complaint on July 29, 2024, alleging claims for "
     "continuing nuisance and continuing trespass. Plaintiffs claim property damage, "
     "lost rent, and lost wages due to early retirement allegedly caused by the stress "
     "of the dispute.", False),
])

# ============================== III. PROCEDURAL STATUS ==============================
h1("III.", "PROCEDURAL STATUS")
body([
    ("Trial is scheduled for July 13, 2026. The depositions of Plaintiffs Ron Dudash "
     "and Mimi Dudash were taken on January 28, 2026. Defendant’s Motion for "
     "Summary Judgment is scheduled to be heard on __________.", False),
])

# ============================== IV. LIABILITY ==============================
h1("IV.", "LIABILITY")
body([
    ("Plaintiffs cannot establish liability for their nuisance and trespass claims at "
     "trial. The evidentiary record demonstrates that Plaintiffs' claims are barred by "
     "the statute of limitations, precluded by decades of implied consent, and fatally "
     "undermined by physical evidence proving the alleged water damage was "
     "self-inflicted.", False),
])

h2("A.", "Plaintiffs' Claims Are Barred By The Three-Year Statute Of Limitations For "
         "Permanent Nuisance And Trespass")
body([
    ("Under California law, a three-year statute of limitations applies to claims for "
     "injury to real property, which encompasses causes of action for both nuisance and "
     "trespass. (Code of Civil Procedure section 338(b).) The application of this "
     "limitation period depends on whether the alleged condition is classified as "
     "continuing or permanent. A permanent nuisance or trespass involves a solid "
     "structure or a completed landscape modification, such as a building foundation, a "
     "buried pipe, or settled land grading, which is intended to remain indefinitely. "
     "(Bookout v. State of California (2010) 186 Cal.App.4th 1478; Field-Escandon v. "
     "DeMann (1988) 204 Cal.App.3d 228.) When a nuisance or trespass is permanent, the "
     "three-year statute of limitations begins to run upon the creation of the condition "
     "or as soon as the physical damage to the property becomes appreciable.", False),
])
body([
    ("Plaintiffs base their lawsuit on two specific property conditions: (1) a steel "
     "tubing bumper attached to the east wall of the Dudash Property; and (2) a "
     "difference in land grading on the unpaved portion of the Matthews Property. Both "
     "of these conditions are quintessential permanent structures and completed "
     "landscape alterations that were created decades ago.", False),
])
body([
    ("Regarding the steel tubing, the evidence confirms this structure is a permanent "
     "fixture installed between 1989 and 2001:", False),
])
body([("[photo of tubing]", False)])
body([
    ("At his deposition, Plaintiff Ron Dudash testified that the steel bumper was "
     "already bolted to the exterior wall when Plaintiffs purchased the Dudash Property "
     "in 2004. A heavy, bolted steel structure that has remained in place for over "
     "twenty-five years is undoubtedly a permanent condition, entirely analogous to the "
     "permanent sewer pipe analyzed in Field-Escandon. ", False),
    ("The defense engineer's measurements confirm the permanence of this fixture: a "
     "nominal 5-inch square steel tube, 56 feet 9 inches long, bolted to the face of the "
     "wall, with a painted 2x12 wood member running approximately 46 feet directly "
     "beneath it. ", True),
    ("Because the structure was present and known to Plaintiffs upon their acquisition of "
     "the property in 2004, the three-year statute of limitations to demand its removal "
     "or claim trespass expired nearly two decades before this lawsuit was filed.", False),
])
body([
    ("Likewise, the grading difference between the two properties is a permanent "
     "condition. Plaintiff testified that Defendant's mother removed a section of asphalt "
     "on the Matthews Property in March 2009, creating the unpaved dirt area that "
     "Plaintiffs now allege causes a \"hydraulic effect.\" The removal of the asphalt "
     "was a completed, discrete event. Plaintiff was immediately aware of this condition "
     "and testified that he formally complained to the City of Gardena regarding the "
     "creation of this specific mud patch in March 2009.", False),
])
body([("[quote]", False)])
body([
    ("Most critically, Plaintiff admitted during his deposition that he experienced "
     "appreciable water intrusion into his building three to five times between 2012 and "
     "2023. He testified that muddy water entered through the east wall during these "
     "early incidents, yet he chose not to investigate the root cause, hire an expert, "
     "or take legal action because he felt the flooding was \"not as severe\" as the "
     "incident in 2023.", False),
])
body([("[quote]", False)])
body([
    ("Under Bookout, a plaintiff's decision to ignore appreciable property damage does "
     "not toll the statute of limitations until the damage becomes subjectively severe. "
     "Because the grading alteration was completed in 2009, and Plaintiffs had actual "
     "knowledge of both the condition and the resulting appreciable water intrusion by "
     "2012, the statute of limitations expired no later than 2015. Consequently, "
     "Plaintiffs' 2024 complaint is time-barred as a matter of law.", False),
])

h2("B.", "Plaintiffs Impliedly Consented To The Alleged Trespass And Nuisance Through "
         "Decades Of Active Maintenance")
body([
    ("To establish a claim for private nuisance, a plaintiff must prove that they did "
     "not consent to the defendant's conduct and that the condition constitutes an "
     "unreasonable interference with their property rights. Similarly, a lack of "
     "permission is a fundamental element of any trespass claim. A property owner who "
     "accepts, integrates, and actively maintains a condition on their property cannot "
     "subsequently claim that the condition is an unauthorized trespass or an actionable "
     "nuisance. (citation)", False),
])
body([
    ("Plaintiffs allege that the steel tubing bumper is an unpermitted structure that "
     "Defendant must be ordered to remove. First, Defendant catagorically denies the "
     "allegation that either she or her mother installed this steel tubing on "
     "Plaintiffs’ property. ", False),
    ("Ms. Matthews confirmed to the defense engineer that she did not have the tube "
     "installed and does not know who installed it or when, and the bumper existed as "
     "far back as she can remember. ", True),
    ("Indeed, the tubing had been present on Plaintiffs’ property when they "
     "initially purchased the building. For twenty years, Plaintiffs have treated the "
     "steel tubing not as an unwanted encroachment, but as a maintained fixture of their "
     "own commercial building. During his deposition, Plaintiff testified that he "
     "actively maintained the bumper since purchasing the property in 2004. He admitted "
     "to manually wire brushing, painting, and applying sealant behind the fasteners of "
     "the structure on an annual basis. ", False),
    ("Consistent with this testimony, the defense engineer observed that the interface "
     "between the top of the steel tube and the stucco wall had been re-sealed with a "
     "black sealant-type material.", True),
])
body([("[Quote]", False)])
body([
    ("Furthermore, Plaintiff testified that neither he nor his wife ever asked Defendant "
     "or her mother to remove the steel bumper at any point during their two decades of "
     "ownership. When asked why he never requested its removal, Plaintiff admitted that "
     "removing the bumper would be expensive and difficult due to its weight, and he "
     "simply did not want to spend the money.", False),
])
body([("[quote]", False)])
body([
    ("A plaintiff cannot actively paint, seal, and utilize a structure for twenty years "
     "to save money, only to abruptly recharacterize it as a hostile trespass when "
     "seeking a financial windfall in litigation. Plaintiffs' affirmative maintenance of "
     "the steel tubing constitutes implied consent and acquiescence to its presence, "
     "negating the essential elements of both trespass and nuisance.", False),
])

h2("C.", "Plaintiffs Cannot Establish Causation Due To Self-Inflicted Damage And "
         "Alternative Water Sources")
body([
    ("Liability for property damage requires Plaintiffs to prove that Defendant's "
     "conduct was a substantial factor in causing the harm. The physical evidence, "
     "expert analysis, and testimony from Plaintiffs' own tenants establish that "
     "Defendant's property is not the root cause of the water intrusion. Instead, the "
     "damage is the direct result of Plaintiffs' negligent building maintenance and "
     "independent structural failures.", False),
])
body([
    ("Plaintiffs' central theory is that unpermitted grading on the Matthews Property "
     "forces water into the Dudash Property. This theory is contradicted by independent "
     "engineering data. Defense engineering expert Amor Camatcho, P.E., conducted a site "
     "inspection and measured the slopes of the paving surface on the Matthews Property. "
     "Mr. Camatcho found that the parking lot actually slopes away from the Dudash "
     "Property, averaging a 0.3 percent downward grade toward the east. ", False),
    ("He further measured an average slope of 1.6 percent running downward toward the "
     "north along the edge of the Dudash building—confirming that surface water "
     "drains away from and along the building rather than into it, and that the parking "
     "lot provides adequate drainage. ", True),
    ("Furthermore, the City of Gardena comprehensively investigated Plaintiffs' 2023 "
     "nuisance complaint. Following a site visit and review of aerial photographs, the "
     "City Building Official concluded that no unpermitted grading had occurred on "
     "Defendant's property and declined to issue any violation notice. ", False),
    ("As the defense engineer confirmed through his interview of the City's code "
     "enforcement officer, the City's only concern with the unpaved area was the pooling "
     "of water in the dirt and potential mosquito habitation—not any moisture "
     "intrusion into the Dudash Property—and City of Gardena Violation Notice "
     "C23-0051 did not cite any ingress/egress maintenance violation against the "
     "Matthews Property.", True),
])
body([("[snippit]", False)])
body([
    ("The true cause of the water intrusion was discovered during Mr. Camatcho's "
     "inspection of the Dudash Property's exterior wall. The building's stucco wall is "
     "designed with a metal weep screed at the base, engineered to allow accumulated "
     "moisture to naturally drain out of the wall cavity. Mr. Camatcho observed that the "
     "area below the stucco weep screed had been deliberately sealed with a black "
     "sealant material and concrete. This improper sealing effectively trapped moisture "
     "inside the walls, preventing the building from breathing and directly causing the "
     "interior mold and rot. ", False),
    ("Critically, Mr. Camatcho determined that the height of the suspected fungal growth "
     "on the interior east wall—approximately two feet above the floor—was "
     "consistent with the height of the steel tube bumper on the exterior side of the "
     "wall, tying the moisture to the unmaintained bumper interface rather than to any "
     "ground-level surface water. He also documented that the weep screed conformed to "
     "the 2022 California Building Code (section 2512.1.2) before Plaintiffs sealed it.", True),
])
body([("[snippit]", False)])
body([
    ("At his deposition, Plaintiff admitted that he personally applied these temporary "
     "\"quick patches\" using cement and sealant between 2014 and 2023, and that he did "
     "so without consulting a licensed professional.", False),
])
body([("[quote]", False)])
body([
    ("Plaintiffs' own do-it-yourself repair efforts self-inflicted the water entrapment "
     "they now attempt to blame on Defendant.", False),
])
body([
    ("Notably, the defense engineer also documented conduct suggesting that Plaintiffs "
     "manufactured the very pooling condition they complain of. During the inspection, "
     "the unpaved area was damp with no standing water; a person then walked out of the "
     "Dudash tenant space, dumped a bucket of water onto the soil of the unpaved area, "
     "and walked back inside, creating standing water where none had existed. When the "
     "engineer accessed the interior of the east-end tenant unit, he found no suspected "
     "fungal growth and no moisture-related staining on the adjacent flooring.", True),
])
body([
    ("Moreover, the Dudash Property suffers from a well-documented history of water "
     "intrusion wholly unrelated to the Matthews Property. Former long-term commercial "
     "tenant Orville Cole confirmed that the building experienced severe roof leaks that "
     "required a total roof replacement in 2015. Mr. Cole also reported repeated "
     "instances of water bubbling up from beneath the floor slab and the west wall, "
     "which abuts a completely different property. Current commercial tenant Gina "
     "similarly reported severe water problems originating from beneath the floor slab "
     "on the west wall. Gina's independent plumber inspected the premises and confirmed "
     "the intrusion was caused by sub-slab plumbing failures, not lateral moisture from "
     "the east wall.", False),
])
body([
    ("Plaintiffs' ancillary allegations regarding the parking lot fare no better. "
     "Plaintiffs' own \"code violations\" dossier relied on inapplicable code "
     "provisions: it cited a residential code section (R703.7.2.1) that does not apply "
     "to this commercial building, and an accessibility provision (California Building "
     "Code 1109A.8.3) that governs only marked accessible parking spaces, none of which "
     "exist adjacent to the Dudash building. The defense engineer further confirmed that "
     "the parking lot's pre-cast concrete wheel stops conform to Gardena Municipal Code "
     "section 18.40.060.E, because the proper three-foot measurement runs from the face "
     "of the wall to the front of the wheel stop where a vehicle's wheels actually "
     "stop—a distance of 36 inches.", True),
])
body([
    ("Faced with a deteriorating building, Plaintiffs have developed a documented "
     "pattern of blaming third parties for their structural failures. In a recent small "
     "claims lawsuit involving Mr. Cole, Plaintiff vehemently argued to the court that "
     "the extensive mold in the unit was caused by the tenant's clothing inventory "
     "resting against the walls.", False),
])
body([("[snippit]", False)])
body([
    ("Plaintiff now brings the exact same mold damage to this Court, abandons his prior "
     "sworn position, and attempts to attribute the damage to Defendant. Given "
     "Plaintiffs' admission to sealing their own weep holes, the presence of severe "
     "sub-slab plumbing failures, and Plaintiffs' shifting, contradictory narratives "
     "regarding the source of the mold, Plaintiffs cannot meet their burden of proving "
     "that Defendant's grading was the substantial factor causing their harm.", False),
])

# ============================== V. DAMAGES ==============================
h1("V.", "DAMAGES")
body([
    ("Plaintiffs' damages claims lack objective evidentiary support and are contradicted "
     "by the physical timeline, medical absences, and Plaintiffs' own prior admissions. "
     "Should this case proceed to trial, the evidentiary record will demonstrate that "
     "Plaintiffs cannot meet their burden of proof regarding their highly speculative "
     "wage loss claims, contradictory property damage allegations, or unsupported "
     "demands for emotional distress.", False),
])

h2("A.", "Plaintiffs' Wage Loss And Loss Of Earning Capacity Claims Are Wholly "
         "Speculative And Factually Unsupported")
body([
    ("The most significant financial component of Plaintiffs' claimed damages is their "
     "demand for lost wages and loss of future earning capacity, which totals over "
     "$750,000. Plaintiff Ron Dudash claims $509,760 in past lost wages and $245,000 "
     "annually in future losses, alleging that the stress of the property dispute caused "
     "an inability to concentrate, forcing an early retirement from his position at "
     "Boeing at age 65. Plaintiff Mimi Dudash similarly claims $96,000 in past lost "
     "wages and $4,000 per month in future losses, alleging she can no longer perform "
     "her duties as a property manager. These claims are entirely speculative and lack "
     "any foundational medical or employment evidence.", False),
])
body([
    ("A claim for lost earning capacity based on emotional or psychological "
     "incapacitation requires objective medical evidence demonstrating that the "
     "plaintiff suffers from a condition that genuinely precludes them from working. "
     "Here, there is a complete absence of medical or psychological evidence to support "
     "Plaintiffs' claims. During his deposition, Plaintiff acknowledged that he has "
     "never sought treatment from a therapist, psychologist, or psychiatrist, nor has he "
     "taken any medication for his alleged emotional distress and inability to "
     "concentrate. He further conceded that his claim for lost earnings relies "
     "exclusively on his own self-serving testimony and that of his wife.", False),
])
body([
    ("Furthermore, the objective circumstances of Plaintiff's departure from Boeing "
     "contradict his claim of a forced medical retirement. Plaintiff voluntarily "
     "resigned from his employment with a standard two weeks' notice. He did so at the "
     "customary retirement age of 65. He admitted in his deposition that he never "
     "reported any distress or inability to perform his duties to his employer or "
     "supervisors prior to his departure. Finally, he acknowledged that he is currently "
     "receiving standard retirement benefits from his employer.", False),
])
body([
    ("It is wholly implausible that a property line dispute regarding an adjacent "
     "parking lot incapacitated two individuals from all gainful employment. Because "
     "Plaintiffs possess no medical records, psychological evaluations, or employment "
     "documentation linking their voluntary retirements to the conditions at the "
     "Matthews Property, their exorbitant wage loss claims will be excluded or rejected "
     "at trial as a matter of law.", False),
])

h2("B.", "The Property Damage Claims Are Contradicted By Plaintiffs' Own Prior Lawsuits "
         "And Shifting Narratives")
body([
    ("Plaintiffs seek $75,717 in property damage, alleging that water intrusion and mold "
     "remediation were necessitated exclusively by conditions on the Matthews Property. "
     "However, Plaintiffs' current damages theory is fatally undermined by their own "
     "prior sworn positions in collateral litigation, as well as the severe lack of "
     "documentary evidence supporting their financial demands.", False),
])
body([
    ("First, the physical receipts and invoices produced by Plaintiffs substantiate only "
     "a fraction of their claimed property damage. While demanding over $75,000, the "
     "itemized receipts provided thus far total approximately $28,300. More "
     "problematically, many of the documents comprising this $28,300 are handwritten "
     "receipts or internal spreadsheets created by Plaintiffs themselves, lacking the "
     "necessary foundation of independent contractor verification or proof of actual "
     "payment.", False),
])
body([
    ("Second, and most critically, Plaintiffs have advanced fundamentally contradictory "
     "narratives regarding the cause of the property damage. Faced with severe mold "
     "issues, Plaintiffs initially chose to blame their former commercial tenant, "
     "Orville Cole. When Plaintiffs wrongfully withheld Orville Cole's security deposit, "
     "the tenant filed a small claims lawsuit against them. During those proceedings, "
     "Plaintiff vehemently argued to the court that the extensive mold in the commercial "
     "unit was caused by the tenant's clothing inventory resting against the interior "
     "walls. Plaintiff produced photographs of the store's interior layout in an attempt "
     "to prove the tenant was at fault for the exact same mold damage he now brings "
     "before this Court.", False),
])
body([
    ("During his deposition in the present action, Plaintiff explicitly admitted to this "
     "inconsistency. He acknowledged under oath that he previously blamed his former "
     "tenant for the mold damages in prior litigation, abandoning that theory only when "
     "it became financially advantageous to attribute the damage to Defendant. "
     "Plaintiffs cannot recover damages for property repairs by taking mutually "
     "exclusive factual positions in different judicial forums depending on who they are "
     "trying to extract money from. ", False),
    ("This is particularly so where Defendant's independent forensic engineer has "
     "identified an entirely different, self-inflicted cause for the very same mold.", True),
    (" This documented history of shifting blame destroys Plaintiffs' credibility and "
     "severs the causal link required to recover property damages from Defendant.", False),
])

h2("C.", "The Claimed Lost Rent Is Mathematically Impossible And Refuted By The Tenancy "
         "Timeline")
body([
    ("Plaintiffs assert a claim for $7,800 in lost rental income, calculating this "
     "figure based on an alleged three-month vacancy period at $2,600 per month. "
     "Plaintiffs contend the unit sat vacant because Orville Cole was forced to vacate "
     "due to water and mold issues. The objective tenancy timeline proves this "
     "calculation is false.", False),
])
body([
    ("The evidentiary record establishes that Orville Cole vacated the Dudash Property "
     "on May 31, 2023. By July 2023, a new commercial tenant, Gina, had already moved "
     "into the exact same space and commenced operating her hair salon. Because the "
     "transition between tenants took approximately one month, the mathematical reality "
     "of the tenancy schedule definitively refutes a three-month vacancy loss. "
     "Furthermore, any brief vacancy period between May and July 2023 was the natural "
     "result of standard commercial turnover, including the new tenant's build-out of "
     "the salon space and the removal of interior walls installed by the prior tenant. "
     "Plaintiffs cannot artificially inflate their damages by claiming a quarter-year of "
     "lost rent for a space that was successfully re-leased almost immediately. ", False),
    ("The objective record further undercuts Plaintiffs' claim: the suspected fungal "
     "growth in the east-end unit had been cleaned and the wall repainted less than two "
     "months before the November 2023 engineering inspection, and the salon tenant was "
     "in full operation—demonstrating that the unit was promptly returned to "
     "habitable, income-producing use.", True),
])

h2("D.", "Plaintiffs Are Legally Precluded From Recovering General Damages For Emotional "
         "Distress")
body([
    ("Plaintiffs allege general damages for severe emotional distress, including "
     "anxiety, humiliation, and sleep disturbances, which they attribute to the property "
     "dispute and a related police call. Setting aside the fact that Plaintiffs have "
     "produced zero psychiatric or medical records to substantiate these subjective "
     "complaints, the claims are barred by California law governing property disputes.", False),
])
body([
    ("Under California law, a plaintiff cannot recover general damages for emotional "
     "distress under theories of nuisance or trespass unless the plaintiff physically "
     "occupies the property at issue. The Dudash Property is a commercial building that "
     "Plaintiffs lease to third-party business operators. During the relevant time "
     "periods involving the alleged water intrusions and mold development, the affected "
     "units were physically occupied by commercial tenants, not by Plaintiffs. Because "
     "Plaintiffs operated strictly as absentee commercial landlords, their recovery is "
     "strictly limited to actual, proven property damage and lost use value.", False),
])
body([
    ("Even if Plaintiffs could legally pursue emotional distress damages, their claims "
     "are undermined by their extensive history of aggressive and litigious behavior. "
     "The evidentiary record includes documentation of a physical altercation wherein "
     "Mimi Dudash attempted to punch Defendant, as well as a decades-long pattern of "
     "Plaintiffs wrongfully withholding tenant security deposits, resulting in multiple "
     "small claims judgments against them that they simply refuse to pay. This history "
     "demonstrates that conflict and litigation are Plaintiffs' baseline temperament. "
     "The subjective emotional distress they now claim is neither severe nor uniquely "
     "caused by the conditions of the Matthews Property, and a jury will view these "
     "demands as an overreach unsupported by the facts or the law.", False),
])

# ============================== VI. SETTLEMENT POSITION ==============================
h1("VI.", "SETTLEMENT POSITION")
body([
    ("Defendant is insured under a Nationwide insurance policy, and the applicable "
     "policy limits will be disclosed confidentially to the mediator. On October 25, "
     "2023, Plaintiffs issued an initial settlement demand of $28,300, coupled with a "
     "demand for injunctive relief requiring Defendant to hire a licensed contractor to "
     "regrade and remark the parking spaces. Following the retention of new counsel, "
     "Plaintiffs issued a revised demand of $125,000 in late 2025. In December 2025, "
     "Defendant served statutory offers to compromise pursuant to Code of Civil "
     "Procedure section 998 in the amount of $2,500 to each Plaintiff, for a total offer "
     "of $5,000. Defendant participates in this mediation ready to negotiate in good "
     "faith to reach a reasonable resolution.", False),
])

# ============================== VII. CONCLUSION ==============================
h1("VII.", "CONCLUSION")
body([
    ("Plaintiffs’ claims are legally barred and factually unsupported. The physical "
     "conditions at issue have existed for decades, rendering the nuisance and trespass "
     "causes of action time-barred under the applicable statute of limitations. This "
     "procedural bar is compounded by Plaintiffs’ affirmative, annual maintenance "
     "of the steel bumper over the past twenty years, which establishes implied consent "
     "and defeats the trespass allegations as a matter of law. Furthermore, the physical "
     "evidence and Plaintiffs' own admissions confirm that the water intrusion was "
     "self-inflicted through Plaintiffs' improper application of cement and sealants over "
     "the building's engineered weep screeds, rather than any conditions on the Matthews "
     "Property. ", False),
    ("These conclusions are independently corroborated by the December 22, 2023 forensic "
     "engineering report of Amor Camatcho, P.E., which establishes that the parking lot "
     "drains away from the Dudash Property and that the water intrusion was caused by the "
     "unmaintained steel-bumper interface and Plaintiffs' own sealing of their "
     "engineered weep screeds. ", True),
    ("The value of this case is further diminished by Plaintiffs' speculative, medically "
     "unsubstantiated wage loss demands and a documented history of attributing the exact "
     "same property damage to their former tenants in collateral litigation. In light of "
     "these significant evidentiary hurdles, Defendant views this matter as presenting "
     "limited exposure, but participates in this mediation prepared to engage in "
     "good-faith negotiations to reach a reasonable and final resolution.", False),
])

out = r"E:\geminiterminal2\Defendant's Confidential Mediation Brief - UPDATED.docx"
doc.save(out)
print("Saved:", out)

# ---- verification ----
chk = Document(out)
total = len(chk.paragraphs)
bold_runs = 0
bold_chars = 0
for p in chk.paragraphs:
    for r in p.runs:
        if r.bold and r.text.strip():
            bold_runs += 1
            bold_chars += len(r.text)
print("Paragraphs:", total)
print("Bold runs (additions):", bold_runs, "| bold chars:", bold_chars)
