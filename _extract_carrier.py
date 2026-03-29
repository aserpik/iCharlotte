"""Extract text from Carrier004.docx and write to a temp file."""
from docx import Document
import os

doc_path = r'Z:\Shared\Current Clients\3100 - PARSAC\094 - Haire v. City of Wildomar\STATUS\[Draft] Carrier004.docx'
out_path = os.path.join(os.path.dirname(__file__), '_carrier_text.txt')

doc = Document(doc_path)
lines = []
for p in doc.paragraphs:
    if p.text.strip():
        lines.append(p.text)

# Also check tables
for table in doc.tables:
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
        if cells:
            lines.append(' | '.join(cells))

with open(out_path, 'w', encoding='utf-8') as f:
    f.write('\n'.join(lines))

print(f"Wrote {len(lines)} lines to {out_path}")
