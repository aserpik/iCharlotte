import os
import sys
import logging
import datetime
import glob
import re
from typing import Optional, List, Set
from docx import Document
from docx.shared import Pt

# --- Configuration ---
BASE_PATH_WIN = r"Z:\\Shared\\Current Clients"
LOG_FILE = os.path.join(os.getcwd(), "med_record_scan_activity.log")

# Setup Logging
logging.basicConfig(
    filename=LOG_FILE,
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    datefmt='%Y-%m-%d %H:%M:%S'
)

def log_event(message, level="info"):
    timestamp = datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    print(f"[{timestamp}] {message}")
    sys.stdout.flush()
    if level == "info":
        logging.info(message)
    elif level == "error":
        logging.error(message)
    elif level == "warning":
        logging.warning(message)

def get_case_path(file_num: str) -> Optional[str]:
    """Locates the physical directory for a file number (####.###)."""
    if not re.match(r'^\d{4}\.\d{3}$', file_num):
        log_event(f"Invalid file number format: {file_num}. Expected ####.###", level="error")
        return None

    carrier_num, case_num = file_num.split('.')
    base_path = BASE_PATH_WIN

    if not os.path.exists(base_path):
        log_event(f"Base path not found: {base_path}", level="error")
        return None

    # Find Carrier Folder
    carrier_folders = glob.glob(os.path.join(base_path, f"{carrier_num} - *"))
    if not carrier_folders:
        log_event(f"Carrier folder starting with {carrier_num} not found in {base_path}", level="error")
        return None
    carrier_path = carrier_folders[0]

    # Find Case Folder
    case_folders = glob.glob(os.path.join(carrier_path, f"{case_num} - *"))
    if not case_folders:
        log_event(f"Case folder starting with {case_num} not found in {carrier_path}", level="error")
        return None
    
    return os.path.abspath(case_folders[0])

def find_latest_chronology(records_dir: str) -> Optional[str]:
    """Finds the most recent .docx chronology with 'howell' in name."""
    if not os.path.exists(records_dir):
        log_event(f"RECORDS directory not found: {records_dir}", level="error")
        return None

    candidates = []
    # Date pattern: look for something that looks like a date 
    # Numeric: 12.24.25, 12-24-2025, 2025.12.24, 2025-12-24
    # Also handles potential month names if they are used, though less common in filenames
    date_regex = re.compile(r'(\d{1,4}[.\-/]\d{1,2}[.\-/]\d{1,4})|((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*[.\-\s]\d{1,2}[.\-\s]\d{2,4})', re.IGNORECASE)

    for file in os.listdir(records_dir):
        if file.lower().endswith(".docx") and not file.startswith("~$") and any(k in file.lower() for k in ["howell", "summary", "chronology"]):
            if date_regex.search(file):
                full_path = os.path.join(records_dir, file)
                candidates.append(full_path)

    if not candidates:
        log_event("No chronology documents found matching 'howell' and a date.", level="warning")
        return None

    # Sort by modification time to get the "most recent"
    candidates.sort(key=lambda x: os.path.getmtime(x), reverse=True)
    return candidates[0]

def extract_summarized_records(docx_path: str) -> List[str]:
    """Extracts filenames from 'RECORDS SUMMARIZED' table in docx."""
    summarized_files = []
    try:
        doc = Document(docx_path)
        
        target_table = None
        for table in doc.tables:
            found_header = False
            for row in table.rows[:3]:
                for cell in row.cells:
                    if "RECORDS SUMMARIZED" in cell.text.upper():
                        found_header = True
                        break
                if found_header:
                    break
            
            if found_header:
                target_table = table
        
        if target_table:
            # Try to identify the 'Filename' column index
            filename_col_idx = 0
            header_row = target_table.rows[1] if len(target_table.rows) > 1 else None
            if header_row:
                for idx, cell in enumerate(header_row.cells):
                    if "FILENAME" in cell.text.upper():
                        filename_col_idx = idx
                        break
            
            for row in target_table.rows:
                if len(row.cells) > filename_col_idx:
                    text = row.cells[filename_col_idx].text.strip()
                    if text and "RECORDS SUMMARIZED" not in text.upper() and "FILENAME" not in text.upper():
                        for line in text.split('\n'):
                            clean_line = line.strip()
                            if clean_line:
                                summarized_files.append(clean_line)
        else:
            log_event(f"Could not find 'RECORDS SUMMARIZED' table in {docx_path}", level="warning")
            
    except Exception as e:
        log_event(f"Error reading docx {docx_path}: {e}", level="error")
        
    return summarized_files

def main():
    if len(sys.argv) < 2:
        log_event("Usage: python med_record_scan.py <file_number>", level="error")
        sys.exit(1)
        
    file_num = sys.argv[1]
    log_event(f"--- Starting Med Record Scan for {file_num} ---")
    
    case_path = get_case_path(file_num)
    if not case_path:
        log_event("Could not locate case directory. Aborting.", level="error")
        sys.exit(1)
    
    records_dir = os.path.join(case_path, "RECORDS")
    notes_dir = os.path.join(case_path, "NOTES")
    
    if not os.path.exists(notes_dir):
        os.makedirs(notes_dir)
        log_event(f"Created NOTES directory: {notes_dir}")

    # 1. Find Chronology
    chronology_file = find_latest_chronology(records_dir)
    if not chronology_file:
        log_event("Aborting: No chronology document found.", level="error")
        sys.exit(1)
    
    log_event(f"Found latest chronology: {os.path.basename(chronology_file)}")
    
    # 2. Extract Summarized Records
    summarized_list = extract_summarized_records(chronology_file)
    summarized_lower = [s.lower() for s in summarized_list]
    log_event(f"Extracted {len(summarized_list)} entries from 'RECORDS SUMMARIZED' table.")
    
    # 3. Get all PDF files in RECORDS (recursively)
    pdf_files = []
    for root, _, files in os.walk(records_dir):
        for file in files:
            if file.lower().endswith(".pdf"):
                rel_path = os.path.relpath(os.path.join(root, file), records_dir)
                pdf_files.append(rel_path)
    
    log_event(f"Found {len(pdf_files)} PDF files in RECORDS.")
    
    # 4. Cross-reference
    not_summarized = []
    for pdf_path in pdf_files:
        pdf_filename = os.path.basename(pdf_path)
        pdf_lower = pdf_filename.lower()
        pdf_base = os.path.splitext(pdf_lower)[0]
        
        found = False
        for s_name in summarized_lower:
            # Match if filename is exactly in table, or if table entry is in filename (or vice-versa)
            if s_name == pdf_lower or s_name == pdf_base or s_name in pdf_lower or pdf_base in s_name:
                found = True
                break
        
        if not found:
            not_summarized.append(pdf_path)
    
    log_event(f"Identified {len(not_summarized)} records not summarized.")
    
    # 5. Save Output
    output_path = os.path.join(notes_dir, "Records_Not_Summarized.docx")
    try:
        from docx.shared import RGBColor
        doc = Document()
        
        # Set default style
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(12)
        style.font.color.rgb = RGBColor(0, 0, 0)
        
        def format_heading(heading):
            for run in heading.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)

        h0 = doc.add_heading('Records Scan Report', 0)
        format_heading(h0)
        
        doc.add_paragraph(f"Case File: {file_num}")
        doc.add_paragraph(f"Chronology Used: {os.path.basename(chronology_file)}")
        doc.add_paragraph(f"Date of Scan: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        
        h1_1 = doc.add_heading('Summary', 1)
        format_heading(h1_1)
        doc.add_paragraph(f"Total PDF records found in RECORDS folder: {len(pdf_files)}")
        doc.add_paragraph(f"Total entries extracted from 'RECORDS SUMMARIZED' table: {len(summarized_list)}")
        doc.add_paragraph(f"Records identified as NOT summarized: {len(not_summarized)}")

        h1_2 = doc.add_heading('Records NOT Summarized', 1)
        format_heading(h1_2)
        if not_summarized:
            doc.add_paragraph("The following PDF files were not found in the chronology's summarized table:")
            for item in sorted(not_summarized):
                doc.add_paragraph(item, style='List Bullet')
        else:
            doc.add_paragraph("All identified PDF files appear to be summarized.")

        h1_3 = doc.add_heading('Entries Extracted from Chronology', 1)
        format_heading(h1_3)
        if summarized_list:
            for item in sorted(summarized_list):
                doc.add_paragraph(item, style='List Bullet')
        else:
            doc.add_paragraph("No entries found in the summarized table.")

        h1_4 = doc.add_heading('All PDF Files Found in RECORDS', 1)
        format_heading(h1_4)
        if pdf_files:
            for item in sorted(pdf_files):
                doc.add_paragraph(item, style='List Bullet')
        else:
            doc.add_paragraph("No PDF files found.")
            
        doc.save(output_path)
        log_event(f"Saved results to {output_path}")
    except Exception as e:
        log_event(f"Error saving output docx: {e}", level="error")

if __name__ == "__main__":
    main()
