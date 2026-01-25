import os
import sys
import logging
import json
import datetime
import glob
import re
import subprocess
import shutil
from typing import Optional, Dict, Any, List

# Third-party imports
try:
    from docx import Document
    from docx.shared import Pt
    from pypdf import PdfReader
    import google.generativeai as genai
    import pytesseract
    from pdf2image import convert_from_path
    OCR_AVAILABLE = True
    
    # Windows-specific path configuration
    if os.name == 'nt':
        tesseract_path = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
        if os.path.exists(tesseract_path):
            pytesseract.pytesseract.tesseract_cmd = tesseract_path
        poppler_path = r"C:\Program Files\poppler\Library\bin"
        POPPLER_PATH = poppler_path if os.path.exists(poppler_path) else None
    else:
        POPPLER_PATH = None
except ImportError as e:
    print(f"Critical Error: Missing dependency {e}. Please ensure python-docx, pypdf, google-generativeai, pytesseract, and pdf2image are installed.")
    sys.exit(1)

# --- Configuration ---
PROJECT_ROOT = os.getcwd()
LOG_FILE = os.path.join(PROJECT_ROOT, "Discovery_Request_Activity.log")
GEMINI_DATA_DIR = os.path.join(PROJECT_ROOT, ".gemini", "case_data")
BASE_PATH_WIN = r"Z:\\Shared\\Current Clients"
TEMPLATES_DIR = os.path.join(PROJECT_ROOT, "Scripts", "Discovery Templates", "General PI Discovery")

# Map commonly used folder names to canonical categories (synced with audit agent)
FOLDER_MAP = {
    "DISCOVERY": ["DISCOVERY", "Discovery"],
    "PROPOUNDED": ["PROPOUNDED", "Propounded"],
    "RESPONSES": ["RESPONSES", "Responses"],
}

# Setup Logging
logging.basicConfig(
    filename=LOG_FILE,
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    datefmt='%Y-%m-%d %H:%M:%S'
)

def log_event(message, level="info"):
    timestamp = datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    if "--interactive" in sys.argv:
        print(f"[{timestamp}] {message}")
    if level == "info":
        logging.info(message)
    elif level == "error":
        logging.error(message)
    elif level == "warning":
        logging.warning(message)

def find_folder(base_path, candidates):
    """Finds the actual path for a folder given a list of candidate names (from audit.py)."""
    if not base_path or not os.path.exists(base_path):
        return None
    
    try:
        entries = os.listdir(base_path)
    except OSError:
        return None

    norm_candidates = [c.lower().strip() for c in candidates]

    for entry in entries:
        if entry.lower().strip() in norm_candidates:
             return os.path.join(base_path, entry)
    return None

def get_case_path(file_num: str) -> Optional[str]:
    """Locates the physical directory for a file number (####.###)."""
    if not re.match(r'^\d{4}\.\d{3}$', file_num):
        log_event(f"Invalid file number format: {file_num}.", level="error")
        return None

    carrier_num, case_num = file_num.split('.')
    base_path = BASE_PATH_WIN

    if not os.path.exists(base_path):
        log_event(f"Base path not found: {base_path}", level="error")
        return None

    carrier_folders = glob.glob(os.path.join(base_path, f"{carrier_num} - *"))
    if not carrier_folders:
        log_event(f"Carrier folder {carrier_num} not found.", level="error")
        return None
    carrier_path = carrier_folders[0]

    case_folders = glob.glob(os.path.join(carrier_path, f"{case_num} - *"))
    if not case_folders:
        log_event(f"Case folder {case_num} not found.", level="error")
        return None
    
    return os.path.abspath(case_folders[0])

def load_case_data(file_num: str) -> Dict[str, Any]:
    json_path = os.path.join(GEMINI_DATA_DIR, f"{file_num}.json")
    if os.path.exists(json_path):
        try:
            with open(json_path, 'r') as f:
                return json.load(f)
        except Exception as e:
            log_event(f"Error loading case data: {e}", level="error")
    return {}

def extract_text_from_pdf(file_path: str, max_pages: int = 5) -> str:
    try:
        reader = PdfReader(file_path)
        text = ""
        total_pages = min(len(reader.pages), max_pages)
        
        for i in range(total_pages):
            page_text = reader.pages[i].extract_text() or ""
            
            # Check if this specific page has meaningful text (threshold: 100 chars)
            if len(page_text.strip()) < 100 and OCR_AVAILABLE:
                log_event(f"Page {i+1} of {os.path.basename(file_path)} has insufficient text. Attempting OCR...")
                try:
                    images = convert_from_path(file_path, first_page=i+1, last_page=i+1, poppler_path=POPPLER_PATH)
                    if images:
                        ocr_text = pytesseract.image_to_string(images[0])
                        if len(ocr_text.strip()) > len(page_text.strip()):
                            page_text = ocr_text
                except Exception as ocr_e:
                    log_event(f"OCR failed for page {i+1}: {ocr_e}", level="warning")
            
            text += page_text + "\n"
        return text
    except Exception as e:
        log_event(f"Error reading PDF {file_path}: {e}", level="error")
        return ""

def call_gemini_api(prompt: str, context_text: str) -> Optional[str]:
    api_key = os.environ.get("GEMINI_API_KEY")
    if not api_key:
        log_event("GEMINI_API_KEY not set.", level="error")
        return None
    genai.configure(api_key=api_key)
    
    models_to_try = ["gemini-3-flash-preview", "gemini-2.5-flash"]
    
    for model_name in models_to_try:
        try:
            model = genai.GenerativeModel(model_name)
            response = model.generate_content(
                f"{prompt}\n\n[TEXT]\n{context_text}"
            )
            if response and response.text:
                return response.text
        except Exception as e:
            log_event(f"Model {model_name} failed: {e}", level="warning")
            continue
            
    return None

def analyze_prior_discovery(discovery_dir: str, party_name: str, client_name: str):
    log_event(f"Analyzing prior discovery in {discovery_dir} for {party_name}...")
    
    # Exclude these subfolder names (case-insensitive) as requested
    def is_excluded(folder_name):
        fn = folder_name.lower()
        if fn in {"subpoena", "subpoenas", "transcripts"}:
            return True
        # Ignore variations of "depo notice" or "depo ntc"
        if "depo" in fn and ("notice" in fn or "ntc" in fn):
            return True
        return False
    
    pdf_files = []
    
    # Logic similar to audit agent: look for PROPOUNDED party folder first
    prop_path = find_folder(discovery_dir, FOLDER_MAP["PROPOUNDED"])
    if prop_path:
        try:
            for entry in os.listdir(prop_path):
                entry_path = os.path.join(prop_path, entry)
                if os.path.isdir(entry_path):
                    # Check if this folder belongs to the target party
                    if party_name.lower() in entry.lower() or entry.lower() in party_name.lower():
                        log_event(f"Found party-specific propounded folder: {entry}")
                        for root, dirs, files in os.walk(entry_path):
                            dirs[:] = [d for d in dirs if not is_excluded(d)]
                            for f in files:
                                if f.lower().endswith(".pdf"):
                                    pdf_files.append(os.path.join(root, f))
        except OSError:
            pass

    # If no files found in party-specific propounded folder, or to be thorough, 
    # scan entire discovery dir excluding the forbidden folders.
    if not pdf_files:
        for root, dirs, files in os.walk(discovery_dir):
            dirs[:] = [d for d in dirs if not is_excluded(d)]
            for file in files:
                if file.lower().endswith(".pdf"):
                    pdf_files.append(os.path.join(root, file))
    
    results = []

    for pdf in pdf_files:
        filename = os.path.basename(pdf)
        
        # Primarily use the filename and path for discovery analysis to save tokens/time
        filename_prompt = f"""
        Analyze if this filename/path likely represents discovery served FROM '{client_name}' TO '{party_name}'.
        Path: {pdf}
        
        If it is discovery served TO {party_name}, extract:
        1. Type (SI, RPD, FROG, SROG, etc.)
        2. Set Number (1, 2, etc.)
        3. Date (if present in filename)
        Return JSON: {{"served": true/false, "type": "...", "set": "...", "date": "...", "confidence": "high/low"}}
        """
        
        res_json = call_gemini_api(filename_prompt, f"Filename: {filename}")
        if res_json:
            try:
                clean_json = res_json.strip().strip("```json").strip("```")
                data = json.loads(clean_json)
                if data.get("served") and data.get("confidence") == "high":
                    data["filename"] = filename
                    results.append(data)
                    continue # Skip full text extraction if filename is highly confident
            except:
                pass

        # Fallback to full text analysis if filename is ambiguous
        text = extract_text_from_pdf(pdf, max_pages=3)
        if not text: continue
        
        prompt = f"""
        Analyze if this discovery was served FROM '{client_name}' TO '{party_name}'.
        If yes, extract:
        1. Type (SI, RPD, etc.)
        2. Set Number
        3. Date Served
        4. Any custom definitions used.
        Return JSON: {{"served": true/false, "type": "...", "set": "...", "date": "...", "definitions": "..."}}
        Otherwise return {{"served": false}}
        """
        res_json = call_gemini_api(prompt, text)
        if res_json:
            try:
                clean_json = res_json.strip().strip("```json").strip("```")
                data = json.loads(clean_json)
                if data.get("served"):
                    data["filename"] = filename
                    results.append(data)
            except:
                pass
    
    return results

def get_caption_title_and_body(template_path: str):
    doc = Document(template_path)
    title = ""
    # Find the first non-empty paragraph for the title
    title_idx = -1
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip():
            full_text = p.text.strip()
            # Extract words inside quotation marks (handle both standard and smart quotes)
            match = re.search(r'["“](.+?)["”]', full_text)
            if match:
                title = match.group(1)
            else:
                # Fallback to entire text if no quotes found
                title = full_text
            title_idx = i
            break
    
    body_paragraphs = []
    if title_idx != -1:
        for p in doc.paragraphs[title_idx + 1:]:
            # Stop if we hit the signature block in the template
            if "dated:" in p.text.lower():
                break
            body_paragraphs.append(p)
            
    return title, body_paragraphs

from docx.shared import Pt, Inches

def copy_para_with_formatting(source_p, target_doc, before_para=None, prefix=None, counter=None):
    """Copies a paragraph with basic formatting and runs. Reconstructs numbering if prefix/counter provided."""
    if before_para:
        new_p = before_para.insert_paragraph_before()
    else:
        new_p = target_doc.add_paragraph()
        
    # Force Double Spacing (2.0)
    new_p.paragraph_format.line_spacing = 2.0
    new_p.paragraph_format.alignment = source_p.paragraph_format.alignment
    new_p.paragraph_format.first_line_indent = source_p.paragraph_format.first_line_indent
    new_p.paragraph_format.left_indent = source_p.paragraph_format.left_indent
    new_p.paragraph_format.right_indent = source_p.paragraph_format.right_indent
    new_p.paragraph_format.space_after = source_p.paragraph_format.space_after
    new_p.paragraph_format.space_before = source_p.paragraph_format.space_before
    
    # Handle numbering reconstruction for separators (e.g. SPECIAL INTERROGATORY NO. 1)
    if not source_p.text.strip() and prefix and counter is not None:
        if hasattr(source_p._element, 'pPr') and source_p._element.pPr is not None and source_p._element.pPr.numPr is not None:
            counter[0] += 1
            run = new_p.add_run(f"{prefix} NO. {counter[0]}:")
            run.bold = True
            run.underline = True # Added underline
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            # Indent removed from the separator paragraph itself as requested
            new_p.paragraph_format.first_line_indent = 0 
            return new_p

    # Apply 0.5 inch indent to:
    # 1. Instruction paragraphs (e.g., A., B., C.)
    # 2. Discovery requests and body paragraphs that follow a separator
    # 3. Paragraphs between caption and instructions (usually start with "Pursuant to" or "TO [PARTY]")
    is_instruction = re.match(r'^[A-Z]\.\s+', source_p.text.strip())
    is_following_separator = (counter is not None and counter[0] > 0)
    
    # Check if this is a paragraph that should be indented (excluding headings like "INSTRUCTIONS...")
    should_indent = False
    text_stripped = source_p.text.strip()
    if text_stripped:
        if is_instruction or is_following_separator:
            should_indent = True
        elif "TO " in text_stripped[:10] and " ATTORNEYS" in text_stripped:
            should_indent = True
        elif "Pursuant to " in text_stripped[:20]:
            should_indent = True
        elif "The place of inspection" in text_stripped:
            should_indent = True
        elif "A written response" in text_stripped:
            should_indent = True
        elif "Failure to provide" in text_stripped:
            should_indent = True
        elif "PROPOUNDING PARTY:" in text_stripped:
            should_indent = True
        elif "RESPONDING PARTY:" in text_stripped:
            should_indent = True
        elif "SET NO.:" in text_stripped:
            should_indent = True

    if should_indent:
        new_p.paragraph_format.first_line_indent = Inches(0.5)
    else:
        new_p.paragraph_format.first_line_indent = 0

    for run in source_p.runs:
        new_run = new_p.add_run(run.text)
        new_run.bold = run.bold
        new_run.italic = run.italic
        new_run.underline = run.underline
        if run.font.name:
            new_run.font.name = run.font.name
        if run.font.size:
            new_run.font.size = run.font.size
    return new_p

def replace_placeholders(doc, replacements: Dict[str, str]):
    def do_replace(text):
        for k, v in replacements.items():
            # Check for case matching
            if k.isupper():
                text = text.replace(k, str(v).upper())
            else:
                text = text.replace(k, str(v))
        return text

    def process_paragraphs(paragraphs):
        for para in paragraphs:
            # First, check if the placeholder is in the full text but potentially split across runs
            full_text = para.text
            needs_replacement = any(k in full_text for k in replacements)
            
            if needs_replacement:
                # If there's only one run or it's simple, do direct replacement
                if len(para.runs) <= 1:
                    for run in para.runs:
                        if run.text:
                            run.text = do_replace(run.text)
                else:
                    # Complex case: placeholder might be split. 
                    # We'll do a simple but effective run-merging replacement if needed,
                    # or just replace the full text and clear other runs if we have to.
                    # For discovery, usually we can just join runs, replace, and put back in first run.
                    combined = "".join(r.text for r in para.runs)
                    new_text = do_replace(combined)
                    if combined != new_text:
                        # Clear all runs
                        for i in range(len(para.runs)):
                            para.runs[i].text = ""
                        # Set text of first run
                        para.runs[0].text = new_text

    process_paragraphs(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                process_paragraphs(cell.paragraphs)

def draft_discovery(case_path: str, party_name: str, client_name: str, disc_type_opt: str):
    log_event(f"Drafting discovery for {party_name}...")
    output_dir = os.path.join(case_path, "NOTES", "AI OUTPUT", "AI DISCOVERY")
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    # 1. Locate Caption Template
    case_data = load_case_data(extract_file_num(case_path))
    caption_template = case_data.get("caption_template_path")
    if not caption_template or not os.path.exists(caption_template):
        log_event("Caption template not found in case data. Attempting discovery.", level="info")
        # Try to find it manually
        from complaint import find_caption_doc
        caption_template = find_caption_doc(case_path)
        if not caption_template:
            log_event("No Caption Document found.", level="error")
            return

    # Files to generate
    files_to_gen = [
        {"prefix": "SI(1)", "template": "SI(1)_tPltf.docx"},
        {"prefix": "RPD(1)", "template": "RPD(1)_tPltf.docx"}
    ]

    for item in files_to_gen:
        new_filename = f"{item['prefix']} t{party_name}.docx"
        new_path = os.path.join(output_dir, new_filename)
        shutil.copy2(caption_template, new_path)
        
        templ_path = os.path.join(TEMPLATES_DIR, item['template'])
        if not os.path.exists(templ_path):
            log_event(f"Template {templ_path} not found.", level="error")
            continue
            
        templ_title, templ_body = get_caption_title_and_body(templ_path)
        
        doc = Document(new_path)
        
        # Replace CAPTION PAGE / TITLE with title and remove other words in that area
        found_placeholder = False
        
        # Search paragraphs
        for para in doc.paragraphs:
            if "CAPTION PAGE" in para.text.upper() or "CAPTION TITLE" in para.text.upper():
                para.text = "" # Remove other words in that area
                run = para.add_run(templ_title)
                run.bold = True
                found_placeholder = True

        # Search tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if "CAPTION PAGE" in para.text.upper() or "CAPTION TITLE" in para.text.upper():
                            para.text = "" # Remove other words in that area
                            run = para.add_run(templ_title)
                            run.bold = True
                            found_placeholder = True

        # Determine separator prefix
        prefix_str = ""
        if "SI" in item['prefix']:
            prefix_str = "SPECIAL INTERROGATORY"
        elif "RPD" in item['prefix']:
            prefix_str = "REQUEST FOR PRODUCTION"
        
        counter = [0]

        # Find signature block start to insert body before it
        insertion_para = None
        for p in doc.paragraphs:
            if "dated:" in p.text.lower():
                insertion_para = p
                break
        
        # Insert template body text
        for p in templ_body:
            copy_para_with_formatting(p, doc, before_para=insertion_para, prefix=prefix_str, counter=counter)

        # Replacements
        replacements = {
            "[RESPONDING_PARTY]": party_name,
            "[PROPOUNDING_PARTY]": client_name
        }
        replace_placeholders(doc, replacements)
        
        # Final pass: Ensure the body text is double spaced, but SKIP caption tables and signature block
        in_signature_block = False
        for p in doc.paragraphs:
            if "dated:" in p.text.lower() or in_signature_block:
                in_signature_block = True
                p.paragraph_format.line_spacing = 1.0 # Single space signature block
                continue
            
            p.paragraph_format.line_spacing = 2.0
            
        # Keep tables single spaced (usually caption table)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        p.paragraph_format.line_spacing = 1.0

        # Update Footer: Size 10, Center Aligned, Replaced Title, No "CAPTION PAGE"
        final_footer_text = templ_title
        for k, v in replacements.items():
            final_footer_text = final_footer_text.replace(k, str(v).upper() if k.isupper() else str(v))

        for section in doc.sections:
            # Handle all footer types
            footers = [section.footer, section.first_page_footer, section.even_page_footer]
            for footer in footers:
                if footer is None: continue
                # Clear all paragraphs in this footer
                for p in footer.paragraphs:
                    p.text = ""
                    # Remove any existing runs
                    for run in p.runs:
                        run.text = ""
                
                # If no paragraphs exist, add one
                if not footer.paragraphs:
                    footer.add_paragraph()
                
                # Set the text in the first paragraph
                p = footer.paragraphs[0]
                p.alignment = 1 # Center
                run = p.add_run(final_footer_text)
                run.font.size = Pt(10)
                run.font.name = 'Times New Roman'

        doc.save(new_path)
        log_event(f"Generated: {new_filename}")

def extract_file_num(path: str) -> str:
    # Reverse of get_case_path
    # Path looks like: .../#### - Carrier Name/### - Case Name
    match = re.search(r'(\d{4}) - .+[\\/](\d{3}) -', path)
    if match:
        return f"{match.group(1)}.{match.group(2)}"
    return ""

def interactive_loop(file_num: str):
    case_path = get_case_path(file_num)
    if not case_path:
        print(f"Error: Could not find case folder for {file_num}")
        input("Press Enter to exit...")
        return

    case_data = load_case_data(file_num)
    plaintiffs = case_data.get("plaintiffs", [])
    defendants = case_data.get("defendants", [])
    client_name = case_data.get("client_name", "Our Client")

    all_parties = plaintiffs + defendants
    if not all_parties:
        print("No parties found in case data.")
        input("Press Enter to exit...")
        return

    print(f"\n--- Discovery Request Agent: Case {file_num} ---")
    print("Which Party do you want to propound discovery on?")
    for i, party in enumerate(all_parties, 1):
        print(f"{i}. {party}")

    choice = input("\nEnter number(s) separated by commas (e.g. 1, 3): ")
    selected_indices = [int(x.strip()) - 1 for x in choice.split(",") if x.strip().isdigit()]
    
    selected_parties = [all_parties[i] for i in selected_indices if 0 <= i < len(all_parties)]

    discovery_dir = os.path.join(case_path, "DISCOVERY")
    
    for party in selected_parties:
        print(f"\nChecking prior discovery for: {party}")
        prior = []
        if os.path.exists(discovery_dir):
            prior = analyze_prior_discovery(discovery_dir, party, client_name)
        
        if not prior:
            print(f"No prior discovery served from {client_name} on {party}.")
        else:
            print(f"Found prior discovery for {party}:")
            for p in prior:
                print(f" - {p.get('type')} Set {p.get('set')} served on {p.get('date')}")

        print("\nWhat type of discovery do you want to generate?")
        print("a. Basic PI Discovery")
        print("b. Basic Wrongful Death Discovery")
        print("c. Next Set Discovery")
        
        dtype = input("Choice (a/b/c): ").lower()
        if dtype in ['a', 'b', 'c']:
            draft_discovery(case_path, party, client_name, dtype)
        else:
            print("Invalid choice.")

    print("\nProcessing complete.")
    input("Press Enter to exit...")

def main():
    if len(sys.argv) < 2:
        sys.exit(1)

    file_num = sys.argv[1]
    
    if "--interactive" in sys.argv:
        interactive_loop(file_num)
    else:
        # Spawn new terminal window
        log_event(f"Triggering interactive discovery agent for {file_num}")
        script_path = os.path.abspath(__file__)
        if os.name == 'nt':
            # Use cmd /c start to pop up a new window
            cmd = f'start "Discovery Request Agent - {file_num}" cmd /c python "{script_path}" {file_num} --interactive'
            subprocess.Popen(cmd, shell=True)
        else:
            # For non-windows, we might just run it or use a terminal emulator
            print("Interactive mode only supported on Windows for now.")

if __name__ == "__main__":
    main()
