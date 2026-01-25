import os
import sys
import logging
import json
import datetime
import glob
import re
import subprocess
import shutil
import time
from typing import Optional, Dict, Any

try:
    import google.generativeai as genai
    from docx import Document
    from docx.shared import Pt
    from docx.oxml.shared import OxmlElement, qn
    from docx.opc.constants import RELATIONSHIP_TYPE as RELs
    from pypdf import PdfReader
    import pytesseract
    from pdf2image import convert_from_path
except ImportError as e:
    print(f"Critical Error: Missing dependency {e}. Please ensure google-generativeai, python-docx, pypdf, pytesseract, pdf2image are installed.")
    sys.exit(1)

# --- Configuration ---
# Directories
BASE_PATH_WIN = r"Z:\\Shared\\Current Clients"
GEMINI_DATA_DIR = os.path.join(os.getcwd(), ".gemini", "case_data")
LOG_FILE = os.path.join(os.getcwd(), "Docket_activity.log")
SCRIPTS_DIR = os.path.join(os.getcwd(), "Scripts")
SCRAPER_SCRIPT = os.path.join(SCRIPTS_DIR, "docket_scraper.py")
RIVERSIDE_SCRAPER_SCRIPT = os.path.join(SCRIPTS_DIR, "riverside_docket_scraper.py")
KERN_SCRAPER_SCRIPT = os.path.join(SCRIPTS_DIR, "kern_docket_scraper.py")
SB_SCRAPER_SCRIPT = os.path.join(SCRIPTS_DIR, "san_bernardino_docket_scraper.py")
SD_SCRAPER_SCRIPT = os.path.join(SCRIPTS_DIR, "san_diego_docket_scraper.py")
OC_SCRAPER_SCRIPT = os.path.join(SCRIPTS_DIR, "orange_county_docket_scraper.py")
SACRAMENTO_SCRAPER_SCRIPT = os.path.join(SCRIPTS_DIR, "sacramento_docket_scraper.py")
COMPLAINT_SCRIPT = os.path.join(SCRIPTS_DIR, "complaint.py")
PROCEDURAL_HIST_PROMPT_FILE = os.path.join(SCRIPTS_DIR, "Procedural_History.txt")
MASTER_STATUS_PATH = r"C:\Users\ASerpik.DESKTOP-MRIMK0D\OneDrive - Bordin Semmer LLP\Desktop\MASTER_CASE_STATUS.docx"

# OCR Settings
OCR_AVAILABLE = False
POPPLER_PATH = None

if os.name == 'nt':
    # Tesseract Path
    tesseract_paths = [
        r"C:\\Program Files\\Tesseract-OCR\\tesseract.exe",
        os.path.expanduser(r"~\\AppData\\Local\\Programs\\Tesseract-OCR\\tesseract.exe")
    ]
    for p in tesseract_paths:
        if os.path.exists(p):
            pytesseract.pytesseract.tesseract_cmd = p
            OCR_AVAILABLE = True
            break
    
    # Poppler Path
    poppler_paths = [
        r"C:\\Program Files\\poppler\\Library\\bin",
        r"C:\\Program Files\\poppler-0.68.0\\bin",
        r"C:\\Program Files (x86)\\poppler\\Library\\bin"
    ]
    for p in poppler_paths:
        if os.path.exists(p):
            POPPLER_PATH = p
            break
else:
    if shutil.which("tesseract"):
        OCR_AVAILABLE = True
    if shutil.which("pdftoppm"):
        POPPLER_PATH = None

# Setup Logging
logging.basicConfig(
    filename=LOG_FILE,
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    datefmt='%Y-%m-%d %H:%M:%S'
)

def log_event(message, level="info"):
    timestamp = datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    print(f"[{timestamp}] {message}")
    sys.stdout.flush()
    if level == "info":
        logging.info(message)
    elif level == "error":
        logging.error(message)
    elif level == "warning":
        logging.warning(message)

def parse_file_numbers(input_str: str):
    """Parses a string containing file numbers, ranges, or commas into a list."""
    # Remove spaces and split by comma
    parts = [p.strip() for p in input_str.split(',')]
    files = []
    
    for part in parts:
        part = part.strip("'").strip('"')
        if '-' in part:
            try:
                start_str, end_str = part.split('-')
                start_str = start_str.strip()
                end_str = end_str.strip()
                
                # Assume format Prefix.Suffix (e.g., 5800.004)
                if '.' in start_str:
                    prefix_start, suffix_start = start_str.rsplit('.', 1)
                    
                    # Handle end string formats: "5800.008" or just "008"
                    if '.' in end_str:
                        prefix_end, suffix_end = end_str.rsplit('.', 1)
                        if prefix_start != prefix_end:
                            # If prefixes differ, treat as separate items? 
                            # For safety, let's just log warning or fallback to simple processing if needed.
                            # But assuming user meant a range within the same prefix.
                            pass 
                    else:
                        suffix_end = end_str # e.g. "5800.004-008"
                    
                    s_val = int(suffix_start)
                    e_val = int(suffix_end)
                    padding = len(suffix_start)
                    
                    if s_val > e_val:
                        # Descending range? Swap or just handle
                        s_val, e_val = e_val, s_val
                        
                    for i in range(s_val, e_val + 1):
                        files.append(f"{prefix_start}.{str(i).zfill(padding)}")
                else:
                    # Integer range
                    s_val = int(start_str)
                    e_val = int(end_str)
                    if s_val > e_val:
                         s_val, e_val = e_val, s_val
                    for i in range(s_val, e_val + 1):
                        files.append(str(i))
                        
            except ValueError:
                log_event(f"Failed to parse range: {part}. Treating as literal.", level="warning")
                files.append(part)
        else:
            if part:
                files.append(part)
                
    return files

def get_case_path(file_num: str):
    """Locates the physical directory for a file number (####.###)."""
    if not re.match(r'^\d{4}\.\d{3}$', file_num):
        log_event(f"Invalid file number format: {file_num}. Expected ####.###", level="error")
        return None

    carrier_num, case_num = file_num.split('.')
    base_path = BASE_PATH_WIN

    if not os.path.exists(base_path):
        log_event(f"Base path not found: {base_path}", level="error")
        return None

    # Find Carrier Folder
    carrier_folders = glob.glob(os.path.join(base_path, f"{carrier_num} - *"))
    if not carrier_folders:
        log_event(f"Carrier folder starting with {carrier_num} not found in {base_path}", level="error")
        return None
    carrier_path = carrier_folders[0]

    # Find Case Folder
    case_folders = glob.glob(os.path.join(carrier_path, f"{case_num} - *"))
    if not case_folders:
        log_event(f"Case folder starting with {case_num} not found in {carrier_path}", level="error")
        return None
    
    return os.path.abspath(case_folders[0])

def get_case_var(file_num: str, key: str):
    """Retrieves a variable from the case's JSON state file."""
    json_file = os.path.join(GEMINI_DATA_DIR, f"{file_num}.json")
    if not os.path.exists(json_file):
        return None
    
    try:
        with open(json_file, 'r') as f:
            data = json.load(f)
            return data.get(key)
    except Exception as e:
        log_event(f"Error reading JSON for {file_num}: {e}", level="error")
        return None

def save_case_var(file_num: str, key: str, value: Any):
    """Saves a variable to the case's JSON state file."""
    if not os.path.exists(GEMINI_DATA_DIR):
        os.makedirs(GEMINI_DATA_DIR)
    
    json_file = os.path.join(GEMINI_DATA_DIR, f"{file_num}.json")
    data = {}
    
    if os.path.exists(json_file):
        try:
            with open(json_file, 'r') as f:
                data = json.load(f)
        except json.JSONDecodeError:
            log_event(f"Corrupt JSON for {file_num}, resetting.", level="error")
            data = {}
            
    data[key] = value
    
    with open(json_file, 'w') as f:
        json.dump(data, f, indent=4)

def find_complaint_file(case_path: str) -> Optional[str]:
    """Finds the primary complaint in the PLEADINGS or PLEADING directory and specific subfolders."""
    
    # 1. Identify Base Directories
    base_dirs = []
    for d in ["PLEADINGS", "PLEADING"]:
        path = os.path.join(case_path, d)
        if os.path.exists(path):
            base_dirs.append(path)
    
    if not base_dirs:
        log_event(f"Neither PLEADINGS nor PLEADING directory found at {case_path}", level="warning")
        return None

    # 2. Identify Search Paths (Base + Specific Subfolders)
    search_paths = list(base_dirs) # Start with base directories
    target_subs = ["complaint", "fac", "sac", "s&c"]
    
    for base in base_dirs:
        try:
            # List immediate subdirectories
            with os.scandir(base) as it:
                for entry in it:
                    if entry.is_dir() and entry.name.lower() in target_subs:
                        search_paths.append(entry.path)
        except OSError as e:
            log_event(f"Error scanning directory {base}: {e}", level="warning")

    candidates = []
    # 3. Search for files with "complaint" in name within identified paths
    for search_dir in search_paths:
        try:
            for root, _, files in os.walk(search_dir):
                for file in files:
                    if file.lower().endswith(".pdf") and "complaint" in file.lower():
                        candidates.append(os.path.join(root, file))
        except OSError as e:
             log_event(f"Error walking directory {search_dir}: {e}", level="warning")
    
    # Priority filtering
    filtered_candidates = []
    ignore_terms = ["motion", "response", "demurrer", "answer", "reply"]
    
    for c in candidates:
        name = os.path.basename(c).lower()
        if not any(term in name for term in ignore_terms):
            filtered_candidates.append(c)
    
    if not filtered_candidates:
        if candidates:
            log_event("All complaint candidates filtered out by priority terms. Using first raw candidate.", level="warning")
            return candidates[0]
        log_event("No complaint file found in PLEADINGS/PLEADING or target subfolders.", level="warning")
        return None

    return filtered_candidates[0]

def extract_text_from_file(file_path: str, force_ocr_first_page: bool = False) -> str:
    """Extracts text, falling back to OCR if needed."""
    if not file_path or not os.path.exists(file_path):
        return ""

    log_event(f"Extracting text from: {file_path}")
    text = ""
    ext = os.path.splitext(file_path)[1].lower()

    try:
        if ext == ".pdf":
            reader = PdfReader(file_path)
            
            def ocr_page(page_idx):
                if not OCR_AVAILABLE:
                    return ""
                try:
                    images = convert_from_path(file_path, first_page=page_idx+1, last_page=page_idx+1, poppler_path=POPPLER_PATH)
                    if images:
                        return pytesseract.image_to_string(images[0])
                except Exception as e:
                    log_event(f"OCR failed for page {page_idx+1}: {e}", level="error")
                return ""

            for i, page in enumerate(reader.pages):
                if force_ocr_first_page and i == 0:
                    page_text = ocr_page(i)
                else:
                    page_text = page.extract_text() or ""
                    if len(page_text.strip()) < 100:
                        log_event(f"Page {i+1} insufficient text (<100 chars). Triggering OCR.", level="warning")
                        ocr_txt = ocr_page(i)
                        if len(ocr_txt.strip()) > len(page_text.strip()):
                            page_text = ocr_txt

                text += page_text + "\n"

        elif ext in [".docx", ".doc"]:
            if ext == ".docx":
                doc = Document(file_path)
                for p in doc.paragraphs:
                    text += p.text + "\n"
        
        return text

    except Exception as e:
        log_event(f"Error extracting text from {file_path}: {e}", level="error")
        return ""

def apply_markdown(cell, text: str):
    """Parses basic Markdown (bold only) and adds to cell."""
    cell.text = ""
    lines = text.split('\n')
    for line in lines:
        paragraph = cell.add_paragraph()
        parts = line.split('**')
        for i, part in enumerate(parts):
            run = paragraph.add_run(part)
            if i % 2 == 1:
                run.font.bold = True

def update_variables_docx(case_dir: str, data: Dict):
    """Compiles extracted variables into variables.docx in NOTES/AI OUTPUT folder."""
    notes_dir = os.path.join(case_dir, "NOTES", "AI OUTPUT")
    if not os.path.exists(notes_dir):
        os.makedirs(notes_dir)
    
    docx_path = os.path.join(notes_dir, "variables.docx")
    
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    
    doc.add_heading(f"Case Variables: {data.get('case_number', 'Unknown')}", 0)
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Variable'
    hdr_cells[1].text = 'Value'
    
    for k, v in data.items():
        if v is None or v == "None": continue
        if k == "hidden_complaint_text": continue # Skip hidden text variable
        row_cells = table.add_row().cells
        row_cells[0].text = str(k)
        if k in ["factual_background", "procedural_history"]:
            apply_markdown(row_cells[1], str(v))
        else:
            row_cells[1].text = str(v)
        
    doc.save(docx_path)
    log_event(f"Saved variables to {docx_path}")

def call_gemini_api(prompt: str, context_text: str, models: list = ["gemini-2.5-flash", "gemini-1.5-flash"]) -> Optional[str]:
    """Calls Gemini with fallback models."""
    api_key = os.environ.get("GEMINI_API_KEY")
    if not api_key:
        log_event("GEMINI_API_KEY not set.", level="error")
        return None

    genai.configure(api_key=api_key)
    # Limit context
    full_prompt = f"{prompt}\n\n[CONTEXT DOCUMENT]\n{context_text[:100000]}" 

    for model_name in models:
        try:
            log_event(f"Querying model: {model_name}")
            model = genai.GenerativeModel(model_name)
            response = model.generate_content(full_prompt)
            if response and response.text:
                return response.text
        except Exception as e:
            log_event(f"Model {model_name} failed: {e}", level="warning")
            continue
    
    log_event("All AI models failed.", level="error")
    return None

def clean_json_string(s: str) -> str:
    """Cleans Markdown code blocks from string to get raw JSON."""
    s = s.strip()
    if s.startswith("```"):
        lines = s.split('\n')
        if len(lines) >= 2:
            s = '\n'.join(lines[1:-1])
    return s.strip()

def add_hyperlink(paragraph, url, text, color="0000FF", underline=True):
    """
    A helper function to place a hyperlink within a paragraph object.
    """
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, RELs.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # Create a w:r element
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Add color if it is given
    if color:
        c = OxmlElement('w:color')
        c.set(qn('w:val'), color)
        rPr.append(c)

    # Remove underlining if it is requested
    if not underline:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'none')
        rPr.append(u)

    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    
    return hyperlink

def update_master_status(file_num: str, data: Dict[str, Any], docket_path: str = None):
    """Updates the Master Case Status Word document."""
    doc_path = MASTER_STATUS_PATH
    
    # Helper: Format dates to MM-DD-YYYY
    def format_date(date_str):
        if not date_str or date_str == "None":
            return date_str
        try:
            dt = datetime.datetime.strptime(date_str, "%Y-%m-%d")
            return dt.strftime("%m-%d-%Y")
        except ValueError:
            return date_str

    # Helper: Clean text (remove []')
    def clean_text(text):
        if not text: return ""
        return str(text).replace("[", "").replace("]", "").replace("'", "")

    # Helper: Format hearings to "Name (MM-DD-YY)"
    def format_hearings(hearings_str):
        if not hearings_str or hearings_str == "None":
            return hearings_str
        
        # Abbreviations
        hearing_map = [
            (r"(?i)Final Status Conference", "FSC"),
            (r"(?i)Trial Setting Conference", "TSC"),
            (r"(?i)Trial Readiness Conference", "TRC"),
            (r"(?i)Order to Show Cause", "OSC"),
            (r"(?i)Motion for Summary Judgment", "MSJ")
        ]
        for pattern, replacement in hearing_map:
            hearings_str = re.sub(pattern, replacement, hearings_str)
        
        # Reformat dates
        # Regex for YYYY-MM-DD
        date_pattern = re.compile(r"(\d{4}-\d{2}-\d{2})")
        
        parts = hearings_str.split(',') # Assume comma separation for multiple
        new_parts = []
        for part in parts:
            part = part.strip()
            match = date_pattern.search(part)
            if match:
                date_str = match.group(1)
                try:
                    dt = datetime.datetime.strptime(date_str, "%Y-%m-%d")
                    formatted_date = dt.strftime("%m-%d-%y")
                    # Remove date and "on" from name
                    name = part.replace(date_str, "").replace(" on ", " ").strip()
                    new_parts.append(f"{name} ({formatted_date})")
                except ValueError:
                    new_parts.append(part)
            else:
                new_parts.append(part)
        
        return ", ".join(new_parts)

    if not os.path.exists(doc_path):
        log_event("Master Status doc not found. Creating new one.")
        doc = Document()
        table = doc.add_table(rows=1, cols=8)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = "File Number"
        hdr[1].text = "Plaintiff Name"
        hdr[2].text = "Case number"
        hdr[3].text = "County"
        hdr[4].text = "Assigned Associate"
        hdr[5].text = "Trial Date"
        hdr[6].text = "Other Hearings"
        hdr[7].text = "Last Updated"
    else:
        try:
            doc = Document(doc_path)
        except Exception as e:
            log_event(f"Failed to open Master Status doc: {e}", level="error")
            return

    if not doc.tables:
        table = doc.add_table(rows=1, cols=8)
        # ... header setup ...
    else:
        table = doc.tables[0]

    # Look for existing row
    target_row = None
    for row in table.rows:
        if row.cells[0].text.strip() == file_num:
            target_row = row
            break
    
    if not target_row:
        target_row = table.add_row()
    
    # Update cells
    try:
        # 1. File Number
        cell_fn = target_row.cells[0]
        cell_fn.text = "" 
        case_folder = get_case_path(file_num)
        if case_folder:
            p = cell_fn.paragraphs[0]
            add_hyperlink(p, case_folder, clean_text(file_num))
        else:
            cell_fn.text = clean_text(file_num)

        # 2. Plaintiff Name
        plaintiffs_raw = clean_text(data.get("plaintiffs", ""))
        if plaintiffs_raw:
            first_pl = plaintiffs_raw.split(',')[0].strip()
            last_name = first_pl.split()[-1] if first_pl else ""
            target_row.cells[1].text = last_name
        else:
            target_row.cells[1].text = ""

        # 3. Case Number
        cell_cn = target_row.cells[2]
        cell_cn.text = ""
        case_num_str = clean_text(data.get("case_number", ""))
        if docket_path:
             p = cell_cn.paragraphs[0]
             add_hyperlink(p, docket_path, case_num_str)
        else:
             cell_cn.text = case_num_str

        # 4. County
        target_row.cells[3].text = clean_text(data.get("venue_county", ""))

        # 5. Associate
        target_row.cells[4].text = ""

        # 6. Trial Date
        target_row.cells[5].text = format_date(clean_text(str(data.get("trial_date", ""))))

        # 7. Other Hearings
        raw_hearings = clean_text(str(data.get("other_hearings", "")))
        target_row.cells[6].text = format_hearings(raw_hearings)

        # 8. Last Updated
        target_row.cells[7].text = datetime.datetime.now().strftime("%m-%d-%Y")

        # Apply Styling (Times New Roman, Size 8)
        for cell in target_row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(8)
                
                # Hyperlinks style
                for child in paragraph._element:
                    if child.tag.endswith('hyperlink'):
                        for subchild in child:
                            if subchild.tag.endswith('r'):
                                from docx.text.run import Run
                                run_obj = Run(subchild, paragraph)
                                run_obj.font.name = 'Times New Roman'
                                run_obj.font.size = Pt(8)

        # Sort Rows
        try:
            all_rows = list(table.rows)
            if len(all_rows) > 1:
                header_row = all_rows[0]
                data_rows = all_rows[1:]
                def get_sort_key(row):
                    try:
                        return float(row.cells[0].text.strip())
                    except ValueError:
                        return float('inf')
                data_rows.sort(key=get_sort_key)
                tbl = table._tbl
                for row in data_rows:
                    tbl.remove(row._tr)
                for row in data_rows:
                    tbl.append(row._tr)
        except Exception as e:
             log_event(f"Error sorting rows: {e}", level="warning")

        doc.save(doc_path)
        log_event(f"Updated Master Case Status at {doc_path}")
    except Exception as e:
        log_event(f"Error saving Master Status doc: {e}", level="error")

def run_complaint_agent(file_num: str):
    """Runs the complaint agent script."""
    log_event(f"Triggering Complaint Agent for {file_num}...")
    try:
        # Capture output to diagnose immediate failures (e.g. import errors)
        result = subprocess.run(
            [sys.executable, COMPLAINT_SCRIPT, file_num], 
            capture_output=True, 
            text=True,
            check=True
        )
        log_event("Complaint Agent finished.")
        # log_event(f"Complaint Agent Output:\n{result.stdout}") # Optional: detailed logging
    except subprocess.CalledProcessError as e:
        log_event(f"Complaint Agent failed with exit code {e.returncode}.", level="error")
        log_event(f"STDOUT: {e.stdout}", level="error")
        log_event(f"STDERR: {e.stderr}", level="error")
    except Exception as e:
        log_event(f"Failed to run Complaint Agent: {e}", level="error")

def create_summary_report(results):
    """Creates a Word document summary of the processed file numbers and their status."""
    timestamp = datetime.datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    report_filename = f"dockets_processed_{timestamp}.docx"
    
    try:
        doc = Document()
        doc.add_heading('Dockets Processed Report', 0)
        doc.add_paragraph(f"Report Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'File Number'
        hdr_cells[1].text = 'Status'

        for fn, status in results:
            row_cells = table.add_row().cells
            row_cells[0].text = str(fn)
            row_cells[1].text = str(status)
            
        doc.save(report_filename)
        log_event(f"Summary report created: {report_filename}")
        return report_filename
    except Exception as e:
        log_event(f"Failed to create summary report: {e}", level="error")
        return None

def main():
    if len(sys.argv) < 2:
        log_event("Usage: python docket.py <file_number> [--headless] [--headful]", level="error")
        sys.exit(1)

    # Check for headless/headful flags
    is_headless = "--headless" in sys.argv
    is_headful = "--headful" in sys.argv
    
    args_to_parse = [arg for arg in sys.argv[1:] if arg not in ["--headless", "--headful"]]
    
    raw_arg = " ".join(args_to_parse)
    file_numbers = parse_file_numbers(raw_arg)
    
    # --- Dispatcher Mode ---
    if len(file_numbers) > 1:
        log_event(f"Detected multiple file numbers: {file_numbers}. Launching separate agents in batches...")
        
        MAX_CONCURRENT = 5
        AGENT_TIMEOUT = 600 # 10 minutes
        results = []
        processes = []
        
        # We use a simple loop to manage concurrent processes
        file_queue = list(file_numbers)
        active_processes = {} # pid -> (file_num, process_obj, start_time)
        
        while file_queue or active_processes:
            # Fill slots
            while file_queue and len(active_processes) < MAX_CONCURRENT:
                fn = file_queue.pop(0)
                log_event(f"Spawning agent for {fn}...")
                try:
                    spawn_args = [sys.executable, sys.argv[0], fn, "--headful"]
                    if is_headless:
                        spawn_args.append("--headless")
                    
                    if os.name == 'nt':
                        p = subprocess.Popen(spawn_args, creationflags=subprocess.CREATE_NO_WINDOW)
                    else:
                        p = subprocess.Popen(spawn_args)
                    
                    active_processes[p.pid] = (fn, p, time.time())
                except Exception as e:
                    log_event(f"Failed to spawn agent for {fn}: {e}", level="error")
                    results.append((fn, f"Failed to Start: {e}"))
            
            # Check for finished or timed-out processes
            if active_processes:
                to_remove = []
                for pid, (fn, p, start_time) in active_processes.items():
                    elapsed = time.time() - start_time
                    if p.poll() is not None:
                        # Process finished
                        if p.returncode == 0:
                            results.append((fn, "Success"))
                        else:
                            results.append((fn, f"Failed (Exit Code: {p.returncode})"))
                        to_remove.append(pid)
                    elif elapsed > AGENT_TIMEOUT:
                        # Timeout
                        log_event(f"Agent for {fn} timed out after {AGENT_TIMEOUT}s. Terminating.", level="warning")
                        p.terminate()
                        results.append((fn, "Timed Out"))
                        to_remove.append(pid)
                
                for pid in to_remove:
                    del active_processes[pid]
            
            if file_queue or active_processes:
                time.sleep(1) # Wait before next check
        
        create_summary_report(results)
        log_event(f"Completed processing {len(file_numbers)} file numbers. Summary report created.")
        sys.exit(0)
        
    # --- Worker Mode ---
    if not file_numbers:
        log_event("No valid file numbers found.", level="error")
        sys.exit(1)
        
    file_num = file_numbers[0]
    log_event(f"--- Starting Docket Agent for File: {file_num} (Headless: {is_headless}) ---")

    # --- Phase 1: Preparation ---
    case_number = get_case_var(file_num, "case_number")
    venue_county = get_case_var(file_num, "venue_county")
    
    missing_case = not case_number or str(case_number).lower() in ["null", "n/a", "none"]
    missing_venue = not venue_county or str(venue_county).lower() in ["null", "n/a", "none"]

    if missing_case or missing_venue:
        log_event(f"Missing critical data (Case #: {case_number}, Venue: {venue_county}) for {file_num}. Running Complaint Agent.")
        run_complaint_agent(file_num)
        
        # Reload
        case_number = get_case_var(file_num, "case_number")
        venue_county = get_case_var(file_num, "venue_county")

        # Validate again
        missing_case = not case_number or str(case_number).lower() in ["null", "n/a", "none"]
        if missing_case:
            log_event("Error: Complaint Agent failed to retrieve case number. Terminating.", level="error")
            sys.exit(1)
            
    log_event(f"Found Case Number: {case_number}. Proceeding to Phase 2.")

    # --- Phase 2: Data Retrieval ---
    if venue_county:
        venue_county = str(venue_county).lower().strip()
    
    selected_scraper = None
    use_interactive_mode = False

    extra_args = []
    
    if venue_county == "riverside":
        log_event("Venue is Riverside. Using Riverside Scraper.")
        selected_scraper = RIVERSIDE_SCRAPER_SCRIPT
        # Even if headless, Riverside might need some handling or it might just work if fully automated
        use_interactive_mode = not is_headless 
    elif venue_county == "san bernardino":
        log_event("Venue is San Bernardino. Using San Bernardino Scraper.")
        selected_scraper = SB_SCRAPER_SCRIPT
    elif venue_county == "kern":
        log_event("Venue is Kern. Using Kern Scraper.")
        selected_scraper = KERN_SCRAPER_SCRIPT
    elif venue_county == "san diego":
        log_event("Venue is San Diego. Using San Diego Scraper.")
        selected_scraper = SD_SCRAPER_SCRIPT
    elif venue_county == "orange":
        log_event("Venue is Orange County. Using Orange County Scraper.")
        selected_scraper = OC_SCRAPER_SCRIPT
        
        # Orange County requires filing year
        filing_date = get_case_var(file_num, "filing_date")
        if filing_date and "-" in str(filing_date):
            year = str(filing_date).split("-")[0]
            extra_args.append(year)
        else:
            current_year = datetime.datetime.now().strftime("%Y")
            log_event(f"Filing date ({filing_date}) missing/invalid for Orange County. Defaulting to {current_year}.", level="warning")
            extra_args.append(current_year)
            
    elif venue_county == "los angeles":
        log_event("Venue is Los Angeles. Using LA Scraper.")
        selected_scraper = SCRAPER_SCRIPT
    elif venue_county == "sacramento":
        log_event("Venue is Sacramento. Using Sacramento Scraper.")
        selected_scraper = SACRAMENTO_SCRAPER_SCRIPT
    else:
        # Prompt: "I want the agent to do this only if the venue_county variable is Los Angeles. 
        # However, if the venue_county variable is riverside..."
        log_event(f"Venue '{venue_county}' is not supported (LA, SB, Riverside, SD, or Orange only). Skipping download.", level="warning")
        sys.exit(1) # Changed from 0 to 1 to indicate it wasn't processed

    log_event(f"Launching scraper: {selected_scraper}...")
    
    # Run Scraper with Retries
    scraper_success = False
    max_scraper_attempts = 3
    for attempt in range(1, max_scraper_attempts + 1):
        try:
            log_event(f"Scraper attempt {attempt} of {max_scraper_attempts}...")
            # Prepare scraper arguments
            scraper_args = [sys.executable, selected_scraper, case_number] + extra_args
            
            # Logic: Orange County always headful. 
            # Others: Headless by default unless --headful is passed to docket.py
            if venue_county == "orange":
                scraper_args.append("--headful")
                # Also ensure we don't accidentally pass --headless if it was passed to docket.py
            else:
                if is_headful:
                    scraper_args.append("--headful")
                else:
                    scraper_args.append("--headless")

            # Run Scraper
            if use_interactive_mode:
                # Interactive mode: Do not capture output
                result = subprocess.run(scraper_args, check=False)
                if result.returncode == 0:
                    scraper_success = True
                    break
                else:
                    log_event(f"Scraper attempt {attempt} failed with exit code {result.returncode}.", level="warning")
            else:
                # Standard mode: Capture output
                result = subprocess.run(scraper_args, capture_output=True, text=True)
                
                if result.returncode == 0:
                    log_event(f"Scraper {selected_scraper} completed successfully.")
                    scraper_success = True
                    break
                else:
                    log_event(f"Scraper attempt {attempt} failed.\nSTDOUT: {result.stdout}\nSTDERR: {result.stderr}", level="warning")

        except Exception as e:
            log_event(f"Failed to execute Docket Scraper on attempt {attempt}: {e}", level="error")
        
        if attempt < max_scraper_attempts:
            log_event("Retrying scraper in 5 seconds...")
            time.sleep(5)
    
    # --- Post-Processing: Move File ---
    if scraper_success:
        # Find the generated PDF (docket_<case_number>_*.pdf)
        # We look for all candidates and pick the newest one to avoid issues with locked or old files.
        date_str = datetime.datetime.now().strftime("%Y.%m.%d")
        
        # Use case_number in glob for uniqueness
        file_pattern = f"docket_{case_number}_*.pdf"
        candidates = glob.glob(file_pattern)
        
        # Fallback to old pattern if case-specific not found (for backward compatibility during migration)
        if not candidates:
            candidates = glob.glob("docket_*.pdf")
            
        if not candidates:
            log_event(f"Error: Scraper reported success but no '{file_pattern}' found.", level="error")
            sys.exit(1)
        else:
            # Sort by modification time, newest first
            candidates.sort(key=os.path.getmtime, reverse=True)
            source_file = candidates[0]
            
            log_event(f"Found source file: {source_file}")

            case_path = get_case_path(file_num)
            if case_path:
                target_dir = os.path.join(case_path, "NOTES", "AI OUTPUT")
                if not os.path.exists(target_dir):
                    os.makedirs(target_dir)
                
                new_filename = f"Docket_{date_str}.pdf"
                target_path = os.path.join(target_dir, new_filename)
                
                try:
                    shutil.copy2(source_file, target_path)
                    os.remove(source_file)
                    log_event(f"COMPLETED: Docket downloaded and saved to {target_path}")
                    
                    # --- Phase 3: Extraction & Reporting ---
                    log_event("--- Phase 3: Extracting Hearing Data ---")
                    
                    docket_text = extract_text_from_file(target_path)
                    if docket_text:
                        prompt = """
                        Analyze this docket text.
                        1. Identify the 'Trial Date' if explicitly scheduled. Format: YYYY-MM-DD (or "None" if not found).
                        2. Identify 'Other Hearings' (any future scheduled hearings other than the trial). Summarize them briefly (e.g., "CMC on 2025-01-15, Motion on 2025-02-20"). If none, return "None".
                        
                        Return JSON:
                        {
                            "trial_date": "string",
                            "other_hearings": "string"
                        }
                        """
                        
                        json_str = call_gemini_api(prompt, docket_text)
                        if json_str:
                            try:
                                extracted = json.loads(clean_json_string(json_str))
                                trial_date = extracted.get("trial_date", "None")
                                other_hearings = extracted.get("other_hearings", "None")
                                
                                log_event(f"Extracted Trial Date: {trial_date}")
                                log_event(f"Extracted Other Hearings: {other_hearings}")
                                
                                save_case_var(file_num, "trial_date", trial_date)
                                save_case_var(file_num, "other_hearings", other_hearings)
                                
                                # Prepare data for Master Status
                                master_data = {
                                    "file_number": file_num,
                                    "plaintiffs": get_case_var(file_num, "plaintiffs"),
                                    "case_number": get_case_var(file_num, "case_number"),
                                    "venue_county": get_case_var(file_num, "venue_county"),
                                    "trial_date": trial_date,
                                    "other_hearings": other_hearings
                                }
                                
                                log_event("Updating Master Case Status Document...")
                                update_master_status(file_num, master_data, target_path)
                                
                            except json.JSONDecodeError:
                                log_event("Failed to parse Gemini response for hearings.", level="error")
                        
                        # --- Phase 4: Procedural History ---
                        log_event("--- Phase 4: Generating Procedural History ---")
                        
                        # Try to get cached text from Complaint Agent first
                        complaint_text = get_case_var(file_num, "hidden_complaint_text")
                        
                        if complaint_text:
                            log_event("Using cached complaint text from hidden variable.")
                        else:
                            log_event("Cached complaint text not found. Attempting to extract from file.")
                            complaint_file = find_complaint_file(case_path)
                            if complaint_file:
                                log_event(f"Found Complaint for Procedural History: {complaint_file}")
                                complaint_text = extract_text_from_file(complaint_file)
                            else:
                                log_event("Complaint file not found for Procedural History context.", level="warning")

                        if os.path.exists(PROCEDURAL_HIST_PROMPT_FILE):
                            with open(PROCEDURAL_HIST_PROMPT_FILE, 'r', encoding='utf-8') as f:
                                ph_prompt = f.read()
                            
                            combined_context = f"[CASE DOCKET TEXT]\n{docket_text}\n\n[COMPLAINT TEXT]\n{complaint_text}"
                            
                            log_event("Calling AI for Procedural History (Docket + Complaint)...")
                            ph_result = call_gemini_api(ph_prompt, combined_context, ["gemini-3-flash-preview", "gemini-2.5-flash"])
                            
                            if ph_result:
                                save_case_var(file_num, "procedural_history", ph_result)
                                log_event("Generated Procedural History.")
                            else:
                                log_event("AI failed to generate Procedural History.", level="warning")
                        else:
                            log_event(f"Prompt file missing: {PROCEDURAL_HIST_PROMPT_FILE}", level="warning")

                        # Final update to variables.docx with all accumulated data
                        log_event("Updating variables.docx with all case data...")
                        json_file = os.path.join(GEMINI_DATA_DIR, f"{file_num}.json")
                        if os.path.exists(json_file):
                            try:
                                with open(json_file, 'r') as f:
                                    all_vars = json.load(f)
                                update_variables_docx(case_path, all_vars)
                                log_event("Successfully updated variables.docx")
                            except Exception as e:
                                log_event(f"Error updating variables.docx: {e}", level="error")
                    else:
                        log_event("Failed to extract text from saved Docket PDF.", level="error")
                        sys.exit(1)

                except Exception as e:
                    log_event(f"Error saving file to {target_path}: {e}", level="error")
                    sys.exit(1)
            else:
                log_event("Could not resolve case path. Leaving file in current directory.", level="warning")
    else:
        log_event("Skipping extraction/analysis due to scraper failure.", level="warning")
        sys.exit(1)

if __name__ == "__main__":
    main()
