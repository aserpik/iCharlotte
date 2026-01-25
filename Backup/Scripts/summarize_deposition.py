import os
import sys
import logging
import datetime
import re
import subprocess
import google.generativeai as genai
from docx import Document
from docx.shared import Pt, Inches
from pypdf import PdfReader
try:
    import pytesseract
    from pdf2image import convert_from_path
    OCR_AVAILABLE = True
    
    # Windows-specific path configuration
    if os.name == 'nt':
        # Tesseract Path
        tesseract_path = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
        if os.path.exists(tesseract_path):
            pytesseract.pytesseract.tesseract_cmd = tesseract_path
        
        # Poppler Path
        poppler_path = r"C:\Program Files\poppler\Library\bin"
        if not os.path.exists(poppler_path):
             poppler_path = None
        
        POPPLER_PATH = poppler_path
    else:
        POPPLER_PATH = None

except ImportError:
    OCR_AVAILABLE = False
    POPPLER_PATH = None


# --- Configuration ---
LOG_FILE = r"C:\GeminiTerminal\Summarize_Deposition_activity.log"
PROMPT_FILE = r"C:\GeminiTerminal\Scripts\SUMMARIZE_DEPOSITION_PROMPT.txt"

# Set up logging
logging.basicConfig(
    filename=LOG_FILE,
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    datefmt='%Y-%m-%d %H:%M:%S'
)

def log_event(message, level="info"):
    print(message)
    sys.stdout.flush()
    if level == "info":
        logging.info(message)
    elif level == "error":
        logging.error(message)
    elif level == "warning":
        logging.warning(message)

def extract_text(file_path):
    """Extracts text from PDF, DOCX, or plain text files."""
    log_event(f"Extracting text from: {file_path}")
    ext = os.path.splitext(file_path)[1].lower()
    text = ""

    try:
        if ext == ".pdf":
            reader = PdfReader(file_path)
            
            def get_page_image(page_index):
                try:
                    images = convert_from_path(file_path, first_page=page_index+1, last_page=page_index+1, poppler_path=POPPLER_PATH)
                    return images[0] if images else None
                except Exception as e:
                    log_event(f"Error converting page {page_index+1} to image: {e}", level="warning")
                    return None

            for i, page in enumerate(reader.pages):
                page_text = page.extract_text() or ""
                if len(page_text.strip()) < 50:
                    if OCR_AVAILABLE:
                        image = get_page_image(i)
                        if image:
                            try:
                                ocr_page_text = pytesseract.image_to_string(image)
                                text += ocr_page_text + "\n"
                            except:
                                text += page_text + "\n"
                        else:
                             text += page_text + "\n"
                    else:
                        text += page_text + "\n"
                else:
                    text += page_text + "\n"

        elif ext == ".docx":
            doc = Document(file_path)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
        else:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
        
        return text

    except Exception as e:
        log_event(f"Error extracting text from {file_path}: {e}", level="error")
        return None

def call_gemini(prompt, text):
    """Calls Gemini API with fallback logic."""
    api_key = os.environ.get("GEMINI_API_KEY")
    if not api_key:
        log_event("Error: GEMINI_API_KEY environment variable not set.", level="error")
        return None

    genai.configure(api_key=api_key)
    full_prompt = f"{prompt}\n\nDOCUMENT CONTENT:\n{text}"
    model_sequence = ["gemini-3-pro-preview", "gemini-2.5-pro"]

    for model_name in model_sequence:
        log_event(f"Attempting to use model: {model_name}")
        try:
            model = genai.GenerativeModel(model_name)
            response = model.generate_content(full_prompt)
            if response and response.text:
                return response.text
        except BaseException as e:
            log_event(f"Failed with model {model_name}: {e}", level="warning")
            continue
    return None

def add_markdown_to_doc(doc, content):
    """Parses basic Markdown and adds it to the docx Document with specific formatting."""
    lines = content.split('\n')
    for line in lines:
        stripped = line.strip()
        if not stripped: continue
        
        if stripped.startswith('#'):
            text = stripped.lstrip('#').strip()
            if not text.endswith('.'): text += "."
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.first_line_indent = Inches(0.5)
            run = p.add_run(text)
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            continue
        
        if stripped.startswith('* ') or stripped.startswith('- '):
            text = stripped[2:].strip()
            p = doc.add_paragraph(text, style='List Bullet')
            p.paragraph_format.line_spacing = 1.0
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
            continue
        
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.first_line_indent = Inches(0.5)
        parts = re.split(r'(\*\*.*\*\*)', stripped)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
            else:
                run = p.add_run(part)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

def save_to_docx(content, output_path, title_text):
    """Saves the content to a DOCX file with specific subheading formatting."""
    base_name, ext = os.path.splitext(output_path)
    counter = 1
    current_output_path = output_path
    
    while True:
        try:
            if os.path.exists(current_output_path):
                doc = Document(current_output_path)
            else:
                doc = Document()

            style = doc.styles['Normal']
            style.font.name = 'Times New Roman'
            style.font.size = Pt(12)
            style.paragraph_format.line_spacing = 1.0
            
            next_num = 1
            for para in doc.paragraphs:
                if re.match(r'^\d+\.\tDeposition of', para.text):
                    next_num += 1

            deponent_name = title_text
            for skip in [".pdf", ".docx", "deposition", "depo", "transcript"]:
                deponent_name = deponent_name.replace(skip, "").replace(skip.upper(), "")
            deponent_name = deponent_name.strip(" _-")

            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.left_indent = Inches(1.5)
            p.paragraph_format.first_line_indent = Inches(-0.5)
            tab_stops = p.paragraph_format.tab_stops
            tab_stops.add_tab_stop(Inches(1.5))

            run_num = p.add_run(f"{next_num}.\t")
            run_num.bold = True
            run_num.font.name = 'Times New Roman'
            run_num.font.size = Pt(12)

            run_title = p.add_run(f"Deposition of {deponent_name}")
            run_title.bold = True
            run_title.underline = True
            run_title.font.name = 'Times New Roman'
            run_title.font.size = Pt(12)
            
            add_markdown_to_doc(doc, content)
            doc.save(current_output_path)
            return True

        except (PermissionError, IOError):
            counter += 1
            current_output_path = f"{base_name} v.{counter}{ext}"
            if counter > 10: return False
        except Exception as e:
            log_event(f"Error saving to DOCX: {e}", level="error")
            return False

def main():
    if len(sys.argv) < 2:
        log_event("Error: No file path provided.", level="error")
        sys.exit(1)

    # Use sys.argv[1:] directly to allow space-separated file paths.
    # Users must quote paths containing spaces.
    file_paths = sys.argv[1:]
    
    # --- Dispatcher Mode ---
    if len(file_paths) > 1:
        log_event(f"Detected multiple file paths: {len(file_paths)} items. Launching separate agents...")
        for path in file_paths:
            # We must quote paths with spaces when re-launching
            quoted_path = f'"{path}"' if ' ' in path and not path.startswith('"') else path
            log_event(f"Spawning agent for: {path}")
            try:
                # Spawn independent process
                if os.name == 'nt':
                    # Use CREATE_NO_WINDOW (0x08000000) for headless execution
                    subprocess.Popen([sys.executable, sys.argv[0], path], creationflags=0x08000000)
                else:
                    subprocess.Popen([sys.executable, sys.argv[0], path])
            except Exception as e:
                log_event(f"Failed to spawn agent for {path}: {e}", level="error")
        
        print(f"Launched {len(file_paths)} summarize deposition agents.")
        sys.exit(0)

    # --- Worker Mode ---
    # At this point, we assume we are processing a SINGLE path (which might be a directory or file)
    
    if not file_paths:
        log_event("No valid file paths found.", level="error")
        sys.exit(1)
        
    input_path = file_paths[0]
    
    # Remove surrounding quotes if present (common in CLI args)
    if (input_path.startswith('"') and input_path.endswith('"')) or (input_path.startswith("'") and input_path.endswith("'")):
        input_path = input_path[1:-1]
    
    # Handle absolute path conversion
    input_path = os.path.abspath(input_path)
    if not os.path.exists(input_path):
        sys.exit(1)

    if os.path.isdir(input_path):
        files_to_process = []
        for root, _, files in os.walk(input_path):
            for file in files:
                if file.lower().endswith(('.pdf', '.docx')):
                    if "Deposition_summaries" in file: continue
                    files_to_process.append(os.path.join(root, file))
        
        for file_path in files_to_process:
            try:
                subprocess.run([sys.executable, sys.argv[0], file_path], check=True)
            except: pass
        sys.exit(0)
    
    text = extract_text(input_path)
    if not text: sys.exit(1)

    try:
        with open(PROMPT_FILE, "r", encoding="utf-8") as f:
            prompt_instruction = f.read()
    except Exception as e:
        log_event(f"Error reading prompt file {PROMPT_FILE}: {e}", level="error")
        sys.exit(1)

    summary = call_gemini(prompt_instruction, text)
    if not summary: sys.exit(1)

    parts = input_path.split(os.sep)
    output_dir = None
    for i, part in enumerate(parts):
        if part.lower() == "current clients":
            if i + 2 < len(parts):
                case_root_parts = parts[:i+3]
                output_dir = os.path.join(os.sep.join(case_root_parts), "NOTES", "AI OUTPUT")
            break
            
    if not output_dir:
        for i in range(len(parts) - 1, -1, -1):
            if parts[i].upper() == "NOTES":
                output_dir = os.path.join(os.sep.join(parts[:i+1]), "AI OUTPUT")
                break
    
    if not output_dir:
        input_dir = os.path.dirname(input_path)
        parent_dir = os.path.dirname(input_dir)
        output_dir = os.path.join(parent_dir, "NOTES", "AI OUTPUT")

    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    output_file = os.path.join(output_dir, "Deposition_summaries.docx")
    save_to_docx(summary, output_file, os.path.basename(input_path))

if __name__ == '__main__':
    main()
