import os
import sys
import logging
import datetime
import re
import subprocess
import google.generativeai as genai
from docx import Document
from docx.shared import Pt, Inches
from pypdf import PdfReader
try:
    import pytesseract
    from pdf2image import convert_from_path
    OCR_AVAILABLE = True
    
    # Windows-specific path configuration
    if os.name == 'nt':
        # Tesseract Path
        tesseract_path = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
        if os.path.exists(tesseract_path):
            pytesseract.pytesseract.tesseract_cmd = tesseract_path
        
        # Poppler Path
        poppler_path = r"C:\Program Files\poppler\Library\bin"
        if not os.path.exists(poppler_path):
             poppler_path = None
        
        POPPLER_PATH = poppler_path
    else:
        POPPLER_PATH = None

except ImportError:
    OCR_AVAILABLE = False
    POPPLER_PATH = None


# --- Configuration ---
LOG_FILE = r"C:\GeminiTerminal\Med_Chron_activity.log"
PROMPT_FILE = r"C:\GeminiTerminal\Scripts\MED_CHRON_PROMPT.txt"

# Set up logging
logging.basicConfig(
    filename=LOG_FILE,
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    datefmt='%Y-%m-%d %H:%M:%S'
)

def log_event(message, level="info"):
    print(message)
    sys.stdout.flush()
    if level == "info":
        logging.info(message)
    elif level == "error":
        logging.error(message)
    elif level == "warning":
        logging.warning(message)

def extract_text(file_path):
    """Extracts text from PDF, DOCX, or plain text files."""
    log_event(f"Extracting text from: {file_path}")
    ext = os.path.splitext(file_path)[1].lower()
    text = ""

    try:
        if ext == ".pdf":
            reader = PdfReader(file_path)
            total_pages = len(reader.pages)
            log_event(f"PDF has {total_pages} pages.")
            
            def get_page_image(page_index):
                try:
                    images = convert_from_path(file_path, first_page=page_index+1, last_page=page_index+1, poppler_path=POPPLER_PATH)
                    return images[0] if images else None
                except Exception as e:
                    log_event(f"Error converting page {page_index+1} to image: {e}", level="warning")
                    return None

            for i, page in enumerate(reader.pages):
                page_text = page.extract_text() or ""
                
                if len(page_text.strip()) < 50:
                    log_event(f"Page {i+1} has insufficient text. Attempting OCR...", level="warning")
                    if OCR_AVAILABLE:
                        image = get_page_image(i)
                        if image:
                            try:
                                ocr_page_text = pytesseract.image_to_string(image)
                                if len(ocr_page_text.strip()) > len(page_text.strip()):
                                    text += ocr_page_text + "\n"
                                else:
                                    text += page_text + "\n"
                            except Exception:
                                text += page_text + "\n"
                        else:
                             text += page_text + "\n"
                    else:
                        text += page_text + "\n"
                else:
                    text += page_text + "\n"

        elif ext == ".docx":
            doc = Document(file_path)
            # Only iterate over paragraphs to ignore tables
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
        else:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
        
        if not text.strip():
            log_event(f"Warning: Extracted text is empty for {file_path}", level="warning")
            return None
        
        return text

    except Exception as e:
        log_event(f"Error extracting text from {file_path}: {e}", level="error")
        return None

def call_gemini(prompt, text):
    """Calls Gemini API."""
    api_key = os.environ.get("GEMINI_API_KEY")
    if not api_key:
        log_event("Error: GEMINI_API_KEY environment variable not set.", level="error")
        return None

    genai.configure(api_key=api_key)

    full_prompt = f"{prompt}\n\nDOCUMENT CONTENT:\n{text}"
    
    model_sequence = [
        "gemini-3-flash-preview", 
        "gemini-2.5-flash"
    ]

    for model_name in model_sequence:
        log_event(f"Attempting to use model: {model_name}")
        try:
            model = genai.GenerativeModel(model_name)
            response = model.generate_content(full_prompt)
            if response and response.text:
                log_event(f"Success with model: {model_name}")
                return response.text
        except Exception as e:
            log_event(f"Failed with model {model_name}: {e}", level="warning")
            continue
    
    log_event("Error: All model attempts failed.", level="error")
    return None

def add_markdown_to_doc(doc, content):
    """Parses basic Markdown and applies formatting."""
    lines = content.split('\n')
    active_paragraph = None
    
    # Regex to identify dates at the beginning of a paragraph/sentence
    # Matches: (Optional "On ") (Date String)
    # Date String: Month DD, YYYY
    date_pattern = re.compile(r"^(On )?((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]* \d{1,2},? \d{4})")

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        
        if stripped.startswith('#'):
            text = stripped.lstrip('#').strip()
            if not text.endswith('.'):
                text += "."
            active_paragraph = doc.add_paragraph()
            run = active_paragraph.add_run(text + " ")
            run.bold = True
            continue
        
        if stripped.startswith('* ') or stripped.startswith('- '):
            text = stripped[2:].strip()
            doc.add_paragraph(text, style='List Bullet')
            active_paragraph = None
            continue
        
        # Regular paragraph
        if active_paragraph:
            p = active_paragraph
        else:
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Inches(0.5)
        
        # Check for date at start of line
        match = date_pattern.match(stripped)
        if match:
            on_prefix = match.group(1) # "On " or None
            date_str = match.group(2)  # "January 1, 2024"
            
            total_match_len = len(match.group(0))
            remaining_text = stripped[total_match_len:]
            
            if on_prefix:
                p.add_run(on_prefix)
            
            run = p.add_run(date_str)
            run.underline = True
            
            # Process remaining text for bold markers
            parts = re.split(r'(\**.*?\**)', remaining_text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
        else:
            # Standard markdown parsing
            parts = re.split(r'(\**.*?\**)', stripped)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
        
        active_paragraph = None

def extract_provider_from_filename(filename):
    """Extracts provider name from filename based on patterns."""
    # Pattern 1: 12345-001_ PROVIDER NAME (1).pdf
    # Split by underscore
    parts = filename.split('_')
    if len(parts) > 1:
        # Take the part after the first underscore
        potential_name = parts[1]
        # Remove extension
        potential_name = os.path.splitext(potential_name)[0]
        # Remove trailing parentheses like (1)
        potential_name = re.sub(r'\s*\(\d+\)$', '', potential_name)
        return potential_name.strip()
    
    # Fallback: Use filename without extension and sanitize
    name = os.path.splitext(filename)[0]
    return name.strip()

def sanitize_filename(name):
    """Sanitizes a string to be safe for filenames."""
    # Remove invalid characters
    name = re.sub(r'[<>:"/\\|?*]', '', name)
    # Trim whitespace
    name = name.strip()
    return name

def save_to_docx(content, output_dir, provider_name, original_filename):
    """Saves to DOCX with fixed filename med_chron.docx."""
    
    # User requested specific filename
    filename = "med_chron.docx"
    output_path = os.path.join(output_dir, filename)

    try:
        if os.path.exists(output_path):
             # If exists, we append
             doc = Document(output_path)
             doc.add_page_break()
             log_event(f"Appending to existing file: {output_path}")
        else:
             doc = Document()
             log_event(f"Creating new file: {output_path}")

        # Styles
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(12)
        style.paragraph_format.line_spacing = 1.0
        
        for i in range(1, 10):
            if f'Heading {i}' in doc.styles:
                h = doc.styles[f'Heading {i}']
                h.font.name = 'Times New Roman'
                h.font.size = Pt(12)
                h.paragraph_format.line_spacing = 1.0

        for s in ['List Bullet', 'List Number']:
            if s in doc.styles:
                l = doc.styles[s]
                l.font.name = 'Times New Roman'
                l.font.size = Pt(12)
                l.paragraph_format.line_spacing = 1.0

        # Title
        p = doc.add_paragraph()
        run = p.add_run(f"Medical Record Chronology - {provider_name}")
        run.bold = True
        run.underline = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

        doc.add_paragraph(f"Source: {original_filename}")
        doc.add_paragraph(f"Generated on: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        
        add_markdown_to_doc(doc, content)
            
        doc.save(output_path)
        log_event(f"Saved to: {output_path}")
        return True

    except Exception as e:
        log_event(f"Error saving DOCX: {e}", level="error")
        return False

def filter_content(text):
    """Filters text to only include content under specific headings."""
    headings = [
        "BRIEF SYNOPSIS OF PRE-INJURY MEDICAL RECORD:",
        "BRIEF SYNOPSIS OF POST-INJURY MEDICAL RECORD:"
    ]
    
    indices = []
    for h in headings:
        # Case insensitive search might be safer, but user used uppercase.
        # Strict adherence to user request first.
        idx = text.find(h)
        if idx != -1:
            indices.append((idx, h))
    
    # Sort by position in text
    indices.sort(key=lambda x: x[0])
    
    if not indices:
        log_event("Target headings not found. Processing skipped.", level="warning")
        return None

    filtered_chunks = []
    for i, (start_idx, header) in enumerate(indices):
        content_start = start_idx + len(header)
        
        # Determine end of this section
        if i + 1 < len(indices):
            content_end = indices[i+1][0]
        else:
            content_end = len(text)
            
        chunk = text[content_start:content_end].strip()
        filtered_chunks.append(f"{header}\n{chunk}")
        
    return "\n\n".join(filtered_chunks)

def main():
    if len(sys.argv) < 2:
        log_event("Error: No file path provided.", level="error")
        sys.exit(1)

    # Attempt to handle unquoted paths with spaces by joining all arguments
    input_path = " ".join(sys.argv[1:])
    
    # If the joined path doesn't exist, but the first argument does, 
    # it might be a single file with tricky chars or we're in a directory where 
    # sys.argv[1] was actually correct and the rest were garbage.
    if not os.path.exists(input_path) and os.path.exists(sys.argv[1]):
        input_path = sys.argv[1]
    
    # Handle absolute path conversion
    input_path = os.path.abspath(input_path)

    if not os.path.exists(input_path):
        log_event(f"Error: File not found: {input_path}", level="error")
        sys.exit(1)

    log_event(f"--- Starting Med Chron Agent for: {input_path} ---")

    if os.path.isdir(input_path):
        log_event(f"Input is a directory. Scanning: {input_path}")
        files_to_process = []
        for root, _, files in os.walk(input_path):
            for file in files:
                if file.lower().endswith(('.pdf', '.docx')):
                    if "med_chron" in file.lower(): continue
                    files_to_process.append(os.path.join(root, file))
        
        if not files_to_process:
            log_event("No suitable files found.")
            sys.exit(0)
        
        for file_path in files_to_process:
            try:
                subprocess.run([sys.executable, sys.argv[0], file_path], check=True)
            except subprocess.CalledProcessError as e:
                log_event(f"Subprocess failed for {file_path}: {e}", level="error")
        sys.exit(0)
    
    # Single File
    text = extract_text(input_path)
    if not text:
        sys.exit(1)
        
    # Apply Filtering
    filtered_text = filter_content(text)
    if not filtered_text:
        log_event("No valid content found under specified headings.")
        sys.exit(0)

    try:
        with open(PROMPT_FILE, "r", encoding="utf-8") as f:
            prompt_instruction = f.read()
    except Exception as e:
        log_event(f"Error reading prompt file: {e}", level="error")
        sys.exit(1)

    # Determine Provider Name from Filename
    filename = os.path.basename(input_path)
    provider_name = extract_provider_from_filename(filename)
    log_event(f"Identified Provider from Filename: {provider_name}")

    final_content = call_gemini(prompt_instruction, filtered_text)
    if not final_content:
        sys.exit(1)

    # Determine Output Directory
    parts = input_path.split(os.sep)
    output_dir = None
    for i, part in enumerate(parts):
        if part.lower() == "current clients":
            if i + 2 < len(parts):
                output_dir = os.sep.join(parts[:i+3] + ["NOTES", "AI OUTPUT"])
            break
            
    if not output_dir:
        for i in range(len(parts) - 1, -1, -1):
            if parts[i].upper() == "NOTES":
                output_dir = os.path.join(os.sep.join(parts[:i+1]), "AI OUTPUT")
                break
    
    if not output_dir:
        input_dir = os.path.dirname(input_path)
        parent_dir = os.path.dirname(input_dir)
        output_dir = os.path.join(parent_dir, "NOTES", "AI OUTPUT")

    if not os.path.exists(output_dir):
        try:
            os.makedirs(output_dir)
        except Exception:
            pass

    save_to_docx(final_content, output_dir, provider_name, filename)
    log_event("--- Agent Finished ---")

if __name__ == '__main__':
    main()
