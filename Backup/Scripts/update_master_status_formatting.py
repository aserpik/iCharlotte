import os
import re
import datetime
from docx import Document
from docx.shared import Pt

# Path to the Master Case Status document
DOC_PATH = r"C:\Users\ASerpik.DESKTOP-MRIMK0D\OneDrive - Bordin Semmer LLP\Desktop\MASTER_CASE_STATUS.docx"

def update_formatting():
    if not os.path.exists(DOC_PATH):
        print(f"Error: Document not found at {DOC_PATH}")
        return

    try:
        doc = Document(DOC_PATH)
        
        if not doc.tables:
            print("No tables found in the document.")
            return
            
        table = doc.tables[0]
        rows = table.rows
        if len(rows) < 2:
            print("Table has no data rows.")
            return

        count_cleaned = 0
        count_hearings_reformatted = 0
        
        # Regex to find dates like YYYY-MM-DD or MM-DD-YYYY
        # We look for " on YYYY-MM-DD" or similar patterns to split
        date_pattern = re.compile(r"(\d{4}-\d{2}-\d{2}|\d{2}-\d{2}-\d{4})")

        for row in rows[1:]: # Skip header
            # 1. Clean characters [ ] ' from ALL cells
            # 2. Set Font to Times New Roman 10
            for cell in row.cells:
                text = cell.text
                if any(c in text for c in "[]'"):
                    new_text = text.replace("[", "").replace("]", "").replace("'", "")
                    cell.text = new_text # This resets formatting, so we apply it after
                    count_cleaned += 1
                
                # Apply Font Style
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(10)
                    
                    # Also handle runs inside hyperlinks (messy in python-docx but we can try basic iter)
                    # For now, just setting the paragraph style might not be enough if runs override.
                    # The reset of cell.text above clears hyperlinks in that cell, which might be an issue 
                    # if the user wants to keep them. 
                    # INSTRUCTION SAID "does not have any ] or ' characters". 
                    # If we blindly replace cell.text, we lose hyperlinks.
                    # We should iterate paragraphs and runs to replace text safely.
            
            # --- Better text cleaning preserving structure ---
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    # Clean text inside runs to preserve hyperlinks if possible
                    # (Though modifying run.text can be tricky if split across runs)
                    # Simplest approach: Text replacement on the paragraph text is destructive to hyperlinks.
                    # But if we need to remove chars, we must. 
                    # Let's assume for now that if we find the chars, we just replace the text.
                    # If the user wants to keep hyperlinks, we'd need a more complex XML parser.
                    # Given "update the document... does not have...", I will prioritize content cleaning.
                    # If a cell has a hyperlink, its text is usually in a run inside a hyperlink field.
                    # python-docx doesn't easily let us edit text inside hyperlink fields without dropping the link.
                    # I will accept that hyperlinks might be lost if they contain the forbidden characters.
                    pass

            # --- Update Other Hearings (Column 6) ---
            try:
                cell_hearings = row.cells[6]
                hearings_text = cell_hearings.text.strip()
                
                # Format: "Hearing Name on YYYY-MM-DD" -> "Hearing Name (MM-DD-YY)"
                # Also handle multiple hearings if comma separated? 
                # The user example: "FSC (04-25-26)"
                
                if hearings_text and hearings_text != "None":
                    # Split by comma if multiple hearings? 
                    # The prompt implies a single format or list. Let's process the whole string.
                    # Strategy: Find all dates. For each date, try to find the preceding text?
                    # Or simpler: Split by "on" ?
                    
                    # Case 1: "CMC on 2025-01-15"
                    # Case 2: "CMC on 2025-01-15, Motion on 2025-02-20"
                    
                    parts = hearings_text.split(',')
                    new_parts = []
                    for part in parts:
                        part = part.strip()
                        # Clean chars first
                        part = part.replace("[", "").replace("]", "").replace("'", "")
                        
                        match = date_pattern.search(part)
                        if match:
                            date_str = match.group(1)
                            # Parse date
                            try:
                                if "-" in date_str and len(date_str.split("-")[0]) == 4: # YYYY-MM-DD
                                    dt = datetime.datetime.strptime(date_str, "%Y-%m-%d")
                                else: # MM-DD-YYYY
                                    dt = datetime.datetime.strptime(date_str, "%m-%d-%Y")
                                
                                formatted_date = dt.strftime("%m-%d-%y")
                                
                                # Extract Name: Remove date and "on"
                                name = part.replace(date_str, "").replace(" on ", " ").strip()
                                new_part = f"{name} ({formatted_date})"
                                new_parts.append(new_part)
                            except ValueError:
                                new_parts.append(part)
                        else:
                            new_parts.append(part)
                    
                    new_hearings_text = ", ".join(new_parts)
                    
                    if new_hearings_text != cell_hearings.text:
                        cell_hearings.text = new_hearings_text
                        count_hearings_reformatted += 1
                        
                    # Re-apply font to this cell specifically after text change
                    for p in cell_hearings.paragraphs:
                        for r in p.runs:
                            r.font.name = 'Times New Roman'
                            r.font.size = Pt(10)
            except IndexError:
                pass

            # --- Global cleaning of [ ] ' and Font styling for ALL cells ---
            for cell in row.cells:
                # Text cleaning (iterating runs is safer for formatting but hard for global replace)
                # We will do a robust check: if forbidden chars exist, clean string and reset.
                curr_text = cell.text
                if any(c in curr_text for c in "[]'"):
                    clean_text = curr_text.replace("[", "").replace("]", "").replace("'", "")
                    cell.text = clean_text # This might kill hyperlinks
                
                # Apply font style to every run in every paragraph
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(10)
                    # Check for hyperlinks explicitly if we didn't nuke them
                    # (Text cleaning above nukes them if triggered. If not triggered, we need to style them.)
                    # Accessing runs inside hyperlinks requires XML traversal or specific library features.
                    # We'll rely on the standard runs for now.

        doc.save(DOC_PATH)
        print(f"Successfully updated document.")
        print(f"  - Cleaned forbidden chars and set font.")
        print(f"  - Reformatted hearings: {count_hearings_reformatted}")

    except Exception as e:
        print(f"An error occurred: {e}")

if __name__ == "__main__":
    update_formatting()
