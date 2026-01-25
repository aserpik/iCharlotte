import os
import sys
import logging
import datetime
import re
import subprocess
import google.generativeai as genai
from docx import Document
from docx.shared import Pt, Inches
from pypdf import PdfReader
try:
    import pytesseract
    from pdf2image import convert_from_path
    OCR_AVAILABLE = True
    
    # Windows-specific path configuration
    if os.name == 'nt':
        # Tesseract Path
        tesseract_path = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
        if os.path.exists(tesseract_path):
            pytesseract.pytesseract.tesseract_cmd = tesseract_path
        
        # Poppler Path
        # We found it at C:\Program Files\poppler\Library\bin\pdftoppm.exe
        # pdf2image needs the folder containing the binaries
        poppler_path = r"C:\Program Files\poppler\Library\bin"
        if not os.path.exists(poppler_path):
             # Fallback attempt or standard location
             poppler_path = None
        
        # We will use this variable in extract_text
        POPPLER_PATH = poppler_path
    else:
        POPPLER_PATH = None

except ImportError:
    OCR_AVAILABLE = False
    POPPLER_PATH = None


# --- Configuration ---
# Hardcoded log path as per instructions
LOG_FILE = r"C:\GeminiTerminal\Summarize_Discovery_activity.log"
PROMPT_FILE = r"C:\GeminiTerminal\Scripts\SUMMARIZE_DISCOVERY_PROMPT.txt"

# Set up logging
logging.basicConfig(
    filename=LOG_FILE,
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    datefmt='%Y-%m-%d %H:%M:%S'
)

def log_event(message, level="info"):
    print(message)  # Also print to stdout for debugging if run manually
    sys.stdout.flush() # Force flush
    if level == "info":
        logging.info(message)
    elif level == "error":
        logging.error(message)
    elif level == "warning":
        logging.warning(message)

def extract_text(file_path):
    """Extracts text from PDF, DOCX, or plain text files. Falls back to OCR for specific PDF pages if text is insufficient."""
    log_event(f"Extracting text from: {file_path}")
    ext = os.path.splitext(file_path)[1].lower()
    text = ""

    try:
        if ext == ".pdf":
            reader = PdfReader(file_path)
            total_pages = len(reader.pages)
            log_event(f"PDF has {total_pages} pages.")
            
            # Helper to convert specific page index to image if needed
            def get_page_image(page_index):
                try:
                    # convert_from_path returns a list of images. We need to fetch just the specific page.
                    # 'first_page' and 'last_page' are 1-based indices in pdf2image
                    # Pass poppler_path if explicitly set (mainly for Windows)
                    images = convert_from_path(file_path, first_page=page_index+1, last_page=page_index+1, poppler_path=POPPLER_PATH)
                    return images[0] if images else None
                except Exception as e:
                    log_event(f"Error converting page {page_index+1} to image: {e}", level="warning")
                    return None

            for i, page in enumerate(reader.pages):
                page_text = page.extract_text() or ""
                
                # Check if this specific page has meaningful text (threshold: 50 chars)
                if len(page_text.strip()) < 50:
                    log_event(f"Page {i+1} has insufficient text ({len(page_text.strip())} chars). Attempting OCR...", level="warning")
                    
                    if OCR_AVAILABLE:
                        image = get_page_image(i)
                        if image:
                            try:
                                ocr_page_text = pytesseract.image_to_string(image)
                                if len(ocr_page_text.strip()) > len(page_text.strip()):
                                    log_event(f"OCR successful for page {i+1}. Extracted {len(ocr_page_text)} chars.")
                                    text += ocr_page_text + "\n"
                                else:
                                    log_event(f"OCR for page {i+1} did not yield better results. Using original.", level="warning")
                                    text += page_text + "\n"
                            except Exception as ocr_e:
                                log_event(f"OCR failed for page {i+1}: {ocr_e}. Using original.", level="error")
                                text += page_text + "\n"
                        else:
                             log_event(f"Could not render image for page {i+1}. Using original.", level="warning")
                             text += page_text + "\n"
                    else:
                        log_event(f"OCR not available. Skipping OCR for page {i+1}.", level="warning")
                        text += page_text + "\n"
                else:
                    # Page text is sufficient
                    text += page_text + "\n"

        elif ext == ".docx":
            doc = Document(file_path)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
        else:
            # Assume plain text for other extensions
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
        
        if not text.strip():
            log_event(f"Warning: Extracted text is empty for {file_path}", level="warning")
            return None
        
        log_event(f"Successfully extracted {len(text)} characters total.")
        log_event(f"Snippet: {text[:100].replace(chr(10), ' ')}...") # Log snippet
        return text

    except Exception as e:
        log_event(f"Error extracting text from {file_path}: {e}", level="error")
        return None

def call_gemini(prompt, text):
    """Calls Gemini API with fallback logic."""
    api_key = os.environ.get("GEMINI_API_KEY")
    if not api_key:
        log_event("Error: GEMINI_API_KEY environment variable not set.", level="error")
        return None

    genai.configure(api_key=api_key)

    full_prompt = f"{prompt}\n\nDOCUMENT CONTENT:\n{text}"
    
    models_to_try = ["gemini-1.5-pro", "gemini-2.0-flash-exp", "gemini-1.5-flash"]
    
    model_sequence = [
        "gemini-3-flash-preview", 
        "gemini-2.5-flash", 
        "gemini-1.5-pro", 
        "gemini-1.5-flash"
    ]

    for model_name in model_sequence:
        log_event(f"Attempting to use model: {model_name}")
        try:
            model = genai.GenerativeModel(model_name)
            response = model.generate_content(full_prompt)
            if response and response.text:
                log_event(f"Success with model: {model_name}")
                return response.text
        except BaseException as e: # Catch ALL exceptions including system exits
            log_event(f"Failed with model {model_name}: {e}", level="warning")
            continue
    
    log_event("Error: All model attempts failed.", level="error")
    return None

def add_markdown_to_doc(doc, content):
    """Parses basic Markdown and adds it to the docx Document."""
    lines = content.split('\n')
    
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        
        # Headings: Convert to bold text at start of paragraph, ending with period
        if stripped.startswith('#'):
            text = stripped.lstrip('#').strip()
            if not text.endswith('.'):
                text += "."
            
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.first_line_indent = Inches(0.5)
            run = p.add_run(text)
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            continue
        
        # List items
        if stripped.startswith('* ') or stripped.startswith('- '):
            text = stripped[2:].strip()
            p = doc.add_paragraph(text, style='List Bullet')
            p.paragraph_format.line_spacing = 1.0
            # Ensure font is correct for list items
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
            continue
        
        # Normal text (with bold support)
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.first_line_indent = Inches(0.5)
        
        # Simple bold parsing: **text**
        parts = re.split(r'(\*\*.*?\*\*)', stripped)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
            else:
                run = p.add_run(part)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

def save_to_docx(content, output_path, title_text):
    """Saves the content to a DOCX file. Appends if exists. Handles locking."""
    base_name, ext = os.path.splitext(output_path)
    counter = 1
    current_output_path = output_path
    
    while True:
        try:
            if os.path.exists(current_output_path):
                try:
                    doc = Document(current_output_path)
                    log_event(f"Appending to existing file: {current_output_path}")
                except Exception:
                    # If opening fails, assume locked/corrupt and force next version
                    raise PermissionError("Cannot open existing file.")
            else:
                doc = Document()
                log_event(f"Creating new file: {current_output_path}")

            # Apply Styles (Times New Roman, Size 12, Single Spaced)
            style = doc.styles['Normal']
            style.font.name = 'Times New Roman'
            style.font.size = Pt(12)
            style.paragraph_format.line_spacing = 1.0
            
            # Count existing subheadings to determine the next number
            # We look for paragraphs that match our subheading pattern
            next_num = 1
            for para in doc.paragraphs:
                if re.match(r'^\d+\.\t.*Responses to Written Discovery', para.text):
                    next_num += 1

            # Determine "Responding Party" from title_text (filename)
            # Remove extension and potentially "discovery" or "responses" keywords
            party_name = title_text
            for skip in [".pdf", ".docx", "discovery", "responses", "response"]:
                party_name = party_name.replace(skip, "").replace(skip.upper(), "")
            party_name = party_name.strip(" _-")

            # Subheading: [Number]. [Party Name]’s Responses to Written Discovery
            # Indentation: Number at 1.0", text at 1.5"
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.0
            # Left indent 1.5", First line indent -0.5" (relative to 1.5") = 1.0" for the number
            p.paragraph_format.left_indent = Inches(1.5)
            p.paragraph_format.first_line_indent = Inches(-0.5)
            
            # Add tab stop at 1.5" for the text following the number
            tab_stops = p.paragraph_format.tab_stops
            tab_stops.add_tab_stop(Inches(1.5))

            # Number and Title
            run_num = p.add_run(f"{next_num}.\t")
            run_num.bold = True
            run_num.font.name = 'Times New Roman'
            run_num.font.size = Pt(12)

            run_title = p.add_run(f"{party_name}’s Responses to Written Discovery")
            run_title.bold = True
            run_title.underline = True
            run_title.font.name = 'Times New Roman'
            run_title.font.size = Pt(12)
            
            add_markdown_to_doc(doc, content)
                
            doc.save(current_output_path)
            log_event(f"Saved summary to: {current_output_path}")
            return True

        except (PermissionError, IOError) as e:
            log_event(f"File locked or inaccessible: {current_output_path}. Trying next version. Error: {e}", level="warning")
            counter += 1
            # Format: Original v.2.docx, Original v.3.docx
            current_output_path = f"{base_name} v.{counter}{ext}"
            
            if counter > 10:
                log_event("Failed to save after 10 attempts.", level="error")
                return False
        except Exception as e:
            log_event(f"Error saving to DOCX: {e}", level="error")
            return False

def main():
    if len(sys.argv) < 2:
        log_event("Error: No file path provided.", level="error")
        sys.exit(1)

    # Clean up empty arguments that can result from double-quoting issues ("")
    raw_args = [arg for arg in sys.argv[1:] if arg.strip()]
    
    if not raw_args:
        log_event("Error: No valid file paths provided.", level="error")
        sys.exit(1)

    # Heuristic to handle split paths: 
    # If multiple args are provided but the first doesn't exist, 
    # check if joining them all back together forms a valid path.
    if len(raw_args) > 1:
        combined_path = " ".join(raw_args).strip("\"'")
        if os.path.exists(combined_path) and not os.path.exists(raw_args[0]):
            file_paths = [combined_path]
        else:
            file_paths = raw_args
    else:
        file_paths = raw_args
    
    # --- Dispatcher Mode ---
    if len(file_paths) > 1:
        log_event(f"Detected multiple file paths: {len(file_paths)} items. Launching separate agents...")
        for path in file_paths:
            # We must quote paths with spaces when re-launching
            quoted_path = f'"{path}"' if ' ' in path and not path.startswith('"') else path
            log_event(f"Spawning agent for: {path}")
            try:
                # Spawn independent process
                if os.name == 'nt':
                    # Use CREATE_NO_WINDOW (0x08000000) for headless execution
                    subprocess.Popen([sys.executable, sys.argv[0], path], creationflags=0x08000000)
                else:
                    subprocess.Popen([sys.executable, sys.argv[0], path])
            except Exception as e:
                log_event(f"Failed to spawn agent for {path}: {e}", level="error")
        
        print(f"Launched {len(file_paths)} summarize discovery agents.")
        sys.exit(0)

    # --- Worker Mode ---
    # At this point, we assume we are processing a SINGLE path (which might be a directory or file)
    
    if not file_paths:
        log_event("No valid file paths found.", level="error")
        sys.exit(1)
        
    input_path = file_paths[0]
    
    # Remove surrounding quotes if present (common in CLI args)
    if (input_path.startswith('"') and input_path.endswith('"')) or (input_path.startswith("'") and input_path.endswith("'")):
        input_path = input_path[1:-1]
    
    # Handle absolute path conversion
    input_path = os.path.abspath(input_path)

    if not os.path.exists(input_path):
        log_event(f"Error: File not found: {input_path}", level="error")
        sys.exit(1)

    log_event(f"--- Starting Agent for input: {input_path} ---")

    # --- Directory Handling ---
    if os.path.isdir(input_path):
        log_event(f"Input is a directory. Scanning for files in: {input_path}")
        files_to_process = []
        for root, _, files in os.walk(input_path):
            for file in files:
                if file.lower().endswith(('.pdf', '.docx')):
                    # Exclude the output file itself if it happens to be in the scan path
                    if "Discovery_Response_Summaries" in file:
                        continue
                    
                    # Filter out documents with "RFP", "prod", or "production" in the filename
                    lower_file = file.lower()
                    if any(keyword in lower_file for keyword in ["rfp", "prod", "production"]):
                        log_event(f"Skipping filtered file: {file}")
                        continue

                    files_to_process.append(os.path.join(root, file))
        
        if not files_to_process:
            log_event("No suitable files (.pdf, .docx) found in directory.")
            sys.exit(0)
        
        log_event(f"Found {len(files_to_process)} files to process.")
        
        # Spin up a new CLI instance for each document
        # We run them sequentially to avoid file locking issues on the single output file
        for file_path in files_to_process:
            log_event(f"Spawning subprocess for: {file_path}")
            try:
                # Call this script again with the specific file path
                subprocess.run([sys.executable, sys.argv[0], file_path], check=True)
            except subprocess.CalledProcessError as e:
                log_event(f"Subprocess failed for {file_path}: {e}", level="error")
        
        log_event("--- Directory Processing Complete ---")
        sys.exit(0)
    
    # --- Single File Processing ---

    # 1. Extract Text
    text = extract_text(input_path)
    if not text:
        sys.exit(1)

    # 2. Load Prompt
    try:
        with open(PROMPT_FILE, "r", encoding="utf-8") as f:
            prompt_instruction = f.read()
    except Exception as e:
        log_event(f"Error reading prompt file {PROMPT_FILE}: {e}", level="error")
        sys.exit(1)

    # 3. Call LLM
    summary = call_gemini(prompt_instruction, text)
    if not summary:
        sys.exit(1)

    # 4. Save Output
    # Strategy: Locate "Current Clients" in the path.
    # The Case Root is typically 2 levels deeper: Current Clients/<Client>/<Case>
    # We want to save to <Case Root>/NOTES
    
    parts = input_path.split(os.sep)
    output_dir = None

    # Find "Current Clients" (case-insensitive)
    for i, part in enumerate(parts):
        if part.lower() == "current clients":
            # Check if we have enough depth for Client and Case
            if i + 2 < len(parts):
                # Construct path up to Case Root: Current Clients (i) -> Client (i+1) -> Case (i+2)
                case_root_parts = parts[:i+3]
                output_dir = os.sep.join(case_root_parts + ["NOTES", "AI OUTPUT"])
            break
            
    if not output_dir:
        # Fallback 1: If "NOTES" is already in the path, use it.
        for i in range(len(parts) - 1, -1, -1):
            if parts[i].upper() == "NOTES":
                output_dir = os.path.join(os.sep.join(parts[:i+1]), "AI OUTPUT")
                break
    
    if not output_dir:
        # Fallback 2: Sibling NOTES to the file's parent folder (Original fallback)
        input_dir = os.path.dirname(input_path)
        parent_dir = os.path.dirname(input_dir)
        output_dir = os.path.join(parent_dir, "NOTES", "AI OUTPUT")

    if not os.path.exists(output_dir):
        try:
            os.makedirs(output_dir)
            log_event(f"Created directory: {output_dir}")
        except Exception as e:
            log_event(f"Error creating directory {output_dir}: {e}", level="error")
            sys.exit(1)

    output_file = os.path.join(output_dir, "Discovery_Response_Summaries.docx")
    save_to_docx(summary, output_file, os.path.basename(input_path))
    
    log_event("--- Agent Finished Successfully ---")

if __name__ == '__main__':
    main()
