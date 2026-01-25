import argparse
import os
import sys
import re
import glob
import json
import subprocess
import logging
import platform
import shutil
import tempfile
import time
import concurrent.futures
import io
import google.generativeai as genai

# Word Document Generation
try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    # We will handle the missing import gracefully when index creation is called
    pass

# --- Dependency Setup ---
# Attempt to import PyMuPDF (fitz) for fast processing
try:
    import fitz
except ImportError:
    print("Error: 'pymupdf' is required for this optimized version.")
    print("Please run: pip install pymupdf")
    sys.exit(1)

# PyPDF2 is still used for the final split to minimize write logic changes
try:
    import pypdf # Prefer pypdf
    PyPDF2 = pypdf
except ImportError:
    try:
        import PyPDF2
    except ImportError:
        print("Error: pypdf or PyPDF2 is required.")
        sys.exit(1)

# OCR Setup
OCR_AVAILABLE = False
try:
    import pytesseract
    from PIL import Image
    
    # Tesseract Path Logic
    if os.name == 'nt':
        tesseract_paths = [
            r"C:\\Program Files\\Tesseract-OCR\\tesseract.exe",
            os.path.expanduser(r"~\\AppData\\Local\\Programs\\Tesseract-OCR\\tesseract.exe")
        ]
        for p in tesseract_paths:
            if os.path.exists(p):
                pytesseract.pytesseract.tesseract_cmd = p
                OCR_AVAILABLE = True
                break
    else:
        if shutil.which("tesseract"):
            OCR_AVAILABLE = True
            
except ImportError:
    pass

# --- Logging Setup ---
PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
LOG_FILE = os.path.join(PROJECT_ROOT, "separator_activity.log")

logging.basicConfig(
    level=logging.INFO,
    format='[%(asctime)s] %(levelname)s: %(message)s',
    handlers=[
        logging.FileHandler(LOG_FILE),
        logging.StreamHandler(sys.stdout)
    ]
)
logger = logging.getLogger("Separator")

# --- Constants ---
PRIMARY_MODEL = "gemini-3-flash-preview"
FALLBACK_MODEL = "gemini-2.5-flash"
BASE_PATH_WIN = r"Z:\\Shared\\Current Clients"

# --- Functions ---

def get_case_path(file_num: str):
    """Locates the physical directory for a file number (####.###)."""
    if not re.match(r'^\d{4}\.\d{3}$', file_num):
        return None

    carrier_num, case_num = file_num.split('.')
    base_path = BASE_PATH_WIN

    if not os.path.exists(base_path):
        return None

    # Find Carrier Folder
    carrier_folders = glob.glob(os.path.join(base_path, f"{carrier_num} - *"))
    if not carrier_folders:
        return None
    carrier_path = carrier_folders[0]

    # Find Case Folder
    case_folders = glob.glob(os.path.join(carrier_path, f"{case_num} - *"))
    if not case_folders:
        return None
    
    return os.path.abspath(case_folders[0])

def get_page_text_fast(pdf_path, page_num):
    """
    Extracts text from a single page using PyMuPDF (fitz).
    If text is sparse, it renders the page image and uses OCR (Tesseract).
    
    page_num is 1-based.
    """
    try:
        # Open the PDF for this specific task (thread/process safe if opened here)
        doc = fitz.open(pdf_path)
        # fitz uses 0-based indexing
        page = doc.load_page(page_num - 1)
        
        # 1. Try direct text extraction (VERY FAST)
        text = page.get_text("text")
        
        # 2. If sparse, fall back to OCR
        if len(text.strip()) < 50 and OCR_AVAILABLE:
            # Render page to image (pixmap)
            # matrix=fitz.Matrix(2, 2) roughly doubles resolution (~144 dpi -> ~300 dpi) for better OCR
            # But for speed on headers, 1.5 or even 1.0 (72 dpi) might suffice. 
            # 2.0 is safer for accuracy.
            mat = fitz.Matrix(2, 2) 
            
            # Crop to top 25% of the page to save time
            rect = page.rect
            clip = fitz.Rect(0, 0, rect.width, rect.height * 0.25)
            
            pix = page.get_pixmap(matrix=mat, clip=clip)
            
            # Convert to PIL Image
            img_data = pix.tobytes("png")
            img = Image.open(io.BytesIO(img_data))
            
            # Run OCR
            ocr_text = pytesseract.image_to_string(img)
            
            if len(ocr_text.strip()) > len(text.strip()):
                text = ocr_text

        doc.close() 
        
        # Clean up text
        clean_text = " ".join(text[:1000].split())[:500]
        return f"Page {page_num}: {clean_text}" 
        
    except Exception as e:
        logger.error(f"Error processing page {page_num}: {e}")
        return f"Page {page_num}: [Error processing page]"

def extract_headers(pdf_path):
    """
    Extracts headers from all pages using parallel processing.
    """
    headers_map = {}
    
    try:
        # Open once to get page count (fast)
        doc = fitz.open(pdf_path)
        num_pages = len(doc)
        doc.close()
    except Exception as e:
        logger.error(f"Failed to read PDF structure: {e}")
        sys.exit(1)
        
    logger.info(f"Extracting headers from {num_pages} pages in {pdf_path} using parallel processing...")
    
    # Use ThreadPoolExecutor. 
    # ProcessPoolExecutor is better for CPU, but Tesseract invokes a subprocess anyway,
    # and fitz releases GIL for many ops. Threads are lighter and avoid pickling issues.
    max_workers = os.cpu_count() or 4
    if max_workers > 16: max_workers = 16 # Cap it reasonably
    
    with concurrent.futures.ThreadPoolExecutor(max_workers=max_workers) as executor:
        future_to_page = {
            executor.submit(get_page_text_fast, pdf_path, i + 1): i + 1 
            for i in range(num_pages)
        }
        
        processed_count = 0
        for future in concurrent.futures.as_completed(future_to_page):
            page_num = future_to_page[future]
            try:
                result = future.result()
                headers_map[page_num] = result
            except Exception as exc:
                logger.error(f"Page {page_num} generated an exception: {exc}")
                headers_map[page_num] = f"Page {page_num}: [Exception]"
            
            processed_count += 1
            if processed_count % 50 == 0 or processed_count == num_pages:
                logger.info(f"Processed {processed_count}/{num_pages} pages")
                
    # Sort headers by page number
    sorted_headers = [headers_map[i] for i in sorted(headers_map.keys())]
    return sorted_headers

def call_gemini(prompt, model):
    """Calls Gemini API directly."""
    api_key = os.environ.get("GEMINI_API_KEY")
    if not api_key:
        raise Exception("GEMINI_API_KEY environment variable not set.")
        
    genai.configure(api_key=api_key)
    
    try:
        # Run API call
        m = genai.GenerativeModel(model)
        response = m.generate_content(prompt)
        return response.text.strip()
    except Exception as e:
        raise Exception(f"Gemini API error: {e}")

def analyze_headers(headers):
    """Sends headers to AI to find boundaries."""
    prompt = (
        "I have a single PDF containing multiple distinct legal or administrative documents. "
        "Below is the text from the top of each page. "
        "Identify the start and end page of each distinct document. "
        "For each document, also extract the document's date (e.g., '2023-05-15') if it is visible on the first page of that document. "
        "Return the result ONLY as a list of lines, one per document, using the pipe character '|' as a separator.\n"
        "Format: ID|Title|Date|StartPage|EndPage\n"
        "Provide a detailed and descriptive title for each document (e.g., 'Plaintiff's Motion to Compel Discovery' instead of just 'Motion', or 'Exhibit A - Medical Report' instead of 'Exhibit').\n"
        "Example:\n"
        "1|Plaintiff's Complaint for Damages|2023-05-15|1|5\n"
        "2|Proof of Service of Summons||6|7\n"
        "3|Exhibit A - Medical Report from Dr. Smith|2022-12-01|8|10\n"
        "\n\nPDF HEADERS:\n" + "\n".join(headers)
    )
    
    response = ""
    try:
        logger.info(f"Analyzing structure with {PRIMARY_MODEL}...")
        response = call_gemini(prompt, PRIMARY_MODEL)
    except Exception as e:
        logger.warning(f"Primary model failed: {e}. Trying fallback {FALLBACK_MODEL}...")
        try:
            response = call_gemini(prompt, FALLBACK_MODEL)
        except Exception as e2:
            logger.error(f"Fallback model failed: {e2}")
            sys.exit(1)
            
    # Parse Response
    docs = []
    lines = response.split('\n')
    for line in lines:
        line = line.strip()
        if not line or '|' not in line: continue
        if line.startswith("```"): continue # Skip markdown
        
        parts = line.split('|')
        if len(parts) >= 5:
            try:
                docs.append({
                    "id": parts[0].strip(),
                    "title": parts[1].strip(),
                    "date": parts[2].strip(),
                    "start": int(parts[3].strip()),
                    "end": int(parts[4].strip())
                })
            except ValueError:
                pass
    
    if not docs:
        logger.error("No documents identified by AI.")
        logger.debug(f"Raw AI response: {response}")
        sys.exit(1)
        
    return docs

def sanitize_filename(name):
    return "".join(c for c in name if c.isalnum() or c in (' ', '-', '_')).strip()

def split_pdf_pages(pdf_path, output_folder, selection):
    """Splits the PDF based on selection."""
    try:
        reader = PyPDF2.PdfReader(pdf_path)
    except Exception as e:
        print(f"Error reading source PDF: {e}")
        return

    base_name = os.path.splitext(os.path.basename(pdf_path))[0]
    
    print(f"\nProcessing {len(selection)} items...")
    
    for doc in selection:
        try:
            writer = PyPDF2.PdfWriter()
            start = doc['start'] - 1 # 1-based to 0-based
            end = doc['end']     # exclusive in python slice, inclusive in prompt? Prompt said "1-5", usually means 1,2,3,4,5. 
                                 # Range(start, end) where end is 5 would be 0,1,2,3,4. 
                                 # If page 5 is included, we need index 4. Range(0, 5) covers 0,1,2,3,4.
                                 # So if doc says end=5, we want loop range(0, 5). 
            
            # Verify bounds
            if start < 0: start = 0
            if end > len(reader.pages): end = len(reader.pages)
            
            for i in range(start, end):
                writer.add_page(reader.pages[i])
            
            safe_title = sanitize_filename(doc['title'])
            if len(safe_title) > 100:
                safe_title = safe_title[:100].strip()
            # Format: MyPacket- 01 - Complaint.pdf
            # User example: "MyPacket- 01 - Complaint.pdf" (original + id + title)
            # Assuming ID is numeric.
            try:
                 id_num = int(doc['id'])
                 id_str = f"{id_num:02d}"
            except:
                 id_str = str(doc['id'])
                 
            out_filename = f"{base_name} - {id_str} - {safe_title}.pdf"
            out_path = os.path.join(output_folder, out_filename)
            
            with open(out_path, "wb") as f:
                writer.write(f)
                
            print(f"  [OK] Saved: {out_filename}")
            logger.info(f"Saved {out_filename}")
            
        except Exception as e:
            print(f"  [ERROR] Failed to save {doc['title']}: {e}")
            logger.error(f"Failed to save {doc['title']}: {e}")

# --- Modes ---

def run_analysis(pdf_path):
    if not os.path.exists(pdf_path):
        logger.error(f"File not found: {pdf_path}")
        return

    logger.info(f"Starting analysis for: {pdf_path}")
    
    # 1. Extract
    headers = extract_headers(pdf_path)
    
    # 2. Analyze
    docs = analyze_headers(headers)
    logger.info(f"identified {len(docs)} documents.")
    
    # 3. Create Index Document
    create_index_word(pdf_path, docs)
    
    # 4. Save Temp Map
    fd, temp_path = tempfile.mkstemp(suffix=".json", text=True)
    with os.fdopen(fd, 'w') as f:
        json.dump(docs, f)
    
    logger.info(f"Map saved to {temp_path}. Launching interactive window...")
    
    # 4. Launch Interactive Window
    # We recursively call this script with different arguments in a new console
    current_script = os.path.abspath(__file__)
    
    cmd = [sys.executable, current_script, "--interactive", temp_path, "--original-pdf", pdf_path]
    
    if os.name == 'nt':
        subprocess.Popen(cmd, creationflags=subprocess.CREATE_NEW_CONSOLE)
    else:
        # Linux fallback (not primary target per prompt, but good practice)
        subprocess.Popen(cmd)

def get_notes_dir(pdf_path):
    """
    Locates the NOTES/AI OUTPUT folder within the case root.
    Uses logic similar to complaint.py by extracting file number from path.
    """
    norm_path = os.path.abspath(pdf_path)
    
    # 1. Try to extract ####.### pattern from path (Standardized logic)
    match = re.search(r'(\d{4})\.(\d{3})', norm_path)
    if match:
        file_num = f"{match.group(1)}.{match.group(2)}"
        case_root = get_case_path(file_num)
        if case_root:
            notes_dir = os.path.join(case_root, "NOTES", "AI OUTPUT")
            return notes_dir

    # 2. Fallback to 'Current Clients' path segment logic
    parts = norm_path.split(os.sep)
    try:
        idx = -1
        for i, part in enumerate(parts):
            if part.lower() == "current clients":
                idx = i
                break
        
        if idx != -1 and len(parts) > idx + 2:
            # Case root is usually: ...\Current Clients\#### - Carrier\### - Case
            case_root = os.sep.join(parts[:idx+3])
            notes_dir = os.path.join(case_root, "NOTES", "AI OUTPUT")
            return notes_dir
    except Exception as e:
        logger.warning(f"Failed to parse case root from path: {e}")

    # 3. Fallback: climb up the tree looking for a NOTES folder
    current_dir = os.path.dirname(norm_path)
    for _ in range(5):
        potential_notes = os.path.join(current_dir, "NOTES")
        if os.path.exists(potential_notes) and os.path.isdir(potential_notes):
            return os.path.join(potential_notes, "AI OUTPUT")
        parent = os.path.dirname(current_dir)
        if parent == current_dir: break
        current_dir = parent
        
    # Final Fallback
    return os.path.join(os.path.dirname(norm_path), "NOTES", "AI OUTPUT")

def create_index_word(pdf_path, docs):
    """Creates a Word document index of identified sub-documents."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        logger.error("python-docx not installed. Skipping index creation.")
        return

    try:
        notes_ai_dir = get_notes_dir(pdf_path)
        indexes_dir = os.path.join(notes_ai_dir, "INDEXES")
        
        if not os.path.exists(indexes_dir):
            os.makedirs(indexes_dir, exist_ok=True)
            logger.info(f"Created INDEXES folder: {indexes_dir}")
        
        base_name = os.path.splitext(os.path.basename(pdf_path))[0]
        docx_filename = f"Index_{base_name}.docx"
        docx_path = os.path.join(indexes_dir, docx_filename)
        
        doc = Document()
        
        # Set default style to Times New Roman, 12, Black
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        font.color.rgb = RGBColor(0, 0, 0)
        
        # Header paragraph
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"INDEX OF DOCUMENTS - {base_name}")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)

        # Create Table
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        table.autofit = False # Disable autofit to allow manual widths

        # Define Column Widths (Total ~6.5 inches for standard margins)
        # Id: ~0.4", Date: ~1.2", Pages: ~1.2", Title: Remaining (~3.7")
        col_widths = [Inches(0.4), Inches(3.7), Inches(1.2), Inches(1.2)]
        
        # Setup Headers
        headers = ['Id', 'Document Title', 'Document Date', 'Page Ranges']
        hdr_cells = table.rows[0].cells
        for i, h_text in enumerate(headers):
            cell = hdr_cells[i]
            cell.width = col_widths[i]
            cell_p = cell.paragraphs[0]
            cell_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cell_p.add_run(h_text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
            
        # Add Data Rows
        for d in docs:
            row_cells = table.add_row().cells
            
            # Formatting Page Ranges
            start_pg = d.get('start', '')
            end_pg = d.get('end', '')
            page_range = f"{start_pg} - {end_pg}" if start_pg != end_pg else str(start_pg)

            row_data = [
                str(d.get('id', '')),
                str(d.get('title', '')),
                str(d.get('date', '')),
                page_range
            ]
            for i, val in enumerate(row_data):
                cell = row_cells[i]
                cell.width = col_widths[i]
                cell_p = cell.paragraphs[0]
                run = cell_p.add_run(val)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(0, 0, 0)
        
        doc.save(docx_path)
        logger.info(f"Index saved to: {docx_path}")
        print(f"\n[SUCCESS] Index created: {docx_filename} in INDEXES folder.")

    except Exception as e:
        logger.error(f"Error creating index: {e}")
        print(f"\n[ERROR] Failed to create index document: {e}")

def run_interactive(json_path, pdf_path):
    # Load Docs
    docs = []
    try:
        with open(json_path, 'r') as f:
            docs = json.load(f)
    except Exception as e:
        print(f"Error loading document map: {e}")
        input("Press Enter to exit...")
        return

    # Display Loop
    while True:
        os.system('cls' if os.name == 'nt' else 'clear')
        print(f"--- Document Separator: {os.path.basename(pdf_path)} ---\n")
        print(f"{ 'ID':<4} | { 'Date':<12} | { 'Pages':<8} | {'Title'}")
        print("-" * 60)
        for doc in docs:
            print(f"{doc['id']:<4} | {doc['date']:<12} | {doc['start']}-{doc['end']:<6} | {doc['title']}")
        print("-" * 60)
        
        print("\nEnter IDs to keep (e.g., '1, 3', '1-5', 'all') or 'q' to quit.")
        choice = input("> ").strip().lower()
        
        if choice == 'q':
            break
            
        selection = []
        if choice == 'all':
            selection = docs
        else:
            try:
                # Parse range/list
                ids_to_keep = set()
                parts = choice.split(',')
                for part in parts:
                    part = part.strip()
                    if '-' in part:
                        s, e = map(int, part.split('-'))
                        ids_to_keep.update(range(s, e + 1))
                    else:
                        ids_to_keep.add(str(part)) # Try string first
                        ids_to_keep.add(str(int(part))) # Also numeric string
                
                selection = [d for d in docs if str(d['id']) in ids_to_keep]
            except Exception as e:
                print(f"Invalid input: {e}")
                time.sleep(1)
                continue
        
        if not selection:
            print("No documents selected.")
            time.sleep(1)
            continue
            
        # Process Selection
        base_dir = os.path.dirname(os.path.abspath(pdf_path))
        source_name = os.path.splitext(os.path.basename(pdf_path))[0]
        output_folder = os.path.join(base_dir, f"PULLED-{source_name}")
        
        if not os.path.exists(output_folder):
            os.makedirs(output_folder)
            
        split_pdf_pages(pdf_path, output_folder, selection)
        
        print("\nDone! Files saved to:")
        print(output_folder)
        print("\nPress Enter to close this window.")
        input()
        
        # Cleanup temp file
        try:
            os.remove(json_path)
        except:
            pass
        break

def main():
    if len(sys.argv) < 2:
        # Check if we're in a special mode
        if "--interactive" not in sys.argv:
            print("Usage: separate.py <pdf_path>")
            sys.exit(1)

    # Handle path with spaces if passed without quotes
    if "--interactive" not in sys.argv:
        # Join all arguments except the script name to handle spaces
        # But handle cases where shell might have stripped outer quotes
        raw_path = " ".join(sys.argv[1:])
        
        # Strip potential outer quotes added by drag-and-drop or shell
        pdf_path = raw_path.strip().strip('"').strip("'")
        
        # If the path doesn't exist, try to see if the first arg was actually the path
        # (Fallback for specific shell behaviors)
        if not os.path.exists(pdf_path) and os.path.exists(sys.argv[1]):
            pdf_path = sys.argv[1]
        
        run_analysis(pdf_path)
        return

    parser = argparse.ArgumentParser()
    parser.add_argument("pdf_path", help="PDF File Path", nargs='?')
    parser.add_argument("--interactive", help="Path to temp json for interactive mode")
    parser.add_argument("--original-pdf", help="Original PDF for interactive mode")
    args = parser.parse_args()

    if args.interactive and args.original_pdf:
        run_interactive(args.interactive, args.original_pdf)
    elif args.pdf_path:
        run_analysis(args.pdf_path)
    else:
        print("Usage: separate.py <pdf_path>")

if __name__ == "__main__":
    main()