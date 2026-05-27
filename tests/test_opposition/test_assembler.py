import pytest
from docx import Document

from icharlotte_core.opposition.assembler import assemble_opposition_preview
from icharlotte_core.opposition.models import DraftDocument
from icharlotte_core.word_validator import validate_opposition_docx


def test_assemble_body_only_preview_and_validate(tmp_path):
    output_path = tmp_path / "opposition_preview.docx"
    draft = DraftDocument(
        title="Opposition to Motion for Summary Judgment",
        body_text=(
            "I. INTRODUCTION\n"
            "The motion should be denied.\n\n"
            "II. ARGUMENT\n"
            "Triable issues exist."
        ),
    )

    result = assemble_opposition_preview(
        draft=draft,
        output_path=str(output_path),
        caption_path="",
    )

    assert result == str(output_path)
    validation = validate_opposition_docx(str(output_path))
    assert validation.has_errors is False
    doc = Document(str(output_path))
    assert "Triable issues exist" in "\n".join(p.text for p in doc.paragraphs)


def test_assemble_reuses_caption_docx_when_available(tmp_path):
    caption_path = tmp_path / "Caption Page.docx"
    caption = Document()
    caption.add_paragraph("CAPTION PAGE")
    caption.save(str(caption_path))
    output_path = tmp_path / "preview.docx"

    assemble_opposition_preview(
        draft=DraftDocument(title="Opposition", body_text="Argument body"),
        output_path=str(output_path),
        caption_path=str(caption_path),
    )

    doc = Document(str(output_path))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Opposition" in text
    assert "Argument body" in text
    assert "CAPTION PAGE" not in text


def test_assemble_replaces_caption_markers_in_headers_and_footers(tmp_path):
    caption_path = tmp_path / "Caption Page.docx"
    caption = Document()
    caption.add_paragraph("CAPTION PAGE")
    caption.sections[0].header.paragraphs[0].text = "CAPTION PAGE"
    caption.sections[0].footer.paragraphs[0].text = "CAPTION PAGE"
    caption.save(str(caption_path))
    output_path = tmp_path / "preview.docx"

    assemble_opposition_preview(
        draft=DraftDocument(title="Opposition", body_text="Argument body"),
        output_path=str(output_path),
        caption_path=str(caption_path),
    )

    validation = validate_opposition_docx(str(output_path))
    assert validation.has_errors is False
    doc = Document(str(output_path))
    all_text = "\n".join(
        [
            *(p.text for p in doc.paragraphs),
            *(p.text for p in doc.sections[0].header.paragraphs),
            *(p.text for p in doc.sections[0].footer.paragraphs),
        ]
    )
    assert "CAPTION PAGE" not in all_text
    assert all_text.count("Opposition") >= 3


def test_validate_opposition_docx_catches_header_caption_marker(tmp_path):
    output_path = tmp_path / "bad_opposition.docx"
    doc = Document()
    doc.add_paragraph("Opposition body")
    doc.sections[0].header.paragraphs[0].text = "CAPTION PAGE"
    doc.save(str(output_path))

    validation = validate_opposition_docx(str(output_path))

    assert validation.has_errors is True
    assert any(finding.rule == "caption_marker" for finding in validation.findings)


def test_assemble_replaces_split_run_caption_marker(tmp_path):
    caption_path = tmp_path / "Split Caption Page.docx"
    caption = Document()
    paragraph = caption.add_paragraph()
    paragraph.add_run("CAPTION ")
    paragraph.add_run("PAGE")
    caption.save(str(caption_path))
    output_path = tmp_path / "preview.docx"

    assemble_opposition_preview(
        draft=DraftDocument(title="Opposition", body_text="Argument body"),
        output_path=str(output_path),
        caption_path=str(caption_path),
    )

    doc = Document(str(output_path))
    assert "CAPTION PAGE" not in "\n".join(p.text for p in doc.paragraphs)
    assert "Opposition" in "\n".join(p.text for p in doc.paragraphs)


def test_validate_opposition_docx_catches_split_run_caption_marker(tmp_path):
    output_path = tmp_path / "bad_split_caption.docx"
    doc = Document()
    paragraph = doc.add_paragraph()
    paragraph.add_run("CAPTION ")
    paragraph.add_run("PAGE")
    doc.add_paragraph("Opposition body")
    doc.save(str(output_path))

    validation = validate_opposition_docx(str(output_path))

    assert validation.has_errors is True
    assert any(finding.rule == "caption_marker" for finding in validation.findings)


def test_assemble_rejects_caption_output_same_path(tmp_path):
    caption_path = tmp_path / "Caption Page.docx"
    caption = Document()
    caption.add_paragraph("CAPTION PAGE")
    caption.save(str(caption_path))

    with pytest.raises(ValueError):
        assemble_opposition_preview(
            draft=DraftDocument(title="Opposition", body_text="Argument body"),
            output_path=str(caption_path),
            caption_path=str(caption_path),
        )

    doc = Document(str(caption_path))
    assert "CAPTION PAGE" in "\n".join(p.text for p in doc.paragraphs)
