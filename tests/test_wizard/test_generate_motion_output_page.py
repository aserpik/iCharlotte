"""Citation-review surface on the Generate Motion output page (parity with
the oppose output page)."""
from __future__ import annotations

import os

import pytest

pytest.importorskip("PySide6")
pytest.importorskip("pytestqt")

from PySide6.QtCore import QUrl  # noqa: E402

from icharlotte_core.opposition.models import (  # noqa: E402
    CitationVerification,
    DraftDocument,
)
from icharlotte_core.ui.wizard.pages.generate_motion_page import (  # noqa: E402
    GenerateMotionOutputPage,
)


def _draft(citations):
    return DraftDocument(
        title="Motion to Compel",
        body_text="See *Smith v. Jones* (2010) 50 Cal.4th 100 for support.",
        citations=citations,
    )


def test_supported_citation_renders_green(qtbot):
    page = GenerateMotionOutputPage()
    qtbot.addWidget(page)
    page.show_result(_draft([
        CitationVerification(
            citation_text="Smith v. Jones (2010) 50 Cal.4th 100",
            verdict="SUPPORTED", kind="case",
        )
    ]))
    html = page.editor.toHtml().lower()
    assert "#1e8e3e" in html  # green underline color


def test_not_supported_citation_renders_red(qtbot):
    page = GenerateMotionOutputPage()
    qtbot.addWidget(page)
    page.show_result(_draft([
        CitationVerification(
            citation_text="Smith v. Jones (2010) 50 Cal.4th 100",
            verdict="NOT_SUPPORTED", kind="case",
        )
    ]))
    assert "#c5221f" in page.editor.toHtml().lower()  # red


def test_summary_banner_counts_per_verdict(qtbot):
    page = GenerateMotionOutputPage()
    qtbot.addWidget(page)
    page.show_result(DraftDocument(
        title="M", body_text="Body.",
        citations=[
            CitationVerification(citation_text="a", verdict="SUPPORTED"),
            CitationVerification(citation_text="b", verdict="SUPPORTED"),
            CitationVerification(citation_text="c", verdict="PARTIAL"),
            CitationVerification(citation_text="d", verdict="NOT_SUPPORTED"),
        ],
    ))
    banner = page.summary_banner.text().lower()
    assert "supported" in banner
    assert "2" in banner


def test_clicking_anchor_updates_detail_panel(qtbot):
    page = GenerateMotionOutputPage()
    qtbot.addWidget(page)
    page.show_result(DraftDocument(
        title="M",
        body_text="First *A v. B* (2010) 1 Cal.5th 1; then *C v. D* (2011) 2 Cal.5th 2.",
        citations=[
            CitationVerification(citation_text="A v. B (2010) 1 Cal.5th 1",
                                 case_name="A v. B", verdict="SUPPORTED", kind="case"),
            CitationVerification(citation_text="C v. D (2011) 2 Cal.5th 2",
                                 case_name="C v. D", verdict="NOT_SUPPORTED", kind="case"),
        ],
    ))
    page._on_anchor_clicked(QUrl("citation:1"))
    assert "C v. D" in page.detail_panel.header_label.text()
    assert "NOT SUPPORTED" in page.detail_panel.header_label.text()


def test_empty_citations_shows_motion_specific_message(qtbot):
    page = GenerateMotionOutputPage()
    qtbot.addWidget(page)
    page.show_result(DraftDocument(title="M", body_text="Body.", citations=[]))
    assert "motion" in page.detail_panel.body_html.lower()


def test_save_warns_on_red_verdicts(qtbot, monkeypatch, tmp_path):
    from PySide6.QtWidgets import QFileDialog, QMessageBox

    page = GenerateMotionOutputPage()
    qtbot.addWidget(page)
    preview = tmp_path / "preview.docx"
    preview.write_bytes(b"dummy")
    page.show_result(DraftDocument(
        title="M", body_text="b", preview_path=str(preview),
        citations=[CitationVerification(citation_text="x", verdict="NOT_SUPPORTED")],
    ))

    warned = {"yes": False}

    def fake_question(parent, title, text, *args, **kwargs):
        warned["yes"] = True
        return QMessageBox.StandardButton.Cancel  # user cancels

    monkeypatch.setattr(QMessageBox, "question", fake_question)
    monkeypatch.setattr(QFileDialog, "getSaveFileName", lambda *a, **k: ("", ""))

    page.save_as()
    assert warned["yes"]


def test_open_in_word_button_present(qtbot):
    page = GenerateMotionOutputPage()
    qtbot.addWidget(page)
    assert hasattr(page, "open_btn")
    assert page.open_btn.text() == "Open in Word"


def test_open_in_word_enabled_only_with_existing_preview(qtbot, tmp_path):
    page = GenerateMotionOutputPage()
    qtbot.addWidget(page)

    # No preview path → disabled.
    page.show_result(DraftDocument(title="M", body_text="Body.", citations=[]))
    assert not page.open_btn.isEnabled()

    # preview_path that does not exist on disk → still disabled.
    page.show_result(DraftDocument(
        title="M", body_text="Body.", citations=[],
        preview_path=str(tmp_path / "nope.docx"),
    ))
    assert not page.open_btn.isEnabled()

    # Real on-disk .docx preview → enabled.
    preview = tmp_path / "preview.docx"
    preview.write_bytes(b"dummy")
    page.show_result(DraftDocument(
        title="M", body_text="Body.", citations=[],
        preview_path=str(preview),
    ))
    assert page.open_btn.isEnabled()

    # load_output on a real file also enables it.
    page2 = GenerateMotionOutputPage()
    qtbot.addWidget(page2)
    page2.load_output(str(preview))
    assert page2.open_btn.isEnabled()


def test_reopen_restores_verified_citations_from_sidecar(qtbot, tmp_path):
    # Generating with verified citations must persist a sidecar beside the
    # preview .docx so that reopening (load_output) restores the SAME verdicts
    # instead of showing "no citations detected".
    from docx import Document

    preview = tmp_path / "Motion Preview.docx"
    doc = Document()
    doc.add_paragraph("See Smith v. Jones (2010) 50 Cal.4th 100 for support.")
    doc.save(str(preview))

    page = GenerateMotionOutputPage()
    qtbot.addWidget(page)
    page.show_result(DraftDocument(
        title="Motion to Compel",
        body_text="See *Smith v. Jones* (2010) 50 Cal.4th 100 for support.",
        preview_path=str(preview),
        citations=[CitationVerification(
            citation_text="Smith v. Jones (2010) 50 Cal.4th 100",
            case_name="Smith v. Jones", verdict="SUPPORTED", kind="case",
            evidence="The court held the duty applies.",
        )],
    ))
    # show_result persisted the verdicts beside the preview.
    assert os.path.isfile(str(preview) + ".citations.json")

    # Reopen in a fresh page → verdicts restored, panel shows the citation.
    page2 = GenerateMotionOutputPage()
    qtbot.addWidget(page2)
    page2.load_output(str(preview))
    assert len(page2.draft.citations) == 1
    assert page2.draft.citations[0].verdict == "SUPPORTED"
    assert page2.draft.citations[0].citation_text == "Smith v. Jones (2010) 50 Cal.4th 100"
    assert "SUPPORTED" in page2.detail_panel.header_label.text()


def test_diagnostics_panel_renders_and_reopens_from_sidecar(qtbot, tmp_path):
    from docx import Document

    preview = tmp_path / "Motion Preview.docx"
    doc = Document()
    doc.add_paragraph("Generated motion body.")
    doc.save(str(preview))

    diagnostics = {
        "task": "generate_motion",
        "research": {
            "target_count": 9,
            "retrieved_authorities": 27,
            "source": "local_corpus",
        },
        "citations": {
            "found": 12,
            "verdicts": {"SUPPORTED": 10, "NOT_SUPPORTED": 2},
            "replacement_candidates": 2,
        },
        "phase_seconds": {"research": 42.25, "total": 50.0},
    }

    page = GenerateMotionOutputPage()
    qtbot.addWidget(page)
    page.show_result(DraftDocument(
        title="Motion",
        body_text="Generated motion body.",
        preview_path=str(preview),
        diagnostics=diagnostics,
    ))

    assert not page.diagnostics_browser.isHidden()
    text = page.diagnostics_browser.toPlainText().lower()
    assert "research targets: 9" in text
    assert "replacement candidates: 2" in text
    assert os.path.isfile(str(preview) + ".diagnostics.json")

    page2 = GenerateMotionOutputPage()
    qtbot.addWidget(page2)
    page2.load_output(str(preview))
    assert page2.draft.diagnostics["research"]["target_count"] == 9
    assert "retrieved authorities: 27" in page2.diagnostics_browser.toPlainText().lower()


def test_accepting_replacement_candidate_updates_draft_body(qtbot):
    page = GenerateMotionOutputPage()
    qtbot.addWidget(page)
    old = CitationVerification(
        citation_text="Smith v. Jones (2010) 50 Cal.4th 100",
        normalized_citation="Smith v. Jones (2010) 50 Cal.4th 100",
        verdict="NOT_SUPPORTED",
        kind="case",
        replacement_candidates=[
            CitationVerification(
                citation_text="Brown v. Davis (2015) 60 Cal.App.4th 200",
                normalized_citation="Brown v. Davis (2015) 60 Cal.App.4th 200",
                verdict="SUPPORTED",
                kind="case",
                note="Direct replacement.",
            ).to_dict()
        ],
    )
    page.show_result(DraftDocument(
        title="M",
        body_text="See Smith v. Jones (2010) 50 Cal.4th 100.",
        citations=[old],
    ))

    page.show_citation(0)
    assert "Replacement candidates" in page.detail_panel.body_html
    page._on_replacement_requested(old, 0)

    assert "Brown v. Davis (2015) 60 Cal.App.4th 200" in page.draft.body_text
    assert page.draft.citations[0].verdict == "SUPPORTED"
    assert page.draft.citations[0].citation_text == "Brown v. Davis (2015) 60 Cal.App.4th 200"
