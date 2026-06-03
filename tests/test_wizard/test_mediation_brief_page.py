"""Tests for the Mediation Brief Wizard task."""
import os

import pytest


# ---- Registry / routing (pure logic, no Qt) ----

def test_mediation_brief_registered():
    from icharlotte_core.ui.wizard.registry import TASK_REGISTRY

    spec = TASK_REGISTRY["mediation_brief"]
    assert spec.title == "Mediation Brief"
    assert spec.category == "Motions & Drafting"
    assert spec.script_name == ""
    assert "mediation" in spec.keywords


def test_mediation_brief_has_valid_category():
    from icharlotte_core.ui.wizard.registry import CATEGORY_ORDER, TASK_REGISTRY

    assert TASK_REGISTRY["mediation_brief"].category in CATEGORY_ORDER


def test_mediation_brief_is_in_process():
    from icharlotte_core.ui.wizard.task_routing import (
        get_in_process_task_builder_name,
        is_in_process_task,
        requires_initial_file_picker,
    )

    assert get_in_process_task_builder_name("mediation_brief") == "build_mediation_brief_tab"
    assert is_in_process_task("mediation_brief") is True
    # In-process tasks own their source selection — no pre-Settings picker.
    assert requires_initial_file_picker("mediation_brief") is False


def test_build_mediation_brief_tab_attribute_exists():
    from icharlotte_core.ui.wizard import in_process_task_tab

    assert hasattr(in_process_task_tab, "build_mediation_brief_tab")


# ---- Document reader ----

pytestqt = pytest.importorskip("pytestqt")  # ensures PySide6/Qt present for module import


def test_read_documents_txt_and_docx(tmp_path):
    from docx import Document

    from icharlotte_core.ui.wizard.pages.mediation_brief_page import _read_documents

    txt = tmp_path / "note.txt"
    txt.write_text("Plain text body", encoding="utf-8")

    docx_path = tmp_path / "summary.docx"
    doc = Document()
    doc.add_paragraph("Intro paragraph text")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "CellA"
    table.rows[0].cells[1].text = "CellB"
    doc.save(str(docx_path))

    content, warnings = _read_documents([str(txt), str(docx_path)])

    assert "Plain text body" in content
    assert "Intro paragraph text" in content
    # Table-aware extraction must surface cell text (no silent data loss).
    assert "CellA" in content and "CellB" in content
    assert warnings == []


def test_read_documents_reports_missing_file(tmp_path):
    from icharlotte_core.ui.wizard.pages.mediation_brief_page import _read_documents

    content, warnings = _read_documents([str(tmp_path / "nope.pdf")])

    assert content == ""
    assert any("not found" in w.lower() for w in warnings)


# ---- Save helper ----

def test_save_brief_writes_to_target(tmp_path):
    from icharlotte_core.ui.wizard.pages.mediation_brief_page import _save_brief

    class FakeGen:
        def assemble_document(self, caption, out):
            with open(out, "w", encoding="utf-8") as fh:
                fh.write("brief")

    dest = str(tmp_path / "MEDIATION")
    path = _save_brief(FakeGen(), "caption.docx", dest, "Brief.docx")

    assert os.path.basename(path) == "Brief.docx"
    assert os.path.isfile(path)


def test_save_brief_falls_back_when_destination_locked(tmp_path):
    from icharlotte_core.ui.wizard.pages.mediation_brief_page import _save_brief

    class LockGen:
        def assemble_document(self, caption, out):
            if os.path.basename(out) == "Brief.docx":
                raise PermissionError("file is open in Word")
            with open(out, "w", encoding="utf-8") as fh:
                fh.write("brief")

    dest = str(tmp_path / "MEDIATION")
    path = _save_brief(LockGen(), "caption.docx", dest, "Brief.docx")

    assert os.path.basename(path) == "Brief (2).docx"
    assert os.path.isfile(path)


# ---- Worker ----

class _FakeGen:
    """Stand-in for MediationBriefGenerator that records the pipeline calls."""

    def __init__(self):
        self.sections = {}
        self.is_active = False
        self.caption_template_path = None
        self.document_content = ""
        self.calls = []

    def get_style_excerpts(self):
        self.calls.append("style")
        return {}

    def run_planning_pass(self):
        self.calls.append("planning")
        return "plan"

    def generate_section(self, name):
        self.calls.append(f"section:{name}")
        return f"{name} body"

    def assemble_document(self, caption, out):
        self.calls.append("assemble")
        with open(out, "w", encoding="utf-8") as fh:
            fh.write("brief")


def _run_worker(qtbot, monkeypatch, tmp_path, *, cancel_before=False, gen=None):
    from icharlotte_core.ui.wizard.pages import mediation_brief_page as mod
    from icharlotte_core.mediation_brief import GENERATION_ORDER

    fake = gen or _FakeGen()
    monkeypatch.setattr(mod, "MediationBriefGenerator", lambda: fake)
    monkeypatch.setattr(mod, "_read_documents", lambda paths: ("document text", []))

    caption = tmp_path / "Caption.docx"
    caption.write_text("caption", encoding="utf-8")
    case_path = str(tmp_path)
    settings = {"files": ["x.pdf"], "caption_path": str(caption)}

    worker = mod.MediationBriefWizardWorker(case_path, "001", settings)
    progress = []
    results = []
    worker.progress.connect(progress.append)
    worker.finished_result.connect(lambda ok, payload: results.append((ok, payload)))
    if cancel_before:
        worker.cancel()
    worker.run()  # run synchronously in the test thread
    return fake, progress, results, GENERATION_ORDER


def test_worker_success_runs_full_pipeline(qtbot, monkeypatch, tmp_path):
    fake, progress, results, order = _run_worker(qtbot, monkeypatch, tmp_path)

    assert fake.calls[0] == "style"
    assert fake.calls[1] == "planning"
    for name in order:
        assert f"section:{name}" in fake.calls
    assert fake.is_active is True
    assert len(results) == 1
    ok, path = results[0]
    assert ok is True
    assert os.path.join("NOTES", "AI OUTPUT", "MEDIATION") in path
    assert os.path.isfile(path)


def test_worker_cancel_stops_before_sections(qtbot, monkeypatch, tmp_path):
    fake, progress, results, order = _run_worker(
        qtbot, monkeypatch, tmp_path, cancel_before=True
    )

    assert results == [(False, "Generation cancelled.")]
    assert not any(c.startswith("section:") for c in fake.calls)


def test_worker_reports_engine_error(qtbot, monkeypatch, tmp_path):
    class BoomGen(_FakeGen):
        def run_planning_pass(self):
            raise RuntimeError("boom")

    fake, progress, results, order = _run_worker(
        qtbot, monkeypatch, tmp_path, gen=BoomGen()
    )

    assert len(results) == 1
    ok, payload = results[0]
    assert ok is False
    assert "boom" in payload


def test_worker_fails_on_empty_content(qtbot, monkeypatch, tmp_path):
    from icharlotte_core.ui.wizard.pages import mediation_brief_page as mod

    monkeypatch.setattr(mod, "MediationBriefGenerator", lambda: _FakeGen())
    monkeypatch.setattr(mod, "_read_documents", lambda paths: ("", ["nothing"]))
    caption = tmp_path / "Caption.docx"
    caption.write_text("caption", encoding="utf-8")

    worker = mod.MediationBriefWizardWorker(
        str(tmp_path), "001",
        {"files": ["x.pdf"], "caption_path": str(caption)},
    )
    results = []
    worker.finished_result.connect(lambda ok, payload: results.append((ok, payload)))
    worker.run()

    assert results and results[0][0] is False
    assert "could not read" in results[0][1].lower()


# ---- Settings page ----

def test_settings_page_autodetects_caption(qtbot, tmp_path):
    from icharlotte_core.ui.wizard.pages.mediation_brief_page import (
        MediationBriefSettingsPage,
    )

    (tmp_path / "Caption Page.docx").write_text("x", encoding="utf-8")
    page = MediationBriefSettingsPage(str(tmp_path), "001")
    qtbot.addWidget(page)

    assert page.caption_edit.text().endswith("Caption Page.docx")


def test_settings_page_add_dedupes_and_roundtrips(qtbot, tmp_path):
    from icharlotte_core.ui.wizard.pages.mediation_brief_page import (
        MediationBriefSettingsPage,
    )

    page = MediationBriefSettingsPage(str(tmp_path), "001")
    qtbot.addWidget(page)

    page.files_list.addItem("a.pdf")
    page.files_list.addItem("a.pdf")  # duplicate added manually
    page.from_dict({"files": ["b.pdf", "c.pdf"], "caption_path": "cap.docx"})

    assert page.current_files() == ["b.pdf", "c.pdf"]
    assert page.to_dict() == {"files": ["b.pdf", "c.pdf"], "caption_path": "cap.docx"}


def test_settings_page_validation_blocks_without_files(qtbot, tmp_path, monkeypatch):
    from icharlotte_core.ui.wizard.pages import mediation_brief_page as mod

    page = mod.MediationBriefSettingsPage(str(tmp_path), "001")
    qtbot.addWidget(page)
    monkeypatch.setattr(mod.QMessageBox, "warning", lambda *a, **k: None)

    emitted = []
    page.run_requested.connect(emitted.append)
    page._on_generate()  # no files, no caption

    assert emitted == []


def test_settings_page_emits_run_requested(qtbot, tmp_path):
    from icharlotte_core.ui.wizard.pages.mediation_brief_page import (
        MediationBriefSettingsPage,
    )

    caption = tmp_path / "Caption.docx"
    caption.write_text("x", encoding="utf-8")
    page = MediationBriefSettingsPage(str(tmp_path), "001")
    qtbot.addWidget(page)
    page.files_list.addItem("doc1.pdf")
    page.caption_edit.setText(str(caption))

    with qtbot.waitSignal(page.run_requested, timeout=500) as blocker:
        page._on_generate()

    payload = blocker.args[0]
    assert payload["files"] == ["doc1.pdf"]
    assert payload["caption_path"] == str(caption)
    assert payload["suggested_filename"].endswith(".docx")
    assert payload["save_default_dir"].endswith(os.path.join("AI OUTPUT", "MEDIATION"))
