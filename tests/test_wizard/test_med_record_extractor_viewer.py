"""Tests for the Med Record Extractor chronology viewer."""
from __future__ import annotations

import tempfile
from pathlib import Path
from unittest.mock import patch

import pytest
from docx import Document
from PySide6.QtCore import QSettings, Qt
from PySide6.QtGui import QColor, QPixmap
from PySide6.QtWidgets import QWidget

pytest.importorskip("pytestqt")

from tests.test_med_record_extractor import _build_chronology_docx


def _install_recording_pdf_viewer(monkeypatch):
    import icharlotte_core.ui.pdf_viewer_widget as viewer_mod

    calls = []

    class FakeViewer(QWidget):
        def __init__(self):
            super().__init__()
            self.current_page = 1

        def load_pdf(self, path):
            calls.append(("load", path))

        def go_to_page(self, page_number):
            calls.append(("go", page_number))
            self.current_page = page_number

        def get_current_page(self):
            return self.current_page

    monkeypatch.setattr(viewer_mod, "PdfViewerWidget", FakeViewer)
    return calls


@pytest.fixture(autouse=True)
def isolated_qsettings(tmp_path):
    previous_default_format = QSettings.defaultFormat()
    QSettings.setPath(
        QSettings.Format.IniFormat,
        QSettings.Scope.UserScope,
        str(tmp_path),
    )
    QSettings.setDefaultFormat(QSettings.Format.IniFormat)
    yield
    QSettings("iCharlotte", "iCharlotte").sync()
    QSettings.setDefaultFormat(previous_default_format)


def test_selection_page_loads_chronology_document(qtbot):
    from icharlotte_core.ui.wizard.pages.med_record_extractor_page import (
        MedChronologySelectionPage,
    )

    with tempfile.TemporaryDirectory() as td:
        source = _build_chronology_docx(Path(td) / "chronology.docx")
        page = MedChronologySelectionPage(
            case_path=td,
            file_number="5800.013",
            chronology_path=str(source),
        )
        qtbot.addWidget(page)

        assert page.selected_count_label.text() == "0 rows selected"
        assert page.synopsis_panel.count() == 2
        assert page.table_panel.count() == 2


def test_selection_page_uses_compact_top_layout(qtbot, tmp_path):
    from PySide6.QtWidgets import QLabel

    from icharlotte_core.ui.wizard import theme
    from icharlotte_core.ui.wizard.pages.med_record_extractor_page import (
        MedChronologySelectionPage,
    )

    source = _build_chronology_docx(tmp_path / "chronology.docx")
    page = MedChronologySelectionPage(str(tmp_path), "5800.013", str(source))
    qtbot.addWidget(page)

    margins = page.layout().contentsMargins()

    assert margins.top() <= theme.SPACE_SM
    assert page.layout().indexOf(page.preview_splitter) <= 2
    assert all(
        label.text() != "Select Medical Chronology Records"
        for label in page.findChildren(QLabel)
    )


def test_chronology_column_widths_persist_globally(qtbot, tmp_path):
    from icharlotte_core.ui.wizard.pages.med_record_extractor_page import (
        MedChronologySelectionPage,
    )

    first_source = _build_chronology_docx(tmp_path / "first_chronology.docx")
    second_source = _build_chronology_docx(tmp_path / "second_chronology.docx")
    first = MedChronologySelectionPage(
        case_path=str(tmp_path / "first_case"),
        file_number="5800.013",
        chronology_path=str(first_source),
    )
    qtbot.addWidget(first)

    first.table_panel.setColumnWidth(4, 640)
    first.table_panel.save_column_widths()

    second = MedChronologySelectionPage(
        case_path=str(tmp_path / "second_case"),
        file_number="5800.014",
        chronology_path=str(second_source),
    )
    qtbot.addWidget(second)

    assert second.table_panel.columnWidth(4) == 640


def test_chronology_row_heights_auto_fit_wrapped_text(qtbot, tmp_path):
    from icharlotte_core.ui.wizard.pages.med_record_extractor_page import (
        MedChronologySelectionPage,
    )

    source = tmp_path / "chronology.docx"
    doc = Document()
    doc.add_paragraph("BRIEF SYNOPSIS OF POST-INJURY MEDICAL RECORD:")
    doc.add_paragraph("On 09/21/2020, Test Plaintiff saw Kaiser Permanente.")
    table = doc.add_table(rows=1, cols=5)
    for index, header in enumerate([
        "DATE",
        "PAGE NO",
        "PROVIDER",
        "DESCRIPTION",
        "Red Flags/Comments",
    ]):
        table.rows[0].cells[index].text = header
    row = table.add_row().cells
    row[0].text = "09/21/2020"
    row[1].text = "Record\n\nPg No: 1/2"
    row[2].text = "Kaiser Permanente"
    row[3].text = (
        "EMERGENCY DEPARTMENT NOTE\n"
        "This description is intentionally long enough to wrap across several "
        "visual lines when the description column is narrowed for review."
    )
    row[4].text = ""
    doc.save(source)

    page = MedChronologySelectionPage(str(tmp_path), "5800.013", str(source))
    qtbot.addWidget(page)
    page.table_panel.setColumnWidth(4, 120)
    page.table_panel.resizeRowsToContents()

    assert (
        page.table_panel.rowHeight(0)
        > page.table_panel.verticalHeader().defaultSectionSize()
    )


def test_brief_synopsis_entries_wrap_and_fit_full_text(qtbot, tmp_path):
    from icharlotte_core.ui.wizard.pages.med_record_extractor_page import (
        MedChronologySelectionPage,
    )

    long_synopsis = (
        "On 09/22/2020, Test Plaintiff returned to Kaiser Permanente and "
        "reported ongoing right ankle pain, difficulty bearing weight, swelling, "
        "and trouble using crutches while completing basic activities of daily "
        "living after the fall."
    )
    source = _build_chronology_docx(
        tmp_path / "chronology.docx",
        synopsis_extra_paragraphs=(long_synopsis,),
    )
    page = MedChronologySelectionPage(str(tmp_path), "5800.013", str(source))
    qtbot.addWidget(page)

    assert page.synopsis_panel.wordWrap() is True
    assert page.synopsis_panel.textElideMode() == Qt.TextElideMode.ElideNone

    page.synopsis_panel.setFixedWidth(240)
    page.synopsis_panel.doItemsLayout()

    wrapped_height = page.synopsis_panel.sizeHintForRow(1)
    single_line_height = page.synopsis_panel.fontMetrics().height()
    assert wrapped_height > single_line_height * 2


def test_checked_synopsis_entry_gets_yellow_full_item_highlight(qtbot, tmp_path):
    from icharlotte_core.ui.wizard.pages.med_record_extractor_page import (
        SELECTION_HIGHLIGHT_COLOR,
        MedChronologySelectionPage,
    )

    source = _build_chronology_docx(tmp_path / "chronology.docx")
    page = MedChronologySelectionPage(str(tmp_path), "5800.013", str(source))
    qtbot.addWidget(page)

    item = page.synopsis_panel.item(0)
    page.set_paragraph_checked(page.document.synopsis_paragraphs[0].id, True)

    assert item.background().color().name() == SELECTION_HIGHLIGHT_COLOR

    page.set_paragraph_checked(page.document.synopsis_paragraphs[0].id, False)

    assert item.background().style() == Qt.BrushStyle.NoBrush


def test_checked_synopsis_entry_paints_full_row_and_checkbox_yellow(qtbot, tmp_path):
    from icharlotte_core.ui.wizard.pages.med_record_extractor_page import (
        SELECTION_HIGHLIGHT_COLOR,
        MedChronologySelectionPage,
    )

    source = _build_chronology_docx(tmp_path / "chronology.docx")
    page = MedChronologySelectionPage(str(tmp_path), "5800.013", str(source))
    qtbot.addWidget(page)
    page.resize(900, 500)
    page.show()
    qtbot.waitExposed(page)

    item = page.synopsis_panel.item(0)
    page.set_paragraph_checked(page.document.synopsis_paragraphs[0].id, True)
    page.synopsis_panel.viewport().update()
    qtbot.wait(50)

    rect = page.synopsis_panel.visualItemRect(item)
    pixmap = QPixmap(page.synopsis_panel.viewport().size())
    page.synopsis_panel.viewport().render(pixmap)
    image = pixmap.toImage()

    expected = QColor(SELECTION_HIGHLIGHT_COLOR)
    row_pixel = QColor(image.pixel(rect.right() - 20, rect.center().y()))
    checkbox_pixels = [
        QColor(image.pixel(rect.left() + x, rect.center().y()))
        for x in range(5, 31)
    ]

    assert row_pixel.name() == expected.name()
    assert any(pixel.name() == expected.name() for pixel in checkbox_pixels)
    assert all(pixel.name() != "#1976d2" for pixel in checkbox_pixels)


def test_checked_chronology_row_gets_yellow_full_row_highlight(qtbot, tmp_path):
    from icharlotte_core.ui.wizard.pages.med_record_extractor_page import (
        SELECTION_HIGHLIGHT_COLOR,
        MedChronologySelectionPage,
    )

    source = _build_chronology_docx(tmp_path / "chronology.docx")
    page = MedChronologySelectionPage(str(tmp_path), "5800.013", str(source))
    qtbot.addWidget(page)

    row_id = page.document.rows[0].id
    page.set_row_checked(row_id, True)

    assert [
        page.table_panel.item(0, column).background().color().name()
        for column in range(page.table_panel.columnCount())
    ] == [SELECTION_HIGHLIGHT_COLOR] * page.table_panel.columnCount()

    page.set_row_checked(row_id, False)

    assert all(
        page.table_panel.item(0, column).background().style()
        == Qt.BrushStyle.NoBrush
        for column in range(page.table_panel.columnCount())
    )


def test_chronology_column_resize_debounces_save_and_row_autofit(
    qtbot,
    tmp_path,
    monkeypatch,
):
    from icharlotte_core.ui.wizard.pages.med_record_extractor_page import (
        MedChronologySelectionPage,
    )

    source = _build_chronology_docx(tmp_path / "chronology.docx")
    page = MedChronologySelectionPage(str(tmp_path), "5800.013", str(source))
    qtbot.addWidget(page)

    calls = []
    monkeypatch.setattr(
        page.table_panel,
        "save_column_widths",
        lambda: calls.append("save"),
    )
    monkeypatch.setattr(
        page.table_panel,
        "resizeRowsToContents",
        lambda: calls.append("resize"),
    )

    page.table_panel._on_section_resized(4, 520, 640)

    assert calls == []
    assert page.table_panel._resize_commit_timer.isActive()

    page.table_panel._commit_column_resize()

    assert calls == ["save", "resize"]


def test_selection_page_uses_brief_synopsis_and_chronology_row_tabs(qtbot):
    from icharlotte_core.ui.wizard.pages.med_record_extractor_page import (
        MedChronologySelectionPage,
    )

    with tempfile.TemporaryDirectory() as td:
        source = _build_chronology_docx(Path(td) / "chronology.docx")
        page = MedChronologySelectionPage(
            case_path=td,
            file_number="5800.013",
            chronology_path=str(source),
        )
        qtbot.addWidget(page)

        assert page.tab_widget.count() == 2
        assert page.tab_widget.tabText(0) == "Brief Synopsis"
        assert page.tab_widget.tabText(1) == "Chronology Rows"
        assert page.synopsis_panel.count() == 2
        assert page.table_panel.count() == 2


def test_ctrl_f_opens_find_box_and_searches_synopsis_entries(qtbot, tmp_path):
    from icharlotte_core.ui.wizard.pages.med_record_extractor_page import (
        MedChronologySelectionPage,
    )

    source = _build_chronology_docx(tmp_path / "chronology.docx")
    page = MedChronologySelectionPage(str(tmp_path), "5800.013", str(source))
    qtbot.addWidget(page)
    page.show()
    qtbot.waitExposed(page)

    assert not page.find_bar.isVisible()

    qtbot.keyClick(page, Qt.Key.Key_F, modifier=Qt.KeyboardModifier.ControlModifier)

    assert page.find_bar.isVisible()
    qtbot.waitUntil(page.find_input.hasFocus, timeout=500)

    page.find_input.setText("x-ray")

    assert page.synopsis_panel.currentRow() == 1
    assert page.find_status_label.text() == "1 of 1"


def test_find_highlights_matching_synopsis_text_and_clears_on_close(qtbot, tmp_path):
    from icharlotte_core.ui.wizard.pages.med_record_extractor_page import (
        MedChronologySelectionPage,
        search_highlight_html,
    )

    source = _build_chronology_docx(tmp_path / "chronology.docx")
    page = MedChronologySelectionPage(str(tmp_path), "5800.013", str(source))
    qtbot.addWidget(page)
    page.show()
    qtbot.waitExposed(page)

    qtbot.keyClick(page, Qt.Key.Key_F, modifier=Qt.KeyboardModifier.ControlModifier)
    page.find_input.setText("x-ray")

    assert page.synopsis_panel.find_highlight_query == "x-ray"
    assert page.synopsis_panel.find_highlight_rows == {1}
    assert page.synopsis_panel.active_find_highlight_row == 1
    assert page.table_panel.find_highlight_query == ""
    assert "background-color" in search_highlight_html(
        page.synopsis_panel.item(1).text(),
        "x-ray",
        active=True,
    )

    qtbot.keyPress(page.find_input, Qt.Key.Key_Escape)

    assert page.synopsis_panel.find_highlight_query == ""
    assert page.synopsis_panel.find_highlight_rows == set()
    assert page.synopsis_panel.active_find_highlight_row == -1


def test_find_searches_active_chronology_row_text_and_cycles_matches(
    qtbot,
    tmp_path,
):
    from icharlotte_core.ui.wizard.pages.med_record_extractor_page import (
        MedChronologySelectionPage,
    )

    source = _build_chronology_docx(tmp_path / "chronology.docx")
    page = MedChronologySelectionPage(str(tmp_path), "5800.013", str(source))
    qtbot.addWidget(page)
    page.show()
    qtbot.waitExposed(page)
    page.tab_widget.setCurrentIndex(1)

    qtbot.keyClick(page, Qt.Key.Key_F, modifier=Qt.KeyboardModifier.ControlModifier)
    page.find_input.setText("Kaiser Permanente")

    assert page.table_panel.currentRow() == 0
    assert page.find_status_label.text() == "1 of 2"
    assert page.table_panel.find_highlight_query == "kaiser permanente"
    assert page.table_panel.find_highlight_rows == {0, 1}
    assert page.table_panel.active_find_highlight_row == 0
    assert page.selected_count_label.text() == "0 rows selected"
    assert all(not page.is_row_checked(row.id) for row in page.document.rows)

    qtbot.keyPress(page.find_input, Qt.Key.Key_Return)

    assert page.table_panel.currentRow() == 1
    assert page.find_status_label.text() == "2 of 2"
    assert page.table_panel.active_find_highlight_row == 1
    assert page.selected_count_label.text() == "0 rows selected"
    assert all(not page.is_row_checked(row.id) for row in page.document.rows)

    page.find_input.setText("Kaiser Permanente James")

    assert page.table_panel.currentRow() == 1
    assert page.find_status_label.text() == "1 of 1"


def test_clicking_synopsis_text_toggles_entry_selection(qtbot):
    from icharlotte_core.ui.wizard.pages.med_record_extractor_page import (
        MedChronologySelectionPage,
    )

    with tempfile.TemporaryDirectory() as td:
        source = _build_chronology_docx(Path(td) / "chronology.docx")
        page = MedChronologySelectionPage(
            case_path=td,
            file_number="5800.013",
            chronology_path=str(source),
        )
        qtbot.addWidget(page)

        item = page.synopsis_panel.item(0)
        click_pos = page.synopsis_panel.visualItemRect(item).center()

        qtbot.mouseClick(
            page.synopsis_panel.viewport(),
            Qt.MouseButton.LeftButton,
            pos=click_pos,
        )

        assert item.checkState() == Qt.CheckState.Checked
        assert page.is_row_checked(page.document.rows[0].id)

        qtbot.mouseClick(
            page.synopsis_panel.viewport(),
            Qt.MouseButton.LeftButton,
            pos=click_pos,
        )

        assert item.checkState() == Qt.CheckState.Unchecked
        assert not page.is_row_checked(page.document.rows[0].id)


def test_right_clicking_synopsis_text_does_not_toggle_entry_selection(qtbot):
    from icharlotte_core.ui.wizard.pages.med_record_extractor_page import (
        MedChronologySelectionPage,
    )

    with tempfile.TemporaryDirectory() as td:
        source = _build_chronology_docx(Path(td) / "chronology.docx")
        page = MedChronologySelectionPage(
            case_path=td,
            file_number="5800.013",
            chronology_path=str(source),
        )
        qtbot.addWidget(page)

        item = page.synopsis_panel.item(0)
        click_pos = page.synopsis_panel.visualItemRect(item).center()

        qtbot.mouseClick(
            page.synopsis_panel.viewport(),
            Qt.MouseButton.RightButton,
            pos=click_pos,
        )

        assert item.checkState() == Qt.CheckState.Unchecked
        assert not page.is_row_checked(page.document.rows[0].id)


def test_clicking_chronology_row_text_toggles_row_selection(qtbot):
    from icharlotte_core.ui.wizard.pages.med_record_extractor_page import (
        MedChronologySelectionPage,
    )

    with tempfile.TemporaryDirectory() as td:
        source = _build_chronology_docx(Path(td) / "chronology.docx")
        page = MedChronologySelectionPage(
            case_path=td,
            file_number="5800.013",
            chronology_path=str(source),
        )
        qtbot.addWidget(page)
        page.tab_widget.setCurrentIndex(1)

        click_pos = page.table_panel.visualRect(
            page.table_panel.model().index(0, 4)
        ).center()

        qtbot.mouseClick(
            page.table_panel.viewport(),
            Qt.MouseButton.LeftButton,
            pos=click_pos,
        )

        assert page.is_row_checked(page.document.rows[0].id)
        assert page.selected_count_label.text() == "1 row selected"

        qtbot.mouseClick(
            page.table_panel.viewport(),
            Qt.MouseButton.LeftButton,
            pos=click_pos,
        )

        assert not page.is_row_checked(page.document.rows[0].id)
        assert page.selected_count_label.text() == "0 rows selected"


def test_right_clicking_chronology_row_text_does_not_toggle_row_selection(qtbot):
    from icharlotte_core.ui.wizard.pages.med_record_extractor_page import (
        MedChronologySelectionPage,
    )

    with tempfile.TemporaryDirectory() as td:
        source = _build_chronology_docx(Path(td) / "chronology.docx")
        page = MedChronologySelectionPage(
            case_path=td,
            file_number="5800.013",
            chronology_path=str(source),
        )
        qtbot.addWidget(page)
        page.tab_widget.setCurrentIndex(1)

        click_pos = page.table_panel.visualRect(
            page.table_panel.model().index(0, 4)
        ).center()

        qtbot.mouseClick(
            page.table_panel.viewport(),
            Qt.MouseButton.RightButton,
            pos=click_pos,
        )

        assert not page.is_row_checked(page.document.rows[0].id)
        assert page.selected_count_label.text() == "0 rows selected"


def test_double_clicking_chronology_row_loads_cited_pdf_in_embedded_preview(
    qtbot,
    tmp_path,
    monkeypatch,
):
    from icharlotte_core.ui.wizard.pages.med_record_extractor_page import (
        MedChronologySelectionPage,
    )

    records = tmp_path / "RECORDS"
    records.mkdir()
    pdf = records / "Hall - Doc Produced HALL 000001 to 002530 7-21-2023.pdf"
    pdf.write_bytes(b"%PDF-1.4\n")
    source = _build_chronology_docx(tmp_path / "chronology.docx")
    calls = _install_recording_pdf_viewer(monkeypatch)

    page = MedChronologySelectionPage(str(tmp_path), "5800.013", str(source))
    qtbot.addWidget(page)
    page.tab_widget.setCurrentIndex(1)

    click_pos = page.table_panel.visualRect(
        page.table_panel.model().index(0, 4)
    ).center()
    qtbot.mouseDClick(
        page.table_panel.viewport(),
        Qt.MouseButton.LeftButton,
        pos=click_pos,
    )

    assert calls == [("load", str(pdf)), ("go", 599)]
    assert page.pdf_preview_status_label.text() == f"{pdf.name} - page 599"


def test_single_clicking_chronology_row_does_not_load_pdf_preview(
    qtbot,
    tmp_path,
    monkeypatch,
):
    from icharlotte_core.ui.wizard.pages.med_record_extractor_page import (
        MedChronologySelectionPage,
    )

    records = tmp_path / "RECORDS"
    records.mkdir()
    pdf = records / "Hall - Doc Produced HALL 000001 to 002530 7-21-2023.pdf"
    pdf.write_bytes(b"%PDF-1.4\n")
    source = _build_chronology_docx(tmp_path / "chronology.docx")
    calls = _install_recording_pdf_viewer(monkeypatch)

    page = MedChronologySelectionPage(str(tmp_path), "5800.013", str(source))
    qtbot.addWidget(page)
    page.tab_widget.setCurrentIndex(1)

    click_pos = page.table_panel.visualRect(
        page.table_panel.model().index(0, 4)
    ).center()
    qtbot.mouseClick(
        page.table_panel.viewport(),
        Qt.MouseButton.LeftButton,
        pos=click_pos,
    )

    assert page.is_row_checked(page.document.rows[0].id)
    assert calls == []
    assert page.pdf_preview_status_label.text() == "Double-click a chronology entry to preview its source PDF."


def test_opening_multiple_entries_reuses_file_index_cache(
    qtbot,
    tmp_path,
    monkeypatch,
):
    from icharlotte_core.ui.wizard.pages import med_record_extractor_page as mod
    from icharlotte_core.ui.wizard.pages.med_record_extractor_page import (
        MedChronologySelectionPage,
    )

    records = tmp_path / "RECORDS"
    records.mkdir()
    pdf = records / "Hall - Doc Produced HALL 000001 to 002530 7-21-2023.pdf"
    pdf.write_bytes(b"%PDF-1.4\n")
    source = _build_chronology_docx(tmp_path / "chronology.docx")
    calls = _install_recording_pdf_viewer(monkeypatch)
    index_calls = []
    file_index = {pdf.stem.lower(): str(pdf)}

    def fake_build_file_index(case_path):
        index_calls.append(case_path)
        return file_index

    monkeypatch.setattr(mod, "_build_file_index", fake_build_file_index)
    page = MedChronologySelectionPage(str(tmp_path), "5800.013", str(source))
    qtbot.addWidget(page)

    page._open_pdf_for_row(page.document.rows[0])
    page._open_pdf_for_row(page.document.rows[1])

    assert index_calls == [str(tmp_path)]
    assert calls == [
        ("load", str(pdf)),
        ("go", 599),
        ("load", str(pdf)),
        ("go", 598),
    ]


def test_opening_entry_refreshes_file_index_after_cache_miss(
    qtbot,
    tmp_path,
    monkeypatch,
):
    from icharlotte_core.ui.wizard.pages import med_record_extractor_page as mod
    from icharlotte_core.ui.wizard.pages.med_record_extractor_page import (
        MedChronologySelectionPage,
    )

    records = tmp_path / "RECORDS"
    records.mkdir()
    pdf = records / "Hall - Doc Produced HALL 000001 to 002530 7-21-2023.pdf"
    pdf.write_bytes(b"%PDF-1.4\n")
    source = _build_chronology_docx(tmp_path / "chronology.docx")
    calls = _install_recording_pdf_viewer(monkeypatch)
    warnings = []
    index_calls = []
    indexes = [{}, {pdf.stem.lower(): str(pdf)}]

    def fake_build_file_index(case_path):
        index_calls.append(case_path)
        return indexes.pop(0)

    monkeypatch.setattr(mod, "_build_file_index", fake_build_file_index)
    monkeypatch.setattr(
        mod.QMessageBox,
        "warning",
        lambda *args: warnings.append(args),
    )
    page = MedChronologySelectionPage(str(tmp_path), "5800.013", str(source))
    qtbot.addWidget(page)

    page._open_pdf_for_row(page.document.rows[0])

    assert index_calls == [str(tmp_path), str(tmp_path)]
    assert calls == [("load", str(pdf)), ("go", 599)]
    assert warnings == []


def test_embedded_pdf_preview_requests_target_page_without_startup_delay(
    qtbot,
    tmp_path,
    monkeypatch,
):
    from icharlotte_core.ui.wizard.pages.med_record_extractor_page import (
        MedChronologySelectionPage,
    )

    pdf = tmp_path / "record.pdf"
    pdf.write_bytes(b"%PDF-1.4\n")
    source = _build_chronology_docx(tmp_path / "chronology.docx")
    calls = _install_recording_pdf_viewer(monkeypatch)

    page = MedChronologySelectionPage(str(tmp_path), "5800.013", str(source))
    qtbot.addWidget(page)
    page._show_pdf_preview(str(pdf), 599)

    assert calls == [("load", str(pdf)), ("go", 599)]


def test_collapsed_pdf_preview_reopens_when_entry_is_opened(
    qtbot,
    tmp_path,
    monkeypatch,
):
    from icharlotte_core.ui.wizard.pages.med_record_extractor_page import (
        MedChronologySelectionPage,
    )

    pdf = tmp_path / "record.pdf"
    pdf.write_bytes(b"%PDF-1.4\n")
    source = _build_chronology_docx(tmp_path / "chronology.docx")
    calls = _install_recording_pdf_viewer(monkeypatch)

    page = MedChronologySelectionPage(str(tmp_path), "5800.013", str(source))
    qtbot.addWidget(page)

    page.toggle_pdf_preview_collapse()

    assert page._pdf_preview_collapsed is True
    assert page.pdf_preview_panel.isHidden()
    assert page.pdf_preview_toggle_btn.text() == "Show Preview"

    page._show_pdf_preview(str(pdf), 599)

    assert page._pdf_preview_collapsed is False
    assert not page.pdf_preview_panel.isHidden()
    assert page.pdf_preview_toggle_btn.text() == "Hide Preview"
    assert calls == [("load", str(pdf)), ("go", 599)]


def test_double_clicking_synopsis_loads_confident_match_pdf_in_embedded_preview(
    qtbot,
    tmp_path,
    monkeypatch,
):
    from icharlotte_core.ui.wizard.pages.med_record_extractor_page import (
        MedChronologySelectionPage,
    )

    records = tmp_path / "RECORDS"
    records.mkdir()
    pdf = records / "Hall - Doc Produced HALL 000001 to 002530 7-21-2023.pdf"
    pdf.write_bytes(b"%PDF-1.4\n")
    source = _build_chronology_docx(tmp_path / "chronology.docx")
    calls = _install_recording_pdf_viewer(monkeypatch)

    page = MedChronologySelectionPage(str(tmp_path), "5800.013", str(source))
    qtbot.addWidget(page)

    item = page.synopsis_panel.item(0)
    click_pos = page.synopsis_panel.visualItemRect(item).center()
    qtbot.mouseDClick(
        page.synopsis_panel.viewport(),
        Qt.MouseButton.LeftButton,
        pos=click_pos,
    )

    assert calls == [("load", str(pdf)), ("go", 599)]
    assert page.pdf_preview_status_label.text() == f"{pdf.name} - page 599"


def test_chronology_row_tab_displays_all_source_columns_verbatim(qtbot):
    from icharlotte_core.ui.wizard.pages.med_record_extractor_page import (
        MedChronologySelectionPage,
    )

    with tempfile.TemporaryDirectory() as td:
        source = _build_chronology_docx(Path(td) / "chronology.docx")
        page = MedChronologySelectionPage(
            case_path=td,
            file_number="5800.013",
            chronology_path=str(source),
        )
        qtbot.addWidget(page)

        assert page.table_panel.columnCount() == 6
        assert page.table_panel.horizontalHeaderItem(1).text() == "DATE"
        assert page.table_panel.horizontalHeaderItem(2).text() == "PAGE NO"
        assert page.table_panel.horizontalHeaderItem(3).text() == "PROVIDER"
        assert page.table_panel.horizontalHeaderItem(4).text() == "DESCRIPTION"
        assert page.table_panel.horizontalHeaderItem(5).text() == "Red Flags/Comments"
        assert page.table_panel.item(0, 2).text() == (
            "Hall - Doc Produced HALL 000001 to 002530 7-21-2023\n\n"
            "Pg No: 599-604/2530"
        )
        assert page.table_panel.item(0, 3).text() == "Kaiser Permanente\nHenry Louis Marr, DO"
        assert page.table_panel.item(0, 4).text() == (
            "EMERGENCY DEPARTMENT NOTE\nChief Complaint: Ankle injury."
        )


def test_direct_row_selection_emits_selected_rows(qtbot):
    from icharlotte_core.ui.wizard.pages.med_record_extractor_page import (
        MedChronologySelectionPage,
    )

    with tempfile.TemporaryDirectory() as td:
        source = _build_chronology_docx(Path(td) / "chronology.docx")
        page = MedChronologySelectionPage(
            case_path=td,
            file_number="5800.013",
            chronology_path=str(source),
        )
        qtbot.addWidget(page)
        page.set_row_checked(page.document.rows[0].id, True)

        with qtbot.waitSignal(page.run_requested, timeout=500) as blocker:
            page.extract_btn.click()

        settings = blocker.args[0]
        assert settings["chronology_path"] == str(source)
        assert [row.id for row in settings["selected_rows"]] == [page.document.rows[0].id]


def test_synopsis_selection_auto_selects_confident_row(qtbot):
    from icharlotte_core.ui.wizard.pages.med_record_extractor_page import (
        MedChronologySelectionPage,
    )

    with tempfile.TemporaryDirectory() as td:
        source = _build_chronology_docx(Path(td) / "chronology.docx")
        page = MedChronologySelectionPage(
            case_path=td,
            file_number="5800.013",
            chronology_path=str(source),
        )
        qtbot.addWidget(page)
        page.set_paragraph_checked(page.document.synopsis_paragraphs[0].id, True)

        assert page.is_row_checked(page.document.rows[0].id)
        assert page.selected_count_label.text() == "1 row selected"


def test_deselecting_synopsis_owned_row_keeps_visual_selection_in_sync(qtbot):
    from icharlotte_core.ui.wizard.pages.med_record_extractor_page import (
        MedChronologySelectionPage,
    )

    with tempfile.TemporaryDirectory() as td:
        source = _build_chronology_docx(Path(td) / "chronology.docx")
        page = MedChronologySelectionPage(
            case_path=td,
            file_number="5800.013",
            chronology_path=str(source),
        )
        qtbot.addWidget(page)

        row_id = page.document.rows[0].id
        page.set_paragraph_checked(page.document.synopsis_paragraphs[0].id, True)
        page.set_row_checked(row_id, False)

        assert page.is_row_checked(row_id)
        assert page.selected_count_label.text() == "1 row selected"


def test_ambiguous_synopsis_selection_does_not_auto_select_rows(qtbot):
    from icharlotte_core.ui.wizard.pages.med_record_extractor_page import (
        MedChronologySelectionPage,
    )

    with tempfile.TemporaryDirectory() as td:
        source = _build_chronology_docx(
            Path(td) / "chronology.docx",
            synopsis_extra_paragraphs=(
                "On 09/21/2020, she presented to Kaiser Permanente for ankle care.",
            ),
        )
        page = MedChronologySelectionPage(
            case_path=td,
            file_number="5800.013",
            chronology_path=str(source),
        )
        qtbot.addWidget(page)

        ambiguous = page.document.synopsis_paragraphs[1]
        page.set_paragraph_checked(ambiguous.id, True)

        assert all(not page.is_row_checked(row.id) for row in page.document.rows)
        assert page.selected_count_label.text() == "0 rows selected"


def test_open_original_uses_os_startfile(qtbot):
    from icharlotte_core.ui.wizard.pages.med_record_extractor_page import (
        MedChronologySelectionPage,
    )

    with tempfile.TemporaryDirectory() as td:
        source = _build_chronology_docx(Path(td) / "chronology.docx")
        page = MedChronologySelectionPage(
            case_path=td,
            file_number="5800.013",
            chronology_path=str(source),
        )
        qtbot.addWidget(page)

        with patch(
            "icharlotte_core.ui.wizard.pages.med_record_extractor_page.os.startfile"
        ) as startfile:
            page.open_original_btn.click()

        startfile.assert_called_once_with(str(source))


def test_blocking_parse_errors_disable_extract(qtbot):
    from icharlotte_core.ui.wizard.pages.med_record_extractor_page import (
        MedChronologySelectionPage,
    )

    with tempfile.TemporaryDirectory() as td:
        source = Path(td) / "no_table.docx"
        doc = Document()
        doc.add_paragraph("BRIEF SYNOPSIS OF POST-INJURY MEDICAL RECORD:")
        doc.add_paragraph("On 09/21/2020, Test Plaintiff saw Kaiser Permanente.")
        doc.save(source)

        page = MedChronologySelectionPage(
            case_path=td,
            file_number="5800.013",
            chronology_path=str(source),
        )
        qtbot.addWidget(page)

        assert page.extract_btn.isEnabled() is False
        assert page.selected_count_label.text() == "0 rows selected"


def test_non_extractable_rows_are_not_emitted(qtbot):
    from icharlotte_core.ui.wizard.pages.med_record_extractor_page import (
        MedChronologySelectionPage,
    )

    with tempfile.TemporaryDirectory() as td:
        source = _build_chronology_docx(
            Path(td) / "chronology.docx",
            first_page_no=(
                "Hall - Doc Produced HALL 000001 to 002530 7-21-2023\n\nPg No: 17-"
            ),
        )
        page = MedChronologySelectionPage(
            case_path=td,
            file_number="5800.013",
            chronology_path=str(source),
        )
        qtbot.addWidget(page)

        page.set_row_checked(page.document.rows[0].id, True)
        page.set_row_checked(page.document.rows[1].id, True)

        with qtbot.waitSignal(page.run_requested, timeout=500) as blocker:
            page.extract_btn.click()

        assert [row.id for row in blocker.args[0]["selected_rows"]] == [
            page.document.rows[1].id
        ]


def test_non_extractable_rows_cannot_remain_checked(qtbot):
    from icharlotte_core.ui.wizard.pages.med_record_extractor_page import (
        MedChronologySelectionPage,
    )

    with tempfile.TemporaryDirectory() as td:
        source = _build_chronology_docx(
            Path(td) / "chronology.docx",
            first_page_no=(
                "Hall - Doc Produced HALL 000001 to 002530 7-21-2023\n\nPg No: 17-"
            ),
        )
        page = MedChronologySelectionPage(
            case_path=td,
            file_number="5800.013",
            chronology_path=str(source),
        )
        qtbot.addWidget(page)

        page.set_row_checked(page.document.rows[0].id, True)

        assert page.is_row_checked(page.document.rows[0].id) is False
        assert page.selected_count_label.text() == "0 rows selected"
        assert page.extract_btn.isEnabled() is False


def test_synopsis_owned_non_extractable_row_is_not_selected(qtbot):
    from icharlotte_core.ui.wizard.pages.med_record_extractor_page import (
        MedChronologySelectionPage,
    )

    with tempfile.TemporaryDirectory() as td:
        source = _build_chronology_docx(
            Path(td) / "chronology.docx",
            first_page_no=(
                "Hall - Doc Produced HALL 000001 to 002530 7-21-2023\n\nPg No: 17-"
            ),
        )
        page = MedChronologySelectionPage(
            case_path=td,
            file_number="5800.013",
            chronology_path=str(source),
        )
        qtbot.addWidget(page)

        page.set_paragraph_checked(page.document.synopsis_paragraphs[0].id, True)

        assert page.is_row_checked(page.document.rows[0].id) is False
        assert page.selected_count_label.text() == "0 rows selected"
        assert page.extract_btn.isEnabled() is False
        assert "not extractable" in page.match_status_label.text()


def test_failed_synopsis_match_updates_visible_status(qtbot):
    from icharlotte_core.ui.wizard.pages.med_record_extractor_page import (
        MedChronologySelectionPage,
    )

    with tempfile.TemporaryDirectory() as td:
        source = _build_chronology_docx(
            Path(td) / "chronology.docx",
            synopsis_extra_paragraphs=(
                "She returned home with crutches and follow-up instructions.",
            ),
        )
        page = MedChronologySelectionPage(
            case_path=td,
            file_number="5800.013",
            chronology_path=str(source),
        )
        qtbot.addWidget(page)

        no_date = page.document.synopsis_paragraphs[1]
        page.set_paragraph_checked(no_date.id, True)

        assert "No date found" in page.match_status_label.text()
        assert page.match_status_label.isHidden() is False


def test_build_med_extractor_tab_opens_blank_page_without_picker(qtbot, tmp_path):
    from icharlotte_core.ui.wizard.in_process_task_tab import build_med_extractor_tab
    from icharlotte_core.ui.wizard.pages.med_record_extractor_page import (
        MedChronologySelectionPage,
    )
    from icharlotte_core.ui.wizard.registry import get_task

    summary_dir = tmp_path / "RECORDS" / "Medical Summary - DO NOT PRODUCE"
    summary_dir.mkdir(parents=True)

    with patch(
        "icharlotte_core.ui.wizard.in_process_task_tab.QFileDialog.getOpenFileName",
        side_effect=AssertionError("launch should not open the chronology picker"),
    ) as picker:
        tab = build_med_extractor_tab(
            get_task("med_record_extractor"),
            case_path=str(tmp_path),
            file_number="5800.013",
            parent=None,
        )

    qtbot.addWidget(tab)
    assert isinstance(tab.settings_page, MedChronologySelectionPage)
    assert picker.call_count == 0
    assert tab.settings_page.chronology_path == ""
    assert tab.settings_page.select_chronology_btn.text() == "Please select medical chronology"
    assert not tab.settings_page.empty_state.isHidden()
    assert tab.settings_page.preview_splitter.isHidden()


def test_blank_selection_page_button_loads_chronology_from_summary_folder(
    qtbot,
    tmp_path,
):
    from icharlotte_core.ui.wizard.pages.med_record_extractor_page import (
        MedChronologySelectionPage,
    )

    summary_dir = tmp_path / "RECORDS" / "Medical Summary - DO NOT PRODUCE"
    summary_dir.mkdir(parents=True)
    source = _build_chronology_docx(summary_dir / "chronology.docx")
    page = MedChronologySelectionPage(str(tmp_path), "5800.013", "")
    qtbot.addWidget(page)

    with patch(
        "icharlotte_core.ui.wizard.pages.med_record_extractor_page.QFileDialog.getOpenFileName",
        return_value=(str(source), "Word Documents (*.docx)"),
    ) as picker:
        page.select_chronology_btn.click()

    assert picker.call_args.args[2] == str(summary_dir)
    assert page.chronology_path == str(source)
    assert page.to_dict()["chronology_path"] == str(source)
    assert page.synopsis_panel.count() == 2
    assert page.table_panel.count() == 2
    assert page.empty_state.isHidden()
    assert not page.preview_splitter.isHidden()


def test_build_med_extractor_tab_hides_redundant_task_header(qtbot, tmp_path):
    from icharlotte_core.ui.wizard.in_process_task_tab import build_med_extractor_tab
    from icharlotte_core.ui.wizard.registry import get_task

    tab = build_med_extractor_tab(
        get_task("med_record_extractor"),
        case_path=str(tmp_path),
        file_number="5800.013",
        parent=None,
    )

    qtbot.addWidget(tab)

    assert tab.header.isHidden()


def test_build_med_extractor_tab_without_chronology_stays_blank(tmp_path):
    from icharlotte_core.ui.wizard.in_process_task_tab import build_med_extractor_tab
    from icharlotte_core.ui.wizard.registry import get_task

    with patch(
        "icharlotte_core.ui.wizard.in_process_task_tab.QFileDialog.getOpenFileName",
        side_effect=AssertionError("launch should not open the chronology picker"),
    ):
        tab = build_med_extractor_tab(
            get_task("med_record_extractor"),
            case_path=str(tmp_path),
            file_number="5800.013",
            parent=None,
        )

    assert tab is not None
    assert tab.settings_page.chronology_path == ""


def test_selection_page_to_dict_and_from_dict_restore_selection(qtbot, tmp_path):
    from icharlotte_core.ui.wizard.pages.med_record_extractor_page import (
        MedChronologySelectionPage,
    )

    source = _build_chronology_docx(tmp_path / "chronology.docx")
    page = MedChronologySelectionPage(str(tmp_path), "5800.013", str(source))
    qtbot.addWidget(page)
    page.set_row_checked(page.document.rows[0].id, True)
    saved = page.to_dict()

    restored = MedChronologySelectionPage(str(tmp_path), "5800.013", str(source))
    qtbot.addWidget(restored)
    restored.from_dict(saved)

    assert restored.is_row_checked(restored.document.rows[0].id)
    assert restored.selected_count_label.text() == "1 row selected"


def test_selection_page_restore_keeps_synopsis_row_ownership(qtbot, tmp_path):
    from icharlotte_core.ui.wizard.pages.med_record_extractor_page import (
        MedChronologySelectionPage,
    )

    source = _build_chronology_docx(tmp_path / "chronology.docx")
    page = MedChronologySelectionPage(str(tmp_path), "5800.013", str(source))
    qtbot.addWidget(page)
    paragraph_id = page.document.synopsis_paragraphs[0].id
    row_id = page.document.rows[0].id
    page.set_paragraph_checked(paragraph_id, True)
    saved = page.to_dict()

    restored = MedChronologySelectionPage(str(tmp_path), "5800.013", str(source))
    qtbot.addWidget(restored)
    restored.from_dict(saved)
    restored.set_paragraph_checked(paragraph_id, False)

    assert not restored.is_row_checked(row_id)
    assert restored.selected_count_label.text() == "0 rows selected"


def test_selection_page_restore_preserves_mixed_row_ownership(qtbot, tmp_path):
    from icharlotte_core.ui.wizard.pages.med_record_extractor_page import (
        MedChronologySelectionPage,
    )

    source = _build_chronology_docx(tmp_path / "chronology.docx")
    page = MedChronologySelectionPage(str(tmp_path), "5800.013", str(source))
    qtbot.addWidget(page)
    paragraph_id = page.document.synopsis_paragraphs[0].id
    row_id = page.document.rows[0].id
    page.set_row_checked(row_id, True)
    page.set_paragraph_checked(paragraph_id, True)
    saved = page.to_dict()

    restored = MedChronologySelectionPage(str(tmp_path), "5800.013", str(source))
    qtbot.addWidget(restored)
    restored.from_dict(saved)
    restored.set_paragraph_checked(paragraph_id, False)

    assert restored.is_row_checked(row_id)
    assert restored.selected_count_label.text() == "1 row selected"


def test_selection_page_from_dict_replaces_existing_selection(qtbot, tmp_path):
    from icharlotte_core.ui.wizard.pages.med_record_extractor_page import (
        MedChronologySelectionPage,
    )

    source = _build_chronology_docx(tmp_path / "chronology.docx")
    page = MedChronologySelectionPage(str(tmp_path), "5800.013", str(source))
    qtbot.addWidget(page)
    old_row_id = page.document.rows[0].id
    new_row_id = page.document.rows[1].id
    page.set_row_checked(old_row_id, True)

    page.from_dict({
        "chronology_path": str(source),
        "selected_row_ids": [new_row_id],
    })

    assert not page.is_row_checked(old_row_id)
    assert page.is_row_checked(new_row_id)
    assert page.selected_count_label.text() == "1 row selected"


def test_build_med_extractor_tab_restore_uses_saved_chronology_without_picker(qtbot, tmp_path):
    from icharlotte_core.med_record_chronology import parse_chronology_document
    from icharlotte_core.ui.wizard.in_process_task_tab import build_med_extractor_tab
    from icharlotte_core.ui.wizard.registry import get_task

    source = _build_chronology_docx(tmp_path / "chronology.docx")
    parsed = parse_chronology_document(str(source))
    selected_row_id = parsed.rows[0].id

    with patch(
        "icharlotte_core.ui.wizard.in_process_task_tab.QFileDialog.getOpenFileName",
        side_effect=AssertionError("restore path should not open the picker"),
    ):
        tab = build_med_extractor_tab(
            get_task("med_record_extractor"),
            case_path=str(tmp_path),
            file_number="5800.013",
            parent=None,
            chronology_path=str(source),
            initial_settings={
                "chronology_path": str(source),
                "selected_row_ids": [selected_row_id],
            },
        )

    qtbot.addWidget(tab)
    assert tab.settings_page.is_row_checked(selected_row_id)
    assert tab.settings_page.selected_count_label.text() == "1 row selected"


def test_reopen_recent_med_extractor_uses_saved_chronology_without_picker(
    qtbot,
    tmp_path,
    monkeypatch,
):
    import iCharlotte
    from PySide6.QtCore import Signal
    from PySide6.QtWidgets import QTabWidget
    from icharlotte_core.ui.wizard import in_process_task_tab

    win = iCharlotte.MainWindow.__new__(iCharlotte.MainWindow)
    win.tabs = QTabWidget()
    qtbot.addWidget(win.tabs)
    win.case_path = str(tmp_path)
    win.file_number = "5800.013"
    win._on_task_completed = lambda *a, **k: None
    win._hide_fixed_close_buttons = lambda *a, **k: None

    class FakeTab(QTabWidget):
        task_completed = Signal(dict)

    calls = {}
    sentinel = FakeTab()
    saved_settings = {
        "chronology_path": "RECORDS/chronology.docx",
        "selected_row_ids": ["row-a"],
    }

    def fake_builder(
        spec,
        case_path,
        file_number,
        parent,
        *,
        chronology_path="",
        initial_settings=None,
    ):
        calls["task_id"] = spec.task_id
        calls["chronology_path"] = chronology_path
        calls["initial_settings"] = initial_settings
        return sentinel

    monkeypatch.setattr(in_process_task_tab, "build_med_extractor_tab", fake_builder)

    win._on_reopen_recent_task({
        "task_id": "med_record_extractor",
        "files": ["RECORDS/chronology.docx"],
        "settings": saved_settings,
    })

    assert calls["task_id"] == "med_record_extractor"
    assert calls["chronology_path"] == str(tmp_path / "RECORDS" / "chronology.docx")
    assert calls["initial_settings"]["chronology_path"] == calls["chronology_path"]
    assert calls["initial_settings"]["selected_row_ids"] == ["row-a"]
    assert win.tabs.widget(0) is sentinel


def test_restore_open_med_extractor_uses_saved_chronology_without_picker(
    qtbot,
    tmp_path,
    monkeypatch,
):
    import iCharlotte
    from PySide6.QtCore import Signal
    from PySide6.QtWidgets import QTabWidget
    from icharlotte_core.ui.wizard import in_process_task_tab
    from icharlotte_core.ui.wizard.persistence import WizardStatePersistence

    win = iCharlotte.MainWindow.__new__(iCharlotte.MainWindow)
    win.tabs = QTabWidget()
    qtbot.addWidget(win.tabs)
    win.case_path = str(tmp_path)
    win.file_number = "5800.013"
    win._on_task_completed = lambda *a, **k: None
    win._hide_fixed_close_buttons = lambda *a, **k: None

    class FakeTab(QTabWidget):
        task_completed = Signal(dict)

    calls = {}
    sentinel = FakeTab()
    saved_settings = {
        "chronology_path": "RECORDS/chronology.docx",
        "selected_row_ids": ["row-a"],
    }

    def fake_builder(
        spec,
        case_path,
        file_number,
        parent,
        *,
        chronology_path="",
        initial_settings=None,
    ):
        calls["task_id"] = spec.task_id
        calls["chronology_path"] = chronology_path
        calls["initial_settings"] = initial_settings
        return sentinel

    monkeypatch.setattr(in_process_task_tab, "build_med_extractor_tab", fake_builder)
    persistence = WizardStatePersistence(str(tmp_path))
    persistence.set_open_tabs([{
        "task_id": "med_record_extractor",
        "instance_suffix": "",
        "files": ["RECORDS/chronology.docx"],
        "settings": saved_settings,
        "page": "settings",
    }])
    persistence.save()

    iCharlotte.MainWindow._restore_task_tabs_for_case(win)

    assert calls["task_id"] == "med_record_extractor"
    assert calls["chronology_path"] == str(tmp_path / "RECORDS" / "chronology.docx")
    assert calls["initial_settings"]["chronology_path"] == calls["chronology_path"]
    assert calls["initial_settings"]["selected_row_ids"] == ["row-a"]
    assert win.tabs.widget(0) is sentinel


def test_snapshot_med_extractor_output_page_without_output_path(qtbot, tmp_path):
    import iCharlotte
    from PySide6.QtCore import Signal
    from PySide6.QtWidgets import QTabWidget, QWidget
    from icharlotte_core.ui.wizard.in_process_task_tab import (
        InProcessTaskTab,
        MedExtractorOutputPage,
        PAGE_OUTPUT,
    )
    from icharlotte_core.ui.wizard.registry import get_task

    class Settings(QWidget):
        run_requested = Signal(dict)

        def to_dict(self):
            return {"chronology_path": str(tmp_path / "chronology.docx")}

    win = iCharlotte.MainWindow.__new__(iCharlotte.MainWindow)
    win.tabs = QTabWidget()
    qtbot.addWidget(win.tabs)
    win.case_path = str(tmp_path)

    tab = InProcessTaskTab(
        spec=get_task("med_record_extractor"),
        case_path=str(tmp_path),
        file_number="5800.013",
        settings_widget=Settings(),
        output_widget=MedExtractorOutputPage(str(tmp_path)),
        worker_factory=lambda cp, fn, settings, parent: None,
    )
    tab.setProperty("wizard_task_id", "med_record_extractor")
    tab.setCurrentIndex(PAGE_OUTPUT)
    qtbot.addWidget(tab)
    win.tabs.addTab(tab, "Med Record Extractor")

    snapshots = iCharlotte.MainWindow._snapshot_open_task_tabs(
        win,
        cancel_running=False,
    )

    assert snapshots[0]["task_id"] == "med_record_extractor"
    assert snapshots[0]["page"] == "output"
    assert snapshots[0]["output_path"] is None


def test_selection_page_blocks_extract_without_table(qtbot, tmp_path):
    from docx import Document
    from icharlotte_core.ui.wizard.pages.med_record_extractor_page import (
        MedChronologySelectionPage,
    )

    source = tmp_path / "bad.docx"
    doc = Document()
    doc.add_paragraph("BRIEF SYNOPSIS OF POST-INJURY MEDICAL RECORD:")
    doc.add_paragraph("On 09/21/2020, she saw Kaiser Permanente.")
    doc.save(source)

    page = MedChronologySelectionPage(str(tmp_path), "5800.013", str(source))
    qtbot.addWidget(page)

    assert "No usable 5-column chronology table found." in page.warning_label.text()
    assert not page.extract_btn.isEnabled()


def test_selection_page_reports_malformed_page_no(qtbot, tmp_path):
    from docx import Document
    from icharlotte_core.ui.wizard.pages.med_record_extractor_page import (
        MedChronologySelectionPage,
    )

    source = tmp_path / "bad_page.docx"
    doc = Document()
    doc.add_paragraph("BRIEF SYNOPSIS OF POST-INJURY MEDICAL RECORD:")
    doc.add_paragraph("On 09/21/2020, she saw Kaiser Permanente.")
    table = doc.add_table(rows=1, cols=5)
    for index, header in enumerate([
        "DATE",
        "PAGE NO",
        "PROVIDER",
        "DESCRIPTION",
        "Red Flags/Comments",
    ]):
        table.rows[0].cells[index].text = header
    row = table.add_row().cells
    row[0].text = "09/21/2020"
    row[1].text = "source\n\nPg no: 17-"
    row[2].text = "Kaiser Permanente"
    row[3].text = "Bad page reference"
    row[4].text = ""
    doc.save(source)

    page = MedChronologySelectionPage(str(tmp_path), "5800.013", str(source))
    qtbot.addWidget(page)
    page.set_row_checked(page.document.rows[0].id, True)

    assert page.selected_count_label.text() == "0 rows selected"
    assert not page.extract_btn.isEnabled()
    assert page.document.rows[0].warning.startswith("Could not parse record/pages")


def test_match_status_reflects_checked_failed_synopsis_paragraphs(qtbot):
    from icharlotte_core.ui.wizard.pages.med_record_extractor_page import (
        MedChronologySelectionPage,
    )

    with tempfile.TemporaryDirectory() as td:
        source = _build_chronology_docx(
            Path(td) / "chronology.docx",
            synopsis_extra_paragraphs=(
                "She returned home with crutches and follow-up instructions.",
            ),
        )
        page = MedChronologySelectionPage(
            case_path=td,
            file_number="5800.013",
            chronology_path=str(source),
        )
        qtbot.addWidget(page)

        failed = page.document.synopsis_paragraphs[1]
        confident = page.document.synopsis_paragraphs[0]
        page.set_paragraph_checked(failed.id, True)
        page.set_paragraph_checked(confident.id, True)

        assert "No date found" in page.match_status_label.text()

        page.set_paragraph_checked(failed.id, False)

        assert page.match_status_label.text() == ""
        assert page.match_status_label.isHidden()
