"""Reopening/restoring a drafted motion must restore the SAME color-coded
citations, not show "no citations detected". The .docx stores no verdicts, so
the citation-review output page persists them to a JSON sidecar beside the
preview and reloads it in load_output(). Shared base behavior → also covers the
oppose-motion output page.
"""
from __future__ import annotations

import os

import pytest

pytest.importorskip("PySide6")
pytest.importorskip("pytestqt")

from icharlotte_core.opposition.models import (  # noqa: E402
    CitationVerification,
    DraftDocument,
)
from icharlotte_core.ui.wizard.pages.generate_motion_page import (  # noqa: E402
    GenerateMotionOutputPage,
)


def _make_preview(tmp_path):
    from docx import Document

    preview = tmp_path / "Motion Preview.docx"
    doc = Document()
    doc.add_paragraph("See Smith v. Jones (2010) 50 Cal.4th 100 for support.")
    doc.save(str(preview))
    return preview


def test_show_result_persists_citation_sidecar(qtbot, tmp_path):
    preview = _make_preview(tmp_path)
    page = GenerateMotionOutputPage()
    qtbot.addWidget(page)
    page.show_result(DraftDocument(
        title="Motion to Compel",
        body_text="See *Smith v. Jones* (2010) 50 Cal.4th 100 for support.",
        preview_path=str(preview),
        citations=[CitationVerification(
            citation_text="Smith v. Jones (2010) 50 Cal.4th 100",
            case_name="Smith v. Jones", verdict="SUPPORTED", kind="case",
            evidence="The court held the duty applies.",
        )],
    ))
    assert os.path.isfile(str(preview) + ".citations.json")


def test_reopen_restores_verified_citations(qtbot, tmp_path):
    preview = _make_preview(tmp_path)

    # Generate (fresh) → persists the sidecar.
    page = GenerateMotionOutputPage()
    qtbot.addWidget(page)
    page.show_result(DraftDocument(
        title="Motion to Compel",
        body_text="See *Smith v. Jones* (2010) 50 Cal.4th 100 for support.",
        preview_path=str(preview),
        citations=[CitationVerification(
            citation_text="Smith v. Jones (2010) 50 Cal.4th 100",
            case_name="Smith v. Jones", verdict="SUPPORTED", kind="case",
            evidence="The court held the duty applies.",
        )],
    ))

    # Reopen in a fresh page via load_output → verdicts restored.
    page2 = GenerateMotionOutputPage()
    qtbot.addWidget(page2)
    page2.load_output(str(preview))
    assert len(page2.draft.citations) == 1
    assert page2.draft.citations[0].verdict == "SUPPORTED"
    assert page2.draft.citations[0].citation_text == "Smith v. Jones (2010) 50 Cal.4th 100"
    # Panel shows the citation, not the empty-state message.
    assert "SUPPORTED" in page2.detail_panel.header_label.text()


def test_reopen_without_sidecar_shows_empty(qtbot, tmp_path):
    # No sidecar (e.g. older saved motion) → graceful empty state, no crash.
    preview = _make_preview(tmp_path)
    page = GenerateMotionOutputPage()
    qtbot.addWidget(page)
    page.load_output(str(preview))
    assert page.draft.citations == []
