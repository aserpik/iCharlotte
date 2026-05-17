"""Tests for docx <-> HTML helpers used by the Output Page."""
import os
import pytest

from icharlotte_core.ui.wizard.docx_io import load_docx_as_html, save_qtextdocument_as_docx


def test_load_docx_as_html_returns_html(tmp_path):
    # Build a tiny .docx using python-docx for the test.
    from docx import Document
    p = tmp_path / "hello.docx"
    doc = Document()
    doc.add_heading("Title", level=1)
    doc.add_paragraph("Hello, ").add_run("world").bold = True
    doc.save(str(p))

    html = load_docx_as_html(str(p))
    assert "<h1" in html.lower() or "<h1>" in html.lower()
    assert "world" in html
    assert "<strong>" in html.lower() or "<b>" in html.lower()


def test_save_qtextdocument_as_docx_roundtrips_basic_text(tmp_path, qtbot):
    pytest.importorskip("pytestqt")
    from PySide6.QtGui import QTextDocument
    from docx import Document

    qdoc = QTextDocument()
    qdoc.setHtml("<h1>Header</h1><p>Hello <b>bold</b> world.</p>")
    out_path = str(tmp_path / "out.docx")
    save_qtextdocument_as_docx(qdoc, out_path)
    assert os.path.exists(out_path)
    # Re-read with python-docx and check basic content.
    d = Document(out_path)
    all_text = "\n".join(p.text for p in d.paragraphs)
    assert "Header" in all_text
    assert "bold" in all_text
    assert "Hello" in all_text and "world" in all_text
