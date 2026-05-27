"""Tests for the per-row context-documents UI feature in MedChronConfigForm."""

import json
import sys
from pathlib import Path

import pytest

pytest.importorskip("pytestqt")  # NOTE: no underscore — pytest_qt silently skips

PROJECT_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(PROJECT_ROOT))


@pytest.fixture(autouse=True)
def _isolated_custom_analyses_store(tmp_path, monkeypatch):
    """Redirect the global custom-analyses JSON to a per-test tmp path so
    tests don't see (or pollute) the developer's real saved analyses."""
    from icharlotte_core.med_chron import custom_analyses_store
    monkeypatch.setattr(
        custom_analyses_store,
        "_STORE_PATH",
        tmp_path / "store" / "med_chron_custom_analyses.json",
    )
    yield


# -----------------------------
# Task 3: sniff_text_layer tests
# -----------------------------

def test_sniff_text_layer_txt_has_text(tmp_path):
    from icharlotte_core.ui.med_chron_config_form import sniff_text_layer
    p = tmp_path / "ok.txt"
    p.write_text("This is a useful status report with content.", encoding="utf-8")
    has_text, _ = sniff_text_layer(str(p))
    assert has_text is True


def test_sniff_text_layer_txt_empty(tmp_path):
    from icharlotte_core.ui.med_chron_config_form import sniff_text_layer
    p = tmp_path / "empty.txt"
    p.write_text("   \n  ", encoding="utf-8")
    has_text, _ = sniff_text_layer(str(p))
    assert has_text is False


def test_sniff_text_layer_docx_has_text(tmp_path):
    from docx import Document
    from icharlotte_core.ui.med_chron_config_form import sniff_text_layer
    p = tmp_path / "doc.docx"
    doc = Document()
    doc.add_paragraph("Some legible paragraph content for the sniff to find.")
    doc.save(str(p))
    has_text, _ = sniff_text_layer(str(p))
    assert has_text is True


def test_sniff_text_layer_docx_empty(tmp_path):
    from docx import Document
    from icharlotte_core.ui.med_chron_config_form import sniff_text_layer
    p = tmp_path / "blank.docx"
    Document().save(str(p))
    has_text, _ = sniff_text_layer(str(p))
    assert has_text is False


def test_sniff_text_layer_unreadable_returns_false(tmp_path):
    from icharlotte_core.ui.med_chron_config_form import sniff_text_layer
    # File that does not exist.
    has_text, reason = sniff_text_layer(str(tmp_path / "ghost.pdf"))
    assert has_text is False
    assert reason  # non-empty reason string


def test_sniff_text_layer_docx_with_table_content_only(tmp_path):
    """python-docx's doc.paragraphs skips tables — sniff must also sample
    table cells so docs whose content lives in tables (legal chronological
    summaries, intake forms) aren't falsely flagged as 'no text layer'."""
    from docx import Document
    from icharlotte_core.ui.med_chron_config_form import sniff_text_layer
    p = tmp_path / "tables_only.docx"
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Date"
    table.cell(0, 1).text = "Provider"
    table.cell(1, 0).text = "2024-02-01"
    table.cell(1, 1).text = "Acme PT"
    doc.save(str(p))
    has_text, _ = sniff_text_layer(str(p))
    assert has_text is True


def test_sniff_text_layer_pdf_with_text_layer(tmp_path):
    """A PDF whose first page has > 200 chars of extractable text returns
    (True, '')."""
    from pypdf import PdfWriter
    from icharlotte_core.ui.med_chron_config_form import sniff_text_layer

    try:
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
    except ImportError:
        pytest.skip("reportlab not installed — needed to build a text-layer PDF")

    p = tmp_path / "with_text.pdf"
    c = canvas.Canvas(str(p), pagesize=letter)
    # > 200 chars of extractable text. Spread across a few lines.
    long_line = "The defense theory rests on plaintiff's pre-existing degenerative changes documented across multiple imaging studies."
    y = 750
    for _ in range(3):
        c.drawString(72, y, long_line)
        y -= 20
    c.save()

    has_text, _ = sniff_text_layer(str(p))
    assert has_text is True


def test_sniff_text_layer_pdf_below_threshold(tmp_path):
    """A PDF with very little text (under 200 chars) returns (False, reason)."""
    from icharlotte_core.ui.med_chron_config_form import sniff_text_layer

    try:
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
    except ImportError:
        pytest.skip("reportlab not installed — needed to build a text-layer PDF")

    p = tmp_path / "tiny_text.pdf"
    c = canvas.Canvas(str(p), pagesize=letter)
    c.drawString(72, 750, "Tiny.")  # only 5 chars
    c.save()

    has_text, reason = sniff_text_layer(str(p))
    assert has_text is False
    assert reason  # non-empty reason


def test_sniff_text_layer_unsupported_extension(tmp_path):
    """An .rtf or .png is reported as unsupported with a descriptive reason."""
    from icharlotte_core.ui.med_chron_config_form import sniff_text_layer
    p = tmp_path / "image.png"
    p.write_bytes(b"\x89PNG\r\n")
    has_text, reason = sniff_text_layer(str(p))
    assert has_text is False
    assert ".png" in reason


# -----------------------------
# Task 4: ContextDropTextEdit
# -----------------------------

def test_context_drop_textedit_accepts_supported_file_urls(qtbot, tmp_path):
    from PySide6.QtCore import QMimeData, QPoint, Qt, QUrl
    from PySide6.QtGui import QDragEnterEvent
    from icharlotte_core.ui.med_chron_config_form import ContextDropTextEdit

    p1 = tmp_path / "a.pdf"
    p1.write_bytes(b"%PDF-1.4 stub")
    p2 = tmp_path / "b.docx"
    p2.write_bytes(b"docx stub")

    w = ContextDropTextEdit()
    qtbot.addWidget(w)

    received: list[list[str]] = []
    w.files_dropped.connect(lambda paths: received.append(paths))

    mime = QMimeData()
    mime.setUrls([QUrl.fromLocalFile(str(p1)), QUrl.fromLocalFile(str(p2))])
    event = QDragEnterEvent(QPoint(10, 10), Qt.DropAction.CopyAction, mime,
                            Qt.MouseButton.LeftButton, Qt.KeyboardModifier.NoModifier)
    w.dragEnterEvent(event)
    assert event.isAccepted()


def test_context_drop_textedit_rejects_unsupported_file_urls(qtbot, tmp_path):
    from PySide6.QtCore import QMimeData, QPoint, Qt, QUrl
    from PySide6.QtGui import QDragEnterEvent
    from icharlotte_core.ui.med_chron_config_form import ContextDropTextEdit

    p = tmp_path / "bad.png"
    p.write_bytes(b"\x89PNG\r\n")

    w = ContextDropTextEdit()
    qtbot.addWidget(w)

    mime = QMimeData()
    mime.setUrls([QUrl.fromLocalFile(str(p))])
    event = QDragEnterEvent(QPoint(10, 10), Qt.DropAction.CopyAction, mime,
                            Qt.MouseButton.LeftButton, Qt.KeyboardModifier.NoModifier)
    w.dragEnterEvent(event)
    # Not accepted as a "file drop", so the base class falls through; since
    # the base QPlainTextEdit doesn't know how to handle a PNG URL either,
    # the event is not accepted.
    assert not event.isAccepted()


def test_context_drop_textedit_emits_files_dropped(qtbot, tmp_path):
    from PySide6.QtCore import QMimeData, QPoint, Qt, QUrl
    from PySide6.QtGui import QDropEvent
    from icharlotte_core.ui.med_chron_config_form import ContextDropTextEdit

    p = tmp_path / "good.txt"
    p.write_text("hi", encoding="utf-8")

    w = ContextDropTextEdit()
    qtbot.addWidget(w)

    received: list[list[str]] = []
    w.files_dropped.connect(lambda paths: received.append(paths))

    mime = QMimeData()
    mime.setUrls([QUrl.fromLocalFile(str(p))])
    event = QDropEvent(QPoint(10, 10), Qt.DropAction.CopyAction, mime,
                       Qt.MouseButton.LeftButton, Qt.KeyboardModifier.NoModifier)
    w.dropEvent(event)
    assert received == [[str(p)]]


def test_context_drop_textedit_plaintext_drop_still_works(qtbot):
    """Dropping plain text (not file URLs) should fall through to the base
    QPlainTextEdit so the user can still drop a snippet."""
    from PySide6.QtCore import QMimeData, QPoint, Qt
    from PySide6.QtGui import QDropEvent
    from icharlotte_core.ui.med_chron_config_form import ContextDropTextEdit

    w = ContextDropTextEdit()
    qtbot.addWidget(w)

    received: list[list[str]] = []
    w.files_dropped.connect(lambda paths: received.append(paths))

    mime = QMimeData()
    mime.setText("some snippet")
    event = QDropEvent(QPoint(10, 10), Qt.DropAction.CopyAction, mime,
                       Qt.MouseButton.LeftButton, Qt.KeyboardModifier.NoModifier)
    w.dropEvent(event)
    # files_dropped should NOT fire for plain text
    assert received == []


def test_context_drop_textedit_mixed_drop_rejects_all(qtbot, tmp_path):
    """Mixed drop (PDF + PNG) is ambiguous; the whole drop is rejected
    and no files_dropped signal fires."""
    from PySide6.QtCore import QMimeData, QPoint, Qt, QUrl
    from PySide6.QtGui import QDropEvent
    from icharlotte_core.ui.med_chron_config_form import ContextDropTextEdit

    good = tmp_path / "ok.pdf"
    good.write_bytes(b"%PDF-1.4")
    bad = tmp_path / "bad.png"
    bad.write_bytes(b"\x89PNG\r\n")

    w = ContextDropTextEdit()
    qtbot.addWidget(w)

    received: list[list[str]] = []
    w.files_dropped.connect(lambda paths: received.append(paths))

    mime = QMimeData()
    mime.setUrls([QUrl.fromLocalFile(str(good)), QUrl.fromLocalFile(str(bad))])
    event = QDropEvent(QPoint(10, 10), Qt.DropAction.CopyAction, mime,
                       Qt.MouseButton.LeftButton, Qt.KeyboardModifier.NoModifier)
    w.dropEvent(event)
    assert received == []


# -----------------------------
# Task 5: CustomAnalysisRow chip strip
# -----------------------------

def _make_row(qtbot):
    from icharlotte_core.ui.med_chron_config_form import CustomAnalysisRow
    row = CustomAnalysisRow(None, on_remove=lambda r: None)
    qtbot.addWidget(row)
    return row


def test_custom_row_starts_with_no_context_files(qtbot):
    row = _make_row(qtbot)
    assert row.context_files() == []


def test_custom_row_add_context_files_appends_and_dedupes(qtbot, tmp_path):
    p1 = tmp_path / "a.txt"
    p1.write_text("hi", encoding="utf-8")
    p2 = tmp_path / "b.txt"
    p2.write_text("hi", encoding="utf-8")

    row = _make_row(qtbot)
    row.add_context_files([str(p1)])
    row.add_context_files([str(p1), str(p2)])  # p1 is a dup
    assert row.context_files() == [str(p1), str(p2)]


def test_custom_row_remove_context_file_clears_chip(qtbot, tmp_path):
    p = tmp_path / "a.txt"
    p.write_text("hi", encoding="utf-8")

    row = _make_row(qtbot)
    row.add_context_files([str(p)])
    assert row.context_files() == [str(p)]
    row._remove_context_file(str(p))
    assert row.context_files() == []


def test_custom_row_is_empty_ignores_context_files(qtbot, tmp_path):
    """A row with only files attached (no label, no instruction) is still
    considered empty so the form doesn't try to persist it."""
    p = tmp_path / "a.txt"
    p.write_text("hi", encoding="utf-8")

    row = _make_row(qtbot)
    row.add_context_files([str(p)])
    assert row.is_empty() is True


def test_custom_row_chip_strip_renders_one_chip_per_file(qtbot, tmp_path):
    p1 = tmp_path / "a.txt"
    p1.write_text("hi", encoding="utf-8")
    p2 = tmp_path / "b.txt"
    p2.write_text("hi", encoding="utf-8")

    row = _make_row(qtbot)
    row.add_context_files([str(p1), str(p2)])

    # The strip lays chips out in row._chip_strip_layout — each chip has the
    # filename in its child QLabel. Iterate widgets and collect labels.
    from PySide6.QtWidgets import QLabel
    chip_texts = []
    for i in range(row._chip_strip_layout.count()):
        w = row._chip_strip_layout.itemAt(i).widget()
        if w is None:
            continue
        # Each chip's filename QLabel uses objectName "chip_filename"
        for child in w.findChildren(QLabel):
            if child.objectName() == "chip_filename":
                chip_texts.append(child.text())
    assert "a.txt" in chip_texts
    assert "b.txt" in chip_texts


# -----------------------------
# Task 6: + Add context button
# -----------------------------

def test_add_context_button_opens_filedialog_and_attaches(qtbot, tmp_path, monkeypatch):
    """Clicking '+ Add context' opens QFileDialog.getOpenFileNames and any
    selected paths are appended to context_files."""
    p = tmp_path / "picked.pdf"
    p.write_bytes(b"%PDF-1.4")

    # Stub QFileDialog.getOpenFileNames to return our test file.
    from icharlotte_core.ui import med_chron_config_form
    monkeypatch.setattr(
        med_chron_config_form.QFileDialog,
        "getOpenFileNames",
        staticmethod(lambda *a, **kw: ([str(p)], "")),
    )

    row = _make_row(qtbot)
    assert row.context_files() == []
    row._add_ctx_btn.click()
    assert row.context_files() == [str(p)]


def test_add_context_button_cancelled_does_nothing(qtbot, monkeypatch):
    from icharlotte_core.ui import med_chron_config_form
    monkeypatch.setattr(
        med_chron_config_form.QFileDialog,
        "getOpenFileNames",
        staticmethod(lambda *a, **kw: ([], "")),
    )

    row = _make_row(qtbot)
    row._add_ctx_btn.click()
    assert row.context_files() == []


# -----------------------------
# Task 7: warning label
# -----------------------------

def test_warning_label_visible_when_file_lacks_text_layer(qtbot, tmp_path):
    # An empty .txt fails sniff_text_layer.
    p = tmp_path / "blank.txt"
    p.write_text("", encoding="utf-8")

    row = _make_row(qtbot)
    row.add_context_files([str(p)])
    assert row._context_warning_label.isHidden() is False
    assert "blank.txt" in row._context_warning_label.text()


def test_warning_label_hidden_when_all_files_have_text(qtbot, tmp_path):
    p = tmp_path / "good.txt"
    p.write_text("real content", encoding="utf-8")

    row = _make_row(qtbot)
    row.add_context_files([str(p)])
    assert row._context_warning_label.isHidden() is True


def test_warning_label_hides_after_bad_file_removed(qtbot, tmp_path):
    bad = tmp_path / "blank.txt"
    bad.write_text("", encoding="utf-8")
    good = tmp_path / "good.txt"
    good.write_text("real content", encoding="utf-8")

    row = _make_row(qtbot)
    row.add_context_files([str(bad), str(good)])
    assert row._context_warning_label.isHidden() is False

    row._remove_context_file(str(bad))
    assert row._context_warning_label.isHidden() is True


# -----------------------------
# Task 8: dual-shape commit
# -----------------------------


def _write_session(tmp_path, *, narrative_missing=False):
    cache = tmp_path / ".med_chron" / "abc123"
    cache.mkdir(parents=True)
    session_path = cache / "session.json"
    session_path.write_text(json.dumps({
        "version": 1,
        "phase": "awaiting_input",
        "input_path": str(tmp_path / "rec.docx"),
        "narrative_text_path": str(cache / "narrative.txt"),
        "full_text_path": str(cache / "full.txt"),
        "narrative_missing": narrative_missing,
        "provider_name": "Acme PT",
        "file_number": "1234.567",
        "catalog": [
            {"id": "rewrite_chronology", "title": "Rewrite Chronology",
             "description": "...", "uses_tables": False,
             "default_selected": True},
        ],
        "user_config": None,
    }, indent=2), encoding="utf-8")
    return session_path


def test_commit_writes_context_files_to_session(qtbot, tmp_path):
    """When a custom analysis has attached context files, commit_user_config
    must include them in the session JSON's user_config.custom_analyses."""
    from icharlotte_core.ui.med_chron_config_form import MedChronConfigForm
    session_path = _write_session(tmp_path)
    ctx = tmp_path / "status.txt"
    ctx.write_text("ctx", encoding="utf-8")

    form = MedChronConfigForm(session_path)
    qtbot.addWidget(form)
    row = form.add_custom_row()
    row.label_edit.setText("Defense targets")
    row.instruction_edit.setPlainText("Identify providers.")
    row.add_context_files([str(ctx)])

    assert form.commit_user_config() is True

    written = json.loads(session_path.read_text(encoding="utf-8"))
    customs = written["user_config"]["custom_analyses"]
    assert len(customs) == 1
    assert customs[0]["context_files"] == [str(ctx)]
    assert customs[0]["label"] == "Defense targets"


def test_commit_does_not_persist_context_files_to_global_store(qtbot, tmp_path):
    """The global custom_analyses_store must continue to hold only
    {label, instruction} — never context_files."""
    from icharlotte_core.ui.med_chron_config_form import MedChronConfigForm
    from icharlotte_core.med_chron import custom_analyses_store
    session_path = _write_session(tmp_path)
    ctx = tmp_path / "status.txt"
    ctx.write_text("ctx", encoding="utf-8")

    form = MedChronConfigForm(session_path)
    qtbot.addWidget(form)
    row = form.add_custom_row()
    row.label_edit.setText("Defense targets")
    row.instruction_edit.setPlainText("Identify providers.")
    row.add_context_files([str(ctx)])

    form.commit_user_config()

    saved = custom_analyses_store.load()
    assert len(saved) == 1
    assert "context_files" not in saved[0]
    assert saved[0] == {"label": "Defense targets", "instruction": "Identify providers."}


def test_reopening_form_loads_saved_analysis_with_empty_context(qtbot, tmp_path):
    """After a commit, opening a NEW form against a NEW session shows the
    persisted label/instruction but starts with no attached context files."""
    from icharlotte_core.ui.med_chron_config_form import MedChronConfigForm
    session1 = _write_session(tmp_path)
    ctx = tmp_path / "status.txt"
    ctx.write_text("ctx", encoding="utf-8")

    form1 = MedChronConfigForm(session1)
    qtbot.addWidget(form1)
    row = form1.add_custom_row()
    row.label_edit.setText("Defense targets")
    row.instruction_edit.setPlainText("Identify providers.")
    row.add_context_files([str(ctx)])
    form1.commit_user_config()

    # Build a second session in a different tmp subdir.
    session2_dir = tmp_path / "session2"
    session2_dir.mkdir()
    session2 = _write_session(session2_dir)

    form2 = MedChronConfigForm(session2)
    qtbot.addWidget(form2)

    # Pre-populated row should exist with persisted text, but no context files.
    assert len(form2.custom_rows) == 1
    row2 = form2.custom_rows[0]
    assert row2.label() == "Defense targets"
    assert row2.instruction() == "Identify providers."
    assert row2.context_files() == []


def test_commit_omits_unchecked_rows_from_session(qtbot, tmp_path):
    """If the include checkbox is unchecked, the row is persisted globally
    but NOT included in session.json's run-shape list."""
    from icharlotte_core.ui.med_chron_config_form import MedChronConfigForm
    from icharlotte_core.med_chron import custom_analyses_store
    session_path = _write_session(tmp_path)
    ctx = tmp_path / "status.txt"
    ctx.write_text("ctx", encoding="utf-8")

    form = MedChronConfigForm(session_path)
    qtbot.addWidget(form)
    row = form.add_custom_row()
    row.label_edit.setText("Defense targets")
    row.instruction_edit.setPlainText("Identify providers.")
    row.add_context_files([str(ctx)])
    row.include_cb.setChecked(False)

    # Need at least ONE thing checked, or commit fails validation.
    # The default rewrite_chronology checkbox is already checked.
    form.commit_user_config()

    written = json.loads(session_path.read_text(encoding="utf-8"))
    assert written["user_config"]["custom_analyses"] == []

    # But persisted to global store.
    saved = custom_analyses_store.load()
    assert len(saved) == 1
