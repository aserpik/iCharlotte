from pathlib import Path

import pytest
from docx import Document

from Scripts.depo_prep_lib.render_docx import render_outline_docx


def _outline_payload():
    return {
        "deponent_name": "Jane Doe",
        "deponent_role": "Plaintiff",
        "topics": [
            {"topic_id": "t01", "title": "Pre-existing conditions",
             "strategic_note": "Establish baseline.",
             "questions": [
                 {"n": 1, "text": "Before 2024, did you have back pain?",
                  "purpose": "Establish baseline.",
                  "source_facts": ["RFA #7 denied prior pain", "2019 PT intake notes chronic LBP"],
                  "impeachment_hook": "Confront with RFA #7"},
                 {"n": 2, "text": "When did you first see a chiropractor?"},
             ]},
            {"topic_id": "t02", "title": "Treatment timeline",
             "strategic_note": "Highlight gaps.",
             "questions": [{"n": 1, "text": "What treatment did you receive?"}]},
        ],
        "coverage_gaps": ["No question addresses 2019 chiropractor visits."],
    }


def test_render_creates_docx_with_expected_structure(tmp_path):
    out = tmp_path / "outline.docx"
    render_outline_docx(outline=_outline_payload(), output_path=out)
    assert out.exists()

    doc = Document(str(out))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Jane Doe" in text
    assert "Pre-existing conditions" in text
    assert "Before 2024" in text
    assert "Coverage" in text  # coverage notes section


def test_render_skips_optional_fields_when_absent(tmp_path):
    payload = {
        "deponent_name": "X", "deponent_role": "Y",
        "topics": [{"topic_id": "t01", "title": "T", "strategic_note": "",
                    "questions": [{"n": 1, "text": "Q only"}]}],
        "coverage_gaps": [],
    }
    out = tmp_path / "outline.docx"
    render_outline_docx(outline=payload, output_path=out)
    doc = Document(str(out))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Q only" in text
    assert "Purpose:" not in text
    assert "Source facts:" not in text
    assert "Coverage" not in text  # no gaps -> no section


def test_render_no_empty_paragraphs_for_spacing(tmp_path):
    """Spacing must use space_after, not empty paragraphs (MEMORY.md rule)."""
    out = tmp_path / "outline.docx"
    render_outline_docx(outline=_outline_payload(), output_path=out)
    doc = Document(str(out))
    empty_paras = [p for p in doc.paragraphs if not p.text.strip()]
    # Allow a handful of truly structural empties (e.g., spacing line under title);
    # but the bulk must use space_after. Cap loosely at 2.
    assert len(empty_paras) <= 2, (
        f"Found {len(empty_paras)} empty paragraphs - use space_after instead.")
