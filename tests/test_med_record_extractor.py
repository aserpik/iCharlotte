"""Tests for med record extractor page parsing and chronology document parsing."""

import tempfile
import unittest
from pathlib import Path

from docx import Document

from icharlotte_core.med_record_extractor import _parse_page_no


class TestParsePageNo(unittest.TestCase):
    def test_lowercase_pg_no_single_page(self):
        filename, start, end = _parse_page_no("Exhibits\n\nPg no: 112/170")
        self.assertEqual(filename, "Exhibits")
        self.assertEqual(start, 112)
        self.assertEqual(end, 112)

    def test_lowercase_pg_no_page_range(self):
        filename, start, end = _parse_page_no("Exhibits\n\nPg no: 27-28/170")
        self.assertEqual(filename, "Exhibits")
        self.assertEqual(start, 27)
        self.assertEqual(end, 28)

    def test_lowercase_pg_no_multi_page_range(self):
        filename, start, end = _parse_page_no("Exhibits\n\nPg no: 161-163/170")
        self.assertEqual(filename, "Exhibits")
        self.assertEqual(start, 161)
        self.assertEqual(end, 163)

    def test_capital_pg_no_still_works(self):
        filename, start, end = _parse_page_no(
            "SALTARELLI000001-SALTARELLI000772\n\nPg. No: 501-505/772"
        )
        self.assertEqual(filename, "SALTARELLI000001-SALTARELLI000772")
        self.assertEqual(start, 501)
        self.assertEqual(end, 505)

    def test_pg_colon_form_still_works(self):
        filename, start, end = _parse_page_no(
            "60337-0014_ ORTHOPEDICS & SPORTS MEDICINE.pdf\n\nPg: 23-24/69"
        )
        self.assertEqual(filename, "60337-0014_ ORTHOPEDICS & SPORTS MEDICINE.pdf")
        self.assertEqual(start, 23)
        self.assertEqual(end, 24)

    def test_bare_number_slash_total_still_works(self):
        filename, start, end = _parse_page_no("homedepot\n\n93/535")
        self.assertEqual(filename, "homedepot")
        self.assertEqual(start, 93)
        self.assertEqual(end, 93)

    def test_malformed_trailing_dash_fails_gracefully(self):
        # "Pg no: 17-" has no end page and no /total -- unparseable.
        # Must not crash; start==0 signals caller to surface an error.
        filename, start, end = _parse_page_no(
            "Comprehensive Spine and Pain medical and billing\n\nPg no: 17-"
        )
        self.assertEqual(filename, "Comprehensive Spine and Pain medical and billing")
        self.assertEqual(start, 0)
        self.assertEqual(end, 0)


class TestNormalizeDate(unittest.TestCase):
    def test_two_digit_year_uses_1950s_1960s_as_past_dates(self):
        from icharlotte_core.med_record_extractor import _normalize_date

        self.assertEqual(_normalize_date("09/21/68"), "09/21/1968")
        self.assertEqual(_normalize_date("September 21, 68"), "09/21/1968")


class TestBuildFileIndex(unittest.TestCase):
    def test_indexes_demand_folder_pdfs(self):
        from icharlotte_core.med_record_extractor import _build_file_index, _lookup_file

        with tempfile.TemporaryDirectory() as td:
            demand = Path(td) / "DEMAND"
            demand.mkdir()
            pdf = demand / "Richard Goulart demand exhibits.pdf"
            pdf.write_bytes(b"%PDF-1.4\n")

            index = _build_file_index(td)

            self.assertEqual(
                _lookup_file(index, "Richard Goulart demand exhibits.pdf"),
                str(pdf),
            )


def _build_chronology_docx(
    path: Path,
    *,
    include_synopsis: bool = True,
    flags_header: str = "Red Flags/Comments",
    synopsis_extra_paragraphs: tuple[str, ...] = (),
    first_page_no: str = (
        "Hall - Doc Produced HALL 000001 to 002530 7-21-2023\n\n"
        "Pg No: 599-604/2530"
    ),
    include_decoy_table: bool = False,
    after_table_paragraphs: tuple[str, ...] = (),
) -> Path:
    doc = Document()
    doc.add_paragraph("CHRONOLOGICAL MEDICAL SUMMARY")
    doc.add_paragraph("Client Name: Test Plaintiff")
    if include_synopsis:
        doc.add_paragraph("BRIEF SYNOPSIS OF POST-INJURY MEDICAL RECORD:")
        doc.add_paragraph(
            "On 09/21/2020, Test Plaintiff presented to Kaiser Permanente. "
            "She was evaluated by Henry Louis Marr, DO, for a right ankle injury."
        )
        for paragraph in synopsis_extra_paragraphs:
            doc.add_paragraph(paragraph)
        doc.add_paragraph(
            "On 09/21/2020, she had an X-ray performed by James Michael Erskine, MD."
        )
    if include_decoy_table:
        table = doc.add_table(rows=1, cols=5)
        headers = [
            "Updated Date",
            "Pageno Notes",
            "Provider",
            "Description",
            "Red Flags/Comments",
        ]
        for index, header in enumerate(headers):
            table.rows[0].cells[index].text = header
        row = table.add_row().cells
        row[0].text = "01/01/1999"
        row[1].text = "Decoy Record\n\nPg No: 1/1"
        row[2].text = "Wrong Provider"
        row[3].text = "Not a chronology table."
        row[4].text = ""
    table = doc.add_table(rows=1, cols=5)
    headers = ["DATE", "PAGE NO", "PROVIDER", "DESCRIPTION", flags_header]
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    section = table.add_row().cells
    for cell in section:
        cell.text = "POST-INJURY MEDICAL RECORDS"
    row = table.add_row().cells
    row[0].text = "09/21/2020"
    row[1].text = first_page_no
    row[2].text = "Kaiser Permanente\nHenry Louis Marr, DO"
    row[3].text = "EMERGENCY DEPARTMENT NOTE\nChief Complaint: Ankle injury."
    row[4].text = ""
    row = table.add_row().cells
    row[0].text = "09/21/2020"
    row[1].text = "Hall - Doc Produced HALL 000001 to 002530 7-21-2023\n\nPg No: 598/2530"
    row[2].text = "Kaiser Permanente\nJames Michael Erskine, MD"
    row[3].text = "RADIOLOGY REPORT\nX-ray right ankle."
    row[4].text = ""
    for paragraph in after_table_paragraphs:
        doc.add_paragraph(paragraph)
    doc.save(path)
    return path


class TestChronologyDocumentParser(unittest.TestCase):
    def _build_verbatim_synopsis_docx(self, path: Path) -> Path:
        doc = Document()
        doc.add_paragraph("BRIEF SYNOPSIS OF POST-INJURY MEDICAL RECORD:")
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(
            "On 09/21/2020, Test Plaintiff presented to Kaiser Permanente."
        )
        run.add_break()
        paragraph.add_run("She was evaluated by Henry Louis Marr, DO.")
        table = doc.add_table(rows=1, cols=5)
        headers = ["DATE", "PAGE NO", "PROVIDER", "DESCRIPTION", "Red Flags/Comments"]
        for index, header in enumerate(headers):
            table.rows[0].cells[index].text = header
        row = table.add_row().cells
        row[0].text = "09/21/2020"
        row[1].text = "Record\n\nPg No: 1/2"
        row[2].text = "Kaiser Permanente\nHenry Louis Marr, DO"
        row[3].text = "Emergency note"
        row[4].text = ""
        doc.save(path)
        return path

    def test_parse_chronology_document_extracts_synopsis_and_rows(self):
        from icharlotte_core.med_record_chronology import parse_chronology_document

        with tempfile.TemporaryDirectory() as td:
            source = _build_chronology_docx(Path(td) / "chronology.docx")
            parsed = parse_chronology_document(str(source))

        self.assertEqual(len(parsed.synopsis_paragraphs), 2)
        self.assertIn("Henry Louis Marr", parsed.synopsis_paragraphs[0].text)
        self.assertEqual(len(parsed.rows), 2)
        self.assertEqual(parsed.rows[0].date, "09/21/2020")
        self.assertEqual(parsed.rows[0].page_start, 599)
        self.assertEqual(parsed.rows[0].page_end, 604)
        self.assertEqual(parsed.rows[0].record_filename, "Hall - Doc Produced HALL 000001 to 002530 7-21-2023")
        self.assertFalse(parsed.blocking_errors)

    def test_parse_chronology_document_collects_synopsis_continuation_paragraph(self):
        from icharlotte_core.med_record_chronology import parse_chronology_document

        with tempfile.TemporaryDirectory() as td:
            source = _build_chronology_docx(
                Path(td) / "chronology.docx",
                synopsis_extra_paragraphs=(
                    "She returned home with crutches and follow-up instructions.",
                ),
            )
            parsed = parse_chronology_document(str(source))

        self.assertEqual(len(parsed.synopsis_paragraphs), 3)
        self.assertEqual(
            parsed.synopsis_paragraphs[1].text,
            "She returned home with crutches and follow-up instructions.",
        )

    def test_parse_chronology_document_stops_synopsis_at_table_boundary(self):
        from icharlotte_core.med_record_chronology import parse_chronology_document

        with tempfile.TemporaryDirectory() as td:
            source = _build_chronology_docx(
                Path(td) / "chronology.docx",
                after_table_paragraphs=(
                    "On 10/01/2020, this footer-style note should not be selectable.",
                ),
            )
            parsed = parse_chronology_document(str(source))

        self.assertEqual(len(parsed.synopsis_paragraphs), 2)
        self.assertNotIn(
            "10/01/2020",
            "\n".join(paragraph.text for paragraph in parsed.synopsis_paragraphs),
        )

    def test_parse_chronology_document_warns_when_page_no_is_malformed(self):
        from icharlotte_core.med_record_chronology import parse_chronology_document

        with tempfile.TemporaryDirectory() as td:
            source = _build_chronology_docx(
                Path(td) / "chronology.docx",
                first_page_no=(
                    "Hall - Doc Produced HALL 000001 to 002530 7-21-2023\n\nPg No: 17-"
                ),
            )
            parsed = parse_chronology_document(str(source))

        self.assertEqual(len(parsed.rows), 2)
        self.assertFalse(parsed.rows[0].extractable)
        self.assertIn("Could not parse record/pages from PAGE NO:", parsed.rows[0].warning)

    def test_parse_chronology_document_skips_decoy_table_before_real_chronology(self):
        from icharlotte_core.med_record_chronology import parse_chronology_document

        with tempfile.TemporaryDirectory() as td:
            source = _build_chronology_docx(
                Path(td) / "chronology.docx",
                include_decoy_table=True,
            )
            parsed = parse_chronology_document(str(source))

        self.assertEqual(len(parsed.rows), 2)
        self.assertEqual(parsed.rows[0].date, "09/21/2020")
        self.assertEqual(
            parsed.rows[0].record_filename,
            "Hall - Doc Produced HALL 000001 to 002530 7-21-2023",
        )
        self.assertFalse(parsed.blocking_errors)

    def test_parse_chronology_document_warns_when_page_range_is_invalid(self):
        from icharlotte_core.med_record_chronology import parse_chronology_document

        with tempfile.TemporaryDirectory() as td:
            source = _build_chronology_docx(
                Path(td) / "chronology.docx",
                first_page_no=(
                    "Hall - Doc Produced HALL 000001 to 002530 7-21-2023\n\n"
                    "Pg No: 10-5/20"
                ),
            )
            parsed = parse_chronology_document(str(source))

        self.assertEqual(parsed.rows[0].page_start, 10)
        self.assertEqual(parsed.rows[0].page_end, 5)
        self.assertFalse(parsed.rows[0].extractable)
        self.assertIn("Could not parse record/pages from PAGE NO:", parsed.rows[0].warning)

    def test_parse_chronology_document_warns_when_synopsis_missing(self):
        from icharlotte_core.med_record_chronology import parse_chronology_document

        with tempfile.TemporaryDirectory() as td:
            source = _build_chronology_docx(Path(td) / "chronology.docx", include_synopsis=False)
            parsed = parse_chronology_document(str(source))

        self.assertEqual(parsed.synopsis_paragraphs, [])
        self.assertEqual(len(parsed.rows), 2)
        self.assertIn("No Brief Synopsis section found.", parsed.warnings)

    def test_parse_chronology_document_blocks_when_table_missing(self):
        from icharlotte_core.med_record_chronology import parse_chronology_document

        with tempfile.TemporaryDirectory() as td:
            source = Path(td) / "no_table.docx"
            doc = Document()
            doc.add_paragraph("BRIEF SYNOPSIS OF POST-INJURY MEDICAL RECORD:")
            doc.add_paragraph("On 09/21/2020, Test Plaintiff saw Kaiser Permanente.")
            doc.save(source)
            parsed = parse_chronology_document(str(source))

        self.assertEqual(parsed.rows, [])
        self.assertIn("No usable 5-column chronology table found.", parsed.blocking_errors)

    def test_parse_chronology_document_preserves_verbatim_synopsis_text(self):
        from icharlotte_core.med_record_chronology import parse_chronology_document

        with tempfile.TemporaryDirectory() as td:
            source = self._build_verbatim_synopsis_docx(Path(td) / "chronology.docx")
            parsed = parse_chronology_document(str(source))

        self.assertEqual(
            parsed.synopsis_paragraphs[0].text,
            "On 09/21/2020, Test Plaintiff presented to Kaiser Permanente.\n"
            "She was evaluated by Henry Louis Marr, DO.",
        )

    def test_parse_chronology_document_preserves_verbatim_row_cells(self):
        from icharlotte_core.med_record_chronology import parse_chronology_document

        with tempfile.TemporaryDirectory() as td:
            source = _build_chronology_docx(Path(td) / "chronology.docx")
            parsed = parse_chronology_document(str(source))

        self.assertEqual(
            parsed.rows[0].page_no,
            "Hall - Doc Produced HALL 000001 to 002530 7-21-2023\n\n"
            "Pg No: 599-604/2530",
        )
        self.assertEqual(
            parsed.rows[0].provider,
            "Kaiser Permanente\nHenry Louis Marr, DO",
        )
        self.assertEqual(
            parsed.rows[0].description,
            "EMERGENCY DEPARTMENT NOTE\nChief Complaint: Ankle injury.",
        )

    def test_synopsis_matching_still_handles_verbatim_line_breaks(self):
        from icharlotte_core.med_record_chronology import (
            match_synopsis_to_rows,
            parse_chronology_document,
        )

        with tempfile.TemporaryDirectory() as td:
            source = self._build_verbatim_synopsis_docx(Path(td) / "chronology.docx")
            parsed = parse_chronology_document(str(source))

        result = match_synopsis_to_rows(parsed.synopsis_paragraphs[0], parsed.rows)

        self.assertEqual(result.status, "confident")
        self.assertEqual(result.row_ids, (parsed.rows[0].id,))

    def test_parse_chronology_document_blocks_when_flags_header_missing(self):
        from icharlotte_core.med_record_chronology import parse_chronology_document

        with tempfile.TemporaryDirectory() as td:
            source = _build_chronology_docx(
                Path(td) / "chronology.docx",
                flags_header="Notes",
            )
            parsed = parse_chronology_document(str(source))

        self.assertEqual(parsed.rows, [])
        self.assertIn("No usable 5-column chronology table found.", parsed.blocking_errors)


class TestSynopsisMatching(unittest.TestCase):
    def _parsed(self):
        from icharlotte_core.med_record_chronology import parse_chronology_document

        td = tempfile.TemporaryDirectory()
        self.addCleanup(td.cleanup)
        source = _build_chronology_docx(Path(td.name) / "chronology.docx")
        return parse_chronology_document(str(source))

    def test_synopsis_match_confident_by_date_and_provider(self):
        from icharlotte_core.med_record_chronology import match_synopsis_to_rows

        parsed = self._parsed()
        result = match_synopsis_to_rows(parsed.synopsis_paragraphs[0], parsed.rows)

        self.assertEqual(result.status, "confident")
        self.assertEqual(result.row_ids, (parsed.rows[0].id,))
        self.assertEqual(result.candidate_row_ids, ())

    def test_synopsis_match_ambiguous_when_provider_is_not_distinct(self):
        from icharlotte_core.med_record_chronology import SynopsisParagraph, match_synopsis_to_rows

        parsed = self._parsed()
        paragraph = SynopsisParagraph(
            id="syn-x",
            order=0,
            text="On 09/21/2020, she presented to Kaiser Permanente for ankle care.",
        )
        result = match_synopsis_to_rows(paragraph, parsed.rows)

        self.assertEqual(result.status, "ambiguous")
        self.assertEqual(set(result.candidate_row_ids), {row.id for row in parsed.rows})

    def test_synopsis_match_none_without_same_date(self):
        from icharlotte_core.med_record_chronology import SynopsisParagraph, match_synopsis_to_rows

        parsed = self._parsed()
        paragraph = SynopsisParagraph(
            id="syn-x",
            order=0,
            text="On 12/31/2021, she saw Kaiser Permanente for unrelated care.",
        )
        result = match_synopsis_to_rows(paragraph, parsed.rows)

        self.assertEqual(result.status, "none")
        self.assertEqual(result.row_ids, ())

    def test_synopsis_match_accepts_two_digit_year(self):
        from icharlotte_core.med_record_chronology import SynopsisParagraph, match_synopsis_to_rows

        parsed = self._parsed()
        paragraph = SynopsisParagraph(
            id="syn-x",
            order=0,
            text="On 09/21/20, she was evaluated by Henry Louis Marr, DO.",
        )
        result = match_synopsis_to_rows(paragraph, parsed.rows)

        self.assertEqual(result.status, "confident")
        self.assertEqual(result.row_ids, (parsed.rows[0].id,))

    def test_synopsis_match_accepts_month_name_date(self):
        from icharlotte_core.med_record_chronology import SynopsisParagraph, match_synopsis_to_rows

        parsed = self._parsed()
        paragraph = SynopsisParagraph(
            id="syn-x",
            order=0,
            text="On September 21, 2020, she was evaluated by Henry Louis Marr, DO.",
        )
        result = match_synopsis_to_rows(paragraph, parsed.rows)

        self.assertEqual(result.status, "confident")
        self.assertEqual(result.row_ids, (parsed.rows[0].id,))

    def test_synopsis_match_accepts_doctor_before_verb(self):
        from icharlotte_core.med_record_chronology import SynopsisParagraph, match_synopsis_to_rows

        parsed = self._parsed()
        paragraph = SynopsisParagraph(
            id="syn-x",
            order=0,
            text="On 09/21/2020, Dr. Marr evaluated her for ankle pain.",
        )
        result = match_synopsis_to_rows(paragraph, parsed.rows)

        self.assertEqual(result.status, "confident")
        self.assertEqual(result.row_ids, (parsed.rows[0].id,))

    def test_synopsis_match_does_not_confidently_select_generic_provider_token(self):
        from dataclasses import replace

        from icharlotte_core.med_record_chronology import SynopsisParagraph, match_synopsis_to_rows

        parsed = self._parsed()
        rows = [
            replace(parsed.rows[0], provider="Kaiser Clinic\nHenry Louis Marr, DO"),
            parsed.rows[1],
        ]
        paragraph = SynopsisParagraph(
            id="syn-x",
            order=0,
            text="On 09/21/2020, she presented to clinic for follow-up care.",
        )
        result = match_synopsis_to_rows(paragraph, rows)

        self.assertEqual(result.status, "ambiguous")
        self.assertEqual(set(result.candidate_row_ids), {row.id for row in rows})


class TestSelectionState(unittest.TestCase):
    def test_selection_state_removes_only_paragraph_owned_rows(self):
        from icharlotte_core.med_record_chronology import SelectionState

        state = SelectionState()
        state.select_row("row-manual", source="manual")
        state.select_row("row-shared", source="syn-a")
        state.select_row("row-shared", source="syn-b")
        state.select_row("row-only-a", source="syn-a")

        state.clear_source("syn-a")

        self.assertTrue(state.is_row_selected("row-manual"))
        self.assertTrue(state.is_row_selected("row-shared"))
        self.assertFalse(state.is_row_selected("row-only-a"))
        self.assertEqual(state.selected_row_ids(), ["row-manual", "row-shared"])

    def test_selection_state_manual_deselect_preserves_synopsis_owned_row(self):
        from icharlotte_core.med_record_chronology import SelectionState

        state = SelectionState()
        state.select_row("row-shared", source="syn-a")
        state.select_row("row-shared", source="manual")

        state.deselect_row("row-shared", source="manual")

        self.assertTrue(state.is_row_selected("row-shared"))
        state.clear_source("syn-a")
        self.assertFalse(state.is_row_selected("row-shared"))


class TestSelectedRowExtractionWorker(unittest.TestCase):
    def test_worker_uses_selected_rows_without_llm_or_auto_chron_lookup(self):
        from unittest.mock import patch

        from icharlotte_core.med_record_chronology import SelectableChronologyRow
        from icharlotte_core.med_record_extractor import MedRecordExtractorWorker

        row = SelectableChronologyRow(
            id="row-1",
            order=0,
            date="09/21/2020",
            page_no="source\n\nPg No: 2/3",
            provider="Kaiser Permanente",
            description="Emergency department note",
            flags="",
            record_filename="source",
            page_start=2,
            page_end=2,
        )
        worker = MedRecordExtractorWorker(
            case_path="C:/case",
            file_number="5800.013",
            chronology_path="C:/case/RECORDS/summary.docx",
            selected_rows=[row],
        )

        with patch.object(worker, "_parse_entries", side_effect=AssertionError("LLM path should not run")), \
             patch("icharlotte_core.med_record_extractor._find_most_recent_med_chron", side_effect=AssertionError("Auto chronology lookup should not run")), \
             patch("icharlotte_core.med_record_extractor._parse_med_chron_table", side_effect=AssertionError("Chronology table parsing should not run")), \
             patch("icharlotte_core.med_record_extractor._build_file_index", return_value={"source": "C:/case/RECORDS/source.pdf"}), \
             patch("icharlotte_core.med_record_extractor._extract_pages") as extract_pages, \
             patch("icharlotte_core.med_record_extractor._ocr_pdf_if_needed", return_value=False), \
             patch("icharlotte_core.med_record_extractor._update_index_document"):
            results = worker._execute()

        self.assertEqual(len(results), 1)
        self.assertEqual(
            results[0].output_path.replace("\\", "/"),
            "C:/case/NOTES/AI OUTPUT/Med Record Extracts/09-21-2020 - Kaiser Permanente - p2.pdf",
        )
        extract_pages.assert_called_once()

    def test_selected_row_output_names_are_windows_safe_and_unique(self):
        from pathlib import PureWindowsPath
        from unittest.mock import patch

        from icharlotte_core.med_record_chronology import SelectableChronologyRow
        from icharlotte_core.med_record_extractor import MedRecordExtractorWorker

        rows = [
            SelectableChronologyRow(
                id="row-1",
                order=0,
                date="09/21/2020",
                page_no="source one\n\nPg No: 2/3",
                provider='Kaiser: Permanente? "Clinic"',
                description="Emergency department note",
                flags="",
                record_filename="source one",
                page_start=2,
                page_end=2,
            ),
            SelectableChronologyRow(
                id="row-2",
                order=1,
                date="09/21/2020",
                page_no="source two\n\nPg No: 2/3",
                provider='Kaiser: Permanente? "Clinic"',
                description="Follow-up note",
                flags="",
                record_filename="source two",
                page_start=2,
                page_end=2,
            ),
        ]
        worker = MedRecordExtractorWorker(
            case_path="C:/case",
            file_number="5800.013",
            chronology_path="C:/case/RECORDS/summary.docx",
            selected_rows=rows,
        )

        file_index = {
            "source one": "C:/case/RECORDS/source one.pdf",
            "source two": "C:/case/RECORDS/source two.pdf",
        }
        with patch("icharlotte_core.med_record_extractor._build_file_index", return_value=file_index), \
             patch("icharlotte_core.med_record_extractor._extract_pages") as extract_pages, \
             patch("icharlotte_core.med_record_extractor._ocr_pdf_if_needed", return_value=False), \
             patch("icharlotte_core.med_record_extractor._update_index_document"):
            worker._execute()

        output_names = [
            PureWindowsPath(call.args[3]).name
            for call in extract_pages.call_args_list
        ]
        self.assertEqual(len(output_names), 2)
        self.assertEqual(len(set(output_names)), 2)
        invalid_chars = set('<>:"/\\|?*')
        for name in output_names:
            self.assertFalse(invalid_chars.intersection(name))

    def test_update_index_document_validates_after_save(self):
        from unittest.mock import patch

        from icharlotte_core.med_record_chronology import SelectableChronologyRow
        from icharlotte_core.med_record_extractor import _update_index_document

        row = SelectableChronologyRow(
            id="row-1",
            order=0,
            date="09/21/2020",
            page_no="source\n\nPg No: 2/3",
            provider="Kaiser Permanente",
            description="Emergency department note",
            flags="",
            record_filename="source",
            page_start=2,
            page_end=2,
        )
        with tempfile.TemporaryDirectory() as td, \
             patch("icharlotte_core.med_record_extractor.validate_index_docx") as validate:
            validate.return_value.has_errors = False
            _update_index_document(td, [row])

        validate.assert_called_once()

    def test_empty_selected_rows_still_bypasses_pasted_text_path(self):
        from unittest.mock import patch

        from icharlotte_core.med_record_extractor import MedRecordExtractorWorker

        worker = MedRecordExtractorWorker(
            case_path="C:/case",
            file_number="5800.013",
            user_text="09/21/2020 Kaiser Permanente",
            selected_rows=[],
        )

        with patch.object(worker, "_parse_entries", side_effect=AssertionError("LLM path should not run")), \
             patch("icharlotte_core.med_record_extractor._build_file_index", return_value={}):
            results = worker._execute()

        self.assertEqual(results, [])

    def test_run_formats_row_only_failures(self):
        from unittest.mock import patch

        from icharlotte_core.med_record_chronology import SelectableChronologyRow
        from icharlotte_core.med_record_extractor import ExtractionResult, MedRecordExtractorWorker

        row = SelectableChronologyRow(
            id="row-1",
            order=0,
            date="09/21/2020",
            page_no="source\n\nPg No: 2/3",
            provider="Kaiser Permanente",
            description="Emergency department note",
            flags="",
            record_filename="source",
            page_start=2,
            page_end=2,
        )
        worker = MedRecordExtractorWorker(
            case_path="C:/case",
            file_number="5800.013",
            chronology_path="C:/case/RECORDS/summary.docx",
            selected_rows=[row],
        )
        emitted = []
        worker.finished_result.connect(lambda success, message: emitted.append((success, message)))

        with patch.object(
            worker,
            "_execute",
            return_value=[ExtractionResult(matched_row=row, error="PDF not found: source")],
        ):
            worker.run()

        self.assertEqual(emitted, [(True, "1 failed\n  - 09/21/2020 Kaiser Permanente: PDF not found: source")])


if __name__ == "__main__":
    unittest.main()
