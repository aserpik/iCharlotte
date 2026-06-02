"""Tests for the Generate Motion Word assembler (spec #7)."""
import os

from docx import Document

from icharlotte_core.motion_generation.assembler import assemble_motion_preview
from icharlotte_core.motion_generation.config import get_motion_config
from icharlotte_core.opposition.models import DraftDocument


def _text(path):
    return "\n".join(p.text for p in Document(path).paragraphs)


def test_writes_docx_with_title_body_and_placeholders(tmp_path):
    cfg = get_motion_config("compel")
    draft = DraftDocument(
        title="Motion to Compel Further Responses",
        body_text="Introduction\n\nThe motion should be granted.",
    )
    out = tmp_path / "motion.docx"
    path = assemble_motion_preview(draft=draft, output_path=str(out), config=cfg)
    assert os.path.isfile(path)
    text = _text(path)
    assert "Motion to Compel Further Responses" in text
    assert "The motion should be granted." in text
    # Per-type placeholder attachments appear, clearly marked.
    assert "Meet and Confer" in text
    assert "Separate Statement" in text
    assert "to be completed" in text.lower()


def test_generic_motion_has_no_placeholder_sections(tmp_path):
    cfg = get_motion_config("generic")
    draft = DraftDocument(title="Motion", body_text="Body text.")
    out = tmp_path / "generic.docx"
    assemble_motion_preview(draft=draft, output_path=str(out), config=cfg)
    text = _text(str(out))
    assert "to be completed" not in text.lower()


def test_includes_notice_of_motion_heading(tmp_path):
    cfg = get_motion_config("demurrer")
    draft = DraftDocument(title="Demurrer", body_text="Argument.")
    out = tmp_path / "dem.docx"
    assemble_motion_preview(draft=draft, output_path=str(out), config=cfg)
    text = _text(str(out)).lower()
    assert "notice of motion" in text
