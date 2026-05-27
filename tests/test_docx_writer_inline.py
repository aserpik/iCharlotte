"""Tests for ``render_inline_markdown`` and the integration with each script.

These guard against three pre-existing bugs that were fixed when the inline
parser was centralised:

  1. ``re.split(r'(\\**.*?\\**)', ...)`` in med_record/med_chron/exposure used
     ``\\**`` (zero-or-more) instead of ``\\*\\*`` (two literal asterisks).
     The regex matched everywhere and ``**bold**`` markers leaked through.
  2. ``re.split(r'(\\*\\*.*\\*\\*)', ...)`` in liability was greedy, so
     ``**A** and **B**`` collapsed into a single bold span containing the
     middle ``** and **`` markers.
  3. No script handled ``*italic*`` or `` `code` ``, so those rendered as
     literal asterisks/backticks in the Output page.
"""
import unittest

from docx import Document

from icharlotte_core.docx_writer import render_inline_markdown


def _runs(para):
    """Return list of ``(text, bold, italic, font_name)`` for the runs in ``para``."""
    return [
        (r.text, r.bold, r.italic, r.font.name) for r in para.runs
    ]


class TestRenderInlineMarkdown(unittest.TestCase):
    def setUp(self):
        self.doc = Document()
        self.para = self.doc.add_paragraph()

    def test_bold_double_star(self):
        render_inline_markdown(self.para, "Patient **John Doe** today")
        self.assertEqual(_runs(self.para), [
            ("Patient ", None, None, "Times New Roman"),
            ("John Doe", True, None, "Times New Roman"),
            (" today", None, None, "Times New Roman"),
        ])

    def test_multiple_bold_spans_not_swallowed(self):
        """Liability used ``\\*\\*.*\\*\\*`` (greedy) and ate everything between
        the first and last ``**`` on the line. With the lazy helper, each pair
        is its own bold span."""
        render_inline_markdown(self.para, "**A** and **B** today")
        self.assertEqual(_runs(self.para), [
            ("A", True, None, "Times New Roman"),
            (" and ", None, None, "Times New Roman"),
            ("B", True, None, "Times New Roman"),
            (" today", None, None, "Times New Roman"),
        ])

    def test_italic_single_star(self):
        render_inline_markdown(self.para, "See *Roe v. Wade* for context.")
        self.assertEqual(_runs(self.para), [
            ("See ", None, None, "Times New Roman"),
            ("Roe v. Wade", None, True, "Times New Roman"),
            (" for context.", None, None, "Times New Roman"),
        ])

    def test_inline_code(self):
        render_inline_markdown(self.para, "Use `case_id` here.")
        # The code run uses Consolas; the others use Times New Roman.
        runs = _runs(self.para)
        self.assertEqual(runs[0], ("Use ", None, None, "Times New Roman"))
        self.assertEqual(runs[1][0], "case_id")
        self.assertEqual(runs[1][3], "Consolas")
        self.assertEqual(runs[2], (" here.", None, None, "Times New Roman"))

    def test_no_markdown_is_plain_run(self):
        render_inline_markdown(self.para, "Just a plain sentence.")
        self.assertEqual(_runs(self.para), [
            ("Just a plain sentence.", None, None, "Times New Roman"),
        ])

    def test_empty_text_is_noop(self):
        render_inline_markdown(self.para, "")
        self.assertEqual(self.para.runs, [])

    def test_bold_takes_precedence_over_italic(self):
        """``**foo**`` must parse as bold, not as italic-italic-empty-italic-italic."""
        render_inline_markdown(self.para, "**only bold**")
        self.assertEqual(_runs(self.para), [
            ("only bold", True, None, "Times New Roman"),
        ])

    def test_no_match_on_list_bullet_remnant(self):
        """If a caller forgets to strip a leading ``* `` from a list item, the
        single-asterisk italic regex must NOT match it as ``*  some text*``."""
        render_inline_markdown(self.para, "* unstripped bullet")
        # The whole string stays plain.
        self.assertEqual(_runs(self.para), [
            ("* unstripped bullet", None, None, "Times New Roman"),
        ])

    def test_unbalanced_asterisk_is_left_alone(self):
        """A stray ``*`` with no closing partner must not eat content."""
        render_inline_markdown(self.para, "5 * 3 equals 15")
        self.assertEqual(_runs(self.para), [
            ("5 * 3 equals 15", None, None, "Times New Roman"),
        ])

    def test_mixed_bold_italic_code(self):
        render_inline_markdown(
            self.para,
            "Plaintiff **John Doe** cites *Marbury v. Madison*, see `§ 1983`.",
        )
        run_texts = [r.text for r in self.para.runs]
        self.assertIn("John Doe", run_texts)
        self.assertIn("Marbury v. Madison", run_texts)
        self.assertIn("§ 1983", run_texts)
        # Asterisks/backticks must not appear in any run text.
        for r in self.para.runs:
            self.assertNotIn("*", r.text)
            self.assertNotIn("`", r.text)


class TestScriptInlineMarkdownIntegration(unittest.TestCase):
    """End-to-end: each script's ``add_markdown_to_doc`` must not leak
    ``**`` / ``*`` / ``\\`` markers into the resulting Word runs.
    """

    MARKDOWN = (
        "Plaintiff **John Doe** alleges *severe* back pain.\n"
        "See `§ 1983` for the statute. Defendant **disputes** all allegations.\n"
    )

    def _runs_text(self, doc):
        """Concatenate text from all paragraph runs in document order."""
        out = []
        for para in doc.paragraphs:
            for r in para.runs:
                out.append(r.text)
        return "".join(out)

    def _assert_no_markers_leaked(self, module_name: str):
        import importlib
        mod = importlib.import_module(module_name)
        func = getattr(mod, "add_markdown_to_doc")
        doc = Document()
        func(doc, self.MARKDOWN)
        text = self._runs_text(doc)
        self.assertNotIn(
            "**", text,
            f"{module_name}: literal '**' leaked through to the Word runs:\n{text!r}",
        )
        self.assertNotIn(
            "`", text,
            f"{module_name}: literal backtick leaked through:\n{text!r}",
        )
        # Bold/italic must actually be applied somewhere.
        bold_runs = [
            r for para in doc.paragraphs for r in para.runs if r.bold
        ]
        self.assertGreater(
            len(bold_runs), 0,
            f"{module_name}: no bold runs produced for **bold** input",
        )

    def test_summarize(self):
        self._assert_no_markers_leaked("Scripts.summarize")

    def test_summarize_discovery(self):
        self._assert_no_markers_leaked("Scripts.summarize_discovery")

    def test_summarize_deposition(self):
        self._assert_no_markers_leaked("Scripts.summarize_deposition")

    def test_med_record(self):
        self._assert_no_markers_leaked("Scripts.med_record")

    def test_med_chron(self):
        self._assert_no_markers_leaked("Scripts.med_chron")

    def test_liability(self):
        self._assert_no_markers_leaked("Scripts.liability")

    def test_exposure(self):
        self._assert_no_markers_leaked("Scripts.exposure")

    def test_docx_writer_internal(self):
        from icharlotte_core.docx_writer import _add_markdown_to_doc
        doc = Document()
        _add_markdown_to_doc(doc, self.MARKDOWN)
        text = "".join(
            r.text for para in doc.paragraphs for r in para.runs
        )
        self.assertNotIn("**", text)
        self.assertNotIn("`", text)


class TestMedRecordDatePreserved(unittest.TestCase):
    """The med_record / med_chron date-underline branch had its own buggy
    regex (``\\**`` zero-or-more) for the post-date text. The new helper must
    keep the date underline AND render the bold in the rest of the line."""

    def test_med_record_date_branch_renders_bold(self):
        import importlib
        mod = importlib.import_module("Scripts.med_record")
        doc = Document()
        mod.add_markdown_to_doc(
            doc,
            "On January 5, 2024, Dr. Smith prescribed **Norco 10/325 mg**.",
        )
        # Find the underline run (the date) and a bold run (the medication).
        all_runs = [r for para in doc.paragraphs for r in para.runs]
        underline_runs = [r for r in all_runs if r.underline]
        bold_runs = [r for r in all_runs if r.bold]
        self.assertTrue(any("January 5, 2024" in r.text for r in underline_runs))
        self.assertTrue(any("Norco 10/325 mg" in r.text for r in bold_runs))
        # No leftover markers.
        for r in all_runs:
            self.assertNotIn("**", r.text)


if __name__ == "__main__":
    unittest.main()
