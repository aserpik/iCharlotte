import json
import subprocess
import sys
from pathlib import Path

import pytest

pytest.importorskip("docx")
from docx import Document


from icharlotte_core.subpoena_tracker import SubpoenaTrackerWorker


def _write_pdf(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_bytes(b"%PDF-1.4\n% subpoena tracker fixture\n")


def _make_case(tmp_path: Path) -> Path:
    case_path = tmp_path / "5800.999 - Fixture"
    (case_path / "DISCOVERY" / "Subpoenas").mkdir(parents=True)
    (case_path / "RECORDS" / "Subpoenaed").mkdir(parents=True)
    (case_path / "NOTES").mkdir()
    return case_path


def _make_case_without_subpoena_folder(tmp_path: Path) -> Path:
    case_path = tmp_path / "5800.999 - Fixture"
    (case_path / "DISCOVERY").mkdir(parents=True)
    (case_path / "RECORDS").mkdir(parents=True)
    (case_path / "NOTES").mkdir()
    return case_path


def _report_rows(docx_path: Path) -> dict[str, list[str]]:
    doc = Document(str(docx_path))
    table = doc.tables[0]
    rows = {}
    for row in table.rows[1:]:
        cells = [cell.text.strip() for cell in row.cells]
        if cells and cells[0]:
            rows[cells[0]] = cells
    return rows


def _doc_text(docx_path: Path) -> str:
    doc = Document(str(docx_path))
    paragraphs = [p.text for p in doc.paragraphs]
    table_text = []
    for table in doc.tables:
        for row in table.rows:
            table_text.extend(cell.text for cell in row.cells)
    return "\n".join(paragraphs + table_text)


def test_unnumbered_received_record_matches_issued_subpoena_by_facility(tmp_path):
    case_path = _make_case(tmp_path)
    _write_pdf(
        case_path
        / "DISCOVERY"
        / "Subpoenas"
        / "61106-0019 TEMECULA VALLEY FAMILY PHYSICIANS.pdf"
    )
    received_path = (
        case_path
        / "RECORDS"
        / "Subpoenaed"
        / "Temecula Valley Family Physicians.pdf"
    )
    _write_pdf(received_path)

    worker = SubpoenaTrackerWorker(str(case_path), file_number="5800.999")
    warnings = []
    worker.warning.connect(warnings.append)

    subpoenas, ok = worker._scan_issued_subpoenas()
    received = worker._scan_received_records()
    worker._match_received_by_facility(subpoenas, received)

    assert ok
    assert received["61106-0019"]["status"] == "Yes"
    assert Path(received["61106-0019"]["path"]) == received_path
    assert not any(
        "Could not parse ID from received item: Temecula Valley Family Physicians.pdf"
        in message
        for message in warnings
    )


def test_date_first_cnr_received_record_matches_by_cleaned_facility(tmp_path):
    case_path = _make_case(tmp_path)
    _write_pdf(
        case_path
        / "DISCOVERY"
        / "Subpoenas"
        / "62849-0043 PROGRESSIVE CASUALTY INSURANCE COMPANY.pdf"
    )
    received_path = (
        case_path
        / "RECORDS"
        / "Subpoenaed"
        / "2025-08-12-CNR Progressive Casualty Insurance.pdf"
    )
    _write_pdf(received_path)

    worker = SubpoenaTrackerWorker(str(case_path), file_number="5800.999")
    subpoenas, ok = worker._scan_issued_subpoenas()
    received = worker._scan_received_records()
    worker._match_received_by_facility(subpoenas, received)

    assert ok
    assert received["62849-0043"]["status"] == "CNR"
    assert Path(received["62849-0043"]["path"]) == received_path


def test_empty_subpoena_folder_generates_clean_report(tmp_path):
    case_path = _make_case(tmp_path)
    output_root = tmp_path / "sandbox-output"

    worker = SubpoenaTrackerWorker(
        str(case_path),
        file_number="5800.999",
        output_root=str(output_root),
    )
    finished = []
    worker.finished_result.connect(lambda success, result: finished.append((success, result)))

    worker._run_phases()

    assert finished
    success, output_path = finished[-1]
    assert success
    output_docx = Path(output_path)
    assert output_docx.exists()
    assert output_root in output_docx.parents
    text = "\n".join(p.text for p in Document(str(output_docx)).paragraphs)
    assert "No issued subpoena PDFs were found" in text


def test_alternate_records_folder_surfaces_possible_match(tmp_path):
    case_path = _make_case(tmp_path)
    _write_pdf(
        case_path
        / "DISCOVERY"
        / "Subpoenas"
        / "61106-0019 TEMECULA VALLEY FAMILY PHYSICIANS.pdf"
    )
    alternate_path = (
        case_path
        / "RECORDS"
        / "Medical Docs"
        / "Temecula Valley Family Physicians Medical Records.pdf"
    )
    _write_pdf(alternate_path)

    worker = SubpoenaTrackerWorker(str(case_path), file_number="5800.999")
    finished = []
    worker.finished_result.connect(lambda success, result: finished.append((success, result)))

    worker._run_phases()

    assert finished[-1][0]
    output_docx = Path(finished[-1][1])
    rows = _report_rows(output_docx)
    assert rows["61106-0019"][2] == "Possible Match"
    text = _doc_text(output_docx)
    assert "Potential records found outside RECORDS/Subpoenaed" in text
    assert "Medical Docs" in text
    assert "Temecula Valley Family Physicians Medical Records.pdf" in text


def test_standard_received_record_takes_precedence_over_alternate_match(tmp_path):
    case_path = _make_case(tmp_path)
    _write_pdf(
        case_path
        / "DISCOVERY"
        / "Subpoenas"
        / "61106-0019 TEMECULA VALLEY FAMILY PHYSICIANS.pdf"
    )
    standard_path = (
        case_path
        / "RECORDS"
        / "Subpoenaed"
        / "61106-0019 Temecula Valley Family Physicians.pdf"
    )
    alternate_path = (
        case_path
        / "RECORDS"
        / "Medical Docs"
        / "Temecula Valley Family Physicians Medical Records.pdf"
    )
    _write_pdf(standard_path)
    _write_pdf(alternate_path)

    worker = SubpoenaTrackerWorker(str(case_path), file_number="5800.999")
    finished = []
    worker.finished_result.connect(lambda success, result: finished.append((success, result)))

    worker._run_phases()

    assert finished[-1][0]
    output_docx = Path(finished[-1][1])
    rows = _report_rows(output_docx)
    assert rows["61106-0019"][2] == "Yes"
    assert "Potential records found outside RECORDS/Subpoenaed" not in _doc_text(output_docx)


def test_missing_subpoenas_folder_generates_clean_report(tmp_path):
    case_path = _make_case_without_subpoena_folder(tmp_path)
    output_root = tmp_path / "sandbox-output"

    worker = SubpoenaTrackerWorker(
        str(case_path),
        file_number="5800.999",
        output_root=str(output_root),
    )
    finished = []
    worker.finished_result.connect(lambda success, result: finished.append((success, result)))

    worker._run_phases()

    assert finished
    success, output_path = finished[-1]
    assert success
    output_docx = Path(output_path)
    assert output_docx.exists()
    text = "\n".join(p.text for p in Document(str(output_docx)).paragraphs)
    assert "No DISCOVERY/Subpoenas folder was found" in text


def test_unparseable_subpoena_pdfs_generate_clean_report(tmp_path):
    case_path = _make_case(tmp_path)
    _write_pdf(
        case_path
        / "DISCOVERY"
        / "Subpoenas"
        / "Plaintiff Notice of Subpoena to Gifford Backhoe.pdf"
    )
    output_root = tmp_path / "sandbox-output"

    worker = SubpoenaTrackerWorker(
        str(case_path),
        file_number="5800.999",
        output_root=str(output_root),
    )
    finished = []
    worker.finished_result.connect(lambda success, result: finished.append((success, result)))

    worker._run_phases()

    assert finished
    success, output_path = finished[-1]
    assert success
    text = "\n".join(p.text for p in Document(output_path).paragraphs)
    assert "No parseable subpoena IDs were found" in text


def test_records_summarized_skips_rows_with_missing_cells(monkeypatch):
    class Cell:
        def __init__(self, text):
            self.text = text

    class Row:
        def __init__(self, *values):
            self.cells = tuple(Cell(value) for value in values)

    class Table:
        rows = [
            Row("FILENAME", "PAGES"),
            Row("61106-0019 Kaiser Permanente.pdf", "12"),
            Row("merged note row"),
            Row("TOTAL", "12"),
        ]

    class FakeDocument:
        tables = [Table()]

    import icharlotte_core.subpoena_tracker as tracker_module

    monkeypatch.setattr(tracker_module, "Document", lambda _path: FakeDocument())
    worker = SubpoenaTrackerWorker("C:/case")

    entries = worker._read_records_summarized("ignored.docx", "chron.docx")

    assert entries == [
        {
            "sid": "61106-0019",
            "filename": "61106-0019 Kaiser Permanente.pdf",
            "pages": "12",
        }
    ]


def test_generate_report_validates_saved_docx(tmp_path, monkeypatch):
    import icharlotte_core.subpoena_tracker as tracker_module

    case_path = _make_case(tmp_path)
    calls = []

    class Result:
        has_errors = False
        error_count = 0

        def print_summary(self):
            raise AssertionError("passing validation should not print")

    def fake_validate(path):
        calls.append(Path(path))
        return Result()

    monkeypatch.setattr(
        tracker_module, "validate_subpoena_tracker_docx", fake_validate, raising=False
    )
    worker = SubpoenaTrackerWorker(str(case_path), file_number="5800.999")

    worker._generate_report(
        {"11111-0001": {"facility": "Kaiser Permanente", "subpoena_path": "subpoena.pdf"}},
        {"11111-0001": {"status": "Yes", "path": "records.pdf"}},
        {},
        {},
        {},
        {},
    )

    assert len(calls) == 1
    assert calls[0].name.startswith("Tracked_Subpoenas")


def test_cli_respects_output_root_and_emits_json(tmp_path):
    case_path = _make_case(tmp_path)
    _write_pdf(
        case_path / "DISCOVERY" / "Subpoenas" / "11111-0001 Kaiser Permanente.pdf"
    )
    output_root = tmp_path / "cli-output"

    completed = subprocess.run(
        [
            sys.executable,
            "Scripts/subpoena_tracker.py",
            "--case-path",
            str(case_path),
            "--file-number",
            "5800.999",
            "--output-root",
            str(output_root),
            "--json",
        ],
        cwd=Path(__file__).resolve().parents[1],
        text=True,
        capture_output=True,
        check=False,
    )

    assert completed.returncode == 0, completed.stderr
    payload = json.loads(completed.stdout)
    output_docx = Path(payload["result"])
    assert payload["success"] is True
    assert output_docx.exists()
    assert output_root in output_docx.parents
    assert _report_rows(output_docx)["11111-0001"][2] == "No"
