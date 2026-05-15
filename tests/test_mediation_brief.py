import json
import os
import tempfile
import unittest
from unittest.mock import patch, MagicMock


class TestStyleExtraction(unittest.TestCase):

    def test_extract_sections_from_text(self):
        from icharlotte_core.mediation_brief import MediationBriefGenerator
        sample_text = (
            "I.     INTRODUCTION\n"
            "This is the introduction paragraph.\n\n"
            "II.     STATEMENT OF FACTS\n"
            "These are the facts of the case.\n"
            "More facts here.\n\n"
            "III.     LIABILITY\n"
            "Liability arguments here.\n"
        )
        gen = MediationBriefGenerator.__new__(MediationBriefGenerator)
        sections = gen._extract_sections_from_text(sample_text)
        self.assertIn("introduction", sections)
        self.assertIn("statement_of_facts", sections)
        self.assertIn("liability", sections)
        self.assertIn("introduction paragraph", sections["introduction"])
        self.assertIn("facts of the case", sections["statement_of_facts"])

    def test_extract_sections_handles_missing_sections(self):
        from icharlotte_core.mediation_brief import MediationBriefGenerator
        sample_text = (
            "I.     INTRODUCTION\n"
            "Intro text.\n\n"
            "VII.     CONCLUSION\n"
            "Conclusion text.\n"
        )
        gen = MediationBriefGenerator.__new__(MediationBriefGenerator)
        sections = gen._extract_sections_from_text(sample_text)
        self.assertIn("introduction", sections)
        self.assertIn("conclusion", sections)
        self.assertNotIn("liability", sections)

    def test_cache_structure(self):
        from icharlotte_core.mediation_brief import MediationBriefGenerator
        gen = MediationBriefGenerator.__new__(MediationBriefGenerator)
        gen._sample_dir = "C:\\AI\\Mediation Briefs"
        with patch.object(gen, '_read_sample_pdfs') as mock_read:
            mock_read.return_value = {
                "hashes": {"sample1.pdf": "abc123"},
                "sections": {
                    "introduction": ["Intro text from sample 1"],
                    "liability": ["Liability text from sample 1"],
                }
            }
            with tempfile.TemporaryDirectory() as tmpdir:
                cache_path = os.path.join(tmpdir, "style_cache.json")
                gen._cache_path = cache_path
                gen._save_style_cache(mock_read.return_value)
                with open(cache_path, 'r') as f:
                    cached = json.load(f)
                self.assertIn("source_hashes", cached)
                self.assertIn("sections", cached)
                self.assertEqual(cached["source_hashes"]["sample1.pdf"], "abc123")


class TestCaptionHandling(unittest.TestCase):

    def _make_caption_doc(self, tmpdir, include_sig_block=False):
        from docx import Document
        doc = Document()
        doc.add_paragraph("LAW FIRM NAME")
        doc.add_paragraph("CAPTION PAGE")
        doc.add_paragraph("Some caption content")
        if include_sig_block:
            doc.add_paragraph("")
            doc.add_paragraph("DATED: April 10, 2026")
            doc.add_paragraph("")
            doc.add_paragraph("By: ____________________")
            doc.add_paragraph("John Smith, Esq.")
            doc.add_paragraph("State Bar No. 123456")
        path = os.path.join(tmpdir, "case_caption.docx")
        doc.save(path)
        return path

    def test_find_caption_in_folder(self):
        from docx import Document
        from icharlotte_core.mediation_brief import MediationBriefGenerator
        gen = MediationBriefGenerator()
        with tempfile.TemporaryDirectory() as tmpdir:
            self._make_caption_doc(tmpdir)
            doc2 = Document()
            doc2.add_paragraph("Not a caption")
            doc2.save(os.path.join(tmpdir, "other_doc.docx"))
            result = gen.find_caption_template(tmpdir)
            self.assertIsNotNone(result)
            self.assertIn("caption", os.path.basename(result).lower())

    def test_find_caption_returns_none_when_missing(self):
        from docx import Document
        from icharlotte_core.mediation_brief import MediationBriefGenerator
        gen = MediationBriefGenerator()
        with tempfile.TemporaryDirectory() as tmpdir:
            doc = Document()
            doc.add_paragraph("Not a caption")
            doc.save(os.path.join(tmpdir, "other.docx"))
            result = gen.find_caption_template(tmpdir)
            self.assertIsNone(result)

    def test_replace_caption_page_text(self):
        from docx import Document
        from icharlotte_core.mediation_brief import MediationBriefGenerator
        gen = MediationBriefGenerator()
        with tempfile.TemporaryDirectory() as tmpdir:
            caption_path = self._make_caption_doc(tmpdir)
            output_path = os.path.join(tmpdir, "output.docx")
            gen.prepare_caption_template(caption_path, output_path)
            doc = Document(output_path)
            all_text = "\n".join(p.text for p in doc.paragraphs)
            self.assertNotIn("CAPTION PAGE", all_text)
            self.assertIn("DEFENDANT'S CONFIDENTIAL MEDIATION BRIEF", all_text)

    def test_signature_block_detection(self):
        from icharlotte_core.mediation_brief import MediationBriefGenerator
        gen = MediationBriefGenerator()
        with tempfile.TemporaryDirectory() as tmpdir:
            caption_path = self._make_caption_doc(tmpdir, include_sig_block=True)
            output_path = os.path.join(tmpdir, "output.docx")
            sig_paras = gen.prepare_caption_template(caption_path, output_path)
            self.assertTrue(len(sig_paras) > 0)
            sig_text = " ".join(p.text for p in sig_paras)
            self.assertIn("DATED", sig_text)

    def test_no_signature_block(self):
        from icharlotte_core.mediation_brief import MediationBriefGenerator
        gen = MediationBriefGenerator()
        with tempfile.TemporaryDirectory() as tmpdir:
            caption_path = self._make_caption_doc(tmpdir, include_sig_block=False)
            output_path = os.path.join(tmpdir, "output.docx")
            sig_paras = gen.prepare_caption_template(caption_path, output_path)
            self.assertEqual(len(sig_paras), 0)

    def test_footer_replacement_handles_run_nested_in_hyperlink(self):
        # Repro for "Assembly error: Element is not a child of this node":
        # caption-page footers commonly contain a w:r nested inside w:hyperlink
        # (e.g. the case-name hyperlink in pleading-paper templates). The footer
        # replacement code must not assume runs are direct children of w:p.
        from docx import Document
        from lxml import etree
        from icharlotte_core.mediation_brief import MediationBriefGenerator

        W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        gen = MediationBriefGenerator()

        with tempfile.TemporaryDirectory() as tmpdir:
            doc = Document()
            doc.add_paragraph("Body content")

            footer = doc.sections[0].footer
            footer_para = footer.paragraphs[0]
            para_elem = footer_para._element

            hyperlink = etree.SubElement(para_elem, f"{{{W_NS}}}hyperlink")
            run = etree.SubElement(hyperlink, f"{{{W_NS}}}r")
            t = etree.SubElement(run, f"{{{W_NS}}}t")
            t.text = "Smith v. Jones - Case No. 12345"

            path = os.path.join(tmpdir, "caption_with_nested_footer.docx")
            doc.save(path)

            reopened = Document(path)
            # Must not raise "Element is not a child of this node"
            gen._replace_caption_page_footers(reopened)


class TestSectionGeneration(unittest.TestCase):

    def test_build_section_prompt_includes_planning_output(self):
        from icharlotte_core.mediation_brief import MediationBriefGenerator
        gen = MediationBriefGenerator()
        gen.planning_output = "KEY FACTS:\n- Accident on Jan 1, 2025"
        gen.sections = {}
        gen.document_content = "Document text here"
        gen._style_cache = {"sections": {}}
        prompt = gen._build_section_prompt("statement_of_facts")
        self.assertIn("Accident on Jan 1, 2025", prompt)

    def test_build_section_prompt_includes_prior_sections(self):
        from icharlotte_core.mediation_brief import MediationBriefGenerator
        gen = MediationBriefGenerator()
        gen.planning_output = "Planning data"
        gen.sections = {
            "statement_of_facts": "The plaintiff was injured on Main St.",
            "procedural_status": "Trial is set for June 2027.",
        }
        gen.document_content = "Doc text"
        gen._style_cache = {"sections": {}}
        prompt = gen._build_section_prompt("liability")
        self.assertIn("injured on Main St", prompt)
        self.assertIn("Trial is set for June 2027", prompt)

    def test_build_introduction_prompt_includes_all_sections(self):
        from icharlotte_core.mediation_brief import MediationBriefGenerator
        gen = MediationBriefGenerator()
        gen.planning_output = "Planning data"
        gen.sections = {
            "statement_of_facts": "Facts text",
            "procedural_status": "Status text",
            "liability": "Liability text",
            "damages": "Damages text",
            "settlement_position": "Settlement text",
            "conclusion": "Conclusion text",
        }
        gen.document_content = "Doc text"
        gen._style_cache = {"sections": {}}
        prompt = gen._build_section_prompt("introduction")
        self.assertIn("Liability text", prompt)
        self.assertIn("Damages text", prompt)
        self.assertIn("Conclusion text", prompt)

    def test_generation_order(self):
        from icharlotte_core.mediation_brief import GENERATION_ORDER
        self.assertEqual(GENERATION_ORDER[-1], "introduction")
        self.assertEqual(len(GENERATION_ORDER), 7)


class TestTextParsing(unittest.TestCase):

    def test_parse_body_paragraphs(self):
        from icharlotte_core.mediation_brief import MediationBriefGenerator
        gen = MediationBriefGenerator()
        text = "First paragraph of the section.\n\nSecond paragraph here."
        elements = gen._parse_section_text(text, "statement_of_facts")
        body_elements = [e for e in elements if e["type"] == "body"]
        self.assertEqual(len(body_elements), 2)
        self.assertIn("First paragraph", body_elements[0]["text"])

    def test_parse_subsection_headings(self):
        from icharlotte_core.mediation_brief import MediationBriefGenerator
        gen = MediationBriefGenerator()
        text = (
            "Introduction paragraph.\n\n"
            "SUBSECTION: Defendant Owed No Duty Of Care\n"
            "Content under first subsection.\n\n"
            "SUBSECTION: Plaintiff Was Comparatively At Fault\n"
            "Content under second subsection.\n"
        )
        elements = gen._parse_section_text(text, "liability")
        headings = [e for e in elements if e["type"] == "l2_heading"]
        self.assertEqual(len(headings), 2)
        self.assertEqual(headings[0]["text"], "Defendant Owed No Duty Of Care")
        self.assertEqual(headings[1]["text"], "Plaintiff Was Comparatively At Fault")

    def test_parse_deposition_quotes(self):
        from icharlotte_core.mediation_brief import MediationBriefGenerator
        gen = MediationBriefGenerator()
        text = (
            "The plaintiff testified as follows:\n\n"
            "I was not paying attention to the road at the time of the accident. "
            "(Smith Depo Trns., at p. 45:12.)\n\n"
            "This admission is significant."
        )
        elements = gen._parse_section_text(text, "liability")
        quotes = [e for e in elements if e["type"] == "depo_quote"]
        self.assertEqual(len(quotes), 1)
        self.assertIn("not paying attention", quotes[0]["text"])


class TestWordAssembly(unittest.TestCase):

    def test_assemble_creates_docx(self):
        from docx import Document
        from icharlotte_core.mediation_brief import MediationBriefGenerator
        gen = MediationBriefGenerator()
        gen.sections = {
            "introduction": "This is the introduction.",
            "statement_of_facts": "These are the facts.",
            "procedural_status": "Trial is set for June 2027.",
            "liability": "SUBSECTION: No Duty\nDefendant owed no duty.",
            "damages": "SUBSECTION: No Causation\nNo evidence of causation.",
            "settlement_position": "Policy limits are $1M.",
            "conclusion": "We will prevail at trial.",
        }
        with tempfile.TemporaryDirectory() as tmpdir:
            caption = Document()
            caption.add_paragraph("LAW FIRM")
            caption.add_paragraph("CAPTION PAGE")
            caption_path = os.path.join(tmpdir, "caption.docx")
            caption.save(caption_path)
            output_path = os.path.join(tmpdir, "brief.docx")
            gen.assemble_document(caption_path, output_path)
            self.assertTrue(os.path.exists(output_path))
            doc = Document(output_path)
            all_text = "\n".join(p.text for p in doc.paragraphs)
            self.assertIn("INTRODUCTION", all_text)
            self.assertIn("STATEMENT OF FACTS", all_text)
            self.assertIn("introduction", all_text.lower())


class TestPipeline(unittest.TestCase):

    def test_generate_all_sections_order(self):
        from icharlotte_core.mediation_brief import MediationBriefGenerator, GENERATION_ORDER
        gen = MediationBriefGenerator()
        gen.document_content = "Test document content"
        gen._style_cache = {"sections": {}}
        call_order = []
        def mock_generate(section_name, refinement_instruction=""):
            call_order.append(section_name)
            return f"Generated {section_name}"
        with patch.object(gen, 'generate_section', side_effect=mock_generate):
            with patch.object(gen, 'run_planning_pass', return_value="Planning output"):
                gen.generate_all_sections()
        self.assertEqual(call_order, GENERATION_ORDER)
        self.assertEqual(gen.sections["introduction"], "Generated introduction")
        self.assertEqual(gen.sections["conclusion"], "Generated conclusion")

    def test_pipeline_sets_active_flag(self):
        from icharlotte_core.mediation_brief import MediationBriefGenerator
        gen = MediationBriefGenerator()
        gen.document_content = "Test doc"
        gen._style_cache = {"sections": {}}
        with patch.object(gen, 'generate_section', return_value="text"):
            with patch.object(gen, 'run_planning_pass', return_value="plan"):
                gen.generate_all_sections()
        self.assertTrue(gen.is_active)

    def test_generate_all_sections_progress_callback(self):
        from icharlotte_core.mediation_brief import MediationBriefGenerator, GENERATION_ORDER
        gen = MediationBriefGenerator()
        gen.document_content = "Test doc"
        gen._style_cache = {"sections": {}}
        calls = []
        def progress(section_name, index, total):
            calls.append((section_name, index, total))
        with patch.object(gen, 'generate_section', return_value="text"):
            with patch.object(gen, 'run_planning_pass', return_value="plan"):
                gen.generate_all_sections(progress_callback=progress)
        # Should be called once for planning + once per section
        expected_total = len(GENERATION_ORDER) + 1
        self.assertEqual(len(calls), expected_total)
        # First call is planning at index 0
        self.assertEqual(calls[0], ("planning", 0, expected_total))
        # Remaining calls are sections at indices 1..N
        for i, section_name in enumerate(GENERATION_ORDER, start=1):
            self.assertEqual(calls[i], (section_name, i, expected_total))

    def test_reset_clears_state(self):
        from icharlotte_core.mediation_brief import MediationBriefGenerator
        gen = MediationBriefGenerator()
        gen.sections = {"introduction": "some text"}
        gen.planning_output = "some planning"
        gen.document_content = "some content"
        gen.caption_template_path = "/some/path.docx"
        gen.is_active = True
        gen.reset()
        self.assertEqual(gen.sections, {})
        self.assertEqual(gen.planning_output, "")
        self.assertEqual(gen.document_content, "")
        self.assertIsNone(gen.caption_template_path)
        self.assertFalse(gen.is_active)

    def test_worker_emits_signals(self):
        from icharlotte_core.mediation_brief import MediationBriefGenerator, MediationBriefWorker, GENERATION_ORDER
        import sys
        # PySide6 requires a QApplication to exist for QThread
        try:
            from PySide6.QtWidgets import QApplication
            app = QApplication.instance() or QApplication(sys.argv)
        except Exception:
            self.skipTest("PySide6 QApplication not available")

        gen = MediationBriefGenerator()
        gen.document_content = "Test doc"
        gen._style_cache = {"sections": {}}

        section_started_calls = []
        section_complete_calls = []
        all_complete_calls = []
        error_calls = []

        with patch.object(gen, 'generate_section', return_value="Generated text"):
            with patch.object(gen, 'run_planning_pass', return_value="Planning output"):
                worker = MediationBriefWorker(gen)
                worker.section_started.connect(lambda name, idx, total: section_started_calls.append((name, idx, total)))
                worker.section_complete.connect(lambda name, text: section_complete_calls.append((name, text)))
                worker.all_complete.connect(lambda d: all_complete_calls.append(d))
                worker.error.connect(lambda msg: error_calls.append(msg))
                worker.run()  # run synchronously in test

        self.assertEqual(len(error_calls), 0)
        self.assertEqual(len(all_complete_calls), 1)
        result_dict = all_complete_calls[0]
        for section in GENERATION_ORDER:
            self.assertIn(section, result_dict)
        self.assertEqual(len(section_complete_calls), len(GENERATION_ORDER))
        self.assertTrue(gen.is_active)

    def test_worker_stop_requested(self):
        from icharlotte_core.mediation_brief import MediationBriefGenerator, MediationBriefWorker
        import sys
        try:
            from PySide6.QtWidgets import QApplication
            app = QApplication.instance() or QApplication(sys.argv)
        except Exception:
            self.skipTest("PySide6 QApplication not available")

        gen = MediationBriefGenerator()
        gen.document_content = "Test doc"
        gen._style_cache = {"sections": {}}

        all_complete_calls = []
        error_calls = []
        section_complete_calls = []

        call_count = [0]
        def mock_generate(section_name, refinement_instruction=""):
            call_count[0] += 1
            return "text"

        with patch.object(gen, 'generate_section', side_effect=mock_generate):
            with patch.object(gen, 'run_planning_pass', return_value="plan"):
                worker = MediationBriefWorker(gen)
                worker.all_complete.connect(lambda d: all_complete_calls.append(d))
                worker.error.connect(lambda msg: error_calls.append(msg))
                worker.section_complete.connect(lambda name, text: section_complete_calls.append(name))
                worker.request_stop()
                worker.run()

        # After stop requested, no sections should be generated
        self.assertEqual(call_count[0], 0)
        self.assertEqual(len(section_complete_calls), 0)


class TestRefinement(unittest.TestCase):

    def test_route_sections_parses_comma_list(self):
        from icharlotte_core.mediation_brief import MediationBriefGenerator
        gen = MediationBriefGenerator()
        result = gen._parse_routing_response("damages,liability")
        self.assertEqual(result, ["damages", "liability"])

    def test_route_sections_handles_none(self):
        from icharlotte_core.mediation_brief import MediationBriefGenerator
        gen = MediationBriefGenerator()
        result = gen._parse_routing_response("none")
        self.assertEqual(result, [])

    def test_route_sections_filters_invalid(self):
        from icharlotte_core.mediation_brief import MediationBriefGenerator
        gen = MediationBriefGenerator()
        result = gen._parse_routing_response("damages,invalid_section,liability")
        self.assertEqual(result, ["damages", "liability"])

    def test_refine_regenerates_introduction_when_liability_changes(self):
        from icharlotte_core.mediation_brief import MediationBriefGenerator
        gen = MediationBriefGenerator()
        gen.sections = {
            "introduction": "Old intro",
            "statement_of_facts": "Facts",
            "procedural_status": "Status",
            "liability": "Old liability",
            "damages": "Old damages",
            "settlement_position": "Settlement",
            "conclusion": "Conclusion",
        }
        gen.planning_output = "Planning"
        gen.document_content = "Doc"
        gen._style_cache = {"sections": {}}
        gen.is_active = True
        call_log = []
        def mock_generate(section_name, refinement_instruction=""):
            call_log.append(section_name)
            return f"New {section_name}"
        with patch.object(gen, 'generate_section', side_effect=mock_generate):
            gen.refine_sections(["liability"], "Make it stronger")
        self.assertIn("liability", call_log)
        self.assertIn("introduction", call_log)
        self.assertEqual(gen.sections["liability"], "New liability")
        self.assertEqual(gen.sections["introduction"], "New introduction")


if __name__ == '__main__':
    unittest.main()
