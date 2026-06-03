"""Tests for validate_index_docx (offline separator-index validation)."""
import pytest

pytest.importorskip("docx")
from docx import Document

from icharlotte_core.word_validator import validate_index_docx


def _make_index(path, n_docs):
    doc = Document()
    doc.add_paragraph("INDEX OF DOCUMENTS - test")
    table = doc.add_table(rows=1, cols=4)
    hdr = table.rows[0].cells
    for i, h in enumerate(["Id", "Document Title", "Document Date", "Page Ranges"]):
        hdr[i].text = h
    for i in range(n_docs):
        cells = table.add_row().cells
        cells[0].text = str(i + 1)
        cells[1].text = f"Doc {i + 1}"
    doc.save(str(path))


def test_missing_file_is_error(tmp_path):
    result = validate_index_docx(str(tmp_path / "nope.docx"), expected_doc_count=3)
    assert result.has_errors


def test_valid_index_passes(tmp_path):
    p = tmp_path / "Index_test.docx"
    _make_index(p, 3)
    result = validate_index_docx(str(p), expected_doc_count=3)
    assert not result.has_errors


def test_row_count_mismatch_warns(tmp_path):
    p = tmp_path / "Index_test.docx"
    _make_index(p, 2)
    result = validate_index_docx(str(p), expected_doc_count=5)
    assert result.has_warnings


def test_header_only_table_is_error(tmp_path):
    p = tmp_path / "Index_test.docx"
    _make_index(p, 0)  # header row only, no data rows
    result = validate_index_docx(str(p))
    assert result.has_errors
