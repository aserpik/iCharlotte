"""Tests for the response assembler (Word document generation)."""
import os
import tempfile
import unittest

from docx import Document as DocxDocument

from icharlotte_core.discovery.response_assembler import (
    ResponseAssembler,
    build_response_title,
    build_response_filename,
    parse_response_text,
    _fill_firm_contact_placeholders,
)
from icharlotte_core.discovery.response_parser import ParsedDiscovery
from icharlotte_core.discovery.response_rules import ResponseRules


class TestBuildResponseTitle(unittest.TestCase):
    def test_fi_title(self):
        title = build_response_title(
            responding_name="USA WASTE OF CALIFORNIA, INC.",
            responding_role="Defendant",
            propounding_role="Plaintiff",
            disc_type="FI",
            set_word="ONE",
        )
        self.assertIn("USA WASTE", title.upper())
        self.assertIn("FORM INTERROGATORIES", title.upper())
        self.assertIn("SET ONE", title.upper())
        self.assertIn("DEFENDANT", title.upper())
        self.assertIn("PLAINTIFF", title.upper())

    def test_rfa_title(self):
        title = build_response_title(
            responding_name="ACME CORP",
            responding_role="Defendant",
            propounding_role="Plaintiff",
            disc_type="RFA",
            set_word="TWO",
        )
        self.assertIn("REQUESTS FOR ADMISSION", title.upper())
        self.assertIn("SET TWO", title.upper())

    def test_si_title(self):
        title = build_response_title(
            responding_name="SMITH LLC",
            responding_role="Defendant",
            propounding_role="Plaintiff",
            disc_type="SI",
            set_word="ONE",
        )
        self.assertIn("SPECIAL INTERROGATORIES", title.upper())

    def test_rpd_title(self):
        title = build_response_title(
            responding_name="DOE INC",
            responding_role="Defendant",
            propounding_role="Plaintiff",
            disc_type="RPD",
            set_word="THREE",
        )
        self.assertIn("REQUEST FOR PRODUCTION", title.upper())


class TestBuildResponseFilename(unittest.TestCase):
    def test_fi_filename(self):
        fn = build_response_filename("USA Waste", "FI", 1)
        self.assertEqual(fn, "Def USA Waste's Resp to FI(1).docx")

    def test_rpd_filename(self):
        fn = build_response_filename("Acme", "RPD", 2)
        self.assertEqual(fn, "Def Acme's Resp to RPD(2).docx")

    def test_si_filename(self):
        fn = build_response_filename("Smith", "SI", 3)
        self.assertEqual(fn, "Def Smith's Resp to SI(3).docx")


class TestParseResponseText(unittest.TestCase):
    def test_parse_si_responses(self):
        text = (
            "SPECIAL INTERROGATORY NO. 1:\n"
            "Describe the incident.\n\n"
            "RESPONSE TO SPECIAL INTERROGATORY NO. 1:\n"
            "Objection. Subject to objections, response here.\n\n"
            "SPECIAL INTERROGATORY NO. 2:\n"
            "State all facts.\n\n"
            "RESPONSE TO SPECIAL INTERROGATORY NO. 2:\n"
            "Objection. Response two.\n"
        )
        pairs = parse_response_text(text, "SI")
        self.assertEqual(len(pairs), 2)
        self.assertEqual(pairs[0]["number"], "1")
        self.assertIn("Describe the incident", pairs[0]["request"])
        self.assertIn("response here", pairs[0]["response"])

    def test_parse_fi_responses(self):
        text = (
            "FORM INTERROGATORY NO. 1.1:\n"
            "State the name...\n\n"
            "RESPONSE TO FORM INTERROGATORY NO. 1.1:\n"
            "Attorney info.\n\n"
            "FORM INTERROGATORY NO. 3.1:\n"
            "Are you a corporation?\n\n"
            "RESPONSE TO FORM INTERROGATORY NO. 3.1:\n"
            "Yes.\n"
        )
        pairs = parse_response_text(text, "FI")
        self.assertEqual(len(pairs), 2)
        self.assertEqual(pairs[0]["number"], "1.1")

    def test_parse_rfa_responses(self):
        text = (
            "REQUEST FOR ADMISSION NO. 1:\n"
            "Admit that...\n\n"
            "RESPONSE TO REQUEST FOR ADMISSION NO. 1:\n"
            "Deny.\n"
        )
        pairs = parse_response_text(text, "RFA")
        self.assertEqual(len(pairs), 1)
        self.assertEqual(pairs[0]["number"], "1")

    def test_parse_rpd_responses(self):
        text = (
            "REQUEST FOR PRODUCTION NO. 1:\n"
            "All documents...\n\n"
            "RESPONSE TO REQUEST NO. 1:\n"
            "Will comply.\n"
        )
        pairs = parse_response_text(text, "RPD")
        self.assertEqual(len(pairs), 1)
        self.assertEqual(pairs[0]["number"], "1")

    def test_empty_text(self):
        pairs = parse_response_text("", "SI")
        self.assertEqual(len(pairs), 0)


class TestInsertResponsePairs(unittest.TestCase):
    def test_no_objections_does_not_insert_waiver(self):
        doc = DocxDocument()
        assembler = ResponseAssembler.__new__(ResponseAssembler)
        rules = ResponseRules()

        assembler._insert_response_pairs(
            doc,
            [
                {
                    "number": "1.1",
                    "request": "State who prepared the responses.",
                    "response": "Attorney information.",
                }
            ],
            "FI",
            "ONE",
            rules,
        )

        full_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
        self.assertNotIn(rules.waiver_language, full_text)
        self.assertIn("Attorney information.", full_text)

    def test_objections_still_insert_waiver_inline(self):
        doc = DocxDocument()
        assembler = ResponseAssembler.__new__(ResponseAssembler)
        rules = ResponseRules()

        assembler._insert_response_pairs(
            doc,
            [
                {
                    "number": "15.1",
                    "request": "Identify denials and defenses.",
                    "response": (
                        "Objection text. "
                        f"{rules.waiver_language}\n"
                        "General denial response."
                    ),
                }
            ],
            "FI",
            "ONE",
            rules,
        )

        full_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
        self.assertIn(f"Objection text. {rules.waiver_language}", full_text)
        self.assertIn("General denial response.", full_text)


class TestFirmContactPlaceholders(unittest.TestCase):
    def test_fills_firm_contact_placeholders_from_caption_text(self):
        text = (
            "Responding Party and its attorneys of record, {firm_name}, "
            "{firm_address}; {firm_phone}."
        )
        caption_text = (
            "Dated: April 20, 2026\t\t\tBORDIN SEMMER LLP\n"
            "BORDIN SEMMER LLP\n"
            "Joshua Bordin-Wosk, State Bar No. 241077\n"
            "jbordinwosk@bordinsemmer.com\n"
            "101 Continental Blvd., Suite 700\n"
            "El Segundo, CA 90245\n"
            "Phone:\t(323) 457-2110\n"
            "Fax:\t(323) 457-2120\n"
        )

        filled = _fill_firm_contact_placeholders(text, caption_text)

        self.assertIn("BORDIN SEMMER LLP, 101 Continental", filled)
        self.assertNotIn("Dated:", filled)
        self.assertIn("101 Continental Blvd., Suite 700, El Segundo, CA 90245", filled)
        self.assertIn("(323) 457-2110", filled)
        self.assertNotIn("{firm_", filled)


if __name__ == "__main__":
    unittest.main()
