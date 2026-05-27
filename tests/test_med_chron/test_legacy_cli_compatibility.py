"""Tests that ``python med_chron.py <file>`` (no --phase) still works.

This is the IndexTab agent-runner path. It must produce
``med_chron_<filename>.docx`` exactly like the pre-refactor agent.
"""

import os
import sys
from pathlib import Path
from unittest.mock import patch

import pytest
from docx import Document

PROJECT_ROOT = Path(__file__).resolve().parent.parent.parent
sys.path.insert(0, str(PROJECT_ROOT / "Scripts"))

import med_chron  # noqa: E402


def _make_chrono_docx(path: Path, *, pre="Pre.", post="Post.") -> Path:
    doc = Document()
    doc.add_paragraph(f"BRIEF SYNOPSIS OF PRE-INJURY MEDICAL RECORD: {pre}")
    doc.add_paragraph(f"BRIEF SYNOPSIS OF POST-INJURY MEDICAL RECORD: {post}")
    doc.save(str(path))
    return path


def test_legacy_dispatcher_dispatches_to_legacy_when_no_phase(tmp_path, monkeypatch):
    src = _make_chrono_docx(tmp_path / "1234-001_ ACME.docx")

    called = {"legacy": False, "prep": False, "run": False}

    monkeypatch.setattr(med_chron, "process_legacy",
                        lambda p, **kw: called.__setitem__("legacy", True) or 0)
    monkeypatch.setattr(med_chron, "process_prep",
                        lambda p, o, **kw: called.__setitem__("prep", True) or 0)
    monkeypatch.setattr(med_chron, "process_run",
                        lambda p, o, **kw: called.__setitem__("run", True) or 0)

    monkeypatch.setattr(sys, "argv", ["med_chron.py", str(src)])
    with pytest.raises(SystemExit) as exc:
        med_chron.main()
    assert exc.value.code == 0
    assert called == {"legacy": True, "prep": False, "run": False}


def test_dispatcher_routes_prep_phase(tmp_path, monkeypatch):
    src = _make_chrono_docx(tmp_path / "rec.docx")
    called = []
    monkeypatch.setattr(med_chron, "process_prep",
                        lambda p, o, **kw: called.append(("prep", p, o)) or 0)
    monkeypatch.setattr(sys, "argv",
                        ["med_chron.py", "--phase=prep", str(src)])
    with pytest.raises(SystemExit) as exc:
        med_chron.main()
    assert exc.value.code == 0
    assert called and called[0][0] == "prep"


def test_dispatcher_routes_run_phase(tmp_path, monkeypatch):
    fake_session = tmp_path / "s.json"
    fake_session.write_text("{}", encoding="utf-8")
    called = []
    monkeypatch.setattr(med_chron, "process_run",
                        lambda p, o, **kw: called.append(("run", p, o)) or 0)
    monkeypatch.setattr(sys, "argv",
                        ["med_chron.py", "--phase=run", str(fake_session)])
    with pytest.raises(SystemExit) as exc:
        med_chron.main()
    assert exc.value.code == 0
    assert called and called[0][0] == "run"


def test_legacy_mode_writes_existing_filename_pattern(tmp_path):
    src = _make_chrono_docx(tmp_path / "1234-001_ ACME PT.docx")
    out_dir = tmp_path / "NOTES" / "AI OUTPUT"
    out_dir.mkdir(parents=True)

    with patch.object(med_chron.LLMCaller, "call", return_value="# Rewrite\nbody"):
        rc = med_chron.process_legacy(str(src), output_dir_override=str(out_dir))

    assert rc == 0
    # The existing sanitizer keeps dashes (regex is [^a-zA-Z0-9_\-]).
    # Input "1234-001_ ACME PT.docx" → safe_name "1234-001__ACME_PT".
    expected = out_dir / "med_chron_1234-001__ACME_PT.docx"
    assert expected.exists(), f"missing expected output: {expected}; have: {list(out_dir.iterdir())}"
