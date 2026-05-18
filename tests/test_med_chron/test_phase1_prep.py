"""Tests for Phase 1 (prep) of the Med-Cron agent."""

import json
import os
import sys
from pathlib import Path
from unittest.mock import patch

import pytest
from docx import Document

PROJECT_ROOT = Path(__file__).resolve().parent.parent.parent
sys.path.insert(0, str(PROJECT_ROOT / "Scripts"))

import med_chron  # noqa: E402


def _make_chrono_docx(path: Path, narrative_pre: str, narrative_post: str,
                     table_rows: list[list[str]]) -> Path:
    doc = Document()
    doc.add_paragraph(f"BRIEF SYNOPSIS OF PRE-INJURY MEDICAL RECORD: {narrative_pre}")
    doc.add_paragraph(f"BRIEF SYNOPSIS OF POST-INJURY MEDICAL RECORD: {narrative_post}")
    table = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
    for r, row in enumerate(table_rows):
        for c, cell in enumerate(row):
            table.cell(r, c).text = cell
    doc.save(str(path))
    return path


def test_phase1_writes_session_with_both_text_caches(tmp_path):
    src = tmp_path / "1234-001_ ACME PT.docx"
    _make_chrono_docx(
        src,
        narrative_pre="No prior back issues.",
        narrative_post="Treated for lumbar strain following the accident.",
        table_rows=[["Date", "Provider"], ["2024-01-15", "Acme PT"]],
    )
    out_dir = tmp_path / "NOTES" / "AI OUTPUT"

    rc = med_chron.process_prep(str(src), str(out_dir))
    assert rc == 0

    from icharlotte_core.med_chron import session_manager
    paths = session_manager.compute_session_paths(str(src), str(out_dir))

    data = session_manager.read_session(paths.session_path)
    assert data["phase"] == "awaiting_input"
    assert data["user_config"] is None
    assert data["narrative_missing"] is False
    assert paths.narrative_text_path.exists()
    assert paths.full_text_path.exists()


def test_phase1_narrative_only_excludes_table_content(tmp_path):
    src = tmp_path / "rec.docx"
    _make_chrono_docx(
        src,
        narrative_pre="No prior back issues.",
        narrative_post="Treated for lumbar strain.",
        table_rows=[["Date", "Provider"], ["2024-01-15", "UNIQUE_TABLE_TOKEN"]],
    )
    out_dir = tmp_path / "out"

    rc = med_chron.process_prep(str(src), str(out_dir))
    assert rc == 0

    from icharlotte_core.med_chron import session_manager
    paths = session_manager.compute_session_paths(str(src), str(out_dir))
    narrative = paths.narrative_text_path.read_text(encoding="utf-8")
    assert "UNIQUE_TABLE_TOKEN" not in narrative
    assert "lumbar strain" in narrative


def test_phase1_full_text_includes_table_rows(tmp_path):
    src = tmp_path / "rec.docx"
    _make_chrono_docx(
        src,
        narrative_pre="No prior issues.",
        narrative_post="Treated.",
        table_rows=[["Date", "Provider"], ["2024-01-15", "UNIQUE_TABLE_TOKEN"]],
    )
    out_dir = tmp_path / "out"

    rc = med_chron.process_prep(str(src), str(out_dir))
    assert rc == 0

    from icharlotte_core.med_chron import session_manager
    paths = session_manager.compute_session_paths(str(src), str(out_dir))
    full = paths.full_text_path.read_text(encoding="utf-8")
    assert "UNIQUE_TABLE_TOKEN" in full


def test_phase1_missing_synopsis_marks_narrative_missing(tmp_path):
    src = tmp_path / "no_synopsis.docx"
    doc = Document()
    doc.add_paragraph("This document has no BRIEF SYNOPSIS sections.")
    doc.add_paragraph("Just regular medical content.")
    doc.save(str(src))
    out_dir = tmp_path / "out"

    rc = med_chron.process_prep(str(src), str(out_dir))
    assert rc == 0  # still succeeds

    from icharlotte_core.med_chron import session_manager
    paths = session_manager.compute_session_paths(str(src), str(out_dir))
    data = session_manager.read_session(paths.session_path)
    assert data["narrative_missing"] is True
    assert paths.narrative_text_path.exists()
    assert paths.narrative_text_path.read_text(encoding="utf-8").strip() == ""


def test_phase1_prints_awaiting_input_token_on_success(tmp_path, capsys):
    src = tmp_path / "rec.docx"
    _make_chrono_docx(
        src,
        narrative_pre="Pre.",
        narrative_post="Post.",
        table_rows=[["A", "B"], ["1", "2"]],
    )
    out_dir = tmp_path / "out"

    med_chron.process_prep(str(src), str(out_dir))
    out = capsys.readouterr().out
    awaiting = [ln for ln in out.splitlines() if ln.startswith("AWAITING_INPUT:")]
    assert awaiting, f"AWAITING_INPUT token not printed; stdout: {out!r}"


def test_phase1_session_includes_catalog_snapshot(tmp_path):
    src = tmp_path / "rec.docx"
    _make_chrono_docx(
        src,
        narrative_pre="P.",
        narrative_post="Q.",
        table_rows=[["A", "B"], ["1", "2"]],
    )
    out_dir = tmp_path / "out"

    med_chron.process_prep(str(src), str(out_dir))
    from icharlotte_core.med_chron import session_manager
    paths = session_manager.compute_session_paths(str(src), str(out_dir))
    data = session_manager.read_session(paths.session_path)

    assert isinstance(data["catalog"], list)
    rewrite = next(e for e in data["catalog"] if e["id"] == "rewrite_chronology")
    assert rewrite["default_selected"] is True
    assert rewrite["uses_tables"] is False


def test_phase1_reuses_cache_on_unchanged_input(tmp_path):
    src = tmp_path / "rec.docx"
    _make_chrono_docx(
        src,
        narrative_pre="P.",
        narrative_post="Q.",
        table_rows=[["A", "B"], ["1", "2"]],
    )
    out_dir = tmp_path / "out"

    med_chron.process_prep(str(src), str(out_dir))
    from icharlotte_core.med_chron import session_manager
    paths = session_manager.compute_session_paths(str(src), str(out_dir))

    sentinel = "SENTINEL_TEXT_DO_NOT_OVERWRITE"
    paths.full_text_path.write_text(sentinel, encoding="utf-8")

    rc = med_chron.process_prep(str(src), str(out_dir))
    assert rc == 0
    assert paths.full_text_path.read_text(encoding="utf-8") == sentinel
