"""Tests for ResponseAssembler output formatting (footer + objection layout)."""
import os
import tempfile
import unittest

from docx import Document as DocxDocument

from icharlotte_core.discovery.response_assembler import ResponseAssembler
from icharlotte_core.discovery.response_parser import ParsedDiscovery, ParsedRequest
from icharlotte_core.discovery.response_review_state import ReviewState, RequestReview
from icharlotte_core.discovery.response_rules import ResponseRules

_REPO_ROOT = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
CAPTION_PAGE = os.path.join(_REPO_ROOT, "discovery", "Caption Page (AS FM).docx")
_SKIP_MSG = "Caption page template not found — skipping ResponseAssembler tests"


def _all_footer_text(doc) -> str:
    chunks = []
    for section in doc.sections:
        for footer in (
            section.first_page_footer,
            section.footer,
            section.even_page_footer,
        ):
            for para in footer.paragraphs:
                if para.text:
                    chunks.append(para.text)
    return "\n".join(chunks)


def _assemble(review: RequestReview) -> str:
    parsed = ParsedDiscovery(
        discovery_type="FI",
        propounding_party="Plaintiff SMITH",
        responding_party="Defendant JONES",
        set_number=1,
        set_word="ONE",
        case_number="1",
        requests=[ParsedRequest(number=review.number, text=review.request_text)],
    )
    rules = ResponseRules()
    response_text = ReviewState([review]).to_plain_text(parsed, rules)
    out = os.path.join(tempfile.mkdtemp(), "response.docx")
    ResponseAssembler(CAPTION_PAGE).assemble(
        parsed=parsed, response_text=response_text, rules=rules, output_path=out,
    )
    return out


@unittest.skipUnless(os.path.isfile(CAPTION_PAGE), _SKIP_MSG)
class ResponseAssemblerOutputTests(unittest.TestCase):
    def test_footer_has_no_pleading_title_placeholder(self):
        out = _assemble(
            RequestReview(
                number="2.1",
                request_text="State your name.",
                proposed_substantive_response="My name is John Jones.",
                approved=True,
            )
        )
        try:
            footer_text = _all_footer_text(DocxDocument(out))
        finally:
            os.remove(out)
        self.assertNotIn("PLEADING TITLE", footer_text.upper())
        self.assertIn("RESPONSES TO", footer_text.upper())

    def test_multiple_objections_render_as_single_paragraph(self):
        review = RequestReview(
            number="12.4",
            request_text="State the names of witnesses.",
            proposed_objections=(
                "Responding Party objects to this Request on the grounds that it "
                "calls for speculation and is vague, ambiguous, and overbroad.\n\n"
                "Responding Party objects to this Request on the grounds that it is "
                "unduly burdensome and overly broad as to time and scope."
            ),
            proposed_substantive_response="The witnesses are A and B.",
            approved=True,
        )
        out = _assemble(review)
        try:
            paras = [p.text for p in DocxDocument(out).paragraphs if p.text.strip()]
            obj_paras = [t for t in paras if "objects to this Request" in t]
            # All objections collapse into ONE paragraph.
            self.assertEqual(len(obj_paras), 1, "expected a single objection paragraph")
            self.assertIn("speculation", obj_paras[0])
            self.assertIn("unduly burdensome", obj_paras[0])
            # Only the substantive response begins a new paragraph.
            subst_paras = [t for t in paras if "The witnesses are A and B." in t]
            self.assertEqual(len(subst_paras), 1)
            self.assertNotIn("objects to this Request", subst_paras[0])
        finally:
            os.remove(out)


if __name__ == "__main__":
    unittest.main()
