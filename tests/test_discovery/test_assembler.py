"""Tests for the discovery document assembler."""
import os
import unittest
import tempfile
import shutil

from icharlotte_core.discovery.models import (
    DiscoveryType,
    Party,
    PartyRole,
    DiscoveryRequest,
    DiscoverySet,
)

# Caption page path (relative to repo root)
_REPO_ROOT = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
CAPTION_PAGE = os.path.join(_REPO_ROOT, "discovery", "Caption Page (AS FM).docx")

_SKIP_MSG = "Caption page not found — skipping assembler tests"


def _has_caption_page():
    return os.path.isfile(CAPTION_PAGE)


def _make_party(name, role, is_ours=False, abbreviation=""):
    return Party(
        name=name,
        role=role,
        is_our_client=is_ours,
        abbreviation=abbreviation,
    )


def _make_si_set(num_requests=5, previous_count=0, set_number=1):
    """Build a DiscoverySet of Special Interrogatories for testing."""
    pltf = _make_party("John Smith", PartyRole.PLAINTIFF, abbreviation="Pltf")
    deft = _make_party("Acme Corp", PartyRole.DEFENDANT, is_ours=True, abbreviation="Def")
    start = previous_count + 1
    requests = [
        DiscoveryRequest(
            number=start + i,
            text=f"State all facts relating to topic {start + i}.",
            definitions=[f'("Topic {start + i}" means the subject described above.)']
            if i % 3 == 0
            else [],
        )
        for i in range(num_requests)
    ]
    return DiscoverySet(
        discovery_type=DiscoveryType.SI,
        set_number=set_number,
        directed_to=pltf,
        propounding_party=deft,
        requests=requests,
        definitions_block="DEFINITIONS\n\n\"INCIDENT\" means the accident described in the complaint.",
        instructions_block="INSTRUCTIONS TO ANSWERING PARTY\n\nAnswer each interrogatory fully.",
        previous_count=previous_count,
    )


def _make_rpd_set(num_requests=5, set_number=1):
    """Build a DiscoverySet of Requests for Production for testing."""
    pltf = _make_party("John Smith", PartyRole.PLAINTIFF, abbreviation="Pltf")
    deft = _make_party("Acme Corp", PartyRole.DEFENDANT, is_ours=True, abbreviation="Def")
    requests = [
        DiscoveryRequest(
            number=i + 1,
            text=f"Produce all documents relating to document category {i + 1}.",
        )
        for i in range(num_requests)
    ]
    return DiscoverySet(
        discovery_type=DiscoveryType.RPD,
        set_number=set_number,
        directed_to=pltf,
        propounding_party=deft,
        requests=requests,
        definitions_block="DEFINITIONS\n\n\"DOCUMENT\" means any writing as defined in Evidence Code section 250.",
        instructions_block="INSTRUCTIONS TO ANSWERING PARTY\n\nProduce all responsive documents.",
    )


def _extract_all_text(docx_path):
    """Extract all paragraph text from a .docx file."""
    from docx import Document
    doc = Document(docx_path)
    return "\n".join(p.text for p in doc.paragraphs)


@unittest.skipUnless(_has_caption_page(), _SKIP_MSG)
class TestAssembleCreateDocx(unittest.TestCase):
    """test_assemble_creates_docx: creates file, contains 'SPECIAL INTERROGATORY NO. 1:'."""

    def setUp(self):
        self.tmpdir = tempfile.mkdtemp(prefix="test_assembler_")

    def tearDown(self):
        shutil.rmtree(self.tmpdir, ignore_errors=True)

    def test_assemble_creates_docx(self):
        from icharlotte_core.discovery.assembler import DiscoveryAssembler

        ds = _make_si_set(num_requests=5)
        output_path = os.path.join(self.tmpdir, "output", "test_si.docx")

        assembler = DiscoveryAssembler(caption_page_path=CAPTION_PAGE)
        result = assembler.assemble(
            ds, output_path, attorney_name="Jane Attorney"
        )

        # File was created
        self.assertTrue(os.path.isfile(result))
        self.assertEqual(result, output_path)

        # Contains request header
        text = _extract_all_text(output_path)
        self.assertIn("SPECIAL INTERROGATORY NO. 1:", text)


@unittest.skipUnless(_has_caption_page(), _SKIP_MSG)
class TestAssembleIncludesDefinitions(unittest.TestCase):
    """test_assemble_includes_definitions: definitions text appears in output."""

    def setUp(self):
        self.tmpdir = tempfile.mkdtemp(prefix="test_assembler_")

    def tearDown(self):
        shutil.rmtree(self.tmpdir, ignore_errors=True)

    def test_assemble_includes_definitions(self):
        from icharlotte_core.discovery.assembler import DiscoveryAssembler

        ds = _make_si_set(num_requests=5)
        output_path = os.path.join(self.tmpdir, "test_si_defs.docx")

        assembler = DiscoveryAssembler(caption_page_path=CAPTION_PAGE)
        assembler.assemble(ds, output_path, attorney_name="Jane Attorney")

        text = _extract_all_text(output_path)
        # Definitions block should be present
        self.assertIn("INCIDENT", text)
        # Inline definitions should also be present
        self.assertIn("Topic 1", text)


@unittest.skipUnless(_has_caption_page(), _SKIP_MSG)
class TestAssembleRPD(unittest.TestCase):
    """test_assemble_rpd: RPD requests appear correctly."""

    def setUp(self):
        self.tmpdir = tempfile.mkdtemp(prefix="test_assembler_")

    def tearDown(self):
        shutil.rmtree(self.tmpdir, ignore_errors=True)

    def test_assemble_rpd(self):
        from icharlotte_core.discovery.assembler import DiscoveryAssembler

        ds = _make_rpd_set(num_requests=3)
        output_path = os.path.join(self.tmpdir, "test_rpd.docx")

        assembler = DiscoveryAssembler(caption_page_path=CAPTION_PAGE)
        assembler.assemble(ds, output_path, attorney_name="Jane Attorney")

        text = _extract_all_text(output_path)
        self.assertIn("REQUEST FOR PRODUCTION NO. 1:", text)
        self.assertIn("REQUEST FOR PRODUCTION NO. 2:", text)
        self.assertIn("REQUEST FOR PRODUCTION NO. 3:", text)
        # Should contain preamble reference to CCP 2031
        self.assertIn("2031.010", text)


@unittest.skipUnless(_has_caption_page(), _SKIP_MSG)
class TestAssembleWithDeclaration(unittest.TestCase):
    """test_assemble_with_declaration: SI with >35 requests includes 'DECLARATION' and '2030.070'."""

    def setUp(self):
        self.tmpdir = tempfile.mkdtemp(prefix="test_assembler_")

    def tearDown(self):
        shutil.rmtree(self.tmpdir, ignore_errors=True)

    def test_assemble_with_declaration(self):
        from icharlotte_core.discovery.assembler import DiscoveryAssembler

        ds = _make_si_set(num_requests=40, previous_count=0)
        output_path = os.path.join(self.tmpdir, "test_si_decl.docx")

        assembler = DiscoveryAssembler(caption_page_path=CAPTION_PAGE)
        assembler.assemble(ds, output_path, attorney_name="Jane Attorney")

        text = _extract_all_text(output_path)
        self.assertIn("DECLARATION", text)
        self.assertIn("2030.070", text)


@unittest.skipUnless(_has_caption_page(), _SKIP_MSG)
class TestAssembleNoDeclarationUnder35(unittest.TestCase):
    """SI with <= 35 requests should NOT include a declaration."""

    def setUp(self):
        self.tmpdir = tempfile.mkdtemp(prefix="test_assembler_")

    def tearDown(self):
        shutil.rmtree(self.tmpdir, ignore_errors=True)

    def test_no_declaration_under_35(self):
        from icharlotte_core.discovery.assembler import DiscoveryAssembler

        ds = _make_si_set(num_requests=10, previous_count=0)
        output_path = os.path.join(self.tmpdir, "test_si_no_decl.docx")

        assembler = DiscoveryAssembler(caption_page_path=CAPTION_PAGE)
        assembler.assemble(ds, output_path, attorney_name="Jane Attorney")

        text = _extract_all_text(output_path)
        self.assertNotIn("2030.070", text)


@unittest.skipUnless(_has_caption_page(), _SKIP_MSG)
class TestAssembleFromPlainText(unittest.TestCase):
    """Test assemble_from_plain_text re-parses and delegates to assemble."""

    def setUp(self):
        self.tmpdir = tempfile.mkdtemp(prefix="test_assembler_")

    def tearDown(self):
        shutil.rmtree(self.tmpdir, ignore_errors=True)

    def test_assemble_from_plain_text(self):
        from icharlotte_core.discovery.assembler import DiscoveryAssembler

        ds = _make_si_set(num_requests=0)  # Start with no requests
        plain = (
            "SPECIAL INTERROGATORY NO. 1:\n\n"
            "\tState all facts relating to the incident.\n\n"
            "SPECIAL INTERROGATORY NO. 2:\n\n"
            "\tIdentify all witnesses.\n"
        )
        output_path = os.path.join(self.tmpdir, "test_from_text.docx")

        assembler = DiscoveryAssembler(caption_page_path=CAPTION_PAGE)
        result = assembler.assemble_from_plain_text(
            plain, ds, output_path, attorney_name="Jane Attorney"
        )

        self.assertTrue(os.path.isfile(result))
        text = _extract_all_text(output_path)
        self.assertIn("SPECIAL INTERROGATORY NO. 1:", text)
        self.assertIn("SPECIAL INTERROGATORY NO. 2:", text)
        self.assertIn("incident", text.lower())
        self.assertIn("witnesses", text.lower())


@unittest.skipUnless(_has_caption_page(), _SKIP_MSG)
class TestFindCaptionPage(unittest.TestCase):
    """Test find_caption_page static method finds caption page in a folder."""

    def setUp(self):
        self.tmpdir = tempfile.mkdtemp(prefix="test_assembler_")

    def tearDown(self):
        shutil.rmtree(self.tmpdir, ignore_errors=True)

    def test_find_caption_page_found(self):
        from icharlotte_core.discovery.assembler import DiscoveryAssembler

        # Copy caption page into a temp folder with the expected name
        dest = os.path.join(self.tmpdir, "Caption Page (AS FM).docx")
        shutil.copy2(CAPTION_PAGE, dest)

        result = DiscoveryAssembler.find_caption_page(self.tmpdir)
        self.assertIsNotNone(result)
        self.assertTrue(os.path.isfile(result))

    def test_find_caption_page_not_found(self):
        from icharlotte_core.discovery.assembler import DiscoveryAssembler

        # Empty folder — no caption page
        result = DiscoveryAssembler.find_caption_page(self.tmpdir)
        self.assertIsNone(result)


@unittest.skipUnless(_has_caption_page(), _SKIP_MSG)
class TestAssembleDocumentTitle(unittest.TestCase):
    """Document title should appear in the output."""

    def setUp(self):
        self.tmpdir = tempfile.mkdtemp(prefix="test_assembler_")

    def tearDown(self):
        shutil.rmtree(self.tmpdir, ignore_errors=True)

    def test_document_title_in_output(self):
        from icharlotte_core.discovery.assembler import DiscoveryAssembler

        ds = _make_si_set(num_requests=5)
        output_path = os.path.join(self.tmpdir, "test_title.docx")

        assembler = DiscoveryAssembler(caption_page_path=CAPTION_PAGE)
        assembler.assemble(ds, output_path, attorney_name="Jane Attorney")

        text = _extract_all_text(output_path)
        # Title should mention the propounding party and set number
        self.assertIn("SPECIAL INTERROGATORIES", text)
        self.assertIn("SET ONE", text)


@unittest.skipUnless(_has_caption_page(), _SKIP_MSG)
class TestAssembleSignatureBlock(unittest.TestCase):
    """Signature block should contain firm name and attorney name."""

    def setUp(self):
        self.tmpdir = tempfile.mkdtemp(prefix="test_assembler_")

    def tearDown(self):
        shutil.rmtree(self.tmpdir, ignore_errors=True)

    def test_signature_block(self):
        from icharlotte_core.discovery.assembler import DiscoveryAssembler

        ds = _make_si_set(num_requests=5)
        output_path = os.path.join(self.tmpdir, "test_sig.docx")

        assembler = DiscoveryAssembler(caption_page_path=CAPTION_PAGE)
        assembler.assemble(
            ds, output_path,
            attorney_name="Jane Attorney",
            firm_name="Smith & Jones LLP",
        )

        text = _extract_all_text(output_path)
        self.assertIn("SMITH & JONES LLP", text)
        self.assertIn("Jane Attorney", text)


@unittest.skipUnless(_has_caption_page(), _SKIP_MSG)
class TestAssembleSectionHeading(unittest.TestCase):
    """Section heading should appear centered and formatted correctly."""

    def setUp(self):
        self.tmpdir = tempfile.mkdtemp(prefix="test_assembler_")

    def tearDown(self):
        shutil.rmtree(self.tmpdir, ignore_errors=True)

    def test_section_heading_centered(self):
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from icharlotte_core.discovery.assembler import DiscoveryAssembler

        ds = _make_si_set(num_requests=3)
        output_path = os.path.join(self.tmpdir, "test_heading.docx")

        assembler = DiscoveryAssembler(caption_page_path=CAPTION_PAGE)
        assembler.assemble(ds, output_path, attorney_name="Jane Attorney")

        doc = Document(output_path)
        # Find the section heading paragraph
        heading_para = None
        for p in doc.paragraphs:
            if "SPECIAL INTERROGATORIES, SET ONE" in p.text:
                heading_para = p
                break

        self.assertIsNotNone(heading_para, "Section heading not found in document")
        # The style "Center Double Bold Und" handles centering — check style name
        style_name = heading_para.style.name if heading_para.style else ""
        is_centered = (
            heading_para.alignment == WD_ALIGN_PARAGRAPH.CENTER
            or "center" in style_name.lower()
        )
        self.assertTrue(is_centered, f"Heading should be centered, got style={style_name}")


if __name__ == "__main__":
    unittest.main()
