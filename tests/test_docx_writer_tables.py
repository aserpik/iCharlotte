"""Tests for the markdown-table helpers in ``icharlotte_core.docx_writer``.

The LLM emits markdown pipe-tables (``| a | b |``) in wizard-task output.
Without these helpers the script wrote each row as a literal text paragraph,
so the Output page showed raw pipes instead of a Word table. These tests
prove:

  1. ``try_consume_markdown_table`` recognises pipe-tables and ignores
     non-table content.
  2. ``add_markdown_table_to_doc`` emits a real ``w:tbl`` with the right
     header cells and rows.
  3. The ``add_markdown_to_doc`` patched into each Scripts/* module
     produces a table in the resulting .docx (round-trip via python-docx).
"""
import os
import tempfile
import unittest

from docx import Document

from icharlotte_core.docx_writer import (
    add_markdown_table_to_doc,
    try_consume_markdown_table,
)


class TestTryConsumeMarkdownTable(unittest.TestCase):
    def test_recognises_simple_pipe_table(self):
        text = (
            "| Date | Provider | Notes |\n"
            "| ---- | -------- | ----- |\n"
            "| 09/09/2019 | Dr. Mama | Norco 10/325 mg bid |\n"
            "| 10/10/2019 | Dr. Papa | Gabapentin |\n"
        )
        lines = text.split("\n")
        consumed, header, rows = try_consume_markdown_table(lines, 0)
        self.assertEqual(consumed, 4)
        self.assertEqual(header, ["Date", "Provider", "Notes"])
        self.assertEqual(rows, [
            ["09/09/2019", "Dr. Mama", "Norco 10/325 mg bid"],
            ["10/10/2019", "Dr. Papa", "Gabapentin"],
        ])

    def test_recognises_table_without_outer_pipes(self):
        lines = [
            "Date | Provider | Notes",
            "--- | --- | ---",
            "09/09/2019 | Dr. Mama | Norco",
        ]
        consumed, header, rows = try_consume_markdown_table(lines, 0)
        self.assertEqual(consumed, 3)
        self.assertEqual(header, ["Date", "Provider", "Notes"])
        self.assertEqual(rows, [["09/09/2019", "Dr. Mama", "Norco"]])

    def test_recognises_alignment_separators(self):
        lines = [
            "| Date | Provider |",
            "| :--- | ---: |",
            "| 09/09/2019 | Dr. Mama |",
        ]
        consumed, _, rows = try_consume_markdown_table(lines, 0)
        self.assertEqual(consumed, 3)
        self.assertEqual(rows, [["09/09/2019", "Dr. Mama"]])

    def test_rejects_pipe_line_with_no_separator(self):
        # No separator row → not a table; classifier must NOT consume.
        lines = [
            "| Date | Provider | Notes |",
            "Some prose follows.",
        ]
        consumed, header, rows = try_consume_markdown_table(lines, 0)
        self.assertEqual(consumed, 0)
        self.assertIsNone(header)
        self.assertIsNone(rows)

    def test_rejects_plain_text(self):
        lines = ["This is a paragraph.", "Another line."]
        consumed, header, rows = try_consume_markdown_table(lines, 0)
        self.assertEqual(consumed, 0)
        self.assertIsNone(header)

    def test_table_stops_at_blank_line(self):
        lines = [
            "| A | B |",
            "| --- | --- |",
            "| 1 | 2 |",
            "",
            "| 3 | 4 |",  # belongs to a different block
        ]
        consumed, _, rows = try_consume_markdown_table(lines, 0)
        # Header + separator + one data row = 3 consumed; blank line stops it.
        self.assertEqual(consumed, 3)
        self.assertEqual(rows, [["1", "2"]])

    def test_pads_short_rows_to_header_width(self):
        lines = [
            "| A | B | C |",
            "| --- | --- | --- |",
            "| 1 | 2 |",  # only 2 cells — should be padded
        ]
        _, header, rows = try_consume_markdown_table(lines, 0)
        self.assertEqual(len(header), 3)
        self.assertEqual(rows, [["1", "2", ""]])

    def test_trims_overlong_rows_to_header_width(self):
        lines = [
            "| A | B |",
            "| --- | --- |",
            "| 1 | 2 | 3 |",  # one extra cell — should be trimmed
        ]
        _, header, rows = try_consume_markdown_table(lines, 0)
        self.assertEqual(len(header), 2)
        self.assertEqual(rows, [["1", "2"]])


class TestAddMarkdownTableToDoc(unittest.TestCase):
    def test_emits_real_word_table_with_header_and_rows(self):
        doc = Document()
        add_markdown_table_to_doc(
            doc,
            header=["Date", "Provider", "Notes"],
            rows=[
                ["09/09/2019", "Dr. Mama", "Norco"],
                ["10/10/2019", "Dr. Papa", "Gabapentin"],
            ],
        )
        # python-docx exposes added tables on doc.tables
        self.assertEqual(len(doc.tables), 1)
        tbl = doc.tables[0]
        self.assertEqual(len(tbl.rows), 3)  # 1 header + 2 data
        self.assertEqual(len(tbl.columns), 3)
        self.assertEqual(tbl.rows[0].cells[0].text, "Date")
        self.assertEqual(tbl.rows[1].cells[1].text, "Dr. Mama")
        self.assertEqual(tbl.rows[2].cells[2].text, "Gabapentin")

    def test_header_runs_are_bold(self):
        doc = Document()
        add_markdown_table_to_doc(doc, header=["H1"], rows=[["v1"]])
        hdr_run = doc.tables[0].rows[0].cells[0].paragraphs[0].runs[0]
        self.assertTrue(hdr_run.bold)

    def test_round_trip_through_docx_save(self):
        doc = Document()
        add_markdown_table_to_doc(
            doc, header=["A", "B"], rows=[["1", "2"], ["3", "4"]]
        )
        fd, path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        try:
            doc.save(path)
            reloaded = Document(path)
            self.assertEqual(len(reloaded.tables), 1)
            self.assertEqual(reloaded.tables[0].rows[1].cells[0].text, "1")
        finally:
            try:
                os.remove(path)
            except OSError:
                pass


class TestScriptAddMarkdownToDocIntegration(unittest.TestCase):
    """End-to-end: each script's ``add_markdown_to_doc`` must turn a pipe-table
    block in the LLM output into a real Word table, not literal text paragraphs.
    """

    MARKDOWN = (
        "## Medications\n"
        "\n"
        "| Date | Provider | Notes |\n"
        "| ---- | -------- | ----- |\n"
        "| 09/09/2019 | Dr. Mama | Norco 10/325 mg bid |\n"
        "| 10/10/2019 | Dr. Papa | Gabapentin |\n"
        "\n"
        "Follow-up scheduled.\n"
    )

    def _assert_renders_table(self, module_name: str, attr: str = "add_markdown_to_doc"):
        import importlib
        mod = importlib.import_module(module_name)
        func = getattr(mod, attr)
        doc = Document()
        func(doc, self.MARKDOWN)
        self.assertEqual(
            len(doc.tables), 1,
            f"{module_name}.{attr} did not emit a Word table for a pipe-table block",
        )
        tbl = doc.tables[0]
        self.assertEqual(tbl.rows[0].cells[0].text, "Date")
        self.assertEqual(tbl.rows[1].cells[1].text, "Dr. Mama")

    def test_summarize(self):
        self._assert_renders_table("Scripts.summarize")

    def test_summarize_discovery(self):
        self._assert_renders_table("Scripts.summarize_discovery")

    def test_summarize_deposition(self):
        self._assert_renders_table("Scripts.summarize_deposition")

    def test_med_record(self):
        self._assert_renders_table("Scripts.med_record")

    def test_med_chron(self):
        self._assert_renders_table("Scripts.med_chron")

    def test_liability(self):
        self._assert_renders_table("Scripts.liability")

    def test_exposure(self):
        self._assert_renders_table("Scripts.exposure")

    def test_docx_writer_internal(self):
        # The private helper inside docx_writer is what safe_append_to_docx uses
        # as its default ``format_func``. Confirm it also gets tables right.
        from icharlotte_core.docx_writer import _add_markdown_to_doc
        doc = Document()
        _add_markdown_to_doc(doc, self.MARKDOWN)
        self.assertEqual(len(doc.tables), 1)


if __name__ == "__main__":
    unittest.main()
