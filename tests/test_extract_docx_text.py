"""Tests for extract_docx_text — table-aware .docx extraction.

python-docx's ``doc.paragraphs`` silently skips tables, headers/footers, and
nested content. Legal chronological summaries and medical timelines live
entirely inside tables, so missing them means the LLM sees only the title.
These tests prove the helper recovers every kind of content.
"""
import os
import tempfile
import unittest

from docx import Document

from icharlotte_core.document_processor import (
    DocumentProcessor,
    extract_docx_text,
)


def _make_docx(build_fn) -> str:
    """Write a fresh .docx to a temp file, apply ``build_fn(doc)``, return path."""
    doc = Document()
    build_fn(doc)
    fd, path = tempfile.mkstemp(suffix=".docx")
    os.close(fd)
    doc.save(path)
    return path


class TestExtractDocxText(unittest.TestCase):

    def _cleanup(self, path):
        try:
            os.remove(path)
        except OSError:
            pass

    def test_paragraphs_only(self):
        def build(doc):
            doc.add_paragraph("First paragraph.")
            doc.add_paragraph("Second paragraph.")
        path = _make_docx(build)
        self.addCleanup(self._cleanup, path)

        text = extract_docx_text(path)
        self.assertIn("First paragraph.", text)
        self.assertIn("Second paragraph.", text)

    def test_table_cells_are_extracted(self):
        """Regression test: old extractor missed ALL table content."""
        def build(doc):
            doc.add_paragraph("Summary title")
            t = doc.add_table(rows=3, cols=3)
            t.rows[0].cells[0].text = "Date"
            t.rows[0].cells[1].text = "Facility"
            t.rows[0].cells[2].text = "Description"
            t.rows[1].cells[0].text = "01/02/2026"
            t.rows[1].cells[1].text = "Barstow Hosp"
            t.rows[1].cells[2].text = "Chest pain"
            t.rows[2].cells[0].text = "02/03/2026"
            t.rows[2].cells[1].text = "Pain Mgmt"
            t.rows[2].cells[2].text = "Prescribed oxycodone 10mg"
        path = _make_docx(build)
        self.addCleanup(self._cleanup, path)

        text = extract_docx_text(path)
        # Title is present
        self.assertIn("Summary title", text)
        # Every cell is extracted
        self.assertIn("Date", text)
        self.assertIn("Facility", text)
        self.assertIn("Description", text)
        self.assertIn("Barstow Hosp", text)
        self.assertIn("Chest pain", text)
        self.assertIn("oxycodone 10mg", text)

    def test_table_rows_use_pipe_separator(self):
        def build(doc):
            t = doc.add_table(rows=1, cols=3)
            t.rows[0].cells[0].text = "A"
            t.rows[0].cells[1].text = "B"
            t.rows[0].cells[2].text = "C"
        path = _make_docx(build)
        self.addCleanup(self._cleanup, path)

        text = extract_docx_text(path)
        self.assertIn("A | B | C", text)

    def test_preserves_document_order(self):
        """Paragraphs and tables are emitted in the order they appear."""
        def build(doc):
            doc.add_paragraph("Intro paragraph.")
            t = doc.add_table(rows=1, cols=1)
            t.rows[0].cells[0].text = "Middle table cell."
            doc.add_paragraph("Trailing paragraph.")
        path = _make_docx(build)
        self.addCleanup(self._cleanup, path)

        text = extract_docx_text(path)
        intro = text.index("Intro paragraph.")
        middle = text.index("Middle table cell.")
        trailing = text.index("Trailing paragraph.")
        self.assertLess(intro, middle)
        self.assertLess(middle, trailing)

    def test_nested_paragraphs_inside_cells(self):
        """A cell with multiple paragraphs should join them into one cell."""
        def build(doc):
            t = doc.add_table(rows=1, cols=1)
            cell = t.rows[0].cells[0]
            cell.paragraphs[0].text = "Line one"
            cell.add_paragraph("Line two")
            cell.add_paragraph("Line three")
        path = _make_docx(build)
        self.addCleanup(self._cleanup, path)

        text = extract_docx_text(path)
        self.assertIn("Line one", text)
        self.assertIn("Line two", text)
        self.assertIn("Line three", text)

    def test_header_is_extracted(self):
        def build(doc):
            section = doc.sections[0]
            section.header.paragraphs[0].text = "FORNEY v. BARSTOW"
            doc.add_paragraph("Body content")
        path = _make_docx(build)
        self.addCleanup(self._cleanup, path)

        text = extract_docx_text(path)
        self.assertIn("[HEADER]", text)
        self.assertIn("FORNEY v. BARSTOW", text)
        # Header comes before body
        self.assertLess(text.index("FORNEY v. BARSTOW"), text.index("Body content"))

    def test_include_headers_false_skips_header(self):
        def build(doc):
            doc.sections[0].header.paragraphs[0].text = "HEADER TEXT"
            doc.add_paragraph("Body content")
        path = _make_docx(build)
        self.addCleanup(self._cleanup, path)

        text = extract_docx_text(path, include_headers=False)
        self.assertNotIn("[HEADER]", text)
        self.assertNotIn("HEADER TEXT", text)
        self.assertIn("Body content", text)

    def test_large_table_all_rows_extracted(self):
        """Stand-in for the 219-row Forney chronology."""
        ROW_COUNT = 50
        def build(doc):
            t = doc.add_table(rows=ROW_COUNT, cols=2)
            for i in range(ROW_COUNT):
                t.rows[i].cells[0].text = f"Date-{i}"
                t.rows[i].cells[1].text = f"Event-{i}"
        path = _make_docx(build)
        self.addCleanup(self._cleanup, path)

        text = extract_docx_text(path)
        for i in range(ROW_COUNT):
            self.assertIn(f"Date-{i}", text, f"Row {i} date missing")
            self.assertIn(f"Event-{i}", text, f"Row {i} event missing")

    def test_document_processor_uses_table_extractor(self):
        """DocumentProcessor._extract_from_docx now includes table content."""
        def build(doc):
            doc.add_paragraph("Title only")
            t = doc.add_table(rows=1, cols=2)
            t.rows[0].cells[0].text = "key"
            t.rows[0].cells[1].text = "value"
        path = _make_docx(build)
        self.addCleanup(self._cleanup, path)

        result = DocumentProcessor().extract_text(path)
        self.assertTrue(result.success)
        self.assertIn("Title only", result.text)
        self.assertIn("key", result.text)
        self.assertIn("value", result.text)


if __name__ == "__main__":
    unittest.main()
