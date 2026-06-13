import sys
import os
import argparse
import ctypes
import copy

# Workaround: WMI queries hang indefinitely on some Windows machines.
# platform.system(), platform.platform(), etc. all call _wmi_query internally.
# Pre-populate the caches and disable _wmi_query before any library imports.
import platform as _platform_mod
_platform_mod._uname_cache = _platform_mod.uname_result(
    system='Windows',
    node=os.environ.get('COMPUTERNAME', ''),
    release='11',
    version='10.0.26200',
    machine=os.environ.get('PROCESSOR_ARCHITECTURE', 'AMD64'),
)
_platform_mod.win32_ver = lambda *a, **kw: ('11', '10.0.26200', '', 'Multiprocessor Free')
_orig_wmi_query = getattr(_platform_mod, '_wmi_query', None)
_platform_mod._wmi_query = lambda *a, **kw: (_ for _ in ()).throw(OSError('WMI disabled'))

# Set Windows App User Model ID BEFORE any Qt imports for proper taskbar icon
try:
    ctypes.windll.shell32.SetCurrentProcessExplicitAppUserModelID('iCharlotte.LegalSuite.1')
except Exception:
    pass

# Load environment variables from .env file
from dotenv import load_dotenv
load_dotenv()

# MUST register schemes before QApplication is created
from icharlotte_core.bridge import register_custom_schemes
register_custom_schemes()

import re
import glob
import json
import uuid
import subprocess
import datetime
import time
from functools import partial

# --- Imports ---
try:
    from PySide6.QtWidgets import (
        QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
        QPushButton, QTreeWidgetItem, QHeaderView, QMessageBox, QLabel,
        QFrame, QSplitter, QAbstractItemView, QLineEdit,
        QTreeWidgetItemIterator, QTabWidget, QScrollArea, QMenu, QDialog,
        QFileIconProvider, QToolButton, QGroupBox, QCheckBox, QComboBox,
        QInputDialog
    )
    from PySide6.QtCore import Qt, QThread, Signal, QFileInfo, QMetaObject, Q_ARG, QSettings, QTimer, QThreadPool, QRunnable
    from PySide6.QtGui import QAction, QShortcut, QKeySequence, QIcon, QCursor
    from PySide6.QtWebEngineCore import QWebEngineUrlScheme
except ImportError:
    print("Error: PySide6 or its components are not installed. Please run: pip install PySide6 PySide6-WebEngine")
    sys.exit(1)

# Global hotkey support
try:
    import keyboard
    KEYBOARD_AVAILABLE = True
except (ImportError, OSError):
    KEYBOARD_AVAILABLE = False
    print("Warning: 'keyboard' library not available. Global hotkeys (Win+F, Win+C) disabled.")

# --- Core Modules ---
from icharlotte_core.config import SCRIPTS_DIR, GEMINI_DATA_DIR, BASE_PATH_WIN
from icharlotte_core.gc_guard import no_gc
from icharlotte_core.utils import (
    log_event, get_case_path, sanitize_filename, format_date_to_mm_dd_yyyy
)
from icharlotte_core.ui.widgets import (
    StatusWidget, AgentRunner, FileTreeWidget
)
from icharlotte_core.ui.case_view_enhanced import (
    EnhancedAgentButton, AgentSettingsDB, AgentSettingsDialog,
    AdvancedFilterWidget, FilePreviewWidget, OutputBrowserWidget,
    ProcessingLogWidget, ProcessingLogDB, FileTagsDB, EnhancedFileTreeWidget
)
from icharlotte_core.ui import theme
from icharlotte_core.ui.dialogs import FileNumberDialog, VariablesDialog, PromptsDialog
from icharlotte_core.ui.report_generator_dialog import ReportGeneratorDialog, ReportPipelineWorker
from icharlotte_core.subpoena_tracker import SubpoenaTrackerWorker
from icharlotte_core.ui.tabs import ChatTab, IndexTab
from icharlotte_core.ui.master_case_tab import MasterCaseTab
from icharlotte_core.master_db import MasterCaseDatabase
from icharlotte_core.ui.templates_resources_tab import TemplatesResourcesTab
from icharlotte_core.ui.discovery_tab import DiscoveryTab
from icharlotte_core.ui.wizard.summary_browser import SummaryBrowserTab
from icharlotte_core.ui.wizard.summary_outputs import (
    summary_browser_title,
    task_id_for_summary_action,
)
from icharlotte_core.word_hotkey import init_word_hotkey, stop_word_hotkey
from icharlotte_core.ui.zoom_handler import ZoomEventFilter
from icharlotte_core.app_crash_handler import (
    install_crash_handler, checkpoint, add_context,
    log_info, log_warning, log_error, log_debug, safe_slot
)

class QuickOpenDialog(QDialog):
    """Lightweight popup dialog for quick file number or plaintiff name entry via double-Ctrl tap."""

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Quick Open Case Folder")
        self.setFixedSize(400, 100)
        self.setWindowFlags(Qt.WindowType.Dialog | Qt.WindowType.FramelessWindowHint | Qt.WindowType.WindowStaysOnTopHint)
        self.setAttribute(Qt.WidgetAttribute.WA_TranslucentBackground, False)

        # Load recent cases for autocomplete
        self.recent_file = os.path.join(GEMINI_DATA_DIR, "recent_cases.json")
        self.recent_cases = self._load_recent_cases()

        # Load all cases from database for plaintiff name lookup
        self.db = MasterCaseDatabase()
        self.all_cases = self.db.get_all_cases()
        # Build a mapping of plaintiff names to file numbers (lowercase for matching)
        self.plaintiff_to_file = {}
        for case in self.all_cases:
            name = case.get('plaintiff_last_name', '').strip()
            file_num = case.get('file_number', '')
            if name and file_num:
                self.plaintiff_to_file[name.lower()] = file_num

        layout = QVBoxLayout(self)
        layout.setContentsMargins(15, 15, 15, 15)

        # Style the dialog (light/white theme)
        self.setStyleSheet(f"""
            QDialog {{
                background-color: {theme.BG};
                border: 2px solid {theme.PRIMARY};
                border-radius: 10px;
            }}
            QComboBox {{
                background-color: {theme.BG_SUBTLE};
                color: {theme.TEXT_BODY};
                border: 1px solid {theme.BORDER};
                border-radius: 5px;
                padding: 8px 12px;
                font-size: 16px;
                min-height: 30px;
            }}
            QComboBox:focus {{
                border: 2px solid {theme.PRIMARY};
                background-color: {theme.BG};
            }}
            QComboBox::drop-down {{
                border: none;
                width: 20px;
            }}
            QComboBox QAbstractItemView {{
                background-color: {theme.BG};
                color: {theme.TEXT_BODY};
                selection-background-color: {theme.PRIMARY};
                selection-color: #fff;
            }}
            QLabel {{
                color: {theme.TEXT_MUTED};
                font-size: 11px;
            }}
        """)

        # File number input with autocomplete (includes both file numbers and plaintiff names)
        self.file_input = QComboBox()
        self.file_input.setEditable(True)
        self.file_input.lineEdit().setPlaceholderText("Enter file number or plaintiff name")

        # Build autocomplete list: recent cases first, then all cases with names
        autocomplete_items = list(self.recent_cases)  # Recent file numbers
        for case in self.all_cases:
            file_num = case.get('file_number', '')
            name = case.get('plaintiff_last_name', '').strip()
            if name and file_num:
                # Add "Name (####.###)" format for easy selection
                display = f"{name} ({file_num})"
                if display not in autocomplete_items and file_num not in autocomplete_items:
                    autocomplete_items.append(display)

        self.file_input.addItems(autocomplete_items)
        self.file_input.setCurrentIndex(-1)  # Start empty
        self.file_input.lineEdit().returnPressed.connect(self._on_enter)
        layout.addWidget(self.file_input)

        # Help text
        help_label = QLabel("Press Enter to open folder • Esc to close")
        help_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(help_label)

        self.result_file_number = None

    def _load_recent_cases(self):
        if os.path.exists(self.recent_file):
            try:
                with open(self.recent_file, 'r') as f:
                    return json.load(f)
            except Exception as e:
                log_event(f"Failed to load recent cases {self.recent_file}: {e}", "error")
                return []
        return []

    def _save_recent_case(self, file_num):
        if not file_num:
            return
        if file_num in self.recent_cases:
            self.recent_cases.remove(file_num)
        self.recent_cases.insert(0, file_num)
        self.recent_cases = self.recent_cases[:10]

        if not os.path.exists(GEMINI_DATA_DIR):
            os.makedirs(GEMINI_DATA_DIR)

        try:
            with open(self.recent_file, 'w') as f:
                json.dump(self.recent_cases, f)
        except Exception as e:
            log_event(f"Failed to save recent cases {self.recent_file}: {e}", "error")

    def _resolve_to_file_number(self, text):
        """Resolve user input to a file number.

        Handles:
        - Direct file number (####.###)
        - "Name (####.###)" format from autocomplete
        - Plaintiff name lookup
        """
        text = text.strip()
        if not text:
            return None

        # Check if it's the "Name (####.###)" format from autocomplete
        if '(' in text and text.endswith(')'):
            # Extract file number from parentheses
            start = text.rfind('(')
            file_num = text[start + 1:-1].strip()
            if re.match(r'^\d{4}[.\-]\d{3}$', file_num):
                return file_num.replace('-', '.')

        # Check if it's a file number format (####.### or ####-###)
        if re.match(r'^\d{4}[.\-]\d{3}$', text):
            return text.replace('-', '.')

        # Otherwise, try to find by plaintiff name
        text_lower = text.lower()

        # Exact match first
        if text_lower in self.plaintiff_to_file:
            return self.plaintiff_to_file[text_lower]

        # Partial match (find first case where name contains the search text)
        for name, file_num in self.plaintiff_to_file.items():
            if text_lower in name:
                return file_num

        # No match found - return input as-is (will fail gracefully in caller)
        return text

    def _on_enter(self):
        user_input = self.file_input.currentText().strip()
        if user_input:
            file_num = self._resolve_to_file_number(user_input)
            self.result_file_number = file_num
            self._save_recent_case(file_num)
            self.accept()

    def keyPressEvent(self, event):
        if event.key() == Qt.Key.Key_Escape:
            self.reject()
        else:
            super().keyPressEvent(event)

    def showEvent(self, event):
        super().showEvent(event)
        # Center on the screen where the mouse cursor is (for multi-monitor support)
        cursor_pos = QCursor.pos()
        screen = QApplication.screenAt(cursor_pos)
        if screen is None:
            screen = QApplication.primaryScreen()
        screen_geom = screen.geometry()
        self.move(
            screen_geom.x() + (screen_geom.width() - self.width()) // 2,
            screen_geom.y() + screen_geom.height() // 3
        )
        # Use a timer to ensure focus is set after window is fully rendered
        # This is necessary for multi-monitor setups where focus can be unreliable
        # 100ms delay needed on Windows for reliable focus stealing
        QTimer.singleShot(100, self._ensure_focus)

    def _ensure_focus(self):
        """Ensure the input field has focus after the window is shown."""
        # On Windows, we need to use native APIs to force foreground focus
        if sys.platform == 'win32':
            try:
                import ctypes
                hwnd = int(self.winId())
                # Attach to the foreground window's thread to get permission to set foreground
                user32 = ctypes.windll.user32
                foreground_hwnd = user32.GetForegroundWindow()
                foreground_thread = user32.GetWindowThreadProcessId(foreground_hwnd, None)
                current_thread = ctypes.windll.kernel32.GetCurrentThreadId()
                # Attach threads to allow focus stealing
                user32.AttachThreadInput(foreground_thread, current_thread, True)
                user32.SetForegroundWindow(hwnd)
                user32.AttachThreadInput(foreground_thread, current_thread, False)
            except Exception:
                pass  # Fall back to Qt methods if native approach fails

        self.activateWindow()
        self.raise_()
        self.file_input.setFocus()
        self.file_input.lineEdit().setFocus()
        self.file_input.lineEdit().selectAll()


class DirectoryTreeWorker(QThread):
    data_ready = Signal(list) # Emits (root, dirs, files) tuples
    finished = Signal()

    def __init__(self, root_path):
        super().__init__()
        self.root_path = root_path
        self.running = True
        
    def run(self):
        import time as _time
        _start = _time.monotonic()
        _batch_count = 0
        _total_dirs = 0
        _total_files = 0
        log_debug(f"DirectoryTreeWorker: scanning {self.root_path}")
        try:
            batch = []
            for root, dirs, files in os.walk(self.root_path):
                if not self.running:
                    log_debug(f"DirectoryTreeWorker: stopped early after {_time.monotonic()-_start:.1f}s")
                    break
                # Skip hidden files/dirs
                dirs[:] = [d for d in dirs if not d.startswith('.') and not d.startswith('~$')]

                file_data = []
                for f in files:
                    if not self.running:
                        break
                    if f.startswith('.') or f.startswith('~$'):
                        continue
                    path = os.path.join(root, f)
                    try:
                        stat = os.stat(path)
                        size = stat.st_size
                        mtime = stat.st_mtime
                        file_data.append((f, size, mtime))
                    except Exception:
                        file_data.append((f, 0, 0))

                _total_dirs += len(dirs)
                _total_files += len(file_data)
                batch.append((root, dirs, file_data))
                if len(batch) >= 10:
                    _batch_count += 1
                    self.data_ready.emit(batch)
                    batch = []
                    self.msleep(10) # Yield to UI

            if batch:
                _batch_count += 1
                self.data_ready.emit(batch)
            log_debug(f"DirectoryTreeWorker: done in {_time.monotonic()-_start:.1f}s, {_batch_count} batches, {_total_dirs} dirs, {_total_files} files")
            self.finished.emit()
        except Exception as e:
            log_error(f"DirectoryTreeWorker CRASHED after {_time.monotonic()-_start:.1f}s: {e}", exc_info=True)
            self.finished.emit()

    def stop(self):
        self.running = False


class _LibraryCaptureJob(QRunnable):
    """Best-effort, off-UI-thread capture of a finished task's source text
    into the document library. Never raises (capture_from_task_entry swallows).

    Emits ``done_signal(case_root)`` after capture so any open Chat tab can
    refresh its Saved Documents list without the user clicking Refresh.
    """

    def __init__(self, case_root, entry, done_signal=None):
        super().__init__()
        self._case_root, self._entry = case_root, entry
        self._done_signal = done_signal

    def run(self):
        from icharlotte_core.doc_library.capture import capture_from_task_entry
        try:
            capture_from_task_entry(self._case_root, self._entry)
        finally:
            if self._done_signal is not None:
                try:
                    self._done_signal.emit(self._case_root)
                except Exception:
                    pass


class MainWindow(QMainWindow):
    # Signals for thread-safe hotkey callbacks
    open_file_signal = Signal()
    change_file_signal = Signal()
    quick_open_signal = Signal()  # For double-Ctrl quick open
    ctrl_press_signal = Signal()  # For thread-safe Ctrl press handling
    ctrl_release_signal = Signal()  # For thread-safe Ctrl release handling
    library_captured = Signal(str)  # case_root; emitted after a task's source text is captured into the doc library

    def __init__(self, file_number=None, case_path=None, initial_tab=None):
        super().__init__()
        checkpoint("MainWindow.__init__ starting", file_number=file_number)

        # Connect signals for global hotkeys (thread-safe)
        self.open_file_signal.connect(self._on_open_file_hotkey)
        self.change_file_signal.connect(self._on_change_file_hotkey)
        self.quick_open_signal.connect(self._on_quick_open_hotkey)
        self.ctrl_press_signal.connect(self._handle_ctrl_press)
        self.ctrl_release_signal.connect(self._handle_ctrl_release)

        # Double-Ctrl tap detection state
        self._last_ctrl_tap_time = None  # Time of last valid tap (quick press+release), None = no recent tap
        self._ctrl_press_time = 0  # When Ctrl was pressed
        self._ctrl_tap_threshold = 0.3  # Max seconds between taps for double-tap (tightened to avoid accidental triggers)
        self._ctrl_hold_threshold = 0.2  # Max seconds Ctrl can be held to count as tap
        self.file_number = file_number
        self.case_path = case_path
        add_context('current_file_number', file_number)
        add_context('current_case_path', case_path)
        self._update_window_title()
        self.resize(1200, 800)

        # Set application icon
        icon_path = os.path.join(os.path.dirname(__file__), 'icharlotte.ico')
        if os.path.exists(icon_path):
            self.setWindowIcon(QIcon(icon_path))

        # Set AppUserModelID on window handle for proper taskbar grouping
        try:
            from ctypes import wintypes
            hwnd = int(self.winId())
            shell32 = ctypes.windll.shell32
            propvariant = ctypes.create_unicode_buffer('iCharlotte.LegalSuite.1')
            shell32.SHGetPropertyStoreForWindow.argtypes = [wintypes.HWND, ctypes.POINTER(ctypes.c_void_p), ctypes.POINTER(ctypes.c_void_p)]
            shell32.SHGetPropertyStoreForWindow.restype = ctypes.HRESULT
        except Exception:
            pass

        self.agent_runners = [] # Keep references to prevent GC
        self._agent_queue = []  # Queued agents waiting to run (when concurrency limit reached)
        self.MAX_CONCURRENT_AGENTS = 4  # Max subprocess agents running simultaneously
        self._tree_generation = 0  # Incremented on each populate_tree to ignore stale worker callbacks
        self._populating_tree = False  # Re-entrancy guard for populate_tree
        self._tree_refresh_timer = QTimer()
        self._tree_refresh_timer.setSingleShot(True)
        self._tree_refresh_timer.timeout.connect(self.populate_tree)
        self.cached_models = {} # Cache for models: {provider: [list]}
        self.fetcher = None

        log_info(f"Initializing MainWindow for {file_number} at {case_path}")
        self.icon_provider = QFileIconProvider()
        # Cache icons to avoid slow network file access
        self._icon_cache = {}
        self._folder_icon = self.icon_provider.icon(QFileIconProvider.IconType.Folder)
        self._file_icon = self.icon_provider.icon(QFileIconProvider.IconType.File)
        # Wizard mode controller (global Advanced/Wizard toggle).
        from icharlotte_core.ui.wizard.mode_controller import ModeController
        self.mode_controller = ModeController(parent=self)
        self.setup_ui()
        self._restore_window_state()

        # Apply current mode and react to future mode changes.
        self.mode_controller.mode_changed.connect(self._apply_mode_visibility)
        self._apply_mode_visibility(self.mode_controller.mode)

        # Restore tab: explicit --tab arg wins, else last session's active tab.
        if initial_tab is not None and 0 <= initial_tab < self.tabs.count():
            self.tabs.setCurrentIndex(initial_tab)
        else:
            self._restore_last_tab()

        # Only populate tree and check docket if a case is loaded
        if self.case_path:
            # A loaded case always opens in Wizard mode on the Wizard tab,
            # overriding any restored initial_tab above.
            self._default_to_wizard_mode()
            self.populate_tree()
            self.load_status_history()
            # Restore wizard task tabs for the initial case (if any).
            try:
                self._restore_task_tabs_for_case()
            except Exception as e:
                log_event(f"[wizard] startup restore failed: {e}")
            # Refresh Wizard tab's Recent Tasks list for the startup case.
            if hasattr(self, "wizard_tab") and self.wizard_tab is not None:
                from icharlotte_core.ui.wizard.persistence import WizardStatePersistence
                try:
                    p = WizardStatePersistence(self.case_path)
                    self.wizard_tab.refresh_recent_tasks(p.get_recent_tasks())
                except Exception as e:
                    log_event(f"[wizard] startup refresh recent_tasks failed: {e}")

        # Register global hotkeys (Win+F for Open File, Win+C for Change File)
        self._setup_global_hotkeys()

    def _setup_global_hotkeys(self):
        """Register global hotkeys for Open File (Win+F), Change File (Win+C), double-Ctrl quick open, and Win+V for Word AI."""
        if not KEYBOARD_AVAILABLE:
            return

        try:
            # Register Win+F for Open File
            keyboard.add_hotkey('win+f', lambda: self.open_file_signal.emit(), suppress=True)
            # Register Win+C for Change File
            keyboard.add_hotkey('win+c', lambda: self.change_file_signal.emit(), suppress=True)
            # Register Ctrl key press and release for double-tap detection
            keyboard.on_press_key('ctrl', self._on_ctrl_press)
            keyboard.on_release_key('ctrl', self._on_ctrl_release)
            log_event("Global hotkeys registered: Win+F (Open File), Win+C (Change File), Double-Ctrl (Quick Open)")
        except Exception as e:
            log_event(f"Failed to register global hotkeys: {e}", "error")

        # Register Word AI hotkey (Win+V)
        try:
            if init_word_hotkey(self):
                log_event("Word AI hotkey registered: Win+V")
        except Exception as e:
            log_event(f"Failed to register Word AI hotkey: {e}", "error")

    def _on_ctrl_press(self, event):
        """Called from keyboard library thread - emit signal for thread safety."""
        self.ctrl_press_signal.emit()

    def _on_ctrl_release(self, event):
        """Called from keyboard library thread - emit signal for thread safety."""
        self.ctrl_release_signal.emit()

    def _handle_ctrl_press(self):
        """Handle Ctrl press on main thread - record time for tap duration calculation."""
        self._ctrl_press_time = time.time()

    def _handle_ctrl_release(self):
        """Handle Ctrl release on main thread - detect double-tap for quick open dialog.

        Only triggers if:
        1. Ctrl was held briefly (< 200ms) - this is a "tap", not a held key
        2. Two taps occurred within 500ms of each other
        """
        current_time = time.time()
        hold_duration = current_time - self._ctrl_press_time

        # Only count as a tap if Ctrl was held briefly (not held for shortcuts)
        if hold_duration > self._ctrl_hold_threshold:
            # Ctrl was held too long - this was probably a shortcut, not a tap
            self._last_ctrl_tap_time = None
            return

        # This is a valid tap - check if it's a double-tap
        if self._last_ctrl_tap_time is not None:
            time_since_last_tap = current_time - self._last_ctrl_tap_time
            if time_since_last_tap < self._ctrl_tap_threshold:
                # Double-tap detected - emit signal (thread-safe)
                self.quick_open_signal.emit()
                self._last_ctrl_tap_time = None  # Reset to prevent triple-tap triggering
                return

        # Record this tap for potential double-tap detection
        self._last_ctrl_tap_time = current_time

    def _on_open_file_hotkey(self):
        """Handle Win+F hotkey - bring window to front and open file."""
        self.activateWindow()
        self.raise_()
        self.showNormal()
        self.open_root_folder()

    def _on_change_file_hotkey(self):
        """Handle Win+C hotkey - bring window to front and change file."""
        self.activateWindow()
        self.raise_()
        self.showNormal()
        self.change_file()

    def _on_quick_open_hotkey(self):
        """Handle double-Ctrl hotkey - show quick open dialog to open case folder."""
        dialog = QuickOpenDialog(None)  # No parent for independent focus
        dialog.setWindowFlag(Qt.WindowType.WindowStaysOnTopHint, True)
        if dialog.exec() == QDialog.DialogCode.Accepted and dialog.result_file_number:
            file_num = dialog.result_file_number
            case_path = get_case_path(file_num)
            if case_path and os.path.exists(case_path):
                try:
                    os.startfile(case_path)
                    log_event(f"Quick open: Opened folder for {file_num}")
                except Exception as e:
                    QMessageBox.warning(self, "Error", f"Could not open folder: {e}")
            else:
                QMessageBox.warning(self, "Error", f"Folder not found for {file_num}")

    def _update_window_title(self):
        """Update window title based on current file_number and case_path."""
        if self.file_number and self.case_path:
            self.setWindowTitle(f"iCharlotte - {self.file_number} - {os.path.basename(self.case_path)}")
        else:
            self.setWindowTitle("iCharlotte")

    def setup_ui(self):
        checkpoint("setup_ui starting - creating tabs")
        self.tabs = QTabWidget()
        self.setCentralWidget(self.tabs)
        self.tabs.setTabsClosable(True)
        self.tabs.tabCloseRequested.connect(self._on_tab_close_requested)
        # Refresh open Chat tabs' Saved Documents when a background capture lands.
        self.library_captured.connect(self._on_library_captured)
        # Hide close buttons on the fixed Master List / Wizard / Advanced tabs.
        # (We'll re-hide after every addTab via _hide_fixed_close_buttons.)

        # Make tabs expand to fill available width (no scroll arrows needed)
        self.tabs.tabBar().setExpanding(True)
        self.tabs.setUsesScrollButtons(False)

        # Larger variant of the global underline tab treatment (theme.py).
        # Scoped to this widget so nested tab widgets keep the compact default.
        self.tabs.setStyleSheet(f"""
            QTabWidget::pane {{
                border: none;
                border-top: 1px solid {theme.BORDER};
            }}
            QTabBar::tab {{
                background: transparent;
                color: {theme.TEXT_MUTED};
                font-size: 13px;
                font-weight: 500;
                padding: 10px 6px;
                border: none;
                border-bottom: 3px solid transparent;
            }}
            QTabBar::tab:selected {{
                color: {theme.PRIMARY};
                font-weight: 600;
                border-bottom: 3px solid {theme.PRIMARY};
            }}
            QTabBar::tab:hover:!selected {{
                color: {theme.TEXT};
                background-color: {theme.BG_SUBTLE};
            }}
        """)

        # --- Tab 0: Master List ---
        checkpoint("Creating MasterCaseTab")
        self.master_tab = MasterCaseTab(self, mode_controller=self.mode_controller)
        self.tabs.addTab(self.master_tab, "Master List")

        # --- Tab 1 (Wizard Mode only): Wizard ---
        from icharlotte_core.ui.wizard.wizard_tab import WizardTab
        self.wizard_tab = WizardTab(self)
        self.tabs.addTab(self.wizard_tab, "Wizard")
        self.wizard_tab.task_requested.connect(self._open_task_tab)
        self.wizard_tab.reopen_requested.connect(self._on_reopen_recent_task)
        self.wizard_tab.card_action_requested.connect(self._on_card_action)

        # --- Tab 2: Case View ---
        case_view_widget = QWidget()
        self.tabs.addTab(case_view_widget, "Case View")

        main_layout = QVBoxLayout(case_view_widget)

        # Top Toolbar
        toolbar_layout = QHBoxLayout()

        btn_view_docket = theme.secondary_button("ViewDocket")
        btn_view_docket.clicked.connect(self.view_docket)
        toolbar_layout.addWidget(btn_view_docket)

        btn_notes = theme.secondary_button("Notes")
        btn_notes.clicked.connect(self.open_notes)
        toolbar_layout.addWidget(btn_notes)

        btn_vars = theme.secondary_button("Variables")
        btn_vars.clicked.connect(self.manage_variables)
        toolbar_layout.addWidget(btn_vars)

        # Output Browser Button
        btn_outputs = theme.secondary_button("Output Browser")
        btn_outputs.clicked.connect(self.open_output_browser)
        toolbar_layout.addWidget(btn_outputs)

        # Processing Log Button
        btn_proc_log = theme.secondary_button("Processing Log")
        btn_proc_log.clicked.connect(self.open_processing_log)
        toolbar_layout.addWidget(btn_proc_log)

        btn_generate_report = theme.primary_button("Generate Report")
        btn_generate_report.clicked.connect(self.open_report_generator)
        toolbar_layout.addWidget(btn_generate_report)

        toolbar_layout.addStretch()

        # Wrapper for vertical layout of Case View
        wrapper_layout = QVBoxLayout()
        wrapper_layout.addLayout(toolbar_layout)

        # Main horizontal splitter (agents | tree | preview)
        self.main_splitter = QSplitter(Qt.Orientation.Horizontal)
        wrapper_layout.addWidget(self.main_splitter)
        self.main_splitter.splitterMoved.connect(self._on_splitter_moved)

        main_layout.addLayout(wrapper_layout)

        # Initialize agent settings database
        self.agent_settings_db = AgentSettingsDB()
        self.agent_buttons = {}  # Track enhanced agent buttons
        self.running_agents = {}  # Track which agents are running {script: file_number}

        # Left Panel (Case Agents)
        left_panel = QFrame()
        left_panel.setFrameShape(QFrame.Shape.StyledPanel)
        left_panel.setFixedWidth(180)
        left_layout = QVBoxLayout(left_panel)
        left_layout.setAlignment(Qt.AlignmentFlag.AlignTop)
        left_layout.setContentsMargins(
            theme.SPACE_SM, theme.SPACE_SM, theme.SPACE_SM, theme.SPACE_SM
        )
        left_layout.setSpacing(theme.SPACE_XS + 2)

        title_label = QLabel("Case Agents")
        title_label.setStyleSheet(
            f"font-weight: 600; font-size: {theme.FONT_H3}px; color: {theme.TEXT}; margin-bottom: 5px;"
        )
        left_layout.addWidget(title_label)

        # Enhanced Agent Buttons with running indicator, status, and settings
        self.create_enhanced_agent_button("Docket Agent", "docket.py", left_layout, arg_type="file_number")
        self.create_enhanced_agent_button("Complaint Agent", "complaint.py", left_layout, arg_type="file_number")
        self.create_enhanced_agent_button("Report Agent", "report.py", left_layout, arg_type="file_number")
        # Subpoena Tracker — in-process QThread worker (not a subprocess agent)
        self._subpoena_btn = EnhancedAgentButton("Subpoena Tracker", "subpoena_tracker")
        self._subpoena_btn.clicked.connect(self._run_subpoena_tracker)
        self.agent_buttons["subpoena_tracker"] = self._subpoena_btn
        left_layout.addWidget(self._subpoena_btn)
        # Med Record Extractor — in-process QThread worker
        self._med_extract_btn = EnhancedAgentButton("Med Record Extractor", "med_record_extractor")
        self._med_extract_btn.clicked.connect(self._run_med_record_extractor)
        self.agent_buttons["med_record_extractor"] = self._med_extract_btn
        left_layout.addWidget(self._med_extract_btn)

        # Document Agents
        left_layout.addSpacing(8)
        new_label = QLabel("Document Agents")
        new_label.setStyleSheet(
            f"font-weight: 600; font-size: {theme.FONT_BODY}px; color: {theme.TEXT_MUTED};"
        )
        left_layout.addWidget(new_label)

        left_layout.addStretch()
        self.main_splitter.addWidget(left_panel)
        
        # Center Panel (File Tree with Enhanced Features)
        center_panel = QFrame()
        center_layout = QVBoxLayout(center_panel)

        # Header Layout (Status Label + Expand/Collapse Button)
        header_layout = QHBoxLayout()

        self.refresh_btn = QPushButton()
        self.refresh_btn.setIcon(self.style().standardIcon(self.style().StandardPixmap.SP_BrowserReload))
        self.refresh_btn.setToolTip("Refresh Tree")
        self.refresh_btn.clicked.connect(self._schedule_tree_refresh)
        header_layout.addWidget(self.refresh_btn)

        self.search_input = QLineEdit()
        self.search_input.setPlaceholderText("Filter files...")
        self.search_input.textChanged.connect(self.filter_tree)
        header_layout.addWidget(self.search_input)

        # Advanced Filter Toggle
        self.filter_toggle_btn = QPushButton("▼ Filters")
        self.filter_toggle_btn.setCheckable(True)
        self.filter_toggle_btn.clicked.connect(self.toggle_advanced_filters)
        header_layout.addWidget(self.filter_toggle_btn)

        self.expand_btn = QPushButton("Expand All")
        self.expand_btn.setCheckable(True)
        self.expand_btn.clicked.connect(self.toggle_expand)
        header_layout.addWidget(self.expand_btn)

        # Preview Toggle
        self.preview_toggle_btn = QPushButton("Preview")
        self.preview_toggle_btn.setCheckable(True)
        self.preview_toggle_btn.clicked.connect(self.toggle_preview_pane)
        header_layout.addWidget(self.preview_toggle_btn)

        self.clear_all_btn = QPushButton("Clear All")
        self.clear_all_btn.clicked.connect(self.clear_all_checkboxes)
        header_layout.addWidget(self.clear_all_btn)

        center_layout.addLayout(header_layout)

        # Advanced Filter Widget (Hidden by default)
        self.advanced_filter = AdvancedFilterWidget()
        self.advanced_filter.filter_changed.connect(self.apply_advanced_filters)
        self.advanced_filter.hide()
        center_layout.addWidget(self.advanced_filter)

        self.status_label = QLabel("Ready")
        center_layout.addWidget(self.status_label)

        # AGENTS definition for the Task Queue
        self.AGENTS = [
            {"id": "separate", "name": "Separate", "script": "separate.py", "color": "#e91e63", "short": "SEP"},
            {"id": "summarize", "name": "Summarize", "script": "summarize.py", "color": "#2196f3", "short": "SUM"},
            {"id": "sum_disc", "name": "Sum. Disc.", "script": "summarize_discovery.py", "color": "#4caf50", "short": "DISC"},
            {"id": "sum_depo", "name": "Sum. Depo.", "script": "summarize_deposition.py", "color": "#ff9800", "short": "DEPO"},
            {"id": "med_rec", "name": "Med Rec", "script": "med_record.py", "color": "#9c27b0", "short": "MED"},
            {"id": "med_chron", "name": "Med Chron", "script": "med_chron.py", "color": "#00bcd4", "short": "CHRON"},
        ]

        # Enhanced File Tree with additional columns
        self.tree = EnhancedFileTreeWidget()
        self.tree.item_moved.connect(lambda: self._schedule_tree_refresh())
        self.tree.folder_created.connect(lambda p: self._schedule_tree_refresh())
        self.tree.setHeaderLabels([
            "Category / File",
            "Queued Tasks (Click to Add ➕)",
            "Size",
            "Date Modified",
            "Status"
        ])
        self.tree.setSortingEnabled(True)
        self.tree.setColumnWidth(0, 300)
        self.tree.setColumnWidth(1, 200)
        self.tree.setColumnWidth(2, 70)
        self.tree.setColumnWidth(3, 100)
        self.tree.setColumnWidth(4, 80)
        self.tree.setSelectionMode(QAbstractItemView.SelectionMode.ExtendedSelection)
        self.tree.setAlternatingRowColors(True)
        self.tree.itemSelectionChanged.connect(self.on_tree_selection_changed)
        self.tree.itemDoubleClicked.connect(self.on_tree_double_click)
        self.tree.itemClicked.connect(self.on_tree_item_clicked)
        self.tree.tasks_column_clicked.connect(self.on_tree_item_clicked)
        center_layout.addWidget(self.tree)

        btn_layout = QHBoxLayout()
        self.process_btn = theme.primary_button("Process All Queued Tasks")
        self.process_btn.setMinimumHeight(36)
        self.process_btn.clicked.connect(self.process_checked_items)
        btn_layout.addWidget(self.process_btn)

        center_layout.addLayout(btn_layout)

        self.main_splitter.addWidget(center_panel)

        # Right Panel (File Preview - Hidden by default)
        self.preview_pane = FilePreviewWidget()
        self.preview_pane.hide()
        self.main_splitter.addWidget(self.preview_pane)

        self.main_splitter.setSizes([180, 800, 0])
        self.main_splitter.setStretchFactor(0, 0)
        self.main_splitter.setStretchFactor(1, 1)
        self.main_splitter.setStretchFactor(2, 0)
        
        # --- Tab 2: Status ---
        self.status_tab = QWidget()
        self.tabs.addTab(self.status_tab, "Status")
        status_layout = QVBoxLayout(self.status_tab)
        
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        self.status_container = QWidget()
        self.status_list_layout = QVBoxLayout(self.status_container)
        self.status_list_layout.setAlignment(Qt.AlignmentFlag.AlignTop)
        
        scroll.setWidget(self.status_container)
        status_layout.addWidget(scroll)
        
        clear_btn = theme.secondary_button("Clear Completed")
        clear_btn.clicked.connect(self.clear_completed_status)
        status_layout.addWidget(clear_btn)

        # --- Tab 3: Index ---
        self.index_tab = IndexTab(self)
        self.tabs.addTab(self.index_tab, "Index")
        if self.file_number:
            self.index_tab.load_data(self.file_number)

        # --- Tab 4: Chat ---
        # Persistent chat tab (visible in Advanced mode; hidden in Wizard mode).
        # Wizard mode spawns additional ChatTab instances via _open_task_tab.
        checkpoint("Creating ChatTab")
        self.chat_tab = ChatTab()
        self.tabs.addTab(self.chat_tab, "Chat")
        if self.file_number:
            self.chat_tab.load_case(self.file_number)

        # --- Tab: Discovery ---
        self.discovery_tab = DiscoveryTab()
        self.tabs.addTab(self.discovery_tab, "Discovery")
        if self.file_number:
            self.discovery_tab.load_case(self.file_number)

        # --- Tab: Templates / Resources ---
        self.templates_tab = TemplatesResourcesTab(main_window=self)
        self.tabs.addTab(self.templates_tab, "Templates / Resources")

        # --- Lazy tab loading ---------------------------------------------
        # On a case switch the active tab becomes "Case View", so none of the
        # data tabs below are visible. Instead of eagerly reloading every one
        # of them (JSON/DB reads + widget rebuilds, much of it on Z:\), we
        # register a per-tab loader and defer it until that tab is actually
        # shown. See load_case_by_number / _on_tab_changed.
        self._register_lazy_tabs()
        self.tabs.currentChanged.connect(self._on_tab_changed)

        # Add corner buttons next to the tabs
        self.corner_widget = QWidget()
        self.corner_layout = QHBoxLayout(self.corner_widget)
        self.corner_layout.setContentsMargins(5, 5, 10, 5)
        self.corner_layout.setSpacing(8)

        # Corner buttons share a fixed height so they align with the tab bar.
        corner_btn_height = 34

        # Notes button (secondary) — opens/creates the case AS NOTES document.
        # Mirrors the Notes button in the Case View toolbar so it stays
        # reachable in Wizard mode, where the Case View tab is hidden. Placed
        # before the Open File button so it sits to its left in the corner.
        self.btn_notes_corner = theme.secondary_button("Notes")
        self.btn_notes_corner.setFixedHeight(corner_btn_height)
        self.btn_notes_corner.setToolTip("Open or create the AS NOTES document for the current case")
        self.btn_notes_corner.clicked.connect(self.open_notes)
        self.corner_layout.addWidget(self.btn_notes_corner)

        # Open File button (blue - primary)
        self.btn_open_root = theme.primary_button("Open File")
        self.btn_open_root.setFixedHeight(corner_btn_height)
        self.btn_open_root.setToolTip("Open case folder in Explorer (Win+F)")
        self.btn_open_root.clicked.connect(self.open_root_folder)
        self.corner_layout.addWidget(self.btn_open_root)

        # Change File button removed in favor of Master List mode toggle (Wizard).

        # View menu button
        self.setup_view_menu()
        self.corner_layout.addWidget(self.view_btn)

        # Prompts button (secondary)
        self.prompts_btn = theme.secondary_button("Prompts")
        self.prompts_btn.setFixedHeight(corner_btn_height)
        self.prompts_btn.setToolTip("Open Prompt Engineering Workbench")
        self.prompts_btn.clicked.connect(self.manage_prompts)
        self.corner_layout.addWidget(self.prompts_btn)

        # Settings button with dropdown menu
        self.settings_btn = QToolButton()
        self.settings_btn.setText("Settings ▾")
        self.settings_btn.setToolTip("Application settings")
        self.settings_btn.setPopupMode(QToolButton.ToolButtonPopupMode.InstantPopup)
        theme.style_tool_button(self.settings_btn)
        self.settings_btn.setFixedHeight(corner_btn_height)
        self.settings_menu = QMenu(self)
        self.docket_refresh_action = self.settings_menu.addAction("Auto Docket Refresh")
        self.docket_refresh_action.setCheckable(True)
        self.settings_btn.setMenu(self.settings_menu)
        self.corner_layout.addWidget(self.settings_btn)

        # Wire settings actions to master_case_tab's toggle methods
        self.master_tab.docket_refresh_action = self.docket_refresh_action
        self.docket_refresh_action.toggled.connect(self.master_tab.toggle_docket_refresh)

        # Restart button (red - danger)
        self.restart_btn = theme.danger_button("Restart")
        self.restart_btn.setFixedHeight(corner_btn_height)
        self.restart_btn.setToolTip("Restart iCharlotte (open work is saved first)")
        self.restart_btn.clicked.connect(self.restart_app)
        self.corner_layout.addWidget(self.restart_btn)

        self.tabs.setCornerWidget(self.corner_widget, Qt.Corner.TopRightCorner)

        # Ctrl+Scroll zoom for all tabs (skips PDF viewers)
        self._zoom_filter = ZoomEventFilter(self.tabs, self)
        QApplication.instance().installEventFilter(self._zoom_filter)

        self._hide_fixed_close_buttons()
        checkpoint("setup_ui complete - all tabs created")

    # --- Wizard Mode: tab visibility orchestration ---

    # Names of tabs to HIDE when in Wizard Mode.
    # Master List is always visible. The Wizard tab and any task tabs are
    # added/managed separately in Phase 2+.
    _WIZARD_HIDDEN_TABS = {
        "Case View",
        "Status",
        "Index",
        "Chat",
        "Discovery",
        "Templates / Resources",
    }

    def _apply_mode_visibility(self, mode: str) -> None:
        """Show/hide tabs based on current mode."""
        is_wizard = (mode == "wizard")
        for i in range(self.tabs.count()):
            tab_text = self.tabs.tabText(i)
            if tab_text in self._WIZARD_HIDDEN_TABS:
                self.tabs.setTabVisible(i, not is_wizard)
            elif tab_text == "Wizard":
                self.tabs.setTabVisible(i, is_wizard)
            # Master List + task tabs (managed separately) stay visible.
        if not self.tabs.isTabVisible(self.tabs.currentIndex()):
            self.tabs.setCurrentIndex(0)
        self._hide_fixed_close_buttons()

    def _default_to_wizard_mode(self) -> None:
        """A freshly loaded case always opens in Wizard mode on the Wizard tab.

        Called from every case-load path (startup, Master List double-click,
        Change File dialog). ``set_mode`` flips tab visibility via
        ``_apply_mode_visibility``; we then explicitly select the Wizard tab
        (visibility changes can bounce the selection to Master List).
        """
        from icharlotte_core.ui.wizard.mode_controller import MODE_WIZARD
        self.mode_controller.set_mode(MODE_WIZARD)
        wizard_idx = self._index_of_tab("Wizard")
        if wizard_idx >= 0:
            self.tabs.setCurrentIndex(wizard_idx)

    # --- Wizard Mode: per-case task-tab snapshot/restore ---

    def _iter_task_tabs(self) -> list:
        """Return (index, widget) for every TaskTab currently in self.tabs."""
        out = []
        for i in range(self.tabs.count()):
            w = self.tabs.widget(i)
            if w is not None and w.property("wizard_task_id") is not None:
                out.append((i, w))
        return out

    def _iter_summary_browser_tabs(self) -> list:
        """Return case-scoped summary browser/viewer utility tabs."""
        out = []
        for i in range(self.tabs.count()):
            w = self.tabs.widget(i)
            if w is None:
                continue
            if (
                w.property("summary_browser_task_id") is not None
                or w.property("summary_output_viewer_path") is not None
            ):
                out.append((i, w))
        return out

    def _relpath_under(self, root: str, path: str) -> str:
        try:
            return os.path.relpath(path, root)
        except ValueError:
            return path

    def _snapshot_open_task_tabs(self, cancel_running: bool = True) -> list:
        """Build the persistence-ready snapshot of currently-open task tabs.

        cancel_running: at shutdown (True) a tab caught mid-run on the status
        page has its worker cancelled cleanly on the way out. Opportunistic
        mid-session saves pass False so they OBSERVE state without killing an
        in-flight task on this or any other open tab.
        """
        if not self.case_path:
            return []
        snapshots = []
        for _, tab in self._iter_task_tabs():
            if isinstance(tab, ChatTab):
                settings = {}
                conv_id = getattr(tab, "current_conversation_id", None)
                if conv_id:
                    settings["current_conversation_id"] = conv_id
                snapshots.append({
                    "task_id": "chat",
                    "instance_suffix": tab.property("wizard_instance_suffix") or "",
                    "files": [],
                    "settings": settings,
                    "page": "settings",
                    "output_path": None,
                })
                continue
            task_id = tab.spec.task_id
            # Determine page label.
            page_idx = tab.currentIndex()
            case_intake_metadata = {}
            if task_id == "case_intake_docket":
                from icharlotte_core.ui.wizard.pages.case_intake_docket_page import (
                    TASK_PAGE_COMPLAINT_STATUS,
                    TASK_PAGE_DOCKET_STATUS,
                    TASK_PAGE_REVIEW,
                    TASK_PAGE_SETTINGS as CASE_INTAKE_DOCKET_PAGE_SETTINGS,
                    TASK_PAGE_OUTPUT as CASE_INTAKE_DOCKET_PAGE_OUTPUT,
                )

                case_intake_metadata = dict(getattr(tab, "_last_metadata", {}) or {})
                if not case_intake_metadata and page_idx in (
                    TASK_PAGE_REVIEW,
                    CASE_INTAKE_DOCKET_PAGE_OUTPUT,
                ) and hasattr(tab, "review_page"):
                    try:
                        case_intake_metadata = dict(tab.review_page.to_dict())
                    except Exception:
                        case_intake_metadata = {}
                has_review_metadata = any(
                    bool(value)
                    for key, value in case_intake_metadata.items()
                    if key != "complaint_file"
                )

                if page_idx == CASE_INTAKE_DOCKET_PAGE_OUTPUT:
                    page = "output"
                elif page_idx == TASK_PAGE_REVIEW:
                    page = "review"
                elif page_idx == TASK_PAGE_DOCKET_STATUS:
                    if cancel_running and getattr(tab, "_worker", None) is not None:
                        try:
                            tab._worker.cancel()
                        except Exception:
                            pass
                    page = "review" if has_review_metadata else "settings"
                elif page_idx == TASK_PAGE_COMPLAINT_STATUS:
                    if cancel_running and getattr(tab, "_worker", None) is not None:
                        try:
                            tab._worker.cancel()
                        except Exception:
                            pass
                    page = "settings"
                elif page_idx == CASE_INTAKE_DOCKET_PAGE_SETTINGS:
                    page = "settings"
                else:
                    page = "settings"
            elif page_idx == 1:  # PAGE_STATUS — store as settings (no live worker
                                 # survives a restart, so it can't be restored).
                if cancel_running and getattr(tab, "_worker", None) is not None:
                    try:
                        tab._worker.cancel()
                    except Exception:
                        pass
                page = "settings"
            elif page_idx == 2:  # PAGE_OUTPUT
                page = "output"
            else:
                page = "settings"

            files_rel = [self._relpath_under(self.case_path, f) for f in tab.files]
            output_path = (
                getattr(tab.output_page, "output_path", None)
                if page == "output"
                else None
            )
            output_path_rel = self._relpath_under(self.case_path, output_path) if output_path else None
            settings = tab.settings_page.to_dict()
            snapshot = {
                "task_id": task_id,
                "instance_suffix": tab.property("wizard_instance_suffix") or "",
                "files": files_rel,
                "settings": settings,
                "page": page,
                "output_path": output_path_rel,
            }
            if task_id == "case_intake_docket":
                snapshot["settings"] = dict(case_intake_metadata)
                snapshot["metadata"] = dict(case_intake_metadata)
                snapshot["summary"] = dict(getattr(tab.output_page, "summary", {}) or {})
            snapshots.append(snapshot)
        return snapshots

    def _remove_all_task_tabs(self) -> None:
        # Iterate in reverse so indices stay stable.
        summary_tabs = (
            self._iter_summary_browser_tabs()
            if hasattr(self, "_iter_summary_browser_tabs")
            else []
        )
        for idx, widget in reversed(summary_tabs):
            self.tabs.removeTab(idx)
            widget.deleteLater()
        for idx, widget in reversed(self._iter_task_tabs()):
            if isinstance(widget, ChatTab):
                try:
                    widget.save_current_state()
                except Exception as e:
                    log_event(f"[wizard] chat save_current_state failed: {e}")
            self.tabs.removeTab(idx)
            widget.deleteLater()

    def _save_wizard_state_for_current_case(self, cancel_running: bool = True) -> None:
        if not self.case_path:
            return
        from icharlotte_core.ui.wizard.persistence import WizardStatePersistence
        p = WizardStatePersistence(self.case_path)
        p.set_open_tabs(self._snapshot_open_task_tabs(cancel_running=cancel_running))
        p.save()

    def _persist_open_tabs(self) -> None:
        """Save the open-tab snapshot WITHOUT cancelling any running worker.

        Called after the open-tab set or a tab's contents change (open / close /
        complete) so the session survives a force-kill or freeze — both of which
        bypass closeEvent and restart_app, leaving nothing saved otherwise.
        """
        try:
            self._save_wizard_state_for_current_case(cancel_running=False)
        except Exception as e:
            log_event(f"[wizard] opportunistic state save failed: {e}")

    def _persist_open_tabs_soon(self) -> None:
        """Deferred variant of _persist_open_tabs.

        A task tab emits ``task_completed`` BEFORE it finishes switching to its
        output page, so deferring one event-loop turn lets the snapshot capture
        the output page (and its output_path) rather than the stale settings
        page. Used on task completion.
        """
        QTimer.singleShot(0, self._persist_open_tabs)

    def _on_library_captured(self, case_root: str) -> None:
        """A background capture finished — refresh any open Chat tab's library tree.

        Runs on the GUI thread (queued from the capture worker thread), so it is
        safe to touch the Chat tab's QTreeWidget here.
        """
        try:
            from icharlotte_core.ui.tabs import ChatTab
            ChatTab.refresh_open_library_trees(self.tabs)
        except Exception:
            pass

    def _on_task_completed(self, entry: dict) -> None:
        if not self.case_path:
            return
        # Best-effort: capture the finished task's source text into the
        # document library (off the UI thread). Uses the original entry with
        # absolute file paths, before they are rewritten case-relative below.
        # Capture original absolute-path entry into the document library (off-UI-thread, best-effort).
        QThreadPool.globalInstance().start(
            _LibraryCaptureJob(self.case_path, copy.deepcopy(entry), self.library_captured)
        )
        from icharlotte_core.ui.wizard.persistence import WizardStatePersistence
        # Store files & output_path as case-relative.
        entry = dict(entry)
        entry["files"] = [self._relpath_under(self.case_path, f) for f in entry.get("files", [])]
        if entry.get("output_path"):
            entry["output_path"] = self._relpath_under(self.case_path, entry["output_path"])
        p = WizardStatePersistence(self.case_path)
        p.add_recent_task(entry)
        p.save()
        # Tell the Wizard tab to refresh its Recent Tasks list.
        if hasattr(self, "wizard_tab") and self.wizard_tab is not None:
            if hasattr(self.wizard_tab, "refresh_recent_tasks"):
                self.wizard_tab.refresh_recent_tasks(p.get_recent_tasks())
        # Re-snapshot the open tabs so the just-produced output is restorable
        # even if the session ends by force-kill (deferred so the tab has
        # finished switching to its output page).
        self._persist_open_tabs_soon()

    def _on_reopen_recent_task(self, entry: dict) -> None:
        from icharlotte_core.ui.wizard.registry import get_task, TASK_REGISTRY
        from icharlotte_core.ui.wizard.task_tab import TaskTab, PAGE_OUTPUT, PAGE_SETTINGS
        from icharlotte_core.ui.wizard.task_routing import get_in_process_task_builder_name
        from icharlotte_core.ui.wizard.instance_naming import next_instance_suffix

        task_id = entry.get("task_id")
        if task_id not in TASK_REGISTRY:
            QMessageBox.warning(self, "Unknown task", f"This case references an unknown task: {task_id!r}")
            return
        spec = get_task(task_id)

        out_rel = entry.get("output_path") or ""
        out_abs = os.path.join(self.case_path, out_rel) if out_rel and self.case_path else out_rel
        files = [
            os.path.join(self.case_path, f) if self.case_path and not os.path.isabs(f) else f
            for f in entry.get("files", [])
        ]

        existing_titles = [self.tabs.tabText(i) for i in range(self.tabs.count())]
        suffix = next_instance_suffix(spec.title, existing_titles)
        title = f"{spec.title} {suffix}".strip()

        builder_name = get_in_process_task_builder_name(task_id)
        if builder_name and builder_name not in (
            "build_oppose_motion_tab",
            "build_motion_drafting_tab",
            "build_mediation_brief_tab",
            "build_generate_motion_tab",
            "build_case_intake_docket_tab",
        ):
            # In-process custom tabs (e.g. Separate) own their source selection;
            # reopening re-runs the builder's picker. Analysis output is
            # ephemeral, so there's nothing to restore beyond the tab itself.
            # (Mediation Brief is excluded above so its saved brief is reloaded.)
            from icharlotte_core.ui.wizard import in_process_task_tab
            builder = getattr(in_process_task_tab, builder_name)
            builder_kwargs = {
                "spec": spec,
                "case_path": self.case_path,
                "file_number": self.file_number,
                "parent": self,
            }
            if builder_name == "build_med_extractor_tab":
                settings = dict(entry.get("settings") or {})
                chronology_path = settings.get("chronology_path") or (files[0] if files else "")
                if chronology_path and self.case_path and not os.path.isabs(chronology_path):
                    chronology_path = os.path.join(self.case_path, chronology_path)
                if chronology_path:
                    chronology_path = os.path.normpath(chronology_path)
                if chronology_path:
                    settings["chronology_path"] = chronology_path
                builder_kwargs["chronology_path"] = chronology_path
                builder_kwargs["initial_settings"] = settings
            task_tab = builder(**builder_kwargs)
            if task_tab is None:
                return
            task_tab.setProperty("wizard_task_id", spec.task_id)
            task_tab.setProperty("wizard_instance_suffix", suffix)
            task_tab.task_completed.connect(self._on_task_completed)
            new_index = self.tabs.addTab(task_tab, title)
            self.tabs.setCurrentIndex(new_index)
            self._hide_fixed_close_buttons()
            return

        restored_summary = False
        if builder_name == "build_oppose_motion_tab":
            from icharlotte_core.ui.wizard.pages.oppose_motion_page import (
                OpposeMotionTaskTab,
                TASK_PAGE_OUTPUT,
                TASK_PAGE_SETTINGS,
            )

            settings = dict(entry.get("settings") or {})
            motion_file = settings.get("motion_file") or (files[0] if files else "")
            context_files = settings.get("context_files") or files[1:]
            task_tab = OpposeMotionTaskTab(
                spec=spec,
                case_path=self.case_path,
                file_number=self.file_number,
                motion_file=motion_file,
                context_files=context_files,
                parent=self,
            )
            task_tab.settings_page.from_dict(settings)
            output_page = TASK_PAGE_OUTPUT
            settings_page = TASK_PAGE_SETTINGS
        elif builder_name == "build_motion_drafting_tab":
            from icharlotte_core.ui.wizard.pages.motion_drafting_page import (
                MotionDraftingTaskTab,
                TASK_PAGE_OUTPUT,
                TASK_PAGE_SETTINGS,
            )

            settings = dict(entry.get("settings") or {})
            task_tab = MotionDraftingTaskTab(
                spec=spec,
                case_path=self.case_path,
                file_number=self.file_number,
                parent=self,
            )
            task_tab.settings_page.from_dict(settings)
            output_page = TASK_PAGE_OUTPUT
            settings_page = TASK_PAGE_SETTINGS
        elif builder_name == "build_generate_motion_tab":
            from icharlotte_core.ui.wizard.pages.generate_motion_page import (
                GenerateMotionTaskTab,
                TASK_PAGE_OUTPUT,
                TASK_PAGE_SETTINGS,
            )

            settings = dict(entry.get("settings") or {})
            task_tab = GenerateMotionTaskTab(
                spec=spec,
                case_path=self.case_path,
                file_number=self.file_number,
                parent=self,
            )
            task_tab.settings_page.from_dict(settings)
            output_page = TASK_PAGE_OUTPUT
            settings_page = TASK_PAGE_SETTINGS
        elif builder_name == "build_mediation_brief_tab":
            from icharlotte_core.ui.wizard.pages.mediation_brief_page import (
                MediationBriefTaskTab,
                TASK_PAGE_OUTPUT,
                TASK_PAGE_SETTINGS,
            )

            task_tab = MediationBriefTaskTab(
                spec=spec,
                case_path=self.case_path,
                file_number=self.file_number,
                parent=self,
            )
            output_page = TASK_PAGE_OUTPUT
            settings_page = TASK_PAGE_SETTINGS
        elif builder_name == "build_case_intake_docket_tab":
            from icharlotte_core.ui.wizard.pages.case_intake_docket_page import (
                CaseIntakeDocketTaskTab,
                TASK_PAGE_OUTPUT,
                TASK_PAGE_REVIEW,
                TASK_PAGE_SETTINGS,
            )

            metadata = dict(entry.get("metadata") or entry.get("settings") or {})
            task_tab = CaseIntakeDocketTaskTab(
                spec=spec,
                case_path=self.case_path,
                file_number=self.file_number,
                parent=self,
            )
            if metadata:
                task_tab.load_review_state(metadata)
            summary = dict(entry.get("summary") or {})
            for key in ("docket_pdf", "variables_docx"):
                path = str(summary.get(key) or "")
                if path and self.case_path and not os.path.isabs(path):
                    summary[key] = os.path.join(self.case_path, path)
            if summary:
                task_tab.load_output_summary(summary, metadata=metadata or None)
                restored_summary = True
            output_page = TASK_PAGE_OUTPUT
            settings_page = TASK_PAGE_REVIEW if metadata else TASK_PAGE_SETTINGS
        else:
            task_tab = TaskTab(
                spec=spec,
                files=files,
                case_path=self.case_path,
                file_number=self.file_number,
                parent=self,
            )
            output_page = PAGE_OUTPUT
            settings_page = PAGE_SETTINGS
        task_tab.setProperty("wizard_task_id", spec.task_id)
        task_tab.setProperty("wizard_instance_suffix", suffix)
        try:
            if get_in_process_task_builder_name(task_id) != "build_oppose_motion_tab":
                task_tab.settings_page.from_dict(entry.get("settings") or {})
        except Exception:
            pass
        new_index = self.tabs.addTab(task_tab, title)
        task_tab.task_completed.connect(self._on_task_completed)

        if out_abs and os.path.exists(out_abs):
            if hasattr(task_tab.output_page, "load_output"):
                task_tab.output_page.load_output(out_abs)
            task_tab.setCurrentIndex(output_page)
        elif restored_summary:
            task_tab.setCurrentIndex(output_page)
        else:
            QMessageBox.information(
                self,
                "Output missing",
                f"The saved output file no longer exists.\nYou can re-run with the saved settings.",
            )
            task_tab.setCurrentIndex(settings_page)

        self.tabs.setCurrentIndex(new_index)
        self._hide_fixed_close_buttons()

    def _restore_task_tabs_for_case(self) -> None:
        if not self.case_path:
            return
        from icharlotte_core.ui.wizard.persistence import WizardStatePersistence
        from icharlotte_core.ui.wizard.registry import get_task, TASK_REGISTRY
        from icharlotte_core.ui.wizard.task_tab import TaskTab, PAGE_OUTPUT, PAGE_SETTINGS
        from icharlotte_core.ui.wizard.task_routing import get_in_process_task_builder_name

        p = WizardStatePersistence(self.case_path)
        for entry in p.get_open_tabs():
            task_id = entry.get("task_id")
            if task_id not in TASK_REGISTRY:
                continue
            spec = get_task(task_id)
            if task_id == "chat":
                suffix = entry.get("instance_suffix", "") or ""
                tab = ChatTab(parent=self)
                tab.setProperty("wizard_task_id", "chat")
                tab.setProperty("wizard_instance_suffix", suffix)
                if self.file_number:
                    tab.load_case(self.file_number)
                conv_id = (entry.get("settings") or {}).get("current_conversation_id")
                if conv_id:
                    try:
                        tab.on_conversation_selected(conv_id)
                    except Exception as e:
                        log_event(f"[wizard] restore chat conversation failed: {e}")
                self.tabs.addTab(tab, f"{spec.title} {suffix}".strip())
                continue
            files_abs = [
                f if os.path.isabs(f) else os.path.join(self.case_path, f)
                for f in entry.get("files", [])
            ]
            settings_dict = entry.get("settings") or {}
            builder_name = get_in_process_task_builder_name(task_id)
            if builder_name and builder_name not in (
                "build_oppose_motion_tab",
                "build_motion_drafting_tab",
                "build_mediation_brief_tab",
                "build_generate_motion_tab",
                "build_case_intake_docket_tab",
            ):
                # In-process custom tabs (e.g. Separate) re-pick their source on
                # restore; skip silently if the user cancels the picker.
                # (Mediation Brief is excluded above so its saved brief is reloaded.)
                from icharlotte_core.ui.wizard import in_process_task_tab
                builder = getattr(in_process_task_tab, builder_name)
                builder_kwargs = {
                    "spec": spec,
                    "case_path": self.case_path,
                    "file_number": self.file_number,
                    "parent": self,
                }
                if builder_name == "build_med_extractor_tab":
                    settings_dict = dict(settings_dict)
                    chronology_path = settings_dict.get("chronology_path") or (
                        files_abs[0] if files_abs else ""
                    )
                    if chronology_path and self.case_path and not os.path.isabs(chronology_path):
                        chronology_path = os.path.join(self.case_path, chronology_path)
                    if chronology_path:
                        chronology_path = os.path.normpath(chronology_path)
                    if chronology_path:
                        settings_dict["chronology_path"] = chronology_path
                    builder_kwargs["chronology_path"] = chronology_path
                    builder_kwargs["initial_settings"] = settings_dict
                tab = builder(**builder_kwargs)
                if tab is None:
                    continue
                suffix = entry.get("instance_suffix", "") or ""
                tab.setProperty("wizard_task_id", spec.task_id)
                tab.setProperty("wizard_instance_suffix", suffix)
                tab.task_completed.connect(self._on_task_completed)
                self.tabs.addTab(tab, f"{spec.title} {suffix}".strip())
                continue

            restored_summary = False
            if builder_name == "build_oppose_motion_tab":
                from icharlotte_core.ui.wizard.pages.oppose_motion_page import (
                    OpposeMotionTaskTab,
                    TASK_PAGE_OUTPUT,
                    TASK_PAGE_SETTINGS,
                )

                motion_file = settings_dict.get("motion_file") or (
                    files_abs[0] if files_abs else ""
                )
                context_files = settings_dict.get("context_files") or files_abs[1:]
                tab = OpposeMotionTaskTab(
                    spec=spec,
                    case_path=self.case_path,
                    file_number=self.file_number,
                    motion_file=motion_file,
                    context_files=context_files,
                    parent=self,
                )
                output_page = TASK_PAGE_OUTPUT
                settings_page = TASK_PAGE_SETTINGS
            elif builder_name == "build_motion_drafting_tab":
                from icharlotte_core.ui.wizard.pages.motion_drafting_page import (
                    MotionDraftingTaskTab,
                    TASK_PAGE_OUTPUT,
                    TASK_PAGE_SETTINGS,
                )

                tab = MotionDraftingTaskTab(
                    spec=spec,
                    case_path=self.case_path,
                    file_number=self.file_number,
                    parent=self,
                )
                output_page = TASK_PAGE_OUTPUT
                settings_page = TASK_PAGE_SETTINGS
            elif builder_name == "build_generate_motion_tab":
                from icharlotte_core.ui.wizard.pages.generate_motion_page import (
                    GenerateMotionTaskTab,
                    TASK_PAGE_OUTPUT,
                    TASK_PAGE_SETTINGS,
                )

                tab = GenerateMotionTaskTab(
                    spec=spec,
                    case_path=self.case_path,
                    file_number=self.file_number,
                    parent=self,
                )
                output_page = TASK_PAGE_OUTPUT
                settings_page = TASK_PAGE_SETTINGS
            elif builder_name == "build_mediation_brief_tab":
                from icharlotte_core.ui.wizard.pages.mediation_brief_page import (
                    MediationBriefTaskTab,
                    TASK_PAGE_OUTPUT,
                    TASK_PAGE_SETTINGS,
                )

                tab = MediationBriefTaskTab(
                    spec=spec,
                    case_path=self.case_path,
                    file_number=self.file_number,
                    parent=self,
                )
                output_page = TASK_PAGE_OUTPUT
                settings_page = TASK_PAGE_SETTINGS
            elif builder_name == "build_case_intake_docket_tab":
                from icharlotte_core.ui.wizard.pages.case_intake_docket_page import (
                    CaseIntakeDocketTaskTab,
                    TASK_PAGE_OUTPUT,
                    TASK_PAGE_REVIEW,
                    TASK_PAGE_SETTINGS,
                )

                metadata = dict(entry.get("metadata") or settings_dict or {})
                tab = CaseIntakeDocketTaskTab(
                    spec=spec,
                    case_path=self.case_path,
                    file_number=self.file_number,
                    parent=self,
                )
                if metadata:
                    tab.load_review_state(metadata)
                summary = dict(entry.get("summary") or {})
                for key in ("docket_pdf", "variables_docx"):
                    path = str(summary.get(key) or "")
                    if path and self.case_path and not os.path.isabs(path):
                        summary[key] = os.path.join(self.case_path, path)
                if summary:
                    tab.load_output_summary(summary, metadata=metadata or None)
                    restored_summary = True
                output_page = TASK_PAGE_OUTPUT
                saved_page = entry.get("page", "settings")
                if saved_page == "review":
                    settings_page = TASK_PAGE_REVIEW
                elif saved_page == "settings":
                    settings_page = TASK_PAGE_SETTINGS
                else:
                    settings_page = TASK_PAGE_REVIEW if metadata else TASK_PAGE_SETTINGS
            else:
                tab = TaskTab(
                    spec=spec,
                    files=files_abs,
                    case_path=self.case_path,
                    file_number=self.file_number,
                    parent=self,
                )
                output_page = PAGE_OUTPUT
                settings_page = PAGE_SETTINGS
            suffix = entry.get("instance_suffix", "") or ""
            tab.setProperty("wizard_task_id", spec.task_id)
            tab.setProperty("wizard_instance_suffix", suffix)
            tab.task_completed.connect(self._on_task_completed)
            title = f"{spec.title} {suffix}".strip()

            # Restore settings dict if present.
            try:
                tab.settings_page.from_dict(settings_dict)
            except Exception:
                pass

            self.tabs.addTab(tab, title)

            # Restore page.
            page = entry.get("page", "settings")
            if page == "output":
                out_rel = entry.get("output_path")
                out_abs = os.path.join(self.case_path, out_rel) if out_rel else None
                if out_abs and os.path.exists(out_abs):
                    if hasattr(tab.output_page, "load_output"):
                        tab.output_page.load_output(out_abs)
                    tab.setCurrentIndex(output_page)
                elif restored_summary:
                    tab.setCurrentIndex(output_page)
                else:
                    tab.setCurrentIndex(settings_page)
            else:
                tab.setCurrentIndex(settings_page)

            # Mirror the open path: per-file tasks run Phase 1 speculatively on
            # the settings page. Without this, a restored tab on PAGE_SETTINGS
            # shows "Discovering topics…" at 0% forever — no worker is alive.
            if (task_id in ("summarize_depositions", "med_chron_analysis")
                    and tab.currentIndex() == PAGE_SETTINGS
                    and tab.files):
                try:
                    tab.start_speculative_run()
                except Exception as e:
                    log_event(f"[wizard] restore speculative_run failed: {e}", "error")
        self._hide_fixed_close_buttons()

    def _index_of_tab(self, tab_text: str) -> int:
        for i in range(self.tabs.count()):
            if self.tabs.tabText(i) == tab_text:
                return i
        return -1

    def _switch_to_status_tab(self) -> None:
        """Switch to the Status tab by name, only if it's currently visible.

        Tab indices shifted when the Wizard tab was inserted at index 1, so
        hardcoded indices silently jump to the wrong tab. Wizard mode hides
        Status; in that mode this is a no-op (wizard pages own their own UI).
        """
        idx = self._index_of_tab("Status")
        if idx >= 0 and self.tabs.isTabVisible(idx):
            self.tabs.setCurrentIndex(idx)

    def setup_view_menu(self):
        self.view_btn = QToolButton()
        self.view_btn.setText("View ▾")
        self.view_btn.setPopupMode(QToolButton.ToolButtonPopupMode.InstantPopup)
        theme.style_tool_button(self.view_btn)
        self.view_btn.setFixedHeight(34)
        
        self.view_menu = QMenu(self)
        self.view_btn.setMenu(self.view_menu)
        
        # Load and apply settings
        settings = self.load_tab_settings()
        
        # Populate menu based on current tabs
        for i in range(self.tabs.count()):
            tab_text = self.tabs.tabText(i)
            
            # Apply saved setting if available
            if tab_text in settings:
                self.tabs.setTabVisible(i, settings[tab_text])
            
            action = QAction(tab_text, self)
            action.setCheckable(True)
            action.setChecked(self.tabs.isTabVisible(i))
            # Use partial to capture the current loop variable 'i'
            action.toggled.connect(partial(self.toggle_tab_visibility, i))
            self.view_menu.addAction(action)

        self._add_debug_console_action()

    def _add_debug_console_action(self):
        """Add the floating task debug console opener to the View menu."""
        for existing in self.view_menu.actions():
            if existing.text() == "Debug Console":
                self.debug_console_action = existing
                return
        if self.view_menu.actions():
            self.view_menu.addSeparator()
        action = QAction("Debug Console", self.view_menu)
        action.triggered.connect(self.show_task_debug_console)
        self.view_menu.addAction(action)
        self.debug_console_action = action

    def show_task_debug_console(self):
        """Show the reusable floating task debug console."""
        if getattr(self, "_task_debug_window", None) is None:
            from icharlotte_core.ui.task_debug_window import TaskDebugWindow

            self._task_debug_window = TaskDebugWindow(parent=self)
        self._task_debug_window.show()
        self._task_debug_window.raise_()
        self._task_debug_window.activateWindow()

    def toggle_tab_visibility(self, index, visible):
        self.tabs.setTabVisible(index, visible)
        self.save_tab_settings()

    def load_tab_settings(self):
        config_dir = os.path.join(GEMINI_DATA_DIR, "..", "config")
        settings_path = os.path.join(config_dir, "view_settings.json")
        if os.path.exists(settings_path):
            try:
                with open(settings_path, 'r') as f:
                    return json.load(f)
            except Exception as e:
                log_event(f"Error loading view settings: {e}", "error")
        return {}

    def save_tab_settings(self):
        config_dir = os.path.join(GEMINI_DATA_DIR, "..", "config")
        if not os.path.exists(config_dir):
            try:
                os.makedirs(config_dir)
            except Exception as e:
                log_event(f"Failed to create config dir {config_dir}: {e}", "error")
            
        settings_path = os.path.join(config_dir, "view_settings.json")
        settings = {}
        for i in range(self.tabs.count()):
            settings[self.tabs.tabText(i)] = self.tabs.isTabVisible(i)
            
        try:
            with open(settings_path, 'w') as f:
                json.dump(settings, f)
        except Exception as e:
            log_event(f"Error saving view settings: {e}", "error")

    def restart_app(self):
        log_event("User requested manual restart. Spawning new process...")
        # Persist session state BEFORE quitting. QApplication.quit() does not
        # dispatch closeEvent, so without this the open wizard tabs and their
        # contents would be lost when restarting via the in-app button.
        try:
            self._persist_session_state()
        except Exception as e:
            log_event(f"[wizard] restart persist failed: {e}", "error")
        # QApplication.quit() does not dispatch closeEvent, so save here too
        try:
            self._save_window_state()
        except Exception as e:
            log_event(f"Error saving window state on restart: {e}", "error")
        # Close all agent runners if any are running
        for runner in self.agent_runners:
            try:
                runner.terminate()
            except Exception:
                pass

        # Get the absolute path to this script
        script_path = os.path.abspath(__file__)

        # Build restart arguments with current state
        args = [script_path]

        # Add current file number if loaded
        if self.file_number:
            args.extend(['--file-number', str(self.file_number)])

        # Add current case path if loaded
        if self.case_path:
            args.extend(['--case-path', str(self.case_path)])

        # Add current tab index
        current_tab = self.tabs.currentIndex()
        args.extend(['--tab', str(current_tab)])

        log_event(f"Restarting with: python={sys.executable}, args={args}")
        log_event(f"Current state: file_number={self.file_number}, case_path={self.case_path}, tab={current_tab}")

        # Spawn new process
        subprocess.Popen([sys.executable] + args)

        # Exit current process
        QApplication.quit()

    def clear_all_checkboxes(self):
        self.tree.blockSignals(True)
        iterator = QTreeWidgetItemIterator(self.tree)
        while iterator.value():
            item = iterator.value()
            item.setData(0, Qt.ItemDataRole.UserRole + 2, [])
            self.update_item_tasks_ui(item)
            iterator += 1
        self.tree.blockSignals(False)

    def _open_task_tab(self, task_id: str) -> None:
        from icharlotte_core.ui.wizard.registry import get_task
        from icharlotte_core.ui.wizard.file_picker import (
            find_medical_summary_folder,
            resolve_default_folder,
        )
        from icharlotte_core.ui.wizard.task_routing import (
            get_in_process_task_builder_name,
            opens_settings_without_picker,
            requires_initial_file_picker,
        )
        from icharlotte_core.ui.wizard.task_tab import TaskTab
        from icharlotte_core.ui.wizard.instance_naming import next_instance_suffix
        from PySide6.QtWidgets import QFileDialog

        if not self.case_path:
            QMessageBox.information(self, "No case loaded", "Open a case from the Master List first.")
            return

        if task_id == "chat":
            spec = get_task(task_id)
            existing_titles = [self.tabs.tabText(i) for i in range(self.tabs.count())]
            suffix = next_instance_suffix(spec.title, existing_titles)
            title = f"{spec.title} {suffix}".strip()
            chat_tab = ChatTab(parent=self)
            chat_tab.setProperty("wizard_task_id", spec.task_id)
            chat_tab.setProperty("wizard_instance_suffix", suffix)
            if self.file_number:
                chat_tab.load_case(self.file_number)
            new_index = self.tabs.addTab(chat_tab, title)
            self.tabs.setCurrentIndex(new_index)
            log_event(f"[wizard] opened new Chat tab '{title}'")
            self._hide_fixed_close_buttons()
            return

        spec = get_task(task_id)
        in_process_builder_name = get_in_process_task_builder_name(task_id)
        if in_process_builder_name:
            from icharlotte_core.ui.wizard import in_process_task_tab

            existing_titles = [self.tabs.tabText(i) for i in range(self.tabs.count())]
            suffix = next_instance_suffix(spec.title, existing_titles)
            title = f"{spec.title} {suffix}".strip()
            builder = getattr(in_process_task_tab, in_process_builder_name)
            task_tab = builder(
                spec=spec,
                case_path=self.case_path,
                file_number=self.file_number,
                parent=self,
            )
            if task_tab is None:
                return
            task_tab.setProperty("wizard_task_id", spec.task_id)
            task_tab.setProperty("wizard_instance_suffix", suffix)
            task_tab.task_completed.connect(self._on_task_completed)
            new_index = self.tabs.addTab(task_tab, title)
            self.tabs.setCurrentIndex(new_index)
            log_event(f"[wizard] opened in-process task tab '{title}'")
            self._hide_fixed_close_buttons()
            return

        if opens_settings_without_picker(task_id):
            # The Settings page manages its own source selection (e.g. Depo Prep's
            # two-bucket pickers). Open it directly with no pre-Settings file
            # picker so the user adds files in the correct bucket on the page.
            existing_titles = [self.tabs.tabText(i) for i in range(self.tabs.count())]
            suffix = next_instance_suffix(spec.title, existing_titles)
            title = f"{spec.title} {suffix}".strip()
            task_tab = TaskTab(
                spec=spec,
                files=[],
                case_path=self.case_path,
                file_number=self.file_number,
                parent=self,
            )
            task_tab.setProperty("wizard_task_id", spec.task_id)
            task_tab.setProperty("wizard_instance_suffix", suffix)
            task_tab.task_completed.connect(self._on_task_completed)
            new_index = self.tabs.addTab(task_tab, title)
            self.tabs.setCurrentIndex(new_index)
            log_event(f"[wizard] opened '{title}' (settings-managed sources)")
            self._hide_fixed_close_buttons()
            return

        if not requires_initial_file_picker(task_id):
            log_event(f"[wizard] task '{task_id}' has no file-picker route")
            return

        start_dir = None
        if task_id == "med_chron_analysis":
            start_dir = find_medical_summary_folder(self.case_path)
        if start_dir is None:
            start_dir = resolve_default_folder(self.case_path, spec.default_folders)
        files, _ = QFileDialog.getOpenFileNames(
            self,
            f"Select files for {spec.title}",
            start_dir,
            "All files (*.*)",
        )
        if not files:
            return  # user cancelled → no tab created

        if task_id in ("summarize_depositions", "med_chron_analysis"):
            # One tab per file so each has its own speculative Phase 1 worker.
            first_new_index = None
            for file_path in files:
                existing_titles = [self.tabs.tabText(i) for i in range(self.tabs.count())]
                suffix = next_instance_suffix(spec.title, existing_titles)
                title = f"{spec.title} {suffix}".strip()
                task_tab = TaskTab(
                    spec=spec,
                    files=[file_path],
                    case_path=self.case_path,
                    file_number=self.file_number,
                    parent=self,
                )
                task_tab.setProperty("wizard_task_id", spec.task_id)
                task_tab.setProperty("wizard_instance_suffix", suffix)
                task_tab.task_completed.connect(self._on_task_completed)
                new_index = self.tabs.addTab(task_tab, title)
                if first_new_index is None:
                    first_new_index = new_index
                task_tab.start_speculative_run()
                log_event(f"[wizard] opened '{title}' (speculative Phase 1)")
            if first_new_index is not None:
                self.tabs.setCurrentIndex(first_new_index)
            self._hide_fixed_close_buttons()
            return

        # All other tasks: bundle files into a single sequential tab.
        existing_titles = [self.tabs.tabText(i) for i in range(self.tabs.count())]
        suffix = next_instance_suffix(spec.title, existing_titles)
        title = f"{spec.title} {suffix}".strip()

        task_tab = TaskTab(
            spec=spec,
            files=files,
            case_path=self.case_path,
            file_number=self.file_number,
            parent=self,
        )
        task_tab.setProperty("wizard_task_id", spec.task_id)
        task_tab.setProperty("wizard_instance_suffix", suffix)
        task_tab.task_completed.connect(self._on_task_completed)
        new_index = self.tabs.addTab(task_tab, title)
        self.tabs.setCurrentIndex(new_index)
        log_event(f"[wizard] opened task tab '{title}' with {len(files)} file(s)")
        self._hide_fixed_close_buttons()

    def _on_card_action(self, action_id: str) -> None:
        """Dispatch a launcher-card corner-button action."""
        if action_id == "open_separate_index":
            self._reveal_index_tab()
            return
        summary_task_id = task_id_for_summary_action(action_id)
        if summary_task_id:
            self._open_summary_browser_tab(summary_task_id)

    def _open_summary_browser_tab(self, task_id: str) -> None:
        """Open or refresh a task-specific summary browser tab."""
        if not self.case_path or not self.file_number:
            QMessageBox.information(
                self, "No case loaded",
                "Open a case from the Master List first.",
            )
            return

        for i in range(self.tabs.count()):
            widget = self.tabs.widget(i)
            if (
                widget is not None
                and widget.property("summary_browser_task_id") == task_id
            ):
                if hasattr(widget, "refresh"):
                    widget.refresh()
                self.tabs.setCurrentIndex(i)
                self._hide_fixed_close_buttons()
                return

        browser = SummaryBrowserTab(
            case_path=self.case_path,
            file_number=self.file_number,
            task_id=task_id,
            parent=self,
        )
        browser.setProperty("summary_browser_task_id", task_id)
        browser.open_requested.connect(self._open_summary_output_tab)
        new_index = self.tabs.addTab(browser, summary_browser_title(task_id))
        self.tabs.setCurrentIndex(new_index)
        self._hide_fixed_close_buttons()

    def _open_summary_output_tab(self, output_path: str) -> None:
        """Open a selected summary in the standard in-app output editor."""
        if not output_path or not os.path.isfile(output_path):
            QMessageBox.warning(
                self,
                "Output not found",
                "The selected summary file could not be found.",
            )
            return

        norm_path = os.path.normcase(os.path.abspath(output_path))
        for i in range(self.tabs.count()):
            widget = self.tabs.widget(i)
            viewer_path = widget.property("summary_output_viewer_path") if widget else None
            if viewer_path and os.path.normcase(os.path.abspath(viewer_path)) == norm_path:
                self.tabs.setCurrentIndex(i)
                self._hide_fixed_close_buttons()
                return

        from icharlotte_core.ui.wizard.pages.output_page import OutputPage

        viewer = OutputPage(self)
        viewer.setProperty("summary_output_viewer_path", output_path)
        viewer.load_output(output_path)
        if hasattr(viewer, "rerun_btn"):
            viewer.rerun_btn.setVisible(False)
        if hasattr(viewer, "edit_settings_btn"):
            viewer.edit_settings_btn.setVisible(False)
        title = f"Output - {os.path.basename(output_path)}"
        new_index = self.tabs.addTab(viewer, title)
        self.tabs.setCurrentIndex(new_index)
        self._hide_fixed_close_buttons()

    def _reveal_index_tab(self) -> None:
        """Wizard Mode: reveal the hidden Index singleton, reloaded from disk so
        it reflects the latest Separate runs (wizard or advanced)."""
        if not self.case_path or not self.file_number:
            QMessageBox.information(
                self, "No case loaded",
                "Open a case from the Master List first.",
            )
            return
        idx = self._index_of_tab("Index")
        if idx < 0:
            return
        if hasattr(self, "index_tab"):
            self.index_tab.load_data(self.file_number)
        self.tabs.setTabVisible(idx, True)
        self.tabs.setCurrentIndex(idx)
        self._hide_fixed_close_buttons()

    def _on_tab_close_requested(self, index: int) -> None:
        """Only TaskTabs are closeable (they carry a 'wizard_task_id' property)."""
        widget = self.tabs.widget(index)
        if widget is None:
            return
        # Wizard Mode: the Index tab is the shared singleton — its "x" hides it,
        # never destroys it.
        if (
            widget is getattr(self, "index_tab", None)
            and getattr(self, "mode_controller", None) is not None
            and self.mode_controller.is_wizard
        ):
            self.tabs.setTabVisible(index, False)
            self._hide_fixed_close_buttons()
            wiz = self._index_of_tab("Wizard")
            if wiz >= 0:
                self.tabs.setCurrentIndex(wiz)
            return
        if (
            widget.property("summary_browser_task_id") is not None
            or widget.property("summary_output_viewer_path") is not None
        ):
            self.tabs.removeTab(index)
            widget.deleteLater()
            self._hide_fixed_close_buttons()
            self._persist_open_tabs()
            return
        if widget.property("wizard_task_id") is None:
            return  # not a task tab; ignore
        # Cancel any running worker before removing the tab.
        worker = getattr(widget, "_worker", None)
        if worker is not None:
            if widget.__class__.__name__ in (
                "OpposeMotionTaskTab", "GenerateMotionTaskTab", "SeparateTaskTab",
                "MediationBriefTaskTab",
            ) and worker.isRunning():
                QMessageBox.information(
                    self,
                    "Task running",
                    "This task is still running. Wait for it to finish before closing this tab.",
                )
                return
            try:
                worker.cancel()
            except Exception:
                pass
        self.tabs.removeTab(index)
        widget.deleteLater()
        # Persist the reduced open-tab set so a later force-kill can't resurrect
        # a tab the user deliberately closed.
        self._persist_open_tabs()

    def _hide_fixed_close_buttons(self) -> None:
        """Hide close buttons on non-TaskTabs.

        Exception: in Wizard Mode the revealed Index singleton gets a visible
        "x" that re-hides (not destroys) it — see _on_tab_close_requested.
        """
        from PySide6.QtWidgets import QTabBar
        bar = self.tabs.tabBar()
        index_tab = getattr(self, "index_tab", None)
        mc = getattr(self, "mode_controller", None)
        is_wizard = bool(mc is not None and mc.is_wizard)
        for i in range(self.tabs.count()):
            widget = self.tabs.widget(i)
            is_task_tab = widget is not None and widget.property("wizard_task_id") is not None
            is_summary_tab = widget is not None and (
                widget.property("summary_browser_task_id") is not None
                or widget.property("summary_output_viewer_path") is not None
            )
            is_rehideable_index = (
                index_tab is not None
                and widget is index_tab
                and is_wizard
                and self.tabs.isTabVisible(i)
            )
            show_close = is_task_tab or is_summary_tab or is_rehideable_index
            for side in (QTabBar.ButtonPosition.RightSide, QTabBar.ButtonPosition.LeftSide):
                btn = bar.tabButton(i, side)
                if btn is not None:
                    btn.setVisible(show_close)

    def on_tree_item_clicked(self, item, column):
        if column == 1:  # Queued Tasks column (index 1)
            # Get click position relative to the tree's viewport
            pos = self.tree.visualItemRect(item).bottomLeft()
            pos.setX(self.tree.columnViewportPosition(1))
            global_pos = self.tree.viewport().mapToGlobal(pos)

            # Get all selected items (for multi-select support)
            selected_items = self.tree.selectedItems()

            # If clicked item is not in selection, use just the clicked item
            if item not in selected_items:
                selected_items = [item]

            self.show_agent_menu(selected_items, global_pos)

    def show_agent_menu(self, items, global_pos):
        """Show agent menu for one or more selected items."""
        menu = QMenu(self)

        # For multiple items, show count in header
        if len(items) > 1:
            header_action = QAction(f"Apply to {len(items)} selected items:", menu)
            header_action.setEnabled(False)
            menu.addAction(header_action)
            menu.addSeparator()

        # Determine which tasks are common across all items
        # A task is "checked" if ALL items have it, "partial" if some have it
        task_states = {}
        for agent in self.AGENTS:
            agent_id = agent["id"]
            has_count = sum(1 for item in items if agent_id in (item.data(0, Qt.ItemDataRole.UserRole + 2) or []))
            if has_count == len(items):
                task_states[agent_id] = "all"
            elif has_count > 0:
                task_states[agent_id] = "partial"
            else:
                task_states[agent_id] = "none"

        for agent in self.AGENTS:
            state = task_states[agent["id"]]
            action = QAction(agent["name"], menu)
            action.setCheckable(True)
            action.setChecked(state in ["all", "partial"])

            # Visual indication for partial selection
            if state == "partial":
                action.setText(f"{agent['name']} (partial)")

            action.triggered.connect(partial(self.toggle_agent_task_multi, items, agent["id"], state))
            menu.addAction(action)

        menu.addSeparator()
        clear_act = QAction("Clear All Tasks", menu)
        clear_act.triggered.connect(lambda: self.clear_tasks_multi(items))
        menu.addAction(clear_act)

        menu.exec(global_pos)

    def toggle_agent_task_multi(self, items, agent_id, current_state):
        """Toggle a task for multiple items. If partial/none, add to all. If all, remove from all."""
        self.tree.blockSignals(True)

        # If all items have it, remove from all. Otherwise, add to all.
        should_add = current_state != "all"

        for item in items:
            try:
                current_tasks = list(item.data(0, Qt.ItemDataRole.UserRole + 2) or [])
            except RuntimeError:
                continue
            if should_add:
                if agent_id not in current_tasks:
                    current_tasks.append(agent_id)
            else:
                if agent_id in current_tasks:
                    current_tasks.remove(agent_id)

            item.setData(0, Qt.ItemDataRole.UserRole + 2, current_tasks)
            self.update_item_tasks_ui(item)

        self.tree.blockSignals(False)

    def clear_tasks_multi(self, items):
        """Clear all tasks from multiple items."""
        self.tree.blockSignals(True)
        for item in items:
            try:
                item.setData(0, Qt.ItemDataRole.UserRole + 2, [])
                self.update_item_tasks_ui(item)
            except RuntimeError:
                continue
        self.tree.blockSignals(False)

    def toggle_agent_task(self, item, agent_id):
        """Toggle a task for a single item (legacy support)."""
        current_tasks = item.data(0, Qt.ItemDataRole.UserRole + 2) or []
        if agent_id in current_tasks:
            current_tasks.remove(agent_id)
        else:
            current_tasks.append(agent_id)
        self.set_item_tasks(item, current_tasks)

    def set_item_tasks(self, item, task_ids, recursive=True):
        self.tree.blockSignals(True)
        item.setData(0, Qt.ItemDataRole.UserRole + 2, task_ids)
        self.update_item_tasks_ui(item)
        
        if recursive:
            for i in range(item.childCount()):
                self._recursive_set_tasks(item.child(i), task_ids)
        self.tree.blockSignals(False)

    def _recursive_set_tasks(self, item, task_ids):
        item.setData(0, Qt.ItemDataRole.UserRole + 2, task_ids)
        self.update_item_tasks_ui(item)
        for i in range(item.childCount()):
            self._recursive_set_tasks(item.child(i), task_ids)

    def update_item_tasks_ui(self, item):
        task_ids = item.data(0, Qt.ItemDataRole.UserRole + 2) or []
        if not task_ids:
            item.setText(1, " [ + Add Tasks ]")
            item.setForeground(1, Qt.GlobalColor.gray)
            return

        # Create a visual string of short tags
        tags = []
        for tid in task_ids:
            agent = next((a for a in self.AGENTS if a["id"] == tid), None)
            if agent:
                tags.append(f"[{agent['short']}]")

        item.setText(1, " ".join(tags))
        item.setForeground(1, Qt.GlobalColor.blue)

    def filter_tree(self, text):
        search_text = text.lower()
        iterator = QTreeWidgetItemIterator(self.tree)
        
        # If empty, show all
        if not search_text:
            while iterator.value():
                item = iterator.value()
                item.setHidden(False)
                iterator += 1
            return

        # First pass: Hide all
        items = []
        while iterator.value():
            items.append(iterator.value())
            iterator += 1
            
        for item in items:
            item.setHidden(True)
            
        # Second pass: Show matches and parents
        for item in items:
            if search_text in item.text(0).lower():
                item.setHidden(False)
                parent = item.parent()
                while parent:
                    parent.setHidden(False)
                    parent.setExpanded(True) # Expand path to match
                    parent = parent.parent()

    def toggle_expand(self):
        if self.expand_btn.isChecked():
            self.tree.expandAll()
            self.expand_btn.setText("Collapse to Root")
        else:
            self.tree.collapseAll()
            # Expand only the top-level root item
            if self.tree.topLevelItemCount() > 0:
                self.tree.topLevelItem(0).setExpanded(True)
            self.expand_btn.setText("Expand All")

    def change_file(self):
        dialog = FileNumberDialog(self)
        if dialog.exec() == QDialog.DialogCode.Accepted:
            new_file_num = dialog.get_file_number()
            new_path = get_case_path(new_file_num)

            if new_path:
                # Keep the Change File dialog/hotkey path aligned with
                # load_case_by_number: wizard task tabs are case-scoped.
                try:
                    self._save_wizard_state_for_current_case()
                except Exception as e:
                    log_event(f"[wizard] snapshot failed: {e}")
                self._remove_all_task_tabs()

                self.save_status_history()
                # Save the persistent chat tab + any wizard-spawned chat tabs
                # against the OLD case before file_number changes.
                if hasattr(self, 'chat_tab') and self.chat_tab:
                    self.chat_tab.save_current_state()
                for _, tab in self._iter_task_tabs():
                    if isinstance(tab, ChatTab):
                        try:
                            tab.save_current_state()
                        except Exception as e:
                            log_event(f"[wizard] chat save_current_state failed: {e}")
                self.file_number = new_file_num
                self.case_path = new_path
                self.setWindowTitle(f"iCharlotte - {self.file_number} - {os.path.basename(self.case_path)}")
                self.populate_tree()
                # Clear status list to reset state
                self.clear_all_status()
                self.load_status_history()

                # Reset agent buttons, then restore running state for agents on this case
                for btn in self.agent_buttons.values():
                    btn.set_running(False)
                for script, case_num in self.running_agents.items():
                    if case_num == new_file_num and script in self.agent_buttons:
                        self.agent_buttons[script].set_running(True)

                # Rebind the persistent chat tab + any open wizard chat tabs
                # to the NEW case.
                if hasattr(self, 'chat_tab'):
                    self.chat_tab.load_case(self.file_number)
                for _, tab in self._iter_task_tabs():
                    if isinstance(tab, ChatTab):
                        try:
                            tab.load_case(self.file_number)
                        except Exception as e:
                            log_event(f"[wizard] chat load_case failed: {e}")

                # Reset Tabs for new case isolation
                if hasattr(self, 'index_tab'):
                    self.index_tab.load_data(self.file_number)
                if hasattr(self, 'discovery_tab'):
                    self.discovery_tab.load_case(new_file_num)

                # A loaded case always defaults to Wizard mode on the Wizard tab.
                self._default_to_wizard_mode()

                try:
                    self._restore_task_tabs_for_case()
                except Exception as e:
                    log_event(f"[wizard] restore failed: {e}")

                if hasattr(self, "wizard_tab") and self.wizard_tab is not None:
                    from icharlotte_core.ui.wizard.persistence import WizardStatePersistence
                    try:
                        p = WizardStatePersistence(self.case_path)
                        self.wizard_tab.refresh_recent_tasks(p.get_recent_tasks())
                    except Exception as e:
                        log_event(f"[wizard] refresh recent_tasks failed: {e}")

                log_event(f"Switched to case {new_file_num}")
            else:
                QMessageBox.critical(self, "Error", f"Could not find case directory for {new_file_num}")

    def open_root_folder(self):
        if not self.case_path:
            QMessageBox.information(self, "Info", "No case loaded.")
            return
        if os.path.exists(self.case_path):
            try:
                os.startfile(self.case_path)
            except Exception as e:
                QMessageBox.warning(self, "Error", f"Could not open directory: {e}")

    def _register_lazy_tabs(self):
        """Register deferred per-tab loaders, keyed by the tab widget.

        Each loader takes the target file_number and refreshes that tab. They
        are invoked lazily from ``_on_tab_changed`` the first time the tab is
        shown after a case switch, rather than eagerly during the switch.
        """
        self._lazy_loaders = {}          # widget -> loader(file_number)
        self._tabs_pending_reload = set()  # widgets awaiting reload for self.file_number

        candidates = [
            ('index_tab', lambda fn: self.index_tab.load_data(fn)),
            ('chat_tab', lambda fn: self.chat_tab.load_case(fn)),
            ('discovery_tab', lambda fn: self.discovery_tab.load_case(fn)),
        ]
        for attr, loader in candidates:
            widget = getattr(self, attr, None)
            if widget is not None:
                self._lazy_loaders[widget] = loader

    def _schedule_lazy_tab_reloads(self):
        """Mark every lazy tab as needing a reload for the current case and
        immediately load whichever lazy tab (if any) is currently visible.

        Loaders read ``self.file_number``, which the caller must already have
        pointed at the new case.
        """
        if not hasattr(self, '_lazy_loaders'):
            return
        self._tabs_pending_reload = set(self._lazy_loaders.keys())
        self._maybe_load_current_tab()

    def _maybe_load_current_tab(self):
        """If the currently-visible tab has a pending reload, run its loader."""
        if not getattr(self, '_tabs_pending_reload', None):
            return
        widget = self.tabs.currentWidget()
        if widget in self._tabs_pending_reload:
            self._tabs_pending_reload.discard(widget)
            loader = self._lazy_loaders.get(widget)
            if loader:
                try:
                    loader(self.file_number)
                except Exception as e:
                    log_event(f"[lazy-tab] load failed: {e}")

    def _on_tab_changed(self, index):
        self._maybe_load_current_tab()

    def load_case_by_number(self, file_number):
        import time as _time
        _t0 = _time.perf_counter()

        def _lap(label, since):
            now = _time.perf_counter()
            log_event(f"[switch-timing] {label}: {(now - since) * 1000:.0f} ms")
            return now

        log_debug(f"load_case_by_number: switching to {file_number}")
        new_path = get_case_path(file_number)
        if not new_path:
            QMessageBox.critical(self, "Error", f"Could not find case directory for {file_number}")
            return

        log_debug(f"load_case_by_number: path={new_path}")
        # ---- WIZARD: snapshot current case's task tabs, then remove them ----
        try:
            self._save_wizard_state_for_current_case()
        except Exception as e:
            log_event(f"[wizard] snapshot failed: {e}")
        self._remove_all_task_tabs()

        self.save_status_history()
        # Wizard-spawned chat tabs were already saved + removed by
        # _remove_all_task_tabs above. Save the persistent singleton too.
        if hasattr(self, 'chat_tab') and self.chat_tab:
            self.chat_tab.save_current_state()
        _t = _lap("save old-case state", _t0)
        self.file_number = file_number
        self.case_path = new_path
        self._update_window_title()
        self.populate_tree()
        _t = _lap("populate_tree", _t)
        self.clear_all_status()

        for btn in self.agent_buttons.values():
            btn.set_running(False)
        for script, case_num in self.running_agents.items():
            if case_num == file_number and script in self.agent_buttons:
                self.agent_buttons[script].set_running(True)

        self.load_status_history()
        _t = _lap("load_status_history", _t)

        # A loaded case always defaults to Wizard mode on the Wizard tab.
        self._default_to_wizard_mode()

        # Defer Index/Chat/Discovery reloads until their tab is actually shown
        # (none are visible right after the switch). Loads whichever lazy tab
        # is currently active, if any, immediately.
        self._schedule_lazy_tab_reloads()
        _t = _lap("schedule lazy tab reloads", _t)

        # ---- WIZARD: restore the new case's task tabs ----
        try:
            self._restore_task_tabs_for_case()
        except Exception as e:
            log_event(f"[wizard] restore failed: {e}")
        _t = _lap("restore task tabs", _t)

        # Refresh Wizard tab's Recent Tasks list for the new case.
        if hasattr(self, "wizard_tab") and self.wizard_tab is not None:
            from icharlotte_core.ui.wizard.persistence import WizardStatePersistence
            try:
                p = WizardStatePersistence(self.case_path)
                self.wizard_tab.refresh_recent_tasks(p.get_recent_tasks())
            except Exception as e:
                log_event(f"[wizard] refresh recent_tasks failed: {e}")

        _lap("TOTAL switch", _t0)
        log_event(f"Switched to case {self.file_number}")

    def view_docket(self):
        if not self.case_path:
            QMessageBox.information(self, "Info", "No case loaded.")
            return
        # Look in NOTES/AI OUTPUT for Docket_*.pdf
        ai_output_dir = os.path.join(self.case_path, "NOTES", "AI OUTPUT")
        if not os.path.exists(ai_output_dir):
            QMessageBox.information(self, "Info", "AI OUTPUT directory not found.")
            return

        dockets = glob.glob(os.path.join(ai_output_dir, "Docket_*.pdf"))
        if dockets:
            # Sort by modification time, newest first
            dockets.sort(key=os.path.getmtime, reverse=True)
            latest = dockets[0]
            try:
                os.startfile(latest)
            except Exception as e:
                QMessageBox.warning(self, "Error", f"Could not open docket: {e}")
        else:
            QMessageBox.information(self, "Info", "No Docket PDF found in AI OUTPUT.")

    def open_notes(self):
        """Open or create the AS NOTES document for the current case."""
        if not hasattr(self, 'case_path') or not self.case_path:
            QMessageBox.information(self, "Info", "No case loaded.")
            return
        if not hasattr(self, 'file_number') or not self.file_number:
            QMessageBox.information(self, "Info", "No case loaded.")
            return

        notes_dir = os.path.join(self.case_path, "NOTES")
        notes_filename = f"AS NOTES - {self.file_number}.docx"
        notes_path = os.path.join(notes_dir, notes_filename)

        # Create NOTES directory if it doesn't exist
        if not os.path.exists(notes_dir):
            try:
                os.makedirs(notes_dir)
            except Exception as e:
                QMessageBox.warning(self, "Error", f"Could not create NOTES directory: {e}")
                return

        # Create the document if it doesn't exist
        if not os.path.exists(notes_path):
            try:
                from docx import Document
                doc = Document()
                doc.add_heading(f"AS Notes - {self.file_number}", level=1)
                doc.add_paragraph("")
                doc.save(notes_path)
            except Exception as e:
                QMessageBox.warning(self, "Error", f"Could not create notes document: {e}")
                return

        # Open the document
        try:
            os.startfile(notes_path)
        except Exception as e:
            QMessageBox.warning(self, "Error", f"Could not open notes document: {e}")

    def manage_variables(self):
        dialog = VariablesDialog(self.file_number, self)
        dialog.exec()

    def manage_prompts(self):
        # Get current tab index to auto-select relevant agent
        current_tab_index = self.tabs.currentIndex()
        current_tab_name = self.tabs.tabText(current_tab_index)
        dialog = PromptsDialog(self, current_tab=current_tab_name)
        dialog.exec()

    def on_tree_double_click(self, item, column):
        path = item.data(0, Qt.ItemDataRole.UserRole)
        if path and os.path.exists(path) and os.path.isfile(path):
            try:
                os.startfile(path)
            except Exception as e:
                QMessageBox.warning(self, "Error", f"Could not open file: {e}")

    def on_tree_selection_changed(self):
        selected = self.tree.selectedItems()
        count = len(selected)
        if count > 1:
            # Count only files (not directories)
            file_count = sum(1 for item in selected if item.data(0, Qt.ItemDataRole.UserRole + 1) == "file")
            if file_count > 0:
                self.status_label.setText(f"{file_count} files selected - Click 'Queued Tasks' column to apply tasks to all selected")
            # Clear preview for multi-selection
            if hasattr(self, 'preview_pane') and self.preview_pane.isVisible():
                self.preview_pane.clear()
        elif count == 1:
            item = selected[0]
            path = item.data(0, Qt.ItemDataRole.UserRole)
            item_type = item.data(0, Qt.ItemDataRole.UserRole + 1)
            if path:
                self.status_label.setText(f"Selected: {os.path.basename(path)}")
                # Update preview pane if visible and it's a file
                if hasattr(self, 'preview_pane') and self.preview_pane.isVisible() and item_type == "file":
                    self.preview_pane.show_file(path)
        else:
            if hasattr(self, 'file_number') and self.file_number:
                self.status_label.setText(f"Case: {self.file_number}")
            if hasattr(self, 'preview_pane') and self.preview_pane.isVisible():
                self.preview_pane.clear()

    def run_separator_path(self, path, sensitivity=2):
        # Switch to Status Tab to show progress
        self._switch_to_status_tab()
        
        status_widget = StatusWidget("Separator Agent", f"Analyzing {os.path.basename(path)}")
        self.status_list_layout.insertWidget(0, status_widget)
        
        script_path = os.path.join(SCRIPTS_DIR, "separate.py")
        args = [script_path, "--headless", "--sensitivity", str(sensitivity), path]
        
        runner = AgentRunner(sys.executable, args, status_widget)
        self.agent_runners.append(runner)
        
        # Capture output for JSON extraction
        separator_output_container = {"text": ""}
        def collect_stdout(text):
            separator_output_container["text"] += text
            
        runner.log_update.connect(collect_stdout)
        
        def on_finished(success):
            if success:
                # Find JSON path
                match = re.search(r"JSON_MAP: (.+)", separator_output_container["text"])
                if match:
                    json_path = match.group(1).strip()
                    try:
                        with open(json_path, 'r') as f:
                            docs = json.load(f)
                        
                        # Add to Index Tab
                        self.index_tab.add_pdf(path, docs)
                        
                        # Cleanup temp file
                        try:
                            os.remove(json_path)
                        except Exception:
                            pass
                        
                    except Exception as e:
                        QMessageBox.critical(self, "Error", f"Failed to load separator result: {e}")
                else:
                     log_event(f"Warning: Could not find JSON output from separator for {path}", "warning")

            self.cleanup_runner(runner)

            # Re-enable sensitivity controls even on failure
            if hasattr(self, 'index_tab') and hasattr(self.index_tab, 'workbench'):
                self.index_tab.workbench.set_busy(False)

        runner.finished.connect(on_finished)
        runner.start()

    def create_agent_button(self, name, script, layout, arg_type="file_number", extra_flags=None):
        btn = QPushButton(name)
        btn.setFixedHeight(35)
        btn.setStyleSheet("font-size: 11px; padding: 2px;")
        btn.clicked.connect(partial(self.run_agent, name, script, arg_type, extra_flags))
        layout.addWidget(btn)

    def create_enhanced_agent_button(self, name, script, layout, arg_type="file_number", extra_flags=None):
        """Create an enhanced agent button with running indicator, status, and settings."""
        enhanced_btn = EnhancedAgentButton(name, script)
        enhanced_btn.clicked.connect(partial(self.run_enhanced_agent, name, script, arg_type, extra_flags, enhanced_btn))
        enhanced_btn.settings_clicked.connect(partial(self.open_agent_settings, script))

        # Store reference for updating status
        self.agent_buttons[script] = enhanced_btn

        # Update last docket download date for docket agent
        if script == "docket.py":
            self.update_docket_agent_status()

        layout.addWidget(enhanced_btn)

    def run_enhanced_agent(self, name, script, arg_type, extra_flags, btn_widget):
        """Run agent with enhanced status tracking."""
        # Set button to running state
        btn_widget.set_running(True)
        # Store the file_number this agent was started for
        started_for_case = self.file_number
        self.running_agents[script] = started_for_case

        # Run the agent
        runner = self.run_agent(name, script, arg_type, extra_flags)

        # Connect finished signal to update button
        if runner:
            runner.finished.connect(partial(self.on_agent_finished, script, btn_widget, started_for_case))
        else:
            # User cancelled (e.g., file picker dialog) - reset button state
            btn_widget.set_running(False)
            if script in self.running_agents:
                del self.running_agents[script]

    def on_agent_finished(self, script, btn_widget, started_for_case, success):
        """Handle agent completion and update button status."""
        try:
            # Clear the running state
            if script in self.running_agents:
                del self.running_agents[script]

            # Only update button UI if we're still on the same case the agent was started for
            if self.file_number == started_for_case:
                btn_widget.set_running(False)
                # Update status message
                if success:
                    btn_widget.set_status("Last: Just now")
                else:
                    btn_widget.set_status("Last: Failed")

            # Update docket download date for the ORIGINAL case (not current case)
            if script == "docket.py" and success and started_for_case:
                from datetime import datetime
                today = datetime.now().strftime("%Y-%m-%d")
                self.master_db.update_last_docket_download(started_for_case, today)
                # Only update button if still on the same case
                if self.file_number == started_for_case:
                    btn_widget.set_last_run(today)
        except Exception as e:
            log_event(f"Error in on_agent_finished: {e}", "error")

    def _run_subpoena_tracker(self):
        """Launch the in-process subpoena tracker worker."""
        if "subpoena_tracker" in self.running_agents:
            return  # Already running

        if not self.case_path:
            from PySide6.QtWidgets import QMessageBox
            QMessageBox.warning(self, "No Case", "No case is currently loaded.")
            return

        btn = self._subpoena_btn
        btn.set_running(True)
        started_for_case = self.file_number
        self.running_agents["subpoena_tracker"] = started_for_case

        worker = SubpoenaTrackerWorker(self.case_path, file_number=self.file_number)
        worker.progress.connect(lambda msg: log_event(f"Subpoena Tracker: {msg}"))
        worker.warning.connect(lambda msg: log_event(f"Subpoena Tracker warning: {msg}", "warning"))
        worker.finished_result.connect(
            lambda success, result: self._on_subpoena_tracker_finished(
                worker, btn, started_for_case, success, result
            )
        )
        self.agent_runners.append(worker)
        worker.start()

    def _on_subpoena_tracker_finished(self, worker, btn_widget, started_for_case, success, result):
        """Handle subpoena tracker completion."""
        try:
            if "subpoena_tracker" in self.running_agents:
                del self.running_agents["subpoena_tracker"]

            # Clean up the worker reference
            if worker in self.agent_runners:
                self.agent_runners.remove(worker)

            if self.file_number == started_for_case:
                btn_widget.set_running(False)
                if success:
                    btn_widget.set_status("Last: Just now")
                    log_event(f"Subpoena Tracker complete: {result}")
                else:
                    btn_widget.set_status("Last: Failed")
                    log_event(f"Subpoena Tracker failed: {result}", "error")
        except Exception as e:
            log_event(f"Error in _on_subpoena_tracker_finished: {e}", "error")

    def _run_med_record_extractor(self):
        """Open dialog and launch med record extraction worker."""
        if "med_record_extractor" in self.running_agents:
            return
        if not self.case_path:
            from PySide6.QtWidgets import QMessageBox
            QMessageBox.warning(self, "No Case", "No case is currently loaded.")
            return

        from icharlotte_core.ui.med_record_extractor_dialog import MedRecordExtractorDialog
        dialog = MedRecordExtractorDialog(parent=self)
        if not dialog.exec():
            return

        user_text = dialog.get_text()
        if not user_text:
            return

        btn = self._med_extract_btn
        btn.set_running(True)
        started_for_case = self.file_number
        self.running_agents["med_record_extractor"] = started_for_case

        # Add status widget to Status tab
        from icharlotte_core.ui.widgets import StatusWidget
        status_widget = StatusWidget("Med Record Extractor", f"Case {self.file_number}")
        self.status_list_layout.insertWidget(0, status_widget)

        from icharlotte_core.med_record_extractor import MedRecordExtractorWorker
        worker = MedRecordExtractorWorker(self.case_path, self.file_number, user_text)
        worker.progress.connect(lambda msg: (
            log_event(f"Med Record Extractor: {msg}"),
            status_widget.append_log(msg + "\n"),
            status_widget.status_text_label.setText(msg),
        ))
        worker.warning.connect(lambda msg: (
            log_event(f"Med Record Extractor warning: {msg}", "warning"),
            status_widget.append_log(f"WARNING: {msg}\n"),
        ))
        worker.finished_result.connect(
            lambda success, result: self._on_med_extract_finished(
                worker, btn, started_for_case, success, result, status_widget
            )
        )
        self.agent_runners.append(worker)
        worker.start()

    def _on_med_extract_finished(self, worker, btn_widget, started_for_case, success, result, status_widget=None):
        """Handle med record extractor completion."""
        try:
            if "med_record_extractor" in self.running_agents:
                del self.running_agents["med_record_extractor"]
            if worker in self.agent_runners:
                self.agent_runners.remove(worker)
            if self.file_number == started_for_case:
                btn_widget.set_running(False)
                if success:
                    btn_widget.set_status("Last: Just now")
                    log_event(f"Med Record Extractor complete: {result}")
                    if status_widget:
                        status_widget.append_log(f"\n{result}\n")
                        status_widget.set_finished(True)
                        # Set output to the extracts folder
                        import os
                        out_dir = os.path.join(self.case_path, "NOTES", "AI OUTPUT", "Med Record Extracts")
                        if os.path.isdir(out_dir):
                            status_widget.set_output_file(out_dir)
                else:
                    btn_widget.set_status("Last: Failed")
                    log_event(f"Med Record Extractor failed: {result}", "error")
                    if status_widget:
                        status_widget.append_log(f"\nFAILED: {result}\n")
                        status_widget.set_finished(False)
        except Exception as e:
            log_event(f"Error in _on_med_extract_finished: {e}", "error")

    def update_docket_agent_status(self):
        """Update the docket agent button with last download date."""
        if not self.file_number or "docket.py" not in self.agent_buttons:
            return

        case_data = self.master_tab.db.get_case(self.file_number)
        if case_data:
            last_download = case_data.get("last_docket_download")
            if last_download:
                self.agent_buttons["docket.py"].set_last_run(last_download)
            else:
                self.agent_buttons["docket.py"].set_status("Never downloaded")

    def open_agent_settings(self, script):
        """Open the settings dialog for an agent."""
        dialog = AgentSettingsDialog(script, self.agent_settings_db, self)
        dialog.exec()

    def open_output_browser(self):
        """Open the output browser dialog."""
        if not self.case_path or not self.file_number:
            QMessageBox.warning(self, "No Case", "Please load a case first.")
            return

        dialog = OutputBrowserWidget(self.case_path, self.file_number, self)
        dialog.exec()

    def open_processing_log(self):
        """Open the processing log dialog."""
        if not self.file_number:
            QMessageBox.warning(self, "No Case", "Please load a case first.")
            return

        dialog = ProcessingLogWidget(self.file_number, self)
        dialog.exec()

    def open_report_generator(self):
        """Open the report generator dialog and start pipeline if accepted."""
        if not self.file_number:
            QMessageBox.warning(self, "No Case", "Please load a case first.")
            return

        dialog = ReportGeneratorDialog(self.file_number, self)
        if dialog.exec() != QDialog.DialogCode.Accepted:
            return

        config = dialog.get_config()

        # Create status widget
        status_widget = StatusWidget(
            "Report Generator",
            f"{config['report_type']} for {self.file_number}"
        )
        self.status_list_layout.insertWidget(0, status_widget)

        # Create and start worker
        worker = ReportPipelineWorker(config, self)
        self.agent_runners.append(worker)

        # Connect signals to status widget
        worker.stage_started.connect(status_widget.update_pass_start)
        worker.stage_completed.connect(status_widget.update_pass_complete)
        worker.stage_failed.connect(status_widget.update_pass_failed)
        worker.progress_update.connect(status_widget.update_progress)
        worker.log_update.connect(lambda msg: status_widget.append_log(msg))
        worker.output_file.connect(status_widget.set_output_file)
        worker.finished_signal.connect(lambda ok: status_widget.set_finished(ok))
        worker.finished_signal.connect(lambda: self._cleanup_report_worker(worker))

        # Switch to status tab area
        self._switch_to_status_tab()

        worker.start()

    def _cleanup_report_worker(self, worker):
        """Clean up finished report pipeline worker."""
        try:
            if worker in self.agent_runners:
                self.agent_runners.remove(worker)
        except Exception as e:
            log_event(f"Error cleaning up report worker: {e}", "error")

    def toggle_advanced_filters(self, checked):
        """Toggle visibility of advanced filter panel."""
        if checked:
            self.advanced_filter.show()
            self.filter_toggle_btn.setText("▲ Filters")
            # Update available tags
            if self.file_number:
                tags_db = FileTagsDB(self.file_number)
                self.advanced_filter.set_available_tags(tags_db.get_all_tags())
        else:
            self.advanced_filter.hide()
            self.filter_toggle_btn.setText("▼ Filters")

    def toggle_preview_pane(self, checked):
        """Toggle visibility of file preview pane."""
        if checked:
            self.preview_pane.show()
            current = self.main_splitter.sizes()
            if len(current) >= 3:
                total = sum(current)
                # Load saved preview width or use default (25% of total)
                saved_width = self._load_preview_width()
                if saved_width and saved_width > 50:
                    preview_width = min(saved_width, total - 250)  # Leave room for left and center
                else:
                    preview_width = int(total * 0.25)
                center_width = total - 180 - preview_width
                self.main_splitter.setSizes([180, center_width, preview_width])
            # Show currently selected file in preview
            selected = self.tree.selectedItems()
            if len(selected) == 1:
                item = selected[0]
                path = item.data(0, Qt.ItemDataRole.UserRole)
                item_type = item.data(0, Qt.ItemDataRole.UserRole + 1)
                if path and item_type == "file":
                    self.preview_pane.show_file(path)
        else:
            self.preview_pane.hide()
            self.preview_pane.clear()

    def _on_splitter_moved(self, pos, index):
        """Save preview pane width when splitter is moved."""
        if self.preview_pane.isVisible():
            sizes = self.main_splitter.sizes()
            if len(sizes) >= 3 and sizes[2] > 50:
                self._save_preview_width(sizes[2])

    def _save_preview_width(self, width):
        """Save preview pane width to settings."""
        settings = QSettings("iCharlotte", "iCharlotte")
        settings.setValue("preview_pane_width", width)

    def _load_preview_width(self):
        """Load preview pane width from settings."""
        settings = QSettings("iCharlotte", "iCharlotte")
        return settings.value("preview_pane_width", type=int)

    def _restore_window_state(self):
        """Restore window geometry, splitter layout, and tree columns from the previous session."""
        settings = QSettings("iCharlotte", "iCharlotte")
        geometry = settings.value("main_window_geometry")
        if geometry is not None:
            self.restoreGeometry(geometry)
        sizes = settings.value("main_splitter_sizes")
        if isinstance(sizes, (list, tuple)) and sizes:
            try:
                self.main_splitter.setSizes([int(s) for s in sizes])
            except (TypeError, ValueError):
                pass
        widths = settings.value("case_view_tree_column_widths")
        if isinstance(widths, (list, tuple)) and hasattr(self, "tree"):
            try:
                for col, w in enumerate(widths[: self.tree.columnCount()]):
                    if int(w) > 0:
                        self.tree.setColumnWidth(col, int(w))
            except (TypeError, ValueError):
                pass

    def _restore_last_tab(self):
        """Reopen the tab that was active when the app last closed.

        Tabs are matched by name, not index — indices shift between Wizard and
        Advanced mode. Skipped when the saved tab is hidden in the current mode.
        """
        settings = QSettings("iCharlotte", "iCharlotte")
        tab_name = settings.value("last_main_tab", type=str)
        if not tab_name:
            return
        idx = self._index_of_tab(tab_name)
        if idx >= 0 and self.tabs.isTabVisible(idx):
            self.tabs.setCurrentIndex(idx)

    def _save_window_state(self):
        """Persist window geometry, splitter layout, tree columns, and active tab."""
        settings = QSettings("iCharlotte", "iCharlotte")
        settings.setValue("main_window_geometry", self.saveGeometry())
        if hasattr(self, "main_splitter"):
            settings.setValue("main_splitter_sizes", self.main_splitter.sizes())
        if hasattr(self, "tree"):
            settings.setValue(
                "case_view_tree_column_widths",
                [self.tree.columnWidth(c) for c in range(self.tree.columnCount())],
            )
        if hasattr(self, "tabs") and self.tabs.currentIndex() >= 0:
            settings.setValue("last_main_tab", self.tabs.tabText(self.tabs.currentIndex()))

    def apply_advanced_filters(self, filters):
        """Apply advanced filters to the file tree."""
        iterator = QTreeWidgetItemIterator(self.tree)

        while iterator.value():
            item = iterator.value()
            file_path = item.data(0, Qt.ItemDataRole.UserRole)
            item_type = item.data(0, Qt.ItemDataRole.UserRole + 1)

            if item_type == "file" and file_path:
                should_show = self._file_matches_filters(file_path, filters)
                item.setHidden(not should_show)
            elif item_type == "dir":
                # Directories stay visible if any child is visible
                pass

            iterator += 1

        # Update directory visibility based on children
        self._update_directory_visibility()

    def _file_matches_filters(self, file_path, filters):
        """Check if a file matches the given filters."""
        from datetime import datetime, timedelta

        # File type filter
        ext = os.path.splitext(file_path)[1].lower().lstrip('.')
        file_types = filters.get("file_types", [])

        if file_types:
            type_match = False
            if ext in ["pdf"] and "pdf" in file_types:
                type_match = True
            elif ext in ["doc", "docx"] and any(t in file_types for t in ["doc", "docx"]):
                type_match = True
            elif ext in ["jpg", "jpeg", "png", "gif", "bmp", "tiff"] and any(t in file_types for t in ["jpg", "jpeg", "png", "gif", "bmp", "tiff"]):
                type_match = True
            elif "other" in file_types and ext not in ["pdf", "doc", "docx", "jpg", "jpeg", "png", "gif", "bmp", "tiff"]:
                type_match = True

            if not type_match:
                return False

        # Processing status filter
        status_filter = filters.get("processing_status", [])
        if status_filter and "all" not in status_filter:
            if self.file_number:
                proc_log = ProcessingLogDB(self.file_number)
                file_status = proc_log.get_file_processing_status(file_path)

                has_summary = any("summar" in log.get("task_type", "").lower() and log.get("status") == "success" for log in file_status)
                is_unprocessed = not file_status or not any(log.get("status") == "success" for log in file_status)

                status_match = False
                if "unprocessed" in status_filter and is_unprocessed:
                    status_match = True
                if "summarized" in status_filter and has_summary:
                    status_match = True

                if not status_match:
                    return False

        # Date filter
        date_filter = filters.get("date_filter", "Any time")
        if date_filter != "Any time":
            try:
                mtime = os.path.getmtime(file_path)
                file_date = datetime.fromtimestamp(mtime)
                now = datetime.now()

                if date_filter == "Today":
                    if file_date.date() != now.date():
                        return False
                elif date_filter == "Last 7 days":
                    if (now - file_date).days > 7:
                        return False
                elif date_filter == "Last 30 days":
                    if (now - file_date).days > 30:
                        return False
                elif date_filter == "Last 90 days":
                    if (now - file_date).days > 90:
                        return False
                elif date_filter == "Custom range...":
                    date_from = filters.get("date_from")
                    date_to = filters.get("date_to")
                    if date_from:
                        from_dt = datetime.strptime(date_from, "%Y-%m-%d")
                        if file_date < from_dt:
                            return False
                    if date_to:
                        to_dt = datetime.strptime(date_to, "%Y-%m-%d")
                        if file_date > to_dt:
                            return False
            except Exception:
                pass

        # Tag filter
        tag_filter = filters.get("tag")
        if tag_filter and self.file_number:
            tags_db = FileTagsDB(self.file_number)
            file_tags = tags_db.get_tags(file_path)
            if tag_filter not in file_tags:
                return False

        return True

    def _update_directory_visibility(self):
        """Update directory visibility based on visible children."""
        def check_item(item):
            if item.data(0, Qt.ItemDataRole.UserRole + 1) == "dir":
                has_visible_child = False
                for i in range(item.childCount()):
                    child = item.child(i)
                    if not child.isHidden():
                        has_visible_child = True
                        break
                    # Recursively check nested directories
                    if child.data(0, Qt.ItemDataRole.UserRole + 1) == "dir":
                        check_item(child)
                        if not child.isHidden():
                            has_visible_child = True

                item.setHidden(not has_visible_child)

        for i in range(self.tree.topLevelItemCount()):
            check_item(self.tree.topLevelItem(i))

    def _count_running_agents(self):
        """Count currently running subprocess agents (excludes queued ones)."""
        return sum(
            1 for r in self.agent_runners
            if getattr(r, 'success', None) is None and r not in self._agent_queue
        )

    def _check_system_memory_ok(self):
        """Return True if system memory is below critical threshold."""
        try:
            import psutil
            mem = psutil.virtual_memory()
            available_gb = mem.available / (1024 ** 3)
            if available_gb < 1.0:
                log_warning(f"System memory critically low: {available_gb:.1f}GB available")
                return False
            return True
        except ImportError:
            return True  # Can't check, allow it

    def _drain_agent_queue(self):
        """Start queued agents if slots are available."""
        while self._agent_queue and self._count_running_agents() < self.MAX_CONCURRENT_AGENTS:
            if not self._check_system_memory_ok():
                log_warning("Pausing agent queue drain — system memory low")
                break
            runner = self._agent_queue.pop(0)
            log_debug(f"Dequeuing agent: {runner.status_widget.task_id} — running={self._count_running_agents()}")
            runner.status_widget.update_progress(0, "Starting...")
            runner.start()

    def add_status_task(self, name, details, command, args, script_name=None, file_path=None):
        status_widget = StatusWidget(name, details)
        self.status_list_layout.insertWidget(0, status_widget) # Add to top

        runner = AgentRunner(command, args, status_widget, task_id=status_widget.task_id, file_number=self.file_number)
        self.agent_runners.append(runner) # Keep alive
        runner.finished.connect(lambda: self.cleanup_runner(runner))
        runner.finished.connect(lambda _=None: self._drain_agent_queue())

        # Wire the READY button for the deposition agent's interactive flow.
        # AgentRunner.awaiting_input fires when phase 1 emits AWAITING_INPUT:<path>,
        # which makes StatusWidget show the READY button. Clicking it emits ready_clicked.
        if script_name == "summarize_deposition.py":
            status_widget.ready_clicked.connect(
                lambda session_path, runner=runner, widget=status_widget:
                    self._open_depo_summary_dialog(session_path, runner, widget)
            )

        # Store retry info on widget for serialization
        status_widget._retry_command = command
        status_widget._retry_args = list(args)
        status_widget._retry_script_name = script_name
        status_widget._retry_file_path = file_path

        # Retry button — re-launches the same task
        status_widget.retry_requested.connect(
            lambda: self.add_status_task(name, details, command, list(args), script_name, file_path)
        )

        # Log processing entry when agent finishes
        if script_name and file_path and self.file_number:
            runner.finished.connect(lambda success: self._log_processing_entry(script_name, file_path, success))

        running = self._count_running_agents()
        log_debug(f"add_status_task: {name} — total_runners={len(self.agent_runners)}, running={running}, queued={len(self._agent_queue)}")

        # Check concurrency limit and system memory
        if running >= self.MAX_CONCURRENT_AGENTS:
            self._agent_queue.append(runner)
            status_widget.update_progress(0, f"Queued (#{len(self._agent_queue)}) — waiting for slot...")
            log_info(f"Agent '{name}' queued — {running} already running (max {self.MAX_CONCURRENT_AGENTS})")
            return runner

        if not self._check_system_memory_ok():
            self._agent_queue.append(runner)
            status_widget.update_progress(0, "Queued — system memory low")
            log_warning(f"Agent '{name}' queued due to low system memory")
            return runner

        runner.start()
        return runner

    def _open_depo_summary_dialog(self, session_path, agent_runner, status_widget):
        """Open the deposition summary config dialog. On Accept, resume phase 2."""
        from icharlotte_core.ui.depo_summary_config_dialog import DepoSummaryConfigDialog
        try:
            dlg = DepoSummaryConfigDialog(session_path, parent=self)
        except Exception as e:
            log_error(f"Failed to open deposition config dialog: {e}")
            return

        if dlg.exec() == QDialog.Accepted:
            status_widget.clear_ready_state()
            try:
                agent_runner.resume_with_config(session_path)
            except Exception as e:
                log_error(f"Failed to resume phase 2: {e}")
        # On Cancel, leave the READY button visible so the user can re-open the dialog.

    def _log_processing_entry(self, script_name, file_path, success):
        """Log a processing entry when an agent finishes.

        Instead of triggering a full tree rebuild (which destroys all QTreeWidgetItems
        and causes segfaults when other code holds stale references), we do a targeted
        update of just the processing status column for the affected file.
        """
        try:
            if not self.file_number:
                return
            proc_log = ProcessingLogDB(self.file_number)
            status = "success" if success else "failed"
            proc_log.add_entry(file_path, script_name, status)
            log_debug(f"_log_processing_entry: {script_name} {status} for {os.path.basename(file_path)} — updating status column (no tree rebuild)")

            # Refresh the cached proc log so status queries see the new entry
            if hasattr(self, '_cached_proc_log') and self._cached_proc_log is not None:
                self._cached_proc_log = ProcessingLogDB(self.file_number)

            # Targeted update: just update column 4 (processing status) for this file
            item = self.tree_item_map.get(file_path)
            if item is not None:
                try:
                    new_status = self._get_file_processing_status(file_path)
                    item.setText(4, new_status)
                    log_debug(f"_log_processing_entry: updated status for {os.path.basename(file_path)} → '{new_status}'")
                except RuntimeError:
                    log_warning(f"_log_processing_entry: stale tree item for {os.path.basename(file_path)}")
            else:
                log_debug(f"_log_processing_entry: file not in tree_item_map, skipping status update")
        except Exception as e:
            log_error(f"Error logging processing entry: {e}")
            log_event(f"Error logging processing entry: {e}", "error")

    def cleanup_runner(self, runner):
        try:
            # Only cleanup if it matches the current file number (meaning the user saw it finish)
            if getattr(runner, 'file_number', None) == self.file_number:
                if runner in self.agent_runners:
                    self.agent_runners.remove(runner)
        except Exception as e:
            log_event(f"Error in cleanup_runner: {e}", "error")

    def clear_completed_status(self):
        for i in range(self.status_list_layout.count() - 1, -1, -1):
            item = self.status_list_layout.itemAt(i)
            widget = item.widget()
            if widget:
                if widget.is_finished:
                    widget.deleteLater()
                    
    def clear_all_status(self):
        # Disconnect all runners from their widgets before deleting
        # to prevent signals being sent to deleted widgets
        for runner in self.agent_runners:
            if hasattr(runner, 'disconnect_widget'):
                runner.disconnect_widget()

        for i in range(self.status_list_layout.count() - 1, -1, -1):
            item = self.status_list_layout.itemAt(i)
            widget = item.widget()
            if widget:
                widget.deleteLater()

    def save_status_history(self):
        if not self.file_number:
            return

        history = []
        # Loop from 0 to count-1 (top to bottom)
        for i in range(self.status_list_layout.count()):
            item = self.status_list_layout.itemAt(i)
            widget = item.widget()
            if isinstance(widget, StatusWidget):
                history.append(widget.to_dict())

        if not os.path.exists(GEMINI_DATA_DIR):
            os.makedirs(GEMINI_DATA_DIR, exist_ok=True)

        save_path = os.path.join(GEMINI_DATA_DIR, f"{self.file_number}_status_history.json")
        try:
            with open(save_path, 'w') as f:
                json.dump(history, f, indent=2)
            log_event(f"Saved status history to {save_path}")
        except Exception as e:
            log_event(f"Error saving status history: {e}", "error")

    def load_status_history(self):
        if not self.file_number:
            return

        save_path = os.path.join(GEMINI_DATA_DIR, f"{self.file_number}_status_history.json")
        if not os.path.exists(save_path):
            return

        try:
            with open(save_path, 'r') as f:
                history = json.load(f)
            
            for item_data in history:
                widget = StatusWidget.from_dict(item_data)

                # Restore retry info from serialized data, or reconstruct from agent name + filename
                retry_cmd = item_data.get("retry_command")
                retry_args = item_data.get("retry_args")
                retry_script = item_data.get("retry_script_name")
                retry_fp = item_data.get("retry_file_path")

                if not retry_cmd or not retry_args:
                    # Reconstruct from agent name and details (filename)
                    agent_name = item_data.get("agent_name", "")
                    filename = item_data.get("details", "")
                    agent = next((a for a in self.AGENTS if a["name"] == agent_name), None)
                    if agent and filename and self.case_path:
                        # Search for the file in the case folder
                        for root, dirs, files in os.walk(self.case_path):
                            if filename in files:
                                retry_fp = os.path.join(root, filename)
                                break
                        if retry_fp:
                            retry_cmd = sys.executable
                            retry_script = agent["script"]
                            retry_args = [os.path.join(SCRIPTS_DIR, retry_script), retry_fp]

                if retry_cmd and retry_args:
                    widget._retry_command = retry_cmd
                    widget._retry_args = retry_args
                    widget._retry_script_name = retry_script
                    widget._retry_file_path = retry_fp
                    name = item_data.get("agent_name", "Unknown")
                    details = item_data.get("details", "")
                    widget.retry_requested.connect(
                        lambda _n=name, _d=details, _c=retry_cmd, _a=retry_args, _s=retry_script, _f=retry_fp:
                            self.add_status_task(_n, _d, _c, list(_a), _s, _f)
                    )

                # Try to reconnect to running agent
                reconnected = False
                if not widget.is_finished:
                    task_id = getattr(widget, 'task_id', None)
                    if task_id:
                        for runner in self.agent_runners:
                            if getattr(runner, 'task_id', None) == task_id and hasattr(runner, 'reconnect_widget'):
                                runner.reconnect_widget(widget)
                                reconnected = True
                                if getattr(runner, 'success', None) is not None:
                                     self.agent_runners.remove(runner)
                                break

                    if not reconnected:
                        # Mark as interrupted if we couldn't find the runner
                        widget.status_text_label.setText(widget.status_text_label.text() + " (Interrupted)")
                        widget.status_text_label.setStyleSheet(f"color: {theme.WARNING}; font-weight: bold;")
                        widget.is_finished = True

                self.status_list_layout.addWidget(widget)
                
            log_event(f"Loaded {len(history)} status items from history")
        except Exception as e:
            log_event(f"Error loading status history: {e}", "error")

    def _persist_session_state(self) -> None:
        """Flush all per-session state to disk.

        Called from BOTH closeEvent and restart_app so the two shutdown paths
        cannot drift. The in-app Restart button uses QApplication.quit(), which
        does NOT dispatch closeEvent — so without calling this first, the open
        wizard task tabs and their contents are silently lost on restart. Each
        save is isolated so one failure can't block the others on the way out.
        """
        try:
            self._save_wizard_state_for_current_case()
        except Exception as e:
            log_event(f"[wizard] state save failed: {e}")
        try:
            self.save_status_history()
        except Exception as e:
            log_event(f"[wizard] status-history save failed: {e}")
        # Save persistent chat conversation + any wizard-spawned chat tabs.
        if hasattr(self, 'chat_tab') and self.chat_tab:
            try:
                self.chat_tab.save_current_state()
            except Exception as e:
                log_event(f"[wizard] chat save_current_state failed: {e}")
        for _, tab in self._iter_task_tabs():
            if isinstance(tab, ChatTab):
                try:
                    tab.save_current_state()
                except Exception as e:
                    log_event(f"[wizard] chat save_current_state failed: {e}")
        # Save discovery tab state.
        if hasattr(self, 'discovery_tab') and self.discovery_tab:
            try:
                self.discovery_tab.save_state()
            except Exception as e:
                log_event(f"[wizard] discovery save_state failed: {e}")

    def closeEvent(self, event):
        self._persist_session_state()
        self._save_window_state()
        # Clean up global hotkeys
        if KEYBOARD_AVAILABLE:
            try:
                keyboard.unhook_all_hotkeys()
                stop_word_hotkey()
                log_event("Global hotkeys unregistered")
            except Exception as e:
                log_event(f"Error unregistering hotkeys: {e}", "error")
        super().closeEvent(event)

    def run_agent(self, name, script, arg_type, extra_flags):
        log_event(f"Button clicked: {name}")
        script_path = os.path.join(SCRIPTS_DIR, script)
        if not os.path.exists(script_path):
            QMessageBox.critical(self, "Error", f"Script not found: {script_path}")
            return

        args = [script_path] # First arg for python is script path
        
        details = ""
        if arg_type == "file_number":
            args.append(self.file_number)
            details = f"File: {self.file_number}"

        elif arg_type == "file_picker":
            # Show file picker dialog, starting in the case directory
            start_dir = self.case_path if hasattr(self, 'case_path') and self.case_path else ""
            files, _ = QFileDialog.getOpenFileNames(
                self,
                f"Select files for {name}",
                start_dir,
                "Documents (*.pdf *.docx *.txt);;All Files (*.*)"
            )
            if not files:
                return  # User cancelled
            args.extend(files)
            details = f"{len(files)} file(s) selected"

        if extra_flags:
            args.extend(extra_flags)
            
        is_interactive = extra_flags and "--interactive" in extra_flags
        
        if is_interactive:
            try:
                creation_flags = 0x00000010 if os.name == 'nt' else 0
                subprocess.Popen([sys.executable] + args, creationflags=creation_flags)
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Failed to launch: {e}")
            return None
        else:
            if script in ["docket.py", "complaint.py"]:
                args.append("--headless")

            return self.add_status_task(name, details, sys.executable, args)

    def process_checked_items(self):
        log_event("Processing checked items...")
        
        count = 0
        iterator = QTreeWidgetItemIterator(self.tree)
        while iterator.value():
            item = iterator.value()
            path = item.data(0, Qt.ItemDataRole.UserRole)
            item_type = item.data(0, Qt.ItemDataRole.UserRole + 1)
            task_ids = item.data(0, Qt.ItemDataRole.UserRole + 2) or []
            
            if not path or item_type != "file" or not task_ids:
                iterator += 1
                continue

            for tid in task_ids:
                agent = next((a for a in self.AGENTS if a["id"] == tid), None)
                if not agent: continue
                
                if tid == "separate":
                    self.run_separator_path(path)
                    count += 1
                else:
                    filename = os.path.basename(path)
                    display_name = agent["name"]
                    details = f"{filename}"
                    script_name = agent["script"]

                    args = [os.path.join(SCRIPTS_DIR, script_name), path]
                    self.add_status_task(display_name, details, sys.executable, args, script_name=script_name, file_path=path)
                    count += 1
            
            iterator += 1

        self.clear_all_checkboxes() # Clears the task data and UI

        if count > 0:
            self._switch_to_status_tab()
        else:
            QMessageBox.warning(self, "No Selection", "No files selected for processing. (Did you queue any tasks?)")

    def load_cache(self):
        cache_path = os.path.join(GEMINI_DATA_DIR, f"{self.file_number}_tree.json")
        if not os.path.exists(cache_path):
            return False

        try:
            with open(cache_path, 'r') as f:
                data = json.load(f)

            entries = sorted(data, key=lambda x: len(x['path']))

            for entry in entries:
                path = entry['path']
                if path == self.case_path: continue
                
                parent_path = os.path.dirname(path)
                parent_item = self.tree_item_map.get(parent_path)
                
                if parent_item:
                    item = QTreeWidgetItem(parent_item)
                    item.setText(0, os.path.basename(path))
                    item.setData(0, Qt.ItemDataRole.UserRole, path)
                    item.setData(0, Qt.ItemDataRole.UserRole + 1, entry['type'])

                    if entry['type'] == 'dir':
                        item.setIcon(0, self._get_cached_icon(path, is_dir=True))
                        item.setExpanded(False)
                    else:
                        item.setIcon(0, self._get_cached_icon(path))
                        item.setText(2, entry.get('size_str', ''))

                        # Convert cached date to standard format
                        date_str = format_date_to_mm_dd_yyyy(entry.get('date_str', ''))
                        item.setText(3, date_str)
                    
                    # Restore tasks
                    task_ids = entry.get('task_ids', [])
                    item.setData(0, Qt.ItemDataRole.UserRole + 2, task_ids)
                    self.update_item_tasks_ui(item)
                        
                    self.tree_item_map[path] = item
            return True
        except Exception as e:
            log_event(f"Error loading cache: {e}", "error")
            return False

    def save_cache(self):
        log_debug(f"save_cache: starting, map_size={len(self.tree_item_map)}, case={self.file_number}")
        data = []
        _deleted_count = 0
        for path, item in self.tree_item_map.items():
            if path == self.case_path: continue

            try:
                entry = {
                    'path': path,
                    'type': item.data(0, Qt.ItemDataRole.UserRole + 1),
                    'size_str': item.text(2),
                    'date_str': item.text(3),
                    'task_ids': item.data(0, Qt.ItemDataRole.UserRole + 2)
                }
                data.append(entry)
            except RuntimeError:
                # Item's C++ object was already deleted (e.g., parent removed)
                _deleted_count += 1
                continue

        if _deleted_count:
            log_warning(f"save_cache: {_deleted_count} deleted C++ items skipped")

        if not os.path.exists(GEMINI_DATA_DIR):
            os.makedirs(GEMINI_DATA_DIR)

        cache_path = os.path.join(GEMINI_DATA_DIR, f"{self.file_number}_tree.json")
        try:
            with open(cache_path, 'w') as f:
                json.dump(data, f)
            log_debug(f"save_cache: wrote {len(data)} items to {cache_path}")
        except Exception as e:
            log_event(f"Error saving cache: {e}", "error")

    def _get_cached_icon(self, file_path, is_dir=False):
        """Get icon from cache based on file extension to avoid network access."""
        if is_dir:
            return self._folder_icon

        # Get extension and check cache
        ext = os.path.splitext(file_path)[1].lower()
        if ext in self._icon_cache:
            return self._icon_cache[ext]

        # For first occurrence of each extension, get icon from a local temp file
        # This avoids accessing network files while still getting proper icons
        if ext in ['.pdf', '.doc', '.docx', '.xls', '.xlsx', '.ppt', '.pptx',
                   '.txt', '.msg', '.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff',
                   '.mp3', '.mp4', '.avi', '.mov', '.zip', '.rar', '.7z']:
            # Use the icon provider with just the extension info
            icon = self.icon_provider.icon(QFileInfo(f"dummy{ext}"))
            self._icon_cache[ext] = icon
            return icon

        # Default file icon for unknown extensions
        self._icon_cache[ext] = self._file_icon
        return self._file_icon

    def _get_file_processing_status(self, file_path):
        """Get processing status for a file."""
        if not hasattr(self, 'file_number') or not self.file_number:
            return ""

        try:
            # Use cached processing log instead of creating new instance per file
            if not hasattr(self, '_cached_proc_log') or self._cached_proc_log is None:
                return ""

            logs = self._cached_proc_log.get_file_processing_status(file_path)
            if not logs:
                return ""

            # Get unique successful task types
            successful = set()
            for log in logs:
                if log.get("status") == "success":
                    task = log.get("task_type", "").lower()
                    if "discovery" in task or "disc" in task:
                        successful.add("S-DISC")
                    elif "deposition" in task or "depo" in task:
                        successful.add("S-DEPO")
                    elif "med_record" in task or "med_rec" in task:
                        successful.add("MED-REC")
                    elif "summar" in task:
                        successful.add("SUM")
                    elif "chron" in task:
                        successful.add("CHR")

            return ", ".join(sorted(successful)) if successful else ""
        except Exception:
            return ""

    def _get_file_tags(self, file_path):
        """Get tags for a file."""
        if not hasattr(self, 'file_number') or not self.file_number:
            return ""

        try:
            # Use cached tags db instead of creating new instance per file
            if not hasattr(self, '_cached_tags_db') or self._cached_tags_db is None:
                return ""

            tags = self._cached_tags_db.get_tags(file_path)
            return ", ".join(tags) if tags else ""
        except Exception:
            return ""

    def _schedule_tree_refresh(self):
        """Debounced tree refresh — coalesces rapid calls into one."""
        self._tree_refresh_timer.start(500)

    def populate_tree(self):
        # Re-entrancy guard: if we're already inside populate_tree (e.g. via
        # processEvents re-entrancy), defer to a scheduled refresh instead.
        if self._populating_tree:
            log_debug("populate_tree: SKIPPED (re-entrant call) — scheduling deferred refresh")
            self._schedule_tree_refresh()
            return

        self._populating_tree = True
        try:
            self._populate_tree_inner()
        finally:
            self._populating_tree = False

    def _populate_tree_inner(self):
        # Increment generation so stale worker callbacks are ignored
        self._tree_generation += 1
        current_gen = self._tree_generation
        log_debug(f"populate_tree: gen={current_gen}, case={self.file_number}, path={self.case_path}")

        self.tree.clear()
        self.status_label.setText("Scanning directory structure... Please wait.")
        self.tree.setEnabled(False)
        self.process_btn.setEnabled(False)
        self.tree.setSortingEnabled(False)

        # Set up enhanced tree databases for the current case
        # Cache these once to avoid re-reading JSON for every file
        if hasattr(self, 'file_number') and self.file_number:
            self.tree.set_databases(self.file_number)
            self._cached_proc_log = ProcessingLogDB(self.file_number)
            self._cached_tags_db = FileTagsDB(self.file_number)
            # Also update docket agent status
            self.update_docket_agent_status()
        else:
            self._cached_proc_log = None
            self._cached_tags_db = None

        self.tree_item_map = {}
        self.visited_paths = set()
        self.visited_paths.add(self.case_path)

        root_text = f"{os.path.basename(self.case_path)} ({self.file_number})"
        root_item = QTreeWidgetItem(self.tree)
        root_item.setText(0, root_text)
        root_item.setIcon(0, self.icon_provider.icon(QFileInfo(self.case_path)))
        root_item.setData(0, Qt.ItemDataRole.UserRole, self.case_path)
        root_item.setExpanded(True)
        self.tree_item_map[self.case_path] = root_item

        if self.load_cache():
             self.tree.setEnabled(True)
             self.process_btn.setEnabled(True)
             self.status_label.setText(f"Loaded from cache. Verifying...")

        if hasattr(self, 'worker') and self.worker is not None:
            old_worker = self.worker
            old_worker.stop()
            try:
                old_worker.disconnect()
            except Exception:
                pass
            # Do NOT block the UI thread waiting for the old scan to unwind.
            # The old worker may be parked inside os.walk on the Z:\ network
            # drive, where it can't observe the stop flag for several seconds;
            # the former self.worker.wait(5000) here was the cause of the
            # multi-second freeze on case switch.  Instead, detach it: keep a
            # reference so the QThread isn't GC'd while still running (the
            # documented segfault guard) and let it self-terminate.  Its late
            # data_ready/finished emissions are already ignored by the
            # _tree_generation guard in _on_tree_batch/_on_scan_complete, so
            # blocking is not needed for correctness.
            if old_worker.isRunning():
                if not hasattr(self, '_detached_workers'):
                    self._detached_workers = []
                self._detached_workers.append(old_worker)
                old_worker.finished.connect(
                    lambda w=old_worker: self._reap_detached_worker(w)
                )
            else:
                old_worker.deleteLater()

        self.worker = DirectoryTreeWorker(self.case_path)
        # Use lambdas that capture generation to ignore stale callbacks after file switch
        self.worker.data_ready.connect(lambda batch, gen=current_gen: self._on_tree_batch(gen, batch))
        self.worker.finished.connect(lambda gen=current_gen: self._on_scan_complete(gen))
        self.worker.start()
        log_debug(f"populate_tree: worker started for gen={current_gen}")

    def _reap_detached_worker(self, worker):
        """Drop our reference to a detached DirectoryTreeWorker once it has
        finished so it can be garbage-collected. Invoked via the worker's
        finished signal (queued to the UI thread). Replaces the old blocking
        wait()-based teardown that froze the UI on case switch."""
        try:
            if hasattr(self, '_detached_workers') and worker in self._detached_workers:
                self._detached_workers.remove(worker)
        except ValueError:
            pass
        try:
            worker.deleteLater()
        except Exception:
            pass

    def _on_tree_batch(self, generation, batch):
        """Wrapper that ignores stale worker callbacks from a previous file."""
        if generation != self._tree_generation:
            log_debug(f"_on_tree_batch: STALE gen={generation}, current={self._tree_generation} — ignoring")
            return
        self.add_tree_batch(batch)

    def _on_scan_complete(self, generation):
        """Wrapper that ignores stale worker callbacks from a previous file."""
        if generation != self._tree_generation:
            log_debug(f"_on_scan_complete: STALE gen={generation}, current={self._tree_generation} — ignoring")
            return
        self.on_scan_complete()

    def add_tree_batch(self, batch):
        # Pause cyclic GC while building Qt tree items. Cyclic GC can fire mid-loop
        # (on this thread, during allocation) and finalize a stray pythoncom COM
        # proxy on the wrong apartment -> RPC_E_WRONG_THREAD (0x8001010e) -> heap
        # corruption -> access violation in Shiboken/Qt. This batch is also pumped
        # by the nested event loop of native file dialogs, which is exactly where
        # the 2026-06-04 11:38 and 16:27 crashes died. See gc_guard.no_gc.
        with no_gc(collect_on_exit=False):
            self._add_tree_batch_inner(batch)

    def _add_tree_batch_inner(self, batch):
        _batch_gen = self._tree_generation
        _batch_dirs = 0
        _batch_files = 0
        try:
            for root, dirs, files in batch:
                self.visited_paths.add(root)
                parent_item = self.tree_item_map.get(root)
                if not parent_item:
                    continue

                dirs.sort(key=str.lower)
                files.sort(key=lambda x: x[0].lower())

                for d in dirs:
                    dir_path = os.path.join(root, d)
                    if dir_path not in self.tree_item_map:
                        d_item = QTreeWidgetItem(parent_item)
                        d_item.setText(0, d)
                        d_item.setIcon(0, self._get_cached_icon(dir_path, is_dir=True))
                        d_item.setData(0, Qt.ItemDataRole.UserRole, dir_path)
                        d_item.setData(0, Qt.ItemDataRole.UserRole + 1, "dir")
                        d_item.setExpanded(False)
                        self.tree_item_map[dir_path] = d_item
                        _batch_dirs += 1

                for f, size, mtime in files:
                    file_path = os.path.join(root, f)
                    self.visited_paths.add(file_path)

                    size_str = f"{size / 1024:.1f} KB" if size < 1024 * 1024 else f"{size / (1024 * 1024):.1f} MB"
                    date_str = format_date_to_mm_dd_yyyy(mtime)

                    # Get processing status (uses cached DB)
                    proc_status = self._get_file_processing_status(file_path) if self.file_number else ""

                    if file_path not in self.tree_item_map:
                        f_item = QTreeWidgetItem(parent_item)
                        f_item.setText(0, f)

                        # Use cached extension-based icons to avoid network access
                        f_item.setIcon(0, self._get_cached_icon(file_path))

                        f_item.setData(0, Qt.ItemDataRole.UserRole, file_path)
                        f_item.setData(0, Qt.ItemDataRole.UserRole + 1, "file")
                        f_item.setText(2, size_str)
                        f_item.setText(3, date_str)
                        f_item.setText(4, proc_status)

                        self.tree_item_map[file_path] = f_item
                        _batch_files += 1
                    else:
                        f_item = self.tree_item_map[file_path]
                        try:
                            f_item.setText(2, size_str)
                            f_item.setText(3, date_str)
                            f_item.setText(4, proc_status)
                        except RuntimeError:
                            log_warning(f"add_tree_batch: RuntimeError on existing item {file_path}")
                            continue

            log_debug(f"add_tree_batch: gen={_batch_gen}, +{_batch_dirs} dirs, +{_batch_files} files, map_size={len(self.tree_item_map)}")
        except Exception as e:
            log_error(f"add_tree_batch CRASHED: gen={_batch_gen}, +{_batch_dirs}d/+{_batch_files}f, map_size={len(self.tree_item_map)}, error={e}", exc_info=True)
            raise

    def on_scan_complete(self):
        log_debug(f"on_scan_complete: gen={self._tree_generation}, case={self.file_number}, map_size={len(self.tree_item_map)}, visited={len(self.visited_paths)}")
        try:
            # Prune items not visited (deleted files/folders)
            to_remove = []
            for path, item in self.tree_item_map.items():
                if path != self.case_path and path not in self.visited_paths:
                    to_remove.append(path)

            log_debug(f"on_scan_complete: pruning {len(to_remove)} stale items")
            for path in to_remove:
                item = self.tree_item_map.pop(path)
                try:
                    parent = item.parent()
                    if parent:
                        parent.removeChild(item)
                except RuntimeError:
                    log_warning(f"on_scan_complete: RuntimeError removing {path}")

            self.save_cache()
            self.tree.setEnabled(True)
            self.process_btn.setEnabled(True)
            self.status_label.setText(f"Scan Complete. Case: {self.file_number}")
            self.tree.setSortingEnabled(True)
            log_debug(f"on_scan_complete: DONE, final map_size={len(self.tree_item_map)}")
        except Exception as e:
            log_error(f"on_scan_complete CRASHED: gen={self._tree_generation}, case={self.file_number}, error={e}", exc_info=True)
            raise

# Note: Exception handling is now managed by app_crash_handler module
# The install_crash_handler() function sets up sys.excepthook automatically
# Legacy exception_hook removed - see icharlotte_core/app_crash_handler.py

if __name__ == "__main__":
    # Enable faulthandler to catch C-level segfaults that bypass Python exceptions
    import faulthandler
    _faulthandler_log = os.path.join(os.path.dirname(__file__), "logs", "crashes", "faulthandler.log")
    os.makedirs(os.path.dirname(_faulthandler_log), exist_ok=True)
    _faulthandler_file = open(_faulthandler_log, "a")
    _faulthandler_file.write(f"\n{'='*60}\nSession start: {__import__('datetime').datetime.now().isoformat()}\n{'='*60}\n")
    _faulthandler_file.flush()
    faulthandler.enable(file=_faulthandler_file, all_threads=True)

    # Install crash handler FIRST before anything else
    crash_handler = install_crash_handler()
    checkpoint("Application starting")

    try:
        # Parse command-line arguments for restart state restoration
        checkpoint("Parsing command-line arguments")
        parser = argparse.ArgumentParser(description='iCharlotte Legal Document Management Suite')
        parser.add_argument('--file-number', type=str, help='File number to load on startup')
        parser.add_argument('--case-path', type=str, help='Case path to load on startup')
        parser.add_argument('--tab', type=int, help='Tab index to open on startup')
        args, remaining = parser.parse_known_args()

        # Add startup args to crash context
        add_context('startup_file_number', args.file_number)
        add_context('startup_case_path', args.case_path)
        add_context('startup_tab', args.tab)

        # Debug: Log received arguments
        log_debug(f"sys.argv = {sys.argv}")
        log_debug(f"Parsed args = file_number={args.file_number}, case_path={args.case_path}, tab={args.tab}")

        # Disable Chromium sandbox and security for local file editing
        checkpoint("Configuring Qt WebEngine")
        os.environ["QTWEBENGINE_CHROMIUM_FLAGS"] = "--no-sandbox --disable-web-security"

        checkpoint("Creating QApplication")
        app = QApplication(sys.argv)

        # App-wide design system (tokens + base widget styles). Accent buttons
        # styled via theme builders override these defaults locally.
        app.setStyleSheet(theme.app_stylesheet())

        # Set application icon (for taskbar)
        icon_path = os.path.join(os.path.dirname(__file__), 'icharlotte.ico')
        if os.path.exists(icon_path):
            app.setWindowIcon(QIcon(icon_path))

        # Update crash handler with Qt references
        crash_handler.set_qt_references(qt_app=app)

        # Log startup parameters
        log_info(f"Starting with args: file_number={args.file_number}, case_path={args.case_path}, tab={args.tab}")

        # Launch Main Window with restored state if provided
        checkpoint("Creating MainWindow")
        window = MainWindow(
            file_number=args.file_number,
            case_path=args.case_path,
            initial_tab=args.tab
        )

        # Update crash handler with window reference
        crash_handler.set_qt_references(main_window=window)

        checkpoint("Showing MainWindow")
        window.show()

        # Start periodic health monitor to capture state before silent crashes
        from PySide6.QtCore import QTimer as _QTimer
        _health_timer = _QTimer()
        def _health_check():
            try:
                import psutil
                proc = psutil.Process(os.getpid())
                mem = proc.memory_info()
                children = proc.children()
                threads = proc.num_threads()
                # Count running agents
                running_agents = 0
                agent_names = []
                if hasattr(window, 'agent_runners'):
                    for runner in window.agent_runners:
                        if getattr(runner, 'success', None) is None:  # Still running
                            running_agents += 1
                            name = os.path.basename(runner.args[0]) if hasattr(runner, 'args') and runner.args else "?"
                            agent_names.append(name)
                log_debug(
                    f"HEALTH: RSS={mem.rss // 1024 // 1024}MB, "
                    f"threads={threads}, child_procs={len(children)}, "
                    f"running_agents={running_agents}"
                    + (f" [{', '.join(agent_names)}]" if agent_names else "")
                )
                # Warn if memory is getting high (>1.5GB)
                if mem.rss > 1_500_000_000:
                    log_warning(f"HIGH MEMORY: {mem.rss // 1024 // 1024}MB RSS — crash risk")
                # Dump faulthandler traceback periodically when agents are running
                if running_agents > 0:
                    _faulthandler_file.write(f"\n--- Health check {__import__('datetime').datetime.now().isoformat()} agents={running_agents} ---\n")
                    faulthandler.dump_traceback(file=_faulthandler_file, all_threads=True)
                    _faulthandler_file.flush()
            except Exception:
                pass
        _health_timer.timeout.connect(_health_check)
        _health_timer.start(30000)  # Every 30 seconds

        checkpoint("Entering Qt event loop")
        sys.exit(app.exec())

    except Exception as e:
        log_error(f"CRITICAL MAIN ERROR: {e}", exc_info=True)
        print(f"CRITICAL MAIN ERROR: {e}")
        import traceback
        traceback.print_exc()
        # Keep window open if possible or wait for input so user can see error
        input("Press Enter to exit...")
