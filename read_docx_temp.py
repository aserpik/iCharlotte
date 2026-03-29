from docx import Document
import os

path = r'Z:\Shared\Current Clients\3100 - PARSAC\094 - Haire v. City of Wildomar\NOTES\AI OUTPUT\Discovery_Responses_Plaintiff_ANDREW_HAIRE.docx'
doc = Document(path)
out = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'docx_output.txt')
with open(out, 'w', encoding='utf-8') as f:
    for i, p in enumerate(doc.paragraphs):
        f.write(f'{i:04d}|{p.text}\n')
    f.write('\n\n=== TABLES ===\n')
    for ti, table in enumerate(doc.tables):
        f.write(f'\nTable {ti}:\n')
        for ri, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            f.write(f'  Row {ri}: {" | ".join(cells)}\n')
