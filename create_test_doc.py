from docx import Document
import os

doc = Document()
doc.add_paragraph("This is a test paragraph for preview generation.")
doc.add_paragraph("The Plaintiff filed a motion.")

out_dir = os.path.join(".gemini", "tmp")
if not os.path.exists(out_dir):
    os.makedirs(out_dir)
    
out_path = os.path.join(out_dir, "test_preview.docx")
doc.save(out_path)
print(f"Created {out_path}")
