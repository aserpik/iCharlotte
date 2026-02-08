"""
Polish Stage (Stage 4) - Final LLM pass to smooth transitions and ensure consistency.

Reads the full assembled draft, identifies places where transitional phrases or
cross-references are needed, and returns targeted edits to apply to the document
without changing the overall structure.

Usage:
    from Scripts.report_generator.polish import polish_report
    final_path = polish_report(assembled_path, gathered_data, refined_sections)
"""

import os
import re
import logging
from typing import Dict, List, Optional, Tuple
from docx import Document

logger = logging.getLogger(__name__)


def polish_report(
    assembled_path: str,
    gathered_data: Dict,
    refined_sections: Dict[str, str],
    llm_caller=None,
) -> str:
    """
    Apply a polish pass to the assembled report.

    Reads the full text, asks the LLM for targeted edits (transitions,
    cross-references, tonal smoothing), and applies them back to the docx.

    Args:
        assembled_path: Path to the assembled .docx from Stage 3
        gathered_data: Original gathered data (for context)
        refined_sections: The refined section texts
        llm_caller: LLMCaller instance (if None, creates one)

    Returns:
        Path to the polished .docx (same file, modified in place)
    """
    if llm_caller is None:
        llm_caller = _get_llm_caller()

    if llm_caller is None:
        logger.warning("No LLM available for polish pass, skipping")
        return assembled_path

    # Extract full text with section markers
    full_text = _extract_marked_text(assembled_path)
    if not full_text or len(full_text.strip()) < 100:
        logger.warning("Not enough text to polish")
        return assembled_path

    # Get targeted edits from LLM
    edits = _get_polish_edits(full_text, gathered_data, llm_caller)
    if not edits:
        logger.info("No polish edits suggested")
        return assembled_path

    # Apply edits to the document
    _apply_edits(assembled_path, edits)

    logger.info(f"Applied {len(edits)} polish edits to {assembled_path}")
    return assembled_path


def _extract_marked_text(docx_path: str) -> str:
    """
    Extract full text from the assembled report with section markers.
    Section markers help the LLM orient its edit suggestions.
    """
    doc = Document(docx_path)
    lines = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            lines.append("")
            continue

        # Check if this is a section heading (all caps, short)
        if text.isupper() and len(text) < 60 and len(text.split()) <= 5:
            lines.append(f"[SECTION: {text}]")
        else:
            lines.append(text)

    return "\n".join(lines)


def _get_polish_edits(
    full_text: str,
    gathered_data: Dict,
    llm_caller,
) -> List[Tuple[str, str]]:
    """
    Ask the LLM to suggest targeted edits for transitions and consistency.

    Returns:
        List of (find_text, replace_text) tuples
    """
    included = gathered_data.get("included_sections", [])
    report_type = gathered_data.get("report_type", "FSR")

    prompt = f"""You are reviewing a litigation report written by defense counsel for an insurance carrier.
Your task is to suggest TARGETED EDITS to improve transitions and cross-references between sections.

REPORT TYPE: {"Full Status Report / Initial Litigation Plan" if report_type == "FSR" else "Status Report"}
SECTIONS INCLUDED: {', '.join(included)}

RULES:
1. Do NOT add new facts, analysis, or substantive content.
2. Do NOT remove or rewrite sections. Only add transitional phrases or cross-references.
3. Focus on:
   - Smooth transitions between the end of one section and the start of the next
   - Cross-references (e.g., Exposure should reference Liability findings; Further Case Handling should reference outstanding items)
   - Tonal consistency across independently written sections
4. Keep edits minimal and surgical. Prefer inserting a sentence over rewriting a paragraph.
5. Return edits in this exact format, one per block:

FIND: [exact text to find in the document]
REPLACE: [replacement text]

Only return FIND/REPLACE blocks. No other text. If no edits are needed, return "NO EDITS NEEDED".

FULL REPORT TEXT:
"""

    try:
        result = llm_caller.call(
            prompt=prompt,
            text=full_text,
            task_type="general",
            agent_id="agent_report_polish",
        )
    except Exception as e:
        logger.error(f"Polish LLM call failed: {e}")
        return []

    if not result or "NO EDITS NEEDED" in result:
        return []

    return _parse_edit_blocks(result)


def _parse_edit_blocks(response: str) -> List[Tuple[str, str]]:
    """Parse FIND/REPLACE blocks from LLM response."""
    edits = []

    # Match FIND: ... REPLACE: ... blocks
    pattern = r'FIND:\s*(.+?)(?:\n|\r\n?)REPLACE:\s*(.+?)(?=\nFIND:|\Z)'
    matches = re.findall(pattern, response, re.DOTALL)

    for find_text, replace_text in matches:
        find_text = find_text.strip().strip("[]")
        replace_text = replace_text.strip().strip("[]")

        if find_text and replace_text and find_text != replace_text:
            edits.append((find_text, replace_text))

    return edits


def _apply_edits(docx_path: str, edits: List[Tuple[str, str]]):
    """
    Apply find/replace edits to the document.

    Works at the paragraph level — finds paragraphs containing the search text
    and replaces the matching portion while preserving formatting.
    """
    doc = Document(docx_path)
    applied = 0

    for find_text, replace_text in edits:
        for para in doc.paragraphs:
            if find_text in para.text:
                # Replace in the paragraph's runs to preserve formatting
                _replace_in_paragraph(para, find_text, replace_text)
                applied += 1
                break  # Only apply each edit once

    if applied > 0:
        doc.save(docx_path)
        logger.info(f"Applied {applied}/{len(edits)} polish edits")


def _replace_in_paragraph(paragraph, find_text: str, replace_text: str):
    """
    Replace text in a paragraph while preserving run formatting.

    Strategy: Concatenate all run texts, find the match position, then
    distribute the replacement across the same runs.
    """
    # Build full text and run mapping
    runs = paragraph.runs
    if not runs:
        return

    full_text = "".join(run.text for run in runs)
    idx = full_text.find(find_text)
    if idx == -1:
        return

    # Build character-to-run mapping
    char_to_run = []
    for run_idx, run in enumerate(runs):
        for _ in run.text:
            char_to_run.append(run_idx)

    # Find which runs are affected
    start_run = char_to_run[idx]
    end_idx = idx + len(find_text) - 1
    end_run = char_to_run[end_idx]

    # Simple case: entire match is within one run
    if start_run == end_run:
        runs[start_run].text = runs[start_run].text.replace(find_text, replace_text, 1)
        return

    # Multi-run case: put replacement in first run, clear text from middle runs
    # Get the prefix before the match in the first run
    first_run_start = sum(len(runs[i].text) for i in range(start_run))
    prefix_in_first = runs[start_run].text[:idx - first_run_start]

    # Get the suffix after the match in the last run
    last_run_start = sum(len(runs[i].text) for i in range(end_run))
    match_end_in_last = end_idx - last_run_start + 1
    suffix_in_last = runs[end_run].text[match_end_in_last:]

    # Set first run to prefix + replacement
    runs[start_run].text = prefix_in_first + replace_text

    # Clear middle runs
    for i in range(start_run + 1, end_run):
        runs[i].text = ""

    # Set last run to just the suffix
    runs[end_run].text = suffix_in_last


def _get_llm_caller():
    """Create an LLMCaller instance."""
    try:
        from icharlotte_core.llm_config import LLMCaller
        return LLMCaller()
    except ImportError:
        logger.warning("LLMCaller not available")
        return None
