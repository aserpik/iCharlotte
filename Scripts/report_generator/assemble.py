"""
Assemble Stage (Stage 3) - Render refined sections into a Word document.

Two-phase approach:
  1. Use docxtpl to render metadata (header, addressee, dates, case numbers)
  2. Post-process with python-docx to expand section content into properly
     formatted multi-paragraph text with bold/underline sub-headings

This avoids the single-paragraph problem where docxtpl crams all section text
into one {{ var }} paragraph.

Usage:
    from Scripts.report_generator.assemble import assemble_report
    output_path = assemble_report(gathered_data, refined_sections)
"""

import os
import re
import copy
import logging
from typing import Dict, List, Optional, Tuple
from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.table import _Cell
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger(__name__)

PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
DEFAULT_TEMPLATE_PATH = os.path.join(PROJECT_ROOT, "templates", "litigation_report.docx")

SECTION_TO_TEMPLATE_VAR = {
    "FACTUAL BACKGROUND": "factual_background",
    "PROCEDURAL HISTORY": "procedural_history",
    "INVESTIGATION": "investigation",
    "DISCOVERY": "discovery",
    "MEDICAL RECORD REVIEW": "medical_record_review",
    "EVALUATION OF LIABILITY": "evaluation_of_liability",
    "EVALUATION OF EXPOSURE": "evaluation_of_exposure",
    "SETTLEMENT STATUS": "settlement_status",
    "FURTHER CASE HANDLING": "further_case_handling",
}

# Marker prefix used to locate section placeholder paragraphs after rendering
SECTION_MARKER = "%%SECTION_CONTENT%%"


def assemble_report(
    gathered_data: Dict,
    refined_sections: Dict[str, str],
    template_path: str = None,
    output_path: str = None,
) -> str:
    """
    Assemble the final report.

    Phase 1: docxtpl renders date/salutation into template with section markers.
    Phase 2: Dynamically build metadata table (replaces template table).
    Phase 3: python-docx expands section markers into formatted paragraphs.
    Phase 4: Clean up empty paragraphs left over from template extraction.
    """
    if template_path is None:
        template_path = DEFAULT_TEMPLATE_PATH

    if not os.path.exists(template_path):
        raise FileNotFoundError(
            f"Template not found: {template_path}\n"
            "Run template_extractor.py first to create it."
        )

    if output_path is None:
        output_path = _default_output_path(gathered_data)

    # Phase 1: Render metadata with docxtpl (section content = markers)
    context = _build_context(gathered_data, refined_sections)
    logger.info(f"Rendering template: {template_path}")
    tpl = DocxTemplate(template_path)
    tpl.render(context)

    # Save intermediate
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    tpl.save(output_path)

    # Phase 2: Build metadata table dynamically (replaces template table)
    logger.info("Building metadata table...")
    doc = Document(output_path)
    metadata = gathered_data.get("metadata", {})
    metadata["report_type"] = gathered_data.get("report_type", "FSR")
    _build_metadata_table(doc, metadata)

    # Phase 3: Expand section markers into formatted paragraphs
    logger.info("Expanding section content into formatted paragraphs...")
    _expand_section_markers(doc, refined_sections)

    # Phase 4: Clean up empty paragraph bloat from template
    _clean_empty_paragraphs(doc)

    doc.save(output_path)
    logger.info(f"Assembled report saved to: {output_path}")
    return output_path


def _build_context(gathered_data: Dict, refined_sections: Dict[str, str]) -> Dict:
    """
    Build docxtpl context. Section content uses markers, not actual text.

    Only include metadata fields that have actual values to avoid showing
    irrelevant fields (e.g., school_site for non-district cases).
    """
    metadata = gathered_data.get("metadata", {})
    included = gathered_data.get("included_sections", [])

    # Required fields (always include)
    context = {
        "report_date": metadata.get("report_date", ""),
        "file_number": metadata.get("file_number", ""),
        "adjuster_name": metadata.get("adjuster_name", ""),
        "adjuster_email": metadata.get("adjuster_email", ""),
        "adjuster_first_name": metadata.get("adjuster_first_name", ""),
        "intro_paragraph": _build_intro_paragraph(gathered_data),
    }

    # Optional metadata fields - only include if they have values
    optional_fields = {
        "case_number": metadata.get("case_number"),
        "claim_number": metadata.get("claim_number"),
        "policy_number": metadata.get("policy_number"),
        "incident_date": metadata.get("incident_date"),
        "school_site": metadata.get("school_site"),
        "client_name": metadata.get("client_name"),
        "case_name": metadata.get("case_name"),
        "report_subject": metadata.get("report_subject"),
        "attorney_names": metadata.get("assigned_attorney"),
    }

    for key, value in optional_fields.items():
        if value and str(value).strip() and str(value) != "None":
            context[key] = value
        else:
            context[key] = ""  # Empty string for missing fields

    # Section content: use markers so we can find and expand them later
    for section_name, var_name in SECTION_TO_TEMPLATE_VAR.items():
        if section_name in included and section_name in refined_sections:
            context[var_name] = f"{SECTION_MARKER}{section_name}{SECTION_MARKER}"
        else:
            context[var_name] = ""

    return context


def _expand_section_markers(doc: Document, refined_sections: Dict[str, str]):
    """
    Find paragraphs containing section markers and expand them into
    properly formatted multi-paragraph content.

    Section markers are placed after section headings (e.g., MEDICAL RECORD REVIEW).
    We insert an empty paragraph first to ensure proper spacing, then the content.
    """
    # Find all marker paragraphs (iterate in reverse so we can insert after)
    marker_paras = []
    for i, para in enumerate(doc.paragraphs):
        if SECTION_MARKER in para.text:
            # Extract section name from marker
            text = para.text
            start = text.find(SECTION_MARKER) + len(SECTION_MARKER)
            end = text.find(SECTION_MARKER, start)
            if end > start:
                section_name = text[start:end]
                marker_paras.append((i, para, section_name))

    # Process in reverse order to preserve indices
    for idx, marker_para, section_name in reversed(marker_paras):
        content = refined_sections.get(section_name, "")
        if not content:
            # Clear the marker but leave an empty paragraph for spacing
            _clear_paragraph(marker_para)
            continue

        # Ensure the section heading (paragraph before the marker) has space_after
        # so there's always a visual gap between heading and content, regardless
        # of whether empty paragraphs survive cleanup.
        _ensure_heading_spacing(doc, idx)

        # Parse the content into structured paragraphs (pass section name for context)
        parsed = _parse_section_content(content, section_name=section_name)

        # Get the formatting reference from the marker paragraph
        ref_format = _get_paragraph_format(marker_para)

        # Insert parsed paragraphs after the marker, then clear the marker
        _insert_paragraphs_after(doc, marker_para, parsed, ref_format)
        _clear_paragraph(marker_para)


def _parse_section_content(text: str, section_name: str = None) -> List[Tuple[str, str]]:
    """
    Parse refined text into a list of (type, content) tuples.

    Args:
        text: The section content text
        section_name: Name of the section (for context-aware classification)

    Types:
        "subheading_l1" - Lettered sub-heading (A. Pleadings) — 0.5" hanging indent
        "subheading_l2" - Numbered sub-heading (1. Duty) — 1.0" hanging indent
        "subheading_plain" - Unlabeled sub-heading (auto-assigned letter)
        "subheading_plain_nested" - Unlabeled nested sub-heading (auto-assigned number)
        "body" - Normal body paragraph
    """
    paragraphs = []
    lines = text.split("\n")

    # Group lines into paragraphs (separated by blank lines)
    current_lines = []
    for line in lines:
        stripped = line.strip()
        if not stripped:
            if current_lines:
                paragraphs.append("\n".join(current_lines))
                current_lines = []
        else:
            current_lines.append(stripped)
    if current_lines:
        paragraphs.append("\n".join(current_lines))

    # Classify each paragraph
    result = []
    for para_text in paragraphs:
        joined = " ".join(para_text.split("\n"))
        # Strip markdown bold
        clean = re.sub(r'\*\*(.+?)\*\*', r'\1', joined).strip()

        para_type = _classify_paragraph(clean, section_name=section_name)
        result.append((para_type, clean))

    # Auto-assign letters/numbers to plain subheadings
    _auto_letter_subheadings(result)

    return result


def _classify_paragraph(text: str, section_name: str = None) -> str:
    """
    Classify a paragraph as body or one of the subheading levels.

    Args:
        text: The paragraph text
        section_name: The section name for context-aware classification
    """
    # Lettered sub-headings: "A. Something", "B. Something"
    if re.match(r'^[A-Z]\.\s+\S', text) and len(text) < 120:
        return "subheading_l1"

    # Numbered sub-headings: "1. Something", "2. Something"
    if re.match(r'^\d+\.\s+\S', text) and len(text) < 120:
        return "subheading_l2"

    # Too long to be a subheading
    if len(text) > 120:
        return "body"

    # Liability section: Element headings should be nested/numbered
    liability_element_patterns = [
        r'^(?:Ownership and Control|Ownership and Creation of Condition|'
        r'Unreasonable Conduct|Creation of (?:Condition|the Condition)|'
        r'Causation|Consent|Duty|Breach|Damages|Scope of Employment|'
        r'Ratification|Statute of Limitations|Comparative (?:Fault|Negligence)|'
        r'Intentional (?:Conduct|Act)|Knowledge|Actual Malice|'
        r'Publication|Defamatory Statement|Identification|'
        r'Trespass|Interference|Tortious Conduct|'
        r'Proximate Cause|Foreseeability|'
        r'Intent and .*Defense|.*Defense$|'
        r'Standing|Substantial (?:Factor|Interference)|'
        r'Unreasonable (?:Interference|Conduct)|'
        r'Assumption of Risk|Contributory Negligence)$',
    ]
    if section_name and "LIABILITY" in section_name:
        for pattern in liability_element_patterns:
            if re.match(pattern, text, re.IGNORECASE):
                return "subheading_plain_nested"

    # Known sub-heading patterns from real reports (top-level sections)
    # These are causes of action and major section headings — always L1.
    # Must be checked BEFORE the broad LIABILITY fallback so that cause-of-action
    # names like "Continuing Nuisance" aren't misclassified as L2 elements.
    subheading_patterns = [
        r'^(?:Government Tort Claim|Complaint|Pleadings|Responsive Pleading|'
        r'Venue|Representation|Upcoming Hearings?|Future Hearings?|'
        r'Written Discovery|Depositions?|Outstanding Discovery|'
        r'Physical Injuries|Emotional.*Damages|General Damages|'
        r'Economic Damages|Non-Economic Damages|Punitive Damages|'
        r'Property Damage|Lost (?:Wages|Earnings|Rental|Rental Income)(?:\s+and\s+.+)?|'
        r'Related Litigation|Summary(?:\s+of\s+\w+)?|Assessment of Exposure|'
        r'Continuing (?:Nuisance|Trespass(?:\s+to\s+Land)?)|Negligence|Strict Liability|'
        r'Premises Liability|Dangerous Condition|'
        r'Intentional Infliction|Negligent Infliction|'
        r'Battery|Assault|False Imprisonment|Defamation|'
        r'Invasion of Privacy|Fraud|Misrepresentation|'
        r'Injunctive Relief|Declaratory Relief)$',
    ]
    for pattern in subheading_patterns:
        if re.match(pattern, text, re.IGNORECASE):
            return "subheading_plain"

    # Broader fallback for LIABILITY: any short, title-case line that
    # doesn't end with a period is likely an element subheading.
    # Runs AFTER subheading_patterns to avoid catching causes of action.
    if section_name and "LIABILITY" in section_name:
        words = text.split()
        if (2 <= len(words) <= 8 and len(text) < 80
                and not text.endswith('.')
                and not text.endswith(':')
                and sum(1 for w in words if w[0].isupper()) >= len(words) * 0.5):
            return "subheading_plain_nested"

    # Lines with "Plaintiff X's Responses" or "Defendant Y's Responses" pattern
    # These should be numbered (nested under discovery section), not lettered
    # Must be checked BEFORE the generic short-title-case check to avoid
    # misclassifying party names as top-level subheadings
    if re.match(r"^(?:Plaintiff|Defendant)\s+.+(?:'s\s+)?(?:Response|Deposition|Interrogator|Request)", text, re.IGNORECASE):
        return "subheading_plain_nested"

    # Short title-case lines (1-6 words, capitalized, not ending in period)
    words = text.split()
    if 1 <= len(words) <= 6 and not text.endswith('.'):
        cap_words = sum(1 for w in words if w[0].isupper())
        if cap_words >= len(words) * 0.7:
            return "subheading_plain"

    return "body"


def _auto_letter_subheadings(parsed: List[Tuple[str, str]]):
    """
    Auto-assign letter prefixes (A., B., ...) to plain subheadings and
    number prefixes (1., 2., ...) to nested plain subheadings.
    """
    letter_idx = 0
    number_idx = 0

    for i, (ptype, text) in enumerate(parsed):
        if ptype == "subheading_l1":
            # Already lettered — track which letter we're at
            m = re.match(r'^([A-Z])\.', text)
            if m:
                letter_idx = ord(m.group(1)) - ord('A') + 1
            # Reset number counter when we hit a new lettered section
            number_idx = 0
        elif ptype == "subheading_l2":
            # Already numbered — track which number we're at
            m = re.match(r'^(\d+)\.', text)
            if m:
                number_idx = int(m.group(1))
        elif ptype == "subheading_plain":
            # Auto-assign letter
            letter = chr(ord('A') + letter_idx)
            parsed[i] = ("subheading_l1", f"{letter}.\t{text}")
            letter_idx += 1
            number_idx = 0  # Reset number counter
        elif ptype == "subheading_plain_nested":
            # Auto-assign number (nested under current lettered section)
            number_idx += 1
            parsed[i] = ("subheading_l2", f"{number_idx}.\t{text}")


def _get_paragraph_format(para) -> Dict:
    """Extract formatting properties from a reference paragraph."""
    fmt = {}
    pf = para.paragraph_format
    fmt["space_after"] = pf.space_after
    fmt["space_before"] = pf.space_before
    fmt["first_line_indent"] = pf.first_line_indent
    fmt["left_indent"] = pf.left_indent

    # Get font info from first run
    if para.runs:
        run = para.runs[0]
        fmt["font_name"] = run.font.name
        fmt["font_size"] = run.font.size
    else:
        fmt["font_name"] = "Times New Roman"
        fmt["font_size"] = Pt(12)

    return fmt


def _ensure_heading_spacing(doc: Document, marker_idx: int):
    """
    Find the section heading paragraph before the marker and ensure it has
    space_after=240 (12pt) so there's a visual gap before the content.
    """
    # Walk backwards from marker to find the heading (all-caps, short paragraph)
    for i in range(marker_idx - 1, max(marker_idx - 5, -1), -1):
        if i < 0:
            break
        para = doc.paragraphs[i]
        text = para.text.strip()
        if text and text.isupper() and len(text) < 60:
            # Found the heading — set space_after via OXML
            pPr = para._element.find(qn('w:pPr'))
            if pPr is None:
                pPr = OxmlElement('w:pPr')
                para._element.insert(0, pPr)
            spacing = pPr.find(qn('w:spacing'))
            if spacing is None:
                spacing = OxmlElement('w:spacing')
                pPr.append(spacing)
            spacing.set(qn('w:after'), '240')  # 12pt = one line gap
            break


def _emu_to_twips(emu):
    """Convert English Metric Units to twips (twentieths of a point)."""
    return int(emu / 635)


# Default formatting constants (twips)
DEFAULT_FIRST_LINE_INDENT = '720'   # 0.5 inches
DEFAULT_SPACE_AFTER = '240'         # 12pt


def _insert_paragraphs_after(
    doc: Document,
    ref_para,
    parsed: List[Tuple[str, str]],
    ref_format: Dict,
):
    """Insert formatted paragraphs after the reference paragraph."""
    ref_element = ref_para._element

    # Insert in reverse order (each insert goes right after ref)
    for para_type, text in reversed(parsed):
        if not text.strip():
            continue

        new_para = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')

        # --- Indentation per paragraph type ---
        ind = OxmlElement('w:ind')

        if para_type == "subheading_l1":
            # Level 1: letter at 0.5", text at 1.0" (hanging indent)
            ind.set(qn('w:left'), '1440')     # 1.0" total
            ind.set(qn('w:hanging'), '720')   # 0.5" hang-back

        elif para_type == "subheading_l2":
            # Level 2: number at 1.0", text at 1.5" (hanging indent)
            ind.set(qn('w:left'), '2160')     # 1.5" total
            ind.set(qn('w:hanging'), '720')   # 0.5" hang-back

        else:
            # Body paragraphs: ALWAYS use 0.5" first-line indent (720 twips)
            # Don't rely on ref_format as it may be corrupted
            ind.set(qn('w:firstLine'), DEFAULT_FIRST_LINE_INDENT)

        pPr.append(ind)

        # Space after: 12pt (240 twips), single line spacing
        spacing = OxmlElement('w:spacing')
        if ref_format.get("space_after"):
            twips = _emu_to_twips(ref_format["space_after"])
            spacing.set(qn('w:after'), str(twips))
        else:
            spacing.set(qn('w:after'), DEFAULT_SPACE_AFTER)
        spacing.set(qn('w:line'), '240')
        pPr.append(spacing)

        new_para.append(pPr)

        # --- Build run(s) with text ---
        is_subheading = para_type.startswith("subheading")

        if is_subheading and '\t' in text:
            # Subheading with tab separator (e.g., "A.\tPleadings")
            # Prefix (letter/number) is bold only; text is bold + underline
            parts = text.split('\t', 1)
            _add_run(new_para, parts[0], ref_format, bold=True, underline=False)
            _add_tab(new_para)
            _add_run(new_para, parts[1], ref_format, bold=True, underline=True)
        elif is_subheading:
            # Subheading without tab — check for "A. " or "1. " prefix
            m = re.match(r'^([A-Z]\.\s+|^\d+\.\s+)', text)
            if m:
                prefix = m.group(1)
                rest = text[len(prefix):]
                _add_run(new_para, prefix, ref_format, bold=True, underline=False)
                _add_run(new_para, rest, ref_format, bold=True, underline=True)
            else:
                _add_run(new_para, text, ref_format, bold=True, underline=True)
        else:
            # Body paragraph: check for topic heading (e.g., "Background. The plaintiff...")
            # Topic headings are 1-6 words (allowing lowercase connectors like "and", "of")
            # followed by a period, at the start of the paragraph
            topic_match = re.match(
                r'^([A-Z][a-zA-Z]+(?:\s+(?:and|of|to|for|the|in|on|at|by|with|[A-Z][a-zA-Z]+)){0,5}\.)\s+(.+)',
                text
            )
            if topic_match:
                # Bold the topic heading, normal text for the rest
                topic = topic_match.group(1)
                rest = topic_match.group(2)
                _add_run(new_para, topic + " ", ref_format, bold=True, underline=False)
                _add_run(new_para, rest, ref_format, bold=False, underline=False)
            else:
                _add_run(new_para, text, ref_format, bold=False, underline=False)

        ref_element.addnext(new_para)


def _add_run(para_element, text: str, ref_format: Dict,
             bold: bool = False, underline: bool = False):
    """Add a formatted text run to a paragraph element."""
    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Font
    font_name = ref_format.get("font_name", "Times New Roman")
    if font_name:
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
        rPr.append(rFonts)

    # Font size
    font_size = ref_format.get("font_size")
    if font_size:
        half_pts = int(font_size / 6350)
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(half_pts))
        rPr.append(sz)
        szCs = OxmlElement('w:szCs')
        szCs.set(qn('w:val'), str(half_pts))
        rPr.append(szCs)

    if bold:
        rPr.append(OxmlElement('w:b'))
    if underline:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)

    run.append(rPr)

    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    run.append(t)

    para_element.append(run)


def _add_tab(para_element):
    """Add a tab character run to a paragraph element."""
    run = OxmlElement('w:r')
    run.append(OxmlElement('w:tab'))
    para_element.append(run)


def _clean_empty_paragraphs(doc: Document):
    """
    Remove empty paragraphs from the document.

    Strategy: Remove ALL empty paragraphs that appear between content
    paragraphs within sections. Paragraph spacing (space_after) handles
    visual separation, so empty spacer paragraphs are not needed.

    Preserve EXACTLY ONE empty paragraph after each section heading for
    proper spacing between heading and content.
    """
    body = doc.element.body
    elements = list(body)

    in_sections = False  # Are we past the first section heading?
    after_section_heading = False  # Was the previous content para a section heading?
    to_remove = []
    consecutive_empty = 0

    for element in elements:
        if element.tag.endswith('}p'):
            # Preserve closing block paragraphs (zClosing style) — they
            # provide signature space and must never be removed
            pPr = element.find(qn('w:pPr'))
            if pPr is not None:
                pStyle = pPr.find(qn('w:pStyle'))
                if pStyle is not None and pStyle.get(qn('w:val')) == 'zClosing':
                    consecutive_empty = 0
                    continue

            text = ""
            for child in element.iter():
                if child.text:
                    text += child.text
                if child.tail:
                    text += child.tail

            text_stripped = text.strip()

            # Detect section headings (all caps, short)
            is_section_heading = (
                text_stripped and text_stripped.isupper()
                and len(text_stripped) < 60 and len(text_stripped.split()) <= 5
            )
            if is_section_heading:
                in_sections = True
                after_section_heading = True
                consecutive_empty = 0
            elif not text_stripped:
                consecutive_empty += 1
                if not in_sections:
                    # In header area: keep at most one empty para in a row
                    if consecutive_empty > 1:
                        to_remove.append(element)
                elif after_section_heading:
                    # Remove ALL empty paragraphs after section headings —
                    # space_after on the heading provides the visual gap
                    to_remove.append(element)
                else:
                    # Not after section heading - remove all empty paragraphs
                    to_remove.append(element)
            else:
                consecutive_empty = 0
                after_section_heading = False
        else:
            consecutive_empty = 0

    for element in to_remove:
        body.remove(element)

    logger.info(f"Removed {len(to_remove)} empty paragraphs")


def _clear_paragraph(paragraph):
    """Remove all content from a paragraph."""
    for run in paragraph.runs:
        run.text = ""
    # Also clear any direct text
    element = paragraph._element
    for child in list(element):
        if child.tag.endswith('}r'):
            for t in child.findall(qn('w:t')):
                t.text = ""


def _build_intro_paragraph(gathered_data: Dict) -> str:
    """
    Build the salutation and intro paragraph.

    Format:
        Dear [FirstName],

        [Intro paragraph]
    """
    metadata = gathered_data.get("metadata", {})
    report_type = gathered_data.get("report_type", "FSR")

    # Build salutation
    adjuster_first_name = metadata.get("adjuster_first_name", "")
    if not adjuster_first_name:
        # Try to extract first name from full adjuster name
        adjuster_name = metadata.get("adjuster_name", "")
        if adjuster_name:
            adjuster_first_name = adjuster_name.split()[0]

    salutation = f"Dear {adjuster_first_name}," if adjuster_first_name else "Dear Adjuster,"

    # Build intro text
    if report_type == "FSR":
        intro = (
            "Please find enclosed our Initial Litigation Plan regarding the "
            "above-referenced matter. We provide this report for your review "
            "and consideration."
        )
    else:
        intro = (
            "Please find enclosed our updated Status Report regarding the "
            "above-referenced matter. We provide this update for your review "
            "and consideration."
        )

    # Return salutation + blank line + intro
    return f"{salutation}\n\n{intro}"


def _build_metadata_table(doc: Document, metadata: Dict):
    """
    Build the metadata section using a TABLE with paragraphs inside cells (Report Agent format).

    Strategy: Delete all content between VIA EMAIL and first section,
    then insert a table with metadata field paragraphs in a single cell.

    Format:
        VIA EMAIL:

        [Table with 2 columns, left column contains:]
            Re: [Case Name]
            Client/Insured: [Client]
            Claim No.: [Claim Number]
            BS File No.: [File Number]
            Case No.: [Case Number]
            Date of Loss: [Date]
            Subject: [Subject]

        Dear [Adjuster],

        [Intro paragraph]
    """
    # Find key paragraphs
    via_email_idx = None
    first_section_idx = None

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip().upper()
        if "VIA EMAIL" in text:
            via_email_idx = i
        elif text in ["FACTUAL BACKGROUND", "PROCEDURAL HISTORY", "INVESTIGATION"]:
            first_section_idx = i
            break

    if via_email_idx is None:
        logger.warning("Could not find VIA EMAIL paragraph")
        return None

    if first_section_idx is None:
        logger.warning("Could not find first section heading")
        return None

    # Delete all paragraphs and tables between VIA EMAIL and first section
    elements_to_delete = []

    # Collect paragraphs to delete
    for i in range(via_email_idx + 1, first_section_idx):
        if i < len(doc.paragraphs):
            elements_to_delete.append(doc.paragraphs[i]._element)

    # Collect tables to delete (any table before first section)
    for table in doc.tables:
        table_pos = doc.element.body.index(table._element)
        via_email_pos = doc.element.body.index(doc.paragraphs[via_email_idx]._element)
        first_section_pos = doc.element.body.index(doc.paragraphs[first_section_idx]._element)
        if via_email_pos < table_pos < first_section_pos:
            elements_to_delete.append(table._element)

    # Delete collected elements
    for element in elements_to_delete:
        element.getparent().remove(element)

    logger.info(f"Deleted {len(elements_to_delete)} elements between VIA EMAIL and first section")

    # Create table with 2 columns
    via_element = doc.paragraphs[via_email_idx]._element
    table = doc.add_table(rows=2, cols=1)  # ISSUE #6: Use 1 column table

    # ISSUE #6 FIX: Set table width to extend to 5.5" from left margin
    # Table starts at 1.5" indent, so width should be 4" to reach 5.5"
    table.columns[0].width = Inches(4.0)

    # Remove table borders
    for row in table.rows:
        for cell in row.cells:
            cell._element.get_or_add_tcPr().append(_create_no_border_element())

    # Set table left indent to 1.5 inches (1440 twips)
    tbl = table._element
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    tblInd = OxmlElement('w:tblInd')
    tblInd.set(qn('w:w'), '1440')
    tblInd.set(qn('w:type'), 'dxa')
    tblPr.append(tblInd)

    # Position table after VIA EMAIL with TWO blank lines before it
    # Remove table from its current position first
    table._element.getparent().remove(table._element)

    # Insert first blank line after VIA EMAIL
    blank1 = OxmlElement('w:p')
    via_element.addnext(blank1)

    # Insert second blank line
    blank2 = OxmlElement('w:p')
    blank1.addnext(blank2)

    # Insert table after the two blank lines
    blank2.addnext(table._element)

    # Get the left cell (where all metadata fields go)
    metadata_cell = table.rows[0].cells[0]

    # Clear default paragraph in cell
    for para in metadata_cell.paragraphs:
        para.clear()

    # Build field list with label and value
    # ISSUE #3 FIX: Convert all-caps names to proper case
    client_name_table = metadata.get("client_name", "")
    if client_name_table and client_name_table.isupper():
        client_name_table = client_name_table.title()

    fields = [
        ("Re:", metadata.get("case_name", "")),
        ("Client/Insured:", client_name_table),
        ("Claim No.:", metadata.get("claim_number", "")),
        ("BS File No.:", metadata.get("file_number", "")),
        ("Case No.:", metadata.get("case_number", "")),
        ("Date of Loss:", metadata.get("incident_date", "")),
        ("Subject:", metadata.get("report_subject", "Initial Litigation Plan")),
    ]

    # Add blank paragraph at the start
    blank_para = metadata_cell.paragraphs[0]
    blank_para.paragraph_format.line_spacing = 1.0
    blank_para.paragraph_format.space_after = Pt(0)
    blank_para.paragraph_format.space_before = Pt(0)

    # Add fields as paragraphs within the cell
    for label, value in fields:
        para = metadata_cell.add_paragraph()

        # ISSUE #5 FIX: Format consistently - label followed by spaces then value
        # No leading spaces, case name immediately after "Re: " (with one space)
        if value:
            # Align values at position ~20 characters
            spacing = max(1, 20 - len(label))
            text = label + " " * spacing + value
        else:
            # No value - just label with trailing spaces
            text = label + " " * 14

        para.add_run(text)
        para.paragraph_format.line_spacing = 1.0
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.space_before = Pt(0)

    # Add paragraphs after table with correct spacing
    # Find the position after the table to insert paragraphs
    table_element = table._element
    insert_after = table_element

    # Salutation - with space_after to create blank line
    adjuster_first_name = metadata.get("adjuster_first_name", "")
    adjuster_name = metadata.get("adjuster_name", "")
    if not adjuster_first_name and adjuster_name:
        adjuster_first_name = adjuster_name.split()[0]
    salutation = f"Dear {adjuster_first_name}," if adjuster_first_name else "Dear Adjuster,"

    sal_para = OxmlElement('w:p')
    sal_pPr = OxmlElement('w:pPr')

    # Add space_after to create blank line (12pt = one line)
    sal_spacing = OxmlElement('w:spacing')
    sal_spacing.set(qn('w:after'), '240')  # 12pt = 240 twips
    sal_pPr.append(sal_spacing)
    sal_para.append(sal_pPr)

    sal_run = OxmlElement('w:r')
    sal_text = OxmlElement('w:t')
    sal_text.text = salutation
    sal_run.append(sal_text)
    sal_para.append(sal_run)
    insert_after.addnext(sal_para)
    insert_after = sal_para

    # Introduction paragraph - standard intro (NOT factual_background)
    # Factual background goes in the FACTUAL BACKGROUND section, not here
    client_name = metadata.get("client_name", "Defendant")
    # ISSUE #3 FIX: Convert to proper case
    if client_name and client_name.isupper():
        client_name = client_name.title()

    report_type = metadata.get("report_type", "FSR")

    if report_type == "FSR":
        intro = (
            f'As you recall, Bordin Semmer LLP has been retained to represent the interests of {client_name} '
            f'("Defendant") in the above-referenced matter. Please allow the following to serve as our '
            f'Litigation Plan in this matter.'
        )
    else:
        intro = (
            f'As you recall, Bordin Semmer LLP has been retained to represent the interests of {client_name} '
            f'("Defendant") in the above-referenced matter. Please allow the following to serve as an update '
            f'in this matter.'
        )

    intro_para = OxmlElement('w:p')
    # ISSUE #2 FIX: Add 0.5" first-line indent
    # ISSUE #4 FIX: Add space_after to create blank line before FACTUAL BACKGROUND
    intro_pPr = OxmlElement('w:pPr')
    intro_ind = OxmlElement('w:ind')
    intro_ind.set(qn('w:firstLine'), '720')  # 720 twips = 0.5 inch
    intro_pPr.append(intro_ind)

    # Add space_after to create blank line (12pt = one line)
    intro_spacing = OxmlElement('w:spacing')
    intro_spacing.set(qn('w:after'), '240')  # 12pt = 240 twips
    intro_pPr.append(intro_spacing)

    intro_para.append(intro_pPr)

    intro_run = OxmlElement('w:r')
    intro_text = OxmlElement('w:t')
    intro_text.text = intro
    intro_run.append(intro_text)
    intro_para.append(intro_run)
    insert_after.addnext(intro_para)
    insert_after = intro_para

    logger.info("Built metadata section with table")
    return table




def _create_no_border_element():
    """Create XML element for borderless table cell."""
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        border.set(qn('w:sz'), '0')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        tcBorders.append(border)
    return tcBorders


def _looks_like_date(text):
    """Check if text looks like a date."""
    import re
    date_patterns = [
        r'(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{4}',
        r'\d{1,2}/\d{1,2}/\d{4}',
    ]
    return any(re.search(p, text, re.IGNORECASE) for p in date_patterns)


def _default_output_path(gathered_data: Dict) -> str:
    """Generate the default output path."""
    output_dir = gathered_data.get("output_dir", "")
    report_type = gathered_data.get("report_type", "FSR")

    if not output_dir:
        output_dir = os.path.join(PROJECT_ROOT, "output")

    from datetime import datetime
    date_str = datetime.now().strftime("%Y-%m-%d")

    type_label = "ILP" if report_type == "FSR" else "Status_Report"
    filename = f"{type_label}_{date_str}.docx"
    path = os.path.join(output_dir, filename)

    if os.path.exists(path):
        counter = 2
        while os.path.exists(path):
            filename = f"{type_label}_{date_str}_v{counter}.docx"
            path = os.path.join(output_dir, filename)
            counter += 1

    return path
