"""
Gather Stage (Stage 1) - Collect case metadata and agent outputs for report generation.

Pulls case metadata from CaseDataManager and MasterCaseDatabase, scans the
AI OUTPUT folder for agent outputs, and determines report type.

Usage:
    from Scripts.report_generator.gather import gather_case_data
    data = gather_case_data("1280.001")
"""

import os
import re
import json
import logging
from datetime import datetime
from typing import Dict, List, Optional, Tuple
from docx import Document

logger = logging.getLogger(__name__)

# Where agents write their output
AI_OUTPUT_SUBDIR = os.path.join("NOTES", "AI OUTPUT")

# Maps agent output filenames to report section names.
# Pattern can be a string (exact prefix match) or compiled regex.
AGENT_OUTPUT_MAP = [
    # (filename pattern, report section, source description)
    ("liability_eval", "EVALUATION OF LIABILITY", "liability.py"),
    ("exposure_eval", "EVALUATION OF EXPOSURE", "exposure.py"),
    ("Med_Record_", "MEDICAL RECORD REVIEW", "med_record.py"),
    ("med_chron", "MEDICAL RECORD REVIEW", "med_chron.py"),
    ("Discovery_Responses_", "DISCOVERY", "summarize_discovery.py"),
    ("Docket_", "PROCEDURAL STATUS", "docket.py"),
    ("variables", "_VARIABLES", "complaint.py"),
    ("caption_template", "_CAPTION", "complaint.py"),
]

# Report section ordering (canonical order in the report)
SECTION_ORDER = [
    "FACTUAL BACKGROUND",
    "PROCEDURAL STATUS",
    "INVESTIGATION",
    "DISCOVERY",
    "EXPERTS",
    "MEDICAL RECORD REVIEW",
    "EVALUATION OF LIABILITY",
    "EVALUATION OF EXPOSURE",
    "SETTLEMENT STATUS",
    "FURTHER CASE HANDLING",
]

# Sections always included in status reports
STATUS_REPORT_ALWAYS_INCLUDE = {
    "FACTUAL BACKGROUND",
    "PROCEDURAL STATUS",
    "EVALUATION OF LIABILITY",
    "EVALUATION OF EXPOSURE",
    "SETTLEMENT STATUS",
    "FURTHER CASE HANDLING",
}

# Sections where multiple source files are date-versioned snapshots (use latest only)
# Other sections (med records, discovery) legitimately have multiple distinct sources
USE_LATEST_ONLY = {"PROCEDURAL STATUS"}

# Keywords that route AI_OUTPUT summaries to EXPERTS instead of INVESTIGATION
EXPERTS_KEYWORDS = re.compile(
    r'(?:\bIME\b|independent\s+medical\s+exam(?:ination)?|preliminary\s+opinions?\b)',
    re.IGNORECASE
)


def gather_case_data(file_number: str, report_type: str = "FSR",
                     prior_report_path: str = None) -> Dict:
    """
    Gather all data needed to generate a litigation report.

    Args:
        file_number: Case file number (e.g., "1280.001")
        report_type: "FSR" (Full Status Report / ILP) or "STATUS" (Status Report)
        prior_report_path: Path to prior report .docx (for status reports)

    Returns:
        Dict with keys:
            metadata: Case metadata (parties, dates, numbers, etc.)
            sections: Dict mapping section name -> raw content text
            section_files: Dict mapping section name -> source file path
            report_type: "FSR" or "STATUS"
            prior_sections: Dict of prior report sections (if status report)
            case_path: Filesystem path to case folder
            output_dir: Path to AI OUTPUT folder
    """
    # Get case path
    case_path = _get_case_path(file_number)
    if not case_path:
        raise FileNotFoundError(f"Could not find case folder for {file_number}")

    output_dir = os.path.join(case_path, AI_OUTPUT_SUBDIR)
    if not os.path.exists(output_dir):
        raise FileNotFoundError(f"AI OUTPUT folder not found: {output_dir}")

    # Collect metadata from CaseDataManager + MasterCaseDatabase
    metadata = _gather_metadata(file_number, case_path)

    # Scan AI OUTPUT folder and extract content per section
    sections, section_files = _scan_agent_outputs(output_dir)

    # Parse variables.docx specially — it contains structured data from complaint.py
    _process_variables_file(output_dir, sections, section_files, metadata)

    # Ingest AI_OUTPUT.docx summaries into INVESTIGATION / EXPERTS
    _ingest_ai_output(output_dir, sections, section_files)

    # For status reports, parse the prior report for delta comparison
    prior_sections = {}
    if report_type == "STATUS" and prior_report_path:
        prior_sections = _parse_prior_report(prior_report_path)

    # Determine which sections to include based on report type
    included = _determine_included_sections(
        sections, report_type, section_files, prior_report_path
    )

    result = {
        "metadata": metadata,
        "sections": {k: v for k, v in sections.items() if k in included},
        "section_files": {k: v for k, v in section_files.items() if k in included},
        "all_available_sections": list(sections.keys()),
        "included_sections": included,
        "report_type": report_type,
        "prior_sections": prior_sections,
        "case_path": case_path,
        "output_dir": output_dir,
    }

    logger.info(f"Gathered data for {file_number}: {len(result['sections'])} sections")
    return result


def _process_variables_file(
    output_dir: str,
    sections: Dict[str, str],
    section_files: Dict[str, str],
    metadata: Dict,
):
    """
    Parse variables.docx (complaint.py output) which contains a table of
    extracted case variables. Some variables are report section content
    (factual_background, procedural_history), others are metadata.
    """
    filepath = os.path.join(output_dir, "variables.docx")
    if not os.path.exists(filepath):
        return

    try:
        doc = Document(filepath)
    except Exception as e:
        logger.warning(f"Could not open variables.docx: {e}")
        return

    # Parse the key-value table
    variables = {}
    for table in doc.tables:
        for row in table.rows:
            if len(row.cells) >= 2:
                key = row.cells[0].text.strip().lower()
                val = row.cells[1].text.strip()
                if key and val and key != "variable":
                    variables[key] = val

    if not variables:
        return

    # Extract section content from variables
    section_var_map = {
        "factual_background": "FACTUAL BACKGROUND",
        "procedural_history": "PROCEDURAL STATUS",
    }

    for var_key, section_name in section_var_map.items():
        content = variables.get(var_key, None)  # Use .get instead of .pop to keep in variables
        if content and len(content) > 50:
            if section_name not in sections:
                sections[section_name] = content
                section_files[section_name] = filepath
            else:
                # Prepend complaint's prose version before raw docket data
                sections[section_name] = (
                    content + "\n\n--- Supplementary docket data ---\n\n"
                    + sections[section_name]
                )

    # Merge remaining variables into metadata (don't overwrite existing)
    skip_keys = {"variable", "value", "file_path", "caption",
                 "caption_template_path"}
    for key, val in variables.items():
        if key in skip_keys:
            continue
        meta_key = key.replace(" ", "_")
        if meta_key not in metadata:
            metadata[meta_key] = val

    logger.info(f"Parsed {len(variables)} variables from variables.docx")


def _get_case_path(file_number: str) -> Optional[str]:
    """Resolve case folder path using existing infrastructure."""
    try:
        from icharlotte_core.utils import get_case_path
        return get_case_path(file_number)
    except ImportError:
        logger.warning("Could not import get_case_path, trying CaseDataManager")

    try:
        from Scripts.case_data_manager import CaseDataManager
        mgr = CaseDataManager()
        path = mgr.get_value(file_number, "file_path")
        if path and os.path.exists(path):
            return os.path.abspath(path)
    except ImportError:
        pass

    return None


def _gather_metadata(file_number: str, case_path: str) -> Dict:
    """Pull case metadata from CaseDataManager and MasterCaseDatabase."""
    metadata = {
        "file_number": file_number,
        "report_date": datetime.now().strftime("%B %d, %Y"),
        "case_path": case_path,
    }

    # Try CaseDataManager first (richer data from agent extractions)
    try:
        from Scripts.case_data_manager import CaseDataManager
        mgr = CaseDataManager()
        all_vars = mgr.get_all_variables(file_number, flatten=True)

        # Map known variable names to metadata fields
        var_mapping = {
            "plaintiffs": "plaintiff_name",
            "defendants": "defendant_name",
            "case_number": "case_number",
            "incident_date": "incident_date",
            "school_site": "school_site",
            "assigned_attorney": "assigned_attorney",
            "adjuster_name": "adjuster_name",
            "adjuster_email": "adjuster_email",
            "adjuster_first_name": "adjuster_first_name",
            "claim_number": "claim_number",
            "policy_number": "policy_number",
            "report_subject": "report_subject",
            "client_name": "client_name",
        }

        for src_key, dest_key in var_mapping.items():
            val = all_vars.get(src_key)
            if val:
                # Lists -> joined string
                if isinstance(val, list):
                    val = "; ".join(str(v) for v in val)
                metadata[dest_key] = val

        # Build case_name from parties if not already set
        if "case_name" not in metadata:
            p = metadata.get("plaintiff_name", "")
            d = metadata.get("defendant_name", "")
            if p and d:
                # Use first plaintiff and first defendant for case name
                p_short = p.split(";")[0].strip() if ";" in p else p
                d_short = d.split(";")[0].strip() if ";" in d else d
                metadata["case_name"] = f"{p_short} v. {d_short}"

    except ImportError:
        logger.warning("CaseDataManager not available")
    except Exception as e:
        logger.warning(f"Error reading CaseDataManager: {e}")

    # Supplement with MasterCaseDatabase
    try:
        from icharlotte_core.master_db import MasterCaseDatabase
        db = MasterCaseDatabase()
        case = db.get_case(file_number)
        if case:
            if not metadata.get("plaintiff_name") and case.get("plaintiff_last_name"):
                metadata["plaintiff_name"] = case["plaintiff_last_name"]
            if not metadata.get("assigned_attorney") and case.get("assigned_attorney"):
                metadata["assigned_attorney"] = case["assigned_attorney"]
            if not metadata.get("next_hearing_date") and case.get("next_hearing_date"):
                metadata["next_hearing_date"] = case["next_hearing_date"]
            if not metadata.get("trial_date") and case.get("trial_date"):
                metadata["trial_date"] = case["trial_date"]
    except ImportError:
        logger.warning("MasterCaseDatabase not available")
    except Exception as e:
        logger.warning(f"Error reading MasterCaseDatabase: {e}")

    return metadata


def _scan_agent_outputs(output_dir: str) -> Tuple[Dict[str, str], Dict[str, str]]:
    """
    Scan the AI OUTPUT folder for agent outputs and extract their text content.

    Returns:
        (sections, section_files) where:
            sections: Dict mapping section name -> extracted text
            section_files: Dict mapping section name -> file path
    """
    sections = {}
    section_files = {}

    if not os.path.exists(output_dir):
        return sections, section_files

    files = os.listdir(output_dir)

    for filename in files:
        filepath = os.path.join(output_dir, filename)
        if not os.path.isfile(filepath):
            continue

        # Skip temp files
        if filename.startswith("~$"):
            continue

        # Match filename against agent output patterns
        section_name = _match_filename_to_section(filename)
        if not section_name or section_name.startswith("_"):
            continue

        # Extract text from the file
        text = _extract_text(filepath)
        if not text or len(text.strip()) < 20:
            continue

        if section_name in sections:
            if section_name in USE_LATEST_ONLY:
                # Use only the most recent file (e.g., docket PDFs are snapshots)
                existing_path = section_files[section_name].split("; ")[0]
                if os.path.getmtime(filepath) > os.path.getmtime(existing_path):
                    sections[section_name] = text
                    section_files[section_name] = filepath
            else:
                # Concatenate for sections with distinct sources (med records, discovery)
                sections[section_name] += f"\n\n--- Source: {filename} ---\n\n{text}"
                section_files[section_name] += f"; {filepath}"
        else:
            sections[section_name] = text
            section_files[section_name] = filepath

    return sections, section_files


def _match_filename_to_section(filename: str) -> Optional[str]:
    """Match a filename to its report section using AGENT_OUTPUT_MAP."""
    name_lower = filename.lower()

    for pattern, section, source in AGENT_OUTPUT_MAP:
        if name_lower.startswith(pattern.lower()):
            return section

    return None


def _extract_text(filepath: str) -> Optional[str]:
    """Extract text content from a .docx or .pdf file."""
    ext = os.path.splitext(filepath)[1].lower()

    if ext == ".docx":
        return _extract_docx_text(filepath)
    elif ext == ".pdf":
        return _extract_pdf_text(filepath)
    elif ext == ".txt":
        try:
            with open(filepath, "r", encoding="utf-8") as f:
                return f.read()
        except Exception as e:
            logger.warning(f"Error reading {filepath}: {e}")
            return None
    return None


def _extract_docx_text(filepath: str) -> Optional[str]:
    """Extract text from a .docx file."""
    try:
        doc = Document(filepath)
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))

        return "\n\n".join(paragraphs)
    except Exception as e:
        logger.warning(f"Error extracting text from {filepath}: {e}")
        return None


def _extract_pdf_text(filepath: str) -> Optional[str]:
    """Extract text from a PDF file."""
    try:
        import pypdf
        reader = pypdf.PdfReader(filepath)
        pages = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages.append(text.strip())
        return "\n\n".join(pages)
    except ImportError:
        logger.warning("pypdf not installed, cannot extract PDF text")
        return None
    except Exception as e:
        logger.warning(f"Error extracting PDF text from {filepath}: {e}")
        return None


def _parse_prior_report(report_path: str) -> Dict[str, str]:
    """Parse a prior report .docx to extract its sections for delta comparison."""
    try:
        from Scripts.report_generator.style_library import extract_sections_from_docx
        return extract_sections_from_docx(report_path)
    except Exception as e:
        logger.warning(f"Error parsing prior report: {e}")
        return {}


def _determine_included_sections(
    available_sections: Dict[str, str],
    report_type: str,
    section_files: Dict[str, str],
    prior_report_path: str = None,
) -> List[str]:
    """
    Determine which sections to include based on report type.

    FSR: Include all sections where content exists + placeholder sections.
    STATUS: Always-include set + sections with new content since last report.
    """
    included = []

    if report_type == "FSR":
        # Include all canonical sections (content or placeholder)
        for section in SECTION_ORDER:
            included.append(section)
    else:
        # Status report: always-include + conditionally include new content
        for section in SECTION_ORDER:
            if section in STATUS_REPORT_ALWAYS_INCLUDE:
                included.append(section)
            elif section in available_sections:
                # Check if content is new since last report
                if _is_new_content(section, section_files, prior_report_path):
                    included.append(section)

    return included


def _is_new_content(section_name: str, section_files: Dict[str, str],
                    prior_report_path: str = None) -> bool:
    """Check if a section's source files are newer than the prior report."""
    if not prior_report_path or not os.path.exists(prior_report_path):
        return True  # No prior report = everything is new

    prior_mtime = os.path.getmtime(prior_report_path)

    file_path = section_files.get(section_name, "")
    # Handle multiple file paths joined by "; "
    for fp in file_path.split("; "):
        fp = fp.strip()
        if fp and os.path.exists(fp):
            if os.path.getmtime(fp) > prior_mtime:
                return True

    return False


def _ingest_ai_output(output_dir: str, sections: Dict[str, str],
                      section_files: Dict[str, str]):
    """
    Parse AI_OUTPUT.docx (summarize.py output) and route summaries to
    INVESTIGATION or EXPERTS based on content keywords.

    Each block starts with a filename header, followed by "Generated on:" timestamp,
    then the summary text. Summaries mentioning IME, independent medical examination,
    or preliminary opinion(s) go to EXPERTS; all others go to INVESTIGATION.
    """
    filepath = os.path.join(output_dir, "AI_OUTPUT.docx")
    if not os.path.exists(filepath):
        return

    try:
        doc = Document(filepath)
    except Exception as e:
        logger.warning(f"Could not open AI_OUTPUT.docx: {e}")
        return

    blocks = _parse_ai_output_blocks(doc)
    if not blocks:
        return

    for title, text in blocks:
        # Route based on keywords in title or first 200 chars of text
        search_text = title + " " + text[:200]
        if EXPERTS_KEYWORDS.search(search_text):
            target = "EXPERTS"
        else:
            target = "INVESTIGATION"

        block_text = f"{title}\n\n{text}"
        if target in sections:
            sections[target] += f"\n\n--- Source: {title} ---\n\n{block_text}"
        else:
            sections[target] = block_text
            section_files[target] = filepath

    logger.info(f"Ingested {len(blocks)} summaries from AI_OUTPUT.docx")


def _parse_ai_output_blocks(doc) -> List[Tuple[str, str]]:
    """Parse AI_OUTPUT.docx into (title, text) blocks."""
    blocks = []
    current_title = None
    current_lines = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            if current_title:
                current_lines.append("")
            continue

        # Detect filename headers (has file extension or is all-bold)
        is_header = False
        if re.search(r'\.\w{2,4}$', text):  # Ends with .pdf, .docx, etc.
            is_header = True
        elif para.runs and all(r.bold for r in para.runs if r.text.strip()):
            is_header = True

        if is_header and text != current_title:
            # Save previous block
            if current_title and current_lines:
                content = "\n".join(current_lines).strip()
                # Remove "Generated on:" timestamp line
                content = re.sub(r'^Generated on:.*\n?', '', content,
                                 flags=re.MULTILINE).strip()
                if content:
                    blocks.append((current_title, content))
            current_title = text
            current_lines = []
        elif current_title:
            current_lines.append(text)

    # Save last block
    if current_title and current_lines:
        content = "\n".join(current_lines).strip()
        content = re.sub(r'^Generated on:.*\n?', '', content,
                         flags=re.MULTILINE).strip()
        if content:
            blocks.append((current_title, content))

    return blocks


if __name__ == "__main__":
    import sys
    logging.basicConfig(level=logging.INFO)

    if len(sys.argv) < 2:
        print("Usage: python -m Scripts.report_generator.gather <file_number> [FSR|STATUS]")
        sys.exit(1)

    file_num = sys.argv[1]
    rtype = sys.argv[2] if len(sys.argv) > 2 else "FSR"

    data = gather_case_data(file_num, rtype)
    print(f"\n--- Gather Results for {file_num} ({rtype}) ---")
    print(f"Case path: {data['case_path']}")
    print(f"Metadata keys: {list(data['metadata'].keys())}")
    print(f"Available sections: {data['all_available_sections']}")
    print(f"Included sections: {data['included_sections']}")
    for section, text in data["sections"].items():
        print(f"  {section}: {len(text)} chars")
