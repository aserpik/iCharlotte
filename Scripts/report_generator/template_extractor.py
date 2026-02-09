"""
Template Extractor - One-time utility to create a docxtpl template from an existing report.

Takes a real report .docx, strips case-specific content, and replaces it with
Jinja2 template variables while preserving all Word formatting.

Usage:
    python -m Scripts.report_generator.template_extractor <source_report.docx> [output_template.docx]
"""

import os
import sys
import re
import copy
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# Section headings found in the reports (ALL CAPS, bold+underline)
REPORT_SECTIONS = [
    "FACTUAL BACKGROUND",
    "PROCEDURAL STATUS",
    "INVESTIGATION",
    "DISCOVERY",
    "MEDICAL RECORD REVIEW",
    "EVALUATION OF LIABILITY",
    "EVALUATION OF EXPOSURE",
    "SETTLEMENT",
    "SETTLEMENT STATUS",
    "FURTHER CASE HANDLING",
]

# Template variable names for each section
SECTION_VARS = {
    "FACTUAL BACKGROUND": "factual_background",
    "PROCEDURAL STATUS": "procedural_history",
    "INVESTIGATION": "investigation",
    "DISCOVERY": "discovery",
    "MEDICAL RECORD REVIEW": "medical_record_review",
    "EVALUATION OF LIABILITY": "evaluation_of_liability",
    "EVALUATION OF EXPOSURE": "evaluation_of_exposure",
    "SETTLEMENT": "settlement_status",
    "SETTLEMENT STATUS": "settlement_status",
    "FURTHER CASE HANDLING": "further_case_handling",
}


def is_section_heading(paragraph):
    """Check if a paragraph is a top-level section heading (ALL CAPS, bold+underline)."""
    text = paragraph.text.strip()
    if not text:
        return False
    # Check if text matches a known section heading
    for heading in REPORT_SECTIONS:
        if text.upper() == heading:
            return True
    return False


def get_section_name(paragraph):
    """Get the normalized section name for a heading paragraph."""
    text = paragraph.text.strip().upper()
    for heading in REPORT_SECTIONS:
        if text == heading:
            return heading
    return None


def find_section_ranges(doc):
    """
    Find the paragraph index ranges for each section.
    Returns list of (section_name, start_idx, end_idx) tuples.
    start_idx is the heading paragraph, end_idx is exclusive (next heading or end).
    """
    sections = []
    para_count = len(doc.paragraphs)

    for i, para in enumerate(doc.paragraphs):
        name = get_section_name(para)
        if name:
            sections.append((name, i))

    # Build ranges
    ranges = []
    for idx, (name, start) in enumerate(sections):
        if idx + 1 < len(sections):
            end = sections[idx + 1][1]
        else:
            # Find the closing/signature block
            end = para_count
            for j in range(start + 1, para_count):
                style_name = doc.paragraphs[j].style.name.lower()
                if 'closing' in style_name or 'zclosing' in style_name:
                    end = j
                    break
        ranges.append((name, start, end))

    return ranges


def extract_template(source_path, output_path=None):
    """
    Extract a docxtpl template from an existing report.

    Strategy: Instead of trying to manipulate the complex formatting in-place,
    we keep the document structure intact and replace content paragraphs with
    Jinja2 template variables. The section headings are preserved as-is.
    """
    if output_path is None:
        output_path = os.path.join(
            os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))),
            "templates",
            "litigation_report.docx"
        )

    doc = Document(source_path)

    # --- Phase 1: Replace metadata in the addressee table ---
    # The first table contains recipient info, case numbers, etc.
    if doc.tables:
        table = doc.tables[0]
        _replace_table_metadata(table)

    # --- Phase 2: Replace the date line ---
    # First non-empty paragraph is typically the date
    for para in doc.paragraphs:
        text = para.text.strip()
        if text and _looks_like_date(text):
            _clear_and_set_text(para, "{{ report_date }}")
            break

    # --- Phase 3: Replace salutation ---
    # Reports may use "Dear Name," or just "FirstName," as salutation
    salutation_idx = None
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        if text.startswith("Dear "):
            _clear_and_set_text(para, "Dear {{ adjuster_first_name }},")
            salutation_idx = i
            break
        # Also match standalone first-name salutations (e.g., "Natalie,")
        if text.endswith(",") and len(text.split()) <= 2 and text.rstrip(",")[0].isupper():
            _clear_and_set_text(para, "{{ adjuster_first_name }},")
            salutation_idx = i
            break

    # --- Phase 4: Replace intro paragraph ---
    # The first non-empty paragraph after the salutation
    if salutation_idx is not None:
        for i in range(salutation_idx + 1, len(doc.paragraphs)):
            text = doc.paragraphs[i].text.strip()
            if text and not is_section_heading(doc.paragraphs[i]):
                _clear_and_set_text(doc.paragraphs[i], "{{ intro_paragraph }}")
                break

    # --- Phase 5: Replace section content with template variables ---
    section_ranges = find_section_ranges(doc)

    for section_name, start_idx, end_idx in section_ranges:
        var_name = SECTION_VARS.get(section_name, section_name.lower().replace(" ", "_"))

        # Keep the heading paragraph (start_idx), clear content paragraphs
        # First content paragraph gets the template variable
        first_content = True
        for i in range(start_idx + 1, end_idx):
            para = doc.paragraphs[i]
            if first_content:
                _clear_and_set_text(para, "{{ " + var_name + " }}")
                first_content = False
            else:
                # Clear remaining content paragraphs
                _clear_paragraph(para)

    # --- Phase 6: Replace closing/signature ---
    _replace_closing(doc)

    # Save
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"Template saved to: {output_path}")
    return output_path


def _replace_table_metadata(table):
    """Replace metadata fields in the addressee table with template variables."""
    # Map of text patterns to template variables
    metadata_patterns = [
        # Recipient info
        (r'(?:Via\s+)?(?:E-?mail|EMAIL).*', None),  # Keep "VIA EMAIL" as-is
        # Look for specific field labels and replace their values
    ]

    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                text = para.text.strip()
                if not text:
                    continue

                # Replace specific metadata fields
                # These patterns match "Label: Value" or "Label Value" formats
                replacements = {
                    r'(BS File No\.?\s*:?\s*)(.+)': r'\1{{ file_number }}',
                    r'(BS Client(?:/Insured)?\.?\s*:?\s*)(.+)': r'\1{{ client_name }}',
                    r'((?:ASCIP|Claim)\s*(?:/Claim)?\s*No\.?\s*:?\s*)(.+)': r'\1{{ claim_number }}',
                    r'(Case No\.?\s*:?\s*)(.+)': r'\1{{ case_number }}',
                    r'(Policy No\.?\s*:?\s*)(.+)': r'\1{{ policy_number }}',
                    r'(Date of (?:Loss|Incident)(?:/DOL)?\.?\s*:?\s*)(.+)': r'\1{{ incident_date }}',
                    r'((?:Date|Time) of Incident\.?\s*:?\s*)(.+)': r'\1{{ incident_date }}',
                    r'(School Site\.?\s*:?\s*)(.+)': r'\1{{ school_site }}',
                    r'(Subject\s*:?\s*)(.+)': r'\1{{ report_subject }}',
                }

                for pattern, replacement in replacements.items():
                    if re.search(pattern, text, re.IGNORECASE):
                        new_text = re.sub(pattern, replacement, text, flags=re.IGNORECASE)
                        _clear_and_set_text(para, new_text)
                        break

                # Replace recipient name/email lines
                # These are typically the first few lines in the table
                if '@' in text and '.' in text:
                    _clear_and_set_text(para, "{{ adjuster_email }}")
                elif text.startswith("Re:") or text.startswith("RE:"):
                    _clear_and_set_text(para, "Re:\t{{ case_name }}")
                elif _looks_like_person_name(text) and not any(
                    label in text.lower() for label in ['file', 'case', 'claim', 'date', 'site', 'subject', 'policy']
                ):
                    # Likely the recipient name
                    _clear_and_set_text(para, "{{ adjuster_name }}")


def _replace_closing(doc):
    """Replace the closing signature block."""
    in_closing = False
    for para in doc.paragraphs:
        style = para.style.name.lower()
        if 'closing' in style or 'zclosing' in style:
            in_closing = True
        if in_closing:
            text = para.text.strip()
            if text.startswith("Very truly yours") or text.startswith("Sincerely"):
                continue  # Keep the closing phrase
            elif text == "BORDIN SEMMER LLP":
                continue  # Keep firm name
            elif text and not text.startswith("BORDIN"):
                # This is likely an attorney name - replace with template var
                _clear_and_set_text(para, "{{ attorney_names }}")
                break  # Only need to replace once


def _clear_and_set_text(paragraph, new_text):
    """Clear all runs in a paragraph and set new text, preserving the first run's formatting."""
    if not paragraph.runs:
        paragraph.text = new_text
        return

    # Save first run's formatting
    first_run = paragraph.runs[0]

    # Clear all runs
    for run in paragraph.runs:
        run.text = ""

    # Set text on first run
    first_run.text = new_text


def _clear_paragraph(paragraph):
    """Remove all content from a paragraph, making it empty."""
    for run in paragraph.runs:
        run.text = ""
    paragraph.text = ""


def _looks_like_date(text):
    """Check if text looks like a date line (e.g., 'December 16, 2025')."""
    date_patterns = [
        r'^(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{4}$',
        r'^\d{1,2}/\d{1,2}/\d{4}$',
    ]
    return any(re.match(p, text.strip(), re.IGNORECASE) for p in date_patterns)


def _looks_like_person_name(text):
    """Heuristic: check if text looks like a person's name."""
    # Simple heuristic: 2-4 words, each capitalized, no special chars
    words = text.strip().split()
    if 2 <= len(words) <= 5:
        return all(w[0].isupper() for w in words if w)
    return False


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python -m Scripts.report_generator.template_extractor <source_report.docx> [output.docx]")
        sys.exit(1)

    source = sys.argv[1]
    output = sys.argv[2] if len(sys.argv) > 2 else None
    extract_template(source, output)
