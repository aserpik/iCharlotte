"""
Report Validator - Check generated reports against gold standard formatting rules.

Compares a generated .docx report against the reference profile extracted from
gold standard reports. Reports formatting violations at the paragraph level.

Usage:
    # Standalone
    python -m Scripts.report_generator.validate path/to/report.docx
    python -m Scripts.report_generator.validate report.docx --verbose
    python -m Scripts.report_generator.validate report.docx --profile custom.json

    # In pipeline
    from Scripts.report_generator.validate import validate_report
    result = validate_report("path/to/report.docx")
    result.print_summary()
"""

import os
import re
import sys
import json
import logging
import argparse
from dataclasses import dataclass, field
from typing import Any, Dict, List, Optional

from docx import Document
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)

PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
DEFAULT_PROFILE_PATH = os.path.join(PROJECT_ROOT, "config", "report_reference_profile.json")

# Known section headings (all caps)
SECTION_HEADINGS = {
    "FACTUAL BACKGROUND", "PROCEDURAL HISTORY", "INVESTIGATION",
    "DISCOVERY", "MEDICAL RECORD REVIEW", "EVALUATION OF LIABILITY",
    "EVALUATION OF EXPOSURE", "EXPERTS", "SETTLEMENT", "SETTLEMENT STATUS",
    "LITIGATION BUDGET", "FURTHER CASE HANDLING",
}


# ---------------------------------------------------------------------------
# Data classes
# ---------------------------------------------------------------------------

@dataclass
class Finding:
    """A single validation finding."""
    level: str          # "PASS", "FAIL", "WARN"
    category: str       # Rule category (e.g., "section_heading", "body_paragraph")
    message: str        # Human-readable description
    paragraph_index: Optional[int] = None
    expected: Any = None
    actual: Any = None

    def __str__(self):
        loc = f" (para {self.paragraph_index})" if self.paragraph_index is not None else ""
        return f"[{self.level}] {self.category}{loc}: {self.message}"


@dataclass
class ValidationResult:
    """Aggregated validation results."""
    doc_path: str
    findings: List[Finding] = field(default_factory=list)

    @property
    def pass_count(self):
        return sum(1 for f in self.findings if f.level == "PASS")

    @property
    def fail_count(self):
        return sum(1 for f in self.findings if f.level == "FAIL")

    @property
    def warn_count(self):
        return sum(1 for f in self.findings if f.level == "WARN")

    def print_summary(self, verbose=False):
        """Print results to console."""
        name = os.path.basename(self.doc_path)
        print(f"\n=== Report Validation: {name} ===")
        for f in self.findings:
            if verbose or f.level != "PASS":
                print(f"  {f}")
        print(f"\nResults: {self.pass_count} PASS, {self.fail_count} FAIL, {self.warn_count} WARN\n")


# ---------------------------------------------------------------------------
# Helper functions
# ---------------------------------------------------------------------------

def _get_twips(emu_value) -> Optional[int]:
    """Convert EMU value to twips. Returns None if value is None."""
    if emu_value is None:
        return None
    return int(emu_value / 635)  # 1 twip = 635 EMU (approx)


def _get_xml_attr(element, tag, attr):
    """Get an XML attribute value from a child element. Returns None if not found."""
    if element is None:
        return None
    child = element.find(qn(tag))
    if child is None:
        return None
    return child.get(qn(attr)) if qn(attr) in child.attrib else child.get(attr)


def _get_pPr(para):
    """Get paragraph properties element, or None."""
    return para._element.find(qn('w:pPr'))


def _get_rPr_from_pPr(pPr):
    """Get run properties from paragraph properties."""
    if pPr is None:
        return None
    return pPr.find(qn('w:rPr'))


def _get_first_run_rPr(para):
    """Get run properties from the first non-empty run."""
    for run in para.runs:
        if run.text and run.text.strip():
            return run._element.find(qn('w:rPr'))
    return None


def _is_bold(rPr):
    """Check if run properties indicate bold."""
    if rPr is None:
        return False
    b = rPr.find(qn('w:b'))
    if b is None:
        return False
    val = b.get(qn('w:val'))
    return val is None or val not in ('0', 'false')


def _is_underline(rPr):
    """Check if run properties indicate underline."""
    if rPr is None:
        return False
    u = rPr.find(qn('w:u'))
    if u is None:
        return False
    val = u.get(qn('w:val'))
    return val is not None and val != 'none'


def _get_indent_firstLine(pPr):
    """Get first-line indent in twips from paragraph properties."""
    if pPr is None:
        return None
    ind = pPr.find(qn('w:ind'))
    if ind is None:
        return None
    val = ind.get(qn('w:firstLine'))
    return int(val) if val else None


def _get_indent_left(pPr):
    """Get left indent in twips from paragraph properties."""
    if pPr is None:
        return None
    ind = pPr.find(qn('w:ind'))
    if ind is None:
        return None
    val = ind.get(qn('w:left'))
    return int(val) if val else None


def _get_indent_hanging(pPr):
    """Get hanging indent in twips from paragraph properties."""
    if pPr is None:
        return None
    ind = pPr.find(qn('w:ind'))
    if ind is None:
        return None
    val = ind.get(qn('w:hanging'))
    return int(val) if val else None


def _get_spacing_after(pPr):
    """Get space-after in twips from paragraph properties."""
    if pPr is None:
        return None
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        return None
    val = spacing.get(qn('w:after'))
    return int(val) if val else None


def _get_spacing_before(pPr):
    """Get space-before in twips from paragraph properties."""
    if pPr is None:
        return None
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        return None
    val = spacing.get(qn('w:before'))
    return int(val) if val else None


def _get_style_name(para):
    """Get the style name of a paragraph."""
    pPr = _get_pPr(para)
    if pPr is None:
        return None
    pStyle = pPr.find(qn('w:pStyle'))
    if pStyle is None:
        return None
    return pStyle.get(qn('w:val'))


def _para_text(para) -> str:
    """Get stripped paragraph text."""
    return para.text.strip() if para.text else ""


def _is_section_heading(text: str) -> bool:
    """Check if text is a known section heading."""
    return text.strip().rstrip(":") in SECTION_HEADINGS


def load_profile(profile_path: str = None) -> Dict:
    """Load the reference profile JSON."""
    if profile_path is None:
        profile_path = DEFAULT_PROFILE_PATH
    if not os.path.exists(profile_path):
        logger.warning(f"Reference profile not found: {profile_path}")
        return {}
    with open(profile_path, 'r', encoding='utf-8') as f:
        return json.load(f)


# ---------------------------------------------------------------------------
# Validation rules
# ---------------------------------------------------------------------------

def check_section_headings(doc: Document, profile: Dict) -> List[Finding]:
    """Check that section headings are bold, underlined, all-caps, flush left."""
    findings = []
    heading_profile = profile.get("section_heading", {})
    expected_bold = heading_profile.get("bold", True)
    expected_underline = heading_profile.get("underline", "single")

    found_headings = []
    for i, para in enumerate(doc.paragraphs):
        text = _para_text(para)
        if not _is_section_heading(text):
            continue

        found_headings.append(text)
        pPr = _get_pPr(para)

        # Check bold - from run properties (either paragraph rPr or first run rPr)
        rPr = _get_rPr_from_pPr(pPr)
        run_rPr = _get_first_run_rPr(para)
        is_bold = _is_bold(rPr) or _is_bold(run_rPr)

        if expected_bold and not is_bold:
            findings.append(Finding(
                "FAIL", "section_heading",
                f'"{text}" is not bold',
                paragraph_index=i, expected="bold", actual="not bold"
            ))

        # Check underline
        is_underlined = _is_underline(rPr) or _is_underline(run_rPr)
        if expected_underline and not is_underlined:
            findings.append(Finding(
                "FAIL", "section_heading",
                f'"{text}" is not underlined',
                paragraph_index=i, expected="underline", actual="no underline"
            ))

        # Check all-caps
        if text != text.upper():
            findings.append(Finding(
                "FAIL", "section_heading",
                f'"{text}" is not all-caps',
                paragraph_index=i, expected="ALL CAPS", actual=text
            ))

        # Check first-line indent = 0 (flush left)
        first_line = _get_indent_firstLine(pPr)
        if first_line is not None and first_line != 0:
            findings.append(Finding(
                "FAIL", "section_heading",
                f'"{text}" has first-line indent {first_line} twips, expected 0 (flush left)',
                paragraph_index=i, expected=0, actual=first_line
            ))

    if found_headings:
        findings.append(Finding(
            "PASS", "section_heading",
            f"Found {len(found_headings)} section headings: {', '.join(found_headings)}"
        ))
    else:
        findings.append(Finding(
            "WARN", "section_heading",
            "No section headings found in document"
        ))

    return findings


def check_body_paragraphs(doc: Document, profile: Dict) -> List[Finding]:
    """Check body paragraph formatting: first-line indent, font."""
    findings = []
    body_profile = profile.get("body_paragraph", {})
    expected_indent = body_profile.get("first_line_indent_twips", 720)
    expected_font = body_profile.get("font_name", "Times New Roman")

    body_paras = []
    wrong_indent = []
    wrong_font = []

    for i, para in enumerate(doc.paragraphs):
        text = _para_text(para)
        if not text or len(text) < 40:
            continue
        if _is_section_heading(text):
            continue
        style = _get_style_name(para)
        if style in ("zClosing", "zDelivery", "zRe", "zAddressee", "FlushLeft"):
            continue

        body_paras.append(i)

        # Check first-line indent
        pPr = _get_pPr(para)
        first_line = _get_indent_firstLine(pPr)

        # If style is Body, first-line indent is inherited (720) unless overridden
        if style == "Body" and first_line is None:
            first_line = 720  # Inherited from style

        # Check explicit indent if present
        if first_line is not None and first_line != expected_indent:
            # Allow 0 indent for subheadings and special paragraphs
            left = _get_indent_left(pPr)
            hanging = _get_indent_hanging(pPr)
            if left or hanging:
                continue  # This is a subheading with hanging indent, skip

            # Allow 0 indent if paragraph starts with tab (manual indent)
            raw_text = para.text or ""
            if first_line == 0 and (raw_text.startswith("\t") or raw_text.startswith("  \t")):
                continue

            wrong_indent.append((i, first_line))

        # Check font on first run
        for run in para.runs:
            if run.text and run.text.strip():
                if run.font.name and run.font.name != expected_font:
                    wrong_font.append((i, run.font.name))
                break

    if body_paras:
        findings.append(Finding(
            "PASS", "body_paragraph",
            f"Checked {len(body_paras)} body paragraphs"
        ))

    if wrong_indent:
        for idx, actual in wrong_indent[:5]:  # Report first 5
            text_preview = _para_text(doc.paragraphs[idx])[:60]
            findings.append(Finding(
                "FAIL", "body_paragraph",
                f'first-line indent={actual} twips, expected {expected_indent}: "{text_preview}..."',
                paragraph_index=idx, expected=expected_indent, actual=actual
            ))
        if len(wrong_indent) > 5:
            findings.append(Finding(
                "FAIL", "body_paragraph",
                f"... and {len(wrong_indent) - 5} more paragraphs with wrong indent"
            ))
    else:
        findings.append(Finding(
            "PASS", "body_paragraph",
            f"All body paragraphs have correct first-line indent ({expected_indent} twips)"
        ))

    if wrong_font:
        fonts = set(f for _, f in wrong_font)
        findings.append(Finding(
            "WARN", "body_paragraph",
            f"Found non-standard fonts in {len(wrong_font)} paragraphs: {', '.join(fonts)}",
            expected=expected_font, actual=list(fonts)
        ))
    else:
        findings.append(Finding(
            "PASS", "body_paragraph",
            f"All body paragraphs use {expected_font}"
        ))

    return findings


def check_metadata_table(doc: Document, profile: Dict) -> List[Finding]:
    """Check metadata table structure: columns, width, borders."""
    findings = []
    table_profile = profile.get("metadata_table", {})

    if not doc.tables:
        findings.append(Finding("FAIL", "metadata_table", "No tables found in document"))
        return findings

    table = doc.tables[0]  # Metadata table is always first
    tbl = table._tbl

    # Check column count
    expected_cols = table_profile.get("columns", 2)
    actual_cols = len(table.columns)
    if actual_cols == expected_cols:
        findings.append(Finding("PASS", "metadata_table", f"Column count: {actual_cols}"))
    else:
        findings.append(Finding(
            "WARN", "metadata_table",
            f"Column count: {actual_cols}, expected {expected_cols}",
            expected=expected_cols, actual=actual_cols
        ))

    # Check table width
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is not None:
        tblW = tblPr.find(qn('w:tblW'))
        if tblW is not None:
            width = tblW.get(qn('w:w'))
            width_type = tblW.get(qn('w:type'))
            if width and width_type == 'dxa':
                expected_width = table_profile.get("total_width_twips", 10579)
                actual_width = int(width)
                tolerance = 200  # Allow ~0.14" variance
                if abs(actual_width - expected_width) <= tolerance:
                    findings.append(Finding(
                        "PASS", "metadata_table",
                        f"Table width: {actual_width} twips (expected ~{expected_width})"
                    ))
                else:
                    findings.append(Finding(
                        "WARN", "metadata_table",
                        f"Table width: {actual_width} twips, expected ~{expected_width}",
                        expected=expected_width, actual=actual_width
                    ))

        # Check table indent
        tblInd = tblPr.find(qn('w:tblInd'))
        expected_indent = table_profile.get("table_indent_twips", 0)
        if tblInd is not None:
            actual_indent = int(tblInd.get(qn('w:w'), '0'))
            if actual_indent != expected_indent:
                findings.append(Finding(
                    "WARN", "metadata_table",
                    f"Table indent: {actual_indent} twips, gold standards have {expected_indent}",
                    expected=expected_indent, actual=actual_indent
                ))
            else:
                findings.append(Finding("PASS", "metadata_table", f"Table indent: {actual_indent}"))
        elif expected_indent == 0:
            findings.append(Finding("PASS", "metadata_table", "No table indent (matches gold standard)"))

    # Check for Re: line
    has_re = False
    for row in table.rows:
        for cell in row.cells:
            if "Re:" in cell.text:
                has_re = True
                break
    if has_re:
        findings.append(Finding("PASS", "metadata_table", "Re: line found in metadata table"))
    else:
        findings.append(Finding("WARN", "metadata_table", "No 'Re:' line found in metadata table"))

    return findings


def check_closing(doc: Document, profile: Dict) -> List[Finding]:
    """Check closing section: style, attorney names, firm name."""
    findings = []
    closing_profile = profile.get("closing", {})
    expected_style = closing_profile.get("style", "zClosing")
    expected_firm = closing_profile.get("firm_name", "BORDIN SEMMER LLP")
    expected_valedictions = closing_profile.get("valediction", ["Sincerely,", "Very truly yours,"])

    # Find closing valediction
    valediction_found = False
    firm_found = False
    closing_style_correct = True

    for i, para in enumerate(doc.paragraphs):
        text = _para_text(para)
        style = _get_style_name(para)

        if any(text.startswith(v.rstrip(",")) for v in expected_valedictions):
            valediction_found = True
            if style != expected_style:
                closing_style_correct = False
                findings.append(Finding(
                    "FAIL", "closing",
                    f'"{text}" uses style "{style}", expected "{expected_style}"',
                    paragraph_index=i, expected=expected_style, actual=style
                ))

        if text == expected_firm:
            firm_found = True
            if style != expected_style:
                closing_style_correct = False
                findings.append(Finding(
                    "FAIL", "closing",
                    f'"{text}" uses style "{style}", expected "{expected_style}"',
                    paragraph_index=i, expected=expected_style, actual=style
                ))

    if valediction_found:
        findings.append(Finding("PASS", "closing", "Valediction found"))
    else:
        findings.append(Finding("FAIL", "closing", "No valediction found (Sincerely/Very truly yours)"))

    if firm_found:
        findings.append(Finding("PASS", "closing", f'Firm name "{expected_firm}" found'))
    else:
        findings.append(Finding("FAIL", "closing", f'Firm name "{expected_firm}" not found'))

    # Check attorney names
    expected_attorneys = closing_profile.get("attorney_names", [])
    for name in expected_attorneys:
        found = any(name in _para_text(p) for p in doc.paragraphs)
        if found:
            findings.append(Finding("PASS", "closing", f'Attorney "{name}" found'))
        else:
            findings.append(Finding("WARN", "closing", f'Attorney "{name}" not found'))

    return findings


def check_no_empty_paragraphs(doc: Document, profile: Dict) -> List[Finding]:
    """Check for excessive consecutive empty paragraphs (template cleanup artifacts)."""
    findings = []
    consecutive_empty = 0
    max_consecutive = 0
    worst_index = 0

    for i, para in enumerate(doc.paragraphs):
        text = _para_text(para)
        style = _get_style_name(para)

        # Allow empty paragraphs in closing section
        if style == "zClosing":
            consecutive_empty = 0
            continue

        if not text:
            consecutive_empty += 1
            if consecutive_empty > max_consecutive:
                max_consecutive = consecutive_empty
                worst_index = i
        else:
            consecutive_empty = 0

    if max_consecutive > 3:
        findings.append(Finding(
            "WARN", "empty_paragraphs",
            f"Found {max_consecutive} consecutive empty paragraphs near paragraph {worst_index}",
            paragraph_index=worst_index, expected="<= 3", actual=max_consecutive
        ))
    else:
        findings.append(Finding(
            "PASS", "empty_paragraphs",
            f"No excessive empty paragraph runs (max consecutive: {max_consecutive})"
        ))

    return findings


def check_salutation(doc: Document, profile: Dict) -> List[Finding]:
    """Check salutation paragraph formatting."""
    findings = []
    sal_profile = profile.get("salutation", {})

    for i, para in enumerate(doc.paragraphs):
        text = _para_text(para)
        if text.startswith("Dear ") and text.endswith(","):
            pPr = _get_pPr(para)

            # Check that first-line indent is 0 or not set
            first_line = _get_indent_firstLine(pPr)
            if first_line and first_line > 0:
                findings.append(Finding(
                    "FAIL", "salutation",
                    f'Salutation "{text}" has first-line indent {first_line}, expected 0',
                    paragraph_index=i, expected=0, actual=first_line
                ))
            else:
                findings.append(Finding("PASS", "salutation", f'Salutation "{text}" formatting OK'))

            # Check spacing
            space_after = _get_spacing_after(pPr)
            expected_after = sal_profile.get("space_after_twips", 0)
            if space_after is not None and space_after != expected_after:
                findings.append(Finding(
                    "WARN", "salutation",
                    f"Salutation space_after={space_after}, gold standard={expected_after}",
                    paragraph_index=i, expected=expected_after, actual=space_after
                ))

            return findings

    findings.append(Finding("WARN", "salutation", "No salutation paragraph found"))
    return findings


def check_intro_paragraph(doc: Document, profile: Dict) -> List[Finding]:
    """Check intro paragraph (As you recall...) formatting."""
    findings = []
    intro_profile = profile.get("intro_paragraph", {})
    expected_indent = intro_profile.get("first_line_indent_twips", 720)

    for i, para in enumerate(doc.paragraphs):
        text = _para_text(para)
        if "As you recall" in text or "As you know" in text:
            pPr = _get_pPr(para)
            style = _get_style_name(para)

            # Check style
            expected_style = intro_profile.get("style", "Body")
            if style == expected_style:
                findings.append(Finding("PASS", "intro_paragraph", f"Style: {style}"))
            elif style is not None:
                findings.append(Finding(
                    "WARN", "intro_paragraph",
                    f'Style: "{style}", expected "{expected_style}"',
                    paragraph_index=i, expected=expected_style, actual=style
                ))

            # Check indent
            first_line = _get_indent_firstLine(pPr)
            if style == "Body" and first_line is None:
                first_line = 720  # Inherited

            if first_line == expected_indent:
                findings.append(Finding(
                    "PASS", "intro_paragraph",
                    f"First-line indent: {first_line} twips"
                ))
            elif first_line is not None:
                findings.append(Finding(
                    "FAIL", "intro_paragraph",
                    f"First-line indent: {first_line} twips, expected {expected_indent}",
                    paragraph_index=i, expected=expected_indent, actual=first_line
                ))

            return findings

    findings.append(Finding("WARN", "intro_paragraph", "No intro paragraph found"))
    return findings


def check_subheadings(doc: Document, profile: Dict) -> List[Finding]:
    """Check subheading formatting (bold, underlined, proper indentation)."""
    findings = []
    l1_profile = profile.get("subheading_l1", {})

    subheadings_found = 0
    subheading_issues = []

    for i, para in enumerate(doc.paragraphs):
        text = _para_text(para)
        if not text or len(text) > 80:
            continue
        if _is_section_heading(text):
            continue

        style = _get_style_name(para)
        pPr = _get_pPr(para)

        # Detect subheadings: short text that is bold+underlined
        rPr = _get_rPr_from_pPr(pPr)
        run_rPr = _get_first_run_rPr(para)
        is_bold = _is_bold(rPr) or _is_bold(run_rPr)
        is_ul = _is_underline(rPr) or _is_underline(run_rPr)

        if is_bold and is_ul and len(text) < 60:
            subheadings_found += 1

            # Check indentation - subheadings should have left + hanging indent
            left = _get_indent_left(pPr)
            hanging = _get_indent_hanging(pPr)

            # Lettered/numbered subheadings should have some form of hanging indent
            has_prefix = bool(re.match(r'^[A-Z]\.\s|^\d+\.\s', text))
            if has_prefix and not hanging and left is None:
                subheading_issues.append((i, text, "no hanging indent"))

    if subheadings_found > 0:
        findings.append(Finding(
            "PASS", "subheading",
            f"Found {subheadings_found} subheadings (bold + underlined)"
        ))
    else:
        findings.append(Finding(
            "WARN", "subheading",
            "No subheadings detected"
        ))

    for idx, text, issue in subheading_issues[:5]:
        findings.append(Finding(
            "WARN", "subheading",
            f'"{text}" ({issue})',
            paragraph_index=idx
        ))

    return findings


def check_delivery_line(doc: Document, profile: Dict) -> List[Finding]:
    """Check VIA EMAIL delivery line."""
    findings = []
    delivery_profile = profile.get("delivery_line", {})
    expected_text = delivery_profile.get("text", "VIA EMAIL")
    expected_style = delivery_profile.get("style", "zDelivery")

    for i, para in enumerate(doc.paragraphs):
        text = _para_text(para)
        if text == expected_text:
            style = _get_style_name(para)
            if style == expected_style:
                findings.append(Finding("PASS", "delivery_line", f'"{text}" with style {style}'))
            else:
                findings.append(Finding(
                    "WARN", "delivery_line",
                    f'"{text}" uses style "{style}", expected "{expected_style}"',
                    paragraph_index=i, expected=expected_style, actual=style
                ))
            return findings

    findings.append(Finding("WARN", "delivery_line", f'"{expected_text}" not found'))
    return findings


def check_all_caps_names(doc: Document, profile: Dict) -> List[Finding]:
    """Check that client/party names are not in ALL CAPS (should be title case)."""
    findings = []

    # Check metadata table for all-caps values
    if doc.tables:
        table = doc.tables[0]
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                # Skip labels and short text
                if ":" in text or len(text) < 5:
                    continue
                # Check for all-caps strings that look like names
                words = text.split()
                all_caps_words = [w for w in words if w.isupper() and len(w) > 2
                                  and w not in ("LLP", "LLC", "INC", "DBA", "USA", "VIA")]
                if len(all_caps_words) >= 2:
                    findings.append(Finding(
                        "FAIL", "all_caps_names",
                        f'All-caps name in metadata: "{text}" - should be title case',
                        expected="Title Case", actual=text
                    ))

    if not any(f.level == "FAIL" for f in findings):
        findings.append(Finding("PASS", "all_caps_names", "No all-caps names found in metadata"))

    return findings


# ---------------------------------------------------------------------------
# Rule registry - add new rules by appending to this list
# ---------------------------------------------------------------------------

RULES = [
    check_section_headings,
    check_body_paragraphs,
    check_metadata_table,
    check_salutation,
    check_intro_paragraph,
    check_subheadings,
    check_closing,
    check_no_empty_paragraphs,
    check_delivery_line,
    check_all_caps_names,
]


# ---------------------------------------------------------------------------
# Main entry point
# ---------------------------------------------------------------------------

def validate_report(doc_path: str, profile_path: str = None) -> ValidationResult:
    """
    Validate a generated report against the reference profile.

    Args:
        doc_path: Path to the .docx report to validate
        profile_path: Path to reference profile JSON (uses default if None)

    Returns:
        ValidationResult with all findings
    """
    if not os.path.exists(doc_path):
        result = ValidationResult(doc_path)
        result.findings.append(Finding("FAIL", "file", f"File not found: {doc_path}"))
        return result

    profile = load_profile(profile_path)
    doc = Document(doc_path)
    result = ValidationResult(doc_path)

    for rule in RULES:
        try:
            findings = rule(doc, profile)
            result.findings.extend(findings)
        except Exception as e:
            result.findings.append(Finding(
                "WARN", rule.__name__,
                f"Rule raised exception: {e}"
            ))
            logger.exception(f"Error in rule {rule.__name__}")

    return result


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(description="Validate report formatting")
    parser.add_argument("doc_path", help="Path to the .docx report to validate")
    parser.add_argument("--profile", default=None, help="Path to reference profile JSON")
    parser.add_argument("--verbose", "-v", action="store_true", help="Show PASS results too")
    args = parser.parse_args()

    result = validate_report(args.doc_path, args.profile)
    result.print_summary(verbose=args.verbose)

    # Exit with non-zero if any failures
    sys.exit(1 if result.fail_count > 0 else 0)


if __name__ == "__main__":
    main()
