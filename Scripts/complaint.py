import os
import sys
import logging
import json
import datetime
import glob
import re
import subprocess
import time
import shutil
import gc
from typing import Optional, Dict, Any

# Third-party imports
try:
    from docx import Document
    from docx.shared import Pt
    from pypdf import PdfReader
    import pytesseract
    from pdf2image import convert_from_path
except ImportError as e:
    print(f"Critical Error: Missing dependency {e}. Please ensure python-docx, pypdf, pytesseract, pdf2image are installed.")
    sys.exit(1)

# Import Case Data Manager and Gemini Utils
try:
    from case_data_manager import CaseDataManager
    import gemini_utils
except ImportError:
    # If run from root, Scripts/ might not be in path for imports
    sys.path.append(os.path.join(os.getcwd(), 'Scripts'))
    from case_data_manager import CaseDataManager
    import gemini_utils

# --- Configuration ---
# Directories
BASE_PATH_WIN = r"Z:\\Shared\\Current Clients"
# GEMINI_DATA_DIR is managed by CaseDataManager
LOG_FILE = os.path.join(os.getcwd(), "Complaint_activity.log")
SCRIPTS_DIR = os.path.join(os.getcwd(), "Scripts")

# Prompts
FACTUAL_BG_PROMPT_FILE = os.path.join(SCRIPTS_DIR, "Factual_Background.txt")

# Initialize Data Manager
data_manager = CaseDataManager()

# OCR Settings
OCR_AVAILABLE = False
POPPLER_PATH = None

if os.name == 'nt':
    # Tesseract Path
    tesseract_paths = [
        r"C:\\Program Files\\Tesseract-OCR\\tesseract.exe",
        os.path.expanduser(r"~\\AppData\\Local\\Programs\\Tesseract-OCR\\tesseract.exe")
    ]
    for p in tesseract_paths:
        if os.path.exists(p):
            pytesseract.pytesseract.tesseract_cmd = p
            OCR_AVAILABLE = True
            break
    
    # Poppler Path
    poppler_paths = [
        r"C:\\Program Files\\poppler\\Library\\bin",
        r"C:\\Program Files\\poppler-0.68.0\\bin",
        r"C:\\Program Files (x86)\\poppler\\Library\\bin" # Common alternative
    ]
    for p in poppler_paths:
        if os.path.exists(p):
            POPPLER_PATH = p
            break
else:
    # Linux/Mac assumptions (usually in path)
    if shutil.which("tesseract"):
        OCR_AVAILABLE = True
    if shutil.which("pdftoppm"):
        POPPLER_PATH = None # pdf2image finds it on PATH

# Setup Logging
logging.basicConfig(
    filename=LOG_FILE,
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    datefmt='%Y-%m-%d %H:%M:%S'
)

# --- Centralized Path Resolution ---
import sys
import os
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from icharlotte_core.utils import get_case_path, log_event

def find_complaint_file(case_path: str) -> Optional[str]:
    """
    Finds the most recent/relevant complaint file using scoring logic.
    Priority: 4AC/TAC > SAC > FAC > Amended Complaint > Complaint > S&C
    Restricted to PLEADINGS or PLEADING directories.
    """
    search_dirs = [os.path.join(case_path, "PLEADINGS"), os.path.join(case_path, "PLEADING")]
    valid_search_paths = [d for d in search_dirs if os.path.exists(d)]
    
    if not valid_search_paths:
        log_event(f"No PLEADINGS or PLEADING directory found at {case_path}", level="warning")
        return None

    # 1. Gather all potential candidates (PDFs in PLEADINGS only)
    candidates = []
    
    # Helper to scan a directory
    def scan_dir(directory, recursive=True):
        found = []
        try:
            if recursive:
                for root, _, files in os.walk(directory):
                    for file in files:
                        if file.lower().endswith(".pdf"):
                            found.append(os.path.join(root, file))
            else:
                with os.scandir(directory) as it:
                    for entry in it:
                        if entry.is_file() and entry.name.lower().endswith(".pdf"):
                            found.append(entry.path)
        except OSError as e:
            log_event(f"Error scanning {directory}: {e}", level="warning")
        return found

    # Scan PLEADINGS (Recursive)
    for d in valid_search_paths:
        candidates.extend(scan_dir(d, recursive=True))
    
    # 2. Filter and Score Candidates
    # Terms to identify relevant files - Expanded for robustness
    valid_terms = [
        "complaint", "s&c", "summons", "fac", "sac", "tac", "3ac", "4ac", "amended", 
        "cmpl", "cmp", "pleading", "pld"
    ]
    ignore_terms = ["motion", "response", "demurrer", "answer", "reply", "cross", "notice", "proof", "pos", "svc", "service"]
    
    scored_files = []
    
    for file_path in candidates:
        name = os.path.basename(file_path).lower()
        
        # Must contain at least one valid term
        if not any(term in name for term in valid_terms):
            continue
            
        # Must NOT contain ignore terms (unless "cross-complaint"? strict for now)
        if any(term in name for term in ignore_terms):
            continue
            
        # Scoring Logic
        score = 0
        if "4ac" in name or "4th amended" in name or "fourth amended" in name:
            score = 6
        elif "tac" in name or "3ac" in name or "3rd amended" in name or "third amended" in name:
            score = 5
        elif "sac" in name or "2nd amended" in name or "second amended" in name:
            score = 4
        elif "fac" in name or "1st amended" in name or "first amended" in name:
            score = 3
        elif "amended" in name: # Generic amended
            score = 2
        elif "complaint" in name or "s&c" in name or "cmpl" in name or "cmp" in name:
            score = 1
            
        # Get modification time for tie-breaking
        try:
            mtime = os.path.getmtime(file_path)
        except OSError:
            mtime = 0
            
        scored_files.append({
            "path": file_path,
            "name": name,
            "score": score,
            "mtime": mtime
        })
    
    if not scored_files:
        log_event(f"No valid complaint files found in PLEADINGS.", level="warning")
        return None
        
    # 3. Sort: Highest Score first, then Newest Date first
    scored_files.sort(key=lambda x: (-x['score'], -x['mtime']))
    
    best_match = scored_files[0]
    log_event(f"Selected Best Complaint: {best_match['name']} (Score: {best_match['score']})")
    
    return best_match['path']

def find_insurance_init_file(case_path: str) -> Optional[str]:
    """Finds insurance initiation document in CLAIM folder (PDF with 'INT')."""
    claim_dir = os.path.join(case_path, "CLAIM")
    if not os.path.exists(claim_dir):
        log_event(f"CLAIM directory not found at {claim_dir}", level="warning")
        return None

    for root, _, files in os.walk(claim_dir):
        for file in files:
            if file.lower().endswith(".pdf") and "int" in file.lower():
                return os.path.join(root, file)
    
    log_event("No Insurance Initiation (INT) document found.", level="warning")
    return None

def find_caption_doc(case_path: str) -> Optional[str]:
    """
    Finds a document with 'CAPTION' in title.
    Priority: Root folder > Subfolders. .docx > .pdf.
    """
    # 1. Check Root Folder First
    root_candidates = []
    try:
        with os.scandir(case_path) as it:
            for entry in it:
                if entry.is_file() and "caption" in entry.name.lower():
                    if entry.name.lower().endswith((".docx", ".doc", ".pdf")):
                        root_candidates.append(entry.path)
    except OSError:
        pass

    if root_candidates:
        # Prioritize docx
        docx_files = [f for f in root_candidates if f.lower().endswith((".docx", ".doc"))]
        if docx_files:
            return docx_files[0]
        return root_candidates[0]

    # 2. Fallback to Recursive Search
    for root, _, files in os.walk(case_path):
        # Skip the root as we just checked it
        if os.path.abspath(root) == os.path.abspath(case_path):
            continue
        for file in files:
            if "caption" in file.lower() and file.lower().endswith((".docx", ".doc", ".pdf")):
                return os.path.join(root, file)
    return None

def extract_text_from_file(file_path: str, force_ocr_first_page: bool = False) -> str:
    """Extracts text, falling back to OCR if needed. Supports DOCX tables."""
    if not file_path or not os.path.exists(file_path):
        return ""

    log_event(f"Extracting text from: {file_path}")
    text = ""
    ext = os.path.splitext(file_path)[1].lower()

    try:
        if ext == ".pdf":
            reader = PdfReader(file_path)
            
            def ocr_page(page_idx):
                if not OCR_AVAILABLE:
                    return ""
                try:
                    images = convert_from_path(file_path, first_page=page_idx+1, last_page=page_idx+1, poppler_path=POPPLER_PATH)
                    if images:
                        return pytesseract.image_to_string(images[0])
                except Exception as e:
                    log_event(f"OCR failed for page {page_idx+1}: {e}", level="error")
                return ""

            for i, page in enumerate(reader.pages):
                if i % 10 == 0:
                    gc.collect()

                if force_ocr_first_page and i == 0:
                    page_text = ocr_page(i)
                else:
                    page_text = page.extract_text() or ""
                    if len(page_text.strip()) < 100:
                        log_event(f"Page {i+1} insufficient text (<100 chars). Triggering OCR.", level="warning")
                        ocr_txt = ocr_page(i)
                        if len(ocr_txt.strip()) > len(page_text.strip()):
                            page_text = ocr_txt

                text += page_text + "\n"

        elif ext in [".docx", ".doc"]:
            if ext == ".docx":
                doc = Document(file_path)
                # Text from paragraphs
                for p in doc.paragraphs:
                    text += p.text + "\n"
                # Text from tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            text += cell.text + "\n"
        
        return text

    except Exception as e:
        log_event(f"Error extracting text from {file_path}: {e}", level="error")
        return ""





def apply_markdown(cell, text: str):
    """Parses basic Markdown (bold only) and adds to cell."""
    # Clear existing text
    cell.text = ""
    
    lines = text.split('\n')
    for line in lines:
        paragraph = cell.add_paragraph()
        parts = line.split('**')
        for i, part in enumerate(parts):
            run = paragraph.add_run(part)
            if i % 2 == 1: # Odd indices are between ** **
                run.font.bold = True

def extract_file_number_from_path(file_path: str) -> Optional[str]:
    """
    Extracts file number in format ####.### from the file path.
    Based on structure: .../Current Clients/<Folder1>/<Folder2>/...
    Folder1 starts with 4 digits. Folder2 starts with 3 digits.
    """
    if not file_path:
        return None
        
    try:
        # Normalize path separators
        norm_path = os.path.normpath(file_path)
        parts = norm_path.split(os.sep)
        
        # Find index of "Current Clients" (case insensitive)
        cc_index = -1
        for i, part in enumerate(parts):
            if part.lower() == "current clients":
                cc_index = i
                break
        
        if cc_index != -1 and cc_index + 2 < len(parts):
            folder1 = parts[cc_index + 1]
            folder2 = parts[cc_index + 2]
            
            # Extract numbers
            # "The first four numbers... first four numbers of the folder"
            match1 = re.match(r'^(\d{4})', folder1)
            # "The last 3 numbers... first 3 numbers of the next embedded folder"
            match2 = re.match(r'^(\d{3})', folder2)
            
            if match1 and match2:
                return f"{match1.group(1)}.{match2.group(1)}"
                
        return None
    except Exception as e:
        log_event(f"Error extracting file number from path: {e}", level="error")
        return None

def clean_caption_docx(input_path: str, output_path: str):
    """
    Cleans a caption docx by:
    1. Deleting text between the first table (caption table) and the signature block.
    2. Deleting text after the signature block (e.g., proof of service).
    """
    if not input_path.lower().endswith(".docx"):
        log_event(f"Caption is not a .docx ({os.path.basename(input_path)}). Skipping cleaning and copying.", level="warning")
        shutil.copy2(input_path, output_path)
        return

    try:
        doc = Document(input_path)
        body = doc.element.body
        
        # 1. Identify the first table
        first_table_idx = -1
        for i, child in enumerate(body):
            if child.tag.endswith('tbl'):
                first_table_idx = i
                break
        
        if first_table_idx == -1:
            log_event("No table found in caption document. Saving as is.", level="warning")
            doc.save(output_path)
            return

        # 2. Identify "Dated:" (Signature block start)
        dated_idx = -1
        for i in range(first_table_idx + 1, len(body)):
            child = body[i]
            if child.tag.endswith('p'):
                text = "".join(child.itertext()).strip().lower()
                if "dated:" in text or "dated " in text:
                    dated_idx = i
                    break
        
        if dated_idx == -1:
            log_event("No 'Dated:' found in caption document after table. Saving as is.", level="warning")
            doc.save(output_path)
            return

        # 3. Identify signature block end
        # We look for "Attorneys for" or similar as the end marker.
        sig_end_idx = dated_idx
        for i in range(dated_idx + 1, len(body)):
            child = body[i]
            if child.tag.endswith('p'):
                text = "".join(child.itertext()).strip().lower()
                if "attorneys for" in text or "attorney for" in text:
                    sig_end_idx = i
                
                # Check for Proof of Service which is definitely after
                if "proof of service" in text or "certificate of service" in text:
                    break
            elif child.tag.endswith('tbl'):
                # Stop if we hit another table
                break
        
        # 4. Perform Deletions (Reverse Order)
        # Delete everything after sig_end_idx
        for i in range(len(body) - 1, sig_end_idx, -1):
            child = body[i]
            if not child.tag.endswith('sectPr'):
                body.remove(child)

        # Delete everything between first_table_idx and dated_idx
        for i in range(dated_idx - 1, first_table_idx, -1):
            body.remove(body[i])

        doc.save(output_path)
        log_event(f"Cleaned caption document saved to {output_path}")

    except Exception as e:
        log_event(f"Error cleaning caption document: {e}", level="error")
        # Fallback to simple copy
        shutil.copy2(input_path, output_path)

def update_variables_docx(case_dir: str, data: Dict):
    """Compiles extracted variables into variables.docx in NOTES/AI OUTPUT folder."""
    notes_dir = os.path.join(case_dir, "NOTES", "AI OUTPUT")
    if not os.path.exists(notes_dir):
        os.makedirs(notes_dir)
    
    docx_path = os.path.join(notes_dir, "variables.docx")
    
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    
    doc.add_heading(f"Case Variables: {data.get('case_number', 'Unknown')}", 0)
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Variable'
    hdr_cells[1].text = 'Value'
    
    # Flatten dictionary for table
    for k, v in data.items():
        row_cells = table.add_row().cells
        row_cells[0].text = str(k)
        
        # Apply markdown formatting for specific fields
        if k in ["factual_background", "procedural_history"]:
            apply_markdown(row_cells[1], str(v))
        else:
            row_cells[1].text = str(v)
        
    doc.save(docx_path)
    log_event(f"Saved variables to {docx_path}")

def parse_file_numbers(input_str: str):
    """Parses a string containing file numbers, ranges, or commas into a list."""
    # Remove spaces and split by comma
    parts = [p.strip() for p in input_str.split(',')]
    files = []
    
    for part in parts:
        part = part.strip("'").strip('"')
        if '-' in part:
            try:
                start_str, end_str = part.split('-')
                start_str = start_str.strip()
                end_str = end_str.strip()
                
                # Assume format Prefix.Suffix (e.g., 5800.004)
                if '.' in start_str:
                    prefix_start, suffix_start = start_str.rsplit('.', 1)
                    
                    # Handle end string formats: "5800.008" or just "008"
                    if '.' in end_str:
                        prefix_end, suffix_end = end_str.rsplit('.', 1)
                        if prefix_start != prefix_end:
                            pass 
                    else:
                        suffix_end = end_str # e.g. "5800.004-008"
                    
                    s_val = int(suffix_start)
                    e_val = int(suffix_end)
                    padding = len(suffix_start)
                    
                    if s_val > e_val:
                        s_val, e_val = e_val, s_val
                        
                    for i in range(s_val, e_val + 1):
                        files.append(f"{prefix_start}.{str(i).zfill(padding)}")
                else:
                    # Integer range
                    s_val = int(start_str)
                    e_val = int(end_str)
                    if s_val > e_val:
                         s_val, e_val = e_val, s_val
                    for i in range(s_val, e_val + 1):
                        files.append(str(i))
                        
            except ValueError:
                log_event(f"Failed to parse range: {part}. Treating as literal.", level="warning")
                files.append(part)
        else:
            if part:
                files.append(part)
                
    return files

def main():
    if len(sys.argv) < 2:
        log_event("Usage: python complaint.py <file_number> [--headless]", level="error")
        sys.exit(1)

    # Check for headless flag
    is_headless = "--headless" in sys.argv
    args_to_parse = [arg for arg in sys.argv[1:] if arg != "--headless"]
    
    raw_arg = " ".join(args_to_parse)
    file_numbers = parse_file_numbers(raw_arg)
    
    # --- Dispatcher Mode ---
    if len(file_numbers) > 1:
        log_event(f"Detected multiple file numbers: {file_numbers}. Launching separate agents...")
        for fn in file_numbers:
            log_event(f"Spawning agent for {fn}...")
            try:
                # Spawn independent process
                spawn_args = [sys.executable, sys.argv[0], fn, "--headless"]
                
                if os.name == 'nt':
                    subprocess.Popen(spawn_args, creationflags=subprocess.CREATE_NO_WINDOW)
                else:
                    subprocess.Popen(spawn_args)
            except Exception as e:
                log_event(f"Failed to spawn agent for {fn}: {e}", level="error")
        
        print(f"Launched {len(file_numbers)} complaint agents in headless mode.")
        sys.exit(0)
        
    # --- Worker Mode ---
    if not file_numbers:
        log_event("No valid file numbers found.", level="error")
        sys.exit(1)
        
    file_num = file_numbers[0]
    log_event(f"--- Starting Complaint Agent for {file_num} ---")
    
    # 1. Environment & Context
    case_path = get_case_path(file_num)
    if not case_path:
        log_event("Could not locate case directory. Aborting.", level="error")
        sys.exit(1)
    
    log_event(f"Case Directory: {case_path}")
    data_manager.save_variable(file_num, "file_path", case_path, auto_tag=False)
    
    # 2. Document Discovery
    complaint_file = find_complaint_file(case_path)
    insurance_file = find_insurance_init_file(case_path)
    
    complaint_text = ""
    insurance_text = ""
    
    if complaint_file:
        log_event(f"Found Complaint: {complaint_file}")
        complaint_text = extract_text_from_file(complaint_file)
        # Save raw text to hidden variable for other agents (e.g., Docket) to use
        # Auto-tagging enabled so it picks up Pleading/Complaint/Party tags
        if complaint_text:
            data_manager.save_variable(file_num, "hidden_complaint_text", complaint_text)
    else:
        log_event("No Complaint file found.", level="error")
        # Continue? Yes, might handle insurance doc only, but 'Factual Background' will fail. 
    
    if insurance_file:
        log_event(f"Found Insurance Init: {insurance_file}")
        insurance_text = extract_text_from_file(insurance_file)
    
    # 3. AI Data Extraction
    # Initialize with empty placeholders as requested
    extracted_data = {
        "case_number": "",
        "venue_county": "",
        "plaintiff_counsel": "",
        "causes_of_action": [],
        "plaintiffs": [],
        "defendants": [],
        "filing_date": "",
        "adjuster_name": "",
        "adjuster_email": "",
        "insured_name": "",
        "claim_number": "",
        "incident_date": "",
        "client_name": "",
        "client_email": "",
        "case_name": ""
    }

    # Extract file_number from path if complaint file exists
    if complaint_file:
        derived_file_num = extract_file_number_from_path(complaint_file)
        if derived_file_num:
            extracted_data["file_number"] = derived_file_num
            log_event(f"Derived File Number from path: {derived_file_num}")
    
    # 3a. Complaint Extraction
    if complaint_text:
        prompt_complaint = """
        Extract the following details from the legal complaint as a JSON object.
        CRITICAL: "case_number" MUST be the Court Case Number (e.g., "21STCV12345", "RIC123456"). 
        DO NOT use form numbers like "PLD-PI-001", "CIV-100", or "SUM-100". If no court case number is stamped, return null.

        For "case_name": Generate a case name in the format: "[First Plaintiff Last Name] v. [First Defendant Last Name OR Business Name]".
        If there is more than one defendant, add " et. al." at the end.
        Examples: 
        - "Smith v. Jones"
        - "Doe v. Walmart Inc. et. al."

        {
            "case_number": "string (or null)",
            "venue_county": "string",
            "plaintiff_counsel": "string",
            "causes_of_action": ["list", "of", "strings"],
            "plaintiffs": ["list", "of", "strings"],
            "defendants": ["list", "of", "strings"],
            "filing_date": "YYYY-MM-DD",
            "case_name": "string"
        }
        Return ONLY the JSON.
        """
        json_str = gemini_utils.call_gemini_api(prompt_complaint, complaint_text, ["gemini-3-flash-preview", "gemini-2.5-flash"])
        if json_str:
            try:
                data = json.loads(gemini_utils.clean_json_string(json_str))
                extracted_data.update(data)
            except json.JSONDecodeError:
                log_event("Failed to parse Complaint JSON from AI.", level="error")
    
    # 3b. Insurance Extraction
    if insurance_text:
        prompt_ins = """
        Extract the following details from the insurance document as a JSON object:
        {
            "adjuster_name": "string",
            "adjuster_email": "string",
            "insured_name": "string",
            "claim_number": "string",
            "incident_date": "YYYY-MM-DD"
        }
        Return ONLY the JSON.
        """
        json_str = gemini_utils.call_gemini_api(prompt_ins, insurance_text, ["gemini-3-flash-preview", "gemini-2.5-flash"])
        if json_str:
            try:
                data = json.loads(gemini_utils.clean_json_string(json_str))
                extracted_data.update(data)
            except json.JSONDecodeError:
                log_event("Failed to parse Insurance JSON from AI.", level="error")
                
    # 3d. Case Number Verification (Cross-referencing)
    # Always attempt to verify with Caption document if available, as it is often more accurate.
    caption_file = find_caption_doc(case_path)
    if caption_file:
        log_event(f"Found Caption Document: {caption_file}")
        
        # Save as template for other agents
        try:
            ai_output_dir = os.path.join(case_path, "NOTES", "AI OUTPUT")
            if not os.path.exists(ai_output_dir):
                os.makedirs(ai_output_dir)
            
            ext = os.path.splitext(caption_file)[1]
            template_path = os.path.join(ai_output_dir, f"caption_template{ext}")
            
            # Clean the caption document before saving
            clean_caption_docx(caption_file, template_path)
            
            data_manager.save_variable(file_num, "caption", template_path, auto_tag=False)
            data_manager.save_variable(file_num, "caption_template_path", template_path, auto_tag=False)
            log_event(f"Saved cleaned caption to: {template_path}")
        except Exception as e:
            log_event(f"Failed to save caption template: {e}", level="warning")

        caption_text = extract_text_from_file(caption_file)
        
        if caption_text:
            verify_prompt = """
            Extract the Case Number, Venue County, and our Client's Name from this Caption Document.
            The name of our client(s) will typically appear on the first page after the words "Attorneys for Defendant", "Attorneys for Defendants", or "Attorneys for".
            Ignore form identifiers like "PLD-PI-001". Look for a Case Number (e.g., 23STCV..., RIC...).

            Also generate "case_name" based on the caption information:
            Format: "[First Plaintiff Last Name] v. [First Defendant Last Name OR Business Name]".
            If there is more than one defendant, add " et. al." at the end.

            Return JSON: {
                "case_number": "found_number", 
                "venue_county": "found_county", 
                "client_name": "found_client_name",
                "case_name": "generated_case_name"
            }
            """
            json_str = gemini_utils.call_gemini_api(verify_prompt, caption_text, ["gemini-2.5-flash"])
            if json_str:
                log_event(f"Caption Verification RAW JSON: {json_str}")
                try:
                    data = json.loads(gemini_utils.clean_json_string(json_str))
                    new_number = data.get("case_number")
                    new_county = data.get("venue_county")
                    new_client = data.get("client_name")
                    new_case_name = data.get("case_name")
                    
                    # Validate: Reject common form numbers if mistakenly picked up
                    if new_number and "PLD-PI-001" in new_number.upper():
                        log_event(f"Rejected invalid case number '{new_number}' (matches form ID).")
                        new_number = None

                    if new_number:
                        old_number = extracted_data.get("case_number")
                        if old_number and old_number != new_number:
                            log_event(f"Case Number Mismatch! Complaint: {old_number} vs Caption: {new_number}. Updating to Caption value.")
                        elif not old_number:
                            log_event(f"Recovered Case Number from Caption: {new_number}")
                        
                        # update with caption's number
                        extracted_data["case_number"] = new_number
                    
                    # Update venue_county if missing
                    if new_county:
                        if not extracted_data.get("venue_county"):
                            extracted_data["venue_county"] = new_county
                            log_event(f"Recovered Venue County from Caption: {new_county}")
                    
                    # Update client_name
                    if new_client:
                        extracted_data["client_name"] = new_client
                        log_event(f"Extracted Client Name from Caption: {new_client}")

                    # Update case_name (Caption is preferred source)
                    if new_case_name:
                        extracted_data["case_name"] = new_case_name
                        log_event(f"Generated Case Name from Caption: {new_case_name}")
                            
                except Exception as e:
                    log_event(f"Error parsing verification JSON: {e}", level="error")
    
    # Fallback to targeted OCR if still missing (even if caption file existed but failed)
    if not extracted_data.get("case_number") and complaint_file:
         # Fallback to targeted OCR if still missing and no caption file
        log_event("Case Number missing. Attempting targeted OCR on first page.")
        page1_text = extract_text_from_file(complaint_file, force_ocr_first_page=True)
        
        recon_prompt = """
        The case number was missing. 
        Here is text from Page 1 of the complaint (OCR'd).
        Find the Case Number. Return JSON: {"case_number": "found_number"}
        """
        json_str = gemini_utils.call_gemini_api(recon_prompt, page1_text, ["gemini-2.5-flash"]) 
        if json_str:
             try:
                data = json.loads(gemini_utils.clean_json_string(json_str))
                if data.get("case_number"):
                    extracted_data["case_number"] = data["case_number"]
                    log_event(f"Recovered Case Number: {data['case_number']}")
             except:
                 pass

    # Save extracted data points
    for k, v in extracted_data.items():
        # Tagging is enabled by default (auto_tag=True)
        # We assume these are Meta Data fields as requested
        data_manager.save_variable(file_num, k, v, extra_tags=["Meta Data"])
    
    update_variables_docx(case_path, extracted_data)

    # 4. Automated Legal Drafting
    log_event("--- Starting Automated Legal Drafting ---")
    try:
        if complaint_text:
            # Factual Background
            if os.path.exists(FACTUAL_BG_PROMPT_FILE):
                log_event(f"Reading prompt from {FACTUAL_BG_PROMPT_FILE}")
                with open(FACTUAL_BG_PROMPT_FILE, 'r', encoding='utf-8') as f:
                    fb_prompt = f.read()
                
                log_event("Calling AI for Factual Background...")
                fb_result = gemini_utils.call_gemini_api(fb_prompt, complaint_text, ["gemini-3-flash-preview", "gemini-2.5-flash"])
                
                if fb_result:
                    data_manager.save_variable(file_num, "factual_background", fb_result, extra_tags=["Meta Data"])
                    extracted_data["factual_background"] = fb_result
                    log_event("Generated Factual Background.")
                else:
                    log_event("AI returned no result for Factual Background.", level="warning")
            else:
                log_event(f"Prompt file missing: {FACTUAL_BG_PROMPT_FILE}", level="warning")
        else:
            log_event("No complaint text available for drafting.", level="warning")

    except Exception as e:
        log_event(f"Critical error in Legal Drafting section: {e}", level="error")
        import traceback
        log_event(traceback.format_exc(), level="error")

    # Final update to docx with drafted sections
    log_event("Updating variables.docx with drafted sections...")
    try:
        update_variables_docx(case_path, extracted_data)
        log_event("Successfully updated variables.docx")
    except Exception as e:
        log_event(f"Error updating docx: {e}", level="error")

    log_event("--- Complaint Agent Completed ---")

if __name__ == "__main__":
    main()
