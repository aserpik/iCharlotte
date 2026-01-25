import os
import sys
import logging
import json
import datetime
import glob
import re
import subprocess
import shutil
import time
from typing import Optional, Dict, Any, List

# Third-party imports
try:
    from docx import Document
    from docx.shared import Pt
    from pypdf import PdfReader
    from google import genai
    import pytesseract
    from pdf2image import convert_from_path
    OCR_AVAILABLE = True
    
    # Windows-specific path configuration
    if os.name == 'nt':
        tesseract_path = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
        if os.path.exists(tesseract_path):
            pytesseract.pytesseract.tesseract_cmd = tesseract_path
        poppler_path = r"C:\Program Files\poppler\Library\bin"
        POPPLER_PATH = poppler_path if os.path.exists(poppler_path) else None
    else:
        POPPLER_PATH = None
except ImportError as e:
    print(f"Critical Error: Missing dependency {e}. Please ensure python-docx, pypdf, google-genai, pytesseract, and pdf2image are installed.")
    sys.exit(1)

# --- Configuration ---
PROJECT_ROOT = os.getcwd()
LOG_FILE = os.path.join(PROJECT_ROOT, "Discovery_Request_Activity.log")
GEMINI_DATA_DIR = os.path.join(PROJECT_ROOT, ".gemini", "case_data")
BASE_PATH_WIN = r"Z:\\Shared\\Current Clients"
TEMPLATES_BASE_DIR = os.path.join(PROJECT_ROOT, "Scripts", "Discovery Templates")

# Map commonly used folder names to canonical categories (synced with audit agent)
FOLDER_MAP = {
    "DISCOVERY": ["DISCOVERY", "Discovery"],
    "PROPOUNDED": ["PROPOUNDED", "Propounded"],
    "RESPONSES": ["RESPONSES", "Responses"],
}

# Setup Logging
logging.basicConfig(
    filename=LOG_FILE,
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    datefmt='%Y-%m-%d %H:%M:%S'
)

def log_event(message, level="info"):
    timestamp = datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    if "--interactive" in sys.argv:
        print(f"[{timestamp}] {message}")
    if level == "info":
        logging.info(message)
    elif level == "error":
        logging.error(message)
    elif level == "warning":
        logging.warning(message)

def find_folder(base_path, candidates):
    """Finds the actual path for a folder given a list of candidate names (from audit.py)."""
    if not base_path or not os.path.exists(base_path):
        return None
    
    try:
        entries = os.listdir(base_path)
    except OSError:
        return None

    norm_candidates = [c.lower().strip() for c in candidates]

    for entry in entries:
        if entry.lower().strip() in norm_candidates:
             return os.path.join(base_path, entry)
    return None

def get_case_path(file_num: str) -> Optional[str]:
    """Locates the physical directory for a file number (####.###)."""
    if not re.match(r'^\d{4}\.\d{3}$', file_num):
        log_event(f"Invalid file number format: {file_num}.", level="error")
        return None

    carrier_num, case_num = file_num.split('.')
    base_path = BASE_PATH_WIN

    if not os.path.exists(base_path):
        log_event(f"Base path not found: {base_path}", level="error")
        return None

    carrier_folders = glob.glob(os.path.join(base_path, f"{carrier_num} - *"))
    if not carrier_folders:
        log_event(f"Carrier folder {carrier_num} not found.", level="error")
        return None
    carrier_path = carrier_folders[0]

    case_folders = glob.glob(os.path.join(carrier_path, f"{case_num} - *"))
    if not case_folders:
        log_event(f"Case folder {case_num} not found.", level="error")
        return None
    
    return os.path.abspath(case_folders[0])

def load_case_data(file_num: str) -> Dict[str, Any]:
    json_path = os.path.join(GEMINI_DATA_DIR, f"{file_num}.json")
    if os.path.exists(json_path):
        try:
            with open(json_path, 'r') as f:
                raw_data = json.load(f)
            
            # Flatten data structure if it uses the new metadata format
            flat_data = {}
            for k, v in raw_data.items():
                if isinstance(v, dict) and "value" in v:
                    flat_data[k] = v["value"]
                else:
                    flat_data[k] = v
            return flat_data
        except Exception as e:
            log_event(f"Error loading case data: {e}", level="error")
    return {}

def extract_text_from_pdf(file_path: str, max_pages: int = 5) -> str:
    try:
        reader = PdfReader(file_path)
        text = ""
        total_pages = len(reader.pages)
        
        # Determine which pages to read
        # If small enough, read all.
        # If large, read first 'max_pages' and last 'max_pages'
        pages_to_read = []
        if total_pages <= max_pages * 2:
            pages_to_read = range(total_pages)
        else:
            pages_to_read = list(range(max_pages)) + list(range(total_pages - max_pages, total_pages))
        
        for i in pages_to_read:
            page_text = reader.pages[i].extract_text() or ""
            
            # Check if this specific page has meaningful text (threshold: 100 chars)
            if len(page_text.strip()) < 100 and OCR_AVAILABLE:
                log_event(f"Page {i+1} of {os.path.basename(file_path)} has insufficient text. Attempting OCR...")
                try:
                    images = convert_from_path(file_path, first_page=i+1, last_page=i+1, poppler_path=POPPLER_PATH)
                    if images:
                        ocr_text = pytesseract.image_to_string(images[0])
                        if len(ocr_text.strip()) > len(page_text.strip()):
                            page_text = ocr_text
                except Exception as ocr_e:
                    log_event(f"OCR failed for page {i+1}: {ocr_e}", level="warning")
            
            text += page_text + "\n"
        return text
    except Exception as e:
        log_event(f"Error reading PDF {file_path}: {e}", level="error")
        return ""

def call_gemini_api(prompt: str, context_text: str) -> Optional[str]:
    api_key = os.environ.get("GEMINI_API_KEY")
    if not api_key:
        log_event("GEMINI_API_KEY not set.", level="error")
        return None
    
    client = genai.Client(api_key=api_key)
    
    models_to_try = ["gemini-3-pro-preview", "gemini-3-flash-preview"]
    
    for model_name in models_to_try:
        try:
            response = client.models.generate_content(
                model=model_name,
                contents=f"{prompt}\n\n[TEXT]\n{context_text}"
            )
            if response and response.text:
                return response.text
        except Exception as e:
            log_event(f"Model {model_name} failed: {e}", level="warning")
            continue
            
    return None

def get_discovery_cache_path(file_num: str) -> str:
    return os.path.join(GEMINI_DATA_DIR, f"{file_num}_discovery_cache.json")

def load_discovery_cache(file_num: str) -> Dict[str, Any]:
    path = get_discovery_cache_path(file_num)
    if os.path.exists(path):
        try:
            with open(path, 'r') as f:
                return json.load(f)
        except:
            return {}
    return {}

def save_discovery_cache(file_num: str, cache: Dict[str, Any]):
    path = get_discovery_cache_path(file_num)
    try:
        with open(path, 'w') as f:
            json.dump(cache, f, indent=2)
    except Exception as e:
        log_event(f"Failed to save cache: {e}", level="warning")

def analyze_prior_discovery(discovery_dir: str, party_name: str, client_name: str, file_num: str):
    log_event(f"Analyzing prior discovery in {discovery_dir} for {party_name}...")
    
    # Cache Setup
    cache = load_discovery_cache(file_num)
    cache_dirty = False
    
    # Exclude these subfolder names (case-insensitive) as requested
    def is_excluded(folder_name):
        fn = folder_name.lower()
        if fn in {"subpoena", "subpoenas", "transcripts", "correspondence", "pleadings"}:
            return True
        # Ignore variations of "depo notice" or "depo ntc"
        if "depo" in fn and ("notice" in fn or "ntc" in fn):
            return True
        return False
    
    pdf_files = []
    
    # Logic similar to audit agent: look for PROPOUNDED party folder first
    prop_path = find_folder(discovery_dir, FOLDER_MAP["PROPOUNDED"])
    if prop_path:
        try:
            for entry in os.listdir(prop_path):
                entry_path = os.path.join(prop_path, entry)
                if os.path.isdir(entry_path):
                    # Check if this folder belongs to the target party
                    if party_name.lower() in entry.lower() or entry.lower() in party_name.lower():
                        log_event(f"Found party-specific propounded folder: {entry}")
                        for root, dirs, files in os.walk(entry_path):
                            dirs[:] = [d for d in dirs if not is_excluded(d)]
                            for f in files:
                                if f.lower().endswith(".pdf"):
                                    pdf_files.append(os.path.join(root, f))
        except OSError:
            pass

    # If no files found in party-specific propounded folder, or to be thorough, 
    # scan entire discovery dir excluding the forbidden folders.
    if not pdf_files:
        for root, dirs, files in os.walk(discovery_dir):
            dirs[:] = [d for d in dirs if not is_excluded(d)]
            for file in files:
                if file.lower().endswith(".pdf"):
                    pdf_files.append(os.path.join(root, file))
    
    results = []
    files_to_analyze = [] # List of (full_path, filename, mtime)

    # Heuristic Patterns for Exclusion (Improvement #3)
    exclude_patterns = [
        r"subpoena", r"notice", r"proof of service", r"pos\b", 
        r"correspondence", r"letter", r"email", r"summary", r"index",
        r"verification", r"extension", r"meet and confer", r"m&c",
        r"receipt", r"invoice", r"billing"
    ]

    for pdf in pdf_files:
        filename = os.path.basename(pdf)
        try:
            file_mtime = os.path.getmtime(pdf)
        except OSError:
            continue

        # Check Cache (Improvement #4)
        cached_entry = cache.get(pdf)
        if cached_entry:
            # Check if file modified (compare timestamps)
            if cached_entry.get("mtime") == file_mtime:
                # Valid cache hit
                if cached_entry.get("is_discovery"):
                    results.append(cached_entry["data"])
                continue # Skip analysis
        
        # Heuristic Filtering (Improvement #3)
        # Skip obvious non-discovery to save tokens and time
        if any(re.search(p, filename.lower()) for p in exclude_patterns):
            # Mark as not discovery in cache so we don't check again
            cache[pdf] = {"mtime": file_mtime, "is_discovery": False}
            cache_dirty = True
            continue

        files_to_analyze.append((pdf, filename, file_mtime))
    
    # 3. Batch Analysis (Improvement #1)
    if files_to_analyze:
        batch_size = 10
        total_batches = (len(files_to_analyze) + batch_size - 1) // batch_size
        
        for i in range(0, len(files_to_analyze), batch_size):
            batch = files_to_analyze[i:i+batch_size]
            batch_filenames = [b[1] for b in batch]
            
            current_batch_num = (i // batch_size) + 1
            log_event(f"Batch analyzing {len(batch)} files ({current_batch_num}/{total_batches})...")
            
            prompt = f"""
            Analyze the following filenames. Identify which are Discovery Requests served FROM '{client_name}' TO '{party_name}'.
            Ignore any that are responses FROM {party_name} or discovery from other parties.
            
            Files:
            {json.dumps(batch_filenames, indent=1)}
            
            Return a JSON list of objects for the matches ONLY:
            [
                {{"filename": "exact_filename_from_list", "type": "SI/RPD/FROG/RFA", "set": "1", "date": "YYYY-MM-DD"}},
            ]
            """
            
            response_text = call_gemini_api(prompt, "")
            
            batch_results_map = {}
            parsed_successfully = False
            
            if response_text:
                try:
                    # Robust JSON extraction
                    match = re.search(r'\[.*\]', response_text, re.DOTALL)
                    if match:
                        json_str = match.group(0)
                        data = json.loads(json_str)
                        if isinstance(data, list):
                            for item in data:
                                if item.get("filename"):
                                    batch_results_map[item["filename"]] = item
                            parsed_successfully = True
                    else:
                        # Fallback for simple clean
                        clean_json = response_text.strip().strip("```json").strip("```")
                        data = json.loads(clean_json)
                        if isinstance(data, list):
                            for item in data:
                                if item.get("filename"):
                                    batch_results_map[item["filename"]] = item
                            parsed_successfully = True
                except Exception as e:
                    log_event(f"Error parsing batch response: {e}. Raw: {response_text[:100]}...", level="warning")

            if parsed_successfully:
                # Update results and cache ONLY if we got a valid response
                for path, fname, mtime in batch:
                    if fname in batch_results_map:
                        item = batch_results_map[fname]
                        item["served"] = True # Implicit from prompt
                        item["path"] = path # helpful
                        
                        # Add to current results
                        results.append(item)
                        
                        # Update cache
                        cache[path] = {
                            "mtime": mtime,
                            "is_discovery": True,
                            "data": item
                        }
                    else:
                        # Negative cache (file analyzed but found not relevant)
                        cache[path] = {
                            "mtime": mtime,
                            "is_discovery": False
                        }
                    cache_dirty = True
            else:
                log_event("Skipping cache update for this batch due to parsing failure.", level="warning")

    if cache_dirty:
        save_discovery_cache(file_num, cache)
    
    return results

def get_caption_title_and_body(template_path: str):
    doc = Document(template_path)
    title = ""
    # Find the first non-empty paragraph for the title
    title_idx = -1
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip():
            full_text = p.text.strip()
            # Extract words inside quotation marks (handle both standard and smart quotes)
            match = re.search(r'["“](.+?)["”]', full_text)
            if match:
                title = match.group(1)
            else:
                # Fallback to entire text if no quotes found
                title = full_text
            title_idx = i
            break
    
    body_paragraphs = []
    if title_idx != -1:
        for p in doc.paragraphs[title_idx + 1:]:
            # Stop if we hit the signature block in the template
            if "dated:" in p.text.lower():
                break
            body_paragraphs.append(p)
            
    return title, body_paragraphs

from docx.shared import Pt, Inches

def copy_para_with_formatting(source_p, target_doc, before_para=None, prefix=None, counter=None):
    """Copies a paragraph with basic formatting and runs. Reconstructs numbering if prefix/counter provided."""
    if before_para:
        new_p = before_para.insert_paragraph_before()
    else:
        new_p = target_doc.add_paragraph()
        
    # Force Double Spacing (2.0)
    new_p.paragraph_format.line_spacing = 2.0
    new_p.paragraph_format.alignment = source_p.paragraph_format.alignment
    new_p.paragraph_format.first_line_indent = source_p.paragraph_format.first_line_indent
    new_p.paragraph_format.left_indent = source_p.paragraph_format.left_indent
    new_p.paragraph_format.right_indent = source_p.paragraph_format.right_indent
    new_p.paragraph_format.space_after = source_p.paragraph_format.space_after
    new_p.paragraph_format.space_before = source_p.paragraph_format.space_before
    
    # Handle numbering reconstruction for separators (e.g. SPECIAL INTERROGATORY NO. 1)
    if not source_p.text.strip() and prefix and counter is not None:
        if hasattr(source_p._element, 'pPr') and source_p._element.pPr is not None and source_p._element.pPr.numPr is not None:
            counter[0] += 1
            run = new_p.add_run(f"{prefix} NO. {counter[0]}:")
            run.bold = True
            run.underline = True # Added underline
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            # Indent removed from the separator paragraph itself as requested
            new_p.paragraph_format.first_line_indent = 0 
            return new_p

    # Apply 0.5 inch indent to:
    # 1. Instruction paragraphs (e.g., A., B., C.)
    # 2. Discovery requests and body paragraphs that follow a separator
    # 3. Paragraphs between caption and instructions (usually start with "Pursuant to" or "TO [PARTY]")
    is_instruction = re.match(r'^[A-Z]\.\s+', source_p.text.strip())
    is_following_separator = (counter is not None and counter[0] > 0)
    
    # Check if this is a paragraph that should be indented (excluding headings like "INSTRUCTIONS...")
    should_indent = False
    text_stripped = source_p.text.strip()
    if text_stripped:
        if is_instruction or is_following_separator:
            should_indent = True
        elif "TO " in text_stripped[:10] and " ATTORNEYS" in text_stripped:
            should_indent = True
        elif "Pursuant to " in text_stripped[:20]:
            should_indent = True
        elif "The place of inspection" in text_stripped:
            should_indent = True
        elif "A written response" in text_stripped:
            should_indent = True
        elif "Failure to provide" in text_stripped:
            should_indent = True
        elif "PROPOUNDING PARTY:" in text_stripped:
            should_indent = True
        elif "RESPONDING PARTY:" in text_stripped:
            should_indent = True
        elif "SET NO.:" in text_stripped:
            should_indent = True

    if should_indent:
        new_p.paragraph_format.first_line_indent = Inches(0.5)
    else:
        new_p.paragraph_format.first_line_indent = 0

    for run in source_p.runs:
        new_run = new_p.add_run(run.text)
        new_run.bold = run.bold
        new_run.italic = run.italic
        new_run.underline = run.underline
        if run.font.name:
            new_run.font.name = run.font.name
        if run.font.size:
            new_run.font.size = run.font.size
    return new_p

def replace_placeholders(doc, replacements: Dict[str, str]):
    def do_replace(text):
        for k, v in replacements.items():
            # Check for case matching
            if k.isupper():
                text = text.replace(k, str(v).upper())
            else:
                text = text.replace(k, str(v))
        return text

    def process_paragraphs(paragraphs):
        for para in paragraphs:
            # First, check if the placeholder is in the full text but potentially split across runs
            full_text = para.text
            needs_replacement = any(k in full_text for k in replacements)
            
            if needs_replacement:
                # If there's only one run or it's simple, do direct replacement
                if len(para.runs) <= 1:
                    for run in para.runs:
                        if run.text:
                            run.text = do_replace(run.text)
                else:
                    # Complex case: placeholder might be split. 
                    # We'll do a simple but effective run-merging replacement if needed,
                    # or just replace the full text and clear other runs if we have to.
                    # For discovery, usually we can just join runs, replace, and put back in first run.
                    combined = "".join(r.text for r in para.runs)
                    new_text = do_replace(combined)
                    if combined != new_text:
                        # Clear all runs
                        for i in range(len(para.runs)):
                            para.runs[i].text = ""
                        # Set text of first run
                        para.runs[0].text = new_text

    process_paragraphs(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                process_paragraphs(cell.paragraphs)

def draft_discovery(case_path: str, party_name: str, client_name: str, disc_type_opt: str):
    log_event(f"Drafting discovery for {party_name}...")
    output_dir = os.path.join(case_path, "NOTES", "AI OUTPUT", "AI DISCOVERY")
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    # 1. Locate Caption Template
    case_data = load_case_data(extract_file_num(case_path))
    caption_template = case_data.get("caption_template_path")
    if not caption_template or not os.path.exists(caption_template):
        log_event("Caption template not found in case data. Attempting discovery.", level="info")
        # Try to find it manually
        from complaint import find_caption_doc
        caption_template = find_caption_doc(case_path)
        if not caption_template:
            log_event("No Caption Document found.", level="error")
            return

    # Determine template subdirectory
    if disc_type_opt == 'b':
        template_subdir = "Wrongful Death Discovery"
    else:
        template_subdir = "General PI Discovery"

    # Files to generate
    files_to_gen = [
        {"prefix": "SI(1)", "template": "SI(1)_tPltf.docx"},
        {"prefix": "RPD(1)", "template": "RPD(1)_tPltf.docx"}
    ]

    for item in files_to_gen:
        new_filename = f"{item['prefix']} t{party_name}.docx"
        new_path = os.path.join(output_dir, new_filename)
        shutil.copy2(caption_template, new_path)
        
        templ_path = os.path.join(TEMPLATES_BASE_DIR, template_subdir, item['template'])
        if not os.path.exists(templ_path):
            log_event(f"Template {templ_path} not found.", level="error")
            continue
            
        templ_title, templ_body = get_caption_title_and_body(templ_path)
        
        doc = Document(new_path)
        
        # Replace CAPTION PAGE / TITLE with title and remove other words in that area
        found_placeholder = False
        
        # Search paragraphs
        for para in doc.paragraphs:
            if "CAPTION PAGE" in para.text.upper() or "CAPTION TITLE" in para.text.upper():
                para.text = "" # Remove other words in that area
                run = para.add_run(templ_title)
                run.bold = True
                found_placeholder = True

        # Search tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if "CAPTION PAGE" in para.text.upper() or "CAPTION TITLE" in para.text.upper():
                            para.text = "" # Remove other words in that area
                            run = para.add_run(templ_title)
                            run.bold = True
                            found_placeholder = True

        # Determine separator prefix
        prefix_str = ""
        if "SI" in item['prefix']:
            prefix_str = "SPECIAL INTERROGATORY"
        elif "RPD" in item['prefix']:
            prefix_str = "REQUEST FOR PRODUCTION"
        
        counter = [0]

        # Find signature block start to insert body before it
        insertion_para = None
        for p in doc.paragraphs:
            if "dated:" in p.text.lower():
                insertion_para = p
                break
        
        # Insert template body text
        for p in templ_body:
            copy_para_with_formatting(p, doc, before_para=insertion_para, prefix=prefix_str, counter=counter)

        # Replacements
        replacements = {
            "[RESPONDING_PARTY]": party_name,
            "[PROPOUNDING_PARTY]": client_name
        }
        
        if disc_type_opt == 'b':
            decedent_name = case_data.get("Decedent_Name")
            if not decedent_name:
                decedent_name = "DECEDENT"
            replacements["[DECEDENT_NAME]"] = decedent_name

        replace_placeholders(doc, replacements)
        
        # Final pass: Ensure the body text is double spaced, but SKIP caption tables and signature block
        in_signature_block = False
        for p in doc.paragraphs:
            if "dated:" in p.text.lower() or in_signature_block:
                in_signature_block = True
                p.paragraph_format.line_spacing = 1.0 # Single space signature block
                continue
            
            p.paragraph_format.line_spacing = 2.0
            
        # Keep tables single spaced (usually caption table)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        p.paragraph_format.line_spacing = 1.0

        # Update Footer: Size 10, Center Aligned, Replaced Title, No "CAPTION PAGE"
        final_footer_text = templ_title
        for k, v in replacements.items():
            final_footer_text = final_footer_text.replace(k, str(v).upper() if k.isupper() else str(v))

        for section in doc.sections:
            # Handle all footer types
            footers = [section.footer, section.first_page_footer, section.even_page_footer]
            for footer in footers:
                if footer is None: continue
                # Clear all paragraphs in this footer
                for p in footer.paragraphs:
                    p.text = ""
                    # Remove any existing runs
                    for run in p.runs:
                        run.text = ""
                
                # If no paragraphs exist, add one
                if not footer.paragraphs:
                    footer.add_paragraph()
                
                # Set the text in the first paragraph
                p = footer.paragraphs[0]
                p.alignment = 1 # Center
                run = p.add_run(final_footer_text)
                run.font.size = Pt(10)
                run.font.name = 'Times New Roman'

        doc.save(new_path)
        log_event(f"Generated: {new_filename}")

def extract_file_num(path: str) -> str:
    # Reverse of get_case_path
    # Path looks like: .../#### - Carrier Name/### - Case Name
    match = re.search(r'(\d{4}) - .+[\\/](\d{3}) -', path)
    if match:
        return f"{match.group(1)}.{match.group(2)}"
    return ""

def interactive_loop(file_num: str):
    case_path = get_case_path(file_num)
    if not case_path:
        print(f"Error: Could not find case folder for {file_num}")
        input("Press Enter to exit...")
        return

    case_data = load_case_data(file_num)
    plaintiffs = case_data.get("plaintiffs", [])
    defendants = case_data.get("defendants", [])
    client_name = case_data.get("client_name", "Our Client")

    all_parties = plaintiffs + defendants
    if not all_parties:
        print("No parties found in case data.")
        input("Press Enter to exit...")
        return

    print(f"\n--- Discovery Request Agent: Case {file_num} ---")
    print("Which Party do you want to propound discovery on?")
    for i, party in enumerate(all_parties, 1):
        print(f"{i}. {party}")

    choice = input("\nEnter number(s) separated by commas (e.g. 1, 3): ")
    selected_indices = [int(x.strip()) - 1 for x in choice.split(",") if x.strip().isdigit()]
    
    selected_parties = [all_parties[i] for i in selected_indices if 0 <= i < len(all_parties)]

    discovery_dir = os.path.join(case_path, "DISCOVERY")
    
    for party in selected_parties:
        print(f"\nChecking prior discovery for: {party}")
        prior = []
        if os.path.exists(discovery_dir):
            prior = analyze_prior_discovery(discovery_dir, party, client_name, file_num)
        
        if not prior:
            print(f"No prior discovery served from {client_name} on {party}.")
        else:
            print(f"Found prior discovery for {party}:")
            for p in prior:
                print(f" - {p.get('type')} Set {p.get('set')} served on {p.get('date')}")

        print("\nWhat type of discovery do you want to generate?")
        print("a. Basic PI Discovery")
        print("b. Basic Wrongful Death Discovery")
        print("c. Next Set Discovery")
        
        dtype = input("Choice (a/b/c): ").lower().strip()
        log_event(f"User selected: '{dtype}'")
        if dtype in ['a', 'b']:
            print(f"Generating option {dtype}...")
            draft_discovery(case_path, party, client_name, dtype)
        elif dtype == 'c':
            print("Generating Next Set Discovery...")
            draft_next_set_discovery(case_path, party, client_name, prior)
        else:
            print(f"Invalid choice: '{dtype}'")

    print("\nProcessing complete.")
    input("Press Enter to exit...")

def get_last_request_number(text: str, disc_type: str) -> int:
    """Scans text for the highest request number (e.g., 'Special Interrogatory No. 35')."""
    max_num = 0
    # Patterns for different types
    patterns = []
    
    # Specific patterns
    if "SI" in disc_type or "INTERROGATORY" in disc_type.upper():
        patterns.append(r'(?:INTERROGATORY|Interrogatory)\s*(?:NO\.|No\.|No|#)\s*:?\s*(\d+)')
    elif "RFA" in disc_type or "ADMISSION" in disc_type.upper():
        patterns.append(r'(?:ADMISSION|Admission)\s*(?:NO\.|No\.|No|#)\s*:?\s*(\d+)')
    elif "RPD" in disc_type or "PRODUCTION" in disc_type.upper():
        patterns.append(r'(?:REQUEST|Request)\s*(?:NO\.|No\.|No|#)\s*:?\s*(\d+)')
    
    # Generic fallback patterns
    patterns.append(r'(?:NO\.|No\.|No|#)\s*:?\s*(\d+)')
    patterns.append(r'^\s*(\d+)\.\s+') # "1. " at start of line (risky but common in some formats)

    for pat in patterns:
        matches = re.findall(pat, text, re.MULTILINE)
        for m in matches:
            try:
                num = int(m)
                # Sanity check: Request numbers are usually < 500. 
                # If we catch a year "2024", ignore it.
                if 0 < num < 1000:
                    if num > max_num:
                        max_num = num
            except ValueError:
                pass
        
        # If we found a strong match with a specific pattern, trust it and stop trying weaker ones
        # But continue scanning matches within that pattern to find the max
        if max_num > 0 and "INTERROGATORY" in pat.upper(): 
            break 
            
    return max_num

def get_next_set_info(prior_list: List[Dict], target_type: str) -> Dict:
    """
    Determines next set number and starting request number.
    Returns: {'next_set': int, 'start_num': int, 'last_file': str}
    """
    # Normalize target type for matching
    target_clean = target_type.upper().strip()
    
    relevant = []
    for p in prior_list:
        p_type = str(p.get("type", "")).upper()
        # Loose matching
        if target_clean == "SI" and ("SI" in p_type or "INTERROGATORY" in p_type):
            relevant.append(p)
        elif target_clean == "RFA" and ("RFA" in p_type or "ADMISSION" in p_type):
            relevant.append(p)
        elif target_clean == "RPD" and ("RPD" in p_type or "PRODUCTION" in p_type):
            relevant.append(p)

    if not relevant:
        return {'next_set': 1, 'start_num': 1, 'last_file': None}

    # Find max set
    max_set = 0
    last_item = None
    for item in relevant:
        try:
            s_num = int(str(item.get("set", "0")).replace("One", "1").replace("Two", "2").strip()) # Basic word handling or just assume digits
            if s_num > max_set:
                max_set = s_num
                last_item = item
        except:
            pass
            
    if max_set == 0:
        return {'next_set': 1, 'start_num': 1, 'last_file': None}

    # Analyze last file for request number
    last_file = last_item.get("filename") # This might just be basename from analyze_prior_discovery
    # We need full path. analyze_prior_discovery results usually don't have full path in the simplified dict unless we added it?
    # Actually analyze_prior_discovery in current script saves "filename" which is os.path.basename(pdf). 
    # But it calculates it from 'pdf' which is the full path. 
    # I need to modify analyze_prior_discovery or re-find the file. 
    # For now, let's assume I can't easily get the full path unless I stored it.
    # WAIT: analyze_prior_discovery in the code I read does: data["filename"] = filename (basename).
    # I should update `analyze_prior_discovery` to store full path if possible, OR search for it.
    # Hack: Search recursively for this filename again.
    
    full_path = None
    if last_file:
         # Try to find it in the discovery dir
         # This is slow but safe
         pass # Handled in draft_next_set_discovery by passing discovery dir

    return {'next_set': max_set + 1, 'start_num': 1, 'last_file_name': last_file} # start_num 1 is placeholder

def draft_next_set_discovery(case_path: str, party_name: str, client_name: str, prior_list: List[Dict]):
    log_event(f"Drafting Next Set Discovery for {party_name}...")
    output_dir = os.path.join(case_path, "NOTES", "AI OUTPUT", "AI DISCOVERY")
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    # 1. Ask for Discovery Types
    print("\nSelect Discovery Types to generate:")
    print("1. Special Interrogatories (SI)")
    print("2. Requests for Admission (RFA)")
    print("3. Requests for Production (RPD)")
    type_choice = input("Enter number(s) (e.g., 1, 3): ")
    
    types_selected = []
    if "1" in type_choice: types_selected.append("SI")
    if "2" in type_choice: types_selected.append("RFA")
    if "3" in type_choice: types_selected.append("RPD")
    
    if not types_selected:
        print("No types selected. Aborting.")
        return

    # 2. Ask for Subject Matter
    print("\nIdentify the subject matter of the discovery requests.")
    print("You can input narrative, list, bullet points, etc. (Press Enter twice to finish):")
    lines = []
    while True:
        line = input()
        if not line:
            break
        lines.append(line)
    subject_matter = "\n".join(lines)
    
    if not subject_matter.strip():
        print("No subject matter entered. Aborting.")
        return

    # 3. Process Each Type
    case_data = load_case_data(extract_file_num(case_path))
    caption_template = case_data.get("caption_template_path")
    if not caption_template or not os.path.exists(caption_template):
        from complaint import find_caption_doc
        caption_template = find_caption_doc(case_path)
    
    if not caption_template:
        log_event("Caption Document not found.", level="error")
        return

    discovery_dir = os.path.join(case_path, "DISCOVERY")

    for dtype in types_selected:
        print(f"\nProcessing {dtype}...")
        
        # Determine Set and Start Number
        info = get_next_set_info(prior_list, dtype)
        next_set = info['next_set']
        last_file_name = info.get('last_file_name')
        
        start_num = 1
        definitions_text = ""
        preamble_text = ""
        
        # Try to find the last file to extract info
        if last_file_name:
            found_path = None
            for root, _, files in os.walk(discovery_dir):
                if last_file_name in files:
                    found_path = os.path.join(root, last_file_name)
                    break
            
            if found_path:
                print(f"Analyzing previous set: {last_file_name}")
                text = extract_text_from_pdf(found_path, max_pages=10)
                
                # Get last request number
                last_num = get_last_request_number(text, dtype)
                if last_num > 0:
                    start_num = last_num + 1
                    print(f"Found last request number: {last_num}. Starting new set at {start_num}.")
                
                # Extract Preamble and Definitions using AI
                def_prompt = f"""
                Extract the Preamble (standard introductory text) and the Definitions section from this discovery document.
                The Preamble usually starts with "PROPOUNDING PARTY:" and ends before the first request.
                The Definitions section usually follows the Preamble.
                Return ONLY the text of the Preamble and Definitions. Do not include the caption or the requests themselves.
                """
                extracted = call_gemini_api(def_prompt, text)
                if extracted:
                    definitions_text = extracted
        
        if not definitions_text:
            print("Could not extract definitions from previous set (or none exists). Using standard fallback.")
            # Fallback: Just ask AI to generate standard ones
            definitions_text = f"PROPOUNDING PARTY: {client_name}\nRESPONDING PARTY: {party_name}\nSET NO.: {next_set}\n\n[Standard California Preamble and Definitions inserted by AI]"

        # Generate Requests
        req_prompt = f"""
        Draft {dtype} (Set {next_set}) for a California legal case.
        Subject Matter:
        {subject_matter}
        
        Constraints:
        1. Start numbering at {start_num}.
        2. Do not duplicate requests from previous sets (if any).
        3. Use standard California objection-proof phrasing.
        4. Return ONLY the list of requests (e.g., "SPECIAL INTERROGATORY NO. {start_num}: ...").
        """
        
        print("Drafting requests with AI...")
        requests_text = call_gemini_api(req_prompt, f"Subject: {subject_matter}")
        
        if not requests_text:
            print("AI failed to generate requests.")
            continue

        # Create Document
        # Copy caption
        type_map = {"SI": "Special Interrogatories", "RFA": "Requests for Admission", "RPD": "Requests for Production"}
        full_type_name = type_map.get(dtype, dtype)
        
        filename = f"{full_type_name} (Set {next_set}) - {party_name}.docx"
        new_path = os.path.join(output_dir, filename)
        shutil.copy2(caption_template, new_path)
        
        doc = Document(new_path)
        
        # Update Caption Title
        # Similar logic to draft_discovery but we might need to be more aggressive if we don't have a template
        title_text = f"{full_type_name.upper()}, SET {num_to_word(next_set).upper()}"
        
        # Replace CAPTION PAGE / TITLE 
        for para in doc.paragraphs:
            if "CAPTION" in para.text.upper():
                para.text = ""
                run = para.add_run(title_text)
                run.bold = True
                
        # Insert Body
        # Find insertion point (before Date/Signature)
        insertion_para = None
        for p in doc.paragraphs:
            if "dated:" in p.text.lower():
                insertion_para = p
                break
        
        # Helper to add text block
        def add_block(text, bold_prefix=None):
            if not text: return
            lines = text.split('\n')
            for line in lines:
                if not line.strip(): continue
                if insertion_para:
                    p = insertion_para.insert_paragraph_before(line)
                else:
                    p = doc.add_paragraph(line)
                p.paragraph_format.line_spacing = 2.0
                if bold_prefix and line.startswith(bold_prefix):
                    # logic to bold specific parts if needed, but simple text is fine for now
                    pass

        # Add Definitions
        add_block(definitions_text)
        
        # Add Requests
        add_block("\n" + requests_text)
        
        doc.save(new_path)
        print(f"Generated: {filename}")

def num_to_word(n):
    # Simple helper
    words = ["Zero", "One", "Two", "Three", "Four", "Five", "Six", "Seven", "Eight", "Nine", "Ten"]
    if 0 <= n <= 10: return words[n]
    return str(n)

def main():
    if len(sys.argv) < 2:
        sys.exit(1)

    file_num = sys.argv[1]
    
    if "--interactive" in sys.argv:
        interactive_loop(file_num)
    else:
        # Spawn new terminal window
        log_event(f"Triggering interactive discovery agent for {file_num}")
        script_path = os.path.abspath(__file__)
        if os.name == 'nt':
            # Use cmd /c start to pop up a new window
            cmd = f'start "Discovery Request Agent - {file_num}" cmd /c python "{script_path}" {file_num} --interactive'
            subprocess.Popen(cmd, shell=True)
        else:
            # For non-windows, we might just run it or use a terminal emulator
            print("Interactive mode only supported on Windows for now.")

if __name__ == "__main__":
    main()
