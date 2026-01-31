import os
import sys
import logging
import datetime
import glob
import re
from typing import Optional, List, Set
from docx import Document
from docx.shared import Pt, Inches

# --- Configuration ---
BASE_PATH_WIN = r"Z:\\Shared\\Current Clients"
LOG_FILE = os.path.join(os.getcwd(), "med_record_scan_activity.log")

# Setup Logging
logging.basicConfig(
    filename=LOG_FILE,
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    datefmt='%Y-%m-%d %H:%M:%S'
)

def log_event(message, level="info"):
    timestamp = datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    try:
        print(f"[{timestamp}] {message}")
    except OSError:
        pass  # stdout pipe broken (common when running multiple agents)
    if level == "info":
        logging.info(message)
    elif level == "error":
        logging.error(message)
    elif level == "warning":
        logging.warning(message)

def get_case_path(file_num: str) -> Optional[str]:
    """Locates the physical directory for a file number (####.###)."""
    if not re.match(r'^\d{4}\.\d{3}$', file_num):
        log_event(f"Invalid file number format: {file_num}. Expected ####.###", level="error")
        return None

    carrier_num, case_num = file_num.split('.')
    base_path = BASE_PATH_WIN

    if not os.path.exists(base_path):
        log_event(f"Base path not found: {base_path}", level="error")
        return None

    # Find Carrier Folder
    carrier_folders = glob.glob(os.path.join(base_path, f"{carrier_num} - *"))
    if not carrier_folders:
        log_event(f"Carrier folder starting with {carrier_num} not found in {base_path}", level="error")
        return None
    carrier_path = carrier_folders[0]

    # Find Case Folder
    case_folders = glob.glob(os.path.join(carrier_path, f"{case_num} - *"))
    if not case_folders:
        log_event(f"Case folder starting with {case_num} not found in {carrier_path}", level="error")
        return None
    
    return os.path.abspath(case_folders[0])

def find_latest_chronology(records_dir: str) -> Optional[str]:
    """Finds the most recent .docx chronology with 'howell' in name."""
    if not os.path.exists(records_dir):
        log_event(f"RECORDS directory not found: {records_dir}", level="error")
        return None

    candidates = []
    # Date pattern: look for something that looks like a date 
    # Numeric: 12.24.25, 12-24-2025, 2025.12.24, 2025-12-24
    # Also handles potential month names if they are used, though less common in filenames
    date_regex = re.compile(r'(\d{1,4}[.\-/]\d{1,2}[.\-/]\d{1,4})|((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*[.\-\s]\d{1,2}[.\-\s]\d{2,4})', re.IGNORECASE)

    for file in os.listdir(records_dir):
        if file.lower().endswith(".docx") and not file.startswith("~$") and any(k in file.lower() for k in ["howell", "summary", "chronology"]):
            if date_regex.search(file):
                full_path = os.path.join(records_dir, file)
                candidates.append(full_path)

    if not candidates:
        log_event("No chronology documents found matching 'howell' and a date.", level="warning")
        return None

    # Sort by modification time to get the "most recent"
    candidates.sort(key=lambda x: os.path.getmtime(x), reverse=True)
    return candidates[0]

def extract_summarized_records(docx_path: str) -> List[str]:
    """Extracts filenames from 'RECORDS SUMMARIZED' table in docx."""
    summarized_files = []
    try:
        doc = Document(docx_path)
        
        target_table = None
        for table in doc.tables:
            found_header = False
            for row in table.rows[:3]:
                for cell in row.cells:
                    if "RECORDS SUMMARIZED" in cell.text.upper():
                        found_header = True
                        break
                if found_header:
                    break
            
            if found_header:
                target_table = table
        
        if target_table:
            # Try to identify the 'Filename' column index
            filename_col_idx = 0
            header_row = target_table.rows[1] if len(target_table.rows) > 1 else None
            if header_row:
                for idx, cell in enumerate(header_row.cells):
                    if "FILENAME" in cell.text.upper():
                        filename_col_idx = idx
                        break
            
            for row in target_table.rows:
                if len(row.cells) > filename_col_idx:
                    text = row.cells[filename_col_idx].text.strip()
                    if text and "RECORDS SUMMARIZED" not in text.upper() and "FILENAME" not in text.upper():
                        for line in text.split('\n'):
                            clean_line = line.strip()
                            if clean_line:
                                summarized_files.append(clean_line)
        else:
            log_event(f"Could not find 'RECORDS SUMMARIZED' table in {docx_path}", level="warning")
            
    except Exception as e:
        log_event(f"Error reading docx {docx_path}: {e}", level="error")
        
    return summarized_files

def extract_id_from_filename(filename: str) -> Optional[str]:
    """Extracts a potential identifier (e.g., 12345-001) from the start of a filename."""
    match = re.match(r'^(\d+[-_]\d+)', filename)
    if match:
        return match.group(1).replace('_', '-')
    return None

def main():
    if len(sys.argv) < 2:
        log_event("Usage: python med_record_scan.py <file_number>", level="error")
        sys.exit(1)
        
    file_num = " ".join(sys.argv[1:])
    log_event(f"--- Starting Med Record Scan for {file_num} ---")
    
    case_path = get_case_path(file_num)
    if not case_path:
        log_event("Could not locate case directory. Aborting.", level="error")
        sys.exit(1)
    
    records_dir = os.path.join(case_path, "RECORDS")
    notes_dir = os.path.join(case_path, "NOTES", "AI OUTPUT")
    
    if not os.path.exists(notes_dir):
        os.makedirs(notes_dir)
        log_event(f"Created NOTES/AI OUTPUT directory: {notes_dir}")

    # 1. Find Chronology
    chronology_file = find_latest_chronology(records_dir)
    if not chronology_file:
        log_event("Aborting: No chronology document found.", level="error")
        # We continue to subpoena logic even if chronology is missing, per requirements? 
        # Actually, let's keep going but handle chronology_file being None.
        log_event("Skipping chronology-based scan.", level="warning")
    
    summarized_list = []
    if chronology_file:
        log_event(f"Found latest chronology: {os.path.basename(chronology_file)}")
        summarized_list = extract_summarized_records(chronology_file)
    
    summarized_lower = [s.lower() for s in summarized_list]
    log_event(f"Extracted {len(summarized_list)} entries from 'RECORDS SUMMARIZED' table.")
    
    # 2. Get all PDF files in RECORDS (recursively) and Identify Subpoenas
    all_pdfs = []
    issued_subpoenas = [] # List of rel_paths
    
    for root, _, files in os.walk(records_dir):
        for file in files:
            if file.lower().endswith(".pdf"):
                full_path = os.path.join(root, file)
                rel_path = os.path.relpath(full_path, records_dir)
                size_kb = os.path.getsize(full_path) / 1024
                
                if "file copy" in file.lower() and size_kb < 2000:
                    issued_subpoenas.append(("RECORDS", rel_path))
                else:
                    all_pdfs.append(rel_path)
    
    # Check DISCOVERY\SUBPOENA(S)
    discovery_dir = os.path.join(case_path, "DISCOVERY")
    for sub in ["SUBPOENA", "SUBPOENAS"]:
        sub_dir = os.path.join(discovery_dir, sub)
        if os.path.exists(sub_dir):
            for root, _, files in os.walk(sub_dir):
                for file in files:
                    if file.lower().endswith(".pdf"):
                        full_path = os.path.join(root, file)
                        rel_path = os.path.relpath(full_path, discovery_dir)
                        size_kb = os.path.getsize(full_path) / 1024
                        if size_kb < 2000:
                            issued_subpoenas.append(("DISCOVERY", rel_path))

    log_event(f"Identified {len(issued_subpoenas)} issued subpoenas.")
    log_event(f"Found {len(all_pdfs)} responsive record PDFs in RECORDS.")
    
    # 3. Cross-reference Chronology vs Records
    not_summarized = []
    for pdf_path in all_pdfs:
        pdf_filename = os.path.basename(pdf_path)
        pdf_lower = pdf_filename.lower()
        pdf_base = os.path.splitext(pdf_lower)[0]
        
        found = False
        for s_name in summarized_lower:
            if s_name == pdf_lower or s_name == pdf_base or s_name in pdf_lower or pdf_base in s_name:
                found = True
                break
        
        if not found:
            not_summarized.append(pdf_path)
    
    # 4. Cross-reference Subpoenas vs Records
    # A subpoena matches a record if they share the same ID (e.g. 12345-001)
    subpoenas_pending = []
    matched_records = []
    
    record_ids = {}
    for pdf_path in all_pdfs:
        rid = extract_id_from_filename(os.path.basename(pdf_path))
        if rid:
            if rid not in record_ids: record_ids[rid] = []
            record_ids[rid].append(pdf_path)

    for loc, sub_path in issued_subpoenas:
        sub_filename = os.path.basename(sub_path)
        sid = extract_id_from_filename(sub_filename)
        
        match_found = False
        if sid and sid in record_ids:
            match_found = True
            for r in record_ids[sid]:
                if r not in matched_records:
                    matched_records.append(r)
        
        if not match_found:
            subpoenas_pending.append((loc, sub_path))

    log_event(f"Identified {len(subpoenas_pending)} subpoenas with no responsive records.")
    
    # 5. Save Output
    output_path = os.path.join(notes_dir, "Records_Not_Summarized.docx")
    try:
        from docx.shared import RGBColor
        doc = Document()
        
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(12)
        style.font.color.rgb = RGBColor(0, 0, 0)
        
        def format_heading(heading):
            for run in heading.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)

        h0 = doc.add_heading('Records Scan Report', 0)
        format_heading(h0)
        
        doc.add_paragraph(f"Case File: {file_num}")
        doc.add_paragraph(f"Chronology Used: {os.path.basename(chronology_file) if chronology_file else 'None'}")
        doc.add_paragraph(f"Date of Scan: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        
        # 1. Subpoenas Issued
        h1 = doc.add_heading('Subpoenas Issued', 1)
        format_heading(h1)
        if issued_subpoenas:
            from docx.shared import Inches
            for loc, item in sorted(issued_subpoenas):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.5)
                p.paragraph_format.first_line_indent = Inches(-0.25)
                run = p.add_run(f"•\t[{loc}] {item}")
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
        else:
            doc.add_paragraph("No issued subpoenas identified.")

        # 2. Records Received
        h2 = doc.add_heading('Records Received', 1)
        format_heading(h2)
        if all_pdfs:
            for item in sorted(all_pdfs):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.5)
                p.paragraph_format.first_line_indent = Inches(-0.25)
                run = p.add_run(f"•\t{item}")
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
        else:
            doc.add_paragraph("No responsive record PDFs found.")

        # 3. Records identified in chronology table
        h3 = doc.add_heading('Records identified in chronology table', 1)
        format_heading(h3)
        if summarized_list:
            for item in sorted(summarized_list):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.5)
                p.paragraph_format.first_line_indent = Inches(-0.25)
                run = p.add_run(f"•\t{item}")
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
        else:
            doc.add_paragraph("No entries found in the summarized table.")

        # 4. Records not Included in chronology
        h4 = doc.add_heading('Records not Included in chronology', 1)
        format_heading(h4)
        if not_summarized:
            doc.add_paragraph("The following PDF files were found in the RECORDS folder but not identified in the chronology's summarized table:")
            for item in sorted(not_summarized):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.5)
                p.paragraph_format.first_line_indent = Inches(-0.25)
                run = p.add_run(f"•\t{item}")
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
        else:
            doc.add_paragraph("All identified PDF files appear to be summarized.")

        # 5. Outstanding Subpoenas (Records not Received)
        h5 = doc.add_heading('Outstanding Subpoenas (Records not Received)', 1)
        format_heading(h5)
        if subpoenas_pending:
            for loc, item in sorted(subpoenas_pending):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.5)
                p.paragraph_format.first_line_indent = Inches(-0.25)
                run = p.add_run(f"•\t[{loc}] {item}")
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
        else:
            doc.add_paragraph("All issued subpoenas have corresponding records in the RECORDS folder.")
            
        doc.save(output_path)
        log_event(f"Saved results to {output_path}")
    except Exception as e:
        log_event(f"Error saving output docx: {e}", level="error")

if __name__ == "__main__":
    main()
