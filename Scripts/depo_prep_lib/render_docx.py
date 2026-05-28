"""Stage D - render outline.docx via python-docx."""
from __future__ import annotations

from pathlib import Path
from typing import Union

from docx import Document
from docx.shared import Inches, Pt


def _para(doc, text, *, bold=False, italic=False, size=12, indent_left=0.0,
          first_line=0.0, space_after_pt=6):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(indent_left)
    if first_line:
        p.paragraph_format.first_line_indent = Inches(first_line)
    p.paragraph_format.space_after = Pt(space_after_pt)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.name = "Times New Roman"
    r.font.size = Pt(size)
    return p


def render_outline_docx(*, outline: dict, output_path: Union[str, Path]) -> None:
    """Write outline.docx at output_path. Overwrites any existing file."""
    output_path = Path(output_path)
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    deponent_name = outline.get("deponent_name") or "Unknown Deponent"
    deponent_role = outline.get("deponent_role") or ""

    # Title block
    _para(doc, f"Depo Prep Outline - {deponent_name}", bold=True, size=16, space_after_pt=4)
    if deponent_role:
        _para(doc, deponent_role, italic=True, size=11, space_after_pt=12)

    for topic in outline.get("topics", []):
        title = topic.get("title", "(Untitled topic)")
        strat = topic.get("strategic_note", "")
        _para(doc, title.upper(), bold=True, size=13, space_after_pt=4)
        if strat:
            _para(doc, f"Strategic: {strat}", italic=True, size=11,
                  indent_left=0.25, space_after_pt=6)

        for q in topic.get("questions", []):
            _para(doc, f"{q['n']}.  {q.get('text', '')}", size=12,
                  indent_left=0.5, first_line=-0.25, space_after_pt=4)

            if q.get("purpose"):
                _para(doc, f"Purpose: {q['purpose']}", italic=True, size=10,
                      indent_left=0.75, space_after_pt=2)
            if q.get("source_facts"):
                _para(doc, "Source facts:", italic=True, size=10,
                      indent_left=0.75, space_after_pt=2)
                for f in q["source_facts"]:
                    _para(doc, f"• {f}", size=10, indent_left=1.0, space_after_pt=2)
            if q.get("impeachment_hook"):
                _para(doc, f"Impeachment: {q['impeachment_hook']}", italic=True,
                      size=10, indent_left=0.75, space_after_pt=2)
            if q.get("objection_alts"):
                _para(doc, "Objection alts:", italic=True, size=10,
                      indent_left=0.75, space_after_pt=2)
                for a in q["objection_alts"]:
                    _para(doc, f"• {a}", size=10, indent_left=1.0, space_after_pt=2)

    gaps = outline.get("coverage_gaps") or []
    if gaps:
        _para(doc, "Coverage notes from the AI", bold=True, size=12, space_after_pt=4)
        for g in gaps:
            _para(doc, f"• {g}", size=11, indent_left=0.25, space_after_pt=2)

    doc.save(str(output_path))
