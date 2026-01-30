"""
Summarize Agent for iCharlotte

Summarizes documents (PDF, DOCX, TXT) using LLM with cross-check verification.

Features:
- Dynamic OCR threshold for text extraction
- Two-pass processing: Summary + Cross-Check
- Multi-provider LLM fallback (Gemini, Claude, OpenAI)
- Memory monitoring for large documents
- Structured progress reporting for UI integration
"""

import os
import sys
import re
import subprocess
import datetime
from docx import Document
from docx.shared import Pt, Inches

# Add parent directory to path for imports
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

# Import shared infrastructure
from icharlotte_core.document_processor import DocumentProcessor, OCRConfig
from icharlotte_core.agent_logger import AgentLogger, create_legacy_log_event
from icharlotte_core.llm_config import LLMCaller, LLMConfig
from icharlotte_core.memory_monitor import MemoryMonitor, track_memory
from icharlotte_core.exceptions import (
    ExtractionError, LLMError, PassFailedError,
    SummaryPassError, CrossCheckPassError, MemoryLimitError
)
from icharlotte_core.docx_writer import safe_append_to_docx, DocxWriteError

# Import Case Data Manager
try:
    from case_data_manager import CaseDataManager
except ImportError:
    sys.path.append(os.path.join(os.getcwd(), 'Scripts'))
    from case_data_manager import CaseDataManager

# Import Document Registry for classification
try:
    from document_registry import classify_and_register
except ImportError:
    from Scripts.document_registry import classify_and_register


# =============================================================================
# Configuration
# =============================================================================

SCRIPTS_DIR = os.path.dirname(os.path.abspath(__file__))
PROMPT_FILE = os.path.join(SCRIPTS_DIR, "SUMMARIZE_PROMPT.txt")
CROSS_CHECK_PROMPT_FILE = os.path.join(SCRIPTS_DIR, "SUMMARIZE_CROSS_CHECK_PROMPT.txt")

# Legacy log file for backward compatibility
LEGACY_LOG_FILE = r"C:\GeminiTerminal\Summarize_activity.log"


# =============================================================================
# File Number Extraction
# =============================================================================

def extract_file_number(path: str) -> str:
    """Extract file number from path.

    Handles paths like:
    - Z:\\Shared\\Current Clients\\3800- NATIONWIDE\\3850\\084 - Dudash\\...
    - Paths containing literal "3850.084"
    """
    # Standard pattern: 1234.567 (with literal dot)
    match = re.search(r"(\d{4}\.\d{3})", path)
    if match:
        return match.group(1)

    # Parse from directory structure
    # Look for pattern: 4-digit folder followed by 3-digit folder
    parts = os.path.normpath(path).split(os.sep)
    try:
        # Find consecutive folders matching client/matter pattern
        for i in range(len(parts) - 1):
            # Look for 4-digit client folder (e.g., "3850" or "3850 - Name")
            client_match = re.match(r"^(\d{4})(?:\D|$)", parts[i])
            if client_match:
                # Check if next folder is 3-digit matter (e.g., "084" or "084 - Name")
                matter_match = re.match(r"^(\d{3})(?:\D|$)", parts[i + 1])
                if matter_match:
                    return f"{client_match.group(1)}.{matter_match.group(1)}"

        # Fallback: Look for "Current Clients" structure
        # Pattern: Current Clients\GroupFolder\ClientFolder\MatterFolder
        cc_index = -1
        for i, part in enumerate(parts):
            if part.lower() == "current clients":
                cc_index = i
                break

        if cc_index != -1 and cc_index + 3 < len(parts):
            # Skip the group folder (e.g., "3800- NATIONWIDE"), use folders at +2 and +3
            client_folder = parts[cc_index + 2]  # e.g., "3850"
            matter_folder = parts[cc_index + 3]  # e.g., "084 - Dudash"

            client_code = re.match(r"^(\d{4})(?:\D|$)", client_folder)
            matter_code = re.match(r"^(\d{3})(?:\D|$)", matter_folder)

            if client_code and matter_code:
                return f"{client_code.group(1)}.{matter_code.group(1)}"

    except Exception:
        pass

    return None


# =============================================================================
# Document Output Functions
# =============================================================================

def add_markdown_to_doc(doc, content):
    """Parses basic Markdown and adds it to the docx Document."""
    lines = content.split('\n')
    active_paragraph = None

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        # Headings: Convert to bold text at start of paragraph, ending with period
        if stripped.startswith('#'):
            text = stripped.lstrip('#').strip()
            if not text.endswith('.'):
                text += "."
            active_paragraph = doc.add_paragraph()
            run = active_paragraph.add_run(text + " ")
            run.bold = True
            continue

        # List items
        if stripped.startswith('* ') or stripped.startswith('- '):
            text = stripped[2:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.first_line_indent = Inches(-0.25)

            run = p.add_run("\t")
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

            # Support bold parsing within list items
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    r = p.add_run(part[2:-2])
                    r.bold = True
                else:
                    r = p.add_run(part)
                r.font.name = 'Times New Roman'
                r.font.size = Pt(12)

            active_paragraph = None
            continue

        # Normal text (with bold support)
        p = active_paragraph if active_paragraph else doc.add_paragraph()

        parts = re.split(r'(\*\*.*?\*\*)', stripped)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
            else:
                p.add_run(part)

        active_paragraph = None


def save_to_docx(content: str, output_path: str, title_text: str, logger: AgentLogger) -> bool:
    """Saves the content to a DOCX file. Appends if exists. Handles locking."""
    base_name, ext = os.path.splitext(output_path)
    counter = 1
    current_output_path = output_path

    while True:
        try:
            if os.path.exists(current_output_path):
                try:
                    doc = Document(current_output_path)
                    doc.add_page_break()
                    logger.info(f"Appending to existing file: {current_output_path}")
                except Exception:
                    raise PermissionError("Cannot open existing file.")
            else:
                doc = Document()
                logger.info(f"Creating new file: {current_output_path}")

            # Apply Styles
            style = doc.styles['Normal']
            style.font.name = 'Times New Roman'
            style.font.size = Pt(12)
            style.paragraph_format.line_spacing = 1.0

            # Update Heading styles
            for i in range(1, 10):
                if f'Heading {i}' in doc.styles:
                    h_style = doc.styles[f'Heading {i}']
                    h_style.font.name = 'Times New Roman'
                    h_style.font.size = Pt(12)
                    h_style.paragraph_format.line_spacing = 1.0

            # Title
            p = doc.add_paragraph()
            run = p.add_run(title_text)
            run.bold = True
            run.underline = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

            doc.add_paragraph(f"Generated on: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

            add_markdown_to_doc(doc, content)

            doc.save(current_output_path)
            logger.output_file(current_output_path)
            return True

        except (PermissionError, IOError) as e:
            logger.warning(f"File locked: {current_output_path}. Trying next version.")
            counter += 1
            current_output_path = f"{base_name} v.{counter}{ext}"

            if counter > 10:
                logger.error("Failed to save after 10 attempts.")
                return False

        except Exception as e:
            logger.error(f"Error saving to DOCX: {e}")
            return False


def get_output_directory(input_path: str) -> str:
    """Determine the output directory based on input path."""
    parts = input_path.split(os.sep)
    output_dir = None
    case_root_parts = None

    # Priority 1: Find folder starting with exactly 3 digits (Case Folder)
    for i in range(len(parts) - 1, -1, -1):
        if re.match(r'^\d{3}(\D|$)', parts[i]):
            case_root_parts = parts[:i+1]
            break

    # Priority 2: Standard "Current Clients" structure
    if not case_root_parts:
        for i, part in enumerate(parts):
            if part.lower() == "current clients":
                if i + 2 < len(parts):
                    case_root_parts = parts[:i+3]
                break

    if case_root_parts:
        output_dir = os.sep.join(case_root_parts + ["NOTES", "AI OUTPUT"])

    if not output_dir:
        # Fallback 1: NOTES already in path
        for i in range(len(parts) - 1, -1, -1):
            if parts[i].upper() == "NOTES":
                output_dir = os.path.join(os.sep.join(parts[:i+1]), "AI OUTPUT")
                break

    if not output_dir:
        # Fallback 2: Sibling NOTES folder
        input_dir = os.path.dirname(input_path)
        parent_dir = os.path.dirname(input_dir)
        output_dir = os.path.join(parent_dir, "NOTES", "AI OUTPUT")

    return output_dir


# =============================================================================
# Main Processing Functions
# =============================================================================

def process_document(input_path: str, logger: AgentLogger) -> bool:
    """
    Process a single document through the summarization pipeline.

    Pipeline:
    1. Extract text (with dynamic OCR)
    2. Pass 1: Initial summary
    3. Pass 2: Cross-check verification
    4. Save to DOCX

    Args:
        input_path: Path to the document.
        logger: AgentLogger instance.

    Returns:
        True if successful.
    """
    # Initialize components
    memory_monitor = MemoryMonitor(warn_threshold_mb=1500, abort_threshold_mb=2000, logger=logger.info)
    llm_caller = LLMCaller(logger=logger)

    # Progress tracking
    logger.progress(2, "Initializing summarization pipeline...")

    # Load prompts
    try:
        with open(PROMPT_FILE, "r", encoding="utf-8") as f:
            summary_prompt = f.read()
    except Exception as e:
        logger.error(f"Error reading prompt file: {e}")
        return False

    cross_check_prompt = None
    if os.path.exists(CROSS_CHECK_PROMPT_FILE):
        try:
            with open(CROSS_CHECK_PROMPT_FILE, "r", encoding="utf-8") as f:
                cross_check_prompt = f.read()
        except Exception:
            logger.warning("Could not load cross-check prompt. Skipping cross-check pass.")

    total_passes = 3 if cross_check_prompt else 2
    logger.progress(5, "Prompts loaded, starting text extraction...")

    # ==========================================================================
    # Pass 1: Text Extraction (5% - 25%)
    # ==========================================================================
    logger.pass_start("Extraction", 1, total_passes)
    logger.progress(7, "Reading document...")

    try:
        with memory_monitor.track_operation("Text Extraction"):
            processor = DocumentProcessor(
                ocr_config=OCRConfig(adaptive=True),
                logger=logger
            )
            logger.progress(10, "Running text extraction (OCR if needed)...")
            result = processor.extract_with_dynamic_ocr(input_path)

            if not result.success:
                raise ExtractionError(f"Failed to extract text: {result.error}", file_path=input_path)

            text = result.text
            logger.progress(22, f"Extracted {result.char_count} chars from {result.page_count} pages")
            logger.info(f"Extracted {result.char_count} chars from {result.page_count} pages")
            logger.info(f"OCR used on {len(result.ocr_pages)} pages ({result.ocr_percentage:.1f}%)")

    except MemoryLimitError as e:
        logger.pass_failed("Extraction", str(e), recoverable=False)
        return False
    except Exception as e:
        logger.pass_failed("Extraction", str(e), recoverable=False)
        return False

    logger.progress(25, "Text extraction complete")
    logger.pass_complete("Extraction", success=True)

    # ==========================================================================
    # Pass 2: Summary Generation (25% - 65%)
    # ==========================================================================
    logger.pass_start("Summary", 2, total_passes)
    logger.progress(28, "Starting summary generation...")

    try:
        with memory_monitor.track_operation("Summary Generation"):
            logger.progress(35, "Sending document to LLM for summarization...")
            summary = llm_caller.call(summary_prompt, text, task_type="summary")

            if not summary:
                raise SummaryPassError("LLM returned empty response")

            logger.progress(62, f"Summary generated: {len(summary)} chars")
            logger.info(f"Generated summary: {len(summary)} chars")

    except Exception as e:
        logger.pass_failed("Summary", str(e), recoverable=True)
        return False

    logger.progress(65, "Summary generation complete")
    logger.pass_complete("Summary", success=True)

    # ==========================================================================
    # Pass 3: Cross-Check Verification (Optional) (65% - 88%)
    # ==========================================================================
    if cross_check_prompt:
        logger.pass_start("CrossCheck", 3, total_passes)
        logger.progress(68, "Starting cross-check verification...")

        try:
            with memory_monitor.track_operation("Cross-Check"):
                # Format the cross-check prompt with summary and original
                formatted_prompt = cross_check_prompt.replace("{summary}", summary).replace("{original}", text[:50000])

                logger.progress(72, "Sending to LLM for cross-check verification...")
                verified_summary = llm_caller.call(formatted_prompt, "", task_type="cross_check")

                if verified_summary:
                    # Check if cross-check made meaningful changes
                    if len(verified_summary) > len(summary) * 0.9:
                        summary = verified_summary
                        logger.progress(85, "Cross-check completed with improvements")
                        logger.info("Cross-check completed with improvements")
                    else:
                        logger.progress(85, "Cross-check returned shorter result, using original")
                        logger.warning("Cross-check returned shorter result. Using original summary.")
                else:
                    logger.progress(85, "Cross-check returned empty, using original")
                    logger.warning("Cross-check returned empty. Using original summary.")

        except Exception as e:
            logger.pass_failed("CrossCheck", str(e), recoverable=True)
            logger.warning("Continuing with unverified summary")

        logger.progress(88, "Cross-check verification complete")
        logger.pass_complete("CrossCheck", success=True)
    else:
        logger.progress(88, "Cross-check skipped (no prompt configured)")

    # ==========================================================================
    # Save Output (with process-safe locking) (88% - 95%)
    # ==========================================================================
    logger.progress(89, "Preparing to save output...")
    output_dir = get_output_directory(input_path)

    if not os.path.exists(output_dir):
        try:
            os.makedirs(output_dir)
            logger.info(f"Created directory: {output_dir}")
        except Exception as e:
            logger.error(f"Error creating directory: {e}")
            return False

    output_file = os.path.join(output_dir, "AI_OUTPUT.docx")
    title = os.path.basename(input_path)

    logger.progress(91, f"Saving to {os.path.basename(output_file)}...")
    try:
        saved_path = safe_append_to_docx(
            output_path=output_file,
            content=summary,
            title=title,
            logger=logger
        )
        logger.output_file(saved_path)
        logger.progress(95, "Document saved successfully")
    except DocxWriteError as e:
        logger.error(f"Failed to save output: {e}")
        return False

    # ==========================================================================
    # Save to Case Data (95% - 98%)
    # ==========================================================================
    logger.progress(96, "Saving to case database...")
    file_num = extract_file_number(input_path)

    try:
        if file_num:
            data_manager = CaseDataManager()
            base_name = os.path.splitext(os.path.basename(input_path))[0]
            clean_var_name = re.sub(r"[^a-zA-Z0-9_]", "_", base_name.lower())
            var_key = f"summary_{clean_var_name}"

            logger.info(f"Saving to case data: {var_key} for {file_num}")
            data_manager.save_variable(file_num, var_key, summary, source="summarize_agent")
    except Exception as e:
        logger.warning(f"Could not save to case data: {e}")

    # ==========================================================================
    # Classify and Register Document (98% - 100%)
    # ==========================================================================
    logger.progress(98, "Registering document...")
    try:
        if file_num:
            base_name = os.path.splitext(os.path.basename(input_path))[0]
            classify_and_register(
                file_number=file_num,
                document_name=base_name,
                summary=summary,
                source_path=input_path,
                summary_location=output_file,
                agent="summarize",
                logger=logger
            )
            logger.info(f"Registered document in case registry")
    except Exception as e:
        logger.warning(f"Could not register document: {e}")

    logger.progress(100, "Summarization complete")
    return True


def process_directory(dir_path: str, logger: AgentLogger):
    """Process all documents in a directory."""
    files_to_process = []

    for root, _, files in os.walk(dir_path):
        for file in files:
            if file.lower().endswith(('.pdf', '.docx')):
                if "AI_OUTPUT" in file:
                    continue
                files_to_process.append(os.path.join(root, file))

    if not files_to_process:
        logger.info("No suitable files found in directory.")
        return

    logger.info(f"Found {len(files_to_process)} files to process.")

    for file_path in files_to_process:
        logger.info(f"Processing: {file_path}")
        try:
            subprocess.run([sys.executable, sys.argv[0], file_path], check=True)
        except subprocess.CalledProcessError as e:
            logger.error(f"Failed to process {file_path}: {e}")


def main():
    """Main entry point."""
    if len(sys.argv) < 2:
        print("Error: No file path provided.", flush=True)
        sys.exit(1)

    # Parse arguments
    raw_args = sys.argv[1:]

    # Handle quoted paths with spaces
    combined_arg = " ".join(raw_args)
    clean_combined = combined_arg.strip().strip('"').strip("'")

    if os.path.exists(clean_combined):
        file_paths = [clean_combined]
    else:
        file_paths = raw_args

    # Dispatcher mode: multiple files
    if len(file_paths) > 1:
        print(f"Detected {len(file_paths)} files. Launching separate agents...", flush=True)
        for path in file_paths:
            try:
                if os.name == 'nt':
                    subprocess.Popen([sys.executable, sys.argv[0], path], creationflags=0x08000000)
                else:
                    subprocess.Popen([sys.executable, sys.argv[0], path])
            except Exception as e:
                print(f"Failed to spawn agent for {path}: {e}", flush=True)
        sys.exit(0)

    # Single file/directory mode
    input_path = file_paths[0].strip().strip('"').strip("'")
    input_path = os.path.abspath(input_path)

    if not os.path.exists(input_path):
        print(f"Error: File not found: {input_path}", flush=True)
        sys.exit(1)

    # Extract file number for logger context
    file_number = extract_file_number(input_path)

    # Initialize logger
    logger = AgentLogger("Summarize", file_number=file_number)

    # Also create legacy log_event for backward compatibility
    log_event = create_legacy_log_event("Summarize", LEGACY_LOG_FILE)

    logger.info(f"Starting agent for: {input_path}")

    # Directory handling
    if os.path.isdir(input_path):
        logger.info(f"Input is a directory. Scanning for files...")
        process_directory(input_path, logger)
        logger.info("Directory processing complete.")
        sys.exit(0)

    # Single file processing
    success = process_document(input_path, logger)

    if success:
        logger.info("Agent finished successfully.")
        sys.exit(0)
    else:
        logger.error("Agent finished with errors.")
        sys.exit(1)


if __name__ == '__main__':
    main()
