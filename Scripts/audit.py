import os
import sys
import json
import datetime
import glob
import re
import logging
import subprocess
from pathlib import Path
from typing import Optional, Dict, Any

# Import Case Data Manager
try:
    from case_data_manager import CaseDataManager
except ImportError:
    sys.path.append(os.path.join(os.getcwd(), 'Scripts'))
    from case_data_manager import CaseDataManager

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Warning: python-docx not installed. Document generation will be skipped.")
    Document = None

# --- Configuration & Constants ---
# CASE_DATA_DIR is managed by CaseDataManager
LOG_FILE = os.path.join(os.getcwd(), "Audit_activity.log")
BASE_PATH_WIN = r"Z:\\Shared\\Current Clients"

# Initialize Data Manager
data_manager = CaseDataManager()

# Map commonly used folder names to canonical categories
FOLDER_MAP = {
    "DISCOVERY": ["DISCOVERY", "Discovery"],
    "PROPOUNDED": ["PROPOUNDED", "Propounded"],
    "RESPONSES": ["RESPONSES", "Responses"],
    "EXPERTS": ["EXPERTS", "Experts"],
    "OUR_EXPERTS": ["OUR Experts", "Our Experts", "Retained Experts"],
    "DESIGNATIONS": ["Designations", "Expert Designation", "Designation"],
    "RECORDS": ["RECORDS", "Records", "Medical Records"],
    "SUBPOENAED": ["SUBPOENAED RECORDS", "Subpoenaed", "SUBPOENAS"],
    "PLEADINGS": ["PLEADING", "Pleadings", "Pleadings & Motions"],
}

# Setup Logging
logging.basicConfig(
    filename=LOG_FILE,
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    datefmt='%Y-%m-%d %H:%M:%S'
)

def log_event(message, level="info"):
    timestamp = datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    print(f"[{timestamp}] {message}")
    sys.stdout.flush()
    if level == "info":
        logging.info(message)
    elif level == "error":
        logging.error(message)
    elif level == "warning":
        logging.warning(message)

def get_case_path(file_num: str) -> Optional[str]:
    """Locates the physical directory for a file number (####.###)."""
    if not re.match(r'^\d{4}\.\d{3}$', file_num):
        return None

    carrier_num, case_num = file_num.split('.')
    base_path = BASE_PATH_WIN

    if not os.path.exists(base_path):
        log_event(f"Base path not found: {base_path}", level="error")
        return None

    # Find Carrier Folder
    carrier_folders = glob.glob(os.path.join(base_path, f"{carrier_num} - *"))
    if not carrier_folders:
        log_event(f"Carrier folder starting with {carrier_num} not found in {base_path}", level="error")
        return None
    carrier_path = carrier_folders[0]

    # Find Case Folder
    case_folders = glob.glob(os.path.join(carrier_path, f"{case_num} - *"))
    if not case_folders:
        log_event(f"Case folder starting with {case_num} not found in {carrier_path}", level="error")
        return None
    
    return os.path.abspath(case_folders[0])

def parse_file_numbers(input_str: str):
    """Parses a string containing file numbers, ranges, or commas into a list."""
    parts = [p.strip() for p in input_str.split(',')]
    files = []
    
    for part in parts:
        part = part.strip("'").strip('"')
        if '-' in part:
            try:
                start_str, end_str = part.split('-')
                start_str = start_str.strip()
                end_str = end_str.strip()
                
                if '.' in start_str:
                    prefix_start, suffix_start = start_str.rsplit('.', 1)
                    if '.' in end_str:
                        prefix_end, suffix_end = end_str.rsplit('.', 1)
                    else:
                        suffix_end = end_str
                    
                    s_val = int(suffix_start)
                    e_val = int(suffix_end)
                    padding = len(suffix_start)
                    
                    if s_val > e_val:
                        s_val, e_val = e_val, s_val
                        
                    for i in range(s_val, e_val + 1):
                        files.append(f"{prefix_start}.{str(i).zfill(padding)}")
                else:
                    s_val = int(start_str)
                    e_val = int(end_str)
                    if s_val > e_val:
                         s_val, e_val = e_val, s_val
                    for i in range(s_val, e_val + 1):
                        files.append(str(i))
            except ValueError:
                log_event(f"Failed to parse range: {part}. Treating as literal.", level="warning")
                files.append(part)
        else:
            if part:
                files.append(part)
    return files

def load_case_data(case_number):
    """Loads the JSON data for a specific case. Locates folder if JSON is missing."""
    # Normalize case number format if needed
    if not "." in case_number and len(case_number) == 3:
        case_number = f"5800.{case_number}"
    
    # Try to get data from manager
    # We get flattened variables for easier access in this script
    data = data_manager.get_all_variables(case_number, flatten=True)
    
    if data and "file_path" in data:
        return data
    else:
        log_event(f"JSON data or file_path for {case_number} not found. Attempting to locate directory...")
        case_path = get_case_path(case_number)
        if case_path:
            log_event(f"Located case directory: {case_path}")
            data_manager.save_variable(case_number, "file_path", case_path, extra_tags=["Meta Data"])
            return {"file_path": case_path}
        else:
            log_event(f"Error: Could not find data file or directory for case {case_number}", level="error")
            return None

def find_folder(base_path, candidates):
    """Finds the actual path for a folder given a list of candidate names."""
    if not base_path or not os.path.exists(base_path):
        return None
    
    try:
        entries = os.listdir(base_path)
    except OSError:
        return None

    norm_candidates = [c.lower().strip() for c in candidates]

    for entry in entries:
        if entry.lower().strip() in norm_candidates:
             return os.path.join(base_path, entry)
    return None

def scan_directory(path, keywords):
    """Scans a directory for files matching keywords (case-insensitive)."""
    matches = []
    if not path or not os.path.exists(path):
        return matches

    try:
        for root, _, files in os.walk(path):
            for file in files:
                if any(k.lower() in file.lower() for k in keywords):
                    matches.append(os.path.join(root, file))
    except OSError:
        pass
    return matches

def count_folders(path):
    """Counts subfolders in a directory."""
    if not path or not os.path.exists(path):
        return 0
    try:
        return len([name for name in os.listdir(path) if os.path.isdir(os.path.join(path, name))])
    except OSError:
        return 0

def check_deadlines(trial_date_str):
    """Calculates critical deadlines based on trial date."""
    if not trial_date_str or trial_date_str.lower() == "none":
        return None

    try:
        # Standardize trial date format or handle multiple formats if necessary
        # Assuming YYYY-MM-DD from complaint agent
        trial_date = datetime.datetime.strptime(trial_date_str, "%Y-%m-%d").date()
    except ValueError:
        try:
             # Try common US format
             trial_date = datetime.datetime.strptime(trial_date_str, "%m/%d/%Y").date()
        except ValueError:
             return None

    deadlines = {
        "Trial Date": trial_date,
        "Discovery Cutoff (T-30)": trial_date - datetime.timedelta(days=30),
        "Expert Exchange (T-50)": trial_date - datetime.timedelta(days=50),
        "MSJ Filing (T-105)": trial_date - datetime.timedelta(days=105),
    }
    return deadlines

def save_to_docx(case_path, report_lines):
    """Saves the audit report to a Word document in NOTES/AI OUTPUT with specific formatting."""
    if not Document:
        return

    output_dir = os.path.join(case_path, "NOTES", "AI OUTPUT")
    if not os.path.exists(output_dir):
        try:
            os.makedirs(output_dir)
        except Exception as e:
            log_event(f"Error creating directory {output_dir}: {e}", level="error")
            return

    doc_path = os.path.join(output_dir, "Case_Audit.docx")
    doc = Document()
    
    # Set default style to Times New Roman, Size 12, Black
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    from docx.shared import RGBColor
    font.color.rgb = RGBColor(0, 0, 0)
    
    # Title (Custom formatting to match requirements)
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("CASE AUDIT REPORT")
    run.bold = True
    run.underline = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    
    timestamp_p = doc.add_paragraph(f"Generated on: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph("-" * 30)

    for line in report_lines:
        p = doc.add_paragraph()
        
        # Identify Headings ([ SECTION ]) or Subheadings (Name: or 1. Name)
        is_heading = line.strip().startswith("[") and line.strip().endswith("]")
        is_subheading = line.strip().endswith(":") or re.match(r'^\s*[A-Z\d]+\.\s+[A-Z\s]+', line)
        
        run = p.add_run(line)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        
        if is_heading or is_subheading:
            run.bold = True
            run.underline = True
            
        if line.strip().startswith("[✓]") or line.strip().startswith("[x]") or line.strip().startswith("[i]") or line.strip().startswith("[?]"):
            p.paragraph_format.left_indent = Inches(0.5)

    try:
        doc.save(doc_path)
        log_event(f"Audit report saved to: {doc_path}")
    except Exception as e:
        log_event(f"Error saving Word document: {e}", level="error")

def audit_case(case_number):
    log_event(f"--- Starting Audit for {case_number} ---")
    data = load_case_data(case_number)
    if not data:
        return

    root_path = data.get("file_path")
    if not root_path:
        log_event(f"Case path not found in data for {case_number}", level="error")
        return

    report_lines = []
    
    case_name = data.get('case_name') or os.path.basename(root_path)
    header = f"AUDIT REPORT: CASE {case_name} ({case_number})"
    log_event(header)
    report_lines.append(header)
    report_lines.append(f"Physical Path: {root_path}")
    report_lines.append("")
    
    # 1. Deadlines
    deadlines = check_deadlines(data.get("trial_date"))
    report_lines.append("[ CRITICAL DEADLINES ]")
    log_event("[ CRITICAL DEADLINES ]")
    if deadlines:
        today = datetime.date.today()
        for name, date in deadlines.items():
            delta = (date - today).days
            status = "PASSED" if delta < 0 else f"{delta} days left"
            line = f"  {name:<25} : {date} ({status})"
            log_event(line)
            report_lines.append(line)
    else:
        line = "  [!] Trial Date not set. No statutory deadlines calculated."
        log_event(line)
        report_lines.append(line)
    report_lines.append("")

    # 2. Pleadings
    report_lines.append("[ PLEADINGS ]")
    log_event("[ PLEADINGS ]")
    plead_path = find_folder(root_path, FOLDER_MAP["PLEADINGS"])
    if plead_path:
        # Categorize and List Files
        categories = {
            "Complaints/Summons": ["Complaint", "S&C", "Summons"],
            "Answers": ["Answer"],
            "Motions/Demurrers": ["Motion", "Demurrer", "MSJ", "Opposition", "Reply"]
        }
        
        for cat_name, keywords in categories.items():
            files = scan_directory(plead_path, keywords)
            if files:
                line = f"   {cat_name}:"
                report_lines.append(line)
                for f in files:
                    report_lines.append(f"     - {os.path.basename(f)}")
            else:
                report_lines.append(f"   {cat_name}: [x] None identified.")
    else:
        report_lines.append("   [x] No 'PLEADINGS' folder found.")
    report_lines.append("")

    # 3. Discovery
    report_lines.append("[ DISCOVERY ]")
    log_event("[ DISCOVERY ]")
    disc_path = find_folder(root_path, FOLDER_MAP["DISCOVERY"])
    
    if disc_path:
        # -- Propounded --
        report_lines.append("   A. DISCOVERY PROPOUNDED")
        prop_path = find_folder(disc_path, FOLDER_MAP["PROPOUNDED"])
        if prop_path:
            parties = [d for d in os.listdir(prop_path) if os.path.isdir(os.path.join(prop_path, d))]
            if parties:
                for party in parties:
                    report_lines.append(f"      1. Discovery Propounded to {party}:")
                    party_dir = os.path.join(prop_path, party)
                    files = [f for f in os.listdir(party_dir) if os.path.isfile(os.path.join(party_dir, f))]
                    if files:
                        for f in files:
                            report_lines.append(f"         - {f}")
                    else:
                        report_lines.append("         [?] No files found.")
            else:
                report_lines.append("      [?] Propounded folder exists but no party subfolders found.")
        else:
            report_lines.append("      [x] No 'Propounded' folder found.")

        # -- Responses --
        report_lines.append("")
        report_lines.append("   B. DISCOVERY RESPONSES")
        resp_path = find_folder(disc_path, FOLDER_MAP["RESPONSES"])
        if resp_path:
            parties = [d for d in os.listdir(resp_path) if os.path.isdir(os.path.join(resp_path, d))]
            # Also check for files directly in Responses (often miscellaneous or single party)
            root_files = [f for f in os.listdir(resp_path) if os.path.isfile(os.path.join(resp_path, f))]
            
            if parties:
                for party in parties:
                    report_lines.append(f"      1. Discovery Responses from {party}:")
                    party_dir = os.path.join(resp_path, party)
                    files = [f for f in os.listdir(party_dir) if os.path.isfile(os.path.join(party_dir, f))]
                    if files:
                        for f in files:
                            report_lines.append(f"         - {f}")
                    else:
                        report_lines.append("         [?] No files found in folder.")
            
            if root_files:
                report_lines.append("      2. Miscellaneous/Uncategorized Responses:")
                for f in root_files:
                    report_lines.append(f"         - {f}")
            
            if not parties and not root_files:
                report_lines.append("      [?] Responses folder is empty.")
        else:
            report_lines.append("      [x] No 'Responses' folder found.")
    else:
        report_lines.append("   [x] No 'DISCOVERY' folder found.")
    report_lines.append("")

    # 4. Experts
    report_lines.append("[ EXPERTS ]")
    log_event("[ EXPERTS ]")
    exp_path = find_folder(root_path, FOLDER_MAP["EXPERTS"])
    if exp_path:
        # List all folders in Experts
        all_entries = os.listdir(exp_path)
        folders = [e for e in all_entries if os.path.isdir(os.path.join(exp_path, e))]
        
        if folders:
            report_lines.append("   Expert Folders:")
            for folder in folders:
                report_lines.append(f"     - {folder}")
        else:
            report_lines.append("   [i] No subfolders found in EXPERTS.")

        # Search for Designation/Exchange files
        desig_keywords = ["designation", "exchange", "discl", "expert"]
        desig_files = scan_directory(exp_path, desig_keywords)
        
        # Filter for PDFs
        desig_pdfs = [f for f in desig_files if f.lower().endswith(".pdf")]
        
        if desig_pdfs:
            report_lines.append("   Designation/Exchange Documents:")
            for f in desig_pdfs:
                report_lines.append(f"     - {os.path.basename(f)}")
        else:
            report_lines.append("   [x] No Expert Designation/Exchange PDFs identified.")
    else:
        report_lines.append("   [x] No 'EXPERTS' folder found.")
    report_lines.append("")

    # Summary Section
    report_lines.append("[ SUMMARY & ACTION ITEMS ]")
    log_event("[ SUMMARY & ACTION ITEMS ]")
    actions = 0
    if not deadlines:
        report_lines.append("   [!] ACTION: Set trial date to calculate deadlines.")
        actions += 1
    
    # Check for missing basic folders
    if not plead_path:
        report_lines.append("   [!] MISSING: 'PLEADINGS' folder not found.")
        actions += 1
    if not disc_path:
        report_lines.append("   [!] MISSING: 'DISCOVERY' folder not found.")
        actions += 1
    if not exp_path:
        report_lines.append("   [!] MISSING: 'EXPERTS' folder not found.")
        actions += 1
    
    if actions == 0:
        report_lines.append("   [✓] Key file structures identified.")

    # Save to Word
    save_to_docx(root_path, report_lines)
    log_event(f"--- Audit Completed for {case_number} ---")

def main():
    if len(sys.argv) < 2:
        log_event("Usage: python audit.py <file_number> [--headless]", level="error")
        sys.exit(1)

    is_headless = "--headless" in sys.argv
    args_to_parse = [arg for arg in sys.argv[1:] if arg != "--headless"]
    
    raw_arg = " ".join(args_to_parse)
    file_numbers = parse_file_numbers(raw_arg)
    
    if len(file_numbers) > 1:
        log_event(f"Detected multiple file numbers: {file_numbers}. Launching separate agents...")
        for fn in file_numbers:
            log_event(f"Spawning audit agent for {fn}...")
            try:
                spawn_args = [sys.executable, sys.argv[0], fn, "--headless"]
                if os.name == 'nt':
                    subprocess.Popen(spawn_args, creationflags=subprocess.CREATE_NO_WINDOW)
                else:
                    subprocess.Popen(spawn_args)
            except Exception as e:
                log_event(f"Failed to spawn agent for {fn}: {e}", level="error")
        sys.exit(0)
        
    if not file_numbers:
        log_event("No valid file numbers found.", level="error")
        sys.exit(1)
        
    audit_case(file_numbers[0])

if __name__ == "__main__":
    main()