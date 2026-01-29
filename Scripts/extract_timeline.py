"""
Timeline Extraction Agent for iCharlotte

Extracts dates and events from legal documents to build case timelines.

Features:
- Multi-document timeline aggregation
- Date normalization and validation
- Conflict detection between sources
- JSON and DOCX output formats
- Integration with CaseDataManager
"""

import os
import sys
import re
import json
import datetime
from typing import List, Dict, Optional, Any
from dataclasses import dataclass, field, asdict
from docx import Document
from docx.shared import Pt, Inches

# Add parent directory to path for imports
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

# Import shared infrastructure
from icharlotte_core.document_processor import DocumentProcessor, OCRConfig
from icharlotte_core.agent_logger import AgentLogger, create_legacy_log_event
from icharlotte_core.llm_config import LLMCaller
from icharlotte_core.memory_monitor import MemoryMonitor
from icharlotte_core.exceptions import ExtractionError, LLMError

# Import Case Data Manager
try:
    from case_data_manager import CaseDataManager
except ImportError:
    sys.path.append(os.path.join(os.getcwd(), 'Scripts'))
    from case_data_manager import CaseDataManager

# Import Document Registry
try:
    from document_registry import DocumentRegistry, get_available_documents
except ImportError:
    from Scripts.document_registry import DocumentRegistry, get_available_documents

import argparse


# =============================================================================
# Configuration
# =============================================================================

SCRIPTS_DIR = os.path.dirname(os.path.abspath(__file__))
PROMPT_FILE = os.path.join(SCRIPTS_DIR, "TIMELINE_EXTRACTION_PROMPT.txt")
LEGACY_LOG_FILE = r"C:\GeminiTerminal\Timeline_activity.log"


# =============================================================================
# Data Classes
# =============================================================================

@dataclass
class TimelineEvent:
    """Represents a single event in the timeline."""
    date: str  # YYYY-MM-DD format
    date_text: str = ""
    precision: str = "exact"  # exact, approximate, range, relative
    range_end: str = ""
    event: str = ""
    details: str = ""
    source: str = ""
    page: str = ""
    category: str = "other"
    significance: str = "medium"
    parties_involved: List[str] = field(default_factory=list)
    quote: str = ""

    def to_dict(self) -> dict:
        return asdict(self)

    @classmethod
    def from_dict(cls, data: dict) -> 'TimelineEvent':
        return cls(**{k: v for k, v in data.items() if k in cls.__dataclass_fields__})

    def sort_key(self) -> tuple:
        """Return a tuple for sorting by date."""
        try:
            dt = datetime.datetime.strptime(self.date, "%Y-%m-%d")
            return (dt, self.significance != "high", self.event)
        except ValueError:
            return (datetime.datetime.max, True, self.event)


@dataclass
class Timeline:
    """Collection of timeline events."""
    events: List[TimelineEvent] = field(default_factory=list)
    file_number: str = ""
    generated: str = ""
    sources: List[str] = field(default_factory=list)

    def add_event(self, event: TimelineEvent):
        """Add an event and maintain sorted order."""
        self.events.append(event)
        self.events.sort(key=lambda e: e.sort_key())

    def merge(self, other: 'Timeline'):
        """Merge another timeline into this one."""
        for event in other.events:
            self.events.append(event)
        self.sources.extend(other.sources)
        self.sources = list(set(self.sources))
        self.events.sort(key=lambda e: e.sort_key())

    def to_dict(self) -> dict:
        return {
            "file_number": self.file_number,
            "generated": self.generated,
            "sources": self.sources,
            "events": [e.to_dict() for e in self.events]
        }

    @classmethod
    def from_dict(cls, data: dict) -> 'Timeline':
        timeline = cls(
            file_number=data.get("file_number", ""),
            generated=data.get("generated", ""),
            sources=data.get("sources", [])
        )
        for event_data in data.get("events", []):
            timeline.events.append(TimelineEvent.from_dict(event_data))
        return timeline

    def find_conflicts(self) -> List[Dict]:
        """Find potential conflicts between events."""
        conflicts = []

        # Group events by date
        by_date = {}
        for event in self.events:
            if event.date not in by_date:
                by_date[event.date] = []
            by_date[event.date].append(event)

        # Look for potential conflicts
        for date, events in by_date.items():
            if len(events) > 1:
                # Multiple events on same date from different sources
                sources = set(e.source for e in events)
                if len(sources) > 1:
                    conflicts.append({
                        "date": date,
                        "events": [e.event for e in events],
                        "sources": list(sources),
                        "type": "same_date_different_source"
                    })

        return conflicts


# =============================================================================
# Date Parsing Utilities
# =============================================================================

class DateParser:
    """Utilities for parsing and normalizing dates."""

    MONTH_NAMES = {
        'january': 1, 'february': 2, 'march': 3, 'april': 4,
        'may': 5, 'june': 6, 'july': 7, 'august': 8,
        'september': 9, 'october': 10, 'november': 11, 'december': 12,
        'jan': 1, 'feb': 2, 'mar': 3, 'apr': 4, 'jun': 6,
        'jul': 7, 'aug': 8, 'sep': 9, 'sept': 9, 'oct': 10,
        'nov': 11, 'dec': 12
    }

    @classmethod
    def parse_date(cls, date_str: str) -> Optional[str]:
        """
        Parse a date string and return YYYY-MM-DD format.

        Returns None if parsing fails.
        """
        date_str = date_str.strip()

        # Already in YYYY-MM-DD format
        if re.match(r'^\d{4}-\d{2}-\d{2}$', date_str):
            return date_str

        # MM/DD/YYYY or M/D/YYYY
        match = re.match(r'^(\d{1,2})/(\d{1,2})/(\d{4})$', date_str)
        if match:
            month, day, year = match.groups()
            return f"{year}-{int(month):02d}-{int(day):02d}"

        # Month DD, YYYY or Month DDth, YYYY
        match = re.match(
            r'^([A-Za-z]+)\s+(\d{1,2})(?:st|nd|rd|th)?,?\s+(\d{4})$',
            date_str, re.IGNORECASE
        )
        if match:
            month_name, day, year = match.groups()
            month = cls.MONTH_NAMES.get(month_name.lower())
            if month:
                return f"{year}-{month:02d}-{int(day):02d}"

        # DD Month YYYY
        match = re.match(
            r'^(\d{1,2})\s+([A-Za-z]+)\s+(\d{4})$',
            date_str, re.IGNORECASE
        )
        if match:
            day, month_name, year = match.groups()
            month = cls.MONTH_NAMES.get(month_name.lower())
            if month:
                return f"{year}-{month:02d}-{int(day):02d}"

        return None

    @classmethod
    def validate_date(cls, date_str: str) -> bool:
        """Check if a date string is valid."""
        try:
            datetime.datetime.strptime(date_str, "%Y-%m-%d")
            return True
        except ValueError:
            return False


# =============================================================================
# Document Output
# =============================================================================

def save_timeline_to_docx(timeline: Timeline, output_path: str, logger: AgentLogger) -> bool:
    """Save timeline to a DOCX file."""
    try:
        doc = Document()

        # Title
        title = doc.add_paragraph()
        run = title.add_run("CASE TIMELINE")
        run.bold = True
        run.underline = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)

        # Metadata
        doc.add_paragraph(f"File Number: {timeline.file_number}")
        doc.add_paragraph(f"Generated: {timeline.generated}")
        doc.add_paragraph(f"Sources: {', '.join(timeline.sources)}")
        doc.add_paragraph()

        # Category headers
        categories = {}
        for event in timeline.events:
            cat = event.category
            if cat not in categories:
                categories[cat] = []
            categories[cat].append(event)

        # Chronological section
        doc.add_heading("Chronological Timeline", level=1)

        current_year = None
        for event in timeline.events:
            try:
                year = event.date[:4]
                if year != current_year:
                    current_year = year
                    doc.add_heading(year, level=2)
            except:
                pass

            # Event entry
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)

            # Date
            date_run = p.add_run(f"{event.date_text or event.date}: ")
            date_run.bold = True
            date_run.font.name = 'Times New Roman'
            date_run.font.size = Pt(12)

            # Event description
            event_run = p.add_run(event.event)
            event_run.font.name = 'Times New Roman'
            event_run.font.size = Pt(12)

            if event.significance == "high":
                event_run.bold = True

            # Details on new line if present
            if event.details:
                p2 = doc.add_paragraph()
                p2.paragraph_format.left_indent = Inches(0.5)
                details_run = p2.add_run(event.details)
                details_run.font.name = 'Times New Roman'
                details_run.font.size = Pt(11)
                details_run.italic = True

            # Source reference
            if event.source:
                p3 = doc.add_paragraph()
                p3.paragraph_format.left_indent = Inches(0.5)
                source_run = p3.add_run(f"Source: {event.source}")
                source_run.font.name = 'Times New Roman'
                source_run.font.size = Pt(10)
                source_run.font.color.rgb = None  # Gray would be nice

        # By Category section
        doc.add_page_break()
        doc.add_heading("Events by Category", level=1)

        category_names = {
            'incident': 'Incident Events',
            'medical': 'Medical Treatment',
            'employment': 'Employment',
            'legal': 'Legal Proceedings',
            'discovery': 'Discovery',
            'deposition': 'Depositions',
            'other': 'Other Events'
        }

        for cat, events in sorted(categories.items()):
            if not events:
                continue

            doc.add_heading(category_names.get(cat, cat.title()), level=2)

            for event in sorted(events, key=lambda e: e.sort_key()):
                p = doc.add_paragraph(style='List Bullet')
                run = p.add_run(f"{event.date_text or event.date}: {event.event}")
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)

        doc.save(output_path)
        logger.output_file(output_path)
        return True

    except Exception as e:
        logger.error(f"Error saving timeline to DOCX: {e}")
        return False


# =============================================================================
# Main Processing
# =============================================================================

def extract_timeline_from_document(
    input_path: str,
    logger: AgentLogger
) -> Optional[Timeline]:
    """
    Extract timeline events from a single document.

    Args:
        input_path: Path to the document.
        logger: AgentLogger instance.

    Returns:
        Timeline object or None on failure.
    """
    memory_monitor = MemoryMonitor(warn_threshold_mb=1500, abort_threshold_mb=2000, logger=logger.info)
    llm_caller = LLMCaller(logger=logger)

    logger.progress(2, "Initializing timeline extraction...")

    # Load prompt
    try:
        with open(PROMPT_FILE, "r", encoding="utf-8") as f:
            prompt = f.read()
    except Exception as e:
        logger.error(f"Error reading prompt file: {e}")
        return None

    logger.progress(5, "Prompt loaded, starting text extraction...")

    # Extract text (5% - 30%)
    logger.pass_start("Text Extraction", 1, 2)
    logger.progress(8, "Reading document...")
    try:
        with memory_monitor.track_operation("Text Extraction"):
            processor = DocumentProcessor(
                ocr_config=OCRConfig(adaptive=True),
                logger=logger
            )
            logger.progress(12, "Running text extraction (OCR if needed)...")
            result = processor.extract_with_dynamic_ocr(input_path)

            if not result.success:
                raise ExtractionError(f"Failed to extract text: {result.error}")

            text = result.text
            logger.progress(28, f"Extracted {result.char_count} chars")
            logger.info(f"Extracted {result.char_count} chars")

    except Exception as e:
        logger.pass_failed("Text Extraction", str(e), recoverable=False)
        return None

    logger.progress(30, "Text extraction complete")
    logger.pass_complete("Text Extraction", success=True)

    # LLM extraction (30% - 90%)
    logger.pass_start("Timeline Extraction", 2, 2)
    logger.progress(35, "Sending document to LLM for timeline extraction...")
    try:
        with memory_monitor.track_operation("Timeline Extraction"):
            logger.progress(45, "Waiting for LLM response...")
            response = llm_caller.call(prompt, text, task_type="extraction")

            if not response:
                raise LLMError("LLM returned empty response")

            logger.progress(75, "Parsing timeline events from response...")

            # Parse JSON from response
            # Try to extract JSON from the response
            json_match = re.search(r'\{[\s\S]*\}', response)
            if json_match:
                data = json.loads(json_match.group())
            else:
                raise ValueError("No JSON found in response")

            # Create timeline
            timeline = Timeline(
                file_number="",
                generated=datetime.datetime.now().isoformat(),
                sources=[os.path.basename(input_path)]
            )

            logger.progress(82, "Validating and normalizing dates...")
            for event_data in data.get("events", []):
                event = TimelineEvent.from_dict(event_data)

                # Validate/normalize date
                if event.date:
                    normalized = DateParser.parse_date(event.date)
                    if normalized:
                        event.date = normalized
                    elif not DateParser.validate_date(event.date):
                        logger.warning(f"Invalid date format: {event.date}")

                timeline.add_event(event)

            logger.progress(88, f"Extracted {len(timeline.events)} events")
            logger.info(f"Extracted {len(timeline.events)} events")

    except json.JSONDecodeError as e:
        logger.pass_failed("Timeline Extraction", f"JSON parse error: {e}", recoverable=True)
        return None
    except Exception as e:
        logger.pass_failed("Timeline Extraction", str(e), recoverable=True)
        return None

    logger.progress(90, "Timeline extraction complete")
    logger.pass_complete("Timeline Extraction", success=True)
    return timeline


def process_documents(input_paths: List[str], logger: AgentLogger) -> Optional[Timeline]:
    """
    Process multiple documents and merge timelines.

    Args:
        input_paths: List of document paths.
        logger: AgentLogger instance.

    Returns:
        Merged Timeline or None on failure.
    """
    merged = Timeline(
        generated=datetime.datetime.now().isoformat()
    )

    for path in input_paths:
        logger.info(f"Processing: {os.path.basename(path)}")
        timeline = extract_timeline_from_document(path, logger)
        if timeline:
            merged.merge(timeline)

    if not merged.events:
        return None

    return merged


def get_output_directory(input_path: str) -> str:
    """Determine output directory based on input path."""
    parts = input_path.split(os.sep)

    # Find case folder
    for i in range(len(parts) - 1, -1, -1):
        if re.match(r'^\d{3}(\D|$)', parts[i]):
            return os.sep.join(parts[:i+1] + ["NOTES", "AI OUTPUT"])

    # Fallback
    return os.path.join(os.path.dirname(input_path), "AI OUTPUT")


def get_output_directory_for_case(file_number: str) -> str:
    """Get output directory for a case by file number."""
    base = r"C:\Current Clients"

    if os.path.exists(base):
        for client_folder in os.listdir(base):
            client_path = os.path.join(base, client_folder)
            if os.path.isdir(client_path):
                for case_folder in os.listdir(client_path):
                    if file_number in case_folder:
                        return os.path.join(client_path, case_folder, "NOTES", "AI OUTPUT")

    return os.path.join(os.getcwd(), "output")


def find_ai_output_docx(file_number: str) -> Optional[str]:
    """Find the AI_OUTPUT.docx file for a case."""
    base = r"C:\Current Clients"

    if not os.path.exists(base):
        return None

    for client_folder in os.listdir(base):
        client_path = os.path.join(base, client_folder)
        if not os.path.isdir(client_path):
            continue

        for case_folder in os.listdir(client_path):
            if file_number in case_folder:
                ai_output_path = os.path.join(
                    client_path, case_folder, "NOTES", "AI OUTPUT", "AI_OUTPUT.docx"
                )
                if os.path.exists(ai_output_path):
                    return ai_output_path

    return None


def gather_summaries_from_docx(docx_path: str, logger: AgentLogger) -> Dict[str, str]:
    """Extract individual document summaries from AI_OUTPUT.docx."""
    summaries = {}

    try:
        doc = Document(docx_path)

        current_title = None
        current_content = []

        for para in doc.paragraphs:
            text = para.text.strip()

            if not text:
                continue

            # Check if this paragraph is a title (bold + underlined)
            is_title = False
            if para.runs:
                first_run = para.runs[0]
                if first_run.bold and first_run.underline:
                    is_title = True

            if is_title:
                # Save previous summary if exists
                if current_title and current_content:
                    summary_text = "\n".join(current_content)
                    if len(summary_text) > 100:
                        summaries[current_title] = summary_text
                        logger.info(f"Extracted summary: {current_title}")

                current_title = text
                current_content = []

            elif current_title:
                if text.startswith("Generated on:"):
                    continue
                current_content.append(text)

        # Don't forget the last summary
        if current_title and current_content:
            summary_text = "\n".join(current_content)
            if len(summary_text) > 100:
                summaries[current_title] = summary_text
                logger.info(f"Extracted summary: {current_title}")

        logger.info(f"Extracted {len(summaries)} summaries from {docx_path}")

    except Exception as e:
        logger.warning(f"Error reading AI_OUTPUT.docx: {e}")

    return summaries


def gather_case_summaries(file_number: str, logger: AgentLogger,
                          selected_summaries: List[str] = None) -> Dict[str, str]:
    """
    Gather all summaries for a case from AI_OUTPUT.docx and CaseDataManager.

    Args:
        file_number: The case file number.
        logger: AgentLogger instance.
        selected_summaries: Optional list of specific summary names to include.

    Returns:
        Dictionary of document_name -> summary_text.
    """
    summaries = {}

    # First, try to load from AI_OUTPUT.docx (primary source)
    docx_path = find_ai_output_docx(file_number)
    if docx_path:
        logger.info(f"Found AI_OUTPUT.docx: {docx_path}")
        docx_summaries = gather_summaries_from_docx(docx_path, logger)

        for name, text in docx_summaries.items():
            if selected_summaries is not None and name not in selected_summaries:
                continue
            summaries[name] = text
    else:
        logger.info("AI_OUTPUT.docx not found, checking case data store...")

    # Also check CaseDataManager for additional summaries
    try:
        data_manager = CaseDataManager()
        variables = data_manager.get_all_variables(file_number, flatten=False)

        for var_name, var_data in variables.items():
            if any(tag in var_name.lower() for tag in ['summary', 'depo', 'extraction']):
                if selected_summaries is not None and var_name not in selected_summaries:
                    continue

                value = var_data.get('value', '')
                if value and len(value) > 100 and var_name not in summaries:
                    summaries[var_name] = value
                    logger.info(f"Loaded summary from case data: {var_name}")

    except Exception as e:
        logger.warning(f"Error loading case summaries from data store: {e}")

    return summaries


def extract_timeline_from_summaries(
    summaries: Dict[str, str],
    file_number: str,
    logger: AgentLogger
) -> Optional['Timeline']:
    """
    Extract timeline events from document summaries.

    Args:
        summaries: Dictionary of document_name -> summary_text.
        file_number: Case file number.
        logger: AgentLogger instance.

    Returns:
        Timeline object or None on failure.
    """
    if not summaries:
        logger.warning("No summaries provided for timeline extraction")
        return None

    memory_monitor = MemoryMonitor(warn_threshold_mb=1500, abort_threshold_mb=2000, logger=logger.info)
    llm_caller = LLMCaller(logger=logger)

    logger.progress(5, "Initializing timeline extraction from summaries...")

    # Load prompt
    try:
        with open(PROMPT_FILE, "r", encoding="utf-8") as f:
            prompt = f.read()
    except Exception as e:
        logger.error(f"Error reading prompt file: {e}")
        return None

    logger.progress(10, f"Building content from {len(summaries)} summaries...")

    # Build document content
    doc_content = ""
    for name, text in summaries.items():
        doc_content += f"\n\n=== {name} ===\n{text[:15000]}"  # Limit per doc

    logger.progress(15, "Content prepared, starting LLM extraction...")
    logger.pass_start("Timeline Extraction from Summaries", 1, 1)

    try:
        with memory_monitor.track_operation("Timeline Extraction"):
            logger.progress(25, "Sending summaries to LLM...")
            response = llm_caller.call(prompt, doc_content, agent_id="agent_timeline")

            if not response:
                raise LLMError("LLM returned empty response")

            logger.progress(70, "Parsing timeline events from response...")

            # Parse JSON from response
            json_match = re.search(r'\{[\s\S]*\}', response)
            if json_match:
                data = json.loads(json_match.group())
            else:
                raise ValueError("No JSON found in response")

            # Build timeline
            timeline = Timeline(
                file_number=file_number,
                generated=datetime.datetime.now().isoformat(),
                sources=list(summaries.keys())
            )

            logger.progress(80, "Building timeline events...")
            for event_data in data.get('events', []):
                event = TimelineEvent(
                    date=event_data.get('date', ''),
                    date_text=event_data.get('date_text', ''),
                    precision=event_data.get('precision', 'exact'),
                    range_end=event_data.get('range_end', ''),
                    event=event_data.get('event', ''),
                    details=event_data.get('details', ''),
                    source=event_data.get('source', ''),
                    page=event_data.get('page', ''),
                    category=event_data.get('category', 'other'),
                    significance=event_data.get('significance', 'medium'),
                    parties_involved=event_data.get('parties_involved', []),
                    quote=event_data.get('quote', '')
                )
                timeline.add_event(event)

            logger.progress(90, f"Extracted {len(timeline.events)} events from summaries")
            logger.info(f"Extracted {len(timeline.events)} events from summaries")

    except json.JSONDecodeError as e:
        logger.pass_failed("Timeline Extraction", f"JSON parse error: {e}", recoverable=True)
        return None
    except Exception as e:
        logger.pass_failed("Timeline Extraction", str(e), recoverable=True)
        return None

    logger.progress(95, "Timeline extraction complete")
    logger.pass_complete("Timeline Extraction from Summaries", success=True)
    return timeline


def list_available_documents(file_number: str) -> None:
    """List available documents for a case (for pre-selection in UI)."""
    print(f"\nAvailable documents for case {file_number}:")
    print("=" * 60)

    # Get from document registry
    docs = get_available_documents(file_number)

    if docs:
        print(f"\nFrom Document Registry ({len(docs)} documents):")
        for doc in docs:
            print(f"  - {doc['name']}")
            print(f"    Type: {doc['document_type']}")
            print(f"    Agent: {doc['agent']}")
            print()
    else:
        print("\n  No documents in registry. Run summarize agents first.")

    # Also show what's in AI_OUTPUT.docx
    docx_path = find_ai_output_docx(file_number)
    if docx_path:
        print(f"\nFrom AI_OUTPUT.docx ({docx_path}):")
        try:
            doc = Document(docx_path)
            titles = []
            for para in doc.paragraphs:
                if para.runs:
                    first_run = para.runs[0]
                    if first_run.bold and first_run.underline and para.text.strip():
                        titles.append(para.text.strip())
            for title in titles:
                print(f"  - {title}")
            print(f"\n  Total: {len(titles)} summaries in DOCX")
        except Exception as e:
            print(f"  Error reading DOCX: {e}")

    print()


def main():
    """Main entry point."""
    parser = argparse.ArgumentParser(
        description="Extract timeline events from documents or summaries."
    )
    parser.add_argument("input", nargs="?", help="File path, folder path, or case file number")
    parser.add_argument("--summaries", type=str, default=None,
                        help="Comma-separated list of summary names to include (implies summary mode)")
    parser.add_argument("--from-summaries", action="store_true",
                        help="Extract timeline from existing summaries instead of raw documents")
    parser.add_argument("--list", action="store_true",
                        help="List available documents for the case without running extraction")
    parser.add_argument("--file-number", type=str, default=None,
                        help="Case file number (required for --from-summaries or --list)")

    args = parser.parse_args()

    # Handle --list mode
    if args.list:
        file_number = args.file_number or args.input
        if not file_number:
            print("Error: File number required for --list mode")
            sys.exit(1)
        list_available_documents(file_number)
        sys.exit(0)

    # Handle --from-summaries mode
    if args.from_summaries or args.summaries:
        file_number = args.file_number or args.input
        if not file_number or not re.match(r'\d{4}\.\d{3}', file_number):
            print("Error: Valid file number required for summary mode (e.g., 2024.123)")
            sys.exit(1)

        # Parse selected summaries if provided
        selected_summaries = None
        if args.summaries:
            selected_summaries = [s.strip() for s in args.summaries.split(",") if s.strip()]

        # Initialize logger
        logger = AgentLogger("Timeline", file_number=file_number)
        logger.info(f"Extracting timeline from summaries for case: {file_number}")

        if selected_summaries:
            logger.info(f"Filtering to {len(selected_summaries)} selected summaries")

        # Gather summaries
        summaries = gather_case_summaries(file_number, logger, selected_summaries)

        if not summaries:
            logger.error("No summaries found for this case. Run summarization agents first.")
            sys.exit(1)

        logger.info(f"Found {len(summaries)} document summaries")

        # Extract timeline from summaries
        timeline = extract_timeline_from_summaries(summaries, file_number, logger)

        if not timeline:
            logger.error("Failed to extract timeline from summaries")
            sys.exit(1)

        # Determine output path
        output_dir = get_output_directory_for_case(file_number)
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)

        # Save outputs
        docx_path = os.path.join(output_dir, "Case_Timeline.docx")
        json_path = os.path.join(output_dir, "Case_Timeline.json")

        # Save DOCX
        save_timeline_to_docx(timeline, docx_path, logger)

        # Save JSON
        try:
            with open(json_path, "w", encoding="utf-8") as f:
                json.dump(timeline.to_dict(), f, indent=2)
            logger.info(f"Saved JSON to: {json_path}")
        except Exception as e:
            logger.warning(f"Could not save JSON: {e}")

        # Save to CaseDataManager
        try:
            data_manager = CaseDataManager()
            data_manager.save_variable(
                file_number,
                "case_timeline",
                json.dumps(timeline.to_dict()),
                source="timeline_agent",
                extra_tags=["Timeline"]
            )
            logger.info("Saved timeline to case data")
        except Exception as e:
            logger.warning(f"Could not save to case data: {e}")

        logger.info(f"Timeline extraction complete: {len(timeline.events)} events")
        sys.exit(0)

    # Original file/folder mode
    if not args.input:
        print("Error: No input provided.", flush=True)
        print("Usage: python extract_timeline.py <file_or_folder>", flush=True)
        print("       python extract_timeline.py <file_number> --from-summaries", flush=True)
        sys.exit(1)

    input_paths = []
    input_arg = args.input.strip().strip('"').strip("'")
    input_path = os.path.abspath(input_arg)

    if os.path.isdir(input_path):
        # Collect all documents in directory
        for root, _, files in os.walk(input_path):
            for file in files:
                if file.lower().endswith(('.pdf', '.docx')):
                    input_paths.append(os.path.join(root, file))
    elif os.path.isfile(input_path):
        input_paths.append(input_path)

    if not input_paths:
        print("Error: No valid files found.", flush=True)
        sys.exit(1)

    # Extract file number
    file_num_match = re.search(r"(\d{4}\.\d{3})", input_paths[0])
    file_number = file_num_match.group(1) if file_num_match else None

    # Initialize logger
    logger = AgentLogger("Timeline", file_number=file_number)
    logger.info(f"Processing {len(input_paths)} documents for timeline extraction")

    # Process documents
    if len(input_paths) == 1:
        timeline = extract_timeline_from_document(input_paths[0], logger)
    else:
        timeline = process_documents(input_paths, logger)

    if not timeline:
        logger.error("Failed to extract timeline")
        sys.exit(1)

    timeline.file_number = file_number or ""

    # Check for conflicts
    conflicts = timeline.find_conflicts()
    if conflicts:
        logger.warning(f"Found {len(conflicts)} potential date conflicts")

    # Determine output path
    output_dir = get_output_directory(input_paths[0])
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    # Save outputs
    docx_path = os.path.join(output_dir, "Case_Timeline.docx")
    json_path = os.path.join(output_dir, "Case_Timeline.json")

    # Save DOCX
    save_timeline_to_docx(timeline, docx_path, logger)

    # Save JSON
    try:
        with open(json_path, "w", encoding="utf-8") as f:
            json.dump(timeline.to_dict(), f, indent=2)
        logger.info(f"Saved JSON to: {json_path}")
    except Exception as e:
        logger.warning(f"Could not save JSON: {e}")

    # Save to CaseDataManager
    if file_number:
        try:
            data_manager = CaseDataManager()
            data_manager.save_variable(
                file_number,
                "case_timeline",
                json.dumps(timeline.to_dict()),
                source="timeline_agent",
                extra_tags=["Timeline"]
            )
            logger.info("Saved timeline to case data")
        except Exception as e:
            logger.warning(f"Could not save to case data: {e}")

    logger.info(f"Timeline extraction complete: {len(timeline.events)} events")
    sys.exit(0)


if __name__ == '__main__':
    main()
