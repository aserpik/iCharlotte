import os
import sys
import re
import json
import glob
import logging
import datetime
import difflib
from typing import Dict, List, Optional
import concurrent.futures

# Import Case Data Manager
try:
    from case_data_manager import CaseDataManager
except ImportError:
    sys.path.append(os.path.join(os.getcwd(), 'Scripts'))
    from case_data_manager import CaseDataManager

try:
    import fitz  # PyMuPDF
except ImportError:
    print("Error: 'pymupdf' is required. Please run: pip install pymupdf")
    sys.exit(1)

try:
    from google import genai
except ImportError:
    print("Error: 'google-genai' is required.")
    sys.exit(1)

import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import qn
from docx.oxml import OxmlElement
import docx.opc.constants as opc_constants

# Helper for Hyperlinks in docx
def add_hyperlink(paragraph, url, text, color="0000FF", underline=True):
    """Adds a hyperlink to a paragraph with specified color and underline."""
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    if color:
        c = OxmlElement('w:color')
        c.set(qn('w:val'), color)
        rPr.append(c)

    if underline:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)
    
    # Set Font and Size for the link
    f = OxmlElement('w:rFonts')
    f.set(qn('w:ascii'), 'Times New Roman')
    f.set(qn('w:hAnsi'), 'Times New Roman')
    rPr.append(f)
    
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '24') # 12pt
    rPr.append(sz)

    new_run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)
    return hyperlink

# --- Configuration ---
LOG_FILE = os.path.join(os.getcwd(), "Subpoena_Tracker_activity.log")
# GEMINI_DATA_DIR managed by CaseDataManager
API_KEY = os.environ.get("GEMINI_API_KEY")
DEFAULT_MODEL = "gemini-3-flash-preview"
FALLBACK_MODEL = "gemini-2.5-flash"

# Initialize Data Manager
data_manager = CaseDataManager()

# Generic terms to ignore when extracting entity from filename
GENERIC_TERMS = [
    "sdt", "affidavit", "records", "response", "production", "file copy", 
    "subpoena", "pos", "proof of service", "billing", "medical", "images", 
    "films", "radiology", "pathology", "statements", "declaration", "x-ray", "mri",
    "summary", "chronology", "report", "invoice", "cost", "receipt", "sdt", "aff",
    "records recv'd", "records received", "subpoenas issued", "records recv d", "records recvd"
]

NO_RECORDS_TERMS = ["cnr", "cnb", "cnx", "no records", "no billing", "certificate of no", "none found"]

# Setup Logging
logging.basicConfig(
    filename=LOG_FILE,
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    datefmt='%Y-%m-%d %H:%M:%S'
)

def log_event(message, level="info"):
    print(message)
    sys.stdout.flush()
    if level == "info":
        logging.info(message)
    elif level == "error":
        logging.error(message)
    elif level == "warning":
        logging.warning(message)

def get_case_path(file_num: str) -> Optional[str]:
    """Retrieves the physical path from the JSON case data."""
    return data_manager.get_value(file_num, "file_path")

def extract_text_from_pdf_first_page(file_path: str) -> str:
    """Fast text extraction from just the first page."""
    try:
        doc = fitz.open(file_path)
        if len(doc) > 0:
            text = doc[0].get_text()
            doc.close()
            return text
        doc.close()
        return ""
    except Exception as e:
        log_event(f"Error reading PDF {os.path.basename(file_path)}: {e}", level="warning")
        return ""

def identify_metadata_via_llm(text: str) -> Dict[str, str]:
    """Asks Gemini to identify the target entity and issuance date from the subpoena text."""
    if not API_KEY:
        return {"entity": "Unknown (No API Key)", "date": "Unknown"}
    
    client = genai.Client(api_key=API_KEY)
    
    prompt = """
    Analyze the following subpoena text and extract:
    1. The name of the entity or person being subpoenaed (Custodian of Records).
    2. The date the subpoena was issued or signed.
    
    Rules:
    - Entity: Return ONLY the name. No "To:", "Attn:", or "Custodian of".
    - Date: Return in YYYY-MM-DD format. If only a month/year is found, use the 1st of the month.
    - Return as a JSON object: {"entity": "...", "date": "..."}
    - If either is not found, use "Unknown".
    
    Document Text:
    """
    
    for model_name in [DEFAULT_MODEL, FALLBACK_MODEL]:
        try:
            response = client.models.generate_content(
                model=model_name,
                contents=prompt + text[:4000]
            )
            raw_text = response.text.strip()
            if "```json" in raw_text:
                raw_text = raw_text.split("```json")[1].split("```")[0].strip()
            elif "```" in raw_text:
                raw_text = raw_text.split("```")[1].split("```")[0].strip()
                
            data = json.loads(raw_text)
            
            prefixes_to_strip = ["custodian of records", "the custodian of records", "custodian of", "to"]
            clean_name = data.get("entity", "Unknown")
            lower_name = clean_name.lower()
            for p in prefixes_to_strip:
                if lower_name.startswith(p):
                    clean_name = clean_name[len(p):].strip().lstrip(", ")
                    lower_name = clean_name.lower()
            
            data["entity"] = clean_name
            return data
        except:
            continue

    return {"entity": "Unknown", "date": "Unknown"}

def get_plaintiffs(file_num: str) -> List[str]:
    """Retrieves plaintiff names from case data."""
    plaintiffs = data_manager.get_value(file_num, "plaintiffs")
    if plaintiffs:
        return plaintiffs
    return []

def clean_name_string(name: str, item_id: str = "", plaintiffs: List[str] = []) -> str:
    """Cleans up a candidate entity name."""
    name = os.path.splitext(name)[0]
    if item_id and item_id in name:
        name = name.replace(item_id, "")
    
    name = re.sub(r'[\(\)\[\{\]].*?[\)\]\}]', '', name)
    name = name.strip(" -_,")
    
    lower_name = name.lower()
    
    # Check against plaintiffs
    for p in plaintiffs:
        p_clean = p.lower().replace(",", "").strip()
        if p_clean in lower_name or lower_name in p_clean:
            if len(lower_name) > 3:
                return "Unknown"

    clean_candidate = re.sub(r'[^a-z]', ' ', lower_name)
    words = clean_candidate.split()
    if not words or all(w in GENERIC_TERMS or len(w) < 2 for w in words):
        return "Unknown"
        
    return name.strip() if name.strip() else "Unknown"

def clean_filename_entity(filename: str, parent_folder: str, item_id: str, plaintiffs: List[str] = []) -> str:
    """Extracts entity name from filename or parent folder."""
    candidate = "Unknown"
    split_chars = ["_", ","]
    for char in split_chars:
        if char in filename:
            parts = filename.split(char, 1)
            if len(parts) > 1:
                candidate = clean_name_string(parts[1], item_id, plaintiffs)
                if candidate != "Unknown":
                    break
    
    if candidate == "Unknown":
        candidate = clean_name_string(filename, item_id, plaintiffs)
        
    if candidate == "Unknown" or len(candidate) < 3:
        folder_candidate = clean_name_string(parent_folder, item_id, plaintiffs)
        if folder_candidate != "Unknown":
            return folder_candidate
            
    return candidate

def scan_case_files(case_path: str, plaintiffs: List[str] = []):
    """Scans the case directory for relevant files."""
    log_event(f"Scanning directory: {case_path}")
    
    tracked_items: Dict[str, Dict] = {}
    all_record_paths = [] 
    wo_level_files = {} 
    
    id_pattern = re.compile(r'([1-9]\d{4}[-.]\d{4})')
    wo_pattern = re.compile(r'WO#?\s*(\d{5})', re.I)

    for root, dirs, files in os.walk(case_path):
        skip_folders = [
            "billing", "pleading", "pleadings", "motion", "correspondence", "corr", 
            "trial", "deft eurostar's experts", "pltf experts"
        ]
        dirs[:] = [d for d in dirs if d.lower() not in skip_folders]
        
        parent_folder = os.path.basename(root)
        root_upper = root.upper()
        
        # Determine Plaintiff from folder hierarchy
        # We look for folders that match a plaintiff name or have specific patterns
        # Better heuristic: Plaintiff folders are usually subfolders of RECORDS or DISCOVERY
        # containing many subfolders or specific IDs.
        folder_plaintiff = "Unknown"
        path_parts = root.split(os.sep)
        for part in path_parts:
            part_upper = part.upper()
            for p in plaintiffs:
                p_short = p.split()[0].upper() # First name match
                if p_short in part_upper and len(p_short) > 2:
                    folder_plaintiff = p
                    break
            if folder_plaintiff != "Unknown":
                break

        for file in files:
            full_path = os.path.join(root, file)
            lower_name = file.lower()
            
            # 1. Collect all record-related files for fuzzy matching
            if any(k in root_upper for k in ["RECORDS", "JJ PHOTO", "DISCOVERY", "SUBPOENA", "TRANSCRIPT"]):
                 all_record_paths.append(full_path)

            if not file.lower().endswith((".pdf", ".zip", ".docx", ".msg")):
                continue
            
            match = id_pattern.search(file)
            wo_match = wo_pattern.search(file) if not match else None
            
            is_in_subpoena_folder = any(k in root_upper for k in ["SUBPOENAS", "SUBPOENAES"]) and "RECORDS" not in root_upper
            is_in_records_folder = "RECORDS" in root_upper
            
            item_id = None
            if match:
                item_id = match.group(1).replace('.', '-')
            elif wo_match:
                item_id = f"WO-{wo_match.group(1)}"
                wo_id = wo_match.group(1)
                if wo_id not in wo_level_files:
                    wo_level_files[wo_id] = {"pos": [], "subpoena": []}
                
                if "pos" in lower_name and not "proposal" in lower_name:
                    wo_level_files[wo_id]["pos"].append(full_path)
                elif any(k in lower_name for k in ["file copy", "subpoena", "subp"]):
                    wo_level_files[wo_id]["subpoena"].append(full_path)
            elif is_in_subpoena_folder and any(k in lower_name for k in ["subp", "subpoena"]):
                pseudo_entity = clean_filename_entity(file, parent_folder, "", plaintiffs)
                if pseudo_entity != "Unknown":
                    item_id = f"NAME-{pseudo_entity.upper()}"

            if item_id:
                if item_id not in tracked_items:
                    tracked_items[item_id] = {
                        "id": item_id,
                        "entity": "Unknown",
                        "plaintiff": folder_plaintiff,
                        "subpoena_file": None,
                        "issued_date": "Unknown",
                        "record_files": [],
                        "pos_files": [],
                        "status": "Unknown",
                        "is_no_records": False,
                        "page_counts": {} 
                    }
                
                if folder_plaintiff != "Unknown" and tracked_items[item_id]["plaintiff"] == "Unknown":
                    tracked_items[item_id]["plaintiff"] = folder_plaintiff

                if "pos" in lower_name and not "proposal" in lower_name:
                    if full_path not in tracked_items[item_id]["pos_files"]:
                        tracked_items[item_id]["pos_files"].append(full_path)
                    
                elif (any(k in lower_name for k in ["file copy", "subpoena", "subp"]) and not is_in_records_folder) or (is_in_subpoena_folder and "CNR" not in file.upper() and "CNB" not in file.upper()):
                    if tracked_items[item_id]["subpoena_file"] is None or "file copy" in lower_name:
                        tracked_items[item_id]["subpoena_file"] = full_path
                        
                else:
                    if any(k in lower_name for k in NO_RECORDS_TERMS):
                        tracked_items[item_id]["is_no_records"] = True
                    
                    if full_path not in tracked_items[item_id]["record_files"]:
                        # Heuristic: only add to record_files if it looks like production
                        # Exclude transcripts and depo notices from primary matching
                        if not any(k in root_upper for k in ["TRANSCRIPT", "DEPO NTC"]):
                            tracked_items[item_id]["record_files"].append(full_path)
                        
                        if lower_name.endswith(".pdf"):
                            try:
                                with fitz.open(full_path) as doc:
                                    tracked_items[item_id]["page_counts"][full_path] = len(doc)
                            except:
                                pass
                    
                if tracked_items[item_id]["entity"] == "Unknown" or tracked_items[item_id]["entity"].lower() in ["records recv'd", "records received", "subpoenas issued"]:
                    clean_id = item_id.split('-')[-1] if '-' in item_id else item_id
                    entity_guess = clean_filename_entity(file, parent_folder, clean_id, plaintiffs)
                    if entity_guess != "Unknown":
                            tracked_items[item_id]["entity"] = entity_guess

    # Associate WO-level files
    for item_id, data in tracked_items.items():
        wo_prefix = item_id.split('-')[0]
        if wo_prefix in wo_level_files:
            for f in wo_level_files[wo_prefix]["pos"]:
                if f not in data["pos_files"]:
                    data["pos_files"].append(f)
            if data["subpoena_file"] is None and wo_level_files[wo_prefix]["subpoena"]:
                data["subpoena_file"] = wo_level_files[wo_prefix]["subpoena"][0]

    return tracked_items, all_record_paths

def fuzzy_match_record(entity_name: str, file_paths: List[str]) -> List[str]:
    """Checks if any path in file_paths roughly matches the entity name. Returns list of matches."""
    if not entity_name or entity_name == "Unknown":
        return []
        
    matches = []
    ignore_words = ["custodian", "records", "of", "inc", "llc", "corp", "corporation", "the", "center", "medical", "imaging", "clinic", "hospital", "md", "m d"]
    
    clean_entity = entity_name.lower()
    for w in ignore_words:
        clean_entity = clean_entity.replace(w, "")
    clean_entity = re.sub(r'[^a-z0-9]', '', clean_entity)
    
    if len(clean_entity) < 3: 
        return []

    for path in file_paths:
        name = os.path.basename(path).lower()
        if any(k in name for k in ["file copy", "subpoena", "pos "]):
            continue
            
        clean_name = re.sub(r'[^a-z0-9]', '', name)
        
        parent_name = os.path.basename(os.path.dirname(path)).lower()
        clean_parent = re.sub(r'[^a-z0-9]', '', parent_name)
        
        # Prioritize folder matches as they are more likely to be the "Actual Production"
        if clean_entity in clean_parent or clean_entity in clean_name:
            matches.append(path)
            continue
            
        ratio = difflib.SequenceMatcher(None, clean_entity, clean_name).ratio()
        if ratio > 0.75:
            matches.append(path)
            
    # Sort matches so that files in "RECORDS" folder come first
    matches.sort(key=lambda x: "RECORDS" not in x.upper())
    return matches

def generate_docx_report(case_path: str, tracked_items: Dict[str, Dict], stats: Dict):
    """Generates a professional Word document in NOTES/AI OUTPUT with specific formatting."""
    doc = Document()
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    font.color.rgb = RGBColor(0, 0, 0)
    
    p_format = style.paragraph_format
    p_format.line_spacing = 1.0
    p_format.space_after = Pt(0)
    p_format.space_before = Pt(0)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run('Subpoena Tracking Report')
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph(f"Case Path: {case_path}")
    doc.add_paragraph(f"Date Generated: {datetime.date.today().strftime('%m-%d-%y')}")
    doc.add_paragraph("")
    
    doc.add_paragraph().add_run('Summary').bold = True
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, txt in enumerate(['Total Items', 'Completed', 'Pending', 'Overdue']):
        run = hdr_cells[i].paragraphs[0].add_run(txt)
        run.bold = True
    
    row_cells = table.add_row().cells
    row_cells[0].text = str(stats['total'])
    row_cells[1].text = str(stats['completed'])
    row_cells[2].text = str(stats['pending'])
    row_cells[3].text = str(stats['overdue'])
    
    doc.add_paragraph("")
    
    doc.add_paragraph().add_run('Subpoena Details').bold = True
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    headers = ['ID', 'Plaintiff', 'Issued', 'Status', 'Entity']
    for i, txt in enumerate(headers):
        run = hdr_cells[i].paragraphs[0].add_run(txt)
        run.bold = True
    
    sorted_ids = sorted(tracked_items.keys())
    today = datetime.date.today()
    
    for item_id in sorted_ids:
        data = tracked_items[item_id]
        
        has_sub = data["subpoena_file"] is not None
        has_rec = len(data["record_files"]) > 0
        is_nr = data.get("is_no_records", False)
        
        # 1. Strip numbers from Plaintiff
        plaintiff_clean = re.sub(r'[\d.-]', '', data['plaintiff']).strip()
        
        # 2. Format Date as MM-DD-YY
        issued_date_str = data.get("issued_date", "Unknown")
        if issued_date_str != "Unknown":
            try:
                idate = datetime.datetime.strptime(issued_date_str, "%Y-%m-%d")
                issued_date_str = idate.strftime("%m-%d-%y")
            except: pass
        
        status_base = "Unknown"
        if has_sub and has_rec:
            status_base = "COMPLETE" if not is_nr else "COMPLETE (No Recs)"
        elif has_sub and not has_rec:
            status_base = "PENDING"
            raw_date = data.get("issued_date", "Unknown")
            if raw_date != "Unknown":
                try:
                    idate = datetime.datetime.strptime(raw_date, "%Y-%m-%d").date()
                    if (today - idate).days > 30: status_base = "OVERDUE"
                except: pass
        elif not has_sub and has_rec:
            status_base = "UNMATCHED RECORD"
            
        row_cells = table.add_row().cells
        
        # Hyperlink ID to subpoena (Blue and Underlined)
        if data["subpoena_file"]:
            add_hyperlink(row_cells[0].paragraphs[0], data["subpoena_file"], item_id, color="0000FF", underline=True)
        else:
            row_cells[0].text = item_id
            
        row_cells[1].text = plaintiff_clean
        row_cells[2].text = issued_date_str
        
        # Hyperlink Status (Handle Multiple Productions)
        status_p = row_cells[3].paragraphs[0]
        if "COMPLETE" in status_base and data["record_files"]:
            if len(data["record_files"]) == 1:
                add_hyperlink(status_p, data["record_files"][0], status_base, color="0000FF", underline=True)
            else:
                status_p.add_run(status_base + ": ")
                for idx, path in enumerate(data["record_files"]):
                    if idx > 0: status_p.add_run("; ")
                    add_hyperlink(status_p, path, f"No.{idx+1}", color="0000FF", underline=True)
        else:
            status_p.add_run(status_base)
            
        row_cells[4].text = data['entity']

    output_dir = os.path.join(case_path, "NOTES", "AI OUTPUT")
    try:
        os.makedirs(output_dir, exist_ok=True)
        timestamp = datetime.datetime.now().strftime("%Y.%m.%d_%H%M")
        report_path = os.path.join(output_dir, f"Subpoena_Tracker_{timestamp}.docx")
        doc.save(report_path)
        log_event(f"DOCX Report saved to: {report_path}")
    except Exception as e:
        log_event(f"Failed to save DOCX report: {e}", level="error")

def analyze_and_report(case_path: str, tracked_items: Dict[str, Dict], all_record_paths: List[str]):
    """Refines data and calls report generation."""
    log_event("--- Analyzing Identified Items ---")
    
    # 1. Resolve Metadata
    for item_id, data in tracked_items.items():
        if data["subpoena_file"]:
            log_event(f"Extracting metadata for {item_id} from PDF...")
            text = extract_text_from_pdf_first_page(data["subpoena_file"])
            if text:
                metadata = identify_metadata_via_llm(text)
                if metadata["entity"] != "Unknown":
                    data["entity"] = metadata["entity"]
                if metadata["date"] != "Unknown":
                    data["issued_date"] = metadata["date"]
    
    # 2. Match Records
    for item_id, data in tracked_items.items():
        if data["entity"] == "Unknown":
            continue
            
        if not data["record_files"]:
            matches = fuzzy_match_record(data["entity"], all_record_paths)
            # Filter to exclude subpoenas/POS/etc
            matches = [m for m in matches if "pos" not in m.lower() and "subpoena" not in m.lower()]
            data["record_files"] = matches
    
    # 3. Report Stats
    completed = 0
    pending = 0
    overdue = 0
    today = datetime.date.today()
    
    for item_id, data in tracked_items.items():
        has_sub = data["subpoena_file"] is not None
        has_rec = len(data["record_files"]) > 0
        
        if has_sub and has_rec:
            completed += 1
        elif has_sub and not has_rec:
            pending += 1
            raw_date = data.get("issued_date", "Unknown")
            if raw_date != "Unknown":
                try:
                    idate = datetime.datetime.strptime(raw_date, "%Y-%m-%d").date()
                    if (today - idate).days > 30: overdue += 1
                except: pass

    stats = {
        'total': len(tracked_items),
        'completed': completed,
        'pending': pending,
        'overdue': overdue
    }
    generate_docx_report(case_path, tracked_items, stats)

def main():
    if len(sys.argv) < 2:
        print("Usage: python subpoena_tracker.py <file_number>")
        sys.exit(1)
        
    file_num = sys.argv[1]
    log_event(f"Starting Subpoena Tracker for File: {file_num}")
    
    case_path = get_case_path(file_num)
    if not case_path:
        log_event("Could not resolve case path. Exiting.", level="error")
        sys.exit(1)
        
    if not os.path.exists(case_path):
        log_event(f"Directory does not exist: {case_path}", level="error")
        sys.exit(1)
        
    plaintiffs = get_plaintiffs(file_num)
    items, paths = scan_case_files(case_path, plaintiffs)
    
    if not items:
        log_event("No subpoena-related files found.")
    else:
        analyze_and_report(case_path, items, paths)

if __name__ == "__main__":
    main()
