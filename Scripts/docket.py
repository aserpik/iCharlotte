import os
import sys
import logging
import json
import datetime
import glob
import re
import subprocess
import shutil
import time
import gc
from typing import Optional, Dict, Any

try:
    from google import genai
    from docx import Document
    from docx.shared import Pt
    from docx.oxml.shared import OxmlElement, qn
    from docx.opc.constants import RELATIONSHIP_TYPE as RELs
    from pypdf import PdfReader
    import pytesseract
    from pdf2image import convert_from_path
except ImportError as e:
    print(f"Critical Error: Missing dependency {e}. Please ensure google-genai, python-docx, pypdf, pytesseract, pdf2image are installed.")
    sys.exit(1)

# Import Case Data Manager
try:
    from case_data_manager import CaseDataManager
except ImportError:
    # If run from root, Scripts/ might not be in path for imports
    sys.path.append(os.path.join(os.getcwd(), 'Scripts'))
    from case_data_manager import CaseDataManager

# --- Configuration ---
# Directories
BASE_PATH_WIN = r"Z:\\Shared\\Current Clients"
# GEMINI_DATA_DIR managed by CaseDataManager
LOG_FILE = os.path.join(os.getcwd(), "Docket_activity.log")
SCRIPTS_DIR = os.path.join(os.getcwd(), "Scripts")
SCRAPER_SCRIPT = os.path.join(SCRIPTS_DIR, "docket_scraper.py")
RIVERSIDE_SCRAPER_SCRIPT = os.path.join(SCRIPTS_DIR, "riverside_docket_scraper.py")
KERN_SCRAPER_SCRIPT = os.path.join(SCRIPTS_DIR, "kern_docket_scraper.py")
SB_SCRAPER_SCRIPT = os.path.join(SCRIPTS_DIR, "san_bernardino_docket_scraper.py")
SD_SCRAPER_SCRIPT = os.path.join(SCRIPTS_DIR, "san_diego_docket_scraper.py")
OC_SCRAPER_SCRIPT = os.path.join(SCRIPTS_DIR, "orange_county_docket_scraper.py")
SACRAMENTO_SCRAPER_SCRIPT = os.path.join(SCRIPTS_DIR, "sacramento_docket_scraper.py")
SC_SCRAPER_SCRIPT = os.path.join(SCRIPTS_DIR, "santa_clara_docket_scraper.py")
ALAMEDA_SCRAPER_SCRIPT = os.path.join(SCRIPTS_DIR, "alameda_docket_scraper.py")
MARIPOSA_SCRAPER_SCRIPT = os.path.join(SCRIPTS_DIR, "mariposa_docket_scraper.py")
COMPLAINT_SCRIPT = os.path.join(SCRIPTS_DIR, "complaint.py")
PROCEDURAL_HIST_PROMPT_FILE = os.path.join(SCRIPTS_DIR, "Procedural_History.txt")
MASTER_STATUS_PATH = r"C:\Users\ASerpik.DESKTOP-MRIMK0D\OneDrive - Bordin Semmer LLP\Desktop\MASTER_CASE_STATUS.docx"

# Initialize Data Manager
data_manager = CaseDataManager()

# OCR Settings
OCR_AVAILABLE = False
POPPLER_PATH = None

if os.name == 'nt':
    # Tesseract Path
    tesseract_paths = [
        r"C:\\Program Files\\Tesseract-OCR\\tesseract.exe",
        os.path.expanduser(r"~\\AppData\\Local\\Programs\\Tesseract-OCR\\tesseract.exe")
    ]
    for p in tesseract_paths:
        if os.path.exists(p):
            pytesseract.pytesseract.tesseract_cmd = p
            OCR_AVAILABLE = True
            break
    
    # Poppler Path
    poppler_paths = [
        r"C:\\Program Files\\poppler\\Library\\bin",
        r"C:\\Program Files\\poppler-0.68.0\\bin",
        r"C:\\Program Files (x86)\\poppler\\Library\\bin"
    ]
    for p in poppler_paths:
        if os.path.exists(p):
            POPPLER_PATH = p
            break
else:
    if shutil.which("tesseract"):
        OCR_AVAILABLE = True
    if shutil.which("pdftoppm"):
        POPPLER_PATH = None

# Setup Logging
logging.basicConfig(
    filename=LOG_FILE,
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    datefmt='%Y-%m-%d %H:%M:%S'
)

def log_event(message, level="info"):
    timestamp = datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    try:
        print(f"[{timestamp}] {message}", flush=True)
    except UnicodeEncodeError:
        try:
             print(f"[{timestamp}] {message}".encode(sys.stdout.encoding or 'utf-8', errors='replace').decode(sys.stdout.encoding or 'utf-8'), flush=True)
        except Exception:
             pass  # Silently ignore if stdout is broken
    except OSError:
        pass  # stdout pipe broken (common when running multiple agents)
    if level == "info":
        logging.info(message)
    elif level == "error":
        logging.error(message)
    elif level == "warning":
        logging.warning(message)

def parse_file_numbers(input_str: str):
    """Parses a string containing file numbers, ranges, or commas into a list."""
    # Remove spaces and split by comma
    parts = [p.strip() for p in input_str.split(',')]
    files = []
    
    for part in parts:
        part = part.strip("'").strip('"')
        if '-' in part:
            try:
                start_str, end_str = part.split('-')
                start_str = start_str.strip()
                end_str = end_str.strip()
                
                # Assume format Prefix.Suffix (e.g., 5800.004)
                if '.' in start_str:
                    prefix_start, suffix_start = start_str.rsplit('.', 1)
                    
                    # Handle end string formats: "5800.008" or just "008"
                    if '.' in end_str:
                        prefix_end, suffix_end = end_str.rsplit('.', 1)
                        if prefix_start != prefix_end:
                            # If prefixes differ, treat as separate items? 
                            # For safety, let's just log warning or fallback to simple processing if needed.
                            # But assuming user meant a range within the same prefix.
                            pass 
                    else:
                        suffix_end = end_str # e.g. "5800.004-008"
                    
                    s_val = int(suffix_start)
                    e_val = int(suffix_end)
                    padding = len(suffix_start)
                    
                    if s_val > e_val:
                        # Descending range? Swap or just handle
                        s_val, e_val = e_val, s_val
                        
                    for i in range(s_val, e_val + 1):
                        files.append(f"{prefix_start}.{str(i).zfill(padding)}")
                else:
                    # Integer range
                    s_val = int(start_str)
                    e_val = int(end_str)
                    if s_val > e_val:
                         s_val, e_val = e_val, s_val
                    for i in range(s_val, e_val + 1):
                        files.append(str(i))
                        
            except ValueError:
                log_event(f"Failed to parse range: {part}. Treating as literal.", level="warning")
                files.append(part)
        else:
            if part:
                files.append(part)
                
    return files

# --- Centralized Path Resolution ---
import sys
import os
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from icharlotte_core.utils import get_case_path
from icharlotte_core.master_db import MasterCaseDatabase

def find_complaint_file(case_path: str) -> Optional[str]:
    """Finds the primary complaint in the PLEADINGS or PLEADING directory and specific subfolders."""
    
    # 1. Identify Base Directories
    base_dirs = []
    for d in ["PLEADINGS", "PLEADING"]:
        path = os.path.join(case_path, d)
        if os.path.exists(path):
            base_dirs.append(path)
    
    if not base_dirs:
        log_event(f"Neither PLEADINGS nor PLEADING directory found at {case_path}", level="warning")
        return None

    # 2. Identify Search Paths (Base + Specific Subfolders)
    search_paths = list(base_dirs) # Start with base directories
    target_subs = ["complaint", "fac", "sac", "s&c"]
    
    for base in base_dirs:
        try:
            # List immediate subdirectories
            with os.scandir(base) as it:
                for entry in it:
                    if entry.is_dir() and entry.name.lower() in target_subs:
                        search_paths.append(entry.path)
        except OSError as e:
            log_event(f"Error scanning directory {base}: {e}", level="warning")

    candidates = []
    # 3. Search for files with "complaint" in name within identified paths
    for search_dir in search_paths:
        try:
            for root, _, files in os.walk(search_dir):
                for file in files:
                    if file.lower().endswith(".pdf") and "complaint" in file.lower():
                        candidates.append(os.path.join(root, file))
        except OSError as e:
             log_event(f"Error walking directory {search_dir}: {e}", level="warning")
    
    # Priority filtering
    filtered_candidates = []
    ignore_terms = ["motion", "response", "demurrer", "answer", "reply"]
    
    for c in candidates:
        name = os.path.basename(c).lower()
        if not any(term in name for term in ignore_terms):
            filtered_candidates.append(c)
    
    if not filtered_candidates:
        if candidates:
            log_event("All complaint candidates filtered out by priority terms. Using first raw candidate.", level="warning")
            return candidates[0]
        log_event("No complaint file found in PLEADINGS/PLEADING or target subfolders.", level="warning")
        return None

    return filtered_candidates[0]

def extract_text_from_file(file_path: str, force_ocr_first_page: bool = False) -> str:
    """Extracts text, falling back to OCR if needed."""
    if not file_path or not os.path.exists(file_path):
        return ""

    log_event(f"Extracting text from: {file_path}")
    text = ""
    ext = os.path.splitext(file_path)[1].lower()

    try:
        if ext == ".pdf":
            reader = PdfReader(file_path)
            
            def ocr_page(page_idx):
                if not OCR_AVAILABLE:
                    return ""
                try:
                    images = convert_from_path(file_path, first_page=page_idx+1, last_page=page_idx+1, poppler_path=POPPLER_PATH)
                    if images:
                        return pytesseract.image_to_string(images[0])
                except Exception as e:
                    log_event(f"OCR failed for page {page_idx+1}: {e}", level="error")
                return ""

            for i, page in enumerate(reader.pages):
                if i % 10 == 0:
                    gc.collect()

                if force_ocr_first_page and i == 0:
                    page_text = ocr_page(i)
                else:
                    page_text = page.extract_text() or ""
                    if len(page_text.strip()) < 100:
                        log_event(f"Page {i+1} insufficient text (<100 chars). Triggering OCR.", level="warning")
                        ocr_txt = ocr_page(i)
                        if len(ocr_txt.strip()) > len(page_text.strip()):
                            page_text = ocr_txt

                text += page_text + "\n"

        elif ext in [".docx", ".doc"]:
            if ext == ".docx":
                doc = Document(file_path)
                # Text from paragraphs
                for p in doc.paragraphs:
                    text += p.text + "\n"
                # Text from tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            text += cell.text + "\n"
        
        return text

    except Exception as e:
        log_event(f"Error extracting text from {file_path}: {e}", level="error")
        return ""

def apply_markdown(cell, text: str):
    """Parses text for specific subheading formatting and basic bold markdown."""
    cell.text = ""
    lines = text.split('\n')
    
    # Regex for "A. Title" style subheadings
    # Matches: Start of line (ignoring leading whitespace), One Capital Letter, Dot, Spaces, Rest of line
    subheading_pattern = re.compile(r"^([A-Z]\.)\s+(.+)$")

    for line in lines:
        stripped_line = line.strip()
        paragraph = cell.add_paragraph()
        
        match = subheading_pattern.match(stripped_line)
        if match:
            letter_part = match.group(1) # e.g. "A."
            # Strip potential markdown from title part if we are forcing formatting
            title_part = match.group(2).replace('**', '') 
            
            # Letter: Bold only
            run_letter = paragraph.add_run(letter_part)
            run_letter.font.bold = True
            run_letter.font.underline = False
            
            # Separator (Tab for clean alignment)
            paragraph.add_run("\t")
            
            # Title: Bold and Underlined
            run_title = paragraph.add_run(title_part)
            run_title.font.bold = True
            run_title.font.underline = True
            
        else:
            # Standard bold processing for other lines
            parts = line.split('**')
            for i, part in enumerate(parts):
                run = paragraph.add_run(part)
                if i % 2 == 1:
                    run.font.bold = True

def update_variables_docx(case_dir: str, data: Dict):
    """Compiles extracted variables into variables.docx in NOTES/AI OUTPUT folder."""
    notes_dir = os.path.join(case_dir, "NOTES", "AI OUTPUT")
    if not os.path.exists(notes_dir):
        os.makedirs(notes_dir)
    
    docx_path = os.path.join(notes_dir, "variables.docx")
    
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    
    doc.add_heading(f"Case Variables: {data.get('case_number', 'Unknown')}", 0)
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Variable'
    hdr_cells[1].text = 'Value'
    
    for k, v in data.items():
        if v is None or v == "None": continue
        if k == "hidden_complaint_text": continue # Skip hidden text variable
        row_cells = table.add_row().cells
        row_cells[0].text = str(k)
        if k in ["factual_background", "procedural_history"]:
            apply_markdown(row_cells[1], str(v))
        else:
            row_cells[1].text = str(v)
        
    doc.save(docx_path)
    log_event(f"Saved variables to {docx_path}")

def call_gemini_api(prompt: str, context_text: str, models: list = ["gemini-2.0-flash", "gemini-1.5-flash"]) -> Optional[str]:
    """Calls Gemini with fallback models."""
    api_key = os.environ.get("GEMINI_API_KEY")
    if not api_key:
        log_event("GEMINI_API_KEY not set.", level="error")
        return None

    client = genai.Client(api_key=api_key)
    # Limit context
    full_prompt = f"{prompt}\n\n[CONTEXT DOCUMENT]\n{context_text[:100000]}" 

    for model_name in models:
        try:
            log_event(f"Querying model: {model_name}")
            response = client.models.generate_content(
                model=model_name,
                contents=full_prompt
            )
            if response and response.text:
                return response.text
        except Exception as e:
            log_event(f"Model {model_name} failed: {e}", level="warning")
            continue
    
    log_event("All AI models failed.", level="error")
    return None

def find_caption_doc(case_path: str) -> Optional[str]:
    """Finds a document with 'CAPTION' in title. Priority: Root folder > Subfolders. .docx > .pdf."""
    root_candidates = []
    try:
        with os.scandir(case_path) as it:
            for entry in it:
                if entry.is_file() and "caption" in entry.name.lower():
                    if entry.name.lower().endswith((".docx", ".doc", ".pdf")):
                        root_candidates.append(entry.path)
    except OSError:
        pass

    if root_candidates:
        docx_files = [f for f in root_candidates if f.lower().endswith((".docx", ".doc"))]
        if docx_files: return docx_files[0]
        return root_candidates[0]

    for root, _, files in os.walk(case_path):
        if os.path.abspath(root) == os.path.abspath(case_path): continue
        for file in files:
            if "caption" in file.lower() and file.lower().endswith((".docx", ".doc", ".pdf")):
                return os.path.join(root, file)
    return None

def re_extract_case_number(file_num: str, complaint_text: str) -> Optional[str]:
    """Uses AI to carefully re-extract case number from complaint text or caption docs when search fails."""
    
    # Priority 1: Check if there's a Caption doc with cleaner text
    case_path = get_case_path(file_num)
    caption_text = ""
    if case_path:
        caption_file = find_caption_doc(case_path)
        if caption_file:
            log_event(f"RECOVERY: Found Caption document: {caption_file}")
            caption_text = extract_text_from_file(caption_file)
    
    combined_context = f"[COMPLAINT TEXT]\n{complaint_text[:20000]}\n\n[CAPTION DOC TEXT]\n{caption_text}"
    
    prompt = """
    The previously extracted case number for this legal document was reported as incorrect or failed to return results from the court search.
    Please carefully re-examine the provided texts (Complaint and/or Caption Document) and extract the correct Court Case Number.
    
    Look for patterns like:
    - 24STCV12345, 25SMCV04926, 23AHCV00123 (Los Angeles)
    - CVPS2401234, CVMV2304567 (Riverside)
    - CIVSB2401234 (San Bernardino)
    - 24CV12345 (Santa Clara)
    
    Ensure you capture the full, exact case number string. 
    Return ONLY the exact case number string. If you absolutely cannot find it, return "None".
    """
    result = call_gemini_api(prompt, combined_context, ["gemini-3-flash-preview", "gemini-2.5-flash"])
    if result:
        clean_result = result.strip().split('\n')[0].strip()
        if clean_result.lower() != "none" and len(clean_result) >= 5:
            return clean_result
    return None

def clean_json_string(s: str) -> str:
    """Cleans Markdown code blocks from string to get raw JSON."""
    s = s.strip()
    if s.startswith("```"):
        lines = s.split('\n')
        if len(lines) >= 2:
            s = '\n'.join(lines[1:-1])
    return s.strip()

def add_hyperlink(paragraph, url, text, color="0000FF", underline=True):
    """
    A helper function to place a hyperlink within a paragraph object.
    """
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, RELs.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # Create a w:r element
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Add color if it is given
    if color:
        c = OxmlElement('w:color')
        c.set(qn('w:val'), color)
        rPr.append(c)

    # Remove underlining if it is requested
    if not underline:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'none')
        rPr.append(u)

    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    
    return hyperlink

def update_master_status(file_num: str, data: Dict[str, Any], docket_path: str = None):
    """Updates the Master Case Status Word document."""
    doc_path = MASTER_STATUS_PATH
    
    # Helper: Format dates to YYYY-MM-DD
    def format_date(date_str):
        if not date_str or date_str == "None":
            return date_str
        try:
            dt = datetime.datetime.strptime(date_str, "%Y-%m-%d")
            return dt.strftime("%Y-%m-%d")
        except ValueError:
            return date_str

    # Helper: Clean text (remove []')
    def clean_text(text):
        if not text: return ""
        return str(text).replace("[", "").replace("]", "").replace("'", "")

    # Helper: Format hearings to "Name (YYYY-MM-DD)"
    def format_hearings(hearings_str):
        if not hearings_str or hearings_str == "None":
            return hearings_str
        
        # Abbreviations
        hearing_map = [
            (r"(?i)Final Status Conference", "FSC"),
            (r"(?i)Trial Setting Conference", "TSC"),
            (r"(?i)Trial Readiness Conference", "TRC"),
            (r"(?i)Order to Show Cause", "OSC"),
            (r"(?i)Motion for Summary Judgment", "MSJ")
        ]
        for pattern, replacement in hearing_map:
            hearings_str = re.sub(pattern, replacement, hearings_str)
        
        # Reformat dates
        # Regex for YYYY-MM-DD
        date_pattern = re.compile(r"(\d{4}-\d{2}-\d{2})")
        
        parts = hearings_str.split(',') # Assume comma separation for multiple
        new_parts = []
        for part in parts:
            part = part.strip()
            match = date_pattern.search(part)
            if match:
                date_str = match.group(1)
                try:
                    dt = datetime.datetime.strptime(date_str, "%Y-%m-%d")
                    # Keep ISO format for consistency
                    formatted_date = dt.strftime("%Y-%m-%d")
                    # Remove date and "on" from name
                    name = part.replace(date_str, "").replace(" on ", " ").strip()
                    new_parts.append(f"{name} ({formatted_date})")
                except ValueError:
                    new_parts.append(part)
            else:
                new_parts.append(part)
        
        return ", ".join(new_parts)

    if not os.path.exists(doc_path):
        log_event("Master Status doc not found. Creating new one.")
        doc = Document()
        table = doc.add_table(rows=1, cols=8)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = "File Number"
        hdr[1].text = "Plaintiff Name"
        hdr[2].text = "Case number"
        hdr[3].text = "County"
        hdr[4].text = "Assigned Associate"
        hdr[5].text = "Trial Date"
        hdr[6].text = "Other Hearings"
        hdr[7].text = "Last Updated"
    else:
        try:
            doc = Document(doc_path)
        except Exception as e:
            log_event(f"Failed to open Master Status doc: {e}", level="error")
            return

    if not doc.tables:
        table = doc.add_table(rows=1, cols=8)
        # ... header setup ...
    else:
        table = doc.tables[0]

    # Look for existing row
    target_row = None
    for row in table.rows:
        if row.cells[0].text.strip() == file_num:
            target_row = row
            break
    
    if not target_row:
        target_row = table.add_row()
    
    # Update cells
    try:
        # 1. File Number
        cell_fn = target_row.cells[0]
        cell_fn.text = "" 
        case_folder = get_case_path(file_num)
        if case_folder:
            p = cell_fn.paragraphs[0]
            add_hyperlink(p, case_folder, clean_text(file_num))
        else:
            cell_fn.text = clean_text(file_num)

        # 2. Plaintiff Name
        plaintiffs_raw = clean_text(data.get("plaintiffs", ""))
        if plaintiffs_raw:
            first_pl = plaintiffs_raw.split(',')[0].strip()
            last_name = first_pl.split()[-1] if first_pl else ""
            target_row.cells[1].text = last_name
        else:
            target_row.cells[1].text = ""

        # 3. Case Number
        cell_cn = target_row.cells[2]
        cell_cn.text = ""
        case_num_str = clean_text(data.get("case_number", ""))
        if docket_path:
             p = cell_cn.paragraphs[0]
             add_hyperlink(p, docket_path, case_num_str)
        else:
             cell_cn.text = case_num_str

        # 4. County
        target_row.cells[3].text = clean_text(data.get("venue_county", ""))

        # 5. Associate
        target_row.cells[4].text = ""

        # 6. Trial Date
        target_row.cells[5].text = format_date(clean_text(str(data.get("trial_date", ""))))

        # 7. Other Hearings
        raw_hearings = clean_text(str(data.get("other_hearings", "")))
        target_row.cells[6].text = format_hearings(raw_hearings)

        # 8. Last Updated
        target_row.cells[7].text = datetime.datetime.now().strftime("%Y-%m-%d")

        # Apply Styling (Times New Roman, Size 8)
        for cell in target_row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(8)
                
                # Hyperlinks style
                for child in paragraph._element:
                    if child.tag.endswith('hyperlink'):
                        for subchild in child:
                            if subchild.tag.endswith('r'):
                                from docx.text.run import Run
                                run_obj = Run(subchild, paragraph)
                                run_obj.font.name = 'Times New Roman'
                                run_obj.font.size = Pt(8)

        # Removed automatic sorting to allow user to sort manually
        # and preserve their sort order if they save it.
        # New entries will be appended to the bottom.

        doc.save(doc_path)
        log_event(f"Updated Master Case Status at {doc_path}")
    except Exception as e:
        log_event(f"Error saving Master Status doc: {e}", level="error")

def run_complaint_agent(file_num: str):
    """Runs the complaint agent script."""
    log_event(f"Triggering Complaint Agent for {file_num}...")
    try:
        # Capture output to diagnose immediate failures (e.g. import errors)
        result = subprocess.run(
            [sys.executable, COMPLAINT_SCRIPT, file_num], 
            capture_output=True, 
            text=True,
            check=True
        )
        log_event("Complaint Agent finished.")
        # log_event(f"Complaint Agent Output:\n{result.stdout}") # Optional: detailed logging
    except subprocess.CalledProcessError as e:
        log_event(f"Complaint Agent failed with exit code {e.returncode}.", level="error")
        log_event(f"STDOUT: {e.stdout}", level="error")
        log_event(f"STDERR: {e.stderr}", level="error")
    except Exception as e:
        log_event(f"Failed to run Complaint Agent: {e}", level="error")

def create_summary_report(results):
    """Creates a Word document summary of the processed file numbers and their status."""
    timestamp = datetime.datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    report_filename = f"dockets_processed_{timestamp}.docx"
    
    try:
        doc = Document()
        doc.add_heading('Dockets Processed Report', 0)
        doc.add_paragraph(f"Report Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'File Number'
        hdr_cells[1].text = 'Status'

        for fn, status in results:
            row_cells = table.add_row().cells
            row_cells[0].text = str(fn)
            row_cells[1].text = str(status)
            
        doc.save(report_filename)
        log_event(f"Summary report created: {report_filename}")
        return report_filename
    except Exception as e:
        log_event(f"Failed to create summary report: {e}", level="error")
        return None

def main():
    if len(sys.argv) < 2:
        log_event("Usage: python docket.py <file_number> [--headless] [--headful]", level="error")
        sys.exit(1)

    # Check for headless/headful flags
    is_headless = "--headless" in sys.argv
    is_headful = "--headful" in sys.argv
    
    args_to_parse = [arg for arg in sys.argv[1:] if arg not in ["--headless", "--headful"]]
    
    raw_arg = " ".join(args_to_parse)
    file_numbers = parse_file_numbers(raw_arg)
    
    # --- Dispatcher Mode ---
    if len(file_numbers) > 1:
        log_event(f"Detected multiple file numbers: {file_numbers}. Launching separate agents in batches...")
        
        MAX_CONCURRENT = 5
        AGENT_TIMEOUT = 600 # 10 minutes
        results = []
        processes = []
        
        # We use a simple loop to manage concurrent processes
        file_queue = list(file_numbers)
        active_processes = {} # pid -> (file_num, process_obj, start_time)
        
        while file_queue or active_processes:
            # Fill slots
            while file_queue and len(active_processes) < MAX_CONCURRENT:
                fn = file_queue.pop(0)
                log_event(f"Spawning agent for {fn}...")
                try:
                    spawn_args = [sys.executable, sys.argv[0], fn, "--headful"]
                    if is_headless:
                        spawn_args.append("--headless")
                    
                    if os.name == 'nt':
                        p = subprocess.Popen(spawn_args, creationflags=subprocess.CREATE_NO_WINDOW)
                    else:
                        p = subprocess.Popen(spawn_args)
                    
                    active_processes[p.pid] = (fn, p, time.time())
                except Exception as e:
                    log_event(f"Failed to spawn agent for {fn}: {e}", level="error")
                    results.append((fn, f"Failed to Start: {e}"))
            
            # Check for finished or timed-out processes
            if active_processes:
                to_remove = []
                for pid, (fn, p, start_time) in active_processes.items():
                    elapsed = time.time() - start_time
                    if p.poll() is not None:
                        # Process finished
                        if p.returncode == 0:
                            results.append((fn, "Success"))
                        else:
                            results.append((fn, f"Failed (Exit Code: {p.returncode})"))
                        to_remove.append(pid)
                    elif elapsed > AGENT_TIMEOUT:
                        # Timeout
                        log_event(f"Agent for {fn} timed out after {AGENT_TIMEOUT}s. Terminating.", level="warning")
                        p.terminate()
                        results.append((fn, "Timed Out"))
                        to_remove.append(pid)
                
                for pid in to_remove:
                    del active_processes[pid]
            
            if file_queue or active_processes:
                time.sleep(1) # Wait before next check
        
        create_summary_report(results)
        log_event(f"Completed processing {len(file_numbers)} file numbers. Summary report created.")
        sys.exit(0)
        
    # --- Worker Mode ---
    if not file_numbers:
        log_event("No valid file numbers found.", level="error")
        sys.exit(1)
        
    file_num = file_numbers[0]
    log_event(f"--- Starting Docket Agent for File: {file_num} (Headless: {is_headless}) ---")
    log_event("Starting preparation phase...")

    # --- Phase 1: Preparation ---
    case_number = data_manager.get_value(file_num, "case_number")
    venue_county = data_manager.get_value(file_num, "venue_county")
    
    missing_case = not case_number or str(case_number).lower() in ["null", "n/a", "none"]
    missing_venue = not venue_county or str(venue_county).lower() in ["null", "n/a", "none"]

    if missing_case or missing_venue:
        log_event(f"Missing critical data (Case #: {case_number}, Venue: {venue_county}) for {file_num}. Running Complaint Agent.")
        run_complaint_agent(file_num)
        
        # Reload
        case_number = data_manager.get_value(file_num, "case_number")
        venue_county = data_manager.get_value(file_num, "venue_county")

        # Validate again
        missing_case = not case_number or str(case_number).lower() in ["null", "n/a", "none"]
        if missing_case:
            log_event("Error: Complaint Agent failed to retrieve case number. Terminating.", level="error")
            sys.exit(1)
            
    log_event(f"Found Case Number: {case_number}. Proceeding to Phase 2.")

    # --- Phase 2: Data Retrieval ---
    if venue_county:
        venue_county = str(venue_county).lower().strip()
    
    selected_scraper = None
    use_interactive_mode = False

    extra_args = []
    
    if venue_county == "riverside":
        log_event("Venue is Riverside. Using Riverside Scraper.")
        selected_scraper = RIVERSIDE_SCRAPER_SCRIPT
        # Even if headless, Riverside might need some handling or it might just work if fully automated
        use_interactive_mode = not is_headless 
    elif venue_county == "san bernardino":
        log_event("Venue is San Bernardino. Using San Bernardino Scraper.")
        selected_scraper = SB_SCRAPER_SCRIPT
    elif venue_county == "kern":
        log_event("Venue is Kern. Using Kern Scraper.")
        selected_scraper = KERN_SCRAPER_SCRIPT
    elif venue_county == "san diego":
        log_event("Venue is San Diego. Using San Diego Scraper.")
        selected_scraper = SD_SCRAPER_SCRIPT
    elif venue_county == "orange":
        log_event("Venue is Orange County. Using Orange County Scraper.")
        selected_scraper = OC_SCRAPER_SCRIPT
        
        # Orange County requires filing year
        filing_date = data_manager.get_value(file_num, "filing_date")
        if filing_date and "-" in str(filing_date):
            year = str(filing_date).split("-")[0]
            extra_args.append(year)
        else:
            current_year = datetime.datetime.now().strftime("%Y")
            log_event(f"Filing date ({filing_date}) missing/invalid for Orange County. Defaulting to {current_year}.", level="warning")
            extra_args.append(current_year)
            
    elif venue_county == "los angeles":
        log_event("Venue is Los Angeles. Using LA Scraper.")
        selected_scraper = SCRAPER_SCRIPT
    elif venue_county == "sacramento":
        log_event("Venue is Sacramento. Using Sacramento Scraper.")
        selected_scraper = SACRAMENTO_SCRAPER_SCRIPT
    elif venue_county == "santa clara":
        log_event("Venue is Santa Clara. Using Santa Clara Scraper.")
        selected_scraper = SC_SCRAPER_SCRIPT
    elif venue_county == "alameda":
        log_event("Venue is Alameda. Using Alameda Scraper.")
        selected_scraper = ALAMEDA_SCRAPER_SCRIPT
    elif venue_county == "mariposa":
        log_event("Venue is Mariposa. Using Mariposa Scraper.")
        selected_scraper = MARIPOSA_SCRAPER_SCRIPT
    else:
        # Prompt: "I want the agent to do this only if the venue_county variable is Los Angeles. 
        # However, if the venue_county variable is riverside..."
        log_event(f"Venue '{venue_county}' is not supported (LA, SB, Riverside, SD, OC, Sacramento, Santa Clara, Alameda, or Mariposa only). Skipping download.", level="warning")
        # Do not exit; proceed to Phase 4
        selected_scraper = None

    scraper_success = False
    corrected_and_retried = False

    while True:
        if selected_scraper:
            log_event(f"Launching scraper: {selected_scraper} for Case: {case_number}...")
            
            # Run Scraper with Retries
            max_scraper_attempts = 3
            for attempt in range(1, max_scraper_attempts + 1):
                try:
                    log_event(f"Scraper attempt {attempt} of {max_scraper_attempts}...")
                    # Prepare scraper arguments
                    scraper_args = [sys.executable, selected_scraper, case_number] + extra_args
                    
                    # Logic: Orange County always headful. 
                    # Others: Headless by default unless --headful is passed to docket.py
                    if venue_county == "orange":
                        scraper_args.append("--headful")
                    else:
                        if is_headful:
                            scraper_args.append("--headful")
                        else:
                            scraper_args.append("--headless")

                    # Run Scraper
                    if use_interactive_mode:
                        # Interactive mode: Do not capture output
                        result = subprocess.run(scraper_args, check=False)
                        if result.returncode == 0:
                            scraper_success = True
                            break
                        else:
                            log_event(f"Scraper attempt {attempt} failed with exit code {result.returncode}.", level="warning")
                    else:
                        # Standard mode: Capture output
                        result = subprocess.run(scraper_args, capture_output=True, text=True)
                        
                        if result.returncode == 0:
                            log_event(f"Scraper {selected_scraper} completed successfully.")
                            scraper_success = True
                            break
                        else:
                            log_event(f"Scraper attempt {attempt} failed.\nSTDOUT: {result.stdout}\nSTDERR: {result.stderr}", level="warning")

                except Exception as e:
                    log_event(f"Failed to execute Docket Scraper on attempt {attempt}: {e}", level="error")
                
                if attempt < max_scraper_attempts:
                    log_event("Retrying scraper in 5 seconds...")
                    time.sleep(5)
            
            # NEW: Post-scraper check for "No Match" PDF content
            if scraper_success:
                # Find the candidate file locally before moving it
                file_pattern = f"docket_{case_number}_*.pdf"
                candidates = glob.glob(file_pattern)
                if not candidates:
                    candidates = glob.glob("docket_*.pdf")
                
                if candidates:
                    candidates.sort(key=os.path.getmtime, reverse=True)
                    check_file = candidates[0]
                    check_text = extract_text_from_file(check_file)
                    
                    no_match_keywords = ["no match found", "no matching cases", "case not found", "invalid case number", "does not match any", "nothing matches"]
                    if any(kw in check_text.lower() for kw in no_match_keywords):
                        log_event(f"DETECTION: Scraper saved a 'No Match' page for {case_number}.")
                        scraper_success = False # Invalidate this run
                        # Remove the useless PDF
                        try: os.remove(check_file)
                        except: pass
            
            # If failed (either exit code or "No Match"), try correction once from complaint
            if not scraper_success and not corrected_and_retried:
                log_event("Scraper failed or returned no results. Attempting to re-extract case number from complaint...")
                complaint_text = data_manager.get_value(file_num, "hidden_complaint_text")
                if not complaint_text:
                    case_path = get_case_path(file_num)
                    if case_path:
                        complaint_file = find_complaint_file(case_path)
                        if complaint_file:
                            complaint_text = extract_text_from_file(complaint_file)
                
                if complaint_text:
                    new_case_number = re_extract_case_number(file_num, complaint_text)
                    if new_case_number and new_case_number != case_number:
                        log_event(f"RECOVERY: Corrected Case Number found: {new_case_number} (Previous: {case_number}). Retrying scraper...")
                        case_number = new_case_number
                        data_manager.save_variable(file_num, "case_number", case_number)
                        corrected_and_retried = True
                        scraper_success = False # Reset for retry loop
                        continue # Loop back to the beginning of the 'while True' block
                    else:
                        log_event("RECOVERY: Re-extraction did not yield a different case number.")
                else:
                    log_event("RECOVERY: Could not retrieve complaint text for re-extraction.")
                
                corrected_and_retried = True # Prevent infinite loops even if continue wasn't hit
        
        break # Exit the while loop if success, or if correction failed/already tried

    # --- Phase 3: Extraction & Reporting ---
    docket_text = ""
    target_path = None
    
    if scraper_success:
        # Find the generated PDF (docket_<case_number>_*.pdf)
        date_str = datetime.datetime.now().strftime("%Y.%m.%d")
        file_pattern = f"docket_{case_number}_*.pdf"
        candidates = glob.glob(file_pattern)
        
        if not candidates:
            candidates = glob.glob("docket_*.pdf")
            
        if candidates:
            candidates.sort(key=os.path.getmtime, reverse=True)
            source_file = candidates[0]
            
            case_path = get_case_path(file_num)
            if case_path:
                target_dir = os.path.join(case_path, "NOTES", "AI OUTPUT")
                if not os.path.exists(target_dir):
                    os.makedirs(target_dir)
                
                new_filename = f"Docket_{date_str}.pdf"
                target_path = os.path.join(target_dir, new_filename)
                
                try:
                    shutil.copy2(source_file, target_path)
                    os.remove(source_file)
                    log_event(f"COMPLETED: Docket downloaded and saved to {target_path}")
                    
                    # Log download date to Master Database
                    try:
                        db = MasterCaseDatabase()
                        db.update_last_docket_download(file_num, datetime.datetime.now().strftime("%Y-%m-%d"))
                    except Exception as db_e:
                        log_event(f"Failed to update Master Database with download date: {db_e}", level="warning")
                    
                    log_event("--- Phase 3: Extracting Hearing Data ---")
                    docket_text = extract_text_from_file(target_path)
                    
                    if docket_text:
                        prompt = """
                        Analyze this docket text.
                        1. Identify the 'Trial Date' if explicitly scheduled. Format: YYYY-MM-DD (or "None" if not found).
                        2. Identify 'Other Hearings' (any future scheduled hearings other than the trial). Summarize them briefly (e.g., "CMC on 2025-01-15, Motion on 2025-02-20"). If none, return "None".
                        
                        Return JSON:
                        {
                            "trial_date": "string",
                            "other_hearings": "string"
                        }
                        """
                        
                        json_str = call_gemini_api(prompt, docket_text)
                        if json_str:
                            try:
                                extracted = json.loads(clean_json_string(json_str))
                                trial_date = extracted.get("trial_date", "None")
                                other_hearings = extracted.get("other_hearings", "None")
                                
                                log_event(f"Extracted Trial Date: {trial_date}")
                                log_event(f"Extracted Other Hearings: {other_hearings}")
                                
                                data_manager.save_variable(file_num, "trial_date", trial_date)
                                data_manager.save_variable(file_num, "other_hearings", other_hearings)
                                
                                # Prepare data for Master Status
                                master_data = {
                                    "file_number": file_num,
                                    "plaintiffs": data_manager.get_value(file_num, "plaintiffs"),
                                    "case_number": data_manager.get_value(file_num, "case_number"),
                                    "venue_county": data_manager.get_value(file_num, "venue_county"),
                                    "trial_date": trial_date,
                                    "other_hearings": other_hearings
                                }
                                
                                log_event("Updating Master Case Status Document...")
                                update_master_status(file_num, master_data, target_path)
                                
                            except json.JSONDecodeError:
                                log_event("Failed to parse Gemini response for hearings.", level="error")
                    else:
                        log_event("Failed to extract text from saved Docket PDF.", level="error")
                except Exception as e:
                    log_event(f"Error processing downloaded docket: {e}", level="error")
    else:
        log_event("Scraper failed or skipped. Proceeding to Procedural History with available information.", level="warning")

    # --- Phase 4: Procedural History (Always attempt) ---
    log_event("--- Phase 4: Generating Procedural History ---")
    case_path = get_case_path(file_num)
    
    if case_path:
        # Try to get cached text from Complaint Agent first
        complaint_text = data_manager.get_value(file_num, "hidden_complaint_text")
        
        if complaint_text:
            log_event("Using cached complaint text from hidden variable.")
        else:
            log_event("Cached complaint text not found. Attempting to extract from file.")
            complaint_file = find_complaint_file(case_path)
            if complaint_file:
                log_event(f"Found Complaint for Procedural History: {complaint_file}")
                complaint_text = extract_text_from_file(complaint_file)
            else:
                log_event("Complaint file not found for Procedural History context.", level="warning")

        if os.path.exists(PROCEDURAL_HIST_PROMPT_FILE):
            with open(PROCEDURAL_HIST_PROMPT_FILE, 'r', encoding='utf-8') as f:
                ph_prompt = f.read()
            
            # If scraper failed, docket_text is empty, which fulfills the user request 
            # (future hearings section will be blank or based only on complaint if mentioned there)
            combined_context = f"[CASE DOCKET TEXT]\n{docket_text}\n\n[COMPLAINT TEXT]\n{complaint_text}"
            
            log_event("Calling AI for Procedural History (Context provided: Docket + Complaint)...")
            ph_result = call_gemini_api(ph_prompt, combined_context, ["gemini-3-flash-preview", "gemini-2.5-flash"])
            
            if ph_result:
                data_manager.save_variable(file_num, "procedural_history", ph_result)
                log_event("Generated Procedural History.")
            else:
                log_event("AI failed to generate Procedural History.", level="warning")
        else:
            log_event(f"Prompt file missing: {PROCEDURAL_HIST_PROMPT_FILE}", level="warning")

        # Final update to variables.docx with all accumulated data
        log_event("Updating variables.docx with all case data...")
        try:
            all_vars = data_manager.get_all_variables(file_num, flatten=True)
            update_variables_docx(case_path, all_vars)
            log_event("Successfully updated variables.docx")
        except Exception as e:
            log_event(f"Error updating variables.docx: {e}", level="error")
    else:
        log_event("Could not resolve case path for Phase 4.", level="error")
        # If we can't find the case path, we really can't do anything useful, so this is a true error.
        sys.exit(1)

    # Note: We exit with 0 (Success) even if scraper failed, provided Phase 4 attempted to run.
    # The logs will indicate the partial failure of the download.
    
    log_event("--- Agent Finished Successfully ---")
    sys.exit(0)


if __name__ == "__main__":
    main()
