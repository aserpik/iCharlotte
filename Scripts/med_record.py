import os
import sys
import datetime
import re
import subprocess
import gc
import time
from google import genai
from docx import Document
from docx.shared import Pt, Inches
import concurrent.futures
import io

# Add parent directory to path for imports
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

# Import crash handler and logger FIRST
try:
    from icharlotte_core.crash_handler import install_crash_handler, checkpoint, MemoryGuard
    from icharlotte_core.agent_logger import AgentLogger, get_logger
    CRASH_HANDLER_AVAILABLE = True
except ImportError:
    CRASH_HANDLER_AVAILABLE = False
    # Fallback implementations
    def install_crash_handler(*args, **kwargs):
        return None
    def checkpoint(*args, **kwargs):
        pass
    class MemoryGuard:
        def __init__(self, *args, **kwargs):
            pass
        def __enter__(self):
            return self
        def __exit__(self, *args):
            pass
        def check(self, force=False):
            return True

# Import process-safe docx writer
try:
    from icharlotte_core.docx_writer import get_docx_lock, LockTimeoutError
except ImportError:
    # Fallback: no locking if module not available
    from contextlib import contextmanager
    @contextmanager
    def get_docx_lock(path, timeout=120):
        yield
    class LockTimeoutError(Exception):
        pass

# --- Dependency Setup ---
try:
    import fitz # PyMuPDF
except ImportError:
    print("ERROR:pymupdf is required for this optimized version.")
    print("Please run: pip install pymupdf")
    sys.exit(1)

try:
    from pypdf import PdfReader
except ImportError:
    try:
        from PyPDF2 import PdfReader
    except ImportError:
        pass # Not strictly needed if fitz works

OCR_AVAILABLE = False
try:
    import pytesseract
    from PIL import Image
    OCR_AVAILABLE = True
except ImportError:
    pass

# --- Path Configuration ---
SCRIPTS_DIR = os.path.dirname(os.path.abspath(__file__))
PROMPT_FILE = os.path.join(SCRIPTS_DIR, "MED_RECORD_PROMPT.txt")

# --- Global Logger ---
_logger = None

def get_agent_logger(file_number: str = None) -> 'AgentLogger':
    """Get or create the agent logger."""
    global _logger
    if _logger is None:
        try:
            _logger = AgentLogger("MedRecord", file_number=file_number)
        except:
            _logger = None
    return _logger

def safe_print(message, file=None, flush=True):
    """
    Print wrapper that handles OSError when stdout/stderr becomes invalid.
    This can happen when running multiple agents simultaneously on Windows.
    """
    try:
        print(message, file=file, flush=flush)
    except OSError:
        # stdout/stderr pipe is broken - silently ignore
        # The AgentLogger will still capture output to log files
        pass

def log_event(message, level="info", progress=None):
    """
    Log an event with structured output for UI parsing.

    Args:
        message: The message to log
        level: Log level (info, error, warning)
        progress: Optional progress percentage (0-100)
    """
    logger = get_agent_logger()

    # Output structured progress first if provided (for UI)
    if progress is not None:
        safe_print(f"PROGRESS:{progress}:{message}")

    # Output human-readable message
    safe_print(f"[MedRecord] {message}")

    # Log to file via AgentLogger
    if logger:
        if level == "error":
            logger.error(message)
        elif level == "warning":
            logger.warning(message)
        else:
            if progress is not None:
                logger.info(message, progress=progress)
            else:
                logger.info(message)

def get_page_text_fast(pdf_path, page_num):
    """
    Extracts text from a single page using PyMuPDF (fitz).
    If text is sparse, it renders the page image and uses OCR (Tesseract).
    page_num is 1-based.
    """
    try:
        # Open PDF (thread-safe if opened locally)
        doc = fitz.open(pdf_path)
        page = doc.load_page(page_num - 1)
        
        # 1. Try direct text extraction (VERY FAST)
        text = page.get_text("text")
        
        # 2. If sparse, fall back to OCR
        if len(text.strip()) < 50 and OCR_AVAILABLE:
            # Render page to image (pixmap)
            # matrix=2.0 (144 dpi) -> 300 dpi approx
            mat = fitz.Matrix(2, 2) 
            pix = page.get_pixmap(matrix=mat)
            
            # Convert to PIL Image
            img_data = pix.tobytes("png")
            img = Image.open(io.BytesIO(img_data))
            
            # Run OCR
            ocr_text = pytesseract.image_to_string(img)
            
            if len(ocr_text.strip()) > len(text.strip()):
                text = ocr_text

        doc.close()
        return (page_num, text)
        
    except Exception as e:
        log_event(f"Error processing page {page_num}: {e}", level="warning")
        return (page_num, "")

def extract_text_parallel(file_path):
    """Extracts text from PDF using parallel processing with memory monitoring."""
    log_event(f"Extracting text from: {file_path} using parallel processing...", progress=5)
    checkpoint("Starting PDF extraction")

    try:
        doc = fitz.open(file_path)
        num_pages = len(doc)
        doc.close()
    except Exception as e:
        log_event(f"Error reading PDF structure: {e}", level="error")
        safe_print(f"PASS_FAILED:Extraction:PDF read error - {e}:recoverable")
        return None

    log_event(f"Total pages: {num_pages}", progress=10)
    checkpoint(f"PDF has {num_pages} pages")

    # Use ThreadPoolExecutor with memory guard
    max_workers = os.cpu_count() or 4
    if max_workers > 16: max_workers = 16

    # Reduce workers for large PDFs to avoid memory issues
    if num_pages > 200:
        max_workers = min(max_workers, 4)
        log_event(f"Large PDF detected, reducing workers to {max_workers}")

    page_texts = {}
    memory_guard = MemoryGuard(
        warn_threshold_mb=1500,
        abort_threshold_mb=2500,
        check_interval_items=20,
        logger=lambda msg: log_event(msg, level="warning")
    )

    with memory_guard:
        with concurrent.futures.ThreadPoolExecutor(max_workers=max_workers) as executor:
            future_to_page = {
                executor.submit(get_page_text_fast, file_path, i + 1): i + 1
                for i in range(num_pages)
            }

            processed_count = 0
            for future in concurrent.futures.as_completed(future_to_page):
                try:
                    p_num, text = future.result()
                    page_texts[p_num] = text
                except Exception as exc:
                    log_event(f"Page generated an exception: {exc}", level="error")

                processed_count += 1

                # Progress updates
                if processed_count % 25 == 0 or processed_count == num_pages:
                    pct = 10 + int((processed_count / num_pages) * 30)  # 10-40%
                    log_event(f"Extracted {processed_count}/{num_pages} pages", progress=pct)
                    checkpoint(f"Extracted page {processed_count}/{num_pages}", pct)

                # Memory check
                try:
                    memory_guard.check()
                except MemoryError as e:
                    log_event(f"Memory limit exceeded: {e}", level="error")
                    safe_print("PASS_FAILED:Extraction:Memory limit exceeded:fatal")
                    # Force garbage collection and continue with what we have
                    gc.collect()
                    break

                # Periodic garbage collection for large documents
                if processed_count % 100 == 0:
                    gc.collect()

    # Reassemble in order
    full_text = ""
    for i in range(1, num_pages + 1):
        if i in page_texts:
            full_text += page_texts[i] + "\n"

    log_event(f"Extraction complete: {len(full_text)} characters", progress=40)
    checkpoint("Extraction complete")

    return full_text

def extract_text(file_path):
    """Wrapper to handle different file types."""
    ext = os.path.splitext(file_path)[1].lower()
    
    if ext == ".pdf":
        return extract_text_parallel(file_path)
        
    elif ext == ".docx":
        try:
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            log_event(f"Error reading DOCX: {e}", level="error")
            return None
            
    else:
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
        except Exception as e:
            log_event(f"Error reading text file: {e}", level="error")
            return None

def chunk_text(text, chunk_size=500000):
    """Splits text into chunks of approximately chunk_size characters."""
    chunks = []
    for i in range(0, len(text), chunk_size):
        chunks.append(text[i:i + chunk_size])
    return chunks

def clean_response(text):
    """
    Extracts content between <rewritten_chronology> tags.
    If tags are missing but a scratchpad is present, it attempts to take everything after the scratchpad.
    """
    # 1. Try to find content within <rewritten_chronology> tags
    chron_pattern = re.compile(r'<rewritten_chronology>(.*?)(?:</rewritten_chronology>|$)', re.DOTALL | re.IGNORECASE)
    chron_match = chron_pattern.search(text)
    if chron_match:
        content = chron_match.group(1).strip()
        if content:
            return content

    # 2. Fallback: If <scratchpad> is present, take everything after </scratchpad>
    if "<scratchpad>" in text.lower() or "</scratchpad>" in text.lower():
        scratch_end_pattern = re.compile(r'</scratchpad>(.*)', re.DOTALL | re.IGNORECASE)
        scratch_match = scratch_end_pattern.search(text)
        if scratch_match:
            return scratch_match.group(1).strip()
            
    # 3. Last resort: Return the text but strip common AI conversational markers if they exist
    # (Though usually the above two cover 99% of cases with the current prompt)
    return text.strip()

def call_gemini(prompt, text):
    """Calls Gemini API, handling large text via chunking with robust error handling."""
    checkpoint("Starting LLM processing")

    api_key = os.environ.get("GEMINI_API_KEY")
    if not api_key:
        log_event("Error: GEMINI_API_KEY environment variable not set.", level="error")
        safe_print("PASS_FAILED:LLM:API key not configured:fatal")
        return None

    try:
        client = genai.Client(api_key=api_key)
    except Exception as e:
        log_event(f"Error initializing Gemini client: {e}", level="error")
        safe_print(f"PASS_FAILED:LLM:Client initialization failed - {e}:fatal")
        return None

    total_chars = len(text)
    log_event(f"Total document length: {total_chars} characters.", progress=45)

    # Chunk if larger than ~800k characters (safety margin for 1M token limit models)
    chunks = chunk_text(text, chunk_size=500000)
    full_response = ""
    total_chunks = len(chunks)

    model_sequence = [
        "gemini-3-flash-preview",
        "gemini-2.5-flash"
    ]

    failed_chunks = 0
    for i, chunk in enumerate(chunks):
        chunk_num = i + 1
        pct = 45 + int((chunk_num / total_chunks) * 35)  # 45-80%
        log_event(f"Processing chunk {chunk_num}/{total_chunks} ({len(chunk)} chars)...", progress=pct)
        checkpoint(f"LLM processing chunk {chunk_num}/{total_chunks}", pct)

        # Add context for chunks
        chunk_prompt = prompt
        if total_chunks > 1:
            chunk_prompt += f"\n\n(Note: This is part {chunk_num} of {total_chunks} of the document. Summarize/extract relevant information from this section.)"

        full_prompt = f"{chunk_prompt}\n\nDOCUMENT CONTENT:\n{chunk}"

        chunk_success = False
        last_error = None

        for model_name in model_sequence:
            log_event(f"Attempting model: {model_name}")
            checkpoint(f"Trying model {model_name}")

            # Retry logic per model
            max_retries = 3
            for attempt in range(max_retries):
                try:
                    response = client.models.generate_content(
                        model=model_name,
                        contents=full_prompt
                    )
                    if response and response.text:
                        log_event(f"Success with model: {model_name}")
                        cleaned = clean_response(response.text)
                        full_response += cleaned + "\n\n"
                        chunk_success = True
                        break  # Exit retry loop
                except Exception as e:
                    last_error = e
                    error_str = str(e)
                    log_event(f"Attempt {attempt+1}/{max_retries} failed with {model_name}: {e}", level="warning")

                    # Rate limit handling
                    if "429" in error_str or "rate" in error_str.lower():
                        wait_time = 30 * (attempt + 1)  # Exponential backoff
                        log_event(f"Rate limit hit. Waiting {wait_time} seconds...")
                        checkpoint(f"Rate limited, waiting {wait_time}s")
                        time.sleep(wait_time)
                        continue

                    # Quota exceeded - try next model
                    if "quota" in error_str.lower():
                        log_event(f"Quota exceeded for {model_name}, trying next model", level="warning")
                        break

                    # Other errors - brief pause and retry
                    if attempt < max_retries - 1:
                        time.sleep(5)

            if chunk_success:
                break  # Exit model loop

        if not chunk_success:
            failed_chunks += 1
            error_msg = str(last_error) if last_error else "Unknown error"
            log_event(f"Error: Failed to process chunk {chunk_num}. Error: {error_msg}", level="error")
            full_response += f"\n[Error: Could not process part {chunk_num} of the document - {error_msg}]\n"

    # Log summary
    if failed_chunks > 0:
        log_event(f"LLM processing completed with {failed_chunks}/{total_chunks} failed chunks", level="warning")
        if failed_chunks == total_chunks:
            safe_print("PASS_FAILED:LLM:All chunks failed:fatal")
            return None
    else:
        log_event("LLM processing completed successfully", progress=80)

    checkpoint("LLM processing complete")
    return full_response.strip()

def add_markdown_to_doc(doc, content):
    """Parses basic Markdown and applies formatting."""
    lines = content.split('\n')
    active_paragraph = None
    
    # Regex to identify dates at the beginning of a paragraph/sentence
    # Matches: (Optional "On ") (Date String)
    # Date String: Month DD, YYYY
    date_pattern = re.compile(r"^(On )?((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]* \d{1,2},? \d{4})")

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        
        if stripped.startswith('#'):
            text = stripped.lstrip('#').strip()
            if not text.endswith('.'):
                text += "."
            active_paragraph = doc.add_paragraph()
            run = active_paragraph.add_run(text + " ")
            run.bold = True
            continue
        
        if stripped.startswith('* ') or stripped.startswith('- '):
            text = stripped[2:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.first_line_indent = Inches(-0.25)
            
            # Use a real bullet character and tab for portability
            run = p.add_run("•\t")
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            
            # Support bold parsing within list items
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    r = p.add_run(part[2:-2])
                    r.bold = True
                else:
                    r = p.add_run(part)
                r.font.name = 'Times New Roman'
                r.font.size = Pt(12)
            active_paragraph = None
            continue
        
        # Regular paragraph
        if active_paragraph:
            p = active_paragraph
        else:
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Inches(0.5)
        
        # Check for date at start of line
        match = date_pattern.match(stripped)
        if match:
            on_prefix = match.group(1) # "On " or None
            date_str = match.group(2)  # "January 1, 2024"
            
            total_match_len = len(match.group(0))
            remaining_text = stripped[total_match_len:]
            
            if on_prefix:
                p.add_run(on_prefix)
            
            run = p.add_run(date_str)
            run.underline = True
            
            # Process remaining text for bold markers
            parts = re.split(r'(\**.*?\**)', remaining_text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
        else:
            # Standard markdown parsing
            parts = re.split(r'(\**.*?\**)', stripped)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
        
        active_paragraph = None

def extract_provider_from_filename(filename):
    """Extracts provider name from filename based on patterns."""
    # Pattern 1: 12345-001_ PROVIDER NAME (1).pdf
    # Split by underscore
    parts = filename.split('_')
    if len(parts) > 1:
        # Take the part after the first underscore
        potential_name = parts[1]
        # Remove extension
        potential_name = os.path.splitext(potential_name)[0]
        # Remove trailing parentheses like (1)
        potential_name = re.sub(r'\s*\(\d+\)$', '', potential_name)
        return potential_name.strip()
    
    # Fallback: Use filename without extension and sanitize
    name = os.path.splitext(filename)[0]
    return name.strip()

def sanitize_filename(name):
    """Sanitizes a string to be safe for filenames."""
    # Remove invalid characters
    name = re.sub(r'[<>:"/\\|?*]', '', name)
    # Trim whitespace
    name = name.strip()
    return name

def save_to_docx(content, output_dir, provider_name, original_filename):
    """Saves to DOCX with dynamic filename and process-safe locking."""
    checkpoint("Starting DOCX save")

    safe_provider = sanitize_filename(provider_name)
    if not safe_provider:
        safe_provider = "Unknown_Provider"

    # Filename format: Med_Record_[name of treatment provider].docx
    filename = f"Med_Record_{safe_provider}.docx"
    output_path = os.path.join(output_dir, filename)

    log_event(f"Saving to: {output_path}", progress=85)

    try:
        # Acquire process-level lock before accessing the file
        with get_docx_lock(output_path):
            log_event(f"Acquired write lock for: {filename}")
            checkpoint("Lock acquired")

            if os.path.exists(output_path):
                 # If exists, we append
                 doc = Document(output_path)
                 doc.add_page_break()
                 log_event(f"Appending to existing file: {output_path}")
            else:
                 doc = Document()
                 log_event(f"Creating new file: {output_path}")

            # Styles
            style = doc.styles['Normal']
            style.font.name = 'Times New Roman'
            style.font.size = Pt(12)
            style.paragraph_format.line_spacing = 1.0
            style.paragraph_format.space_after = Pt(12)

            for i in range(1, 10):
                if f'Heading {i}' in doc.styles:
                    h = doc.styles[f'Heading {i}']
                    h.font.name = 'Times New Roman'
                    h.font.size = Pt(12)
                    h.paragraph_format.line_spacing = 1.0
                    h.paragraph_format.space_after = Pt(12)

            for s in ['List Bullet', 'List Number']:
                if s in doc.styles:
                    l = doc.styles[s]
                    l.font.name = 'Times New Roman'
                    l.font.size = Pt(12)
                    l.paragraph_format.line_spacing = 1.0
                    l.paragraph_format.space_after = Pt(12)

            # Title
            p = doc.add_paragraph()
            run = p.add_run(f"Medical Record Chronology - {provider_name}")
            run.bold = True
            run.underline = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

            doc.add_paragraph(f"Source: {original_filename}")
            doc.add_paragraph(f"Generated on: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

            checkpoint("Adding content to document")
            add_markdown_to_doc(doc, content)

            checkpoint("Saving document")
            doc.save(output_path)

            # Emit structured output file path for UI
            safe_print(f"OUTPUT_FILE:{output_path}")
            log_event(f"Saved to: {output_path}", progress=95)
            checkpoint("Save complete")
            return True

    except LockTimeoutError as e:
        log_event(f"Lock timeout waiting for file: {e}", level="error")
        checkpoint(f"Lock timeout: {e}")
        return False
    except PermissionError as e:
        log_event(f"Permission denied saving file: {e}", level="error")
        checkpoint(f"Permission error: {e}")
        return False
    except Exception as e:
        log_event(f"Error saving DOCX: {e}", level="error")
        checkpoint(f"Save error: {e}")
        return False

def main():
    """Main entry point with crash handling and structured logging."""
    # Parse arguments first for crash context
    raw_args = sys.argv[1:] if len(sys.argv) > 1 else []

    if not raw_args:
        log_event("Error: No file path provided.", level="error")
        print("ERROR:No file path provided")
        sys.exit(1)

    # Heuristic: Check if arguments form a single split path
    combined_arg = " ".join(raw_args)
    clean_combined = combined_arg.strip().strip('"').strip("'")

    if os.path.exists(clean_combined):
        file_paths = [clean_combined]
    else:
        file_paths = raw_args

    # Determine input path for crash handler context
    input_path = file_paths[0].strip().strip('"').strip("'")
    input_path = os.path.abspath(input_path)

    # Install crash handler
    crash_handler = install_crash_handler(
        "MedRecord",
        input_file=input_path,
        file_count=len(file_paths)
    )
    checkpoint("Agent initialized")

    # --- Dispatcher Mode ---
    if len(file_paths) > 1:
        log_event(f"Detected multiple file paths: {len(file_paths)} items. Launching separate agents...")
        spawned = 0
        failed = 0

        for path in file_paths:
            log_event(f"Spawning agent for: {path}")
            try:
                # Spawn with output capture to log file
                if os.name == 'nt':
                    # CREATE_NO_WINDOW (0x08000000) for headless execution
                    proc = subprocess.Popen(
                        [sys.executable, sys.argv[0], path],
                        creationflags=0x08000000,
                        stdout=subprocess.PIPE,
                        stderr=subprocess.PIPE
                    )
                else:
                    proc = subprocess.Popen(
                        [sys.executable, sys.argv[0], path],
                        stdout=subprocess.PIPE,
                        stderr=subprocess.PIPE
                    )
                spawned += 1
                log_event(f"Spawned process PID {proc.pid} for: {os.path.basename(path)}")
            except Exception as e:
                failed += 1
                log_event(f"Failed to spawn agent for {path}: {e}", level="error")

        log_event(f"Launched {spawned} medical record agents ({failed} failed)")
        print(f"PROGRESS:100:Launched {spawned} agents")
        sys.exit(0 if failed == 0 else 1)

    # --- Worker Mode ---
    if not os.path.exists(input_path):
        log_event(f"Error: File not found: {input_path}", level="error")
        print(f"ERROR:File not found: {input_path}")
        sys.exit(1)

    log_event(f"--- Starting Med Record Agent for: {input_path} ---", progress=0)
    checkpoint("Starting worker mode")

    # Handle directory input
    if os.path.isdir(input_path):
        log_event(f"Input is a directory. Scanning: {input_path}")
        files_to_process = []
        for root, _, files in os.walk(input_path):
            for file in files:
                if file.lower().endswith(('.pdf', '.docx')):
                    if "Med_Record_" in file:
                        continue
                    files_to_process.append(os.path.join(root, file))

        if not files_to_process:
            log_event("No suitable files found.")
            sys.exit(0)

        total_files = len(files_to_process)
        for idx, file_path in enumerate(files_to_process):
            log_event(f"Processing file {idx+1}/{total_files}: {os.path.basename(file_path)}")
            try:
                result = subprocess.run(
                    [sys.executable, sys.argv[0], file_path],
                    capture_output=True,
                    text=True,
                    timeout=600  # 10 minute timeout per file
                )
                if result.returncode != 0:
                    log_event(f"Subprocess failed for {file_path}", level="error")
                    if result.stderr:
                        log_event(f"Error output: {result.stderr[:500]}", level="error")
            except subprocess.TimeoutExpired:
                log_event(f"Timeout processing {file_path}", level="error")
            except Exception as e:
                log_event(f"Subprocess error for {file_path}: {e}", level="error")
        sys.exit(0)

    # --- Single File Processing ---
    total_passes = 3  # Extraction, LLM, Write

    # === Pass 1: Extraction ===
    safe_print(f"PASS_START:Extraction:1:{total_passes}")
    checkpoint("Starting extraction pass")
    extraction_start = time.time()

    text = extract_text(input_path)
    if not text:
        log_event("Extraction failed: No text extracted", level="error")
        safe_print("PASS_FAILED:Extraction:No text extracted:fatal")
        sys.exit(1)

    extraction_duration = time.time() - extraction_start
    safe_print(f"PASS_COMPLETE:Extraction:success:{extraction_duration:.1f}")
    log_event(f"Extraction complete ({extraction_duration:.1f}s, {len(text)} chars)")

    # Load prompt file
    try:
        with open(PROMPT_FILE, "r", encoding="utf-8") as f:
            prompt_instruction = f.read()
    except FileNotFoundError:
        log_event(f"Error: Prompt file not found: {PROMPT_FILE}", level="error")
        safe_print(f"ERROR:Prompt file not found: {PROMPT_FILE}")
        sys.exit(1)
    except Exception as e:
        log_event(f"Error reading prompt file: {e}", level="error")
        safe_print(f"ERROR:Failed to read prompt file: {e}")
        sys.exit(1)

    # Determine Provider Name
    filename = os.path.basename(input_path)
    provider_name = extract_provider_from_filename(filename)
    log_event(f"Identified Provider: {provider_name}")
    checkpoint(f"Provider: {provider_name}")

    # === Pass 2: LLM Processing ===
    safe_print(f"PASS_START:LLM:2:{total_passes}")
    checkpoint("Starting LLM pass")
    llm_start = time.time()

    final_content = call_gemini(prompt_instruction, text)
    if not final_content:
        log_event("LLM processing failed: No content generated", level="error")
        safe_print("PASS_FAILED:LLM:No content generated:fatal")
        sys.exit(1)

    llm_duration = time.time() - llm_start
    safe_print(f"PASS_COMPLETE:LLM:success:{llm_duration:.1f}")
    log_event(f"LLM processing complete ({llm_duration:.1f}s)")

    # === Pass 3: Write Output ===
    safe_print(f"PASS_START:Write:3:{total_passes}")
    checkpoint("Starting write pass")
    write_start = time.time()

    # Determine Output Directory
    parts = input_path.split(os.sep)
    output_dir = None
    case_root_parts = None

    # Priority 1: Find folder starting with exactly 3 digits (Case Folder)
    for i in range(len(parts) - 1, -1, -1):
        if re.match(r'^\d{3}(\D|$)', parts[i]):
            log_event(f"Identified Case Folder: {parts[i]}")
            case_root_parts = parts[:i+1]
            break

    # Priority 2: Standard "Current Clients" structure
    if not case_root_parts:
        for i, part in enumerate(parts):
            if part.lower() == "current clients":
                if i + 2 < len(parts):
                    case_root_parts = parts[:i+3]
                break

    if case_root_parts:
        output_dir = os.sep.join(case_root_parts + ["NOTES", "AI OUTPUT"])

    if not output_dir:
        for i in range(len(parts) - 1, -1, -1):
            if parts[i].upper() == "NOTES":
                output_dir = os.path.join(os.sep.join(parts[:i+1]), "AI OUTPUT")
                break

    if not output_dir:
        input_dir = os.path.dirname(input_path)
        parent_dir = os.path.dirname(input_dir)
        output_dir = os.path.join(parent_dir, "NOTES", "AI OUTPUT")

    if not os.path.exists(output_dir):
        try:
            os.makedirs(output_dir)
            log_event(f"Created output directory: {output_dir}")
        except Exception as e:
            log_event(f"Warning: Could not create output directory: {e}", level="warning")
            # Fall back to input file directory
            output_dir = os.path.dirname(input_path)

    success = save_to_docx(final_content, output_dir, provider_name, filename)
    write_duration = time.time() - write_start

    if success:
        safe_print(f"PASS_COMPLETE:Write:success:{write_duration:.1f}")
        log_event(f"Write complete ({write_duration:.1f}s)", progress=100)
        safe_print("PROGRESS:100:Agent completed successfully")
        log_event("--- Agent Finished Successfully ---")
    else:
        safe_print("PASS_FAILED:Write:Save failed:recoverable")
        log_event("Write pass failed", level="error")
        sys.exit(1)

if __name__ == '__main__':
    try:
        main()
    except KeyboardInterrupt:
        safe_print("\nERROR:Agent cancelled by user")
        sys.exit(130)
    except MemoryError as e:
        safe_print(f"ERROR:Memory exhausted - {e}")
        safe_print("PASS_FAILED:Processing:Memory exhausted:fatal")
        log_event(f"FATAL: Memory exhausted: {e}", level="error")
        sys.exit(137)
    except Exception as e:
        # This catches any unhandled exception
        import traceback
        tb = traceback.format_exc()
        safe_print(f"ERROR:Unhandled exception - {type(e).__name__}: {e}")
        safe_print("PASS_FAILED:Processing:Unhandled exception:fatal")
        log_event(f"FATAL: Unhandled exception: {e}\n{tb}", level="error")

        # Also write to stderr for any capture mechanisms
        safe_print(f"\n{'='*60}", file=sys.stderr)
        safe_print("FATAL ERROR in MedRecord Agent", file=sys.stderr)
        safe_print(f"Exception: {type(e).__name__}: {e}", file=sys.stderr)
        safe_print(f"{'='*60}", file=sys.stderr)
        safe_print(tb, file=sys.stderr)

        sys.exit(1)