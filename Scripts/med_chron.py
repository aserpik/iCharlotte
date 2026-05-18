import os
import sys
import logging
import datetime
import re
import subprocess
from pathlib import Path
from concurrent.futures import ThreadPoolExecutor, as_completed
from google import genai
import gc
from docx import Document
from docx.shared import Pt, Inches
from pypdf import PdfReader
from dataclasses import dataclass
from icharlotte_core.llm_config import LLMCaller

# Import Case Data Manager
try:
    from case_data_manager import CaseDataManager
except ImportError:
    sys.path.append(os.path.join(os.getcwd(), 'Scripts'))
    from case_data_manager import CaseDataManager

try:
    import pytesseract
    from pdf2image import convert_from_path
    OCR_AVAILABLE = True
    
    # Windows-specific path configuration
    if os.name == 'nt':
        # Tesseract Path
        tesseract_path = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
        if os.path.exists(tesseract_path):
            pytesseract.pytesseract.tesseract_cmd = tesseract_path
        
        # Poppler Path
        poppler_path = r"C:\Program Files\poppler\Library\bin"
        if not os.path.exists(poppler_path):
             poppler_path = None
        
        POPPLER_PATH = poppler_path
    else:
        POPPLER_PATH = None

except ImportError:
    OCR_AVAILABLE = False
    POPPLER_PATH = None


# --- Configuration ---
LOG_FILE = os.path.join(os.getcwd(), "Med_Chron_activity.log")
PROMPT_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), "MED_CHRON_PROMPT.txt")

# Set up logging
logging.basicConfig(
    filename=LOG_FILE,
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    datefmt='%Y-%m-%d %H:%M:%S'
)

def log_event(message, level="info"):
    try:
        print(message)
    except UnicodeEncodeError:
        try:
            print(message.encode(sys.stdout.encoding or 'utf-8', errors='replace').decode(sys.stdout.encoding or 'utf-8'))
        except Exception:
             pass  # Silently ignore if stdout is broken
    except OSError:
        pass  # stdout pipe broken (common when running multiple agents)
    if level == "info":
        logging.info(message)
    elif level == "error":
        logging.error(message)
    elif level == "warning":
        logging.warning(message)

def extract_text(file_path):
    """Extracts text from PDF, DOCX, or plain text files."""
    log_event(f"Extracting text from: {file_path}")
    ext = os.path.splitext(file_path)[1].lower()
    text = ""

    try:
        if ext == ".pdf":
            reader = PdfReader(file_path)
            total_pages = len(reader.pages)
            log_event(f"PDF has {total_pages} pages.")
            
            def get_page_image(page_index):
                try:
                    images = convert_from_path(file_path, first_page=page_index+1, last_page=page_index+1, poppler_path=POPPLER_PATH)
                    return images[0] if images else None
                except Exception as e:
                    log_event(f"Error converting page {page_index+1} to image: {e}", level="warning")
                    return None

            for i, page in enumerate(reader.pages):
                # Explicit garbage collection every 10 pages to prevent memory buildup
                if i % 10 == 0:
                    gc.collect()

                page_text = page.extract_text() or ""
                
                if len(page_text.strip()) < 50:
                    log_event(f"Page {i+1} has insufficient text. Attempting OCR...", level="warning")
                    if OCR_AVAILABLE:
                        image = get_page_image(i)
                        if image:
                            try:
                                ocr_page_text = pytesseract.image_to_string(image)
                                if len(ocr_page_text.strip()) > len(page_text.strip()):
                                    text += ocr_page_text + "\n"
                                else:
                                    text += page_text + "\n"
                            except Exception:
                                text += page_text + "\n"
                        else:
                             text += page_text + "\n"
                    else:
                        text += page_text + "\n"
                else:
                    text += page_text + "\n"

        elif ext == ".docx":
            doc = Document(file_path)
            # Only iterate over paragraphs to ignore tables
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
        else:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
        
        if not text.strip():
            log_event(f"Warning: Extracted text is empty for {file_path}", level="warning")
            return None
        
        return text

    except Exception as e:
        log_event(f"Error extracting text from {file_path}: {e}", level="error")
        return None

def call_gemini(prompt, text):
    """Calls Gemini API."""
    api_key = os.environ.get("GEMINI_API_KEY")
    if not api_key:
        log_event("Error: GEMINI_API_KEY environment variable not set.", level="error")
        return None

    client = genai.Client(api_key=api_key)

    full_prompt = f"{prompt}\n\nDOCUMENT CONTENT:\n{text}"
    
    model_sequence = [
        "gemini-3-flash-preview", 
        "gemini-2.5-flash"
    ]

    for model_name in model_sequence:
        log_event(f"Attempting to use model: {model_name}")
        try:
            response = client.models.generate_content(
                model=model_name,
                contents=full_prompt
            )
            if response and response.text:
                log_event(f"Success with model: {model_name}")
                return response.text
        except Exception as e:
            log_event(f"Failed with model {model_name}: {e}", level="warning")
            continue
    
    log_event("Error: All model attempts failed.", level="error")
    return None

def add_markdown_to_doc(doc, content):
    """Parses basic Markdown and applies formatting."""
    lines = content.split('\n')
    active_paragraph = None
    
    # Regex to identify dates at the beginning of a paragraph/sentence
    # Matches: (Optional "On ") (Date String)
    # Date String: Month DD, YYYY
    date_pattern = re.compile(r"^(On )?((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]* \d{1,2},? \d{4})")

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        
        if stripped.startswith('#'):
            text = stripped.lstrip('#').strip()
            if not text.endswith('.'):
                text += "."
            active_paragraph = doc.add_paragraph()
            run = active_paragraph.add_run(text + " ")
            run.bold = True
            continue
        
        if stripped.startswith('* ') or stripped.startswith('- '):
            text = stripped[2:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.first_line_indent = Inches(-0.25)
            
            # Use a real bullet character and tab for portability
            run = p.add_run("•\t")
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            
            # Support bold parsing within list items
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    r = p.add_run(part[2:-2])
                    r.bold = True
                else:
                    r = p.add_run(part)
                r.font.name = 'Times New Roman'
                r.font.size = Pt(12)
            active_paragraph = None
            continue
        
        # Regular paragraph
        if active_paragraph:
            p = active_paragraph
        else:
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Inches(0.5)
        
        # Check for date at start of line
        match = date_pattern.match(stripped)
        if match:
            on_prefix = match.group(1) # "On " or None
            date_str = match.group(2)  # "January 1, 2024"
            
            total_match_len = len(match.group(0))
            remaining_text = stripped[total_match_len:]
            
            if on_prefix:
                p.add_run(on_prefix)
            
            run = p.add_run(date_str)
            run.underline = True
            
            # Process remaining text for bold markers
            parts = re.split(r'(\**.*?\**)', remaining_text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
        else:
            # Standard markdown parsing
            parts = re.split(r'(\**.*?\**)', stripped)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
        
        active_paragraph = None

def extract_provider_from_filename(filename):
    """Extracts provider name from filename based on patterns."""
    # Pattern 1: 12345-001_ PROVIDER NAME (1).pdf
    # Split by underscore
    parts = filename.split('_')
    if len(parts) > 1:
        # Take the part after the first underscore
        potential_name = parts[1]
        # Remove extension
        potential_name = os.path.splitext(potential_name)[0]
        # Remove trailing parentheses like (1)
        potential_name = re.sub(r'\s*\(\d+\)$', '', potential_name)
        return potential_name.strip()
    
    # Fallback: Use filename without extension and sanitize
    name = os.path.splitext(filename)[0]
    return name.strip()

def sanitize_filename(name):
    """Sanitizes a string to be safe for filenames."""
    # Remove invalid characters
    name = re.sub(r'[<>:"/\\|?*]', '', name)
    # Trim whitespace
    name = name.strip()
    return name

def save_to_docx(content, output_dir, provider_name, original_filename):
    """Saves to DOCX with unique filename med_chron_{filename}.docx."""
    
    # Use unique filename to avoid collision in parallel execution
    safe_name = re.sub(r"[^a-zA-Z0-9_\-]", "_", os.path.splitext(original_filename)[0])
    filename = f"med_chron_{safe_name}.docx"
    output_path = os.path.join(output_dir, filename)

    try:
        if os.path.exists(output_path):
             # If exists (re-run on same file?), we overwrite or append? 
             # Let's overwrite for a single file agent to avoid duplication on re-runs.
             doc = Document()
             log_event(f"Overwriting/Creating file: {output_path}")
        else:
             doc = Document()
             log_event(f"Creating new file: {output_path}")

        # Styles
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(12)
        style.paragraph_format.line_spacing = 1.0
        
        for i in range(1, 10):
            if f'Heading {i}' in doc.styles:
                h = doc.styles[f'Heading {i}']
                h.font.name = 'Times New Roman'
                h.font.size = Pt(12)
                h.paragraph_format.line_spacing = 1.0

        for s in ['List Bullet', 'List Number']:
            if s in doc.styles:
                l = doc.styles[s]
                l.font.name = 'Times New Roman'
                l.font.size = Pt(12)
                l.paragraph_format.line_spacing = 1.0

        # Title
        p = doc.add_paragraph()
        run = p.add_run(f"Medical Record Chronology - {provider_name}")
        run.bold = True
        run.underline = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

        doc.add_paragraph(f"Source: {original_filename}")
        doc.add_paragraph(f"Generated on: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        
        add_markdown_to_doc(doc, content)
            
        doc.save(output_path)
        log_event(f"Saved to: {output_path}")
        return True

    except Exception as e:
        log_event(f"Error saving DOCX: {e}", level="error")
        return False

def filter_content(text):
    """Filters text to only include content under specific headings."""
    headings = [
        "BRIEF SYNOPSIS OF PRE-INJURY MEDICAL RECORD:",
        "BRIEF SYNOPSIS OF POST-INJURY MEDICAL RECORD:"
    ]
    
    indices = []
    for h in headings:
        # Case insensitive search might be safer, but user used uppercase.
        # Strict adherence to user request first.
        idx = text.find(h)
        if idx != -1:
            indices.append((idx, h))
    
    # Sort by position in text
    indices.sort(key=lambda x: x[0])
    
    if not indices:
        log_event("Target headings not found. Processing skipped.", level="warning")
        return None

    filtered_chunks = []
    for i, (start_idx, header) in enumerate(indices):
        content_start = start_idx + len(header)
        
        # Determine end of this section
        if i + 1 < len(indices):
            content_end = indices[i+1][0]
        else:
            content_end = len(text)
            
        chunk = text[content_start:content_end].strip()
        filtered_chunks.append(f"{header}\n{chunk}")
        
    return "\n\n".join(filtered_chunks)

def _extract_full_text(file_path: str, prefetched_pdf_text: str | None = None) -> str:
    """Extract narrative + table text from a chronology file.

    .docx -> ``icharlotte_core.document_processor.extract_docx_text``
             (canonical extractor that includes tables as pipe-separated rows).
    .pdf  -> ``prefetched_pdf_text`` if provided (avoids double-OCR), else
             falls back to ``extract_text``. PDFs don't have a paragraphs-vs-
             tables split in extraction.
    .doc  -> Word COM read-only, never calls word.Quit().
    """
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".docx":
        from icharlotte_core.document_processor import extract_docx_text
        return extract_docx_text(file_path)
    if ext == ".pdf":
        if prefetched_pdf_text is not None:
            return prefetched_pdf_text
        return extract_text(file_path) or ""
    if ext == ".doc":
        return _extract_doc_via_word_com(file_path)
    return extract_text(file_path) or ""


def _extract_doc_via_word_com(file_path: str) -> str:
    """Read a legacy .doc by attaching to the user's running Word.

    Mirrors ChatTab._extract_doc_text: never set word.Visible and never
    call word.Quit() -- only close the Document we opened. Open ReadOnly
    so the user's session is untouched.
    """
    try:
        import win32com.client  # type: ignore
    except ImportError:
        log_event("win32com not available; cannot extract .doc files", level="warning")
        return ""
    word = None
    doc = None
    try:
        word = win32com.client.Dispatch("Word.Application")
        doc = word.Documents.Open(
            FileName=os.path.abspath(file_path),
            ReadOnly=True,
            AddToRecentFiles=False,
            ConfirmConversions=False,
        )
        return doc.Content.Text or ""
    except Exception as e:
        log_event(f".doc extraction failed for {file_path}: {e}", level="warning")
        return ""
    finally:
        if doc is not None:
            try:
                doc.Close(SaveChanges=False)
            except Exception:
                pass


def _build_catalog_snapshot() -> list:
    """Serialise the curated catalog into a JSON-friendly list."""
    scripts_dir = os.path.dirname(os.path.abspath(__file__))
    if scripts_dir not in sys.path:
        sys.path.insert(0, scripts_dir)
    from MED_CHRON_ANALYSES.catalog import CATALOG
    return [
        {
            "id": d.id,
            "title": d.title,
            "description": d.description,
            "uses_tables": d.uses_tables,
            "default_selected": d.default_selected,
        }
        for d in CATALOG
    ]


def process_prep(input_path: str, output_dir: str) -> int:
    """Phase 1: extract text twice, write session JSON, print AWAITING_INPUT.

    Returns process-style exit code (0 success, non-zero failure). Does
    NOT call sys.exit -- leaves that to main().
    """
    from icharlotte_core.med_chron import session_manager

    paths = session_manager.compute_session_paths(input_path, output_dir)

    # Cache reuse: if session.json + both text files exist, skip extraction.
    if (paths.session_path.exists()
            and paths.narrative_text_path.exists()
            and paths.full_text_path.exists()):
        log_event(f"Reusing cached prep at {paths.cache_dir}")
        print(f"AWAITING_INPUT:{paths.session_path}", flush=True)
        return 0

    paths.cache_dir.mkdir(parents=True, exist_ok=True)

    # --- Narrative-only text ---
    raw_text = extract_text(input_path)
    narrative_missing = False
    if not raw_text:
        # For docx files with no text content, use empty string rather than failing
        log_event(f"Could not extract text from {input_path}", level="warning")
        raw_text = ""

    narrative = filter_content(raw_text)
    if narrative is None:
        narrative_missing = True
        narrative = ""
    paths.narrative_text_path.write_text(narrative, encoding="utf-8")

    # --- Full text (narrative + tables) ---
    full_text = _extract_full_text(input_path, prefetched_pdf_text=raw_text)
    if not full_text:
        # Fall back to the raw_text we already have
        full_text = raw_text or ""
    paths.full_text_path.write_text(full_text, encoding="utf-8")

    # --- Session JSON ---
    filename = os.path.basename(input_path)
    provider_name = extract_provider_from_filename(filename)
    file_num_match = re.search(r"(\d{4}\.\d{3})", input_path)
    file_number = file_num_match.group(1) if file_num_match else None

    session_data = {
        "version": 1,
        "phase": "awaiting_input",
        "input_path": str(input_path),
        "narrative_text_path": str(paths.narrative_text_path),
        "full_text_path": str(paths.full_text_path),
        "narrative_missing": narrative_missing,
        "provider_name": provider_name,
        "file_number": file_number,
        "catalog": _build_catalog_snapshot(),
        "user_config": None,
    }
    session_manager.write_session(paths.session_path, session_data)

    log_event(
        f"Phase 1 complete: cached {len(narrative)} narrative chars "
        f"+ {len(full_text)} full chars; session at {paths.session_path}"
    )
    print(f"AWAITING_INPUT:{paths.session_path}", flush=True)
    return 0


# =============================================================================
# Phase 2: Run selected analyses
# =============================================================================


def _slug(value: str) -> str:
    """Lowercase + sanitize for use in filenames/run ids."""
    cleaned = re.sub(r"[^a-zA-Z0-9_\-]+", "_", value or "")
    return cleaned.strip("_").lower()


def _safe_basename(input_path: str) -> str:
    base = os.path.splitext(os.path.basename(input_path))[0]
    return _slug(base)


@dataclass
class RunSpec:
    id: str
    title: str
    prompt_text: str
    input_text: str
    output_path: str


@dataclass
class RunResult:
    spec: RunSpec
    success: bool
    error: str = ""
    output_chars: int = 0


def _build_run_list(session: dict, narrative: str, full: str,
                    safe_basename: str, output_dir: str) -> list[RunSpec]:
    """Translate user_config + catalog into concrete RunSpec instances."""
    scripts_dir = os.path.dirname(os.path.abspath(__file__))
    if scripts_dir not in sys.path:
        sys.path.insert(0, scripts_dir)
    from MED_CHRON_ANALYSES.catalog import CATALOG_BY_ID, load_prompt

    cfg = session["user_config"]
    runs: list[RunSpec] = []

    for cat_id in cfg.get("selected_catalog_ids", []):
        if cat_id not in CATALOG_BY_ID:
            log_event(f"Skipping unknown catalog id: {cat_id}", level="warning")
            continue
        d = CATALOG_BY_ID[cat_id]
        runs.append(RunSpec(
            id=cat_id,
            title=d.title,
            prompt_text=load_prompt(d.prompt_file),
            input_text=narrative if not d.uses_tables else full,
            output_path=os.path.join(
                output_dir, f"med_chron_{cat_id}_{safe_basename}.docx"
            ),
        ))

    wrapper = None
    for i, c in enumerate(cfg.get("custom_analyses", []), 1):
        if wrapper is None:
            wrapper = load_prompt("_custom_wrapper.txt")
        label_slug = _slug(c["label"])
        runs.append(RunSpec(
            id=f"custom_{i}_{label_slug}",
            title=c["label"],
            prompt_text=wrapper.replace("{user_instruction}", c["instruction"]),
            input_text=full,
            output_path=os.path.join(
                output_dir, f"med_chron_custom_{i}_{label_slug}_{safe_basename}.docx"
            ),
        ))

    return runs


def _drop_rewrite_if_narrative_missing(runs: list[RunSpec], narrative: str) -> list[RunSpec]:
    if narrative.strip():
        return runs
    kept = []
    for r in runs:
        if r.id == "rewrite_chronology":
            log_event(
                "Skipping Rewrite Chronology — no pre/post-injury synopsis "
                "headings found in this document.",
                level="warning",
            )
            continue
        kept.append(r)
    return kept


def _run_one_analysis(spec: RunSpec, llm_caller: LLMCaller,
                       provider_name: str) -> RunResult:
    """Execute a single analysis. Caller MUST NOT let exceptions escape."""
    try:
        log_event(f"[{spec.id}] starting LLM call ({len(spec.input_text)} chars)")
        result = llm_caller.call(
            prompt=spec.prompt_text,
            text=spec.input_text,
            task_type="summary",
        )
        if not result:
            return RunResult(spec=spec, success=False, error="LLM returned empty result")

        os.makedirs(os.path.dirname(spec.output_path), exist_ok=True)
        save_to_docx_at_path(result, spec.output_path, provider_name, spec.title)
        log_event(f"[{spec.id}] done: {len(result)} chars → {spec.output_path}")
        return RunResult(spec=spec, success=True, output_chars=len(result))
    except Exception as e:
        log_event(f"[{spec.id}] failed: {e}", level="error")
        return RunResult(spec=spec, success=False, error=str(e))


def save_to_docx_at_path(content: str, output_path: str,
                         provider_name: str, analysis_title: str) -> None:
    """Write content to output_path with the existing Med-Cron styling.

    If the destination is locked (e.g., open in Word), auto-version up to
    10 attempts: ``out.docx`` -> ``out v.2.docx`` -> ``out v.3.docx``.
    """
    from docx import Document
    from docx.shared import Pt

    base, ext = os.path.splitext(output_path)
    attempt = 1
    last_err = None
    while attempt <= 10:
        candidate = output_path if attempt == 1 else f"{base} v.{attempt}{ext}"
        try:
            doc = Document()
            style = doc.styles['Normal']
            style.font.name = 'Times New Roman'
            style.font.size = Pt(12)
            style.paragraph_format.line_spacing = 1.0

            p = doc.add_paragraph()
            run = p.add_run(f"{analysis_title} — {provider_name}")
            run.bold = True
            run.underline = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

            doc.add_paragraph(f"Generated on: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            add_markdown_to_doc(doc, content)
            doc.save(candidate)
            return
        except (PermissionError, IOError) as e:
            last_err = e
            attempt += 1
    raise RuntimeError(f"Could not save after {attempt - 1} attempts: {last_err}")


def process_run(session_path: str, output_dir: str) -> int:
    """Phase 2: load session, fan out analyses in parallel, write docx each.

    Returns 0 if at least one analysis succeeded, 1 if all failed or the
    session is malformed.
    """
    from icharlotte_core.med_chron import session_manager

    try:
        session = session_manager.read_session(session_path)
    except Exception as e:
        log_event(f"Could not load session at {session_path}: {e}", level="error")
        return 1

    if session.get("phase") != "ready_to_run":
        log_event(
            f"Session phase is {session.get('phase')!r}; expected ready_to_run",
            level="error",
        )
        return 1

    try:
        narrative = Path(session["narrative_text_path"]).read_text(encoding="utf-8")
        full = Path(session["full_text_path"]).read_text(encoding="utf-8")
        safe_basename = _safe_basename(session["input_path"])
    except (KeyError, FileNotFoundError, OSError) as e:
        log_event(f"Session is corrupt or cache files missing: {e}", level="error")
        return 1

    runs = _build_run_list(session, narrative, full, safe_basename, output_dir)
    runs = _drop_rewrite_if_narrative_missing(runs, narrative)

    if not runs:
        log_event("No runs scheduled (after skip rules). Nothing to do.", level="warning")
        return 1

    llm_caller = LLMCaller()
    provider_name = session.get("provider_name") or "Unknown Provider"

    successes = 0
    failures = 0
    total = len(runs)

    log_event(f"Starting {total} analyses (max 4 concurrent)")
    with ThreadPoolExecutor(max_workers=min(total, 4)) as ex:
        futures = {
            ex.submit(_run_one_analysis, r, llm_caller, provider_name): r
            for r in runs
        }
        done = 0
        for f in as_completed(futures):
            result = f.result()
            done += 1
            if result.success:
                successes += 1
            else:
                failures += 1
            pct = int(20 + (done * 70 / total))
            print(f"PROGRESS:{pct}:{done}/{total} done ({failures} failed)", flush=True)

    log_event(f"Phase 2 complete: {successes}/{total} succeeded, {failures} failed")
    print(f"PROGRESS:100:{successes}/{total} analyses complete ({failures} failed)", flush=True)
    return 0 if successes > 0 else 1


def main():
    if len(sys.argv) < 2:
        log_event("Error: No file path provided.", level="error")
        sys.exit(1)

    # Attempt to handle unquoted paths with spaces by joining all arguments
    input_path = " ".join(sys.argv[1:])
    
    # If the joined path doesn't exist, but the first argument does, 
    # it might be a single file with tricky chars or we're in a directory where 
    # sys.argv[1] was actually correct and the rest were garbage.
    if not os.path.exists(input_path) and os.path.exists(sys.argv[1]):
        input_path = sys.argv[1]
    
    # Handle absolute path conversion
    input_path = os.path.abspath(input_path)

    if not os.path.exists(input_path):
        log_event(f"Error: File not found: {input_path}", level="error")
        sys.exit(1)

    log_event(f"--- Starting Med Chron Agent for: {input_path} ---")

    if os.path.isdir(input_path):
        log_event(f"Input is a directory. Scanning: {input_path}")
        files_to_process = []
        for root, _, files in os.walk(input_path):
            for file in files:
                if file.lower().endswith(('.pdf', '.docx')):
                    if "med_chron" in file.lower(): continue
                    files_to_process.append(os.path.join(root, file))
        
        if not files_to_process:
            log_event("No suitable files found.")
            sys.exit(0)
        
        for file_path in files_to_process:
            try:
                subprocess.run([sys.executable, sys.argv[0], file_path], check=True)
            except subprocess.CalledProcessError as e:
                log_event(f"Subprocess failed for {file_path}: {e}", level="error")
        sys.exit(0)
    
    # Single File
    text = extract_text(input_path)
    if not text:
        sys.exit(1)
        
    # Apply Filtering
    filtered_text = filter_content(text)
    if not filtered_text:
        log_event("No valid content found under specified headings.")
        sys.exit(0)

    try:
        with open(PROMPT_FILE, "r", encoding="utf-8") as f:
            prompt_instruction = f.read()
    except Exception as e:
        log_event(f"Error reading prompt file: {e}", level="error")
        sys.exit(1)

    # Determine Provider Name from Filename
    filename = os.path.basename(input_path)
    provider_name = extract_provider_from_filename(filename)
    log_event(f"Identified Provider from Filename: {provider_name}")

    final_content = call_gemini(prompt_instruction, filtered_text)
    if not final_content:
        sys.exit(1)

    # Save to Case Data
    data_manager = CaseDataManager()
    file_num_match = re.search(r"(\d{4}\.\d{3})", input_path)
    if file_num_match:
        file_num = file_num_match.group(1)
        
        # Create a variable key based on provider name if possible
        safe_provider = re.sub(r"[^a-zA-Z0-9_]", "_", provider_name.lower())
        var_key = f"med_chron_{safe_provider}"
        
        log_event(f"Saving medical chronology to case variable: {var_key}")
        data_manager.save_variable(file_num, var_key, final_content, source="med_chron_agent", extra_tags=["Evidence", "Medical Records", "Chronology"])
    else:
        log_event("Could not extract file number from path. Skipping variable save.", level="warning")

    # Determine Output Directory
    parts = input_path.split(os.sep)
    output_dir = None
    case_root_parts = None

    # Priority 1: Find folder starting with exactly 3 digits (Case Folder)
    for i in range(len(parts) - 1, -1, -1):
        if re.match(r'^\d{3}(\D|$)', parts[i]):
            log_event(f"Identified Case Folder by 3-digit pattern: {parts[i]}")
            case_root_parts = parts[:i+1]
            break

    # Priority 2: Standard "Current Clients" structure (Client/Case)
    if not case_root_parts:
        for i, part in enumerate(parts):
            if part.lower() == "current clients":
                if i + 2 < len(parts):
                    log_event("Identified Case Folder by Standard Structure (Current Clients + 2)")
                    case_root_parts = parts[:i+3]
                break
    
    if case_root_parts:
        output_dir = os.sep.join(case_root_parts + ["NOTES", "AI OUTPUT"])

    if not output_dir:
        # Fallback 1: If "NOTES" is already in the path, use it.
        for i in range(len(parts) - 1, -1, -1):
            if parts[i].upper() == "NOTES":
                output_dir = os.path.join(os.sep.join(parts[:i+1]), "AI OUTPUT")
                break
    
    if not output_dir:
        # Fallback 2: Sibling NOTES to the file's parent folder
        input_dir = os.path.dirname(input_path)
        parent_dir = os.path.dirname(input_dir)
        output_dir = os.path.join(parent_dir, "NOTES", "AI OUTPUT")

    if not os.path.exists(output_dir):
        try:
            os.makedirs(output_dir)
        except Exception:
            pass

    save_to_docx(final_content, output_dir, provider_name, filename)
    log_event("--- Agent Finished ---")

if __name__ == '__main__':
    main()
