import os
import sys
import shutil
import logging
import json
import glob
import re
import datetime
from typing import Optional, Dict, Any

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
except ImportError:
    print("Critical Error: python-docx not installed. Please install it.")
    sys.exit(1)

# --- Configuration ---
# Match complaint.py base path structure
BASE_PATH_WIN = r"Z:\\Shared\\Current Clients" 
PROJECT_ROOT = os.getcwd()
LOG_FILE = os.path.join(PROJECT_ROOT, "report_activity.log")
TEMPLATE_FILENAME = "Template Report.docx"
TEMPLATE_PATH = os.path.join(PROJECT_ROOT, TEMPLATE_FILENAME)
GEMINI_DATA_DIR = os.path.join(PROJECT_ROOT, ".gemini", "case_data")

# Setup Logging
logging.basicConfig(
    filename=LOG_FILE,
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    datefmt='%Y-%m-%d %H:%M:%S'
)

try:
    from icharlotte_core.utils import log_event, get_case_path
except ImportError:
    def log_event(message, level="info"):
        try:
            print(f"[{level.upper()}] {message}")
        except OSError:
            pass  # stdout pipe broken
    def get_case_path(file_num):
        return None
$', file_num):
        log_event(f"Invalid file number format: {file_num}. Expected ####.###", level="error")
        return None

    carrier_num, case_num = file_num.split('.')
    base_path = BASE_PATH_WIN

    if not os.path.exists(base_path):
        log_event(f"Base path not found: {base_path}", level="error")
        return None

    # Find Carrier Folder
    carrier_folders = glob.glob(os.path.join(base_path, f"{carrier_num} - *"))
    if not carrier_folders:
        log_event(f"Carrier folder starting with {carrier_num} not found in {base_path}", level="error")
        return None
    carrier_path = carrier_folders[0]

    # Find Case Folder
    case_folders = glob.glob(os.path.join(carrier_path, f"{case_num} - *"))
    if not case_folders:
        log_event(f"Case folder starting with {case_num} not found in {carrier_path}", level="error")
        return None
    
    return os.path.abspath(case_folders[0])

def load_case_data(file_num: str) -> Dict[str, Any]:
    """Loads JSON data for the case and flattens structured items to raw values."""
    json_path = os.path.join(GEMINI_DATA_DIR, f"{file_num}.json")
    raw_data = {}
    if os.path.exists(json_path):
        try:
            with open(json_path, 'r', encoding='utf-8') as f:
                raw_data = json.load(f)
        except Exception as e:
            log_event(f"Error loading case data JSON: {e}", level="error")
    
    # Flatten: extract 'value' from each entry if it's a dict following CaseDataManager structure
    flattened = {}
    for k, v in raw_data.items():
        if isinstance(v, dict) and "value" in v:
            flattened[k] = v["value"]
        else:
            flattened[k] = v
    return flattened

def replace_text_in_paragraph(paragraph, old_text, new_text, clear_formatting=False):
    """
    Replaces text in a paragraph. 
    If clear_formatting is True, removes bold/underline from the entire paragraph 
    (after replacement, which merges runs).
    Also fixes common punctuation spacing issues like 'Name ,' -> 'Name,'.
    """
    if old_text in paragraph.text:
        # Replace the text
        replaced_text = paragraph.text.replace(old_text, str(new_text))
        
        # Cleanup spacing issues (Requirement 2)
        replaced_text = replaced_text.replace(" ,", ",")
        
        paragraph.text = replaced_text

        # Requirement 1: Adjuster name/email should not be bold or underlined.
        if clear_formatting:
            for run in paragraph.runs:
                run.font.bold = False
                run.font.underline = False

def copy_run_formatting(source_run, target_run):
    """Copies font properties from source_run to target_run."""
    try:
        target_run.bold = source_run.bold
        target_run.italic = source_run.italic
        target_run.underline = source_run.underline
        if source_run.font.size:
            target_run.font.size = source_run.font.size
        if source_run.font.name:
            target_run.font.name = source_run.font.name
        if source_run.font.color and source_run.font.color.rgb:
            target_run.font.color.rgb = source_run.font.color.rgb
    except Exception:
        pass

def insert_docs_content(target_doc, placeholder, source_doc_path):
    """
    Replaces a placeholder paragraph with the content of another DOCX file,
    preserving source formatting, adding indentation, and spacing.
    """
    if not os.path.exists(source_doc_path):
        log_event(f"Source document for {placeholder} not found: {source_doc_path}", level="warning")
        return

    # Find the placeholder paragraph
    target_para = None
    for para in target_doc.paragraphs:
        if placeholder in para.text:
            target_para = para
            break
            
    if not target_para:
        log_event(f"Placeholder {placeholder} not found in target document.", level="warning")
        return

    try:
        source_doc = Document(source_doc_path)
        log_event(f"Inserting content from {os.path.basename(source_doc_path)} into {placeholder}...")

        # Insert content *before* the placeholder paragraph
        for p in source_doc.paragraphs:
            if not p.text.strip():
                continue 
            
            # Create new paragraph with same style
            new_p = target_para.insert_paragraph_before("", style=p.style)
            
            # Formatting: Indent + Spacing
            new_p.paragraph_format.first_line_indent = Inches(0.5)
            new_p.paragraph_format.space_before = Pt(0) # Prevent extra space above
            new_p.paragraph_format.space_after = Pt(12) # Blank line between paragraphs

            # Requirement 4: Preserve source formatting (copy runs)
            for run in p.runs:
                new_run = new_p.add_run(run.text)
                copy_run_formatting(run, new_run)
        
        # Remove the placeholder paragraph
        target_para.text = ""
        
    except Exception as e:
        log_event(f"Error reading/inserting source doc {source_doc_path}: {e}", level="error")

def insert_multiline_var(doc, placeholder, text):
    """
    Finds the placeholder and replaces it with multiline text (split by newlines),
    applying indentation and spacing to each new paragraph.
    """
    if not text:
        text = ""
        
    target_para = None
    for para in doc.paragraphs:
        if placeholder in para.text:
            target_para = para
            break
            
    if not target_para:
        return

    # Normalize newlines and split
    clean_text = str(text).replace('\r\n', '\n').replace('\r', '\n')
    paragraphs = clean_text.split('\n')
    
    for line in paragraphs:
        if not line.strip():
            continue
            
        new_p = target_para.insert_paragraph_before(line, style=target_para.style)
        
        # Apply formatting requests
        new_p.paragraph_format.first_line_indent = Inches(0.5)
        new_p.paragraph_format.space_before = Pt(0) # Prevent extra space above
        new_p.paragraph_format.space_after = Pt(12) # Blank line between paragraphs
        
    # Clear the placeholder
    target_para.text = ""

def main():
    if len(sys.argv) < 2:
        log_event("Usage: python report.py <file_number>", level="error")
        sys.exit(1)

    file_num = sys.argv[1]
    log_event(f"--- Starting Report Agent for {file_num} ---")

    # 1. Locate Case Folder
    case_path = get_case_path(file_num)
    if not case_path:
        sys.exit(1)
    
    log_event(f"Case Directory: {case_path}")
    
    status_dir = os.path.join(case_path, "STATUS")
    if not os.path.exists(status_dir):
        os.makedirs(status_dir)
        
    # 2. Check/Create Draft
    draft_pattern = os.path.join(status_dir, "[DRAFT] Carrier*.docx")
    existing_drafts = glob.glob(draft_pattern)
    
    report_doc_path = ""
    
    if existing_drafts:
        report_doc_path = existing_drafts[0]
        log_event(f"Found existing draft: {report_doc_path}")
    else:
        log_event("No draft found. Creating from template...")
        if not os.path.exists(TEMPLATE_PATH):
            log_event(f"Template Report.docx not found at {TEMPLATE_PATH}", level="error")
            sys.exit(1)
            
        new_filename = f"[DRAFT] Carrier {file_num}.docx" 
        report_doc_path = os.path.join(status_dir, new_filename)
        try:
            shutil.copy2(TEMPLATE_PATH, report_doc_path)
            log_event(f"Created draft: {report_doc_path}")
        except Exception as e:
            log_event(f"Error copying template: {e}", level="error")
            sys.exit(1)

    # 3. Load Variables
    case_data = load_case_data(file_num)
    
    # Simple string replacements
    plain_text_vars = [
        "[adjuster_name]",
        "[adjuster_email]",
        "[adjuster_name (first name only)]"
    ]

    replacements = {
        "[adjuster_name]": case_data.get("adjuster_name", ""),
        "[adjuster_email]": case_data.get("adjuster_email", ""),
        "[case_name]": case_data.get("case_name", ""),
        "[insured_name]": case_data.get("insured_name", ""),
        "[client_name]": case_data.get("client_name", ""),
        "[claim_number]": case_data.get("claim_number", ""),
        "[file_number]": case_data.get("file_number", file_num),
        "[case_number]": case_data.get("case_number", ""),
        "[incident_date]": case_data.get("incident_date", ""),
        "[trial_date]": case_data.get("trial_date", ""),
        "[procedural_history]": case_data.get("procedural_history", ""),
        # [factual_background] removed from here to be handled by insert_multiline_var
    }

    # Helper for First Name
    adj_name = replacements["[adjuster_name]"]
    if adj_name:
        replacements["[adjuster_name (first name only)]"] = adj_name.split()[0].strip()
    else:
        replacements["[adjuster_name (first name only)]"] = ""

    # 4. Process Document
    try:
        doc = Document(report_doc_path)
        
        # A. Simple Text Replacements
        for para in doc.paragraphs:
            for placeholder, value in replacements.items():
                if placeholder in para.text:
                    if value is None: value = ""
                    should_clear = placeholder in plain_text_vars
                    replace_text_in_paragraph(para, placeholder, value, clear_formatting=should_clear)
        
        # Tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for placeholder, value in replacements.items():
                            if placeholder in para.text:
                                if value is None: value = ""
                                should_clear = placeholder in plain_text_vars
                                replace_text_in_paragraph(para, placeholder, value, clear_formatting=should_clear)

        # B. Multiline Variables (Factual Background)
        # Requirement: Indent 0.5", Line breaks between paragraphs
        factual_bg = case_data.get("factual_background", "")
        insert_multiline_var(doc, "[factual_background]", factual_bg)

        # C. Content Insertions (Source Docs)
        ai_output_dir = os.path.join(case_path, "NOTES", "AI OUTPUT")
        
        content_map = {
            "[Document_Summary_Output]": "AI_OUTPUT.docx",
            # Discovery output is now dynamic, handled below
            "[Deposition_Summary_Output]": "deposition_summaries.docx",
            "[Liability_Eval_Output]": "liability_eval.docx",
            "[Exposure_Eval_Output]": "exposure_eval.docx"
        }
        
        # 1. Handle dynamic Discovery Responses
        discovery_placeholder = "[Discovery_Summary_Output]"
        # Find all files matching Discovery_Responses_*.docx
        discovery_files = glob.glob(os.path.join(ai_output_dir, "Discovery_Responses_*.docx"))
        
        if discovery_files:
            # Sort files for consistent order (e.g., by name)
            discovery_files.sort()
            log_event(f"Found {len(discovery_files)} discovery response files.")
            
            # Find the placeholder paragraph once
            target_para = None
            for para in doc.paragraphs:
                if discovery_placeholder in para.text:
                    target_para = para
                    break
            
            if target_para:
                # Insert each file's content
                for disc_file in discovery_files:
                    # We reuse insert_docs_content logic but need to adapt it since 
                    # it expects a placeholder which we just found. 
                    # Simpler approach: call insert_docs_content with the SAME placeholder
                    # for each file, but we need to ensure the placeholder isn't removed until the LAST file.
                    # Actually, insert_docs_content removes the placeholder.
                    # Better strategy: Insert all content *before* the placeholder, then remove placeholder at the end.
                    pass 

                # Re-implementing a multi-file insertion loop here for clarity and safety
                try:
                    for disc_file_path in discovery_files:
                        source_doc = Document(disc_file_path)
                        log_event(f"Inserting content from {os.path.basename(disc_file_path)}...")
                        
                        # Add a subheading for the file (optional, but good for separation)
                        # The filenames are like "Discovery_Responses_Party_Name.docx"
                        # We could extract "Party Name" but the doc content usually has headers.
                        
                        for p in source_doc.paragraphs:
                            if not p.text.strip():
                                continue
                            
                            new_p = target_para.insert_paragraph_before("", style=p.style)
                            new_p.paragraph_format.first_line_indent = Inches(0.5)
                            new_p.paragraph_format.space_before = Pt(0)
                            new_p.paragraph_format.space_after = Pt(12)
                            
                            for run in p.runs:
                                new_run = new_p.add_run(run.text)
                                copy_run_formatting(run, new_run)
                        
                        # Add a spacer between documents
                        spacer = target_para.insert_paragraph_before("")
                        # Fix: Ensure blank line doesn't inherit underline from placeholder (appearing as a dash)
                        try:
                            spacer.style = "Normal"
                        except:
                            pass
                        spacer.paragraph_format.space_after = Pt(24)

                    # Finally clear the placeholder
                    target_para.text = ""
                    
                except Exception as e:
                    log_event(f"Error inserting discovery docs: {e}", level="error")
            else:
                log_event(f"Placeholder {discovery_placeholder} not found.", level="warning")
        else:
            # Fallback for old filename if no new ones found
            old_discovery_file = "Discovery_Response_Summaries.docx"
            content_map[discovery_placeholder] = old_discovery_file

        # 2. Handle other static mappings
        for placeholder, filename in content_map.items():
            source_path = os.path.join(ai_output_dir, filename)
            insert_docs_content(doc, placeholder, source_path)

        # Save
        doc.save(report_doc_path)
        log_event(f"Report saved: {report_doc_path}")

        # --- NEW: Auto-apply Rules ---
        rules_path = os.path.join(GEMINI_DATA_DIR, "report_rules.json")
        rule_engine_path = os.path.join(PROJECT_ROOT, "Scripts", "rule_engine.py")
        
        if os.path.exists(rules_path) and os.path.exists(rule_engine_path):
            log_event("Applying formatting rules automatically...")
            import subprocess
            try:
                # Use sys.executable to ensure we use the same environment
                subprocess.run([sys.executable, rule_engine_path, "--apply", report_doc_path, rules_path], check=True)
                log_event("Rules applied successfully.")
            except Exception as e:
                log_event(f"Error applying rules: {e}", level="error")
        else:
             log_event("No report_rules.json found or rule_engine.py missing. Skipping auto-format.")

        log_event(f"--- Report Agent Finished for {file_num} ---")

    except Exception as e:
        log_event(f"Error processing document: {e}", level="error")
        import traceback
        log_event(traceback.format_exc(), level="error")
        sys.exit(1)

if __name__ == "__main__":
    main()