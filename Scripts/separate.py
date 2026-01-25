import argparse
import os
import sys
import re
import glob
import json
import subprocess
import logging
import platform
import shutil
import tempfile
import time
import concurrent.futures
import io
from google import genai

# Word Document Generation
try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    # We will handle the missing import gracefully when index creation is called
    pass

# --- Dependency Setup ---
# Attempt to import PyMuPDF (fitz) for fast processing
try:
    import fitz
except ImportError:
    print("Error: 'pymupdf' is required for this optimized version.")
    print("Please run: pip install pymupdf")
    sys.exit(1)

# PyPDF2 is still used for the final split to minimize write logic changes
try:
    import pypdf # Prefer pypdf
    PyPDF2 = pypdf
except ImportError:
    try:
        import PyPDF2
    except ImportError:
        print("Error: pypdf or PyPDF2 is required.")
        sys.exit(1)

# OCR Setup
OCR_AVAILABLE = False
try:
    import pytesseract
    from PIL import Image
    
    # Tesseract Path Logic
    if os.name == 'nt':
        tesseract_paths = [
            r"C:\\Program Files\\Tesseract-OCR\\tesseract.exe",
            os.path.expanduser(r"~\\AppData\\Local\\Programs\\Tesseract-OCR\\tesseract.exe")
        ]
        for p in tesseract_paths:
            if os.path.exists(p):
                pytesseract.pytesseract.tesseract_cmd = p
                OCR_AVAILABLE = True
                break
    else:
        if shutil.which("tesseract"):
            OCR_AVAILABLE = True
            
except ImportError:
    pass

# --- Logging Setup ---
PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
LOG_FILE = os.path.join(PROJECT_ROOT, "separator_activity.log")

logging.basicConfig(
    level=logging.INFO,
    format='[%(asctime)s] %(levelname)s: %(message)s',
    handlers=[
        logging.FileHandler(LOG_FILE),
        logging.StreamHandler(sys.stdout)
    ]
)
logger = logging.getLogger("Separator")

# --- Constants ---
PRIMARY_MODEL = "gemini-3-flash-preview"
FALLBACK_MODEL = "gemini-2.5-flash"
BASE_PATH_WIN = r"Z:\\Shared\\Current Clients"

# --- Functions ---

try:
    from icharlotte_core.utils import get_case_path
except ImportError:
    def get_case_path(file_num: str):
        """Locates the physical directory for a file number (####.###)."""
        if not re.match(r'^\d{4}\.\d{3}$', file_num):
            return None

        carrier_num, case_num = file_num.split('.')
        base_path = BASE_PATH_WIN

        if not os.path.exists(base_path):
            return None

        # Find Carrier Folder
        carrier_folders = glob.glob(os.path.join(base_path, f"{carrier_num} - *"))
        if not carrier_folders:
            return None
        carrier_path = carrier_folders[0]

        # Find Case Folder
        case_folders = glob.glob(os.path.join(carrier_path, f"{case_num} - *"))
        if not case_folders:
            return None
        
        return os.path.abspath(case_folders[0])

def get_page_text_fast(pdf_path, page_num):
    """
    Extracts text from a single page using PyMuPDF (fitz).
    If text is sparse, it renders the page image and uses OCR (Tesseract).
    
    page_num is 1-based.
    """
    try:
        # Open the PDF for this specific task (thread/process safe if opened here)
        doc = fitz.open(pdf_path)
        # fitz uses 0-based indexing
        page = doc.load_page(page_num - 1)
        
        # 1. Try direct text extraction (VERY FAST)
        text = page.get_text("text")
        
        # 2. If sparse, fall back to OCR
        if len(text.strip()) < 50 and OCR_AVAILABLE:
            # Render page to image (pixmap)
            # matrix=fitz.Matrix(2, 2) roughly doubles resolution (~144 dpi -> ~300 dpi) for better OCR
            # But for speed on headers, 1.5 or even 1.0 (72 dpi) might suffice. 
            # 2.0 is safer for accuracy.
            mat = fitz.Matrix(2, 2) 
            
            # Crop to top 25% of the page to save time
            rect = page.rect
            clip = fitz.Rect(0, 0, rect.width, rect.height * 0.25)
            
            pix = page.get_pixmap(matrix=mat, clip=clip)
            
            # Convert to PIL Image
            img_data = pix.tobytes("png")
            img = Image.open(io.BytesIO(img_data))
            
            # Run OCR
            ocr_text = pytesseract.image_to_string(img)
            
            if len(ocr_text.strip()) > len(text.strip()):
                text = ocr_text

        doc.close() 
        
        # Clean up text
        clean_text = " ".join(text[:1000].split())[:500]
        return f"Page {page_num}: {clean_text}" 
        
    except Exception as e:
        logger.error(f"Error processing page {page_num}: {e}")
        return f"Page {page_num}: [Error processing page]"

def extract_headers(pdf_path):
    """
    Extracts headers from all pages using parallel processing.
    """
    headers_map = {}
    
    try:
        # Open once to get page count (fast)
        doc = fitz.open(pdf_path)
        num_pages = len(doc)
        doc.close()
    except Exception as e:
        logger.error(f"Failed to read PDF structure: {e}")
        sys.exit(1)
        
    logger.info(f"Extracting headers from {num_pages} pages in {pdf_path} using parallel processing...")
    
    # Use ThreadPoolExecutor. 
    # ProcessPoolExecutor is better for CPU, but Tesseract invokes a subprocess anyway,
    # and fitz releases GIL for many ops. Threads are lighter and avoid pickling issues.
    max_workers = os.cpu_count() or 4
    if max_workers > 16: max_workers = 16 # Cap it reasonably
    
    with concurrent.futures.ThreadPoolExecutor(max_workers=max_workers) as executor:
        future_to_page = {
            executor.submit(get_page_text_fast, pdf_path, i + 1): i + 1 
            for i in range(num_pages)
        }
        
        processed_count = 0
        for future in concurrent.futures.as_completed(future_to_page):
            page_num = future_to_page[future]
            try:
                result = future.result()
                headers_map[page_num] = result
            except Exception as exc:
                logger.error(f"Page {page_num} generated an exception: {exc}")
                headers_map[page_num] = f"Page {page_num}: [Exception]"
            
            processed_count += 1
            if processed_count % 50 == 0 or processed_count == num_pages:
                logger.info(f"Processed {processed_count}/{num_pages} pages")
                
    # Sort headers by page number
    sorted_headers = [headers_map[i] for i in sorted(headers_map.keys())]
    return sorted_headers

def call_gemini(prompt, model):
    """Calls Gemini API directly using the new Google GenAI SDK."""
    api_key = os.environ.get("GEMINI_API_KEY")
    if not api_key:
        raise Exception("GEMINI_API_KEY environment variable not set.")
        
    client = genai.Client(api_key=api_key)
    
    try:
        # Run API call
        response = client.models.generate_content(model=model, contents=prompt)
        return response.text.strip()
    except Exception as e:
        raise Exception(f"Gemini API error: {e}")

def analyze_headers_chunk(headers_subset, start_page_num, next_id, prev_doc_context=None):
    """Analyzes a chunk of headers."""
    context_instruction = ""
    if prev_doc_context:
        context_instruction = (
            f"CONTEXT: The previous batch of pages ended with a document titled '{prev_doc_context['title']}' (ID: {prev_doc_context['id']}). "
            f"If Page {start_page_num} appears to be a continuation of this document, "
            f"start your list with an entry using ID {prev_doc_context['id']} and the same title. "
            f"Otherwise, start with ID {next_id}."
        )
    else:
        context_instruction = f"Start numbering new documents with ID {next_id}."

    prompt = (
        "I am processing a large PDF in batches. This batch contains headers from pages "
        f"{start_page_num} to {start_page_num + len(headers_subset) - 1}.\n"
        "Identify distinct legal or administrative documents.\n"
        "Format: ID|Title|Date|StartPage|EndPage\n"
        f"{context_instruction}\n"
        "Rules:\n"
        "1. Return ONLY the list, one document per line.\n"
        "2. Do not use markdown.\n"
        "3. If a document continues to the end of this batch, set EndPage to the last page of this batch.\n"
        "4. Provide a detailed and descriptive title for each document.\n"
        "5. If you identify an Insurance Policy, group all its related parts (Declarations, Endorsements, Conditions, Exclusions, etc.) into a SINGLE document entry. Do not split it.\n"
        "Example:\n"
        "1|Plaintiff's Complaint|2023-01-01|1|5\n"
        "2|Exhibit A|2023-01-02|6|10\n"
        "\nHEADERS:\n" + "\n".join(headers_subset)
    )
    
    response = ""
    try:
        response = call_gemini(prompt, PRIMARY_MODEL)
    except Exception as e:
        logger.warning(f"Primary model chunk analysis failed: {e}. Retrying with fallback...")
        try:
            response = call_gemini(prompt, FALLBACK_MODEL)
        except Exception as e2:
            logger.error(f"Fallback model failed for chunk {start_page_num}: {e2}")
            return []

    docs = []
    lines = response.split('\n')
    for line in lines:
        line = line.strip()
        if not line or '|' not in line: continue
        if line.startswith("```"): continue
        
        parts = line.split('|')
        if len(parts) >= 5:
            try:
                docs.append({
                    "id": parts[0].strip(),
                    "title": parts[1].strip(),
                    "date": parts[2].strip(),
                    "start": int(parts[3].strip()),
                    "end": int(parts[4].strip())
                })
            except ValueError:
                pass
    return docs

def analyze_headers(headers):
    """Sends headers to AI to find boundaries, using chunking for large files."""
    all_docs = []
    chunk_size = 100 # Process 100 pages at a time
    total_pages = len(headers)
    next_id = 1
    
    logger.info(f"Analyzing {total_pages} pages in chunks of {chunk_size}...")
    
    for i in range(0, total_pages, chunk_size):
        chunk = headers[i : i + chunk_size]
        start_page = i + 1
        end_page_batch = i + len(chunk)
        
        logger.info(f"Processing batch: Pages {start_page} to {end_page_batch}")
        
        prev_context = None
        if all_docs:
            prev_context = all_docs[-1]
            
        chunk_docs = analyze_headers_chunk(chunk, start_page, next_id, prev_context)
        
        if not chunk_docs:
            logger.warning(f"No documents found in batch {start_page}-{end_page_batch}")
            continue
            
        for doc in chunk_docs:
            # Check for merge (continuation)
            if all_docs and str(doc['id']) == str(all_docs[-1]['id']):
                # Extend the previous document
                all_docs[-1]['end'] = doc['end']
                logger.info(f"  Extended document '{all_docs[-1]['title']}' to page {doc['end']}")
            else:
                # New document
                all_docs.append(doc)
                
        # Calculate next_id
        current_ids = []
        for d in all_docs:
            try:
                current_ids.append(int(d['id']))
            except:
                pass
        
        if current_ids:
            next_id = max(current_ids) + 1
        else:
            next_id = 1

    if not all_docs:
        logger.error("No documents identified by AI across all batches.")
        sys.exit(1)
        
    return all_docs

def sanitize_filename(name):
    return "".join(c for c in name if c.isalnum() or c in (' ', '-', '_')).strip()

def split_pdf_pages(pdf_path, output_folder, selection):
    """Splits the PDF based on selection."""
    try:
        reader = PyPDF2.PdfReader(pdf_path)
    except Exception as e:
        print(f"Error reading source PDF: {e}")
        return

    base_name = os.path.splitext(os.path.basename(pdf_path))[0]
    
    print(f"\nProcessing {len(selection)} items...")
    
    for doc in selection:
        try:
            writer = PyPDF2.PdfWriter()
            start = doc['start'] - 1 # 1-based to 0-based
            end = doc['end']     # exclusive in python slice, inclusive in prompt? Prompt said "1-5", usually means 1,2,3,4,5. 
                                 # Range(start, end) where end is 5 would be 0,1,2,3,4. 
                                 # If page 5 is included, we need index 4. Range(0, 5) covers 0,1,2,3,4.
                                 # So if doc says end=5, we want loop range(0, 5). 
            
            # Verify bounds
            if start < 0: start = 0
            if end > len(reader.pages): end = len(reader.pages)
            
            for i in range(start, end):
                writer.add_page(reader.pages[i])
            
            safe_title = sanitize_filename(doc['title'])
            if len(safe_title) > 100:
                safe_title = safe_title[:100].strip()
            # Format: MyPacket- 01 - Complaint.pdf
            # User example: "MyPacket- 01 - Complaint.pdf" (original + id + title)
            # Assuming ID is numeric.
            try:
                 id_num = int(doc['id'])
                 id_str = f"{id_num:02d}"
            except:
                 id_str = str(doc['id'])
                 
            out_filename = f"{base_name} - {id_str} - {safe_title}.pdf"
            out_path = os.path.join(output_folder, out_filename)
            
            with open(out_path, "wb") as f:
                writer.write(f)
                
            print(f"  [OK] Saved: {out_filename}")
            logger.info(f"Saved {out_filename}")
            
        except Exception as e:
            print(f"  [ERROR] Failed to save {doc['title']}: {e}")
            logger.error(f"Failed to save {doc['title']}: {e}")

# --- Modes ---

def run_analysis(pdf_path, headless=False):
    if not os.path.exists(pdf_path):
        logger.error(f"File not found: {pdf_path}")
        return

    logger.info(f"Starting analysis for: {pdf_path}")
    
    # 1. Extract
    headers = extract_headers(pdf_path)
    
    # 2. Analyze
    docs = analyze_headers(headers)
    logger.info(f"identified {len(docs)} documents.")
    
    # 3. Create Index Document
    create_index_word(pdf_path, docs)
    
    # 4. Save Temp Map
    fd, temp_path = tempfile.mkstemp(suffix=".json", text=True)
    with os.fdopen(fd, 'w') as f:
        json.dump(docs, f)
    
    logger.info(f"Map saved to {temp_path}")

    if headless:
        print(f"JSON_MAP: {temp_path}")
        return

    logger.info("Launching interactive window...")
    
    # 5. Launch Interactive Window
    # We recursively call this script with different arguments in a new console
    current_script = os.path.abspath(__file__)
    
    cmd = [sys.executable, current_script, "--interactive", temp_path, "--original-pdf", pdf_path]
    
    if os.name == 'nt':
        subprocess.Popen(cmd, creationflags=subprocess.CREATE_NEW_CONSOLE)
    else:
        # Linux fallback (not primary target per prompt, but good practice)
        subprocess.Popen(cmd)

def get_notes_dir(pdf_path):
    """
    Locates the NOTES/AI OUTPUT folder within the case root.
    Uses logic similar to complaint.py by extracting file number from path.
    """
    norm_path = os.path.abspath(pdf_path)
    
    # 1. Try to extract ####.### pattern from path (Standardized logic)
    match = re.search(r'(\d{4})\.(\d{3})', norm_path)
    if match:
        file_num = f"{match.group(1)}.{match.group(2)}"
        case_root = get_case_path(file_num)
        if case_root:
            notes_dir = os.path.join(case_root, "NOTES", "AI OUTPUT")
            return notes_dir

    # 2. Fallback to 'Current Clients' path segment logic
    parts = norm_path.split(os.sep)
    try:
        idx = -1
        for i, part in enumerate(parts):
            if part.lower() == "current clients":
                idx = i
                break
        
        if idx != -1 and len(parts) > idx + 2:
            # Case root is usually: ...\Current Clients\#### - Carrier\### - Case
            case_root = os.sep.join(parts[:idx+3])
            notes_dir = os.path.join(case_root, "NOTES", "AI OUTPUT")
            return notes_dir
    except Exception as e:
        logger.warning(f"Failed to parse case root from path: {e}")

    # 3. Fallback: climb up the tree looking for a NOTES folder
    current_dir = os.path.dirname(norm_path)
    for _ in range(5):
        potential_notes = os.path.join(current_dir, "NOTES")
        if os.path.exists(potential_notes) and os.path.isdir(potential_notes):
            return os.path.join(potential_notes, "AI OUTPUT")
        parent = os.path.dirname(current_dir)
        if parent == current_dir: break
        current_dir = parent
        
    # Final Fallback
    return os.path.join(os.path.dirname(norm_path), "NOTES", "AI OUTPUT")

def create_index_word(pdf_path, docs):
    """Creates a Word document index of identified sub-documents."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        logger.error("python-docx not installed. Skipping index creation.")
        return

    try:
        notes_ai_dir = get_notes_dir(pdf_path)
        indexes_dir = os.path.join(notes_ai_dir, "INDEXES")
        
        if not os.path.exists(indexes_dir):
            os.makedirs(indexes_dir, exist_ok=True)
            logger.info(f"Created INDEXES folder: {indexes_dir}")
        
        base_name = os.path.splitext(os.path.basename(pdf_path))[0]
        docx_filename = f"Index_{base_name}.docx"
        docx_path = os.path.join(indexes_dir, docx_filename)
        
        doc = Document()
        
        # Set default style to Times New Roman, 12, Black
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        font.color.rgb = RGBColor(0, 0, 0)
        
        # Header paragraph
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"INDEX OF DOCUMENTS - {base_name}")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)

        # Create Table
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        table.autofit = False # Disable autofit to allow manual widths

        # Define Column Widths (Total ~6.5 inches for standard margins)
        # Id: ~0.4", Date: ~1.2", Pages: ~1.2", Title: Remaining (~3.7")
        col_widths = [Inches(0.4), Inches(3.7), Inches(1.2), Inches(1.2)]
        
        # Setup Headers
        headers = ['Id', 'Document Title', 'Document Date', 'Page Ranges']
        hdr_cells = table.rows[0].cells
        for i, h_text in enumerate(headers):
            cell = hdr_cells[i]
            cell.width = col_widths[i]
            cell_p = cell.paragraphs[0]
            cell_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cell_p.add_run(h_text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
            
        # Add Data Rows
        for d in docs:
            row_cells = table.add_row().cells
            
            # Formatting Page Ranges
            start_pg = d.get('start', '')
            end_pg = d.get('end', '')
            page_range = f"{start_pg} - {end_pg}" if start_pg != end_pg else str(start_pg)

            row_data = [
                str(d.get('id', '')),
                str(d.get('title', '')),
                str(d.get('date', '')),
                page_range
            ]
            for i, val in enumerate(row_data):
                cell = row_cells[i]
                cell.width = col_widths[i]
                cell_p = cell.paragraphs[0]
                run = cell_p.add_run(val)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(0, 0, 0)
        
        doc.save(docx_path)
        logger.info(f"Index saved to: {docx_path}")
        print(f"\n[SUCCESS] Index created: {docx_filename} in INDEXES folder.")

    except Exception as e:
        logger.error(f"Error creating index: {e}")
        print(f"\n[ERROR] Failed to create index document: {e}")

def merge_pdf_pages(pdf_path, output_folder, selection, merged_title):
    """Merges selected documents into a single PDF."""
    try:
        reader = PyPDF2.PdfReader(pdf_path)
    except Exception as e:
        print(f"Error reading source PDF: {e}")
        return

    print(f"\nMerging {len(selection)} items into '{merged_title}'...")
    
    try:
        writer = PyPDF2.PdfWriter()
        
        for doc in selection:
            start = doc['start'] - 1 # 1-based to 0-based
            end = doc['end']
            
            # Verify bounds
            if start < 0: start = 0
            if end > len(reader.pages): end = len(reader.pages)
            
            for i in range(start, end):
                writer.add_page(reader.pages[i])
        
        safe_title = sanitize_filename(merged_title)
        if not safe_title.lower().endswith('.pdf'):
            safe_title += ".pdf"
            
        out_path = os.path.join(output_folder, safe_title)
        
        with open(out_path, "wb") as f:
            writer.write(f)
            
        print(f"  [OK] Saved merged file: {safe_title}")
        logger.info(f"Saved merged file {safe_title}")
        
    except Exception as e:
        print(f"  [ERROR] Failed to save merged file: {e}")
        logger.error(f"Failed to save merged file: {e}")

def run_interactive(json_path, pdf_path):
    # Load Docs
    docs = []
    try:
        with open(json_path, 'r') as f:
            docs = json.load(f)
    except Exception as e:
        print(f"Error loading document map: {e}")
        input("Press Enter to exit...")
        return

    # Display Loop
    while True:
        os.system('cls' if os.name == 'nt' else 'clear')
        print(f"--- Document Separator: {os.path.basename(pdf_path)} ---\n")
        print(f"{ 'ID':<4} | { 'Date':<12} | { 'Pages':<8} | {'Title'}")
        print("-" * 60)
        for doc in docs:
            print(f"{doc['id']:<4} | {doc['date']:<12} | {doc['start']}-{doc['end']:<6} | {doc['title']}")
        print("-" * 60)
        
        print("\nCommands:")
        print("  <ids>          : Extract specific IDs (e.g., '1, 3', '1-5', 'all')")
        print("  merge <ids>    : Merge specific IDs into one PDF (e.g., 'merge 1, 4, 9')")
        print("  q              : Quit")
        
        choice = input("> ").strip().lower()
        
        if choice == 'q':
            break
            
        selection = []
        is_merge = False
        
        if choice.startswith('merge'):
            is_merge = True
            choice = choice[5:].strip() # Remove 'merge' prefix
        
        if choice == 'all':
            selection = docs
        else:
            try:
                # Parse range/list
                ids_to_keep = set()
                parts = choice.split(',')
                for part in parts:
                    part = part.strip()
                    if not part: continue
                    if '-' in part:
                        s, e = map(int, part.split('-'))
                        # Convert range integers to strings to match ID format
                        ids_to_keep.update(str(i) for i in range(s, e + 1))
                    else:
                        ids_to_keep.add(str(part)) # Try string first
                        ids_to_keep.add(str(int(part))) # Also numeric string
                
                selection = [d for d in docs if str(d['id']) in ids_to_keep]
            except Exception as e:
                print(f"Invalid input: {e}")
                time.sleep(1)
                continue
        
        if not selection:
            print("No documents selected.")
            time.sleep(1)
            continue
            
        # Process Selection
        base_dir = os.path.dirname(os.path.abspath(pdf_path))
        source_name = os.path.splitext(os.path.basename(pdf_path))[0]
        output_folder = os.path.join(base_dir, f"PULLED-{source_name}")
        
        if not os.path.exists(output_folder):
            os.makedirs(output_folder)
            
        if is_merge:
            merged_title = input("Enter title for the merged PDF: ").strip()
            if not merged_title:
                merged_title = f"Merged_Documents_{selection[0]['id']}-{selection[-1]['id']}"
            merge_pdf_pages(pdf_path, output_folder, selection, merged_title)
        else:
            split_pdf_pages(pdf_path, output_folder, selection)
        
        print("\nDone! Files saved to:")
        print(output_folder)
        print("\nPress Enter to close this window.")
        input()
        
        # Cleanup temp file
        try:
            os.remove(json_path)
        except:
            pass
        break

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("pdf_path", help="PDF File Path", nargs='*')
    parser.add_argument("--interactive", help="Path to temp json for interactive mode")
    parser.add_argument("--original-pdf", help="Original PDF for interactive mode")
    parser.add_argument("--headless", action="store_true", help="Run without UI, output JSON path")
    
    # We use parse_known_args to allow for loose argument handling if needed,
    # but standard parse_args is usually safer if we define everything.
    # However, to support "script.py file path with spaces", we need careful handling.
    
    # Check if flags are present that indicate structured usage
    structured_args = "--interactive" in sys.argv or "--headless" in sys.argv or "--original-pdf" in sys.argv
    
    if structured_args:
        args = parser.parse_args()
        if args.interactive and args.original_pdf:
            run_interactive(args.interactive, args.original_pdf)
        elif args.pdf_path:
            # Join if split by spaces (shouldn't happen with correct quoting, but robust)
            path = " ".join(args.pdf_path) if isinstance(args.pdf_path, list) else args.pdf_path
            # Strip quotes
            path = path.strip().strip('"').strip("'")
            run_analysis(path, headless=args.headless)
        else:
             print("Usage: separate.py <pdf_path> [--headless]")
    else:
        # Legacy/Drag-and-Drop Mode: Treat all args as path
        if len(sys.argv) < 2:
            print("Usage: separate.py <pdf_path>")
            sys.exit(1)
            
        raw_path = " ".join(sys.argv[1:])
        pdf_path = raw_path.strip().strip('"').strip("'")
        
        # Fallback check
        if not os.path.exists(pdf_path) and os.path.exists(sys.argv[1]):
            pdf_path = sys.argv[1]
            
        run_analysis(pdf_path, headless=False)

if __name__ == "__main__":
    main()