"""
Summarize Deposition Agent for iCharlotte

Summarizes deposition transcripts using LLM with multi-pass processing.

Features:
- Parallel LLM calls for extraction and summary
- Dynamic topic count based on transcript length
- Automatic deponent name/date extraction from content
- Impeachment material flagging
- Exhibit reference tracking
- Multi-provider LLM fallback (Gemini, Claude, OpenAI)
- Memory monitoring for large transcripts
- Structured progress reporting for UI integration
"""

import os
import sys
import re
import subprocess
import datetime
from docx import Document
from docx.shared import Pt, Inches

# Add parent directory to path for imports
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

# Import shared infrastructure
from icharlotte_core.document_processor import DocumentProcessor, OCRConfig
from icharlotte_core.agent_logger import AgentLogger, create_legacy_log_event
from icharlotte_core.llm_config import LLMCaller, LLMConfig
from icharlotte_core.memory_monitor import MemoryMonitor, track_memory
from icharlotte_core.exceptions import (
    ExtractionError, LLMError,
    SummaryPassError, MemoryLimitError
)
from icharlotte_core.docx_writer import get_docx_lock, LockTimeoutError
from icharlotte_core.deposition import session_manager

# Import Case Data Manager
try:
    from case_data_manager import CaseDataManager
except ImportError:
    sys.path.append(os.path.join(os.getcwd(), 'Scripts'))
    from case_data_manager import CaseDataManager

# Import Document Registry for classification
try:
    from document_registry import DocumentRegistry
except ImportError:
    from Scripts.document_registry import DocumentRegistry


# =============================================================================
# Configuration
# =============================================================================

SCRIPTS_DIR = os.path.dirname(os.path.abspath(__file__))
NARRATIVE_PROMPT_FILE = os.path.join(SCRIPTS_DIR, "SUMMARIZE_DEPOSITION_PROMPT.txt")
CROSS_CHECK_PROMPT_FILE = os.path.join(SCRIPTS_DIR, "DEPOSITION_CROSS_CHECK_PROMPT.txt")
TOPIC_DISCOVERY_PROMPT_FILE = os.path.join(SCRIPTS_DIR, "DEPOSITION_TOPIC_DISCOVERY_PROMPT.txt")

# Legacy log file for backward compatibility
LEGACY_LOG_FILE = r"C:\GeminiTerminal\Summarize_Deposition_activity.log"


# =============================================================================
# File Number Extraction
# =============================================================================

def extract_file_number(path: str) -> str:
    """Extract file number from path.

    Handles paths like:
    - Z:\\Shared\\Current Clients\\3800- NATIONWIDE\\3850\\084 - Dudash\\...
    - Paths containing literal "3850.084"
    """
    # Standard pattern: 1234.567 (with literal dot)
    match = re.search(r"(\d{4}\.\d{3})", path)
    if match:
        return match.group(1)

    # Parse from directory structure
    parts = os.path.normpath(path).split(os.sep)
    try:
        # Find consecutive folders matching client/matter pattern
        for i in range(len(parts) - 1):
            client_match = re.match(r"^(\d{4})(?:\D|$)", parts[i])
            if client_match:
                matter_match = re.match(r"^(\d{3})(?:\D|$)", parts[i + 1])
                if matter_match:
                    return f"{client_match.group(1)}.{matter_match.group(1)}"

        # Fallback: Look for "Current Clients" structure
        cc_index = -1
        for i, part in enumerate(parts):
            if part.lower() == "current clients":
                cc_index = i
                break

        if cc_index != -1 and cc_index + 3 < len(parts):
            client_folder = parts[cc_index + 2]
            matter_folder = parts[cc_index + 3]

            client_code = re.match(r"^(\d{4})(?:\D|$)", client_folder)
            matter_code = re.match(r"^(\d{3})(?:\D|$)", matter_folder)

            if client_code and matter_code:
                return f"{client_code.group(1)}.{matter_code.group(1)}"

    except Exception:
        pass

    return None


# =============================================================================
# Deponent Information Extraction
# =============================================================================

class DeponentExtractor:
    """Extracts deponent information from transcript text."""

    # Common deposition header patterns
    DEPONENT_PATTERNS = [
        # "DEPOSITION OF JOHN SMITH" or "DEPOSITION OF JOHN SMITH, M.D."
        r"DEPOSITION\s+OF\s+([A-Z][A-Za-z]+(?:\s+[A-Z]\.?\s*)?(?:\s+[A-Z][A-Za-z]+)+(?:,?\s*(?:M\.?D\.?|D\.?O\.?|Ph\.?D\.?|D\.?C\.?))?)",
        # "Deponent: John Smith"
        r"[Dd]eponent[:\s]+([A-Z][A-Za-z]+(?:\s+[A-Z]\.?\s*)?(?:\s+[A-Z][A-Za-z]+)+)",
        # "THE VIDEOTAPED DEPOSITION OF JOHN SMITH"
        r"(?:VIDEOTAPED\s+)?DEPOSITION\s+OF\s+([A-Z][A-Za-z]+(?:\s+[A-Z]\.?\s*)?(?:\s+[A-Z][A-Za-z]+)+)",
        # "Examination of John Smith"
        r"[Ee]xamination\s+of\s+([A-Z][A-Za-z]+(?:\s+[A-Z]\.?\s*)?(?:\s+[A-Z][A-Za-z]+)+)",
        # "WITNESS: JOHN SMITH" (in caption)
        r"WITNESS[:\s]+([A-Z][A-Za-z]+(?:\s+[A-Z]\.?\s*)?(?:\s+[A-Z][A-Za-z]+)+)",
    ]

    # Date patterns
    DATE_PATTERNS = [
        # "January 15, 2024" or "January 15th, 2024"
        r"([A-Z][a-z]+\s+\d{1,2}(?:st|nd|rd|th)?,?\s+\d{4})",
        # "01/15/2024" or "1/15/2024"
        r"(\d{1,2}/\d{1,2}/\d{4})",
        # "2024-01-15"
        r"(\d{4}-\d{2}-\d{2})",
        # "15 January 2024"
        r"(\d{1,2}\s+[A-Z][a-z]+\s+\d{4})",
    ]

    # Doctor patterns
    DOCTOR_PATTERNS = [
        r"(?:Dr\.?|Doctor)\s+([A-Z][A-Za-z]+)",
        r"([A-Z][A-Za-z]+),?\s+(?:M\.?D\.?|D\.?O\.?|Ph\.?D\.?)",
    ]

    @classmethod
    def extract_deponent_name(cls, text: str, first_n_chars: int = 5000) -> str:
        """
        Extract deponent name from transcript text.

        Args:
            text: Full transcript text.
            first_n_chars: Number of characters to search (deponent info usually at start).

        Returns:
            Deponent name or empty string if not found.
        """
        search_text = text[:first_n_chars]

        for pattern in cls.DEPONENT_PATTERNS:
            match = re.search(pattern, search_text, re.IGNORECASE)
            if match:
                name = match.group(1).strip()
                # Clean up the name
                name = re.sub(r'\s+', ' ', name)
                # Remove trailing punctuation
                name = name.rstrip('.,;:')
                return name

        return ""

    @classmethod
    def extract_deposition_date(cls, text: str, first_n_chars: int = 3000) -> str:
        """
        Extract deposition date from transcript text.

        Args:
            text: Full transcript text.
            first_n_chars: Number of characters to search.

        Returns:
            Date string or empty string if not found.
        """
        search_text = text[:first_n_chars]

        # Look for date near "taken on" or "held on" or just in the header
        date_context_patterns = [
            r"(?:taken|held|conducted)\s+(?:on\s+)?(" + "|".join(cls.DATE_PATTERNS) + ")",
            r"(?:Date|DATE)[:\s]+(" + "|".join(cls.DATE_PATTERNS) + ")",
        ]

        for pattern in date_context_patterns:
            match = re.search(pattern, search_text, re.IGNORECASE)
            if match:
                return match.group(1).strip()

        # Fallback: just find any date in the header area
        for pattern in cls.DATE_PATTERNS:
            match = re.search(pattern, search_text)
            if match:
                return match.group(1).strip()

        return ""

    @classmethod
    def detect_deponent_type(cls, text: str, name: str = "") -> str:
        """
        Detect if deponent is plaintiff, defendant, doctor, expert, etc.

        Args:
            text: Transcript text (first few thousand chars).
            name: Deponent name if already extracted.

        Returns:
            Deponent type string.
        """
        search_text = text[:10000].upper()
        name_upper = name.upper() if name else ""

        # Check for doctor
        if name:
            for pattern in cls.DOCTOR_PATTERNS:
                if re.search(pattern, text[:5000], re.IGNORECASE):
                    return "Treating Physician"

        if "M.D." in name_upper or "D.O." in name_upper or "PH.D." in name_upper:
            return "Expert/Physician"

        # Check context clues
        if "PLAINTIFF" in search_text and name_upper in search_text:
            # Check if this name appears near "plaintiff"
            if re.search(rf"PLAINTIFF\s*{re.escape(name_upper)}", search_text):
                return "Plaintiff"

        if "DEFENDANT" in search_text and name_upper in search_text:
            if re.search(rf"DEFENDANT\s*{re.escape(name_upper)}", search_text):
                return "Defendant"

        if "EXPERT" in search_text:
            return "Expert Witness"

        return "Witness"

    @classmethod
    def estimate_page_count(cls, text: str, chars_per_page: int = 3000) -> int:
        """Estimate page count from text length."""
        return max(1, len(text) // chars_per_page)

    @classmethod
    def calculate_topic_count(cls, text: str) -> int:
        """
        Calculate appropriate number of topics based on transcript length.

        Returns:
            Recommended number of topics (5-20).
        """
        pages = cls.estimate_page_count(text)

        if pages < 50:
            return min(8, max(5, pages // 6))
        elif pages < 150:
            return min(12, max(8, pages // 12))
        else:
            return min(20, max(12, pages // 15))


# =============================================================================
# Phase 1: Topic Discovery
# =============================================================================

import json as _json  # local alias to avoid shadowing in helpers


def _parse_topic_response(response: str) -> list:
    """Parse the LLM topic-discovery response into a list of topic dicts.

    Falls back to bullet-list parsing if the response is not valid JSON.
    Always returns at most 25 topics, sorted by rank (ascending).
    """
    response = (response or "").strip()
    topics: list = []

    # Strip code fences if present
    if response.startswith("```"):
        lines = response.splitlines()
        # Drop the opening ``` line (and ```json variants) and the closing ``` if any.
        if lines and lines[0].startswith("```"):
            lines = lines[1:]
        if lines and lines[-1].startswith("```"):
            lines = lines[:-1]
        response = "\n".join(lines).strip()

    # Try JSON first
    try:
        parsed = _json.loads(response)
        if isinstance(parsed, list):
            for i, entry in enumerate(parsed, 1):
                if not isinstance(entry, dict):
                    continue
                title = str(entry.get("title", "")).strip()
                if not title:
                    continue
                rank = entry.get("rank", i)
                try:
                    rank = int(rank)
                except (TypeError, ValueError):
                    rank = i
                density = str(entry.get("discussion_density", "medium")).strip().lower()
                if density not in ("high", "medium", "low"):
                    density = "medium"
                topics.append({
                    "id": len(topics) + 1,
                    "title": title,
                    "rank": rank,
                    "discussion_density": density,
                })
    except _json.JSONDecodeError:
        pass

    # Fallback: bullet-list parsing
    if not topics:
        for line in response.splitlines():
            stripped = line.strip()
            if not stripped:
                continue
            for prefix in ("- ", "* ", "• "):
                if stripped.startswith(prefix):
                    stripped = stripped[len(prefix):].strip()
                    break
            # Drop leading numbering like "1." or "1)"
            stripped = re.sub(r"^\d+[\.\)]\s*", "", stripped)
            if not stripped:
                continue
            topics.append({
                "id": len(topics) + 1,
                "title": stripped,
                "rank": len(topics) + 1,
                "discussion_density": "medium",
            })

    # Sort by rank and truncate
    topics.sort(key=lambda t: t["rank"])
    topics = topics[:25]
    # Re-stamp ids and ranks to match final order
    for i, t in enumerate(topics, 1):
        t["id"] = i
        t["rank"] = i
    return topics


def process_topics(input_path: str, logger) -> bool:
    """Phase 1: extract text, discover topics, write session JSON, await input."""
    memory_monitor = MemoryMonitor(warn_threshold_mb=1500, abort_threshold_mb=2000, logger=logger.info)
    llm_caller = LLMCaller(logger=logger)

    logger.progress(2, "Initializing topic discovery...")

    if not os.path.exists(TOPIC_DISCOVERY_PROMPT_FILE):
        logger.error(f"Missing prompt file: {TOPIC_DISCOVERY_PROMPT_FILE}")
        return False
    with open(TOPIC_DISCOVERY_PROMPT_FILE, "r", encoding="utf-8") as f:
        topic_prompt = f.read()

    logger.progress(5, "Extracting transcript text...")
    try:
        with memory_monitor.track_operation("Text Extraction"):
            processor = DocumentProcessor(ocr_config=OCRConfig(adaptive=True), logger=logger)
            result = processor.extract_with_dynamic_ocr(input_path)
            if not result.success:
                raise ExtractionError(f"Failed to extract text: {result.error}", file_path=input_path)
            text = result.text
    except Exception as e:
        logger.pass_failed("Text Extraction", str(e), recoverable=False)
        return False

    logger.progress(20, f"Extracted {len(text)} chars; caching transcript")

    paths = session_manager.compute_session_paths(input_path)
    paths.cached_text_path.write_text(text, encoding="utf-8")

    deponent_name = DeponentExtractor.extract_deponent_name(text) or os.path.splitext(os.path.basename(input_path))[0]
    deposition_date = DeponentExtractor.extract_deposition_date(text)
    deponent_type = DeponentExtractor.detect_deponent_type(text, deponent_name)
    file_number = extract_file_number(input_path)

    logger.progress(40, "Calling LLM for topic discovery...")
    try:
        response = llm_caller.call(topic_prompt, text, task_type="summary")
    except Exception as e:
        logger.pass_failed("Topic Discovery", str(e), recoverable=False)
        return False

    topics = _parse_topic_response(response)
    if not topics:
        logger.error("Topic discovery returned no usable topics")
        return False

    session_data = {
        "version": 1,
        "phase": "awaiting_input",
        "input_path": str(input_path),
        "cached_text_path": str(paths.cached_text_path),
        "deponent_name": deponent_name,
        "deposition_date": deposition_date,
        "deponent_type": deponent_type,
        "file_number": file_number,
        "topics": topics,
        "user_config": None,
    }
    session_manager.write_session(paths.session_path, session_data)

    logger.progress(95, f"Discovered {len(topics)} topics; awaiting user input")
    print(f"AWAITING_INPUT:{paths.session_path}", flush=True)
    return True


# =============================================================================
# Phase 2: Topic-Locked Summary
# =============================================================================

def _build_topic_locked_prompt(base_prompt: str, *, topic_list: list, bullets_per_topic: int,
                                deponent_label: str, custom_rules: str) -> str:
    """Render the topic-locked summary prompt with user-supplied substitutions."""
    rendered_topics = "\n".join(f"- {t}" for t in topic_list)
    return (base_prompt
            .replace("{deponent_label}", deponent_label)
            .replace("{bullets_per_topic}", str(bullets_per_topic))
            .replace("{topic_list}", rendered_topics)
            .replace("{custom_rules}", custom_rules or "(none)"))


def _register_outputs(input_path, summary, deponent_name, deponent_type, output_file, logger):
    """Wrapper around CaseDataManager + DocumentRegistry writes. Best-effort; logs but does not fail the run."""
    try:
        data_manager = CaseDataManager()
        file_num = extract_file_number(input_path)
        if file_num:
            clean_name = re.sub(r"[^a-zA-Z0-9_]", "_", (deponent_name or "unknown").lower())
            var_key = f"depo_summary_{clean_name}"
            data_manager.save_variable(
                file_num, var_key, summary,
                source="deposition_agent",
                extra_tags=["Deposition", deponent_type] if deponent_type else ["Deposition"],
            )
    except Exception as e:
        logger.warning(f"Could not save to case data: {e}")

    try:
        file_num = extract_file_number(input_path)
        if not file_num:
            return
        depo_type_map = {
            "Plaintiff": "Deposition - Plaintiff",
            "Defendant": "Deposition - Defendant",
            "Witness": "Deposition - Witness",
            "Expert Witness": "Deposition - Expert",
            "Expert/Physician": "Deposition - Expert",
            "Treating Physician": "Deposition - Expert",
            "Corporate Representative": "Deposition - Corporate Representative",
        }
        registry_doc_type = depo_type_map.get(deponent_type, "Deposition - Witness")
        from document_registry import DocumentClassifier
        classifier = DocumentClassifier(logger=logger)
        base_name = os.path.splitext(os.path.basename(input_path))[0]
        standardized_name = classifier.generate_name(summary, registry_doc_type, fallback_name=base_name)
        registry = DocumentRegistry()
        registry.register_document(
            file_number=file_num,
            name=standardized_name,
            document_type=registry_doc_type,
            source_path=input_path,
            summary_location=output_file,
            agent="summarize_deposition",
            char_count=len(summary),
        )
    except Exception as e:
        logger.warning(f"Could not register document: {e}")


def process_summary(session_path: str, logger) -> bool:
    """Phase 2: read session, generate topic-locked summary, save docx, cleanup."""
    from pathlib import Path
    memory_monitor = MemoryMonitor(warn_threshold_mb=1500, abort_threshold_mb=2000, logger=logger.info)
    llm_caller = LLMCaller(logger=logger)

    session_path = Path(session_path)
    logger.progress(5, "Loading session...")
    try:
        session = session_manager.read_session(session_path)
    except Exception as e:
        logger.error(f"Failed to load session: {e}")
        return False

    if session.get("phase") != "ready_for_summary" or not session.get("user_config"):
        logger.pass_failed("Phase 2 Init", "Session is not ready_for_summary", recoverable=False)
        return False

    cached_path = Path(session["cached_text_path"])
    if not cached_path.exists():
        logger.pass_failed("Phase 2 Init", f"Cached transcript missing: {cached_path}", recoverable=False)
        return False

    logger.progress(10, "Reading cached transcript...")
    text = cached_path.read_text(encoding="utf-8")

    cfg = session["user_config"]
    final_topics = [t for t in cfg.get("selected_topics", [])] + [t for t in cfg.get("added_topics", [])]
    if not final_topics:
        logger.pass_failed("Phase 2 Init", "No topics selected", recoverable=False)
        return False

    with open(NARRATIVE_PROMPT_FILE, "r", encoding="utf-8") as f:
        base_prompt = f.read()

    prompt = _build_topic_locked_prompt(
        base_prompt,
        topic_list=final_topics,
        bullets_per_topic=cfg.get("bullets_per_topic", 5),
        deponent_label=cfg.get("deponent_label") or session.get("deponent_type", "Deponent"),
        custom_rules=cfg.get("custom_rules", ""),
    )

    logger.progress(30, "Generating summary...")
    try:
        with memory_monitor.track_operation("Narrative Summary"):
            summary = llm_caller.call(prompt, text, task_type="summary")
        if not summary:
            raise SummaryPassError("LLM returned empty summary")
    except Exception as e:
        logger.pass_failed("Narrative Summary", str(e), recoverable=False)
        return False
    logger.progress(70, f"Summary generated: {len(summary)} chars")

    if cfg.get("cross_check_enabled"):
        try:
            with open(CROSS_CHECK_PROMPT_FILE, "r", encoding="utf-8") as f:
                cross_prompt = f.read()
            cross_prompt = cross_prompt.replace("{summary}", summary).replace("{original}", text[:75000])
            logger.progress(75, "Running cross-check...")
            verified = llm_caller.call(cross_prompt, "", task_type="cross_check")
            if verified and len(verified) > len(summary) * 0.8:
                summary = verified
        except Exception as e:
            logger.warning(f"Cross-check failed; using original summary: {e}")
    logger.progress(85, "Cross-check complete")

    output_dir = get_output_directory(session["input_path"])
    os.makedirs(output_dir, exist_ok=True)
    output_file = os.path.join(output_dir, "Deposition_summaries.docx")
    logger.progress(90, f"Saving to {os.path.basename(output_file)}...")
    if not save_to_docx(summary, output_file, session["deponent_name"], session["deposition_date"], logger):
        return False
    logger.progress(95, "Document saved")

    _register_outputs(session["input_path"], summary, session["deponent_name"],
                       session.get("deponent_type"), output_file, logger)

    session_manager.cleanup_session(session_path)
    logger.progress(100, "Deposition summarization complete")
    return True


# =============================================================================
# Document Output Functions
# =============================================================================

def add_markdown_to_doc(doc, content):
    """Parses basic Markdown and adds it to the docx Document with specific formatting."""
    lines = content.split('\n')

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        # Headings (## or #)
        if stripped.startswith('#'):
            text = stripped.lstrip('#').strip()
            if not text.endswith('.'):
                text += "."
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.first_line_indent = Inches(0.5)
            run = p.add_run(text)
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            continue

        # List items (handle *, -, •, and numbered bullets)
        if stripped.startswith('* ') or stripped.startswith('- ') or stripped.startswith('• '):
            # Strip bullet prefix (handle •  which is 1 char + space)
            if stripped.startswith('• '):
                text = stripped[2:].strip()
            else:
                text = stripped[2:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.first_line_indent = Inches(-0.25)

            run = p.add_run("\t")
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

            # Support bold parsing within list items
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    r = p.add_run(part[2:-2])
                    r.bold = True
                else:
                    r = p.add_run(part)
                r.font.name = 'Times New Roman'
                r.font.size = Pt(12)
            continue

        # Normal text (with bold support)
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.first_line_indent = Inches(0.5)

        parts = re.split(r'(\*\*.*?\*\*)', stripped)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
            else:
                run = p.add_run(part)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)


def save_to_docx(content: str, output_path: str, deponent_name: str,
                 deposition_date: str, logger: AgentLogger) -> bool:
    """
    Saves the content to a DOCX file with deposition formatting and process-safe locking.

    Args:
        content: Summary content to save.
        output_path: Path to output file (will append if exists).
        deponent_name: Name of the deponent.
        deposition_date: Date of the deposition.
        logger: AgentLogger instance.

    Returns:
        True if successful.
    """
    base_name, ext = os.path.splitext(output_path)
    counter = 1
    current_output_path = output_path

    while True:
        try:
            # Acquire process-level lock before accessing the file
            with get_docx_lock(current_output_path):
                logger.info(f"Acquired write lock for: {os.path.basename(current_output_path)}")

                if os.path.exists(current_output_path):
                    try:
                        doc = Document(current_output_path)
                        doc.add_page_break()
                        logger.info(f"Appending to existing file: {current_output_path}")
                    except Exception:
                        raise PermissionError("Cannot open existing file.")
                else:
                    doc = Document()
                    logger.info(f"Creating new file: {current_output_path}")

                # Apply styles
                style = doc.styles['Normal']
                style.font.name = 'Times New Roman'
                style.font.size = Pt(12)
                style.paragraph_format.line_spacing = 1.0

                # Count existing depositions for numbering
                next_num = 1
                for para in doc.paragraphs:
                    if re.match(r'^\d+\.\tDeposition of', para.text):
                        next_num += 1

                # Title paragraph with numbering
                p = doc.add_paragraph()
                p.paragraph_format.line_spacing = 1.0
                p.paragraph_format.left_indent = Inches(1.5)
                p.paragraph_format.first_line_indent = Inches(-0.5)
                tab_stops = p.paragraph_format.tab_stops
                tab_stops.add_tab_stop(Inches(1.5))

                run_num = p.add_run(f"{next_num}.\t")
                run_num.bold = True
                run_num.font.name = 'Times New Roman'
                run_num.font.size = Pt(12)

                # Format title with deponent name
                title_name = deponent_name if deponent_name else "Unknown Deponent"
                run_title = p.add_run(f"Deposition of {title_name}")
                run_title.bold = True
                run_title.underline = True
                run_title.font.name = 'Times New Roman'
                run_title.font.size = Pt(12)

                # Add date if available
                if deposition_date:
                    p2 = doc.add_paragraph()
                    p2.paragraph_format.line_spacing = 1.0
                    run_date = p2.add_run(f"Date: {deposition_date}")
                    run_date.font.name = 'Times New Roman'
                    run_date.font.size = Pt(12)
                    run_date.italic = True

                # Add summary content
                add_markdown_to_doc(doc, content)

                doc.save(current_output_path)
                logger.output_file(current_output_path)
                return True

        except LockTimeoutError as e:
            logger.error(f"Lock timeout waiting for file: {e}")
            return False

        except (PermissionError, IOError) as e:
            logger.warning(f"File locked by app: {current_output_path}. Trying next version.")
            counter += 1
            current_output_path = f"{base_name} v.{counter}{ext}"

            if counter > 10:
                logger.error("Failed to save after 10 attempts.")
                return False

        except Exception as e:
            logger.error(f"Error saving to DOCX: {e}")
            return False


def get_output_directory(input_path: str) -> str:
    """Determine the output directory based on input path."""
    parts = input_path.split(os.sep)
    output_dir = None
    case_root_parts = None

    # Priority 1: Find folder starting with exactly 3 digits (Case Folder)
    for i in range(len(parts) - 1, -1, -1):
        if re.match(r'^\d{3}(\D|$)', parts[i]):
            case_root_parts = parts[:i+1]
            break

    # Priority 2: Standard "Current Clients" structure
    if not case_root_parts:
        for i, part in enumerate(parts):
            if part.lower() == "current clients":
                if i + 2 < len(parts):
                    case_root_parts = parts[:i+3]
                break

    if case_root_parts:
        output_dir = os.sep.join(case_root_parts + ["NOTES", "AI OUTPUT"])

    if not output_dir:
        # Fallback 1: NOTES already in path
        for i in range(len(parts) - 1, -1, -1):
            if parts[i].upper() == "NOTES":
                output_dir = os.path.join(os.sep.join(parts[:i+1]), "AI OUTPUT")
                break

    if not output_dir:
        # Fallback 2: Sibling NOTES folder
        input_dir = os.path.dirname(input_path)
        parent_dir = os.path.dirname(input_dir)
        output_dir = os.path.join(parent_dir, "NOTES", "AI OUTPUT")

    return output_dir


# =============================================================================
# Main Entry Point
# =============================================================================

def main():
    """Main entry point. Accepts --phase=topics <input> or --phase=summary <session_path>."""
    if len(sys.argv) < 2:
        print("Error: No file path or session path provided.", flush=True)
        sys.exit(1)

    args = sys.argv[1:]

    # Parse --phase argument (default: topics)
    phase = "topics"
    positional = []
    for a in args:
        if a.startswith("--phase="):
            phase = a.split("=", 1)[1].strip().lower()
        else:
            positional.append(a)

    if not positional:
        print("Error: missing input path.", flush=True)
        sys.exit(1)

    # Handle quoted paths with spaces (preserve existing behavior)
    combined = " ".join(positional).strip().strip('"').strip("'")
    if os.path.exists(combined):
        target = combined
    else:
        target = positional[0]
    target = os.path.abspath(target.strip().strip('"').strip("'"))

    if not os.path.exists(target):
        print(f"Error: path not found: {target}", flush=True)
        sys.exit(1)

    file_number = extract_file_number(target) if phase == "topics" else None
    logger = AgentLogger("Deposition", file_number=file_number)

    if phase == "topics":
        # Directory inputs are not supported in the interactive flow — the
        # detached subprocesses would print AWAITING_INPUT: to nowhere and the
        # iCharlotte UI would never see them. Send single files instead.
        if os.path.isdir(target):
            logger.error(
                "Directory inputs are not supported by the interactive deposition agent. "
                "Run one file at a time, or use the iCharlotte UI which spawns per-file agents."
            )
            sys.exit(2)

        ok = process_topics(target, logger)
        sys.exit(0 if ok else 1)

    if phase == "summary":
        ok = process_summary(target, logger)
        sys.exit(0 if ok else 1)

    print(f"Error: unknown --phase value: {phase}", flush=True)
    sys.exit(2)


if __name__ == '__main__':
    main()
