"""
Summarize Discovery Agent for iCharlotte

Processes discovery response documents with LLM pipeline:
1. Extraction + Summary Pass (parallel) - Extract responses and create summary
2. Cross-Check Pass - Verify completeness

Features:
- Parallel LLM calls for extraction and summary
- Dynamic OCR threshold for text extraction
- Party name extraction from document content
- Document type classification (FROG, SROG, RFP, RFA)
- Auto-consolidation of multiple discovery sets
- Structured progress reporting for UI integration
- Multi-provider LLM fallback
"""

import os
import sys
import re
import subprocess
import datetime
import time
import tempfile
import shutil
from concurrent.futures import ThreadPoolExecutor, as_completed
from docx import Document
from docx.shared import Pt, Inches

# Add parent directory to path for imports
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

# Import shared infrastructure
from icharlotte_core.document_processor import (
    DocumentProcessor, OCRConfig,
    DocumentTypeClassifier, PartyExtractor
)
from icharlotte_core.agent_logger import AgentLogger, create_legacy_log_event
from icharlotte_core.llm_config import LLMCaller, LLMConfig
from icharlotte_core.memory_monitor import MemoryMonitor
from icharlotte_core.exceptions import (
    ExtractionError, LLMError, PassFailedError,
    ExtractionPassError, SummaryPassError, CrossCheckPassError,
    MemoryLimitError, ValidationError
)
from icharlotte_core.docx_writer import get_docx_lock, LockTimeoutError

# Import Case Data Manager
try:
    from case_data_manager import CaseDataManager
except ImportError:
    sys.path.append(os.path.join(os.getcwd(), 'Scripts'))
    from case_data_manager import CaseDataManager

# Import Document Registry for classification
try:
    from document_registry import classify_and_register, DocumentRegistry
except ImportError:
    from Scripts.document_registry import classify_and_register, DocumentRegistry


# =============================================================================
# Configuration
# =============================================================================

SCRIPTS_DIR = os.path.dirname(os.path.abspath(__file__))
PROMPT_FILE = os.path.join(SCRIPTS_DIR, "SUMMARIZE_DISCOVERY_PROMPT.txt")
CONSOLIDATE_PROMPT_FILE = os.path.join(SCRIPTS_DIR, "CONSOLIDATE_DISCOVERY_PROMPT.txt")
EXTRACTION_PROMPT_FILE = os.path.join(SCRIPTS_DIR, "EXTRACTION_PASS_PROMPT.txt")
CROSS_CHECK_PROMPT_FILE = os.path.join(SCRIPTS_DIR, "CROSS_CHECK_PROMPT.txt")

# Legacy log file for backward compatibility
LEGACY_LOG_FILE = r"C:\GeminiTerminal\Summarize_Discovery_activity.log"

# Discovery file filters
SKIP_KEYWORDS = ["rfp", "prod", "production", "rfa", "admission"]


# =============================================================================
# Document Output Functions
# =============================================================================

def add_section_divider(doc, section_title):
    """Adds a visual divider with section title to the document."""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.0

    divider_line = "" * 80
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(divider_line)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.0
    p.alignment = 1  # Center
    run = p.add_run(section_title)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(divider_line)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)

    doc.add_paragraph()


def add_markdown_to_doc(doc, content):
    """Parses basic Markdown and adds it to the docx Document."""
    lines = content.split('\n')

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        # Headings
        if stripped.startswith('#'):
            text = stripped.lstrip('#').strip()
            if not text.endswith('.'):
                text += "."

            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.first_line_indent = Inches(0.5)
            run = p.add_run(text)
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            continue

        # List items
        if stripped.startswith('* ') or stripped.startswith('- '):
            text = stripped[2:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.first_line_indent = Inches(-0.25)

            run = p.add_run("\t")
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    r = p.add_run(part[2:-2])
                    r.bold = True
                else:
                    r = p.add_run(part)
                r.font.name = 'Times New Roman'
                r.font.size = Pt(12)
            continue

        # Normal text
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.first_line_indent = Inches(0.5)

        parts = stripped.split('**')
        for i, part in enumerate(parts):
            run = p.add_run(part)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            if i % 2 == 1:
                run.bold = True


def save_to_docx(extraction_content: str, summary_content: str, output_path: str,
                 title_text: str, discovery_type: str, logger: AgentLogger) -> bool:
    """Saves both extraction and summary content to a DOCX file with process-safe locking."""
    base_name, ext = os.path.splitext(output_path)
    counter = 1
    current_output_path = output_path

    while True:
        try:
            # Acquire process-level lock before accessing the file
            with get_docx_lock(current_output_path):
                logger.info(f"Acquired write lock for: {os.path.basename(current_output_path)}")

                if os.path.exists(current_output_path):
                    try:
                        doc = Document(current_output_path)
                        logger.info(f"Appending to existing file: {current_output_path}")
                    except Exception:
                        raise PermissionError("Cannot open existing file.")
                else:
                    doc = Document()
                    logger.info(f"Creating new file: {current_output_path}")

                # Apply styles
                style = doc.styles['Normal']
                style.font.name = 'Times New Roman'
                style.font.size = Pt(12)
                style.paragraph_format.line_spacing = 1.0

                # Count existing subheadings for numbering
                next_num = 1
                for para in doc.paragraphs:
                    if re.match(r'^\d+\.\t.*Responses to', para.text):
                        next_num += 1

                # Determine party name
                party_name = title_text
                for skip in [".pdf", ".docx", "discovery", "responses", "response"]:
                    party_name = party_name.replace(skip, "").replace(skip.upper(), "")
                party_name = party_name.strip(" _-")

                # Subheading
                p = doc.add_paragraph()
                p.paragraph_format.line_spacing = 1.0
                p.paragraph_format.left_indent = Inches(1.5)
                p.paragraph_format.first_line_indent = Inches(-0.5)

                tab_stops = p.paragraph_format.tab_stops
                tab_stops.add_tab_stop(Inches(1.5))

                run_num = p.add_run(f"{next_num}.\t")
                run_num.bold = True
                run_num.font.name = 'Times New Roman'
                run_num.font.size = Pt(12)

                run_title = p.add_run(f"{party_name}'s Responses to {discovery_type}")
                run_title.bold = True
                run_title.underline = True
                run_title.font.name = 'Times New Roman'
                run_title.font.size = Pt(12)

                # Section A: Extracted Responses
                add_section_divider(doc, "SECTION A: EXTRACTED RESPONSES")
                add_markdown_to_doc(doc, extraction_content)

                # Section B: Narrative Summary
                add_section_divider(doc, "SECTION B: NARRATIVE SUMMARY")
                add_markdown_to_doc(doc, summary_content)

                doc.save(current_output_path)
                logger.output_file(current_output_path)
                return True

        except LockTimeoutError as e:
            logger.error(f"Lock timeout waiting for file: {e}")
            return False

        except (PermissionError, IOError) as e:
            logger.warning(f"File locked by app: {current_output_path}. Trying next version.")
            counter += 1
            current_output_path = f"{base_name} v.{counter}{ext}"

            if counter > 10:
                logger.error("Failed to save after 10 attempts.")
                return False

        except Exception as e:
            logger.error(f"Error saving to DOCX: {e}")
            return False


# =============================================================================
# Party Name Matching
# =============================================================================

def normalize_party_name(name: str) -> str:
    """Normalizes party name for comparison."""
    titles = ["DEFENDANT", "PLAINTIFF", "RESPONDENT", "CROSS-DEFENDANT",
              "CROSS-COMPLAINANT", "THE", "AND"]
    cleaned = name.upper()
    for title in titles:
        cleaned = cleaned.replace(title, "")
    cleaned = re.sub(r"[^A-Z0-9\s]", "", cleaned)
    return " ".join(cleaned.split())


def find_existing_party_file(output_dir: str, new_party_name: str, logger: AgentLogger) -> str:
    """Checks if a file already exists for this party (fuzzy match)."""
    normalized_new = normalize_party_name(new_party_name)
    if len(normalized_new) < 3:
        return None

    try:
        existing_files = [f for f in os.listdir(output_dir)
                         if f.startswith("Discovery_Responses_") and f.endswith(".docx")]
    except OSError:
        return None

    for file in existing_files:
        name_part = file[len("Discovery_Responses_"):-5]
        existing_party_raw = name_part.replace("_", " ")
        normalized_existing = normalize_party_name(existing_party_raw)

        if normalized_new in normalized_existing or normalized_existing in normalized_new:
            logger.info(f"Matched party '{new_party_name}' to existing file '{file}'")
            return file

    return None


def extract_party_from_content(text: str, logger: AgentLogger) -> tuple:
    """
    Extract party name and type from document content.

    Uses multiple strategies:
    1. Regex patterns for common formats
    2. PartyExtractor class
    3. Filename fallback

    Returns:
        Tuple of (party_name, party_type)
    """
    # Strategy 1: Look for common caption patterns
    patterns = [
        r"(Plaintiff|Defendant)\s+([A-Z][A-Za-z]+(?:\s+[A-Z][A-Za-z]+)*)'?s?\s+Response",
        r"Response[s]?\s+(?:of|by)\s+(Plaintiff|Defendant)\s+([A-Z][A-Za-z]+(?:\s+[A-Z][A-Za-z]+)*)",
        r"PROPOUNDING\s+PARTY:\s*(.*?)(?:\n|$)",
        r"RESPONDING\s+PARTY:\s*(.*?)(?:\n|$)",
    ]

    for pattern in patterns:
        match = re.search(pattern, text[:5000], re.IGNORECASE)
        if match:
            groups = match.groups()
            if len(groups) >= 2:
                if groups[0].lower() in ('plaintiff', 'defendant'):
                    party_type = groups[0].title()
                    party_name = groups[1].strip()
                else:
                    party_name = groups[0].strip()
                    party_type = "Unknown"

                logger.info(f"Extracted party from content: {party_name} ({party_type})")
                return party_name, party_type

    # Strategy 2: Use PartyExtractor
    party_name = PartyExtractor.extract_party_name(text)
    if party_name:
        logger.info(f"Extracted party via PartyExtractor: {party_name}")
        return party_name, "Unknown"

    return None, None


# =============================================================================
# Consolidation
# =============================================================================

def consolidate_file(file_path: str, logger: AgentLogger, show_diff: bool = False) -> bool:
    """Reads the current file content, calls LLM to consolidate, and overwrites."""
    logger.info(f"Starting consolidation for: {file_path}")
    logger.progress(98, f"Consolidating: {os.path.basename(file_path)}")

    # Read existing content
    processor = DocumentProcessor(logger=logger)
    result = processor.extract_text(file_path)

    if not result.success:
        logger.error("Failed to extract text for consolidation.")
        return False

    # Load consolidation prompt
    try:
        with open(CONSOLIDATE_PROMPT_FILE, "r", encoding="utf-8") as f:
            prompt = f.read()
    except Exception as e:
        logger.error(f"Error reading consolidation prompt: {e}")
        return False

    # Call LLM
    logger.progress(99, "Running consolidation LLM pass...")
    llm_caller = LLMCaller(logger=logger)
    consolidated = llm_caller.call(prompt, result.text, task_type="summary")

    if not consolidated:
        logger.error("LLM failed to generate consolidated summary.")
        return False

    # Show diff if requested
    if show_diff:
        logger.info("Consolidation changes:")
        # Simple diff - just show length change
        orig_len = len(result.text)
        new_len = len(consolidated)
        logger.info(f"  Original: {orig_len} chars, Consolidated: {new_len} chars")

    # Overwrite file
    try:
        doc = Document()
        logger.info(f"Overwriting with consolidated content: {file_path}")

        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(12)
        style.paragraph_format.line_spacing = 1.0

        add_markdown_to_doc(doc, consolidated)
        doc.save(file_path)
        logger.info(f"Successfully consolidated: {file_path}")
        return True

    except Exception as e:
        logger.error(f"Error overwriting DOCX: {e}")
        return False


# =============================================================================
# Main Processing Functions
# =============================================================================

def process_document(input_path: str, logger: AgentLogger, report_dir: str = None) -> bool:
    """
    Process a single discovery document through the three-pass pipeline.

    Pipeline:
    1. Text Extraction
    2. Pass 1: Extraction + Summary (parallel)
    3. Pass 2: Cross-Check Verification
    4. Save to DOCX
    5. Auto-consolidate if not in batch mode

    Args:
        input_path: Path to the document.
        logger: AgentLogger instance.
        report_dir: Optional directory for batch consolidation.

    Returns:
        True if successful.
    """
    # Initialize components
    memory_monitor = MemoryMonitor(warn_threshold_mb=1500, abort_threshold_mb=2000, logger=logger.info)
    llm_caller = LLMCaller(logger=logger)
    total_passes = 3  # Extraction+Summary (parallel), Cross-Check, Save

    # Progress tracking - define stage weights (total = 100%)
    # Text extraction: 5-15%, LLM calls: 15-75%, Cross-check: 75-90%, Save: 90-100%
    logger.progress(2, "Initializing document processing...")

    # Load prompts
    prompts = {}
    for name, path in [("extraction", EXTRACTION_PROMPT_FILE),
                       ("summary", PROMPT_FILE),
                       ("cross_check", CROSS_CHECK_PROMPT_FILE)]:
        try:
            with open(path, "r", encoding="utf-8") as f:
                prompts[name] = f.read()
        except Exception as e:
            if name == "cross_check":
                logger.warning(f"Could not load cross-check prompt: {e}")
                prompts[name] = None
            else:
                logger.error(f"Error reading {name} prompt: {e}")
                return False

    logger.progress(5, "Prompts loaded, starting text extraction...")

    # ==========================================================================
    # Pass 0: Text Extraction (5% - 15%)
    # ==========================================================================
    logger.pass_start("Extraction", 1, total_passes)
    logger.progress(7, "Reading document and extracting text...")

    try:
        with memory_monitor.track_operation("Text Extraction"):
            processor = DocumentProcessor(
                ocr_config=OCRConfig(adaptive=True),
                logger=logger
            )
            logger.progress(10, "Running text extraction (OCR if needed)...")
            result = processor.extract_with_dynamic_ocr(input_path)

            if not result.success:
                raise ExtractionError(f"Failed to extract text: {result.error}",
                                      file_path=input_path)

            text = result.text
            logger.progress(15, f"Extracted {result.char_count} chars from {result.page_count} pages")
            logger.info(f"Extracted {result.char_count} chars from {result.page_count} pages")

    except Exception as e:
        logger.pass_failed("Extraction", str(e), recoverable=False)
        return False

    logger.pass_complete("Extraction", success=True)

    # ==========================================================================
    # Document Classification (15% - 18%)
    # ==========================================================================
    logger.progress(16, "Classifying document type...")
    doc_type, confidence = DocumentTypeClassifier.classify(text, os.path.basename(input_path))
    discovery_type = DocumentTypeClassifier.get_discovery_type(text, os.path.basename(input_path))

    if discovery_type:
        logger.info(f"Classified as: {discovery_type} (confidence: {confidence:.2f})")
    else:
        discovery_type = "Written Discovery"
        logger.info(f"Document type: {doc_type} (confidence: {confidence:.2f})")

    # Save classified type for registry (before LLM parsing can overwrite it)
    classified_discovery_type = discovery_type
    logger.progress(18, f"Document classified as {discovery_type}")

    # ==========================================================================
    # Pass 1 & 2: Extraction + Summary (Parallel) (18% - 70%)
    # ==========================================================================
    logger.pass_start("LLM Extraction + Summary", 2, total_passes)
    logger.progress(20, "Starting parallel LLM extraction and summary...")
    logger.info("Running extraction and summary passes in parallel...")

    extraction_result = None
    summary_result = None
    extraction_error = None
    summary_error = None
    extraction_done = False
    summary_done = False

    def run_extraction():
        """Run extraction pass."""
        with memory_monitor.track_operation("Extraction Pass"):
            result = llm_caller.call(prompts["extraction"], text, task_type="extraction")
            if not result:
                raise ExtractionPassError("LLM returned empty response")
            return result

    def run_summary():
        """Run summary pass."""
        with memory_monitor.track_operation("Summary Pass"):
            result = llm_caller.call(prompts["summary"], text, task_type="summary")
            if not result:
                raise SummaryPassError("LLM returned empty response")
            # Strip extraction layer if present
            result = re.sub(r'<extraction_layer>.*?</extraction_layer>',
                            '', result, flags=re.DOTALL).strip()
            return result

    try:
        with ThreadPoolExecutor(max_workers=2) as executor:
            extraction_future = executor.submit(run_extraction)
            summary_future = executor.submit(run_summary)

            # Report progress while waiting
            logger.progress(30, "Waiting for LLM extraction response...")

            # Wait for both to complete
            for future in as_completed([extraction_future, summary_future]):
                try:
                    if future == extraction_future:
                        extraction_result = future.result()
                        extraction_done = True
                        # Report progress based on which completed
                        if summary_done:
                            logger.progress(68, "Both LLM passes complete")
                        else:
                            logger.progress(50, "Extraction complete, waiting for summary...")
                        logger.info(f"Extraction complete: {len(extraction_result)} chars")
                    else:
                        summary_result = future.result()
                        summary_done = True
                        if extraction_done:
                            logger.progress(68, "Both LLM passes complete")
                        else:
                            logger.progress(50, "Summary complete, waiting for extraction...")
                        logger.info(f"Summary complete: {len(summary_result)} chars")
                except Exception as e:
                    if future == extraction_future:
                        extraction_error = e
                    else:
                        summary_error = e

    except Exception as e:
        logger.pass_failed("LLM Extraction + Summary", str(e), recoverable=False)
        return False

    # Check for errors
    if extraction_error:
        logger.pass_failed("LLM Extraction", str(extraction_error), recoverable=False)
        return False

    if summary_error:
        logger.pass_failed("Summary", str(summary_error), recoverable=False)
        return False

    logger.progress(70, "LLM extraction and summary passes completed")
    logger.pass_complete("LLM Extraction + Summary", success=True)

    # ==========================================================================
    # Pass 2: Cross-Check (with retry on failure) (70% - 88%)
    # ==========================================================================
    final_summary = summary_result
    cross_check_failed = False

    if prompts["cross_check"]:
        logger.pass_start("CrossCheck", 3, total_passes)
        logger.progress(72, "Starting cross-check verification pass...")

        try:
            with memory_monitor.track_operation("Cross-Check Pass"):
                # Format cross-check prompt
                cross_prompt = f"""{prompts["cross_check"]}

=== PASS 1 (EXTRACTED RESPONSES) ===
{extraction_result}

=== PASS 2 (NARRATIVE SUMMARY) ===
{summary_result}
"""
                logger.progress(75, "Sending cross-check to LLM...")
                cross_check_result = llm_caller.call(cross_prompt, "", task_type="cross_check")

                if cross_check_result:
                    final_summary = cross_check_result
                    logger.progress(85, "Cross-check verification complete")
                    logger.info("Cross-check completed successfully")
                else:
                    raise CrossCheckPassError("Cross-check returned empty")

        except Exception as e:
            logger.pass_failed("CrossCheck", str(e), recoverable=True)
            cross_check_failed = True

            # Retry with different model sequence
            logger.progress(80, "Cross-check failed, retrying with fallback model...")
            logger.info("Retrying cross-check with fallback models...")
            try:
                cross_check_result = llm_caller.call(cross_prompt, "", task_type="quick")
                if cross_check_result:
                    final_summary = cross_check_result
                    cross_check_failed = False
                    logger.progress(85, "Cross-check succeeded on retry")
                    logger.info("Cross-check succeeded on retry")
            except Exception:
                pass

            if cross_check_failed:
                logger.progress(85, "Cross-check skipped, using unverified summary")
                logger.warning("Cross-check failed. Using unverified summary.")

        logger.progress(88, "Cross-check pass finished")
        logger.pass_complete("CrossCheck", success=not cross_check_failed,
                            details="with warnings" if cross_check_failed else None)
    else:
        # No cross-check prompt, skip to 88%
        logger.progress(88, "Cross-check skipped (no prompt configured)")

    # ==========================================================================
    # Parse Metadata (88% - 90%)
    # ==========================================================================
    logger.progress(89, "Parsing metadata from results...")
    responding_party = "Unknown_Party"

    # Parse from extraction result
    extraction_lines = extraction_result.split('\n')
    lines_to_remove = []

    for i in range(min(10, len(extraction_lines))):
        line = extraction_lines[i].strip()
        if line.startswith("RESPONDING_PARTY:"):
            responding_party = line.replace("RESPONDING_PARTY:", "").strip()
            lines_to_remove.append(i)
        elif line.startswith("DISCOVERY_TYPE:"):
            parsed_type = line.replace("DISCOVERY_TYPE:", "").strip()
            if parsed_type:
                discovery_type = parsed_type
            lines_to_remove.append(i)

    for i in sorted(lines_to_remove, reverse=True):
        extraction_lines.pop(i)
    extraction_content = "\n".join(extraction_lines).strip()

    # Fallback: Extract party from document content
    if responding_party == "Unknown_Party":
        extracted_party, party_type = extract_party_from_content(text, logger)
        if extracted_party:
            responding_party = extracted_party
            if party_type and party_type != "Unknown":
                responding_party = f"{party_type} {extracted_party}"

    # Clean summary content
    summary_lines = final_summary.split('\n')
    lines_to_remove = []

    for i in range(min(10, len(summary_lines))):
        line = summary_lines[i].strip()
        if line.startswith("RESPONDING_PARTY:") or line.startswith("DISCOVERY_TYPE:"):
            lines_to_remove.append(i)

    for i in sorted(lines_to_remove, reverse=True):
        summary_lines.pop(i)
    summary_content = "\n".join(summary_lines).strip()

    if responding_party == "Unknown_Party":
        logger.warning("Could not determine responding party.")

    logger.progress(90, f"Metadata parsed: {responding_party}")

    # ==========================================================================
    # Save to Case Data (90% - 92%)
    # ==========================================================================
    logger.progress(91, "Saving to case database...")
    try:
        data_manager = CaseDataManager()
        file_num = extract_file_number(input_path)

        if file_num:
            base_name = os.path.splitext(os.path.basename(input_path))[0]
            clean_var_name = re.sub(r"[^a-zA-Z0-9_]", "_", base_name.lower())
            var_key = f"discovery_summary_{clean_var_name}"

            # Add document type as tag
            tags = ["Discovery"]
            if discovery_type:
                tags.append(discovery_type.replace(" ", "_"))

            logger.info(f"Saving to case data: {var_key} for {file_num}")
            data_manager.save_variable(file_num, var_key, summary_content,
                                       source="discovery_agent", extra_tags=tags)
    except Exception as e:
        logger.warning(f"Could not save to case data: {e}")

    logger.progress(92, "Case data saved")

    # ==========================================================================
    # Save Output (92% - 96%)
    # ==========================================================================
    logger.progress(93, "Preparing output file...")
    output_dir = get_output_directory(input_path)

    if not os.path.exists(output_dir):
        try:
            os.makedirs(output_dir)
            logger.info(f"Created directory: {output_dir}")
        except Exception as e:
            logger.error(f"Error creating directory: {e}")
            return False

    # Check for existing party file
    existing_file = find_existing_party_file(output_dir, responding_party, logger)

    if existing_file:
        output_filename = existing_file
    else:
        safe_party_name = re.sub(r'[\\/*?:"<>|]', "", responding_party).replace(" ", "_")
        output_filename = f"Discovery_Responses_{safe_party_name}.docx"

    output_file = os.path.join(output_dir, output_filename)

    logger.progress(94, f"Saving to {output_filename}...")
    save_success = save_to_docx(extraction_content, summary_content, output_file,
                                responding_party, discovery_type, logger)

    if not save_success:
        return False

    logger.progress(96, "Document saved successfully")

    # ==========================================================================
    # Register Document in Registry (96% - 97%)
    # ==========================================================================
    logger.progress(96, "Registering document...")
    try:
        if file_num:
            base_name = os.path.splitext(os.path.basename(input_path))[0]

            # Map classified discovery_type to registry document type
            # Use classified_discovery_type (from classifier), not discovery_type (may be overwritten by LLM)
            doc_type_map = {
                "Form Interrogatories": "Form Interrogatories - Responses",
                "Special Interrogatories": "Special Interrogatories - Responses",
                "Request for Admissions": "Request for Admissions - Responses",
                "Request for Production": "Request for Production - Responses",
                "Requests for Admission": "Request for Admissions - Responses",
                "Requests for Production": "Request for Production - Responses",
            }
            registry_doc_type = doc_type_map.get(classified_discovery_type, "Form Interrogatories - Responses")

            # Generate standardized document name
            from document_registry import DocumentClassifier
            classifier = DocumentClassifier(logger=logger)
            standardized_name = classifier.generate_name(summary_content, registry_doc_type, fallback_name=base_name)
            logger.info(f"Generated name: '{standardized_name}'")

            registry = DocumentRegistry()
            registry.register_document(
                file_number=file_num,
                name=standardized_name,
                document_type=registry_doc_type,
                source_path=input_path,
                summary_location=output_file,
                agent="summarize_discovery",
                char_count=len(summary_content)
            )
            logger.info(f"Registered document as: {registry_doc_type}")
    except Exception as e:
        logger.warning(f"Could not register document: {e}")

    logger.progress(97, "Document registered")

    # ==========================================================================
    # Consolidation (97% - 100%)
    # ==========================================================================
    if report_dir:
        # Batch mode: Report output file for later consolidation
        logger.progress(98, "Reporting to batch consolidation...")
        import random
        fname = f"worker_{os.getpid()}_{random.randint(0, 1000)}.txt"
        try:
            with open(os.path.join(report_dir, fname), "w") as f:
                f.write(output_file)
            logger.info(f"Reported output file to batch: {output_file}")
        except Exception as e:
            logger.error(f"Failed to write report file: {e}")
        logger.progress(100, "Processing complete")
    else:
        # Single file mode: Auto-consolidate immediately
        logger.progress(98, "Running auto-consolidation...")
        consolidate_file(output_file, logger, show_diff=True)
        logger.progress(100, "Consolidation complete")

    return True


def extract_file_number(path: str) -> str:
    """Extract file number from path.

    Handles paths like:
    - Z:\\Shared\\Current Clients\\3800- NATIONWIDE\\3850\\084 - Dudash\\...
    - Paths containing literal "3850.084"
    """
    # Standard pattern: 1234.567 (with literal dot)
    match = re.search(r"(\d{4}\.\d{3})", path)
    if match:
        return match.group(1)

    # Parse from directory structure
    # Look for pattern: 4-digit folder followed by 3-digit folder
    parts = os.path.normpath(path).split(os.sep)
    try:
        # Find consecutive folders matching client/matter pattern
        for i in range(len(parts) - 1):
            # Look for 4-digit client folder (e.g., "3850" or "3850 - Name")
            client_match = re.match(r"^(\d{4})(?:\D|$)", parts[i])
            if client_match:
                # Check if next folder is 3-digit matter (e.g., "084" or "084 - Name")
                matter_match = re.match(r"^(\d{3})(?:\D|$)", parts[i + 1])
                if matter_match:
                    return f"{client_match.group(1)}.{matter_match.group(1)}"

        # Fallback: Look for "Current Clients" structure
        # Pattern: Current Clients\GroupFolder\ClientFolder\MatterFolder
        cc_index = -1
        for i, part in enumerate(parts):
            if part.lower() == "current clients":
                cc_index = i
                break

        if cc_index != -1 and cc_index + 3 < len(parts):
            # Skip the group folder (e.g., "3800- NATIONWIDE"), use folders at +2 and +3
            client_folder = parts[cc_index + 2]  # e.g., "3850"
            matter_folder = parts[cc_index + 3]  # e.g., "084 - Dudash"

            client_code = re.match(r"^(\d{4})(?:\D|$)", client_folder)
            matter_code = re.match(r"^(\d{3})(?:\D|$)", matter_folder)

            if client_code and matter_code:
                return f"{client_code.group(1)}.{matter_code.group(1)}"

            # Try original positions if new positions don't match
            client_folder = parts[cc_index + 1]
            matter_folder = parts[cc_index + 2]

            client_code = re.match(r"^\d+", client_folder)
            matter_code = re.match(r"^\d+", matter_folder)

            if client_code and matter_code:
                file_num = f"{client_code.group(0)}.{matter_code.group(0)}"
                # Validate format (4 digits.3 digits)
                if re.match(r"^\d{4}\.\d{3}$", file_num):
                    return file_num
    except Exception:
        pass

    return None


def get_output_directory(input_path: str) -> str:
    """Determine the output directory based on input path."""
    parts = input_path.split(os.sep)
    output_dir = None
    case_root_parts = None

    # Priority 1: 3-digit folder pattern
    for i in range(len(parts) - 1, -1, -1):
        if re.match(r'^\d{3}(\D|$)', parts[i]):
            case_root_parts = parts[:i+1]
            break

    # Priority 2: Current Clients structure
    if not case_root_parts:
        for i, part in enumerate(parts):
            if part.lower() == "current clients":
                if i + 2 < len(parts):
                    case_root_parts = parts[:i+3]
                break

    if case_root_parts:
        output_dir = os.sep.join(case_root_parts + ["NOTES", "AI OUTPUT"])

    if not output_dir:
        for i in range(len(parts) - 1, -1, -1):
            if parts[i].upper() == "NOTES":
                output_dir = os.path.join(os.sep.join(parts[:i+1]), "AI OUTPUT")
                break

    if not output_dir:
        input_dir = os.path.dirname(input_path)
        parent_dir = os.path.dirname(input_dir)
        output_dir = os.path.join(parent_dir, "NOTES", "AI OUTPUT")

    return output_dir


def run_concurrently(file_paths: list, extra_args: list = None, batch_size: int = 3,
                     logger: AgentLogger = None):
    """Runs subprocesses for file paths with limited concurrency."""
    if logger:
        logger.info(f"Starting batch processing: {len(file_paths)} files, batch size {batch_size}")

    processes = []

    for path in file_paths:
        while len(processes) >= batch_size:
            processes = [p for p in processes if p.poll() is None]
            if len(processes) >= batch_size:
                time.sleep(1)

        cmd = [sys.executable, sys.argv[0], path]
        if extra_args:
            cmd.extend(extra_args)

        creation_flags = 0x08000000 if os.name == 'nt' else 0

        try:
            if logger:
                logger.info(f"Spawning subprocess for: {os.path.basename(path)}")
            p = subprocess.Popen(cmd, creationflags=creation_flags)
            processes.append(p)
        except Exception as e:
            if logger:
                logger.error(f"Failed to spawn subprocess: {e}")

    while processes:
        processes = [p for p in processes if p.poll() is None]
        if processes:
            time.sleep(1)

    if logger:
        logger.info("All batch processes completed.")


def main():
    """Main entry point."""
    # Parse arguments
    args = sys.argv[1:]
    report_dir = None
    clean_args = []

    i = 0
    while i < len(args):
        arg = args[i]
        if arg == "--report-dir":
            if i + 1 < len(args):
                report_dir = args[i+1]
                i += 2
            else:
                i += 1
        else:
            if arg.strip():
                clean_args.append(arg)
            i += 1

    if not clean_args:
        print("Error: No file paths provided.", flush=True)
        sys.exit(1)

    # Handle split paths
    if len(clean_args) > 1:
        combined_path = " ".join(clean_args).strip().strip('"').strip("'")
        if os.path.exists(combined_path) and not os.path.exists(clean_args[0]):
            file_paths = [combined_path]
        else:
            file_paths = clean_args
    else:
        file_paths = clean_args

    # Extract file number for logger
    file_number = extract_file_number(file_paths[0]) if file_paths else None

    logger = AgentLogger("Discovery", file_number=file_number)

    # Dispatcher mode: multiple files
    if len(file_paths) > 1:
        logger.info(f"Dispatcher mode: {len(file_paths)} files")

        use_temp_dir = False
        if not report_dir:
            report_dir = tempfile.mkdtemp()
            use_temp_dir = True
            logger.info(f"Created temp report directory: {report_dir}")

        run_concurrently(file_paths, extra_args=["--report-dir", report_dir],
                         batch_size=3, logger=logger)

        # Auto-consolidation
        logger.info("Processing consolidation reports...")
        unique_outputs = set()

        if os.path.exists(report_dir):
            for f in os.listdir(report_dir):
                if f.startswith("worker_"):
                    try:
                        with open(os.path.join(report_dir, f), 'r') as rf:
                            content = rf.read().strip()
                            if content:
                                unique_outputs.add(content)
                    except Exception:
                        pass

        logger.info(f"Consolidating {len(unique_outputs)} output files...")

        for doc_path in unique_outputs:
            if os.path.exists(doc_path):
                consolidate_file(doc_path, logger, show_diff=True)

        if use_temp_dir:
            try:
                shutil.rmtree(report_dir)
                logger.info("Cleaned up temp directory.")
            except Exception:
                pass

        logger.info("Batch processing complete.")
        sys.exit(0)

    # Single file/directory mode
    input_path = file_paths[0].strip().strip('"').strip("'")
    input_path = os.path.abspath(input_path)

    if not os.path.exists(input_path):
        logger.error(f"File not found: {input_path}")
        sys.exit(1)

    logger.info(f"Starting agent for: {input_path}")

    # Directory handling
    if os.path.isdir(input_path):
        logger.info(f"Directory mode: scanning {input_path}")
        files_to_process = []

        for root, _, files in os.walk(input_path):
            for file in files:
                if file.lower().endswith((".pdf", ".docx")):
                    if "Discovery_Response" in file:
                        continue

                    lower_file = file.lower()
                    if any(kw in lower_file for kw in SKIP_KEYWORDS):
                        logger.info(f"Skipping filtered file: {file}")
                        continue

                    files_to_process.append(os.path.join(root, file))

        if not files_to_process:
            logger.info("No suitable files found in directory.")
            sys.exit(0)

        logger.info(f"Found {len(files_to_process)} files to process.")

        use_temp_dir = False
        if not report_dir:
            report_dir = tempfile.mkdtemp()
            use_temp_dir = True

        run_concurrently(files_to_process, extra_args=["--report-dir", report_dir],
                         batch_size=3, logger=logger)

        # Auto-consolidation
        if use_temp_dir:
            logger.info("Processing consolidation (Directory Mode)...")
            unique_outputs = set()

            if os.path.exists(report_dir):
                for f in os.listdir(report_dir):
                    if f.startswith("worker_"):
                        try:
                            with open(os.path.join(report_dir, f), 'r') as rf:
                                content = rf.read().strip()
                                if content:
                                    unique_outputs.add(content)
                        except Exception:
                            pass

            for doc_path in unique_outputs:
                if os.path.exists(doc_path):
                    consolidate_file(doc_path, logger, show_diff=True)

            try:
                shutil.rmtree(report_dir)
            except Exception:
                pass

        logger.info("Directory processing complete.")
        sys.exit(0)

    # Single file processing
    input_filename_lower = os.path.basename(input_path).lower()
    if any(kw in input_filename_lower for kw in SKIP_KEYWORDS):
        logger.info(f"Skipping filtered file: {input_path}")
        sys.exit(0)

    success = process_document(input_path, logger, report_dir)

    if success:
        logger.info("Agent finished successfully.")
        sys.exit(0)
    else:
        logger.error("Agent finished with errors.")
        sys.exit(1)


if __name__ == '__main__':
    main()
