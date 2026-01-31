import os
import sys
import logging
import datetime
import re
import subprocess
import gc
from google import genai
from docx import Document
from docx.shared import Pt, Inches
from pypdf import PdfReader

# Import Case Data Manager
try:
    from case_data_manager import CaseDataManager
except ImportError:
    sys.path.append(os.path.join(os.getcwd(), 'Scripts'))
    from case_data_manager import CaseDataManager

try:
    import pytesseract
    from pdf2image import convert_from_path
    OCR_AVAILABLE = True
    
    # Windows-specific path configuration
    if os.name == 'nt':
        # Tesseract Path
        tesseract_path = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
        if os.path.exists(tesseract_path):
            pytesseract.pytesseract.tesseract_cmd = tesseract_path
        
        # Poppler Path
        # We found it at C:\Program Files\poppler\Library\bin\pdftoppm.exe
        # pdf2image needs the folder containing the binaries
        poppler_path = r"C:\Program Files\poppler\Library\bin"
        if not os.path.exists(poppler_path):
             # Fallback attempt or standard location
             poppler_path = None
        
        # We will use this variable in extract_text
        POPPLER_PATH = poppler_path
    else:
        POPPLER_PATH = None

except ImportError:
    OCR_AVAILABLE = False
    POPPLER_PATH = None


# --- Configuration ---
# Hardcoded log path as per instructions
LOG_FILE = r"C:\GeminiTerminal\exposure_activity.log"
PROMPT_FILE = r"C:\GeminiTerminal\Scripts\EXPOSURE_PROMPT.txt"

# Set up logging
logging.basicConfig(
    filename=LOG_FILE,
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    datefmt='%Y-%m-%d %H:%M:%S'
)

def log_event(message, level="info"):
    try:
        print(message, flush=True)  # Also print to stdout for debugging if run manually
    except UnicodeEncodeError:
        try:
             print(message.encode(sys.stdout.encoding or 'utf-8', errors='replace').decode(sys.stdout.encoding or 'utf-8'), flush=True)
        except Exception:
             pass  # Silently ignore if stdout is broken
    except OSError:
        pass  # stdout pipe broken (common when running multiple agents)
    if level == "info":
        logging.info(message)
    elif level == "error":
        logging.error(message)
    elif level == "warning":
        logging.warning(message)

def extract_text(file_path):
    """Extracts text from PDF, DOCX, or plain text files. Falls back to OCR for specific PDF pages if text is insufficient."""
    log_event(f"Extracting text from: {file_path}")
    ext = os.path.splitext(file_path)[1].lower()
    text = ""

    try:
        if ext == ".pdf":
            reader = PdfReader(file_path)
            total_pages = len(reader.pages)
            log_event(f"PDF has {total_pages} pages.")
            
            # Helper to convert specific page index to image if needed
            def get_page_image(page_index):
                try:
                    # convert_from_path returns a list of images. We need to fetch just the specific page.
                    # 'first_page' and 'last_page' are 1-based indices in pdf2image
                    # Pass poppler_path if explicitly set (mainly for Windows)
                    images = convert_from_path(file_path, first_page=page_index+1, last_page=page_index+1, poppler_path=POPPLER_PATH)
                    return images[0] if images else None
                except Exception as e:
                    log_event(f"Error converting page {page_index+1} to image: {e}", level="warning")
                    return None

            for i, page in enumerate(reader.pages):
                if i % 10 == 0:
                    gc.collect()

                page_text = page.extract_text() or ""
                
                # Check if this specific page has meaningful text (threshold: 50 chars)
                if len(page_text.strip()) < 50:
                    log_event(f"Page {i+1} has insufficient text ({len(page_text.strip())} chars). Attempting OCR...", level="warning")
                    
                    if OCR_AVAILABLE:
                        image = get_page_image(i)
                        if image:
                            try:
                                ocr_page_text = pytesseract.image_to_string(image)
                                if len(ocr_page_text.strip()) > len(page_text.strip()):
                                    log_event(f"OCR successful for page {i+1}. Extracted {len(ocr_page_text)} chars.")
                                    text += ocr_page_text + "\n"
                                else:
                                    log_event(f"OCR for page {i+1} did not yield better results. Using original.", level="warning")
                                    text += page_text + "\n"
                            except Exception as ocr_e:
                                log_event(f"OCR failed for page {i+1}: {ocr_e}. Using original.", level="error")
                                text += page_text + "\n"
                        else:
                             log_event(f"Could not render image for page {i+1}. Using original.", level="warning")
                             text += page_text + "\n"
                    else:
                        log_event(f"OCR not available. Skipping OCR for page {i+1}.", level="warning")
                        text += page_text + "\n"
                else:
                    # Page text is sufficient
                    text += page_text + "\n"

        elif ext == ".docx":
            doc = Document(file_path)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
        else:
            # Assume plain text for other extensions
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
        
        if not text.strip():
            log_event(f"Warning: Extracted text is empty for {file_path}", level="warning")
            return None
        
        log_event(f"Successfully extracted {len(text)} characters total.")
        log_event(f"Snippet: {text[:100].replace(chr(10), ' ')}...") # Log snippet
        return text

    except Exception as e:
        log_event(f"Error extracting text from {file_path}: {e}", level="error")
        return None

def call_gemini(prompt, text):
    """Calls Gemini API with fallback logic."""
    api_key = os.environ.get("GEMINI_API_KEY")
    if not api_key:
        log_event("Error: GEMINI_API_KEY environment variable not set.", level="error")
        return None

    client = genai.Client(api_key=api_key)

    full_prompt = f"{prompt}\n\nDOCUMENT CONTENT:\n{text}"
    
    # Updated model sequence for liability agent
    model_sequence = [
        "gemini-3-pro-preview", 
        "gemini-2.5-flash"
    ]

    for model_name in model_sequence:
        log_event(f"Attempting to use model: {model_name}")
        try:
            response = client.models.generate_content(
                model=model_name,
                contents=full_prompt
            )
            if response and response.text:
                log_event(f"Success with model: {model_name}")
                return response.text
        except BaseException as e: # Catch ALL exceptions including system exits
            log_event(f"Failed with model {model_name}: {e}", level="warning")
            continue
    
    log_event("Error: All model attempts failed.", level="error")
    return None

def add_markdown_to_doc(doc, content):
    """Parses basic Markdown and adds it to the docx Document with specific formatting."""
    lines = content.split('\n')
    active_paragraph = None
    
    # Debug: Log the first few lines to see what we are parsing
    log_event("--- content preview ---")
    for i, line in enumerate(lines[:10]):
        log_event(f"Line {i}: {repr(line)}")
    log_event("--- end preview ---")

    # Pattern for subheadings like "A. Duty", "**A. Duty**", "A. **Duty**"
    # Matches: (Non-word chars like *)(Letter.)(Any chars)(Text)
    subheading_pattern = re.compile(r'^\W*([A-Z]\.)\s*(.*)')
    
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        
        # Headings: Convert to bold text at start of paragraph, ending with period
        if stripped.startswith('#'):
            log_event(f"Formatting as Heading: {stripped}")
            # Reset active paragraph as this is a new block
            active_paragraph = None
            
            text = stripped.lstrip('#').strip()
            # Remove bold markers if present in heading
            text = text.replace('**', '').replace('__', '')
            
            if not text.endswith('.'):
                text += "."
            p = doc.add_paragraph()
            run = p.add_run(text + " ")
            run.bold = True
            continue
        
        # Subheadings (A. Duty)
        match = subheading_pattern.match(stripped)
        # Verify it's not just a bullet point or random text starting with a letter
        # The regex requires A. at the start (ignoring non-word chars).
        if match:
            log_event(f"Formatting as Subheading: {stripped}")
            active_paragraph = None
            letter = match.group(1) # "A."
            text_part = match.group(2).strip() # "Duty**" or "**Duty**"
            
            # Clean text_part of trailing/leading bold markers
            text_part = text_part.replace('**', '').replace('__', '')
            
            p = doc.add_paragraph()
            
            # Hanging Indent Strategy:
            # - Left Indent (Body): 1.0 inch
            # - First Line Indent: -0.5 inch (pulls the first line back to 0.5)
            # - Tab Stop: 1.0 inch (to align the text start after the letter)
            p.paragraph_format.left_indent = Inches(1.0)
            p.paragraph_format.first_line_indent = Inches(-0.5)
            p.paragraph_format.tab_stops.add_tab_stop(Inches(1.0))
            
            # Run 1: Letter + Tab (Bold)
            r_letter = p.add_run(letter + "\t")
            r_letter.bold = True
            
            # Run 2: Text (Bold and Underlined)
            r = p.add_run(text_part)
            r.bold = True
            r.underline = True
            continue
        
        # Specific formatting for "EVALUATION OF LIABILITY" or "EVALUATION OF EXPOSURE"
        if "EVALUATION OF LIABILITY" in stripped.upper() or "EVALUATION OF EXPOSURE" in stripped.upper():
            log_event(f"Formatting Title: {stripped}")
            active_paragraph = None
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Inches(0)
            run = p.add_run(stripped.replace('**', '').strip())
            run.bold = True
            run.underline = True
            continue

        log_event(f"Formatting as Normal Text: {stripped}")
        
        # List items
        if stripped.startswith('* ') or stripped.startswith('- '):
            active_paragraph = None
            text = stripped[2:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.first_line_indent = Inches(-0.25)
            
            # Use a real bullet character and tab for portability
            run = p.add_run("•\t")
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            
            # Support bold parsing within list items
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    r = p.add_run(part[2:-2])
                    r.bold = True
                else:
                    r = p.add_run(part)
                r.font.name = 'Times New Roman'
                r.font.size = Pt(12)
            continue
        
        # Normal text (with bold support)
        # "Start of each new paragraph should be indented by 0.5"
        if active_paragraph:
            p = active_paragraph
            p.add_run(" ")
        else:
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Inches(0.5)
            active_paragraph = p
        
        # Simple bold parsing: **text**
        parts = re.split(r'(\**.*\**)', stripped)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
            else:
                p.add_run(part)

def save_to_docx(content, output_path, title_text):
    """Saves the content to a DOCX file. Appends if exists. Handles locking."""
    base_name, ext = os.path.splitext(output_path)
    counter = 1
    current_output_path = output_path
    
    while True:
        try:
            if os.path.exists(current_output_path):
                try:
                    doc = Document(current_output_path)
                    doc.add_page_break() # Separator for new entry
                    log_event(f"Appending to existing file: {current_output_path}")
                except Exception:
                    # If opening fails, assume locked/corrupt and force next version
                    raise PermissionError("Cannot open existing file.")
            else:
                doc = Document()
                log_event(f"Creating new file: {current_output_path}")

            # Apply Styles (Times New Roman, Size 12, Single Spaced)
            style = doc.styles['Normal']
            style.font.name = 'Times New Roman'
            style.font.size = Pt(12)
            style.paragraph_format.line_spacing = 1.0
            
            # Update Heading styles to match requirements
            for i in range(1, 10):
                if f'Heading {i}' in doc.styles:
                    h_style = doc.styles[f'Heading {i}']
                    h_style.font.name = 'Times New Roman'
                    h_style.font.size = Pt(12)
                    h_style.paragraph_format.line_spacing = 1.0

            # Update List styles to match requirements
            for style_name in ['List Bullet', 'List Number']:
                if style_name in doc.styles:
                    l_style = doc.styles[style_name]
                    l_style.font.name = 'Times New Roman'
                    l_style.font.size = Pt(12)
                    l_style.paragraph_format.line_spacing = 1.0

            # Title: Name of document, bold, underlined
            p = doc.add_paragraph()
            run = p.add_run(title_text)
            run.bold = True
            run.underline = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

            doc.add_paragraph(f"Generated on: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            
            add_markdown_to_doc(doc, content)
                
            doc.save(current_output_path)
            log_event(f"Saved liability analysis to: {current_output_path}")
            return True

        except (PermissionError, IOError) as e:
            log_event(f"File locked or inaccessible: {current_output_path}. Trying next version. Error: {e}", level="warning")
            counter += 1
            # Format: Original v.2.docx, Original v.3.docx
            current_output_path = f"{base_name} v.{counter}{ext}"
            
            if counter > 10:
                log_event("Failed to save after 10 attempts.", level="error")
                return False
        except Exception as e:
            log_event(f"Error saving to DOCX: {e}", level="error")
            return False

def main():
    if len(sys.argv) < 2:
        log_event("Error: No file path provided.", level="error")
        sys.exit(1)

    # Attempt to handle unquoted paths with spaces by joining all arguments
    input_path = " ".join(sys.argv[1:])
    
    # If the joined path doesn't exist, but the first argument does, 
    # it might be a single file with tricky chars or we're in a directory where 
    # sys.argv[1] was actually correct and the rest were garbage.
    if not os.path.exists(input_path) and os.path.exists(sys.argv[1]):
        input_path = sys.argv[1]
    
    # Handle absolute path conversion
    input_path = os.path.abspath(input_path)

    if not os.path.exists(input_path):
        log_event(f"Error: File not found: {input_path}", level="error")
        sys.exit(1)

    log_event(f"--- Starting Exposure Agent for input: {input_path} ---")

    # --- Directory Handling ---
    if os.path.isdir(input_path):
        log_event(f"Input is a directory. Scanning for files in: {input_path}")
        files_to_process = []
        for root, _, files in os.walk(input_path):
            for file in files:
                if file.lower().endswith(('.pdf', '.docx')):
                    # Exclude the output file itself if it happens to be in the scan path
                    if "AI_OUTPUT" in file:
                        continue
                    files_to_process.append(os.path.join(root, file))
        
        if not files_to_process:
            log_event("No suitable files (.pdf, .docx) found in directory.")
            sys.exit(0)
        
        log_event(f"Found {len(files_to_process)} files to process.")
        
        # Spin up a new CLI instance for each document
        # We run them sequentially to avoid file locking issues on the single output file
        for file_path in files_to_process:
            log_event(f"Spawning subprocess for: {file_path}")
            try:
                # Call this script again with the specific file path
                subprocess.run([sys.executable, sys.argv[0], file_path], check=True)
            except subprocess.CalledProcessError as e:
                log_event(f"Subprocess failed for {file_path}: {e}", level="error")
        
        log_event("--- Directory Processing Complete ---")
        sys.exit(0)
    
    # --- Single File Processing ---

    # 1. Extract Text
    text = extract_text(input_path)
    if not text:
        sys.exit(1)

    # 2. Load Prompt
    try:
        with open(PROMPT_FILE, "r", encoding="utf-8") as f:
            prompt_instruction = f.read()
    except Exception as e:
        log_event(f"Error reading prompt file {PROMPT_FILE}: {e}", level="error")
        sys.exit(1)

    # 3. Call LLM
    summary = call_gemini(prompt_instruction, text)
    if not summary:
        sys.exit(1)

    # Save to Case Data
    data_manager = CaseDataManager()
    file_num_match = re.search(r"(\d{4}\.\d{3})", input_path)
    if file_num_match:
        file_num = file_num_match.group(1)
        
        log_event(f"Saving exposure analysis to case variable for {file_num}")
        data_manager.save_variable(file_num, "exposure_evaluation", summary, source="exposure_agent", extra_tags=["Reports", "Exposure"])
    else:
        log_event("Could not extract file number from path. Skipping variable save.", level="warning")

    # 4. Save Output
    # Strategy: Locate Case Root based on folder naming conventions or "Current Clients" position
    # The Case Root is typically the folder starting with 3 digits (e.g. "084 - Dudash")
    # OR 2 levels deeper than Current Clients: Current Clients/<Client>/<Case>
    
    parts = input_path.split(os.sep)
    output_dir = None
    case_root_parts = None

    # Priority 1: Find folder starting with exactly 3 digits (Case Folder)
    # We scan from end to start to find the deepest valid case folder
    for i in range(len(parts) - 1, -1, -1):
        # Match exactly 3 digits at start, followed by non-digit or end
        if re.match(r'^\d{3}(\D|$)', parts[i]):
            log_event(f"Identified Case Folder by 3-digit pattern: {parts[i]}")
            case_root_parts = parts[:i+1]
            break

    # Priority 2: Standard "Current Clients" structure (Client/Case)
    if not case_root_parts:
        for i, part in enumerate(parts):
            if part.lower() == "current clients":
                # Check if we have enough depth for Client and Case
                if i + 2 < len(parts):
                    log_event("Identified Case Folder by Standard Structure (Current Clients + 2)")
                    case_root_parts = parts[:i+3]
                break
    
    if case_root_parts:
        output_dir = os.sep.join(case_root_parts + ["NOTES", "AI OUTPUT"])
            
    if not output_dir:
        # Fallback 1: If "NOTES" is already in the path, use it.
        for i in range(len(parts) - 1, -1, -1):
            if parts[i].upper() == "NOTES":
                output_dir = os.path.join(os.sep.join(parts[:i+1]), "AI OUTPUT")
                if i > 0:
                    case_name = parts[i-1]
                break
    
    if not output_dir:
        # Fallback 2: Sibling NOTES to the file's parent folder (Original fallback)
        input_dir = os.path.dirname(input_path)
        parent_dir = os.path.dirname(input_dir)
        output_dir = os.path.join(parent_dir, "NOTES", "AI OUTPUT")
        case_name = os.path.basename(parent_dir)

    if not os.path.exists(output_dir):
        try:
            os.makedirs(output_dir)
            log_event(f"Created directory: {output_dir}")
        except Exception as e:
            log_event(f"Error creating directory {output_dir}: {e}", level="error")
            sys.exit(1)

    output_filename = "exposure_eval.docx"
    output_file = os.path.join(output_dir, output_filename)
    save_to_docx(summary, output_file, os.path.basename(input_path))
    
    log_event("--- Agent Finished Successfully ---")

if __name__ == '__main__':
    main()
