"""
Contradiction Detection Agent for iCharlotte

Analyzes multiple document summaries to identify factual contradictions
and inconsistencies across the case file.

Features:
- Multi-document comparison
- Severity classification
- Suggested resolution actions
- Integration with CaseDataManager
"""

import os
import sys
import re
import json
import datetime
import argparse
from typing import List, Dict, Optional, Any
from dataclasses import dataclass, field, asdict
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_COLOR_INDEX

# Add parent directory to path for imports
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

# Import shared infrastructure
from icharlotte_core.agent_logger import AgentLogger
from icharlotte_core.llm_config import LLMCaller
from icharlotte_core.memory_monitor import MemoryMonitor
from icharlotte_core.exceptions import LLMError

# Import Case Data Manager
try:
    from case_data_manager import CaseDataManager
except ImportError:
    sys.path.append(os.path.join(os.getcwd(), 'Scripts'))
    from case_data_manager import CaseDataManager

# Import Document Registry
try:
    from document_registry import DocumentRegistry, get_available_documents, get_document_type_list
except ImportError:
    from Scripts.document_registry import DocumentRegistry, get_available_documents, get_document_type_list


# =============================================================================
# Configuration
# =============================================================================

SCRIPTS_DIR = os.path.dirname(os.path.abspath(__file__))
PROMPT_FILE = os.path.join(SCRIPTS_DIR, "CONTRADICTION_DETECTION_PROMPT.txt")
LEGACY_LOG_FILE = r"C:\GeminiTerminal\Contradiction_activity.log"


# =============================================================================
# Data Classes
# =============================================================================

@dataclass
class Claim:
    """A factual claim from a document."""
    source: str
    text: str
    context: str = ""

    def to_dict(self) -> dict:
        return asdict(self)


@dataclass
class Contradiction:
    """A contradiction between two or more claims."""
    id: int
    topic: str
    claim_1: Claim
    claim_2: Claim
    additional_claims: List[Claim] = field(default_factory=list)
    severity: str = "medium"  # high, medium, low
    analysis: str = ""
    resolution_needed: bool = True
    suggested_action: str = ""

    def to_dict(self) -> dict:
        result = {
            'id': self.id,
            'topic': self.topic,
            'claim_1': self.claim_1.to_dict(),
            'claim_2': self.claim_2.to_dict(),
            'additional_claims': [c.to_dict() for c in self.additional_claims],
            'severity': self.severity,
            'analysis': self.analysis,
            'resolution_needed': self.resolution_needed,
            'suggested_action': self.suggested_action
        }
        return result


@dataclass
class ContradictionReport:
    """Report containing all detected contradictions."""
    contradictions: List[Contradiction] = field(default_factory=list)
    file_number: str = ""
    generated: str = ""
    documents_analyzed: List[str] = field(default_factory=list)

    def to_dict(self) -> dict:
        high = sum(1 for c in self.contradictions if c.severity == "high")
        medium = sum(1 for c in self.contradictions if c.severity == "medium")
        low = sum(1 for c in self.contradictions if c.severity == "low")

        key_concerns = [
            c.topic for c in self.contradictions
            if c.severity == "high"
        ][:5]

        return {
            'file_number': self.file_number,
            'generated': self.generated,
            'documents_analyzed': self.documents_analyzed,
            'contradictions': [c.to_dict() for c in self.contradictions],
            'summary': {
                'total_contradictions': len(self.contradictions),
                'high_severity': high,
                'medium_severity': medium,
                'low_severity': low,
                'key_concerns': key_concerns
            }
        }

    @classmethod
    def from_dict(cls, data: dict) -> 'ContradictionReport':
        report = cls(
            file_number=data.get('file_number', ''),
            generated=data.get('generated', ''),
            documents_analyzed=data.get('documents_analyzed', [])
        )

        for c_data in data.get('contradictions', []):
            claim_1 = Claim(**c_data.get('claim_1', {}))
            claim_2 = Claim(**c_data.get('claim_2', {}))
            additional = [Claim(**c) for c in c_data.get('additional_claims', [])]

            contradiction = Contradiction(
                id=c_data.get('id', 0),
                topic=c_data.get('topic', ''),
                claim_1=claim_1,
                claim_2=claim_2,
                additional_claims=additional,
                severity=c_data.get('severity', 'medium'),
                analysis=c_data.get('analysis', ''),
                resolution_needed=c_data.get('resolution_needed', True),
                suggested_action=c_data.get('suggested_action', '')
            )
            report.contradictions.append(contradiction)

        return report


# =============================================================================
# Document Output
# =============================================================================

def save_report_to_docx(report: ContradictionReport, output_path: str, logger: AgentLogger) -> bool:
    """Save contradiction report to DOCX."""
    try:
        doc = Document()

        # Title
        title = doc.add_paragraph()
        run = title.add_run("CONTRADICTION ANALYSIS REPORT")
        run.bold = True
        run.underline = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)

        # Metadata
        doc.add_paragraph(f"File Number: {report.file_number}")
        doc.add_paragraph(f"Generated: {report.generated}")
        doc.add_paragraph(f"Documents Analyzed: {len(report.documents_analyzed)}")

        # Summary
        doc.add_heading("Executive Summary", level=1)

        summary = report.to_dict()['summary']
        p = doc.add_paragraph()
        p.add_run(f"Total Contradictions Found: {summary['total_contradictions']}\n")
        p.add_run(f"High Severity: ").bold = True
        p.add_run(f"{summary['high_severity']}\n")
        p.add_run(f"Medium Severity: {summary['medium_severity']}\n")
        p.add_run(f"Low Severity: {summary['low_severity']}\n")

        if summary['key_concerns']:
            p.add_run("\nKey Concerns:\n").bold = True
            for concern in summary['key_concerns']:
                p.add_run(f"  - {concern}\n")

        # Detailed Contradictions
        doc.add_heading("Detailed Findings", level=1)

        # Group by severity
        by_severity = {'high': [], 'medium': [], 'low': []}
        for c in report.contradictions:
            by_severity[c.severity].append(c)

        severity_headers = {
            'high': 'High Severity Issues',
            'medium': 'Medium Severity Issues',
            'low': 'Low Severity Issues'
        }

        for severity in ['high', 'medium', 'low']:
            contradictions = by_severity[severity]
            if not contradictions:
                continue

            doc.add_heading(severity_headers[severity], level=2)

            for c in contradictions:
                # Topic heading
                p = doc.add_paragraph()
                run = p.add_run(f"#{c.id}: {c.topic}")
                run.bold = True
                run.font.size = Pt(12)

                if severity == 'high':
                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW

                # Claims
                table = doc.add_table(rows=2, cols=2)
                table.style = 'Table Grid'

                # Header row
                table.rows[0].cells[0].text = "Source"
                table.rows[0].cells[1].text = "Claim"

                # Claim 1
                table.rows[0].cells[0].paragraphs[0].add_run(c.claim_1.source).bold = True
                table.rows[0].cells[1].text = c.claim_1.text

                # Claim 2
                table.rows[1].cells[0].paragraphs[0].add_run(c.claim_2.source).bold = True
                table.rows[1].cells[1].text = c.claim_2.text

                # Analysis
                p = doc.add_paragraph()
                p.add_run("Analysis: ").bold = True
                p.add_run(c.analysis)

                # Suggested action
                if c.suggested_action:
                    p = doc.add_paragraph()
                    p.add_run("Suggested Action: ").bold = True
                    p.add_run(c.suggested_action)

                doc.add_paragraph()  # Spacing

        doc.save(output_path)
        logger.output_file(output_path)
        return True

    except Exception as e:
        logger.error(f"Error saving report to DOCX: {e}")
        return False


# =============================================================================
# Main Processing
# =============================================================================

def find_ai_output_docx(file_number: str) -> Optional[str]:
    """
    Find the AI_OUTPUT.docx file for a case.

    Args:
        file_number: The case file number (e.g., '3850.084').

    Returns:
        Path to AI_OUTPUT.docx or None if not found.
    """
    # Try multiple possible base paths
    base_paths = [
        r"Z:\Shared\Current Clients",
        r"C:\Current Clients",
    ]

    # Also check CaseDataManager for the file_path variable
    try:
        data_manager = CaseDataManager()
        file_path = data_manager.get_value(file_number, "file_path")
        if file_path and os.path.exists(file_path):
            ai_output_path = os.path.join(file_path, "NOTES", "AI OUTPUT", "AI_OUTPUT.docx")
            if os.path.exists(ai_output_path):
                return ai_output_path
    except Exception:
        pass

    for base in base_paths:
        if not os.path.exists(base):
            continue

        # Search for case folder containing the file number
        for client_folder in os.listdir(base):
            client_path = os.path.join(base, client_folder)
            if not os.path.isdir(client_path):
                continue

            for case_folder in os.listdir(client_path):
                if file_number in case_folder:
                    ai_output_path = os.path.join(
                        client_path, case_folder, "NOTES", "AI OUTPUT", "AI_OUTPUT.docx"
                    )
                    if os.path.exists(ai_output_path):
                        return ai_output_path

    return None


def gather_summaries_from_docx(docx_path: str, logger: AgentLogger) -> Dict[str, str]:
    """
    Extract individual document summaries from AI_OUTPUT.docx.

    The DOCX contains multiple summaries, each starting with a bold+underlined title
    (the original filename) followed by a "Generated on:" line and the summary content.

    Args:
        docx_path: Path to AI_OUTPUT.docx.
        logger: AgentLogger instance.

    Returns:
        Dictionary of document_name -> summary_text.
    """
    summaries = {}

    try:
        doc = Document(docx_path)

        current_title = None
        current_content = []

        for para in doc.paragraphs:
            text = para.text.strip()

            if not text:
                continue

            # Check if this paragraph is a title (bold + underlined)
            is_title = False
            if para.runs:
                first_run = para.runs[0]
                if first_run.bold and first_run.underline:
                    is_title = True

            if is_title:
                # Save previous summary if exists
                if current_title and current_content:
                    summary_text = "\n".join(current_content)
                    if len(summary_text) > 100:  # Skip very short entries
                        summaries[current_title] = summary_text
                        logger.info(f"Extracted summary: {current_title}")

                # Start new summary
                current_title = text
                current_content = []

            elif current_title:
                # Skip "Generated on:" lines
                if text.startswith("Generated on:"):
                    continue
                current_content.append(text)

        # Don't forget the last summary
        if current_title and current_content:
            summary_text = "\n".join(current_content)
            if len(summary_text) > 100:
                summaries[current_title] = summary_text
                logger.info(f"Extracted summary: {current_title}")

        logger.info(f"Extracted {len(summaries)} summaries from {docx_path}")

    except Exception as e:
        logger.warning(f"Error reading AI_OUTPUT.docx: {e}")

    return summaries


def _matches_selection(name: str, selected_summaries: List[str]) -> bool:
    """
    Check if a summary name matches any of the selected summaries.
    Uses multiple matching strategies with careful handling of legal document types.
    """
    if selected_summaries is None:
        return True

    # Legal abbreviations - map abbreviation to (full_form, category)
    # Category helps distinguish between document types
    abbrev_to_full = {
        'frog': ('form interrogatories', 'form_rogs'),
        'srog': ('special interrogatories', 'special_rogs'),
        'rfpd': ('request for production', 'rfpd'),
        'rfa': ('request for admission', 'rfa'),
        'depo': ('deposition', 'deposition'),
    }

    # Also map full forms to categories
    full_to_category = {
        'form interrogatories': 'form_rogs',
        'special interrogatories': 'special_rogs',
        'form interrogatory': 'form_rogs',
        'special interrogatory': 'special_rogs',
        'deposition': 'deposition',
        'request for production': 'rfpd',
        'request for admission': 'rfa',
    }

    name_lower = name.lower()

    def get_doc_category(text):
        """Determine document category from text."""
        text_lower = text.lower()
        # Split on common delimiters
        words = set(re.split(r'[_\s\-\.\'\"()]+', text_lower))
        # Check abbreviations first
        for abbr, (full, cat) in abbrev_to_full.items():
            if abbr in words:
                return cat
        # Check full forms
        for full, cat in full_to_category.items():
            if full in text_lower:
                return cat
        return None

    name_category = get_doc_category(name_lower)

    for selection in selected_summaries:
        sel_lower = selection.lower()

        # Exact match (case-insensitive)
        if name_lower == sel_lower:
            return True

        # Partial match - selection is substring of name or vice versa
        if sel_lower in name_lower or name_lower in sel_lower:
            return True

        # Match without file extension
        name_no_ext = name_lower.rsplit('.', 1)[0] if '.' in name_lower else name_lower
        sel_no_ext = sel_lower.rsplit('.', 1)[0] if '.' in sel_lower else sel_lower
        if name_no_ext == sel_no_ext or sel_no_ext in name_no_ext or name_no_ext in sel_no_ext:
            return True

        # Get selection category
        sel_category = get_doc_category(sel_lower)

        # If both have categories, they must match - this is strict for legal discovery types
        if name_category and sel_category:
            if name_category != sel_category:
                continue  # Category mismatch - skip entirely, don't do keyword matching
            # Categories match! This is strong evidence - just need minimal word overlap
            stop_words = {'the', 'and', 'for', 'with', 'from', 'set', 'one', 'two', 'summary',
                          'discovery', 'responses', 'response', 'plaintiff', 'defendant', 'to'}
            name_words = set(w for w in re.split(r'[_\s\-\.\'\"()]+', name_lower) if len(w) >= 3 and w not in stop_words)
            sel_words = set(w for w in re.split(r'[_\s\-\.\'\"()]+', sel_lower) if len(w) >= 3 and w not in stop_words)
            common_words = name_words & sel_words
            if len(common_words) >= 1:
                return True  # Same category + at least one common word = match
            continue  # Categories match but no common words - still don't fall through to generic matching

        # Extract significant words for keyword matching (only if categories don't apply)
        stop_words = {'the', 'and', 'for', 'with', 'from', 'set', 'one', 'two', 'summary',
                      'discovery', 'responses', 'response', 'plaintiff', 'defendant', 'to'}
        name_words = set(w for w in re.split(r'[_\s\-\.\'\"()]+', name_lower) if len(w) >= 3 and w not in stop_words)
        sel_words = set(w for w in re.split(r'[_\s\-\.\'\"()]+', sel_lower) if len(w) >= 3 and w not in stop_words)
        common_words = name_words & sel_words

        # Special case: deposition matching when only one side has the category
        if name_category == 'deposition' or sel_category == 'deposition':
            if 'depo' in name_lower or 'deposition' in sel_lower:
                if len(common_words) >= 1:
                    return True
            continue  # Deposition type - don't fall through to generic matching

        # For documents where neither has a recognized category, use keyword matching
        if not name_category and not sel_category:
            if len(common_words) >= 2:
                return True
            if len(common_words) >= 1 and len(common_words) >= min(len(name_words), len(sel_words)) * 0.5:
                return True

    return False


def gather_case_summaries(file_number: str, logger: AgentLogger,
                          selected_summaries: List[str] = None) -> Dict[str, str]:
    """
    Gather all summaries for a case from Document Registry, AI_OUTPUT.docx, and CaseDataManager.

    Args:
        file_number: The case file number.
        logger: AgentLogger instance.
        selected_summaries: Optional list of specific summary names to include.
                           If None, includes all available summaries.
                           Matching is case-insensitive and supports partial matches.

    Returns:
        Dictionary of document_name -> summary_text.
    """
    summaries = {}

    # Debug: Log what we're looking for
    if selected_summaries:
        logger.info(f"Looking for summaries: {selected_summaries}")

    # Build a mapping from Document Registry names to CaseDataManager variable names
    # This allows matching UI-friendly names to actual stored variables
    registry_name_to_var = {}
    try:
        docs = get_available_documents(file_number)
        for doc in docs:
            doc_name = doc.get('name', '')
            # The variable name in CaseDataManager follows a pattern based on agent type
            agent = doc.get('agent', '')
            if agent == 'summarize_discovery':
                # Discovery summaries use: discovery_summary_{sanitized_name}
                sanitized = doc_name.lower().replace(' ', '_').replace('.', '_').replace('(', '_').replace(')', '_')
                sanitized = sanitized.rstrip('_')
                var_name = f"discovery_summary_{sanitized}"
                registry_name_to_var[doc_name] = var_name
            elif agent == 'summarize_deposition':
                # Deposition summaries use: depo_summary_{sanitized_name}
                sanitized = doc_name.lower().replace(' ', '_').replace('.', '_').replace('(', '_').replace(')', '_')
                sanitized = sanitized.rstrip('_')
                # Extract just the deponent name part if possible
                if 'deposition of' in doc_name.lower():
                    parts = doc_name.lower().split('deposition of')
                    if len(parts) > 1:
                        deponent = parts[1].strip().replace(' ', '_').replace('.', '_')
                        var_name = f"depo_summary_{deponent}"
                        registry_name_to_var[doc_name] = var_name
                else:
                    var_name = f"depo_summary_{sanitized}"
                    registry_name_to_var[doc_name] = var_name
            elif agent == 'summarize':
                # General summaries might use: summary_{sanitized_name}
                sanitized = doc_name.lower().replace(' ', '_').replace('.', '_').replace('(', '_').replace(')', '_')
                var_name = f"summary_{sanitized}"
                registry_name_to_var[doc_name] = var_name

        logger.info(f"Built registry mapping for {len(registry_name_to_var)} documents")
    except Exception as e:
        logger.warning(f"Could not build registry mapping: {e}")

    # First, try to load from AI_OUTPUT.docx (primary source)
    docx_path = find_ai_output_docx(file_number)
    if docx_path:
        logger.info(f"Found AI_OUTPUT.docx: {docx_path}")
        docx_summaries = gather_summaries_from_docx(docx_path, logger)

        for name, text in docx_summaries.items():
            if not _matches_selection(name, selected_summaries):
                continue
            summaries[name] = text
    else:
        logger.info("AI_OUTPUT.docx not found, checking case data store...")

    # Check CaseDataManager for summaries (discovery, deposition, etc.)
    try:
        data_manager = CaseDataManager()

        # Get all variables for the case (flatten=False to get full objects with value/source/tags)
        variables = data_manager.get_all_variables(file_number, flatten=False)

        # Debug: Log available summary variables
        summary_vars = [v for v in variables.keys() if any(t in v.lower() for t in ['summary', 'depo', 'extraction'])]
        logger.info(f"Available summary variables in CaseDataManager: {summary_vars}")

        for var_name, var_data in variables.items():
            # Filter for summary-type variables
            if any(tag in var_name.lower() for tag in ['summary', 'depo', 'extraction']):
                # Check if this variable matches selection directly or via registry mapping
                matches = _matches_selection(var_name, selected_summaries)

                # Also check if any registry name that maps to this var_name is selected
                if not matches and selected_summaries:
                    for reg_name, mapped_var in registry_name_to_var.items():
                        if mapped_var == var_name or var_name in mapped_var or mapped_var in var_name:
                            if _matches_selection(reg_name, selected_summaries):
                                matches = True
                                logger.info(f"Matched '{reg_name}' -> '{var_name}'")
                                break

                if not matches:
                    logger.info(f"  No match for var '{var_name}'")
                    continue

                value = var_data.get('value', '')
                source = var_data.get('source', var_name)

                if value and len(value) > 100:  # Skip very short entries
                    # Avoid duplicates - use var_name as key to differentiate from docx summaries
                    if var_name not in summaries:
                        summaries[var_name] = value
                        logger.info(f"Loaded summary from case data: {var_name}")

    except Exception as e:
        logger.warning(f"Error loading case summaries from data store: {e}")

    return summaries


def detect_contradictions(
    summaries: Dict[str, str],
    file_number: str,
    logger: AgentLogger
) -> Optional[ContradictionReport]:
    """
    Detect contradictions between document summaries.

    Args:
        summaries: Dictionary of document_name -> summary_text.
        file_number: Case file number.
        logger: AgentLogger instance.

    Returns:
        ContradictionReport or None on failure.
    """
    if len(summaries) < 2:
        logger.warning("Need at least 2 summaries for contradiction detection")
        return None

    memory_monitor = MemoryMonitor(warn_threshold_mb=1500, abort_threshold_mb=2000, logger=logger.info)
    llm_caller = LLMCaller(logger=logger)

    logger.progress(5, "Initializing contradiction detection...")

    # Load prompt
    try:
        with open(PROMPT_FILE, "r", encoding="utf-8") as f:
            prompt = f.read()
    except Exception as e:
        logger.error(f"Error reading prompt file: {e}")
        return None

    logger.progress(10, f"Building content from {len(summaries)} summaries...")

    # Build document content
    doc_content = ""
    for name, text in summaries.items():
        doc_content += f"\n\n=== {name} ===\n{text[:20000]}"  # Limit per doc

    logger.progress(15, "Content prepared, starting analysis...")
    logger.pass_start("Contradiction Analysis", 1, 1)

    try:
        with memory_monitor.track_operation("Contradiction Analysis"):
            logger.progress(20, "Sending summaries to LLM for contradiction analysis...")
            # Use agent_contradict config from LLM settings (default: gemini-3-pro-preview)
            response = llm_caller.call(prompt, doc_content, agent_id="agent_contradict")

            if not response:
                raise LLMError("LLM returned empty response")

            logger.progress(70, "Parsing contradiction findings from response...")

            # Parse JSON from response
            json_match = re.search(r'\{[\s\S]*\}', response)
            if json_match:
                data = json.loads(json_match.group())
            else:
                raise ValueError("No JSON found in response")

            # Build report
            report = ContradictionReport(
                file_number=file_number,
                generated=datetime.datetime.now().isoformat(),
                documents_analyzed=list(summaries.keys())
            )

            logger.progress(80, "Building contradiction report...")

            for c_data in data.get('contradictions', []):
                claim_1_data = c_data.get('claim_1', {})
                claim_2_data = c_data.get('claim_2', {})

                claim_1 = Claim(
                    source=claim_1_data.get('source', ''),
                    text=claim_1_data.get('text', ''),
                    context=claim_1_data.get('context', '')
                )
                claim_2 = Claim(
                    source=claim_2_data.get('source', ''),
                    text=claim_2_data.get('text', ''),
                    context=claim_2_data.get('context', '')
                )

                additional = []
                for ac in c_data.get('additional_claims', []):
                    additional.append(Claim(
                        source=ac.get('source', ''),
                        text=ac.get('text', ''),
                        context=ac.get('context', '')
                    ))

                contradiction = Contradiction(
                    id=c_data.get('id', len(report.contradictions) + 1),
                    topic=c_data.get('topic', ''),
                    claim_1=claim_1,
                    claim_2=claim_2,
                    additional_claims=additional,
                    severity=c_data.get('severity', 'medium'),
                    analysis=c_data.get('analysis', ''),
                    resolution_needed=c_data.get('resolution_needed', True),
                    suggested_action=c_data.get('suggested_action', '')
                )
                report.contradictions.append(contradiction)

            logger.progress(90, f"Found {len(report.contradictions)} contradictions")
            logger.info(f"Found {len(report.contradictions)} contradictions")

    except json.JSONDecodeError as e:
        logger.pass_failed("Contradiction Analysis", f"JSON parse error: {e}", recoverable=True)
        return None
    except Exception as e:
        logger.pass_failed("Contradiction Analysis", str(e), recoverable=True)
        return None

    logger.progress(95, "Analysis complete")
    logger.pass_complete("Contradiction Analysis", success=True)
    return report


def get_output_directory(file_number: str) -> str:
    """Get output directory for the case."""
    # First check CaseDataManager for the file_path variable
    try:
        data_manager = CaseDataManager()
        file_path = data_manager.get_value(file_number, "file_path")
        if file_path and os.path.exists(file_path):
            output_dir = os.path.join(file_path, "NOTES", "AI OUTPUT")
            if os.path.exists(output_dir):
                return output_dir
    except Exception:
        pass

    # Try multiple possible base paths
    base_paths = [
        r"Z:\Shared\Current Clients",
        r"C:\Current Clients",
    ]

    for base in base_paths:
        if not os.path.exists(base):
            continue
        for client_folder in os.listdir(base):
            client_path = os.path.join(base, client_folder)
            if os.path.isdir(client_path):
                for case_folder in os.listdir(client_path):
                    if file_number in case_folder:
                        return os.path.join(client_path, case_folder, "NOTES", "AI OUTPUT")

    # Fallback
    return os.path.join(os.getcwd(), "output")


def list_available_documents(file_number: str) -> None:
    """List available documents for a case (for pre-selection in UI)."""
    print(f"\nAvailable documents for case {file_number}:")
    print("=" * 60)

    # Get from document registry
    docs = get_available_documents(file_number)

    if docs:
        print(f"\nFrom Document Registry ({len(docs)} documents):")
        for doc in docs:
            print(f"  - {doc['name']}")
            print(f"    Type: {doc['document_type']}")
            print(f"    Agent: {doc['agent']}")
            print()
    else:
        print("\n  No documents in registry. Run summarize agents first.")

    # Also show what's in AI_OUTPUT.docx
    docx_path = find_ai_output_docx(file_number)
    if docx_path:
        print(f"\nFrom AI_OUTPUT.docx ({docx_path}):")
        try:
            doc = Document(docx_path)
            titles = []
            for para in doc.paragraphs:
                if para.runs:
                    first_run = para.runs[0]
                    if first_run.bold and first_run.underline and para.text.strip():
                        titles.append(para.text.strip())
            for title in titles:
                print(f"  - {title}")
            print(f"\n  Total: {len(titles)} summaries in DOCX")
        except Exception as e:
            print(f"  Error reading DOCX: {e}")

    print()


def main():
    """Main entry point."""
    parser = argparse.ArgumentParser(
        description="Detect contradictions between document summaries in a case."
    )
    parser.add_argument("file_number", help="The case file number (e.g., 2024.123)")
    parser.add_argument("--summaries", type=str, default=None,
                        help="Comma-separated list of summary names to include (default: all)")
    parser.add_argument("--list", action="store_true",
                        help="List available documents for the case without running analysis")

    args = parser.parse_args()

    file_number = args.file_number.strip()

    # If --list flag, just show available documents and exit
    if args.list:
        list_available_documents(file_number)
        sys.exit(0)

    # Parse selected summaries if provided
    selected_summaries = None
    if args.summaries:
        selected_summaries = [s.strip() for s in args.summaries.split(",") if s.strip()]

    # Validate file number format
    if not re.match(r'\d{4}\.\d{3}', file_number):
        print(f"Warning: File number '{file_number}' may not be in expected format (YYYY.NNN)")

    # Initialize logger
    logger = AgentLogger("Contradiction", file_number=file_number)
    logger.info(f"Starting contradiction detection for case: {file_number}")

    if selected_summaries:
        logger.info(f"Filtering to {len(selected_summaries)} selected summaries")
        for s in selected_summaries:
            logger.info(f"  - Selected: '{s}'")

    # Gather summaries
    summaries = gather_case_summaries(file_number, logger, selected_summaries)

    if not summaries:
        logger.error("No summaries found for this case. Run summarization agents first.")
        sys.exit(1)

    logger.info(f"Found {len(summaries)} document summaries")

    # Detect contradictions (uses agent_contradict config from LLM settings)
    report = detect_contradictions(summaries, file_number, logger)

    if not report:
        logger.error("Failed to detect contradictions")
        sys.exit(1)

    # Get output path
    output_dir = get_output_directory(file_number)
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    # Save outputs
    docx_path = os.path.join(output_dir, "Contradiction_Report.docx")
    json_path = os.path.join(output_dir, "Contradiction_Report.json")

    # Save DOCX
    save_report_to_docx(report, docx_path, logger)

    # Save JSON
    try:
        with open(json_path, "w", encoding="utf-8") as f:
            json.dump(report.to_dict(), f, indent=2)
        logger.info(f"Saved JSON to: {json_path}")
    except Exception as e:
        logger.warning(f"Could not save JSON: {e}")

    # Save to CaseDataManager
    try:
        data_manager = CaseDataManager()
        data_manager.save_variable(
            file_number,
            "contradiction_report",
            json.dumps(report.to_dict()),
            source="contradiction_agent",
            extra_tags=["Contradictions", "Analysis"]
        )
        logger.info("Saved report to case data")
    except Exception as e:
        logger.warning(f"Could not save to case data: {e}")

    # Summary output
    summary = report.to_dict()['summary']
    logger.info(f"Analysis complete: {summary['total_contradictions']} contradictions found")
    logger.info(f"  High: {summary['high_severity']}, Medium: {summary['medium_severity']}, Low: {summary['low_severity']}")

    sys.exit(0)


if __name__ == '__main__':
    main()
