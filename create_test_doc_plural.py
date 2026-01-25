from docx import Document
import os

doc = Document()
doc.add_paragraph("The Plaintiffs filed a motion.") # Already has 'The'
doc.add_paragraph("Plaintiffs argue that...") # Should change to 'the Plaintiffs'

out_dir = os.path.join(".gemini", "tmp")
if not os.path.exists(out_dir):
    os.makedirs(out_dir)
    
out_path = os.path.join(out_dir, "test_preview_plural.docx")
doc.save(out_path)
print(f"Created {out_path}")
