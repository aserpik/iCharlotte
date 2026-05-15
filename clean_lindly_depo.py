"""Convert James Lindly Zoom transcript to standard deposition Q./A. format."""

import json
import os
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

SOURCE_PATH = r"Z:\Shared\Current Clients\5800 - AMTRUST\017 - Mederos\NOTES\James Lindly Depo Rough.docx"
OUTPUT_PATH = r"Z:\Shared\Current Clients\5800 - AMTRUST\017 - Mederos\NOTES\James Lindly Depo Cleaned.docx"
EXTRACTED_JSON = r"C:\Users\ASerpik.DESKTOP-MRIMK0D\.claude\projects\C--geminiterminal2\f02230ce-6941-4d5d-a8d8-3af7c9531e52\tool-results\mcp-word-document-server-get_document_text-1777671036309.txt"


# Speaker mapping: raw name (from "@HH:MM - Name") -> deposition label
# Q./A. for the questioning attorney and the deponent; others are identified by role.
QUESTIONING_ATTORNEY = "Andrei Serpik"
DEPONENT = "James Lindly"

SPEAKER_MAP = {
    "Andrei Serpik": "Q.",
    "James Lindly": "A.",
    "Chris Markarian": "MR. MARKARIAN:",
    "HILDA ZAMORA": "MS. ZAMORA:",
    "Hilda Zamora": "MS. ZAMORA:",
    "Michelle Rubio": "THE REPORTER:",
    "Almamun Alfarah": "THE VIDEOGRAPHER:",
}

# Pattern: "@H:MM - Speaker (optional email/title)" header line
HEADER_RE = re.compile(r"^@\s*\d+:\d+(?::\d+)?\s*-\s*(.+?)\s*$")


def clean_speaker_name(raw: str) -> str:
    """Strip emails, role tags, CSR numbers, etc. from raw header speaker text."""
    # Remove parenthetical content (e.g., email addresses)
    s = re.sub(r"\s*\([^)]*\)", "", raw)
    # Remove pipe-delimited role tags (e.g., "| Videographer")
    s = re.sub(r"\s*\|.*$", "", s)
    # Remove ", CSR ####" and trailing email
    s = re.sub(r",\s*CSR\s*\d+.*$", "", s)
    # Remove standalone email
    s = re.sub(r"\S+@\S+\.\S+", "", s)
    return s.strip()


def map_speaker(raw_name: str) -> str:
    """Return the deposition label for a raw speaker name."""
    cleaned = clean_speaker_name(raw_name)
    if cleaned in SPEAKER_MAP:
        return SPEAKER_MAP[cleaned]
    # Case-insensitive fallback
    for k, v in SPEAKER_MAP.items():
        if k.lower() == cleaned.lower():
            return v
    # Unknown speaker — capitalize and append colon
    return cleaned.upper() + ":"


def parse_transcript(text: str):
    """
    Parse the raw Zoom transcript text into a list of (label, body) tuples.

    Format of each entry in source:
        @0:04 - James Lindly
        Good morning.

    Header lines start with '@'; everything after until the next header is the body.
    """
    lines = text.splitlines()
    entries = []
    current_label = None
    current_body = []

    # Skip preamble lines until we hit the first header
    seen_first_header = False

    for raw_line in lines:
        line = raw_line.rstrip()
        m = HEADER_RE.match(line.strip())
        if m:
            # Flush previous entry
            if current_label is not None:
                body = "\n".join(current_body).strip()
                if body:
                    entries.append((current_label, body))
            current_label = map_speaker(m.group(1))
            current_body = []
            seen_first_header = True
            continue
        if seen_first_header:
            current_body.append(line)
        # Lines before the first header (e.g., title/recording info) are dropped.

    # Flush last entry
    if current_label is not None:
        body = "\n".join(current_body).strip()
        if body:
            entries.append((current_label, body))

    return entries


def merge_consecutive_same_speaker(entries):
    """Combine adjacent entries from the same speaker into a single block."""
    merged = []
    for label, body in entries:
        if merged and merged[-1][0] == label:
            merged[-1] = (label, merged[-1][1] + " " + body)
        else:
            merged.append((label, body))
    return merged


def collapse_whitespace(text: str) -> str:
    """Collapse internal whitespace and normalize spacing around punctuation."""
    # Replace newlines with spaces
    text = re.sub(r"\s*\n\s*", " ", text)
    # Collapse multiple spaces
    text = re.sub(r"[ \t]+", " ", text)
    # Trim space before punctuation
    text = re.sub(r"\s+([,.;:?!])", r"\1", text)
    return text.strip()


def write_depo_docx(entries, out_path: str, case_caption: str = None):
    """Write the cleaned entries to a Word document in standard depo format."""
    doc = Document()

    # Set default style: Times New Roman 12pt, double-spaced
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    pf = normal.paragraph_format
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)

    # Page margins (1" all around — typical for depo transcripts)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("DEPOSITION OF JAMES LINDLY")
    run.bold = True
    run.font.size = Pt(13)
    title.paragraph_format.space_after = Pt(6)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run("April 29, 2026")
    sub_run.italic = True
    subtitle.paragraph_format.space_after = Pt(18)

    # Body
    for label, body in entries:
        clean = collapse_whitespace(body)
        if not clean:
            continue

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.first_line_indent = Inches(0)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)  # hanging indent

        # Bold label, then space, then body text
        run_label = p.add_run(f"{label}\t")
        run_label.bold = True
        p.add_run(clean)

    doc.save(out_path)


def main():
    # Load extracted text from JSON
    with open(EXTRACTED_JSON, "r", encoding="utf-8") as f:
        data = json.load(f)
    text = data[0]["text"]
    print(f"Loaded {len(text):,} chars from extracted transcript JSON")

    entries = parse_transcript(text)
    print(f"Parsed {len(entries)} raw entries")

    merged = merge_consecutive_same_speaker(entries)
    print(f"After merging consecutive same-speaker: {len(merged)} entries")

    # Stats by label
    from collections import Counter
    counter = Counter(label for label, _ in merged)
    print("\nLabel distribution:")
    for label, count in counter.most_common():
        print(f"  {label!r}: {count}")

    # Make sure parent dir exists (should already)
    out_dir = os.path.dirname(OUTPUT_PATH)
    if not os.path.isdir(out_dir):
        raise SystemExit(f"Output directory missing: {out_dir}")

    write_depo_docx(merged, OUTPUT_PATH)
    size = os.path.getsize(OUTPUT_PATH)
    print(f"\nWrote: {OUTPUT_PATH}  ({size:,} bytes)")


if __name__ == "__main__":
    main()
