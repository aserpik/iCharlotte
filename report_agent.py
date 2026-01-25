from docx.enum.text import WD_LINE_SPACING
from docx.shared import Pt, Inches
import sys
import os
import shutil
import glob
from docx import Document
import re
import json
import subprocess
import time

PROJECT_ROOT = os.getcwd()

# Configuration
# Z:\Shared\Current Clients -> /mnt/z/Shared/Current Clients
CASE_ROOT_BASE = "/mnt/z/Shared/Current Clients" 

TEMPLATE_DIR = "/mnt/c/geminitest1/Test"
TEMPLATE_FILENAME = "Template Report v2.docx"

CASE_DATA_DIR = os.path.join(PROJECT_ROOT, ".gemini", "case_data")

def get_case_var_py(file_num, key):
    json_file = os.path.join(CASE_DATA_DIR, f"{file_num}.json")
    if not os.path.exists(json_file):
        return None
    try:
        with open(json_file, 'r') as f:
            data = json.load(f)
        return data.get(key)
    except:
        return None

def find_case_root(case_number):
    try:
        parts = case_number.split('.')
        if len(parts) == 2:
            carrier_num, case_id = parts
        else:
            carrier_num = "*"
            case_id = case_number

        carrier_glob = os.path.join(CASE_ROOT_BASE, f"{carrier_num} - *")
        carrier_folders = glob.glob(carrier_glob)
        
        if not carrier_folders:
            print(f"Error: Carrier folder for {carrier_num} not found.")
            return None

        carrier_path = carrier_folders[0]
        case_glob = os.path.join(carrier_path, f"{case_id} - *")
        case_folders = glob.glob(case_glob)
        
        if not case_folders:
            print(f"Error: Case folder for {case_id} not found in {carrier_path}.")
            return None
            
        return case_folders[0]

    except Exception as e:
        print(f"Error finding case root: {e}")
        return None

def get_paths(case_root):
    return {
        'root': case_root,
        'status': os.path.join(case_root, "STATUS"),
        'notes': os.path.join(case_root, "NOTES", "AI OUTPUT"),
        'template': os.path.join(TEMPLATE_DIR, TEMPLATE_FILENAME)
    }

def find_or_create_draft(paths, case_number):
    if not os.path.exists(paths['status']):
        os.makedirs(paths['status'])
        
    status_pattern = os.path.join(paths['status'], "*Carrier*.docx")
    all_status_docs = glob.glob(status_pattern)
    
    drafts = [d for d in all_status_docs if "[DRAFT]" in os.path.basename(d)]
    
    if drafts:
        print(f"Found existing draft: {drafts[0]}")
        return drafts[0]
    
    if not os.path.exists(paths['template']):
        print(f"Error: Template {paths['template']} not found.")
        return None
            
    new_draft_name = f"[DRAFT] Carrier {case_number}.docx"
    new_draft_path = os.path.join(paths['status'], new_draft_name)
    
    print(f"Creating new draft from template: {new_draft_path}")
    shutil.copy2(paths['template'], new_draft_path)
    return new_draft_path

def read_docx_content(path):
    if not os.path.exists(path):
        return None
    try:
        doc = Document(path)
        return doc
    except Exception as e:
        print(f"Error reading {path}: {e}")
        return None

def parse_ai_output(doc):
    categories = {
        'discovery': [],
        'depositions': [],
        'documents': []
    }
    
    current_category = None 
    current_item = None 
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            if current_item: 
                current_item['content'].append(para)
            continue
            
        text_lower = text.lower()
        is_header = False
        new_category = None
        new_name = "Unknown"
        
        is_style_header = "Heading" in para.style.name or (text.isupper() and len(text) < 100) or "List Paragraph" in para.style.name
        
        if is_style_header or "summary" in text_lower:
            if "discovery" in text_lower or "responses" in text_lower or "interrogatories" in text_lower:
                is_header = True
                new_category = 'discovery'
                name_match = re.search(r'(?:of|from|by)\s+(.+)', text, re.IGNORECASE)
                new_name = name_match.group(1).strip() if name_match else text
                
            elif "deposition" in text_lower:
                is_header = True
                new_category = 'depositions'
                name_match = re.search(r'(?:of|from|by)\s+(.+)', text, re.IGNORECASE)
                new_name = name_match.group(1).strip() if name_match else text

            elif "summary for" in text_lower or "summary of" in text_lower:
                is_header = True
                new_category = 'documents'
                name_match = re.search(r'Summary (?:for|of)\s+(.+)', text, re.IGNORECASE)
                new_name = name_match.group(1).strip() if name_match else text
            
            elif is_style_header and len(current_item['content'] if current_item else []) > 0:
                 is_header = True
                 new_category = 'documents'
                 new_name = text

        if is_header:
            if current_item and current_category:
                categories[current_category].append(current_item)
            
            current_category = new_category if new_category else 'documents'
            current_item = {'name': new_name, 'content': []}
        else:
            if current_item:
                current_item['content'].append(para)
            else:
                current_category = 'documents'
                current_item = {'name': 'General Summary', 'content': [para]}

    if current_item and current_category:
        categories[current_category].append(current_item)
        
    return categories

def get_ordinal(text):
    ordinals = {
        "first": 0, "second": 1, "third": 2, "fourth": 3, "fifth": 4,
        "sixth": 5, "seventh": 6, "eighth": 7, "ninth": 8, "tenth": 9,
        "1st": 0, "2nd": 1, "3rd": 2, "4th": 3, "5th": 4
    }
    text_lower = text.lower()
    for word, val in ordinals.items():
        if f" {word} " in f" {text_lower} " or f"[{word} " in text_lower or f" {word}]" in text_lower or f"{word} " in text_lower:
            return val
    return None

def replace_text_in_runs(para, placeholder, replacement, bold=None, underline=None):
    if placeholder not in para.text:
        return False

    full_match_run = None
    for run in para.runs:
        if placeholder in run.text:
            full_match_run = run
            break
            
    if full_match_run:
        full_match_run.text = full_match_run.text.replace(placeholder, replacement)
        if bold is not None:
            full_match_run.bold = bold
        if underline is not None:
            full_match_run.underline = underline
        return True
        
    para.text = para.text.replace(placeholder, replacement)
    
    if para.runs:
        for run in para.runs:
            if bold is not None: run.bold = bold
            if underline is not None: run.underline = underline
            
    return True

def enforce_single_spacing(doc):
    """
    Iterates through all paragraphs in the document and enforces single spacing.
    """
    for para in doc.paragraphs:
        para.paragraph_format.line_spacing = 1.0
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        # Also ensure space after/before isn't huge? 
        # Requirement was "all be single spaced". 
        # Usually reports have some space after paragraphs. 
        # I will stick to line_spacing = 1.0.

def iter_paragraphs(doc):
    """
    Generator that yields all paragraphs in the document, including those in tables
    and (optionally) headers/footers if we wanted, but sticking to Body + Body Tables
    is usually sufficient for the main report content.
    """
    # 1. Body Paragraphs
    for para in doc.paragraphs:
        yield para

    # 2. Body Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                # Recursive call for nested tables? 
                # python-docx cell.tables is possible but rare in these reports.
                # Just iterating cell paragraphs is safe.
                for para in cell.paragraphs:
                    yield para
                # Check nested tables in cell
                for nested_table in cell.tables:
                    for nrow in nested_table.rows:
                        for ncell in nrow.cells:
                            for npara in ncell.paragraphs:
                                yield npara

def fill_placeholders(doc, parsed_data, liability_doc, exposure_doc, variables):
    used_items = {
        'discovery': set(),
        'depositions': set(),
        'documents': set()
    }

    placeholders_found = []
    
    # Use iter_paragraphs instead of doc.paragraphs
    for i, para in enumerate(iter_paragraphs(doc)):
        text = para.text
        if "[" in text and "]" in text:
            matches = re.findall(r'\[(.*?)\]', text)
            for m in matches:
                placeholders_found.append((i, m, para))

    for idx, content_text, para in placeholders_found:
        full_placeholder = f"[{content_text}]"
        lower_content = content_text.lower()
        
        replacement_text = None
        replacement_content = None 
        
        force_bold = None
        force_underline = None
        
        # --- Variables ---
        
        if "adjuster name" in lower_content:
            raw_name = variables.get("adjuster_name", "Adjuster")
            if "first name" in lower_content:
                 replacement_text = raw_name.split()[0] if raw_name else ""
            else:
                replacement_text = raw_name
            force_bold = False
            force_underline = False

        elif "adjuster email" in lower_content:
            replacement_text = variables.get("adjuster_email", "")
            force_bold = False
            force_underline = False
            
        elif "case name" in lower_content:
            replacement_text = variables.get("case_name", "")
            force_bold = True 
            
        elif "insured name" in lower_content:
            replacement_text = variables.get("insured_name", "")
            
        elif "claim no" in lower_content or "claim number" in lower_content:
            replacement_text = variables.get("claim_number", "")
            
        elif "file no" in lower_content or "file number" in lower_content:
            replacement_text = variables.get("file_number", "")
            
        elif "case no" in lower_content or "court case" in lower_content:
            replacement_text = variables.get("case_number", "")
            
        elif "dol" in lower_content or "date of loss" in lower_content or "incident date" in lower_content:
            replacement_text = variables.get("incident_date", "")

        elif "introduction" in lower_content:
            raw_intro = variables.get("introduction", "")
            clean_intro = re.sub(r'^\s*[\*]*FACTUAL BACKGROUND[\*]*\s*', '', raw_intro, flags=re.IGNORECASE).strip()
            # Split into paragraphs to handle spacing/indentation
            replacement_content = []
            for p_text in clean_intro.split('\n\n'):
                 if p_text.strip():
                     replacement_content.append(p_text.strip())
            
        elif "evaluation of liability" in lower_content:
            if liability_doc:
                replacement_content = liability_doc.paragraphs
            else:
                replacement_text = "[Liability Analysis Not Found]"
        elif "evaluation of exposure" in lower_content:
            if exposure_doc:
                replacement_content = exposure_doc.paragraphs
            else:
                replacement_text = "[Exposure Analysis Not Found]"

        else:
            category = None
            if "discovery" in lower_content or "responding party" in lower_content:
                category = 'discovery'
            elif "deposition" in lower_content or "deponent" in lower_content:
                category = 'depositions'
            elif "document" in lower_content:
                category = 'documents'
                
            if category:
                ordinal = get_ordinal(lower_content)
                if ordinal is not None and ordinal < len(parsed_data[category]):
                    item = parsed_data[category][ordinal]
                    used_items[category].add(ordinal)
                    
                    if "name" in lower_content:
                        replacement_text = item['name']
                        # Requirement: Document names must be underlined and bold
                        force_bold = True
                        force_underline = True
                    elif "summary" in lower_content:
                        replacement_content = item['content']

        # --- Execution ---
        if replacement_text is not None:
            # Clean up replacement text if it's None or "None" string
            if replacement_text is None or replacement_text == "None":
                replacement_text = f"[{content_text}]" # Keep placeholder if missing? Or empty?
            
            replace_text_in_runs(para, full_placeholder, str(replacement_text), bold=force_bold, underline=force_underline)
                
        elif replacement_content is not None:
            # Handle Liability/Exposure Header Stripping
            final_paragraphs_text = []
            
            if isinstance(replacement_content, list) and len(replacement_content) > 0 and isinstance(replacement_content[0], str):
                 # It's a list of strings (like from Introduction)
                 final_paragraphs_text = replacement_content
            else:
                # It's a list of Paragraph objects
                for src_para in replacement_content:
                    src_text = src_para.text.strip()
                    if not src_text: continue
                    
                    src_text_upper = src_text.upper()
                    # Skip if it's the specific header we want to remove
                    if ("EVALUATION OF LIABILITY" in src_text_upper and len(src_text) < 50) or \
                       ("EVALUATION OF EXPOSURE" in src_text_upper and len(src_text) < 50):
                        continue
                    final_paragraphs_text.append(src_text)

            # Insert paragraphs with formatting
            for i, p_text in enumerate(final_paragraphs_text):
                new_p = para.insert_paragraph_before(p_text)
                new_p.paragraph_format.first_line_indent = Inches(0.5)
                new_p.paragraph_format.line_spacing = 1.0
                
                # Add blank line after each paragraph (except maybe the last one if we want to merge with next? No, user said between each)
                # But we are inserting BEFORE 'para'. 
                # So P1, Blank, P2, Blank...
                if i < len(final_paragraphs_text): # Add blank line after every paragraph
                     spacer = para.insert_paragraph_before("")
                     # Fix: Ensure blank line doesn't inherit underline from placeholder
                     try: spacer.style = "Normal"
                     except: pass
                     for run in spacer.runs:
                         run.font.underline = False
                         run.font.bold = False
                     spacer.paragraph_format.line_spacing = 1.0

            
            if full_placeholder == para.text.strip():
                # If the placeholder was the only thing, remove the original paragraph
                # But python-docx doesn't support 'delete()'. 
                # We can clear the text.
                para.text = "" 
                # And maybe remove it from the document? Hard in docx. 
                # Reducing font size to 0 is a hack. 
                # For now, just empty text.
            else:
                 para.text = para.text.replace(full_placeholder, "")

    return used_items

def append_unused_sections(doc, parsed_data, used_items):
    for category in ['discovery', 'depositions', 'documents']:
        for i, item in enumerate(parsed_data[category]):
            if i not in used_items[category]:
                # Requirement: Header must be Bold and Underlined
                p = doc.add_paragraph()
                run = p.add_run(f"{item['name']} ({category.title()})")
                run.bold = True
                run.underline = True
                p.style = 'Heading 1' # Or just normal with bold/underline? Template usually uses Heading 1.
                # If Heading 1 has its own style, bold/underline might be overridden or additive.
                # Let's force it on the run.

                for src_para in item['content']:
                    if src_para.text.strip():
                        doc.add_paragraph(src_para.text, style=src_para.style)

def main():
    sys.stdout.reconfigure(line_buffering=True)

    if len(sys.argv) < 2:
        print("Usage: report_agent.py <file_number>")
        sys.exit(1)
        
    file_number = sys.argv[1] 
    print(f"Processing Case File: {file_number}")
    
    case_root = find_case_root(file_number)
    if not case_root:
        sys.exit(1)
        
    paths = get_paths(case_root)
    draft_path = find_or_create_draft(paths, file_number)
    if not draft_path:
        print("Failed to prepare draft.")
        sys.exit(1)

    print("Loading source documents...")
    liability_doc = read_docx_content(os.path.join(paths['notes'], "Liability_eval.docx"))
    exposure_doc = read_docx_content(os.path.join(paths['notes'], "Exposure_eval.docx"))
    
    ai_output_path = os.path.join(paths['notes'], "AI_OUTPUT.docx")
    parsed_data = {'discovery': [], 'depositions': [], 'documents': []}
    
    if os.path.exists(ai_output_path):
        print(f"Parsing {ai_output_path}...")
        try:
            ai_doc = Document(ai_output_path)
            parsed_data = parse_ai_output(ai_doc)
            print(f"Parsed: {len(parsed_data['discovery'])} discovery, {len(parsed_data['depositions'])} depositions, {len(parsed_data['documents'])} documents.")
        except Exception as e:
            print(f"Error parsing AI_OUTPUT: {e}")
    else:
        print("Warning: AI_OUTPUT.docx not found.")

    print("Checking variables...")
    variables = {
        "file_number": file_number,
        "introduction": get_case_var_py(file_number, "introduction"),
        "adjuster_name": get_case_var_py(file_number, "adjuster_name"),
        "adjuster_email": get_case_var_py(file_number, "adjuster_email"),
        "case_name": get_case_var_py(file_number, "case_name"),
        "insured_name": get_case_var_py(file_number, "insured_name"),
        "claim_number": get_case_var_py(file_number, "claim_number"),
        "case_number": get_case_var_py(file_number, "case_number"), 
        "incident_date": get_case_var_py(file_number, "incident_date")
    }
    
    if not variables["introduction"]:
        print("Introduction not found. Running complaint_agent.sh...")
        complaint_script = os.path.join(PROJECT_ROOT, "scripts", "complaint_agent.sh")
        try:
            subprocess.run([complaint_script, file_number], check=True)
            time.sleep(2) 
            variables["introduction"] = get_case_var_py(file_number, "introduction")
            variables["adjuster_name"] = get_case_var_py(file_number, "adjuster_name")
            variables["adjuster_email"] = get_case_var_py(file_number, "adjuster_email")
            variables["case_name"] = get_case_var_py(file_number, "case_name")
            variables["insured_name"] = get_case_var_py(file_number, "insured_name")
            variables["claim_number"] = get_case_var_py(file_number, "claim_number")
            variables["case_number"] = get_case_var_py(file_number, "case_number")
            variables["incident_date"] = get_case_var_py(file_number, "incident_date")
        except Exception as e:
            print(f"Error running complaint agent: {e}")

    for k, v in variables.items():
        if v is None:
            variables[k] = f"[{k.replace('_', ' ').title()}]"

    print("Merging content into draft...")
    try:
        draft_doc = Document(draft_path)
        used_items = fill_placeholders(draft_doc, parsed_data, liability_doc, exposure_doc, variables)
        append_unused_sections(draft_doc, parsed_data, used_items)
        
        # Enforce single spacing globally
        enforce_single_spacing(draft_doc)
        
        draft_doc.save(draft_path)
        print(f"Report updated successfully: {draft_path}")
        
    except Exception as e:
        print(f"Error merging content: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)

if __name__ == "__main__":
    main()
