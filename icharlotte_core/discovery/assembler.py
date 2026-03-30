"""
Document assembler for the discovery generation pipeline.

Renders DiscoverySet objects into formatted .docx files by appending
content after a caption page template.  Uses the template's built-in
styles (Body Double, Center Double Bold Und, Discovery No., etc.) for
double-spaced legal document formatting.
"""
import os
from datetime import date
from typing import Optional

from docx import Document as DocxDocument
from docx.shared import Pt, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .models import DiscoverySet, DiscoveryType, number_to_word
from .declaration import generate_declaration
from .templates import extract_requests_from_text


# ---------------------------------------------------------------------------
# Style constants — pre-defined in the caption page template
# ---------------------------------------------------------------------------
STYLE_BODY_DOUBLE = "Body Double"                   # Double-spaced, 0.5" first-line indent
STYLE_CENTER_DOUBLE = "Center Double"               # Centered double-spaced
STYLE_CENTER_DOUBLE_BOLD_UND = "Center Double Bold Und"  # Centered bold underline
STYLE_FLUSH_LEFT_DOUBLE = "Flush Left Double"        # Left-aligned double-spaced, no indent
STYLE_DISCOVERY_NO = "Discovery No."                 # Request headers (no indent, bold underline)


def _safe_style(doc, preferred, fallback=STYLE_BODY_DOUBLE):
    """Return the preferred style if it exists, otherwise fallback."""
    try:
        return doc.styles[preferred]
    except KeyError:
        try:
            return doc.styles[fallback]
        except KeyError:
            return None


def _add_para(doc, text="", style_name=STYLE_BODY_DOUBLE,
              bold=False, underline=False):
    """Add a double-spaced paragraph using a named template style.

    Uses the template's built-in styles which already have correct
    line spacing, indentation, and font settings.
    """
    para = doc.add_paragraph()
    style = _safe_style(doc, style_name)
    if style:
        para.style = style

    if text:
        run = para.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        if bold:
            run.bold = True
        if underline:
            run.underline = True
    return para


# ---------------------------------------------------------------------------
# DiscoveryAssembler
# ---------------------------------------------------------------------------

class DiscoveryAssembler:
    """Assembles DiscoverySet objects into formatted .docx files."""

    def __init__(self, caption_page_path: str):
        if not os.path.isfile(caption_page_path):
            raise FileNotFoundError(
                f"Caption page template not found: {caption_page_path}"
            )
        self.caption_page_path = caption_page_path

    # -- public API ---------------------------------------------------------

    def assemble(
        self,
        discovery_set: DiscoverySet,
        output_path: str,
        attorney_name: str = "",
        firm_name: str = "Bordin Semmer LLP",
        date_str: str = "",
    ) -> str:
        """Render a DiscoverySet into a formatted .docx file."""
        ds = discovery_set
        if not date_str:
            date_str = date.today().strftime("%B %d, %Y")

        doc = DocxDocument(self.caption_page_path)

        # Build the document title
        title = self._build_title(ds)

        # (1) Replace "CAPTION PAGE" in caption table, or insert title if empty
        self._set_caption_title(doc, title)

        # (2) Propounding / Responding party block (indented)
        self._insert_party_block(doc, ds)

        # (3) Preamble
        _add_para(doc, "")
        preamble = self._build_preamble(ds)
        _add_para(doc, preamble, STYLE_FLUSH_LEFT_DOUBLE)

        # (4) Instructions
        if ds.instructions_block:
            _add_para(doc, "")
            self._insert_instructions(doc, ds.instructions_block)

        # (4b) Definitions
        if ds.definitions_block:
            _add_para(doc, "")
            self._insert_definitions(doc, ds.definitions_block)

        # (5) Section heading (centered, bold, underline)
        _add_para(doc, "")
        heading = f"{ds.discovery_type.section_heading}, SET {ds.set_word.upper()}"
        p = _add_para(doc, heading, STYLE_CENTER_DOUBLE_BOLD_UND, bold=True, underline=True)
        # Ensure first_line_indent is 0 for centered headings
        p.paragraph_format.first_line_indent = Inches(0)

        # (6) Requests — NO extra blank lines between them
        self._insert_requests(doc, ds)

        # (7) Signature block
        self._insert_signature_block(doc, ds, attorney_name, firm_name, date_str)

        # (8) Declaration (if needed)
        if ds.needs_declaration:
            self._insert_declaration(doc, ds, attorney_name, firm_name, date_str)

        # (9) Footer — document title
        self._set_footer(doc, title)

        # Ensure output directory exists
        output_dir = os.path.dirname(output_path)
        if output_dir:
            os.makedirs(output_dir, exist_ok=True)

        doc.save(output_path)
        return output_path

    def assemble_from_plain_text(
        self,
        plain_text: str,
        discovery_set: DiscoverySet,
        output_path: str,
        attorney_name: str = "",
        firm_name: str = "Bordin Semmer LLP",
        date_str: str = "",
    ) -> str:
        """Re-parse plain text into requests, update the set, and assemble."""
        requests = extract_requests_from_text(plain_text, discovery_set.discovery_type)
        discovery_set.requests = requests
        return self.assemble(
            discovery_set, output_path,
            attorney_name=attorney_name,
            firm_name=firm_name,
            date_str=date_str,
        )

    @staticmethod
    def find_caption_page(case_path: str) -> Optional[str]:
        """Search a case folder for a .docx with 'caption page' in the filename."""
        if not os.path.isdir(case_path):
            return None
        for entry in os.listdir(case_path):
            if entry.lower().endswith(".docx") and "caption page" in entry.lower():
                return os.path.join(case_path, entry)
        # Check one level of subdirectories
        for entry in os.listdir(case_path):
            subdir = os.path.join(case_path, entry)
            if os.path.isdir(subdir):
                for sub_entry in os.listdir(subdir):
                    if sub_entry.lower().endswith(".docx") and "caption page" in sub_entry.lower():
                        return os.path.join(subdir, sub_entry)
        return None

    # -- private helpers ----------------------------------------------------

    def _build_title(self, ds: DiscoverySet) -> str:
        """Build the full document title."""
        return ds.discovery_type.document_title_template.format(
            propounding=f"DEFENDANT {ds.propounding_party.name.upper()}",
            responding=(
                f"{ds.directed_to.role_label.upper()} "
                f"{ds.directed_to.name.upper()}"
            ),
            set_word=ds.set_word.upper(),
        )

    def _set_caption_title(self, doc, title: str):
        """Set the document title in the caption page table.

        Strategy:
        1. Look for "CAPTION PAGE" placeholder text and replace it
        2. If not found, look for the right column of the caption table
           (table with 3 columns, cell [0,2]) and insert title there
        """
        # Strategy 1: Replace "CAPTION PAGE" text
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if "CAPTION PAGE" in para.text.upper():
                            for run in para.runs:
                                run.text = ""
                            if para.runs:
                                para.runs[0].text = title
                                para.runs[0].bold = True
                                para.runs[0].underline = True
                            else:
                                run = para.add_run(title)
                                run.bold = True
                                run.underline = True
                                run.font.name = "Times New Roman"
                                run.font.size = Pt(12)
                            return

        # Strategy 2: Find the 3-column caption table, insert into right cell
        for table in doc.tables:
            if len(table.columns) >= 3:
                cell = table.rows[0].cells[2]
                # Find an empty paragraph or the paragraph after "Case No."
                target_para = None
                for para in cell.paragraphs:
                    if not para.text.strip():
                        target_para = para
                        continue
                    if "case no" in para.text.lower():
                        # Title goes AFTER the case number lines — keep looking
                        target_para = None
                        continue

                # Insert title after case number block
                # Find the right position: after dept/assignment line, before complaint filed
                for i, para in enumerate(cell.paragraphs):
                    text_lower = para.text.strip().lower()
                    if text_lower.startswith("(assigned") or text_lower.startswith("dept"):
                        # Insert title after this paragraph
                        # Add a blank line then the title
                        if i + 1 < len(cell.paragraphs):
                            target_para = cell.paragraphs[i + 1]
                        break

                if target_para is not None and not target_para.text.strip():
                    # Clear and set the title
                    for run in target_para.runs:
                        run.text = ""
                    run = target_para.add_run(title)
                    run.bold = True
                    run.underline = True
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(12)
                    return

                # Last resort: add title as new paragraph in the cell
                para = cell.add_paragraph()
                run = para.add_run(title)
                run.bold = True
                run.underline = True
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)
                return

    def _insert_party_block(self, doc, ds: DiscoverySet):
        """Insert propounding/responding party block.

        Uses the same indent as the sample: left=2400300 EMU (~2.62"),
        first_line_indent=-1943100 EMU (hanging) — this creates the
        tab-aligned layout where labels and values align via tabs.
        """
        LEFT_INDENT = Emu(2400300)
        HANGING_INDENT = Emu(-1943100)

        _add_para(doc, "")

        for text in [
            f"PROPOUNDING PARTY:\t{ds.propounding_party.formal_description}",
            f"RESPONDING PARTY:\t{ds.directed_to.formal_description}",
            f"SET NO.:\t{ds.set_word.upper()} ({ds.set_number})",
        ]:
            p = _add_para(doc, text, STYLE_FLUSH_LEFT_DOUBLE)
            p.paragraph_format.left_indent = LEFT_INDENT
            p.paragraph_format.first_line_indent = HANGING_INDENT

    def _insert_instructions(self, doc, instructions_block: str):
        """Insert the instructions block."""
        lines = instructions_block.strip().split("\n")
        for line in lines:
            stripped = line.strip()
            if not stripped:
                continue
            if "INSTRUCTIONS TO ANSWERING PARTY" in stripped.upper():
                p = _add_para(doc, stripped, STYLE_BODY_DOUBLE, bold=True, underline=True)
                p.paragraph_format.first_line_indent = Inches(0)
            else:
                _add_para(doc, stripped, STYLE_BODY_DOUBLE)

    def _insert_definitions(self, doc, definitions_block: str):
        """Insert the definitions block."""
        for line in definitions_block.strip().split("\n"):
            stripped = line.strip()
            if not stripped:
                continue
            if stripped.upper() in ("DEFINITIONS", "DEFINED TERMS"):
                p = _add_para(doc, stripped, STYLE_CENTER_DOUBLE_BOLD_UND,
                              bold=True, underline=True)
                p.paragraph_format.first_line_indent = Inches(0)
            else:
                _add_para(doc, stripped, STYLE_BODY_DOUBLE)

    def _insert_requests(self, doc, ds: DiscoverySet):
        """Insert numbered discovery requests with NO extra blank lines.

        Request headers use "Discovery No." style (no indent, bold underline).
        Request body and definitions use "Body Double" style (0.5" first-line indent).
        """
        for req in ds.requests:
            # Request header — "Discovery No." style (no 0.5" indent)
            header = ds.discovery_type.request_header_template.format(number=req.number)
            if not header.endswith(":"):
                header += ":"
            _add_para(doc, header, STYLE_DISCOVERY_NO, bold=True, underline=True)

            # Request body text
            _add_para(doc, req.text, STYLE_BODY_DOUBLE)

            # Inline definitions (directly after, no blank line)
            for defn in req.definitions:
                _add_para(doc, defn, STYLE_BODY_DOUBLE)

    def _insert_signature_block(self, doc, ds: DiscoverySet, attorney_name: str,
                                firm_name: str, date_str: str):
        """Insert the signature block."""
        # /// filler lines
        for _ in range(3):
            p = _add_para(doc, "///", STYLE_BODY_DOUBLE)
            p.paragraph_format.first_line_indent = Inches(0)

        _add_para(doc, f"Dated:  {date_str}\t      {firm_name.upper()}",
                  STYLE_FLUSH_LEFT_DOUBLE)
        _add_para(doc, "")
        _add_para(doc, "")
        _add_para(doc, f"By:\t______________________________", STYLE_FLUSH_LEFT_DOUBLE)
        if attorney_name:
            _add_para(doc, f"\t\t{attorney_name}", STYLE_FLUSH_LEFT_DOUBLE)
        _add_para(doc, f"Attorneys for {ds.propounding_party.role_label},",
                  STYLE_FLUSH_LEFT_DOUBLE)
        _add_para(doc, ds.propounding_party.name.upper(), STYLE_FLUSH_LEFT_DOUBLE)

    def _insert_declaration(self, doc, ds: DiscoverySet, attorney_name: str,
                            firm_name: str, date_str: str):
        """Insert CCP declaration with page break before."""
        decl_text = generate_declaration(ds, attorney_name, firm_name)
        if not decl_text:
            return
        decl_text = decl_text.replace("{date}", date_str)

        # Page break
        page_break_para = doc.add_paragraph()
        run = page_break_para.add_run()
        br = OxmlElement("w:br")
        br.set(qn("w:type"), "page")
        run._element.append(br)

        for i, line in enumerate(decl_text.split("\n")):
            stripped = line.strip()
            if not stripped:
                _add_para(doc, "")
                continue
            if i == 0 and stripped.startswith("DECLARATION"):
                p = _add_para(doc, stripped, STYLE_CENTER_DOUBLE_BOLD_UND,
                              bold=True, underline=True)
                p.paragraph_format.first_line_indent = Inches(0)
            elif stripped.startswith("____"):
                _add_para(doc, stripped, STYLE_FLUSH_LEFT_DOUBLE)
            else:
                _add_para(doc, stripped, STYLE_BODY_DOUBLE)

    def _set_footer(self, doc, title: str):
        """Set the document title in the footer of all sections."""
        for section in doc.sections:
            footer = section.footer
            footer.is_linked_to_previous = False

            # Clear existing footer content
            for para in footer.paragraphs:
                for run in para.runs:
                    run.text = ""

            # Set the title in the first footer paragraph
            if footer.paragraphs:
                para = footer.paragraphs[0]
            else:
                para = footer.add_paragraph()

            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(title)
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)

    @staticmethod
    def _build_preamble(ds: DiscoverySet) -> str:
        """Return the statutory preamble text."""
        prop_role = ds.propounding_party.role_label
        prop_name = ds.propounding_party.name.upper()
        resp_role = ds.directed_to.role_label
        resp_name = ds.directed_to.name.upper()
        set_word = number_to_word(ds.set_number)

        if ds.discovery_type == DiscoveryType.SI:
            return (
                f"TO {resp_role.upper()} {resp_name} AND "
                f"ATTORNEYS OF RECORD:\n"
                f"\tPursuant to California Code of Civil Procedure \u00a72030.030, "
                f"{prop_role}, {prop_name} "
                f'("Propounding Party" or "{prop_role}"), hereby '
                f"propounds to {resp_role}, {resp_name} "
                f'("{resp_role}" or "Responding Party"), the following '
                f"{set_word} Set of Special Interrogatories, "
                f"each of which shall be answered fully, separately, in writing, "
                f"under oath, and within thirty (30) days as required by law."
            )
        elif ds.discovery_type == DiscoveryType.RPD:
            return (
                f"TO {resp_role.upper()} {resp_name} AND "
                f"ATTORNEYS OF RECORD:\n"
                f"\tDemand is hereby made by {prop_role}, "
                f"{prop_name} "
                f'("Propounding Party" or "{prop_role}"), pursuant to '
                f"Code of Civil Procedure section 2031.010, et seq., that "
                f"{resp_role}, {resp_name} "
                f'("{resp_role}" or "Responding Party"), produce and permit '
                f"inspection, photographing, and photocopying of the documents "
                f"and/or inspection, photographing, testing, and sampling of "
                f"other tangible things described herein."
            )
        elif ds.discovery_type == DiscoveryType.RFA:
            return (
                f"TO {resp_role.upper()} {resp_name} AND "
                f"ATTORNEYS OF RECORD:\n"
                f"\tPursuant to California Code of Civil Procedure \u00a72033.010, "
                f"{prop_role}, {prop_name} "
                f'("Propounding Party" or "{prop_role}"), hereby '
                f"requests that {resp_role}, {resp_name} "
                f'("{resp_role}" or "Responding Party"), admit the truth of '
                f"the following matters within thirty (30) days as required by law."
            )
        return ""
