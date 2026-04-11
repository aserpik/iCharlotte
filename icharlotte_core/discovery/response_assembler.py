"""
Response assembler for the discovery response pipeline.

Assembles final .docx files from a caption page template + generated response text.
Reuses existing style constants and helper functions from assembler.py.
"""
import os
import re
from typing import List, Optional, Dict

from docx import Document as DocxDocument
from docx.shared import Pt, Inches, Emu
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .assembler import (
    STYLE_BODY_DOUBLE,
    STYLE_FLUSH_LEFT_DOUBLE,
    STYLE_DISCOVERY_NO,
    STYLE_FLUSH_LEFT,
    STYLE_CENTER_DOUBLE_BOLD_UND,
    SIG_LEFT_INDENT,
    SIG_NAME_LEFT_INDENT,
    _add_para,
    _safe_style,
    DiscoveryAssembler,
)
from .response_parser import ParsedDiscovery
from .response_rules import ResponseRules


# ---------------------------------------------------------------------------
# Discovery type → display name map
# ---------------------------------------------------------------------------

_TYPE_NAMES: Dict[str, str] = {
    "FI": "FORM INTERROGATORIES",
    "SI": "SPECIAL INTERROGATORIES",
    "RFA": "REQUESTS FOR ADMISSION",
    "RPD": "REQUEST FOR PRODUCTION OF DOCUMENTS",
}

# Request header regex patterns (group 1 = number).
# Use (?<!RESPONSE TO ) negative lookbehind to avoid matching inside
# "RESPONSE TO SPECIAL INTERROGATORY NO. X:" etc.
_REQUEST_HEADER_PATTERNS: Dict[str, str] = {
    "FI":  r"(?<!RESPONSE TO )FORM INTERROGATORY NO\.\s*([\d.]+)\s*:",
    "SI":  r"(?<!RESPONSE TO )SPECIAL INTERROGATORY NO\.\s*(\d+)\s*:",
    "RFA": r"(?<!RESPONSE TO )REQUEST FOR ADMISSION NO\.\s*(\d+)\s*:",
    "RPD": r"(?<!RESPONSE TO )REQUEST FOR PRODUCTION NO\.\s*(\d+)\s*:",
}

# Response header regex patterns (group 1 = number)
_RESPONSE_HEADER_PATTERNS: Dict[str, str] = {
    "FI":  r"RESPONSE TO FORM INTERROGATORY NO\.\s*([\d.]+)\s*:",
    "SI":  r"RESPONSE TO SPECIAL INTERROGATORY NO\.\s*(\d+)\s*:",
    "RFA": r"RESPONSE TO REQUEST FOR ADMISSION NO\.\s*(\d+)\s*:",
    "RPD": r"RESPONSE TO REQUEST NO\.\s*(\d+)\s*:",
}

# Title-case set-word map
_SET_TITLE_MAP: Dict[str, str] = {
    "ONE": "First", "TWO": "Second", "THREE": "Third", "FOUR": "Fourth",
    "FIVE": "Fifth", "SIX": "Sixth", "SEVEN": "Seventh", "EIGHT": "Eighth",
    "NINE": "Ninth", "TEN": "Tenth",
}


# ---------------------------------------------------------------------------
# Public helper functions
# ---------------------------------------------------------------------------

def build_response_title(
    responding_name: str,
    responding_role: str,
    propounding_role: str,
    disc_type: str,
    set_word: str,
) -> str:
    """Return the all-caps document title for a discovery response.

    Format: "{RESPONDING_ROLE} {RESPONDING_NAME}'S RESPONSES TO
             {PROPOUNDING_ROLE}'S {TYPE_NAME}, SET {SET_WORD}"
    """
    type_name = _TYPE_NAMES.get(disc_type.upper(), disc_type.upper())
    resp_role = responding_role.upper()
    resp_name = responding_name.upper()
    prop_role = propounding_role.upper()
    sw = set_word.upper()
    return (
        f"{resp_role} {resp_name}'S RESPONSES TO {prop_role}'S "
        f"{type_name}, SET {sw}"
    )


def build_response_filename(
    abbreviation: str,
    disc_type: str,
    set_number: int,
) -> str:
    """Return the standard filename for a response Word document.

    Format: "Def {abbreviation}'s Resp to {disc_type}({set_number}).docx"
    """
    return f"Def {abbreviation}'s Resp to {disc_type}({set_number}).docx"


def parse_response_text(text: str, disc_type: str) -> List[Dict]:
    """Parse plain-text LLM response output into request/response pairs.

    Returns a list of dicts:
        {"number": str, "request": str, "response": str}
    """
    if not text or not text.strip():
        return []

    dtype = disc_type.upper()
    req_pat = _REQUEST_HEADER_PATTERNS.get(dtype)
    resp_pat = _RESPONSE_HEADER_PATTERNS.get(dtype)

    if not req_pat or not resp_pat:
        return []

    req_re = re.compile(req_pat, re.IGNORECASE)
    resp_re = re.compile(resp_pat, re.IGNORECASE)

    # Split the text into segments by finding all header positions
    # We interleave request headers and response headers
    tokens = []  # list of (kind, number, start, end)

    for m in req_re.finditer(text):
        tokens.append(("req", m.group(1), m.end(), None))
    for m in resp_re.finditer(text):
        tokens.append(("resp", m.group(1), m.end(), None))

    # Sort by position in the text
    # We need the raw match positions so re-scan to get start positions
    all_positions = []  # (pos, kind, number, content_start)
    for m in req_re.finditer(text):
        all_positions.append((m.start(), "req", m.group(1), m.end()))
    for m in resp_re.finditer(text):
        all_positions.append((m.start(), "resp", m.group(1), m.end()))

    all_positions.sort(key=lambda x: x[0])

    # Extract text segments between headers
    segments = []
    for i, (pos, kind, number, content_start) in enumerate(all_positions):
        next_start = all_positions[i + 1][0] if i + 1 < len(all_positions) else len(text)
        content = text[content_start:next_start].strip()
        segments.append((kind, number, content))

    # Pair up requests with responses
    pairs = []
    pending_requests: Dict[str, str] = {}

    for kind, number, content in segments:
        if kind == "req":
            pending_requests[number] = content
        elif kind == "resp":
            req_text = pending_requests.pop(number, "")
            pairs.append({
                "number": number,
                "request": req_text,
                "response": content,
            })

    return pairs


# ---------------------------------------------------------------------------
# ResponseAssembler class
# ---------------------------------------------------------------------------

class ResponseAssembler:
    """Assembles discovery response .docx files from a caption page template."""

    def __init__(self, caption_page_path: str):
        if not os.path.isfile(caption_page_path):
            raise FileNotFoundError(
                f"Caption page template not found: {caption_page_path}"
            )
        self.caption_page_path = caption_page_path

    def assemble(
        self,
        parsed: ParsedDiscovery,
        response_text: str,
        rules: ResponseRules,
        output_path: str,
        attorney_name: str = "",
        firm_name: str = "Bordin Semmer LLP",
        date_str: str = "",
        verifier_name: str = "",
    ) -> str:
        """Render a discovery response into a formatted .docx file.

        Steps:
        1. Load caption template
        2. Set caption title (replace CAPTION PAGE)
        3. Insert party block
        4. Insert intro paragraph
        5. Insert PRELIMINARY STATEMENT
        6. Insert general objections (RFA/RPD only)
        7. Insert request/response pairs
        8. Insert verification page
        9. Set footer
        10. Save

        Returns the output_path.
        """
        from datetime import date as _date
        if not date_str:
            date_str = _date.today().strftime("%B %d, %Y")

        disc_type = parsed.discovery_type.upper()
        set_word = parsed.set_word.upper() if parsed.set_word else "ONE"

        # Build the document title
        # Parse responding_party — may be "Defendant ACME CORP" or just "ACME CORP"
        resp_role, resp_name = _split_role_name(parsed.responding_party)
        prop_role, prop_name = _split_role_name(parsed.propounding_party)

        title = build_response_title(
            responding_name=resp_name,
            responding_role=resp_role,
            propounding_role=prop_role,
            disc_type=disc_type,
            set_word=set_word,
        )

        # (1) Load caption template
        doc = DocxDocument(self.caption_page_path)

        # (2) Replace CAPTION PAGE with the title
        self._set_caption_title(doc, title)

        # (3) Insert party block
        self._insert_party_block(doc, parsed)

        # (4) Intro paragraph
        self._insert_intro(doc, parsed, rules)

        # (5) Preliminary statement heading + text
        self._insert_preliminary_statement(doc, disc_type, rules)

        # (6) General objections (RFA/RPD only)
        if disc_type in ("RFA", "RPD"):
            self._insert_general_objections(doc, disc_type, rules)

        # (7) Request/response pairs
        pairs = parse_response_text(response_text, disc_type)
        self._insert_response_pairs(doc, pairs, disc_type, set_word, rules)

        # (8) Verification page
        self._insert_verification(doc, parsed, rules, title, verifier_name, date_str)

        # (9) Footer
        self._set_footer(doc, title)

        # (10) Save
        output_dir = os.path.dirname(output_path)
        if output_dir:
            os.makedirs(output_dir, exist_ok=True)
        doc.save(output_path)
        return output_path

    # -- private helpers ----------------------------------------------------

    def _set_caption_title(self, doc, title: str):
        """Replace the 'CAPTION PAGE' placeholder with the response title.

        Identical logic to DiscoveryAssembler._set_caption_title().
        Title is BOLD but NOT underlined.
        """
        W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        W_T = f"{{{W_NS}}}t"
        W_P = f"{{{W_NS}}}p"
        W_R = f"{{{W_NS}}}r"

        for t_elem in doc.element.body.iter(W_T):
            if t_elem.text and "CAPTION PAGE" in t_elem.text.upper():
                para_elem = t_elem.getparent()
                while para_elem is not None and para_elem.tag != W_P:
                    para_elem = para_elem.getparent()
                if para_elem is None:
                    continue

                for run_elem in list(para_elem.iter(W_R)):
                    para_elem.remove(run_elem)

                new_run = OxmlElement("w:r")
                rPr = OxmlElement("w:rPr")
                rPr.append(OxmlElement("w:b"))
                rFonts = OxmlElement("w:rFonts")
                rFonts.set(qn("w:ascii"), "Times New Roman")
                rFonts.set(qn("w:hAnsi"), "Times New Roman")
                rPr.append(rFonts)
                sz = OxmlElement("w:sz")
                sz.set(qn("w:val"), "24")
                rPr.append(sz)
                szCs = OxmlElement("w:szCs")
                szCs.set(qn("w:val"), "24")
                rPr.append(szCs)
                new_run.append(rPr)
                new_t = OxmlElement("w:t")
                new_t.text = title
                new_t.set(qn("xml:space"), "preserve")
                new_run.append(new_t)
                para_elem.append(new_run)
                return

        # Fallback: 3-column table right cell
        for table in doc.tables:
            if len(table.columns) >= 3:
                cell = table.rows[0].cells[2]
                para = cell.add_paragraph()
                run = para.add_run(title)
                run.bold = True
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)
                return

    def _insert_party_block(self, doc, parsed: ParsedDiscovery):
        """Insert PROPOUNDING / RESPONDING / SET NO. block."""
        LEFT_INDENT = Emu(2400300)
        HANGING_INDENT = Emu(-1943100)

        set_word = parsed.set_word.upper() if parsed.set_word else "ONE"
        set_num = parsed.set_number or 1

        for text in [
            f"PROPOUNDING PARTY:\t{parsed.propounding_party}",
            f"RESPONDING PARTY:\t{parsed.responding_party}",
            f"SET NO.:\t{set_word} ({set_num})",
        ]:
            p = _add_para(doc, text, STYLE_FLUSH_LEFT_DOUBLE)
            p.paragraph_format.left_indent = LEFT_INDENT
            p.paragraph_format.first_line_indent = HANGING_INDENT

        _add_para(doc, "", STYLE_FLUSH_LEFT_DOUBLE)

    def _insert_intro(self, doc, parsed: ParsedDiscovery, rules: ResponseRules):
        """Insert the intro paragraph with placeholders filled."""
        disc_type = parsed.discovery_type.upper()
        template_map = {
            "FI": rules.intro_template_fi,
            "SI": rules.intro_template_si,
            "RFA": rules.intro_template_rfa,
            "RPD": rules.intro_template_rpd,
        }
        template = template_map.get(disc_type, rules.intro_template_si)

        set_word = parsed.set_word.upper() if parsed.set_word else "ONE"
        set_word_title = _SET_TITLE_MAP.get(set_word, set_word.capitalize())

        resp_role, resp_name = _split_role_name(parsed.responding_party)
        prop_role, prop_name = _split_role_name(parsed.propounding_party)
        resp_short = _abbreviate(resp_name)
        prop_short = _abbreviate(prop_name)

        intro_text = template.format(
            responding_party=parsed.responding_party,
            responding_short=resp_short,
            propounding_party=parsed.propounding_party,
            propounding_short=prop_short,
            set_word_title=set_word_title,
        )

        _add_para(doc, intro_text, STYLE_BODY_DOUBLE)

    def _insert_preliminary_statement(
        self, doc, disc_type: str, rules: ResponseRules
    ):
        """Insert PRELIMINARY STATEMENT heading and body text."""
        stmt_map = {
            "FI": rules.preliminary_statement_fi,
            "SI": rules.preliminary_statement_si,
            "RFA": rules.preliminary_statement_rfa,
            "RPD": rules.preliminary_statement_rpd,
        }
        stmt_text = stmt_map.get(disc_type, rules.preliminary_statement_si)

        # Heading — centered
        p = _add_para(doc, "PRELIMINARY STATEMENT", STYLE_CENTER_DOUBLE_BOLD_UND,
                      bold=True, underline=True)
        p.paragraph_format.first_line_indent = Inches(0)

        # Body — split on double newlines for paragraph breaks
        for paragraph in stmt_text.split("\n\n"):
            stripped = paragraph.strip()
            if stripped:
                _add_para(doc, stripped, STYLE_BODY_DOUBLE)

    def _insert_general_objections(
        self, doc, disc_type: str, rules: ResponseRules
    ):
        """Insert GENERAL OBJECTIONS heading and body (RFA/RPD only)."""
        objections_map = {
            "RFA": rules.general_objections_rfa,
            "RPD": rules.general_objections_rpd,
        }
        objections_text = objections_map.get(disc_type, "")
        if not objections_text:
            return

        p = _add_para(doc, "GENERAL OBJECTIONS", STYLE_BODY_DOUBLE,
                      bold=True, underline=True)
        p.paragraph_format.first_line_indent = Inches(0)

        for paragraph in objections_text.split("\n\n"):
            stripped = paragraph.strip()
            if stripped:
                _add_para(doc, stripped, STYLE_BODY_DOUBLE)

    def _insert_response_pairs(
        self,
        doc,
        pairs: List[Dict],
        disc_type: str,
        set_word: str,
        rules: ResponseRules = None,
    ):
        """Insert formatted request/response pairs."""
        header_label_map = {
            "FI":  ("FORM INTERROGATORY NO. {number}:",
                    "RESPONSE TO FORM INTERROGATORY NO. {number}:"),
            "SI":  ("SPECIAL INTERROGATORY NO. {number}:",
                    "RESPONSE TO SPECIAL INTERROGATORY NO. {number}:"),
            "RFA": ("REQUEST FOR ADMISSION NO. {number}:",
                    "RESPONSE TO REQUEST FOR ADMISSION NO. {number}:"),
            "RPD": ("REQUEST FOR PRODUCTION NO. {number}:",
                    "RESPONSE TO REQUEST NO. {number}:"),
        }
        req_tmpl, resp_tmpl = header_label_map.get(
            disc_type, header_label_map["SI"]
        )

        waiver = rules.waiver_language.strip()
        reservation = rules.reservation_clause.strip()

        for pair in pairs:
            num = pair["number"]
            req_text = pair.get("request", "")
            resp_text = pair.get("response", "")

            # Request header
            _add_para(doc, req_tmpl.format(number=num),
                      STYLE_DISCOVERY_NO, bold=True, underline=True)

            # Request body
            if req_text:
                _add_para(doc, req_text, STYLE_BODY_DOUBLE)

            # Response header
            _add_para(doc, resp_tmpl.format(number=num),
                      STYLE_DISCOVERY_NO, bold=True, underline=True)

            # Parse the response body into: objections, waiver, substantive, reservation
            # The plain text may contain these as separate lines or merged
            obj_part, subst_part = self._split_response_parts(
                resp_text, waiver, reservation
            )

            # Objections + waiver on the same paragraph (waiver follows objections inline)
            if obj_part:
                _add_para(doc, f"{obj_part} {waiver}", STYLE_BODY_DOUBLE)
            else:
                # No separate objections found — insert waiver as its own paragraph
                _add_para(doc, waiver, STYLE_BODY_DOUBLE)

            # Substantive response — new paragraph with explicit 0.5" first-line indent
            if subst_part:
                sub_paras = [s.strip() for s in subst_part.split("\n\n") if s.strip()]
                for i, sp in enumerate(sub_paras):
                    if i == len(sub_paras) - 1 and reservation:
                        # Last paragraph: append reservation inline
                        p = _add_para(doc, f"{sp} {reservation}", STYLE_BODY_DOUBLE)
                    else:
                        p = _add_para(doc, sp, STYLE_BODY_DOUBLE)
                    # Ensure first line indent on the first substantive paragraph
                    if i == 0:
                        p.paragraph_format.first_line_indent = Inches(0.5)
            elif reservation:
                _add_para(doc, reservation, STYLE_BODY_DOUBLE)

    def _insert_verification(
        self,
        doc,
        parsed: ParsedDiscovery,
        rules: ResponseRules,
        title: str,
        verifier_name: str,
        date_str: str,
    ):
        """Insert verification page on a new page."""
        # Page break
        page_break_para = doc.add_paragraph()
        run = page_break_para.add_run()
        br = OxmlElement("w:br")
        br.set(qn("w:type"), "page")
        run._element.append(br)

        # Heading
        p = _add_para(doc, "VERIFICATION", STYLE_CENTER_DOUBLE_BOLD_UND,
                      bold=True, underline=True)
        p.paragraph_format.first_line_indent = Inches(0)

        if not verifier_name:
            _, verifier_name = _split_role_name(parsed.responding_party)

        verif_text = rules.verification_template.format(
            verifier_name=verifier_name,
            document_title=title,
        )

        for paragraph in verif_text.split("\n\n"):
            stripped = paragraph.strip()
            if stripped:
                # Signature lines stay flush left (no indent)
                if stripped.startswith("___"):
                    p = _add_para(doc, stripped, STYLE_FLUSH_LEFT)
                    p.paragraph_format.left_indent = SIG_NAME_LEFT_INDENT
                else:
                    _add_para(doc, stripped, STYLE_BODY_DOUBLE)

    @staticmethod
    def _split_response_parts(
        resp_text: str,
        waiver: str,
        reservation: str,
    ):
        """Split response body into (objections_part, substantive_part).

        Splits on the waiver boundary ("Subject to and without waiving...")
        so that the assembler can re-insert waiver/reservation with
        correct inline formatting.  Returns (objections, substantive)
        with the waiver and reservation text stripped out.
        """
        text = resp_text.strip()

        # Remove reservation language if present (we re-insert it later)
        if reservation and reservation in text:
            text = text.replace(reservation, "").strip()

        # Split on the waiver marker — everything before is objections,
        # everything after is substantive.  Do NOT remove waiver first,
        # because we need it as the split point.
        waiver_marker = "Subject to and without waiving"
        idx = text.find(waiver_marker)
        if idx > 0:
            obj_part = text[:idx].strip()
            remainder = text[idx:].strip()
            # Strip the waiver sentence itself from the remainder
            # The waiver ends with ":"
            colon_idx = remainder.find(":")
            if colon_idx >= 0:
                subst_part = remainder[colon_idx + 1:].strip()
            else:
                subst_part = remainder[len(waiver):].strip() if waiver else remainder
            return obj_part, subst_part

        # Also try splitting if waiver was already on its own line
        if waiver and waiver in text:
            parts = text.split(waiver, 1)
            obj_part = parts[0].strip()
            subst_part = parts[1].strip() if len(parts) > 1 else ""
            return obj_part, subst_part

        # No clear split — treat entire text as substantive
        return "", text

    def _set_footer(self, doc, title: str):
        """Set the document title in the footer of all sections.

        Removes ALL existing footer content (including [PLEADING TITLE]
        or any other placeholder text) by deleting paragraph XML elements,
        then inserts a single centered paragraph with the title.
        """
        W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        W_P = f"{{{W_NS}}}p"

        for section in doc.sections:
            footer = section.footer
            footer.is_linked_to_previous = False

            # Clear all text from every paragraph in the footer
            # (handles [PLEADING TITLE] and any other placeholder text)
            for para in footer.paragraphs:
                for run in para.runs:
                    run.text = ""
                # Also clear any direct text nodes in the paragraph XML
                for t_elem in list(para._element.iter(f"{{{W_NS}}}t")):
                    t_elem.text = ""

            # Use the first paragraph if it exists, otherwise add one
            if footer.paragraphs:
                para = footer.paragraphs[0]
                # Clear any leftover runs
                for run in list(para.runs):
                    run._element.getparent().remove(run._element)
            else:
                para = footer.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(title)
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)


# ---------------------------------------------------------------------------
# Private utilities
# ---------------------------------------------------------------------------

def _split_role_name(party_string: str):
    """Split "Defendant ACME CORP" into ("Defendant", "ACME CORP").

    If no recognised role prefix, returns ("", party_string).
    """
    if not party_string:
        return ("", "")
    roles = ("plaintiff", "defendant", "cross-defendant", "cross-complainant",
             "petitioner", "respondent", "claimant")
    stripped = party_string.strip()
    for role in roles:
        if stripped.lower().startswith(role):
            name_part = stripped[len(role):].strip()
            return (role.capitalize(), name_part)
    return ("", stripped)


def _abbreviate(name: str) -> str:
    """Return a short abbreviation (first meaningful word) of a party name."""
    if not name:
        return name
    # Strip leading "THE " or "A "
    clean = re.sub(r'^(THE|A)\s+', '', name.strip(), flags=re.IGNORECASE)
    # Take first word
    first_word = clean.split()[0] if clean.split() else clean
    return first_word.rstrip(".,;")
