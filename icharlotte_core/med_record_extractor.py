"""Med Record Extractor — locate and extract medical record pages from chronology references.

Given prose entries (date + provider), finds the matching rows in a medical chronology
document's 5-column table, resolves the referenced PDF and page range, and extracts
those pages to a subfolder.
"""

import logging
import os
import re
from dataclasses import dataclass, field
from datetime import datetime
from difflib import SequenceMatcher
from typing import Any, List, Optional, Tuple

import fitz  # PyMuPDF
from docx import Document
from PySide6.QtCore import QThread, Signal

from icharlotte_core.subpoena_tracker import _find_folder_ci
from icharlotte_core.word_validator import validate_index_docx

logger = logging.getLogger(__name__)

DATE_PATTERN = re.compile(r'(\d{4})[.\-](\d{2})[.\-](\d{2})')
PAGE_SPEC_RE = re.compile(r'Pg\.?\s*(?:No)?:?\s*(\d+)(?:\s*-\s*(\d+))?/(\d+)', re.IGNORECASE)
# Fallback: bare "number/total" with no Pg prefix
BARE_PAGE_RE = re.compile(r'^(\d+)(?:\s*-\s*(\d+))?/(\d+)$')
WINDOWS_FILENAME_INVALID_RE = re.compile(r'[<>:"/\\|?*\x00-\x1f]+')


def _expand_two_digit_year(year: str) -> str:
    value = int(year)
    return str(2000 + value if value < 50 else 1900 + value)


def _safe_filename_part(value: str, *, max_length: int = 40, fallback: str = "record") -> str:
    safe = WINDOWS_FILENAME_INVALID_RE.sub("-", value or "")
    safe = re.sub(r"\s+", " ", safe).strip(" .-")
    if not safe:
        safe = fallback
    return safe[:max_length].strip(" .-") or fallback


def _unique_output_path(output_dir: str, out_name: str, used_paths: set[str], row: Any) -> str:
    out_path = os.path.join(output_dir, out_name)
    if out_path not in used_paths:
        used_paths.add(out_path)
        return out_path

    stem, ext = os.path.splitext(out_name)
    row_suffix = _safe_filename_part(
        getattr(row, "id", "") or f"row-{getattr(row, 'order', 0) + 1}",
        max_length=24,
    )
    candidate = os.path.join(output_dir, f"{stem} - {row_suffix}{ext}")
    counter = 2
    while candidate in used_paths:
        candidate = os.path.join(output_dir, f"{stem} - {row_suffix}-{counter}{ext}")
        counter += 1
    used_paths.add(candidate)
    return candidate


@dataclass
class MedChronRow:
    """One row from the 5-column med chron table."""
    date: str            # raw date string, e.g. "06/12/2003"
    page_no: str         # raw PAGE NO cell text
    provider: str        # raw PROVIDER cell text
    description: str     # raw DESCRIPTION cell text
    flags: str           # raw Red flags / Comments

    # Parsed from page_no
    record_filename: str = ""
    page_start: int = 0
    page_end: int = 0


@dataclass
class ParsedEntry:
    """An entry parsed from user's prose input."""
    date: str         # MM/DD/YYYY
    provider: str     # provider or facility name
    raw_text: str     # original prose paragraph


@dataclass
class ExtractionResult:
    """Result of extracting pages for one entry or selected chronology row."""
    entry: Optional[ParsedEntry] = None
    matched_row: Optional[object] = None
    output_path: Optional[str] = None
    error: Optional[str] = None


def _normalize_date(date_str: str) -> str:
    """Normalize various date formats to MM/DD/YYYY for comparison."""
    date_str = date_str.strip()
    normalized = re.sub(r'\s+', ' ', date_str)
    normalized = re.sub(r'\b([A-Za-z]{3,9})\.', r'\1', normalized)
    normalized = re.sub(r'\bSept\b', 'Sep', normalized, flags=re.IGNORECASE)
    m = re.match(r'^(\d{1,2})[/-](\d{1,2})[/-](\d{2})$', normalized)
    if m:
        return f"{int(m.group(1)):02d}/{int(m.group(2)):02d}/{_expand_two_digit_year(m.group(3))}"
    m = re.match(r'^([A-Za-z]+)\s+(\d{1,2}),?\s+(\d{2})$', normalized)
    if m:
        expanded = f"{m.group(1)} {int(m.group(2))} {_expand_two_digit_year(m.group(3))}"
        for fmt in ('%B %d %Y', '%b %d %Y'):
            try:
                return datetime.strptime(expanded, fmt).strftime('%m/%d/%Y')
            except ValueError:
                pass
    for fmt in (
        '%m/%d/%Y',
        '%m-%d-%Y',
        '%Y-%m-%d',
        '%Y/%m/%d',
        '%Y.%m.%d',
        '%B %d, %Y',
        '%B %d %Y',
        '%b %d, %Y',
        '%b %d %Y',
    ):
        try:
            return datetime.strptime(normalized, fmt).strftime('%m/%d/%Y')
        except ValueError:
            pass
    # Already MM/DD/YYYY
    if re.match(r'^\d{2}/\d{2}/\d{4}$', date_str):
        return date_str
    # YYYY-MM-DD or YYYY.MM.DD
    m = re.match(r'^(\d{4})[.\-/](\d{1,2})[.\-/](\d{1,2})$', date_str)
    if m:
        return f"{int(m.group(2)):02d}/{int(m.group(3)):02d}/{m.group(1)}"
    # M/D/YYYY
    m = re.match(r'^(\d{1,2})/(\d{1,2})/(\d{4})$', date_str)
    if m:
        return f"{int(m.group(1)):02d}/{int(m.group(2)):02d}/{m.group(3)}"
    return date_str


def _parse_page_no(cell_text: str) -> Tuple[str, int, int]:
    """Parse a PAGE NO cell into (filename, start_page, end_page).

    Examples:
        '60337-0014_ ORTHOPEDICS & SPORTS MEDICINE.pdf\\n\\nPg: 23-24/69'
        'SALTARELLI000001-SALTARELLI000772\\n\\nPg. No: 501-505/772'
        'homedepot\\n\\n93/535'
    """
    parts = cell_text.strip().split('\n')
    # Filter out empty parts
    parts = [p.strip() for p in parts if p.strip()]
    if not parts:
        return ("", 0, 0)

    filename = parts[0]
    page_start = 0
    page_end = 0

    # Look for page spec in remaining parts
    for part in parts[1:]:
        m = PAGE_SPEC_RE.search(part)
        if m:
            page_start = int(m.group(1))
            page_end = int(m.group(2)) if m.group(2) else page_start
            break
        m = BARE_PAGE_RE.search(part)
        if m:
            page_start = int(m.group(1))
            page_end = int(m.group(2)) if m.group(2) else page_start
            break

    return (filename, page_start, page_end)


def _build_file_index(case_path: str) -> dict:
    """Walk RECORDS/ and DISCOVERY/RESPONSES/ to build filename-to-path lookup."""
    file_index = {}
    search_folders = [
        _find_folder_ci(case_path, "RECORDS"),
        _find_folder_ci(case_path, "DISCOVERY", "RESPONSES"),
    ]
    for folder in search_folders:
        if folder is None:
            continue
        for root, _dirs, files in os.walk(folder):
            for fname in files:
                key = os.path.splitext(fname)[0].lower().strip()
                if key and key not in file_index:
                    file_index[key] = os.path.join(root, fname)
    return file_index


def _lookup_file(file_index: dict, filename: str) -> Optional[str]:
    """Look up a filename in the index with fuzzy fallback (>0.85 ratio)."""
    lookup_key = os.path.splitext(filename)[0].lower().strip()
    found = file_index.get(lookup_key)
    if found:
        return found
    best_ratio = 0.0
    best_path = None
    for indexed_key, indexed_path in file_index.items():
        ratio = SequenceMatcher(None, lookup_key, indexed_key).ratio()
        if ratio > best_ratio:
            best_ratio = ratio
            best_path = indexed_path
    if best_ratio >= 0.85:
        return best_path
    return None


def _find_most_recent_med_chron(case_path: str) -> Optional[str]:
    """Find the most recent medical chronology .docx in the Medical Summary folder."""
    # Try exact names first, then fall back to pattern match
    for name in (
        "Medical Summary – DO NOT PRODUCE",
        "Medical Summary - DO NOT PRODUCE",
        "Med. Summary - DO NOT PRODUCE",
        "MEDICAL SUMMARY - DO NOT PRODUCE",
    ):
        folder = _find_folder_ci(case_path, "RECORDS", name)
        if folder is not None:
            break
    else:
        # Fallback: scan RECORDS for any subfolder containing "med" and "summary"
        records = _find_folder_ci(case_path, "RECORDS")
        if records:
            for entry in os.scandir(records):
                if entry.is_dir():
                    low = entry.name.lower()
                    if "med" in low and "summ" in low:
                        folder = entry.path
                        break
            else:
                folder = None
        else:
            folder = None
    if folder is None:
        return None

    best_date = ""
    best_path = None
    for root, _dirs, files in os.walk(folder):
        for fname in files:
            if not fname.lower().endswith('.docx') or fname.startswith('~$'):
                continue
            m = DATE_PATTERN.search(fname)
            if m:
                date_str = f"{m.group(1)}-{m.group(2)}-{m.group(3)}"
                if date_str > best_date:
                    best_date = date_str
                    best_path = os.path.join(root, fname)
    return best_path


def _parse_med_chron_table(docx_path: str) -> List[MedChronRow]:
    """Parse the 5-column table from a medical chronology document."""
    doc = Document(docx_path)
    rows_out = []
    for table in doc.tables:
        if not table.rows:
            continue
        if len(table.rows[0].cells) != 5:
            continue
        # Found the 5-column table
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            # Skip header row
            if cells[0].upper() == "DATE":
                continue
            # Skip merged section header rows (all cells identical)
            if len(set(cells)) == 1:
                continue
            # Skip empty date
            if not cells[0]:
                continue
            record_filename, page_start, page_end = _parse_page_no(cells[1])
            rows_out.append(MedChronRow(
                date=cells[0],
                page_no=cells[1],
                provider=cells[2],
                description=cells[3],
                flags=cells[4],
                record_filename=record_filename,
                page_start=page_start,
                page_end=page_end,
            ))
        break  # Use first 5-column table only
    return rows_out


def _match_entry_to_row(entry: ParsedEntry, rows: List[MedChronRow]) -> Optional[MedChronRow]:
    """Match a parsed entry to the best med chron row by date + provider."""
    norm_entry_date = _normalize_date(entry.date)

    # Find all rows with matching date
    date_matches = []
    for row in rows:
        if _normalize_date(row.date) == norm_entry_date:
            date_matches.append(row)

    if not date_matches:
        return None
    if len(date_matches) == 1:
        return date_matches[0]

    # Multiple date matches — disambiguate by provider
    entry_provider = entry.provider.lower()
    best_ratio = 0.0
    best_row = date_matches[0]
    for row in date_matches:
        # Check against all lines in the provider cell
        row_provider = row.provider.lower().replace('\n', ' ')
        ratio = SequenceMatcher(None, entry_provider, row_provider).ratio()
        if ratio > best_ratio:
            best_ratio = ratio
            best_row = row
    return best_row


OCR_CHAR_THRESHOLD = 50  # chars per page below which OCR is triggered


def _ocr_pdf_if_needed(pdf_path: str) -> bool:
    """Add an OCR text layer to a PDF if pages lack selectable text.

    Renders each page to an image, runs tesseract to produce a text-layer PDF,
    then merges the invisible text from tesseract's PDF onto the original page.
    Returns True if OCR was applied.
    """
    try:
        import pytesseract
    except ImportError:
        logger.warning("pytesseract not available — skipping OCR")
        return False

    doc = fitz.open(pdf_path)
    ocr_done = False
    try:
        for page_idx in range(len(doc)):
            page = doc[page_idx]
            text = page.get_text().strip()
            if len(text) >= OCR_CHAR_THRESHOLD:
                continue

            logger.info(f"OCR needed for page {page_idx + 1} of {os.path.basename(pdf_path)}")

            # Render page to image at 300 DPI
            pix = page.get_pixmap(dpi=300)
            from PIL import Image
            import io
            img = Image.open(io.BytesIO(pix.tobytes("png")))

            # Use tesseract's native PDF output — it places text with exact positioning
            pdf_bytes = pytesseract.image_to_pdf_or_hocr(img, extension='pdf')
            if not pdf_bytes:
                continue

            # Open tesseract's PDF and extract text to verify it found something
            tess_doc = fitz.open(stream=pdf_bytes, filetype="pdf")
            tess_text = tess_doc[0].get_text().strip()
            if len(tess_text) < 10:
                tess_doc.close()
                continue

            # Merge tesseract's text layer onto the original page.
            # Tesseract PDF page is image-sized (300 DPI pixels as points).
            # We need to scale it to match the original page dimensions.
            tess_page = tess_doc[0]
            tess_rect = tess_page.rect
            scale_x = page.rect.width / tess_rect.width
            scale_y = page.rect.height / tess_rect.height

            # Extract text dict with positioning from tesseract PDF
            text_dict = tess_page.get_text("dict")
            for block in text_dict.get("blocks", []):
                if block.get("type") != 0:  # text blocks only
                    continue
                for line in block.get("lines", []):
                    for span in line.get("spans", []):
                        span_text = span.get("text", "").strip()
                        if not span_text:
                            continue
                        origin = span.get("origin", (0, 0))
                        fs = span.get("size", 10)

                        # Scale coordinates and font size to match original page
                        scaled_x = origin[0] * scale_x
                        scaled_y = origin[1] * scale_y
                        scaled_fs = fs * min(scale_x, scale_y)

                        # Also scale font to match word width
                        bbox = span.get("bbox", (0, 0, 0, 0))
                        span_width_orig = (bbox[2] - bbox[0]) * scale_x
                        font_obj = fitz.Font("helv")
                        rendered_width = font_obj.text_length(span_text, fontsize=scaled_fs)
                        if rendered_width > 0 and span_width_orig > 0:
                            scaled_fs = scaled_fs * (span_width_orig / rendered_width)

                        try:
                            page.insert_text(
                                (scaled_x, scaled_y),
                                span_text,
                                fontsize=max(scaled_fs, 2),
                                render_mode=3,  # invisible
                                overlay=True,
                            )
                        except Exception:
                            pass

            tess_doc.close()
            ocr_done = True

        if ocr_done:
            import tempfile
            tmp_fd, tmp_path = tempfile.mkstemp(suffix=".pdf", dir=os.path.dirname(pdf_path))
            os.close(tmp_fd)
            try:
                doc.save(tmp_path, garbage=3, deflate=True)
                doc.close()
                doc = None
                os.replace(tmp_path, pdf_path)
            except Exception:
                if os.path.exists(tmp_path):
                    os.remove(tmp_path)
                raise
    finally:
        if doc is not None:
            doc.close()

    return ocr_done


INDEX_DOC_NAME = "Index of Extracted Records.docx"
HEADER_ROW = ["DATE", "PAGE NO", "PROVIDER", "DESCRIPTION"]


def _update_index_document(output_dir: str, new_rows: List[Any]) -> None:
    """Create or append to the cumulative index document in the output folder."""
    index_path = os.path.join(output_dir, INDEX_DOC_NAME)

    from docx.shared import Inches

    # Column widths: Date, Page No, Provider are compact; Description gets the rest
    COL_WIDTHS = [Inches(1.0), Inches(0.8), Inches(1.5), Inches(6.2)]

    def _apply_col_widths(tbl):
        for row in tbl.rows:
            for i, width in enumerate(COL_WIDTHS):
                row.cells[i].width = width

    if os.path.exists(index_path):
        doc = Document(index_path)
        # Find the existing 4-column table (or legacy 5-column)
        table = None
        for t in doc.tables:
            if t.rows and len(t.rows[0].cells) in (4, 5):
                table = t
                break
        if table is None:
            doc = Document()
            doc.add_heading("Index of Extracted Records", level=1)
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            for i, hdr in enumerate(HEADER_ROW):
                table.rows[0].cells[i].text = hdr
            _apply_col_widths(table)
    else:
        doc = Document()
        doc.add_heading("Index of Extracted Records", level=1)
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        for i, hdr in enumerate(HEADER_ROW):
            cell = table.rows[0].cells[i]
            cell.text = hdr
            for run in cell.paragraphs[0].runs:
                run.bold = True
        _apply_col_widths(table)

    # Collect existing dates+page_no to avoid duplicates
    existing = set()
    for row in table.rows[1:]:
        key = (row.cells[0].text.strip(), row.cells[1].text.strip())
        existing.add(key)

    num_cols = len(table.rows[0].cells)
    for med_row in new_rows:
        key = (med_row.date, med_row.page_no)
        if key in existing:
            continue
        row_cells = table.add_row().cells
        row_cells[0].text = med_row.date
        row_cells[1].text = med_row.page_no
        row_cells[2].text = med_row.provider
        row_cells[3].text = med_row.description

    _apply_col_widths(table)
    doc.save(index_path)
    result = validate_index_docx(index_path)
    if result.has_errors:
        result.print_summary()
        raise ValueError(f"Index validation failed: {result.error_count} error(s) in {index_path}")


def _extract_pages(pdf_path: str, page_start: int, page_end: int, output_path: str) -> None:
    """Extract a page range from a PDF and save to output_path."""
    doc = fitz.open(pdf_path)
    try:
        # Convert to 0-indexed
        start_idx = page_start - 1
        end_idx = page_end  # fitz.Document.select uses exclusive end via list
        if start_idx < 0:
            start_idx = 0
        if end_idx > len(doc):
            end_idx = len(doc)
        page_indices = list(range(start_idx, end_idx))
        if not page_indices:
            raise ValueError(f"No valid pages in range {page_start}-{page_end}")
        new_doc = fitz.open()
        new_doc.insert_pdf(doc, from_page=start_idx, to_page=end_idx - 1)
        new_doc.save(output_path)
        new_doc.close()
    finally:
        doc.close()


class MedRecordExtractorWorker(QThread):
    """Background worker that extracts medical record pages based on user entries."""

    progress = Signal(str)
    finished_result = Signal(bool, str)
    warning = Signal(str)

    def __init__(
        self,
        case_path: str,
        file_number: str,
        user_text: str = "",
        parent=None,
        *,
        chronology_path: str = "",
        selected_rows: Optional[List[Any]] = None,
    ):
        super().__init__(parent)
        self.case_path = case_path
        self.file_number = file_number
        self.user_text = user_text
        self.chronology_path = chronology_path
        self._selected_rows_provided = selected_rows is not None
        self.selected_rows = list(selected_rows or [])

    def run(self):
        try:
            results = self._execute()
            successes = [r for r in results if r.output_path]
            failures = [r for r in results if r.error]
            summary_parts = []
            if successes:
                summary_parts.append(f"{len(successes)} record(s) extracted")
            if failures:
                summary_parts.append(f"{len(failures)} failed")
                for f in failures:
                    row = f.matched_row
                    if row is not None:
                        summary_parts.append(f"  - {row.date} {row.provider}: {f.error}")
                    elif f.entry is not None:
                        summary_parts.append(f"  - {f.entry.date} {f.entry.provider}: {f.error}")
                    else:
                        summary_parts.append(f"  - {f.error}")
            self.finished_result.emit(True, "\n".join(summary_parts) if summary_parts else "No entries processed")
        except Exception as e:
            logger.exception("Med record extraction failed")
            self.finished_result.emit(False, str(e))

    def _execute(self) -> List[ExtractionResult]:
        if self._selected_rows_provided:
            return self._execute_selected_rows(self.selected_rows)

        # Step 1: Parse entries from prose using LLM
        self.progress.emit("Parsing entries from text...")
        entries = self._parse_entries(self.user_text)
        if not entries:
            self.finished_result.emit(False, "Could not parse any entries from the input text.")
            return []
        self.progress.emit(f"Found {len(entries)} entries")

        # Step 2: Find most recent med chron
        self.progress.emit("Finding medical chronology...")
        med_chron_path = _find_most_recent_med_chron(self.case_path)
        if not med_chron_path:
            self.finished_result.emit(False, "No medical chronology document found in RECORDS/Medical Summary folder.")
            return []
        self.progress.emit(f"Using: {os.path.basename(med_chron_path)}")

        # Step 3: Parse the 5-column table
        self.progress.emit("Reading chronology table...")
        table_rows = _parse_med_chron_table(med_chron_path)
        if not table_rows:
            self.finished_result.emit(False, "No 5-column table found in the medical chronology document.")
            return []
        self.progress.emit(f"Table has {len(table_rows)} entries")

        # Step 4: Build file index
        self.progress.emit("Building file index...")
        file_index = _build_file_index(self.case_path)

        # Step 5: Output directory
        output_dir = os.path.join(self.case_path, "NOTES", "AI OUTPUT", "Med Record Extracts")
        os.makedirs(output_dir, exist_ok=True)

        # Step 6: Match and extract each entry
        results = []
        for i, entry in enumerate(entries, 1):
            self.progress.emit(f"Processing entry {i}/{len(entries)}: {entry.date} - {entry.provider}")
            result = ExtractionResult(entry=entry)

            # Match to table row
            matched = _match_entry_to_row(entry, table_rows)
            if not matched:
                result.error = "No matching row found in chronology table"
                results.append(result)
                self.warning.emit(f"No match for {entry.date} - {entry.provider}")
                continue
            result.matched_row = matched

            if not matched.record_filename or matched.page_start == 0:
                result.error = f"Could not parse record/pages from PAGE NO: {matched.page_no[:60]}"
                results.append(result)
                continue

            # Find the PDF
            pdf_path = _lookup_file(file_index, matched.record_filename)
            if not pdf_path:
                result.error = f"PDF not found: {matched.record_filename}"
                results.append(result)
                self.warning.emit(f"PDF not found: {matched.record_filename}")
                continue

            # Extract pages
            provider_short = entry.provider[:40].replace('/', '-').replace('\\', '-')
            date_safe = _normalize_date(entry.date).replace('/', '-')
            if matched.page_start == matched.page_end:
                page_label = f"p{matched.page_start}"
            else:
                page_label = f"pp{matched.page_start}-{matched.page_end}"
            out_name = f"{date_safe} - {provider_short} - {page_label}.pdf"
            out_path = os.path.join(output_dir, out_name)

            try:
                _extract_pages(pdf_path, matched.page_start, matched.page_end, out_path)
                result.output_path = out_path
                self.progress.emit(f"Extracted: {out_name}")
                # OCR if needed
                try:
                    if _ocr_pdf_if_needed(out_path):
                        self.progress.emit(f"OCR applied: {out_name}")
                except Exception as ocr_err:
                    self.warning.emit(f"OCR failed for {out_name}: {ocr_err}")
            except Exception as e:
                result.error = f"Page extraction failed: {e}"
                self.warning.emit(f"Extraction error for {entry.date}: {e}")

            results.append(result)

        # Step 7: Update the index document with matched rows
        matched_rows = [r.matched_row for r in results if r.matched_row]
        if matched_rows:
            self.progress.emit("Updating index document...")
            try:
                _update_index_document(output_dir, matched_rows)
            except Exception as e:
                self.warning.emit(f"Failed to update index document: {e}")
                logger.exception("Index document update failed")

        return results

    def _execute_selected_rows(self, selected_rows: List[Any]) -> List[ExtractionResult]:
        self.progress.emit(f"Preparing to extract {len(selected_rows)} selected chronology row(s)...")
        file_index = _build_file_index(self.case_path)
        output_dir = os.path.join(self.case_path, "NOTES", "AI OUTPUT", "Med Record Extracts")
        os.makedirs(output_dir, exist_ok=True)

        results: List[ExtractionResult] = []
        used_output_paths: set[str] = set()
        for i, row in enumerate(selected_rows, 1):
            self.progress.emit(f"Processing row {i}/{len(selected_rows)}: {row.date} - {row.provider}")
            result = ExtractionResult(matched_row=row)
            if not row.record_filename or row.page_start <= 0:
                result.error = f"Could not parse record/pages from PAGE NO: {row.page_no[:60]}"
                results.append(result)
                continue

            pdf_path = _lookup_file(file_index, row.record_filename)
            if not pdf_path:
                result.error = f"PDF not found: {row.record_filename}"
                results.append(result)
                self.warning.emit(f"PDF not found: {row.record_filename}")
                continue

            provider_short = _safe_filename_part(row.provider, max_length=40, fallback="provider")
            date_safe = _normalize_date(row.date).replace("/", "-")
            page_label = f"p{row.page_start}" if row.page_start == row.page_end else f"pp{row.page_start}-{row.page_end}"
            out_name = f"{date_safe} - {provider_short} - {page_label}.pdf"
            out_path = _unique_output_path(output_dir, out_name, used_output_paths, row)
            out_name = os.path.basename(out_path)

            try:
                _extract_pages(pdf_path, row.page_start, row.page_end, out_path)
                result.output_path = out_path
                self.progress.emit(f"Extracted: {out_name}")
                try:
                    if _ocr_pdf_if_needed(out_path):
                        self.progress.emit(f"OCR applied: {out_name}")
                except Exception as ocr_err:
                    self.warning.emit(f"OCR failed for {out_name}: {ocr_err}")
            except Exception as exc:
                result.error = f"Page extraction failed: {exc}"
                self.warning.emit(f"Extraction error for {row.date}: {exc}")
            results.append(result)

        matched_rows = [r.matched_row for r in results if r.matched_row and not r.error]
        if matched_rows:
            self.progress.emit("Updating index document...")
            try:
                _update_index_document(output_dir, matched_rows)
            except Exception as exc:
                self.warning.emit(f"Failed to update index document: {exc}")
                logger.exception("Index document update failed")
        return results

    def _parse_entries(self, text: str) -> List[ParsedEntry]:
        """Use LLM to parse prose entries into structured date + provider pairs."""
        from icharlotte_core.llm_config import LLMCaller

        prompt = """Extract every medical treatment entry from the text below.
For each entry, return a JSON array of objects with these fields:
- "date": the date of treatment in MM/DD/YYYY format
- "provider": the name of the medical provider, doctor, or facility
- "raw": the first 80 characters of the original text for that entry

Return ONLY valid JSON — no markdown, no commentary. Example:
[{"date": "03/22/2016", "provider": "Adventist Health", "raw": "On March 22, 2016, Plaintiff was evaluated by..."}]"""

        caller = LLMCaller()
        response = caller.call(
            prompt=prompt,
            text=text,
            agent_id="agent_med_rec",
        )

        if not response:
            return []

        # Parse JSON from response
        import json
        # Strip markdown code fences if present
        cleaned = response.strip()
        if cleaned.startswith('```'):
            lines = cleaned.split('\n')
            lines = lines[1:]  # drop opening fence
            if lines and lines[-1].strip() == '```':
                lines = lines[:-1]
            cleaned = '\n'.join(lines)

        try:
            items = json.loads(cleaned)
        except json.JSONDecodeError:
            logger.error(f"LLM returned invalid JSON: {cleaned[:200]}")
            return []

        entries = []
        for item in items:
            if isinstance(item, dict) and 'date' in item and 'provider' in item:
                entries.append(ParsedEntry(
                    date=item['date'],
                    provider=item['provider'],
                    raw_text=item.get('raw', ''),
                ))
        return entries
