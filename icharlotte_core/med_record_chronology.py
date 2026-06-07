"""Parse medical chronology summary documents for Med Record Extractor."""
from __future__ import annotations

import hashlib
import os
import re
from collections.abc import Iterator
from dataclasses import dataclass, field
from typing import Literal

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

from icharlotte_core.med_record_extractor import _normalize_date, _parse_page_no


MatchStatus = Literal["confident", "ambiguous", "none"]


@dataclass(frozen=True)
class SynopsisParagraph:
    id: str
    order: int
    text: str
    warning: str = ""


@dataclass(frozen=True)
class SelectableChronologyRow:
    id: str
    order: int
    date: str
    page_no: str
    provider: str
    description: str
    flags: str
    record_filename: str = ""
    page_start: int = 0
    page_end: int = 0
    warning: str = ""

    @property
    def extractable(self) -> bool:
        return _is_extractable_page_range(
            self.record_filename,
            self.page_start,
            self.page_end,
        )


@dataclass(frozen=True)
class ChronologyDocument:
    source_path: str
    synopsis_paragraphs: list[SynopsisParagraph] = field(default_factory=list)
    rows: list[SelectableChronologyRow] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)
    blocking_errors: list[str] = field(default_factory=list)


@dataclass
class SelectionState:
    selected_paragraph_ids: set[str] = field(default_factory=set)
    _row_sources: dict[str, set[str]] = field(default_factory=dict, repr=False)

    def select_paragraph(self, paragraph_id: str) -> None:
        self.selected_paragraph_ids.add(paragraph_id)

    def deselect_paragraph(self, paragraph_id: str) -> None:
        self.selected_paragraph_ids.discard(paragraph_id)
        self.clear_source(paragraph_id)

    def select_row(self, row_id: str, *, source: str = "manual") -> None:
        self._row_sources.setdefault(row_id, set()).add(source)

    def deselect_row(self, row_id: str, *, source: str = "manual") -> None:
        sources = self._row_sources.get(row_id)
        if not sources:
            return
        sources.discard(source)
        if not sources:
            del self._row_sources[row_id]

    def clear_source(self, source: str) -> None:
        self.selected_paragraph_ids.discard(source)
        for row_id, sources in list(self._row_sources.items()):
            sources.discard(source)
            if not sources:
                del self._row_sources[row_id]

    def is_row_selected(self, row_id: str) -> bool:
        return row_id in self._row_sources

    def selected_row_ids(self) -> list[str]:
        return list(self._row_sources)


@dataclass(frozen=True)
class MatchResult:
    status: MatchStatus
    row_ids: tuple[str, ...] = ()
    candidate_row_ids: tuple[str, ...] = ()
    reason: str = ""


_SYNOPSIS_HEADING_RE = re.compile(
    r"^BRIEF\s+SYNOPSIS\s+OF\s+POST[-\s]INJURY\s+MEDICAL\s+RECORD:?\s*$",
    re.I,
)
_CHRON_HEADER_ALIASES = (
    frozenset({"date"}),
    frozenset({"pageno", "pgno", "pagenumber"}),
    frozenset({"provider"}),
    frozenset({"description"}),
    frozenset({"redflagscomments", "redflagcomments"}),
)
_DATE_RE = re.compile(
    r"\b(?:\d{1,2}[/-]\d{1,2}[/-]\d{4}|\d{4}[./-]\d{1,2}[./-]\d{1,2})\b"
)
_CREDENTIAL_RE = re.compile(
    r"\b[A-Z][A-Za-z.'-]+(?:\s+[A-Z][A-Za-z.'-]+){1,4},?\s*"
    r"(?:M\.?D\.?|D\.?O\.?|D\.?C\.?|P\.?A\.?-?C|N\.?P\.?|"
    r"D\.?P\.?T\.?|P\.?T\.?|R\.?N\.?|FNP|PA-C)\b"
)
_PROVIDER_CONTEXT_RE = re.compile(
    r"\b(?:presented to|reported to|returned to|treated at|treated by|"
    r"evaluated by|examined by|seen by|consulted with|followed up with|"
    r"followed up at|visited|saw|with|at|to|by|from)\s+"
    r"([A-Z][^.;:]+?)(?=(?:\s+for\b|\s+regarding\b|\s+where\b|[.;:]|$))",
    re.I,
)
_PROVIDER_TOKEN_STOPWORDS = {
    "a",
    "an",
    "and",
    "by",
    "care",
    "dr",
    "for",
    "from",
    "he",
    "her",
    "him",
    "injury",
    "medical",
    "mr",
    "mrs",
    "ms",
    "on",
    "plaintiff",
    "she",
    "test",
    "the",
    "to",
    "was",
}


def parse_chronology_document(path: str) -> ChronologyDocument:
    warnings: list[str] = []
    blocking_errors: list[str] = []
    synopsis = _parse_synopsis(path)
    rows = _parse_rows(path)
    if not synopsis:
        warnings.append("No Brief Synopsis section found.")
    if not rows:
        blocking_errors.append("No usable 5-column chronology table found.")
    return ChronologyDocument(
        source_path=os.path.normpath(path),
        synopsis_paragraphs=synopsis,
        rows=rows,
        warnings=warnings,
        blocking_errors=blocking_errors,
    )


def match_synopsis_to_rows(
    paragraph: SynopsisParagraph,
    rows: list[SelectableChronologyRow],
) -> MatchResult:
    dates = {_normalize_date(match.group(0)) for match in _DATE_RE.finditer(paragraph.text)}
    dates.discard("")
    if not dates:
        return MatchResult(status="none", reason="No date found in synopsis paragraph.")

    same_date_rows = [row for row in rows if _normalize_date(row.date) in dates]
    if not same_date_rows:
        return MatchResult(status="none", reason="No chronology rows share the synopsis date.")

    candidates = _provider_candidates(paragraph.text)
    if not candidates:
        return MatchResult(
            status="ambiguous",
            candidate_row_ids=tuple(row.id for row in same_date_rows),
            reason="No distinct provider candidate found in synopsis paragraph.",
        )

    scored_rows: list[tuple[int, SelectableChronologyRow]] = []
    for row in same_date_rows:
        score = sum(_provider_score(candidate, row.provider) for candidate in candidates)
        if score > 0:
            scored_rows.append((score, row))

    if not scored_rows:
        return MatchResult(status="none", reason="No same-date row provider matched the synopsis.")

    scored_rows.sort(key=lambda item: (-item[0], item[1].order))
    best_score, best_row = scored_rows[0]
    second_score = scored_rows[1][0] if len(scored_rows) > 1 else 0
    if best_score >= 80 and best_score >= second_score + 40:
        return MatchResult(
            status="confident",
            row_ids=(best_row.id,),
            reason="One same-date provider row was the distinct strongest match.",
        )

    return MatchResult(
        status="ambiguous",
        candidate_row_ids=tuple(row.id for _, row in scored_rows),
        reason="Multiple same-date rows were plausible provider matches.",
    )


def _provider_candidates(text: str) -> tuple[str, ...]:
    candidates: list[str] = []
    seen: set[str] = set()
    for pattern in (_CREDENTIAL_RE, _PROVIDER_CONTEXT_RE):
        for match in pattern.finditer(text or ""):
            candidate = _clean_provider_candidate(match.group(1) if match.groups() else match.group(0))
            key = _provider_key(candidate)
            if candidate and key not in seen:
                seen.add(key)
                candidates.append(candidate)
    return tuple(candidates)


def _clean_provider_candidate(text: str) -> str:
    candidate = _collapse(text)
    candidate = re.sub(r"^(?:the\s+|dr\.?\s+)", "", candidate, flags=re.I)
    candidate = re.sub(r"\s+(?:for|regarding|where)\b.*$", "", candidate, flags=re.I)
    candidate = candidate.strip(" ,.;:-")
    if len(candidate) < 3:
        return ""
    key = _provider_key(candidate)
    tokens = key.split()
    if not tokens or all(token in _PROVIDER_TOKEN_STOPWORDS for token in tokens):
        return ""
    return candidate


def _provider_score(candidate: str, provider: str) -> int:
    candidate_key = _provider_key(candidate)
    provider_key = _provider_key(provider)
    if not candidate_key or not provider_key:
        return 0
    if candidate_key == provider_key:
        return 140
    if candidate_key in provider_key:
        return 100 + min(40, len(candidate_key) // 3)

    candidate_tokens = [
        token
        for token in candidate_key.split()
        if len(token) > 1 and token not in _PROVIDER_TOKEN_STOPWORDS
    ]
    if not candidate_tokens:
        return 0
    provider_tokens = set(provider_key.split())
    overlap = [token for token in candidate_tokens if token in provider_tokens]
    if not overlap:
        return 0
    if len(overlap) == len(candidate_tokens):
        return 80 + (len(overlap) * 5)
    if len(overlap) >= 2 and len(overlap) / len(candidate_tokens) >= 0.67:
        return 50 + int(30 * (len(overlap) / len(candidate_tokens)))
    return 0


def _parse_synopsis(path: str) -> list[SynopsisParagraph]:
    doc = Document(path)
    in_synopsis = False
    paragraphs: list[SynopsisParagraph] = []
    for block in _iter_body_blocks(doc):
        if isinstance(block, Table):
            if in_synopsis and _is_chronology_table(block):
                break
            continue

        text = _collapse(block.text)
        if not text:
            continue
        if _SYNOPSIS_HEADING_RE.match(text):
            in_synopsis = True
            continue
        if not in_synopsis:
            continue

        order = len(paragraphs)
        paragraphs.append(
            SynopsisParagraph(
                id=_stable_id("syn", order, text),
                order=order,
                text=text,
            )
        )
    return paragraphs


def _parse_rows(path: str) -> list[SelectableChronologyRow]:
    doc = Document(path)
    for table in doc.tables:
        if not _is_chronology_table(table):
            continue

        rows: list[SelectableChronologyRow] = []
        for raw_row in table.rows[1:]:
            cells = [_collapse(cell.text) for cell in raw_row.cells]
            if not cells[0]:
                continue
            if len(set(cells)) == 1:
                continue
            record_filename, page_start, page_end = _parse_page_no(raw_row.cells[1].text)
            warning = ""
            if not _is_extractable_page_range(record_filename, page_start, page_end):
                warning = f"Could not parse record/pages from PAGE NO: {cells[1][:80]}"
            order = len(rows)
            rows.append(
                SelectableChronologyRow(
                    id=_stable_id("row", order, "|".join(cells[:4])),
                    order=order,
                    date=cells[0],
                    page_no=raw_row.cells[1].text.strip(),
                    provider=cells[2],
                    description=cells[3],
                    flags=cells[4],
                    record_filename=record_filename,
                    page_start=page_start,
                    page_end=page_end,
                    warning=warning,
                )
            )
        return rows
    return []


def _iter_body_blocks(doc) -> Iterator[Paragraph | Table]:
    for child in doc.element.body.iterchildren():
        if child.tag.endswith("}p"):
            yield Paragraph(child, doc)
        elif child.tag.endswith("}tbl"):
            yield Table(child, doc)


def _is_chronology_table(table: Table) -> bool:
    if not table.rows or len(table.rows[0].cells) != 5:
        return False
    headers = [_normalize_header(cell.text) for cell in table.rows[0].cells]
    return all(
        headers[index] in aliases
        for index, aliases in enumerate(_CHRON_HEADER_ALIASES)
    )


def _is_extractable_page_range(
    record_filename: str,
    page_start: int,
    page_end: int,
) -> bool:
    return bool(record_filename and page_start > 0 and page_end >= page_start)


def _collapse(text: str) -> str:
    return re.sub(r"\s+", " ", text or "").strip()


def _normalize_header(text: str) -> str:
    return re.sub(r"[^a-z0-9]+", "", _collapse(text).lower())


def _provider_key(text: str) -> str:
    return re.sub(r"[^a-z0-9]+", " ", _collapse(text).lower()).strip()


def _stable_id(prefix: str, order: int, text: str) -> str:
    digest = hashlib.sha1(text.encode("utf-8", errors="ignore")).hexdigest()[:12]
    return f"{prefix}-{order}-{digest}"
