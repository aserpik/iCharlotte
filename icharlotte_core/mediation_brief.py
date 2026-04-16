"""
MediationBriefGenerator — style extraction from sample briefs and caching.

Reads sample mediation brief PDFs to extract per-section style excerpts.
These excerpts are used as few-shot examples when generating new briefs.
Cache is hash-validated so it auto-refreshes when samples change.
"""

import hashlib
import json
import logging
import os
import re
import shutil
from typing import Dict, List, Optional, Tuple

from docx.shared import Inches, Pt
from docx.enum.text import WD_BREAK

import fitz  # PyMuPDF
from docx import Document as DocxDocument
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from icharlotte_core.llm_config import LLMCaller
from icharlotte_core.prompt_manager import get_prompt_manager

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Module-level constants
# ---------------------------------------------------------------------------

SAMPLE_BRIEFS_DIR = r"C:\AI\Mediation Briefs"

# Canonical display order (matches roman numeral order in the document)
SECTION_ORDER = [
    "introduction",
    "statement_of_facts",
    "procedural_status",
    "liability",
    "damages",
    "settlement_position",
    "conclusion",
]

# Generation order: introduction last (written after the other sections are done)
GENERATION_ORDER = [
    "statement_of_facts",
    "procedural_status",
    "liability",
    "damages",
    "settlement_position",
    "conclusion",
    "introduction",
]

# Maps canonical section name → (roman_numeral_str, HEADING_TITLE)
SECTION_HEADINGS: Dict[str, tuple] = {
    "introduction":       ("I",   "INTRODUCTION"),
    "statement_of_facts": ("II",  "STATEMENT OF FACTS"),
    "procedural_status":  ("III", "PROCEDURAL STATUS"),
    "liability":          ("IV",  "LIABILITY"),
    "damages":            ("V",   "DAMAGES"),
    "settlement_position":("VI",  "SETTLEMENT POSITION"),
    "conclusion":         ("VII", "CONCLUSION"),
}

# Regex matching roman-numeral section headings as they appear in sample briefs.
# Examples:
#   "I.     INTRODUCTION"
#   "II.  STATEMENT OF FACTS"
#   "IV.   FACTUAL BACKGROUND"
_HEADING_PATTERN = re.compile(
    r"^\s*(I{1,3}|IV|V?I{0,3}|VI{1,2}|VII)\s*\.\s+"
    r"([A-Z][A-Z \t]+[A-Z])\s*$",
    re.MULTILINE,
)

# Maps variant heading titles found in real briefs → canonical section names
_HEADING_TO_SECTION: Dict[str, str] = {
    # introduction variants
    "INTRODUCTION":              "introduction",
    "OVERVIEW":                  "introduction",
    # statement of facts variants
    "STATEMENT OF FACTS":        "statement_of_facts",
    "FACTUAL BACKGROUND":        "statement_of_facts",
    "FACTS":                     "statement_of_facts",
    "BACKGROUND":                "statement_of_facts",
    # procedural status variants
    "PROCEDURAL STATUS":         "procedural_status",
    "PROCEDURAL BACKGROUND":     "procedural_status",
    "PROCEDURAL HISTORY":        "procedural_status",
    "STATUS OF LITIGATION":      "procedural_status",
    # liability variants
    "LIABILITY":                 "liability",
    "LIABILITY ANALYSIS":        "liability",
    "ANALYSIS OF LIABILITY":     "liability",
    # damages variants
    "DAMAGES":                   "damages",
    "DAMAGES ANALYSIS":          "damages",
    "ANALYSIS OF DAMAGES":       "damages",
    "INJURIES AND DAMAGES":      "damages",
    "SPECIAL DAMAGES":           "damages",
    # settlement position variants
    "SETTLEMENT POSITION":       "settlement_position",
    "SETTLEMENT VALUE":          "settlement_position",
    "VALUATION":                 "settlement_position",
    "DEMAND":                    "settlement_position",
    # conclusion variants
    "CONCLUSION":                "conclusion",
    "SUMMARY":                   "conclusion",
}

# Paths for prompt files and style cache
PROMPTS_DIR = os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
    "Scripts", "prompts", "mediation_brief",
)
CACHE_PATH = os.path.join(PROMPTS_DIR, "style_cache.json")


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _make_bold_run(text: str, underline: bool = False):
    """Build an OxmlElement w:r with bold (and optionally underline)."""
    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rPr.append(OxmlElement("w:b"))
    if underline:
        u_elem = OxmlElement("w:u")
        u_elem.set(qn("w:val"), "single")
        rPr.append(u_elem)
    run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    run.append(t)
    return run


# ---------------------------------------------------------------------------
# MediationBriefGenerator
# ---------------------------------------------------------------------------

class MediationBriefGenerator:
    """Generates mediation briefs by extracting style from samples and calling LLM."""

    # ------------------------------------------------------------------
    # Style and formatting constants
    # ------------------------------------------------------------------

    # Sections whose content affects the Introduction — regenerate it when any of these change.
    _INTRO_TRIGGERS = {"liability", "damages", "statement_of_facts", "conclusion"}

    STYLE_GUIDE = """
STYLE AND TONE GUIDE:
- Write as a senior defense litigation attorney addressing a mediator
- Tone: professional, authoritative, and persuasive — but NOT overly argumentative or inflammatory. Avoid hyperbolic language like "exorbitant payout", "highly questionable", "transparent attempt". Let the facts speak for themselves. The goal is to educate and persuade, not to attack.
- Use active voice and strong declarative statements
- Present defense arguments as the clear and logical reading of the evidence
- Point out weaknesses in plaintiff's case through factual analysis, not characterization
- Use specific facts, dates, and evidence — avoid vague generalities
- Always use party labels before names: "Plaintiff [Name]", "Defendant [Name]", or "Co-Defendant [Name]" when referring to parties to the case
- Define ALL party shorthand names and location shorthand on FIRST reference, as early as possible in the brief. For example: "Defendant Kwang Hae Chong D/B/A Pacific Painting ("Pacific Painting")", "the apartment complex located at 14105 Califa Street in Van Nuys, California (the "Premises")". After the definition, use ONLY the shorthand for all subsequent references. IMPORTANT: Do NOT re-define a term that was already defined in a prior section. Each definition should appear exactly ONCE in the entire brief — in whichever section first mentions that party or location
- Do not use placeholder text like [TBD] or [INSERT] — write around missing information naturally
- Be thorough and detailed — length is not a concern
"""

    FORMATTING_RULES = """
FORMATTING RULES:
- Do NOT include level-one section headings (roman numerals) — they are added by the system
- For the LIABILITY and DAMAGES sections: mark each subsection with "SUBSECTION: Title Text" on its own line, followed by the content paragraphs. The system will convert these to properly formatted level-two headings.
- For deposition quotes: format them as Q&A exchanges. Start each quote block on a new line with "DEPO_QUOTE_START" on its own line, then include the verbatim Q&A with "Q." preceding questions and "A." preceding answers, then end with "DEPO_QUOTE_END" on its own line, followed by the citation on the next line in this format: (LastName Depo Trns., at p. PageNum:LineNum.)
  Example:
  DEPO_QUOTE_START
  Q. Were you paying attention to where you were walking?
  A. No, I was looking at my phone at the time.
  DEPO_QUOTE_END
  (Smith Depo Trns., at p. 45:12.)
- Clean deposition quotes of transcript artifacts (line numbers, extra characters) but keep the testimony verbatim. Do NOT remove dashes (--) from testimony — keep them as-is
- CRITICAL — deposition quote source restriction: Deposition quotes may ONLY be drawn from actual deposition transcript documents — i.e., files containing verbatim court-reported Q&A with page and line numbers produced by a court reporter. You MUST NOT create, reconstruct, paraphrase, or extract deposition quotes from deposition summaries, deposition digests, deposition outlines, case summaries, medical records, expert reports, letters, emails, pleadings, or any other secondary or synthesized document — even if such a document describes what a witness said or contains page:line references. If no verbatim deposition transcript for a given witness is present in the source documents, do NOT include any deposition quote for that witness; discuss the witness's testimony in narrative form only, without the DEPO_QUOTE_START/END block. A document is a transcript only if it contains the raw, uninterrupted Q&A flow of a court-reported deposition; if in doubt, treat it as NOT a transcript and omit the quote.
- Write in plain text. Do not use markdown formatting (no **, ##, etc.)
"""

    # Compile-time regexes for section text parsing
    _SUBSECTION_RE = re.compile(r'^SUBSECTION:\s*(.+)$', re.MULTILINE)
    _DEPO_CITE_RE = re.compile(r'\([A-Z][A-Za-z\'\- ]+ Depo Trns\., at p\. \d+:\d+[\-\d:]*\.\)')

    def __init__(self):
        self._sample_dir: str = SAMPLE_BRIEFS_DIR
        self._cache_path: str = CACHE_PATH
        self._style_cache: Optional[Dict] = None

        # Per-section generated content (populated during generation)
        self.sections: Dict[str, str] = {}
        self.planning_output: str = ""
        self.document_content: str = ""
        self.caption_template_path: Optional[str] = None
        self.is_active: bool = False
        self.saved_path: Optional[str] = None  # Last saved file path for overwrite

    # ------------------------------------------------------------------
    # Section parsing
    # ------------------------------------------------------------------

    def _extract_sections_from_text(self, text: str) -> Dict[str, str]:
        """Parse a sample brief's text into a dict of section_name → body_text.

        Identifies section boundaries by roman-numeral heading lines, then maps
        each heading to a canonical section name via ``_HEADING_TO_SECTION``.
        Sections whose headings are unrecognised are silently skipped.
        """
        matches = list(_HEADING_PATTERN.finditer(text))
        if not matches:
            return {}

        sections: Dict[str, str] = {}

        for i, match in enumerate(matches):
            heading_title = match.group(2).strip()
            canonical = _HEADING_TO_SECTION.get(heading_title)
            if canonical is None:
                # Try partial match (heading may have extra words)
                for variant, name in _HEADING_TO_SECTION.items():
                    if heading_title.startswith(variant) or variant in heading_title:
                        canonical = name
                        break

            if canonical is None:
                logger.debug("Unrecognised heading: %r — skipping", heading_title)
                continue

            # Body text runs from end of this heading to start of next heading
            body_start = match.end()
            body_end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
            body = text[body_start:body_end].strip()

            if body:
                sections[canonical] = body

        return sections

    # ------------------------------------------------------------------
    # Hashing / caching
    # ------------------------------------------------------------------

    def _hash_file(self, path: str) -> str:
        """Return the MD5 hex digest of the file at *path*."""
        h = hashlib.md5()
        with open(path, "rb") as fh:
            for chunk in iter(lambda: fh.read(65536), b""):
                h.update(chunk)
        return h.hexdigest()

    def _read_sample_pdfs(self) -> Dict:
        """Read all PDFs in the sample directory, extract sections from each.

        Returns a dict with keys:
          - ``"hashes"``:   {filename: md5_hex}
          - ``"sections"``: {section_name: [excerpt1, excerpt2]}  (best 2 per section)
        """
        if not os.path.isdir(self._sample_dir):
            logger.warning("Sample briefs directory not found: %s", self._sample_dir)
            return {"hashes": {}, "sections": {}}

        pdf_files = sorted(
            f for f in os.listdir(self._sample_dir)
            if f.lower().endswith(".pdf")
        )

        if not pdf_files:
            logger.warning("No PDF files found in %s", self._sample_dir)
            return {"hashes": {}, "sections": {}}

        hashes: Dict[str, str] = {}
        # section_name → list of (text, source_file) sorted by length descending
        raw: Dict[str, List[tuple]] = {s: [] for s in SECTION_ORDER}

        for filename in pdf_files:
            full_path = os.path.join(self._sample_dir, filename)
            try:
                hashes[filename] = self._hash_file(full_path)
                doc = fitz.open(full_path)
                try:
                    text = "\n".join(page.get_text() for page in doc)
                finally:
                    doc.close()

                found = self._extract_sections_from_text(text)
                for section, body in found.items():
                    raw[section].append((body, filename))

            except Exception as exc:
                logger.error("Failed to process sample PDF %s: %s", filename, exc)

        # Keep the 2 longest excerpts per section
        sections: Dict[str, List[str]] = {}
        for section, entries in raw.items():
            if not entries:
                continue
            entries.sort(key=lambda x: len(x[0]), reverse=True)
            sections[section] = [text for text, _ in entries[:2]]

        return {"hashes": hashes, "sections": sections}

    def _save_style_cache(self, data: Dict) -> None:
        """Persist style data to the JSON cache file.

        The on-disk format uses ``source_hashes`` (not ``hashes``) so callers
        can distinguish the raw extraction payload from the persisted format.
        """
        cache_dir = os.path.dirname(self._cache_path)
        os.makedirs(cache_dir, exist_ok=True)

        payload = {
            "source_hashes": data.get("hashes", {}),
            "sections": data.get("sections", {}),
        }
        with open(self._cache_path, "w", encoding="utf-8") as fh:
            json.dump(payload, fh, indent=2, ensure_ascii=False)
        logger.debug("Style cache saved to %s", self._cache_path)

    def _load_style_cache(self) -> Optional[Dict]:
        """Load and validate the style cache from disk.

        Validates that all PDF files recorded in the cache still exist and their
        MD5 hashes match.  Returns ``None`` if the cache is absent, corrupt, or
        stale.
        """
        if not os.path.isfile(self._cache_path):
            return None

        try:
            with open(self._cache_path, "r", encoding="utf-8") as fh:
                cached = json.load(fh)
        except (json.JSONDecodeError, OSError) as exc:
            logger.warning("Could not read style cache: %s", exc)
            return None

        source_hashes = cached.get("source_hashes", {})
        if not source_hashes:
            # Cache with no file references is considered stale
            return None

        # Validate each hash
        for filename, expected_hash in source_hashes.items():
            full_path = os.path.join(self._sample_dir, filename)
            if not os.path.isfile(full_path):
                logger.info("Cache stale: %s no longer exists", filename)
                return None
            try:
                actual_hash = self._hash_file(full_path)
            except OSError as exc:
                logger.warning("Cannot hash %s: %s", filename, exc)
                return None
            if actual_hash != expected_hash:
                logger.info("Cache stale: %s has changed", filename)
                return None

        return cached

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    def get_style_excerpts(self) -> Dict[str, List[str]]:
        """Return style excerpts keyed by section name.

        On first call the sample PDFs are parsed and the result is cached.
        Subsequent calls return the in-memory cache.  The on-disk JSON cache is
        used across sessions; it is invalidated whenever any source PDF changes.
        """
        if self._style_cache is not None:
            return self._style_cache.get("sections", {})

        # Try loading from disk first
        cached = self._load_style_cache()
        if cached is not None:
            logger.debug("Using style cache from disk")
            self._style_cache = cached
            return cached.get("sections", {})

        # Extract fresh from sample PDFs
        logger.info("Extracting style from sample briefs in %s", self._sample_dir)
        data = self._read_sample_pdfs()
        self._save_style_cache(data)

        # Normalise to the same format as the on-disk cache
        self._style_cache = {
            "source_hashes": data.get("hashes", {}),
            "sections": data.get("sections", {}),
        }
        return self._style_cache.get("sections", {})

    # ------------------------------------------------------------------
    # LLM prompt construction and generation
    # ------------------------------------------------------------------

    def _get_section_prompt(self, section_name: str) -> str:
        """Load the main prompt for *section_name* from the Workbench (PromptManager).

        Falls back to reading the file directly from PROMPTS_DIR if the
        PromptManager returns nothing.  Returns an empty string if neither
        source has a prompt for this section.
        """
        pm = get_prompt_manager()
        prompt = pm.get_prompt("mediation_brief", section_name)
        if prompt:
            return prompt

        # Direct file fallback
        fallback_path = os.path.join(PROMPTS_DIR, f"{section_name}_current.txt")
        if os.path.isfile(fallback_path):
            with open(fallback_path, "r", encoding="utf-8") as fh:
                return fh.read()

        logger.warning("No prompt found for section %r", section_name)
        return ""

    def _build_system_prompt(self, section_name: str) -> str:
        """Return the system prompt for the given section.

        Loads style guide and formatting rules from the Workbench
        (PromptManager) if edited, otherwise uses class constants.
        """
        persona = (
            "You are a senior defense litigation attorney with decades of trial experience. "
            "You are writing a confidential mediation brief on behalf of the defendant. "
            "Your goal is to present the strongest possible defense position to the mediator, "
            "supported by the evidence and the facts of the case."
        )
        try:
            from icharlotte_core.prompt_manager import get_prompt
            style = get_prompt("mediation_brief", "style_guide") or self.STYLE_GUIDE
            formatting = get_prompt("mediation_brief", "formatting_rules") or self.FORMATTING_RULES
        except Exception:
            style = self.STYLE_GUIDE
            formatting = self.FORMATTING_RULES
        return f"{persona}\n{style}\n{formatting}"

    def _build_section_prompt(self, section_name: str, refinement_instruction: str = "") -> str:
        """Build the full prompt for *section_name*.

        Assembles:
        1. Main prompt loaded from the Workbench.
        2. Optional refinement instruction.
        3. Style excerpt from cached samples (first excerpt, truncated at 10 000 chars).
        4. Planning pass output.
        5. Previously generated sections — ALL sections when writing the
           introduction; only sections that precede *section_name* in
           GENERATION_ORDER otherwise.
        """
        parts: List[str] = []

        # 1. Main section prompt
        main_prompt = self._get_section_prompt(section_name)
        if main_prompt:
            parts.append(main_prompt)

        # 2. Refinement instruction
        if refinement_instruction:
            parts.append(f"\nREFINEMENT INSTRUCTION:\n{refinement_instruction}")

        # 3. Style excerpt
        excerpts = self.get_style_excerpts()
        section_excerpts = excerpts.get(section_name, [])
        if section_excerpts:
            excerpt = section_excerpts[0][:10000]
            parts.append(
                f"\nSTYLE REFERENCE — example from a prior mediation brief "
                f"(use only as a stylistic guide):\n{excerpt}"
            )

        # 4. Planning pass output
        if self.planning_output:
            parts.append(f"\nPLANNING ANALYSIS:\n{self.planning_output}")

        # 5. Previously generated sections
        if section_name == "introduction":
            # Introduction is written last — include every other section
            prior_sections = {
                s: self.sections[s]
                for s in SECTION_ORDER
                if s != "introduction" and s in self.sections
            }
        else:
            # Include all sections that were generated before this one
            current_index = GENERATION_ORDER.index(section_name) if section_name in GENERATION_ORDER else -1
            prior_names = GENERATION_ORDER[:current_index] if current_index > 0 else []
            prior_sections = {s: self.sections[s] for s in prior_names if s in self.sections}

        if prior_sections:
            prior_text = "\n\n".join(
                f"[{s.upper().replace('_', ' ')}]\n{body}"
                for s, body in prior_sections.items()
            )
            parts.append(f"\nPREVIOUSLY DRAFTED SECTIONS (for context and consistency):\n{prior_text}")
            parts.append(
                "\nIMPORTANT: The sections above have already defined party names, "
                "locations, and other terms with parenthetical shorthand. Do NOT "
                "re-define any of these terms. Use only the shorthand names "
                "(e.g., \"Plaintiff\", \"Pacific Painting\", \"the Premises\") "
                "without repeating the full name and parenthetical definition."
            )

        return "\n".join(parts)

    def generate_section(self, section_name: str, refinement_instruction: str = "") -> Optional[str]:
        """Generate content for *section_name* using the LLM.

        Calls LLMCaller with the section-specific system prompt and the
        assembled section prompt, using the document content as the text
        body.  Stores the result in ``self.sections[section_name]`` and
        returns it.  Returns None if the LLM call fails.
        """
        system_prompt = self._build_system_prompt(section_name)
        section_prompt = self._build_section_prompt(section_name, refinement_instruction)

        caller = LLMCaller()
        # Pass system prompt prepended to the section prompt so the caller
        # receives both in the user-visible prompt parameter.
        full_prompt = f"{system_prompt}\n\n{section_prompt}"

        result = caller.call(
            prompt=full_prompt,
            text=self.document_content,
            agent_id="agent_mediation_brief",
        )

        if result:
            self.sections[section_name] = result
            logger.info("Generated section %r (%d chars)", section_name, len(result))
        else:
            logger.error("LLM returned no content for section %r", section_name)

        return result

    def run_planning_pass(self) -> Optional[str]:
        """Run the planning pass over the document content.

        Loads the planning prompt, calls the LLM, stores the output in
        ``self.planning_output``, and returns it.  Returns None if the
        LLM call fails.
        """
        planning_prompt = self._get_section_prompt("planning")
        if not planning_prompt:
            logger.warning("Planning prompt not found — skipping planning pass")
            return None

        caller = LLMCaller()
        result = caller.call(
            prompt=planning_prompt,
            text=self.document_content,
            agent_id="agent_mediation_brief",
        )

        if result:
            self.planning_output = result
            logger.info("Planning pass complete (%d chars)", len(result))
        else:
            logger.error("Planning pass: LLM returned no content")

        return result

    # ------------------------------------------------------------------
    # Caption template handling
    # ------------------------------------------------------------------

    # XML namespace constants for Word document XML
    _W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    _W_T = f"{{{_W_NS}}}t"
    _W_P = f"{{{_W_NS}}}p"
    _W_R = f"{{{_W_NS}}}r"

    # Signature block indicator phrases (case-insensitive search)
    _SIGNATURE_INDICATORS = [
        "By:",
        "DATED:",
        "Respectfully submitted",
        "State Bar No.",
        "Attorney for",
        "Counsel for",
    ]

    def find_caption_template(self, folder: str) -> Optional[str]:
        """Search *folder* for a .docx file with "caption" in its name.

        Prefers matches in the root of *folder*; falls back to a recursive
        search of subfolders. Returns the full path to the first match
        (case-insensitive), or None if no matching file is found.
        """
        if not os.path.isdir(folder):
            return None
        try:
            for entry in os.listdir(folder):
                if entry.lower().endswith(".docx") and "caption" in entry.lower():
                    return os.path.join(folder, entry)
        except OSError:
            pass
        for dirpath, _, filenames in os.walk(folder):
            if os.path.abspath(dirpath) == os.path.abspath(folder):
                continue
            for fname in filenames:
                if fname.lower().endswith(".docx") and "caption" in fname.lower():
                    return os.path.join(dirpath, fname)
        return None

    def prepare_caption_template(self, caption_path: str, output_path: str) -> list:
        """Copy *caption_path* to *output_path*, process it, and return extracted
        signature paragraph elements.

        Processing steps:
        1. Copy original file to output_path (never modify the original).
        2. Replace "CAPTION PAGE" text in body and footers with the styled
           "DEFENDANT'S CONFIDENTIAL MEDIATION BRIEF" title.
        3. Extract (and remove) any trailing signature block paragraphs.
        4. Save the modified document.

        Returns a list of extracted signature paragraph elements (may be empty).
        """
        shutil.copy2(caption_path, output_path)
        doc = DocxDocument(output_path)

        self._replace_caption_page_body(doc)
        self._replace_caption_page_footers(doc)
        sig_elements = self._extract_signature_block(doc)

        doc.save(output_path)
        return sig_elements

    def _replace_caption_page_body(self, doc) -> None:
        """Search ALL w:t elements in the document body XML for "CAPTION PAGE"
        (including text inside nested tables) and replace the containing paragraph
        with three styled runs:
          - "DEFENDANT'S " (bold)
          - "CONFIDENTIAL" (bold + underline)
          - " MEDIATION BRIEF" (bold)

        Uses raw XML iteration because caption templates typically use nested
        table layouts where doc.paragraphs won't find the text.
        """
        W_T = self._W_T
        W_P = self._W_P
        W_R = self._W_R

        for t_elem in doc.element.body.iter(W_T):
            if t_elem.text and "CAPTION PAGE" in t_elem.text.upper():
                # Walk up to the enclosing w:p
                para_elem = t_elem.getparent()
                while para_elem is not None and para_elem.tag != W_P:
                    para_elem = para_elem.getparent()
                if para_elem is None:
                    continue

                # Remove all existing runs from this paragraph
                for run_elem in list(para_elem.iter(W_R)):
                    parent = run_elem.getparent()
                    if parent is not None:
                        parent.remove(run_elem)

                para_elem.append(_make_bold_run("DEFENDANT'S "))
                para_elem.append(_make_bold_run("CONFIDENTIAL", underline=True))
                para_elem.append(_make_bold_run(" MEDIATION BRIEF"))
                return  # Only replace the first occurrence

    def _replace_caption_page_footers(self, doc) -> None:
        """Replace the pleading title in ALL footers with the mediation brief title.

        Footer text is often inside tables within the footer, so we search
        ALL w:t elements in the footer XML (not just paragraphs). Any
        non-numeric text that isn't a page number gets its parent paragraph
        replaced with the brief title.
        """
        W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        W_T = f"{{{W_NS}}}t"
        W_R = f"{{{W_NS}}}r"
        W_P = f"{{{W_NS}}}p"

        _SKIP_TEXTS = {"", "page", "of"}  # common page number labels

        for section in doc.sections:
            for footer in (
                section.footer,
                section.even_page_footer,
                section.first_page_footer,
            ):
                if footer is None:
                    continue
                # Track which paragraphs we've already replaced
                replaced_paras = set()
                for t_elem in list(footer._element.iter(W_T)):
                    if t_elem.text is None:
                        continue
                    text = t_elem.text.strip()
                    # Skip page numbers and empty text
                    if not text or text.isdigit() or text.lower() in _SKIP_TEXTS:
                        continue
                    # Find the parent paragraph
                    para_elem = t_elem.getparent()
                    while para_elem is not None and para_elem.tag != W_P:
                        para_elem = para_elem.getparent()
                    if para_elem is None or id(para_elem) in replaced_paras:
                        continue
                    replaced_paras.add(id(para_elem))
                    # Clear all runs in this paragraph
                    for run_elem in list(para_elem.iter(W_R)):
                        para_elem.remove(run_elem)
                    # Add the brief title
                    para_elem.append(_make_bold_run("DEFENDANT'S "))
                    para_elem.append(_make_bold_run("CONFIDENTIAL", underline=True))
                    para_elem.append(_make_bold_run(" MEDIATION BRIEF"))

    def _extract_signature_block(self, doc) -> list:
        """Scan backwards through the last 15 paragraphs looking for signature
        block indicators.  When found, extract all paragraphs from that point
        to the end of the document, remove them from the body, and return the
        extracted paragraph objects.

        Returns a list of Paragraph objects (may be empty if no signature block
        is detected).  The elements are removed from the document body so they
        can be re-inserted elsewhere in the assembled brief.
        """
        indicators = self._SIGNATURE_INDICATORS
        all_paras = doc.paragraphs

        # Search the last 15 paragraphs for a signature indicator
        search_paras = all_paras[-15:] if len(all_paras) > 15 else all_paras
        sig_start_index = None

        for i, para in enumerate(search_paras):
            text = para.text
            for indicator in indicators:
                if indicator.lower() in text.lower():
                    # Map back to index in the full paragraph list
                    offset = len(all_paras) - len(search_paras)
                    sig_start_index = offset + i
                    break
            if sig_start_index is not None:
                break

        if sig_start_index is None:
            return []

        # Collect paragraphs from sig_start_index to end
        sig_paras = list(all_paras[sig_start_index:])

        # Remove them from the document body in reverse order to preserve indices
        for para in reversed(sig_paras):
            para_elem = para._element
            parent = para_elem.getparent()
            if parent is not None:
                parent.remove(para_elem)

        return sig_paras

    # ------------------------------------------------------------------
    # Section text parsing
    # ------------------------------------------------------------------

    def _parse_section_text(self, text: str, section_name: str) -> List[Dict]:
        """Parse LLM output into structured elements.

        Handles three element types:
          - ``"l2_heading"``  — line matching ``^SUBSECTION: <title>``
          - ``"depo_quote"``  — block between DEPO_QUOTE_START/END markers,
            or (legacy) a paragraph containing a deposition citation
          - ``"body"``        — everything else

        Returns a list of dicts with keys ``"type"`` and ``"text"``.
        """
        elements: List[Dict] = []
        lines = text.split("\n")

        i = 0
        current_body = []

        def _flush_body():
            """Emit accumulated body lines as body paragraph(s)."""
            if current_body:
                body_text = "\n".join(current_body).strip()
                if body_text:
                    # Split on double-newlines into separate paragraphs
                    for para in body_text.split("\n\n"):
                        para = para.strip()
                        if para:
                            elements.append({"type": "body", "text": para})
                current_body.clear()

        while i < len(lines):
            line = lines[i]
            stripped = line.strip()

            # SUBSECTION heading
            sub_match = self._SUBSECTION_RE.match(stripped)
            if sub_match:
                _flush_body()
                elements.append({"type": "l2_heading", "text": sub_match.group(1).strip()})
                i += 1
                continue

            # DEPO_QUOTE_START block
            if stripped == "DEPO_QUOTE_START":
                _flush_body()
                i += 1
                quote_lines = []
                while i < len(lines) and lines[i].strip() != "DEPO_QUOTE_END":
                    quote_lines.append(lines[i].strip())
                    i += 1
                if i < len(lines):
                    i += 1  # skip DEPO_QUOTE_END
                # Next non-empty line should be the citation
                citation = ""
                while i < len(lines) and not lines[i].strip():
                    i += 1
                if i < len(lines) and self._DEPO_CITE_RE.search(lines[i]):
                    citation = lines[i].strip()
                    i += 1
                quote_text = "\n".join(quote_lines)
                if citation:
                    quote_text += "\n" + citation
                elements.append({"type": "depo_quote", "text": quote_text})
                continue

            # Legacy: paragraph with depo citation (no DEPO_QUOTE_START/END)
            if self._DEPO_CITE_RE.search(stripped):
                _flush_body()
                elements.append({"type": "depo_quote", "text": stripped})
                i += 1
                continue

            # Regular body line
            current_body.append(line)
            i += 1

        _flush_body()
        return elements

    # ------------------------------------------------------------------
    # Word document formatting helpers
    # ------------------------------------------------------------------

    def _add_tab_stop(self, para, position_inches: float = 0.5) -> None:
        """Add a left-aligned tab stop at *position_inches* to *para*."""
        pPr = para._element.get_or_add_pPr()
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "left")
        tab.set(qn("w:pos"), str(int(position_inches * 1440)))  # twips
        tabs.append(tab)
        pPr.append(tabs)

    def _add_l1_heading(self, doc, roman: str, title: str):
        """Add a level-1 section heading formatted as ``I.     INTRODUCTION``.

        Uses a hanging indent (left=0.5", hanging=0.5") with a tab stop at 0.5"
        so the title text aligns neatly after the roman numeral.
        """
        para = doc.add_paragraph()
        pf = para.paragraph_format
        pf.left_indent = Inches(0.5)
        pf.first_line_indent = Inches(-0.5)
        pf.line_spacing = 2.0  # Double spacing
        pf.space_before = Pt(12)
        pf.space_after = Pt(0)
        self._add_tab_stop(para, position_inches=0.5)

        # Run 1: roman numeral — bold only
        r1 = para.add_run(f"{roman}.")
        r1.bold = True

        # Tab character
        para.add_run("\t")

        # Run 2: title — bold + underline
        r2 = para.add_run(title)
        r2.bold = True
        r2.underline = True

        return para

    def _add_l2_heading(self, doc, letter: str, title: str):
        """Add a level-2 subsection heading formatted as ``A.     Title Text``.

        Letter at 0.5" indent, title at 1.0" indent (0.5" after the letter).
        Uses left_indent=1.0" with hanging indent back to 0.5" for the letter,
        and a tab stop at 1.0" for the title.
        """
        para = doc.add_paragraph()
        pf = para.paragraph_format
        pf.left_indent = Inches(1.0)
        pf.first_line_indent = Inches(-0.5)
        pf.line_spacing = 2.0  # Double spacing
        pf.space_before = Pt(10)
        pf.space_after = Pt(0)
        self._add_tab_stop(para, position_inches=1.0)

        # Run 1: letter — bold only
        r1 = para.add_run(f"{letter}.")
        r1.bold = True

        # Tab character
        para.add_run("\t")

        # Run 2: title — bold + underline
        r2 = para.add_run(title)
        r2.bold = True
        r2.underline = True

        return para

    def _add_body_paragraph(self, doc, text: str):
        """Add a body paragraph with double spacing and 0.5 inch first-line indent.

        Deposition citation lines (parenthetical references) get no indent.
        """
        para = doc.add_paragraph(text)
        pf = para.paragraph_format
        pf.line_spacing = 2.0  # Double spacing
        pf.space_after = Pt(0)
        # Citation lines should not be indented
        if self._DEPO_CITE_RE.search(text):
            pf.first_line_indent = Inches(0)
        else:
            pf.first_line_indent = Inches(0.5)
        return para

    def _add_depo_quote(self, doc, text: str):
        """Add deposition quote lines — single-spaced, Q/A with hanging indent.

        Q. and A. lines get a 1.0" left indent with 0.5" hanging indent so the
        letter is at 0.5" and the testimony text starts at 1.0". The citation
        line (parenthetical) is NOT indented — it sits at the left margin.
        """
        lines = text.split("\n")
        first_para = None
        for line in lines:
            line = line.strip()
            if not line:
                continue

            is_citation = self._DEPO_CITE_RE.search(line)
            is_qa = line.startswith("Q.") or line.startswith("A.")

            para = doc.add_paragraph()
            pf = para.paragraph_format
            pf.line_spacing = 1.0  # Single spacing
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)

            if is_citation:
                # Citation line — no indent, with space before to create
                # a visible gap between the last Q/A line and the citation
                pf.line_spacing = 2.0
                pf.space_before = Pt(12)
                para.add_run(line)
            elif is_qa:
                # Q./A. line — letter at 0.5", text at 1.0" via hanging indent + tab
                pf.left_indent = Inches(1.0)
                pf.first_line_indent = Inches(-0.5)
                self._add_tab_stop(para, position_inches=1.0)
                # Split into letter and testimony
                letter = line[:2]  # "Q." or "A."
                testimony = line[2:].strip()
                para.add_run(letter)
                para.add_run("\t")
                para.add_run(testimony)
            else:
                # Continuation line — indented at 0.5"
                pf.left_indent = Inches(0.5)
                para.add_run(line)

            if first_para is None:
                pf.space_before = Pt(6)
                first_para = para

        # Add spacing after the last line
        if first_para:
            para.paragraph_format.space_after = Pt(6)
        return first_para

    # ------------------------------------------------------------------
    # Document assembly
    # ------------------------------------------------------------------

    def assemble_document(
        self,
        caption_path: str,
        output_path: str,
        signature_paragraphs: Optional[list] = None,
    ) -> None:
        """Assemble the full mediation brief Word document.

        Steps:
        1. Call ``prepare_caption_template()`` to copy and process the caption.
        2. Open the processed document.
        3. Add a page break after the caption content.
        4. For each section in SECTION_ORDER, add L1 heading and body elements.
        5. Append the signature block if present.
        6. Save and validate.

        Args:
            caption_path: Path to the source caption .docx file.
            output_path: Destination path for the assembled brief.
            signature_paragraphs: Pre-extracted signature paragraph elements.
                If None, they are extracted from the caption template automatically.
        """
        # Step 1: Prepare caption (copies to output_path, extracts signature)
        extracted_sig = self.prepare_caption_template(caption_path, output_path)
        if signature_paragraphs is None:
            signature_paragraphs = extracted_sig

        # Step 2: Open the processed document
        doc = DocxDocument(output_path)

        # Step 3: Page break after caption
        page_break_para = doc.add_paragraph()
        run = page_break_para.add_run()
        run.add_break(WD_BREAK.PAGE)

        # Step 4: Add each section
        letter_seq = "ABCDEFGHIJKLMNOPQRSTUVWXYZ"

        for section_name in SECTION_ORDER:
            roman, heading_title = SECTION_HEADINGS[section_name]
            self._add_l1_heading(doc, roman, heading_title)

            section_text = self.sections.get(section_name, "")
            if not section_text:
                continue

            elements = self._parse_section_text(section_text, section_name)
            letter_counter = 0  # resets per section

            for elem in elements:
                etype = elem["type"]
                etext = elem["text"]

                if etype == "l2_heading":
                    letter = letter_seq[letter_counter % len(letter_seq)]
                    self._add_l2_heading(doc, letter, etext)
                    letter_counter += 1
                elif etype == "depo_quote":
                    self._add_depo_quote(doc, etext)
                else:
                    self._add_body_paragraph(doc, etext)

        # Step 5: Append signature block
        if signature_paragraphs:
            for sig_para in signature_paragraphs:
                doc.element.body.append(sig_para._element)

        # Step 6: Save
        doc.save(output_path)
        logger.info("Assembled mediation brief saved to %s", output_path)

        # Validate (best-effort)
        try:
            from icharlotte_core.word_validator import validate_report
            result = validate_report(output_path)
            result.print_summary()
        except Exception as e:
            logger.warning("Validation skipped: %s", e)

    # ------------------------------------------------------------------
    # Pipeline orchestration
    # ------------------------------------------------------------------

    def generate_all_sections(self, progress_callback=None) -> None:
        """Run the full generation pipeline: planning pass then all sections.

        Calls ``run_planning_pass()`` first (step 0), then iterates through
        ``GENERATION_ORDER`` generating each section in sequence.

        Args:
            progress_callback: Optional callable ``(section_name, index, total)``
                called before each step (including the planning pass at index 0).
        """
        total = len(GENERATION_ORDER) + 1  # +1 for planning pass

        # Step 0: planning pass
        if progress_callback is not None:
            progress_callback("planning", 0, total)
        self.run_planning_pass()

        # Steps 1..N: generate each section
        for i, section_name in enumerate(GENERATION_ORDER, start=1):
            if progress_callback is not None:
                progress_callback(section_name, i, total)
            result = self.generate_section(section_name)
            if result:
                self.sections[section_name] = result

        self.is_active = True

    # ------------------------------------------------------------------
    # Conversational refinement
    # ------------------------------------------------------------------

    def _parse_routing_response(self, response: str) -> List[str]:
        """Parse the LLM routing response into a list of valid section names.

        Args:
            response: Raw string returned by the routing LLM call.

        Returns:
            A list of canonical section names (those present in SECTION_ORDER).
            Returns an empty list when the response is "none" or contains no
            recognisable section names.
        """
        cleaned = response.strip().lower()
        if cleaned == "none":
            return []
        parts = [p.strip() for p in cleaned.split(",")]
        return [p for p in parts if p in SECTION_ORDER]

    def route_refinement(self, user_message: str) -> List[str]:
        """Ask the LLM which sections of the brief need to be updated.

        Loads the routing prompt, asks the LLM to identify which sections are
        affected by *user_message*, and returns the parsed list.  An empty list
        means the message is not a brief refinement request.

        Args:
            user_message: The user's chat message.

        Returns:
            List of canonical section names to regenerate (may be empty).
        """
        routing_prompt = self._get_section_prompt("routing")
        full_prompt = f"{routing_prompt}\n\nUser message: {user_message}"

        caller = LLMCaller()
        result = caller.call(
            prompt=full_prompt,
            text="",
            agent_id="agent_mediation_brief",
        )

        if not result:
            logger.warning("route_refinement: LLM returned no response")
            return []

        return self._parse_routing_response(result)

    def refine_sections(
        self,
        section_names: List[str],
        instruction: str,
        progress_callback=None,
    ) -> List[str]:
        """Regenerate the specified sections with a refinement instruction.

        When any of the regenerated sections is an _INTRO_TRIGGERS member,
        the Introduction is also regenerated afterwards (without the user's
        instruction — it simply rebuilds itself based on the updated sections).

        Args:
            section_names: Canonical names of the sections to regenerate.
            instruction:   Refinement instruction applied to the requested sections.
            progress_callback: Optional callable ``(section_name,)`` called after
                each section is regenerated.

        Returns:
            List of all section names that were regenerated (including
            "introduction" if it was triggered automatically).
        """
        # Build the regeneration list, adding introduction at the end if needed
        to_regenerate = list(section_names)
        if any(s in self._INTRO_TRIGGERS for s in section_names):
            if "introduction" not in to_regenerate:
                to_regenerate.append("introduction")
            else:
                # Ensure introduction is always last
                to_regenerate.remove("introduction")
                to_regenerate.append("introduction")

        regenerated: List[str] = []
        for section_name in to_regenerate:
            # Apply instruction only to the sections the user asked about;
            # introduction regenerates purely from context, no user instruction.
            if section_name == "introduction" and section_name not in section_names:
                section_instruction = ""
            else:
                section_instruction = instruction

            result = self.generate_section(section_name, refinement_instruction=section_instruction)
            if result:
                self.sections[section_name] = result
                regenerated.append(section_name)

            if progress_callback is not None:
                progress_callback(section_name)

        return regenerated

    def build_refinement_prompts(
        self,
        section_name: str,
        sections_dict: Dict[str, str],
        instruction: str,
    ) -> tuple:
        """Return ``(system_prompt, full_prompt)`` for refining *section_name*.

        Stateless sibling of ``refine_sections`` — used by the Word AI
        assistant popup so refinement can run against a live Word document
        without touching this generator's in-memory state.

        Temporarily swaps ``self.sections`` with *sections_dict* while the
        existing prompt builders run, then restores the original value.

        Args:
            section_name: Canonical section name (e.g. ``"liability"``).
            sections_dict: Mapping of canonical section name → current body
                text, parsed from the live Word document.
            instruction: The user's refinement instruction.

        Returns:
            ``(system_prompt, full_prompt)`` — the two prompt strings ready to
            be submitted through ``TaskLLMWorkerThread``.
        """
        saved_sections = self.sections
        self.sections = dict(sections_dict)
        try:
            system_prompt = self._build_system_prompt(section_name)
            full_prompt = self._build_section_prompt(
                section_name, refinement_instruction=instruction
            )
            return system_prompt, full_prompt
        finally:
            self.sections = saved_sections

    # ------------------------------------------------------------------
    # Quote search and insertion
    # ------------------------------------------------------------------

    # ID-based match format returned by the LLM after reviewing a numbered
    # list of candidate exchanges. The LLM only returns IDs + a relevance
    # sentence; the actual Q/A text and citation are looked up from the
    # parsed TranscriptIndex, so fabrication is impossible.
    _QUOTE_MATCH_RE = re.compile(
        r'MATCH_START\s*\n'
        r'ID:\s*EX_(\d+)\s*\n'
        r'RELEVANCE:\s*(.+?)\n'
        r'MATCH_END',
        re.DOTALL
    )

    # Tokenizer used by the keyword pre-filter. Keeps alphabetic words of
    # length >= 3; stop words are removed to avoid matching on filler.
    _WORD_RE = re.compile(r"[a-zA-Z]+")

    _STOPWORDS = frozenset([
        # Generic English stopwords
        "a", "about", "after", "all", "also", "an", "and", "another", "any",
        "are", "as", "at", "be", "because", "been", "before", "but", "by",
        "can", "could", "did", "do", "does", "down", "during", "each",
        "every", "for", "from", "had", "has", "have", "he", "her", "him",
        "his", "how", "i", "if", "in", "into", "is", "it", "its", "just",
        "may", "me", "might", "my", "no", "not", "now", "of", "on", "one",
        "only", "or", "other", "our", "out", "over", "own", "said", "same",
        "she", "should", "so", "some", "than", "that", "the", "their",
        "them", "then", "there", "these", "they", "this", "those", "to",
        "too", "two", "under", "up", "us", "very", "was", "we", "well",
        "were", "what", "when", "where", "which", "while", "who", "why",
        "will", "with", "would", "you", "your", "yes",
        # Legal / deposition meta-words. These show up constantly in search
        # descriptions ("find testimony that supports the liability
        # argument") but carry no signal about what content to match — in
        # fact they match procedural boilerplate inside the transcript.
        "testimony", "testify", "testified", "depo", "deposition", "depose",
        "brief", "argument", "arguments", "argue", "argued", "support",
        "supports", "supporting", "supported", "case", "matter", "issue",
        "section", "paragraph", "passage", "witness", "witnessed",
        "statement", "statements", "claim", "claims", "claimed",
        "plaintiff", "defendant", "counsel", "attorney", "mediator",
        "find", "finds", "found", "show", "shows", "showed", "prove",
        "proves", "proved", "demonstrate", "demonstrates", "mention",
        "mentions", "mentioned", "discuss", "discusses", "discussed",
        "regarding", "relating", "related", "relates", "concerning",
        "concerns", "concerned",
    ])

    def _tokenize_for_ranking(self, text: str) -> List[str]:
        """Lowercase alphabetic tokens of length >= 3, minus stopwords."""
        tokens = []
        for w in self._WORD_RE.findall(text or ""):
            lw = w.lower()
            if len(lw) >= 3 and lw not in self._STOPWORDS:
                tokens.append(lw)
        return tokens

    def _rank_exchanges_by_keywords(
        self,
        candidates: List[tuple],
        description: str,
        top_n: int = 200,
    ) -> List[tuple]:
        """Score candidate exchanges against the search description.

        Each candidate is a tuple of ``(QAExchange, deponent_last, source)``.
        Score = 10 * (distinct query terms hit) + (total query-term occurrences).
        Coverage breadth dominates; TF is a tie-breaker.

        When the query has content-word signal, return up to *top_n* ranked
        matches. When the query has NO content signal (either empty after
        stopword filtering, or no candidate shares any query term), return
        ALL candidates unchanged — the LLM is then the sole filter. This
        avoids the degenerate case where abstract legal-theory queries
        ("testimony supporting the liability arguments") rank procedural
        boilerplate to the top simply because it appears first in the
        transcript.
        """
        query_terms = set(self._tokenize_for_ranking(description))
        if not query_terms:
            return list(candidates)

        scored = []
        for tup in candidates:
            ex = tup[0]
            doc_tokens = self._tokenize_for_ranking(f"{ex.question} {ex.answer}")
            if not doc_tokens:
                continue
            tf = 0
            distinct_hits = set()
            for t in doc_tokens:
                if t in query_terms:
                    tf += 1
                    distinct_hits.add(t)
            if tf == 0:
                continue
            score = len(distinct_hits) * 10 + tf
            scored.append((score, tup))

        if not scored:
            return list(candidates)

        scored.sort(key=lambda x: x[0], reverse=True)
        return [tup for _, tup in scored[:top_n]]

    def _format_candidates_for_llm(self, candidates: List[tuple]) -> str:
        """Render candidates as a numbered list for the LLM to review.

        Layout is deliberately compact — one exchange per block, with a
        stable ``[EX_NNNN]`` identifier the LLM can echo back.
        """
        lines = []
        for idx, (ex, deponent, source) in enumerate(candidates):
            citation = ex.citation_range()
            lines.append(f"[EX_{idx:04d}] {deponent} Depo Trns., p. {citation}")
            lines.append(f"Q. {ex.question}")
            lines.append(f"A. {ex.answer}")
            lines.append("")
        return "\n".join(lines)

    def _parse_quote_match_results(
        self,
        llm_output: str,
        candidates: List[tuple],
        full_exchanges_by_source: Optional[Dict[Tuple[str, str], Dict[int, object]]] = None,
    ) -> List[Dict]:
        """Parse LLM MATCH_START/END blocks, resolving IDs to real exchanges.

        Any ID not present in ``candidates`` is dropped (grounding check —
        the LLM cannot smuggle in text we didn't supply). Duplicates are
        collapsed.

        When ``full_exchanges_by_source`` is supplied (a mapping of
        ``(deponent, source) -> {exchange_id: QAExchange}`` containing every
        usable exchange from each transcript), the immediately preceding
        exchange is prepended to the result's ``qa_text`` so the quote is
        easier to read in isolation. Context is skipped when no preceding
        usable exchange exists or when the mapping isn't provided (keeps
        unit tests and legacy callers working unchanged).
        """
        if not llm_output or "NO_MATCHES_FOUND" in llm_output:
            return []

        results: List[Dict] = []
        seen: set = set()
        for match in self._QUOTE_MATCH_RE.finditer(llm_output):
            try:
                idx = int(match.group(1))
            except ValueError:
                continue
            if idx < 0 or idx >= len(candidates) or idx in seen:
                continue
            seen.add(idx)

            ex, deponent, source = candidates[idx]
            relevance = match.group(2).strip()
            qa_text = self._build_qa_text_with_context(
                ex, deponent, source, full_exchanges_by_source
            )
            results.append({
                "deponent": deponent,
                "source": source,
                "page_line": ex.citation_range(),
                "relevance": relevance,
                "qa_text": qa_text,
            })
        return results

    def _build_qa_text_with_context(
        self,
        ex,
        deponent: str,
        source: str,
        full_exchanges_by_source: Optional[Dict[Tuple[str, str], Dict[int, object]]],
    ) -> str:
        """Assemble ``qa_text`` for a single quote, optionally prefixed with
        the preceding usable Q/A exchange so the quote reads in context.

        Returns just ``"Q. ...\\nA. ..."`` when no context is available.
        """
        parts: List[str] = []
        if full_exchanges_by_source:
            by_id = full_exchanges_by_source.get((deponent, source), {})
            prev = by_id.get(ex.id - 1)
            if prev is not None:
                parts.append(f"Q. {prev.question}")
                parts.append(f"A. {prev.answer}")
                parts.append("")  # blank line separator between context and main
        parts.append(f"Q. {ex.question}")
        parts.append(f"A. {ex.answer}")
        return "\n".join(parts)

    # Markers that indicate a parse failure bled reporter/colloquy text
    # into the Q or A body. TranscriptParser normally drops these but
    # occasionally leaks them on edge cases, and a leaked colloquy exchange
    # is worthless as a quote candidate.
    _LEAK_MARKERS = (
        "THE REPORTER",
        "THE STENOGRAPHIC REPORTER",
        "THE VIDEOGRAPHER",
        "THE WITNESS",
        "Anyone else?",
        "off the record",
    )

    def _is_usable_exchange(self, ex) -> bool:
        """Quick quality filter to drop exchanges that are unlikely to be
        quoteable — too short, too long, or clearly parse-broken.
        """
        q = (ex.question or "").strip()
        a = (ex.answer or "").strip()
        # Must have real content on both sides
        if len(q) < 10 or len(a) < 3:
            return False
        # Reject runaway exchanges that span too many transcript pages —
        # usually a parser state-machine failure where several Q/A pairs
        # merged into one.
        if (ex.page_end - ex.page_start) > 2:
            return False
        # Reject if the answer is absurdly long (likely multiple merged
        # answers or a transcript section that bled in)
        if len(a) > 1500:
            return False
        # Reject obvious colloquy/reporter leaks inside the text
        haystack = f"{q}\n{a}"
        if any(marker in haystack for marker in self._LEAK_MARKERS):
            return False
        return True

    def _load_transcript_candidates(
        self, transcript_paths: List[str]
    ) -> Tuple[List[tuple], Dict[Tuple[str, str], Dict[int, object]]]:
        """Parse transcripts into a flat list of candidate exchanges.

        Uses :class:`TranscriptParser` which handles both full-size and
        condensed (4-up mini-transcript) formats correctly and caches the
        structured index as JSON next to the PDF. Applies a cheap quality
        filter to drop obvious parser failures before ranking.

        Returns a tuple of:
          * the flat candidate list ``[(ex, deponent, source), ...]`` used
            for ranking + LLM selection, and
          * a per-transcript map ``{(deponent, source): {id: exchange}}``
            covering every *usable* exchange. The map lets
            :meth:`_parse_quote_match_results` look up preceding exchanges
            so a matched quote can be shown with a bit of surrounding
            context.
        """
        from icharlotte_core.deposition.transcript_parser import TranscriptParser

        parser = TranscriptParser()
        out: List[tuple] = []
        by_source: Dict[Tuple[str, str], Dict[int, object]] = {}
        for path in transcript_paths:
            fname = os.path.basename(path)
            ext = os.path.splitext(path)[1].lower()
            if ext != ".pdf":
                logger.warning(
                    "Quote search skipping non-PDF transcript %s — structured "
                    "parser only supports PDFs", fname
                )
                continue
            try:
                index = parser.parse(path)
            except Exception as e:
                logger.warning("Failed to parse transcript %s: %s", fname, e)
                continue
            deponent = index.deponent.last_name or os.path.splitext(fname)[0]
            id_map = by_source.setdefault((deponent, fname), {})
            for ex in index.exchanges:
                if not self._is_usable_exchange(ex):
                    continue
                out.append((ex, deponent, fname))
                id_map[ex.id] = ex
        return out, by_source

    def _format_brief_context(
        self, brief_sections: Optional[Dict[str, str]]
    ) -> str:
        """Render the current brief as a BRIEF CONTEXT block for the LLM.

        Skips empty sections. Section names are rendered using
        :data:`SECTION_HEADINGS` so the LLM sees them as the user would
        ("IV. LIABILITY" etc.). Returns an empty string when no usable
        context was provided — callers should omit the block entirely in
        that case.
        """
        if not brief_sections:
            return ""

        parts: List[str] = []
        for name in SECTION_ORDER:
            body = (brief_sections.get(name) or "").strip()
            if not body:
                continue
            roman, title = SECTION_HEADINGS.get(name, ("", name.upper()))
            header = f"{roman}. {title}" if roman else title
            parts.append(f"--- {header} ---\n{body}")

        # Include any extra sections that aren't in SECTION_ORDER (rare,
        # but harmless — e.g. a variant heading the parser kept under a
        # non-canonical key).
        for name, body in brief_sections.items():
            if name in SECTION_ORDER:
                continue
            body = (body or "").strip()
            if not body:
                continue
            parts.append(f"--- {name.upper()} ---\n{body}")

        if not parts:
            return ""
        return "BRIEF CONTEXT (current text of the mediation brief):\n\n" + "\n\n".join(parts)

    def search_quotes(
        self,
        transcript_paths: List[str],
        description: str,
        brief_sections: Optional[Dict[str, str]] = None,
    ) -> List[Dict]:
        """Search deposition transcripts for Q&A passages matching *description*.

        Pipeline:
          1. Parse each transcript into a :class:`TranscriptIndex` via
             :class:`TranscriptParser` (handles condensed / full-size).
          2. Keyword-rank the exchanges against *description*, keeping the
             top candidates. Prevents large depos from overflowing the LLM
             context and improves signal-to-noise.
          3. Render candidates as a numbered list, optionally prepended
             with a BRIEF CONTEXT block containing the brief's current
             section text (so the LLM can orient on the actual defense
             theory of the case and pick quotes that support the
             arguments already present).
          4. Ask the LLM to return only the IDs of genuine matches.
          5. Resolve the returned IDs back to the parsed exchanges.
             Hallucinations cannot escape the candidate set — text and
             citations always come from the index.

        Args:
            transcript_paths: Deposition PDFs to search.
            description: User's search description.
            brief_sections: Optional mapping of canonical section name →
                current body text. Populated by the chat-tab from
                ``generator.sections`` or by the Word popup from
                ``parse_brief_from_word_doc(doc).sections``. When provided,
                the LLM receives the brief content as orientation context.
                Pass ``None`` to skip this block (e.g. when there is no
                associated brief).
        """
        candidates, full_by_source = self._load_transcript_candidates(transcript_paths)
        if not candidates:
            return []

        ranked = self._rank_exchanges_by_keywords(candidates, description)
        if not ranked:
            return []

        candidate_text = self._format_candidates_for_llm(ranked)
        brief_context = self._format_brief_context(brief_sections)

        # Text sent to the LLM: BRIEF CONTEXT (if any) followed by the
        # numbered candidate list. The search_quotes prompt in
        # Scripts/prompts/mediation_brief/quote_search_current.txt
        # references both blocks by name.
        if brief_context:
            llm_text = f"{brief_context}\n\n---\n\nCANDIDATE EXCHANGES:\n\n{candidate_text}"
        else:
            llm_text = candidate_text

        search_prompt = self._get_section_prompt("quote_search")
        full_prompt = f"{search_prompt}\n\nUSER'S SEARCH DESCRIPTION:\n{description}"
        caller = LLMCaller()
        result = caller.call(
            prompt=full_prompt,
            text=llm_text,
            agent_id="agent_mediation_brief",
        )
        if not result:
            return []
        return self._parse_quote_match_results(result, ranked, full_by_source)

    def insert_quotes_quick(self, quotes: List[Dict], section_name: str,
                            subsection_title: Optional[str] = None) -> None:
        """Insert quote blocks into a section's text (Quick Insert mode)."""
        if section_name not in self.sections:
            return

        quote_blocks = []
        for q in quotes:
            block = (
                f"\n\nDEPO_QUOTE_START\n"
                f"{q['qa_text']}\n"
                f"DEPO_QUOTE_END\n"
                f"({q['deponent']} Depo Trns., at p. {q['page_line']}.)"
            )
            quote_blocks.append(block)
        insert_text = "".join(quote_blocks)

        section_text = self.sections[section_name]

        if subsection_title:
            pattern = re.compile(
                rf'^SUBSECTION:\s*{re.escape(subsection_title)}\s*$',
                re.MULTILINE
            )
            match = pattern.search(section_text)
            if match:
                next_sub = self._SUBSECTION_RE.search(section_text, match.end())
                if next_sub:
                    insert_pos = next_sub.start()
                    section_text = (
                        section_text[:insert_pos].rstrip()
                        + insert_text + "\n\n"
                        + section_text[insert_pos:]
                    )
                else:
                    section_text = section_text.rstrip() + insert_text
            else:
                section_text = section_text.rstrip() + insert_text
        else:
            section_text = section_text.rstrip() + insert_text

        self.sections[section_name] = section_text

    def reset(self) -> None:
        """Clear all generated state, resetting the generator to its initial condition."""
        self.sections = {}
        self.planning_output = ""
        self.document_content = ""
        self.caption_template_path = None
        self.is_active = False
        self.saved_path = None


# ---------------------------------------------------------------------------
# Background worker
# ---------------------------------------------------------------------------

from PySide6.QtCore import QThread, Signal  # noqa: E402 (import after class definition)


class RoutingWorker(QThread):
    """Background worker that routes a refinement message to section names.

    Runs the LLM routing call off the UI thread.

    Signals:
        result(list): list of section names to regenerate, or empty list
        error(str): error message
    """
    result = Signal(list)
    error = Signal(str)

    def __init__(self, generator: MediationBriefGenerator, user_message: str, parent=None):
        super().__init__(parent)
        self._generator = generator
        self._user_message = user_message

    def run(self):
        try:
            sections = self._generator.route_refinement(self._user_message)
            self.result.emit(sections)
        except Exception as exc:
            logger.exception("RoutingWorker: unhandled error")
            self.error.emit(str(exc))


class RefinementWorker(QThread):
    """Background worker for regenerating specific sections.

    Signals:
        section_complete(str, str): section_name, section_text
        all_complete(list): list of regenerated section names
        error(str): error message
    """
    section_complete = Signal(str, str)
    all_complete = Signal(list)
    error = Signal(str)

    def __init__(self, generator: MediationBriefGenerator, section_names: list,
                 instruction: str, parent=None):
        super().__init__(parent)
        self._generator = generator
        self._section_names = section_names
        self._instruction = instruction
        self._stop_requested = False

    def request_stop(self):
        self._stop_requested = True

    def run(self):
        try:
            regenerated = self._generator.refine_sections(
                self._section_names, self._instruction,
            )
            for name in regenerated:
                if self._stop_requested:
                    return
                self.section_complete.emit(name, self._generator.sections[name])
            self.all_complete.emit(regenerated)
        except Exception as exc:
            logger.exception("RefinementWorker: unhandled error")
            self.error.emit(str(exc))


class MediationBriefWorker(QThread):
    """QThread worker that runs the full mediation brief generation pipeline.

    Emits granular signals so a UI can display live progress.

    Signals:
        section_started(str, int, int): section_name, index, total — emitted
            before each step (including planning pass at index 0).
        section_complete(str, str): section_name, section_text — emitted after
            each section (not emitted for the planning pass).
        all_complete(dict): mapping of section_name → text for all sections,
            emitted when generation finishes successfully.
        error(str): error message emitted if an unhandled exception occurs.
    """

    section_started = Signal(str, int, int)
    section_complete = Signal(str, str)
    all_complete = Signal(dict)
    error = Signal(str)

    def __init__(self, generator: MediationBriefGenerator, parent=None):
        super().__init__(parent)
        self._generator = generator
        self._stop_requested: bool = False

    def request_stop(self) -> None:
        """Request that the worker stop after the current step."""
        self._stop_requested = True

    def run(self) -> None:
        """Execute the generation pipeline, emitting signals at each step."""
        try:
            total = len(GENERATION_ORDER) + 1

            # Step 0: planning pass
            self.section_started.emit("planning", 0, total)
            self._generator.run_planning_pass()

            # Steps 1..N: generate each section
            for i, section_name in enumerate(GENERATION_ORDER, start=1):
                if self._stop_requested:
                    break
                self.section_started.emit(section_name, i, total)
                result = self._generator.generate_section(section_name)
                if result:
                    self._generator.sections[section_name] = result
                    self.section_complete.emit(section_name, result)

            self._generator.is_active = True
            self.all_complete.emit(dict(self._generator.sections))

        except Exception as exc:
            logger.exception("MediationBriefWorker: unhandled error")
            self.error.emit(str(exc))
