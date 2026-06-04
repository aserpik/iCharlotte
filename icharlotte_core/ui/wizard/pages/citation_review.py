"""Shared citation-review UI for wizard tasks that draft a brief and verify
its citations (Oppose a Motion, Generate Motion).

Holds the verdict-colored body renderer, the per-citation detail helpers, the
right-side ``CitationDetailPanel`` (and legacy ``CitationDetailDialog``), and a
``CitationReviewOutputPage`` base class. All of it is side-agnostic: it
operates only on ``DraftDocument`` / ``CitationVerification``.
"""
from __future__ import annotations

import html
import os
import re
import shutil

from PySide6.QtCore import QUrl, Qt
from PySide6.QtGui import QDesktopServices
from PySide6.QtWidgets import (
    QDialog,
    QFileDialog,
    QHBoxLayout,
    QLabel,
    QMessageBox,
    QPushButton,
    QTextBrowser,
    QVBoxLayout,
    QWidget,
)

from icharlotte_core.opposition.models import DraftDocument


# ── Body-render block ────────────────────────────────────────────────────────

_HORIZONTAL_RULE_RE = re.compile(r"^[\*\-_]{3,}\s*$")
_MD_HEADING_RE = re.compile(r"^(#{1,6})\s+(.+?)\s*$")
_MD_ITALIC_RE = re.compile(r"\*{1,2}([^\*\n]+?)\*{1,2}")


_VERDICT_COLORS = {
    "SUPPORTED": "#1e8e3e",       # green
    "PARTIAL": "#f9ab00",          # yellow
    "NOT_SUPPORTED": "#c5221f",    # red
    "NOT_FOUND": "#c5221f",        # red (same as NOT_SUPPORTED)
    "UNVERIFIED": "#80868b",       # gray
}


def _color_for_verdict(verdict: str) -> str:
    return _VERDICT_COLORS.get((verdict or "").upper(), "#1a5dbf")  # default blue


def _render_draft_html(draft: DraftDocument) -> str:
    """Render the draft body into HTML with clickable citation anchors."""
    citation_spans = _build_citation_index(draft)
    body_html_lines: list[str] = []

    for raw_line in (draft.body_text or "").splitlines():
        stripped = raw_line.strip()
        if not stripped:
            body_html_lines.append("<p>&nbsp;</p>")
            continue
        if _HORIZONTAL_RULE_RE.match(stripped):
            continue
        heading_match = _MD_HEADING_RE.match(stripped)
        if heading_match:
            level = min(max(len(heading_match.group(1)), 2), 4)
            content = _format_inline_html(heading_match.group(2), citation_spans)
            body_html_lines.append(f"<h{level}>{content}</h{level}>")
            continue
        content = _format_inline_html(stripped, citation_spans)
        body_html_lines.append(f"<p>{content}</p>")

    title = html.escape(draft.title or "Memorandum")
    body = "\n".join(body_html_lines)
    return (
        "<html><body style=\"font-family:'Times New Roman',serif; font-size:13pt;\">"
        f"<h1 style=\"text-align:center;\">{title}</h1>{body}</body></html>"
    )


def _build_citation_index(draft: DraftDocument) -> list[tuple[str, int, str]]:
    """Return [(citation_text, draft_citation_index, verdict), ...] sorted by length desc.

    Sorting by length avoids partial-match collisions (e.g. "62 Cal. 4th 1081"
    must be wrapped before a shorter substring matches inside it).
    """
    spans: list[tuple[str, int, str]] = []
    for index, citation in enumerate(draft.citations or []):
        text = (citation.citation_text or "").strip()
        if text:
            spans.append((text, index, citation.verdict or ""))
    spans.sort(key=lambda triple: len(triple[0]), reverse=True)
    return spans


def _format_inline_html(line: str, citation_spans: list[tuple[str, int, str]]) -> str:
    """Escape, italicize *case names*, and wrap citation texts as clickable anchors."""
    # Mark italic spans with placeholders WITHOUT pre-escaping the inner
    # content. Pre-escaping plus the outer html.escape() double-escapes
    # entities (e.g. "&" → "&amp;" → "&amp;amp;"), which surfaces as
    # literal "&amp;" in the rendered HTML.
    italicized = _MD_ITALIC_RE.sub(
        lambda match: f"\x00ITA{match.group(1)}\x00ITAEND",
        line,
    )
    escaped = html.escape(italicized)
    # Replace closing placeholder before opening — otherwise "\x00ITA" matches
    # the prefix of "\x00ITAEND" and leaves "END" stranded in the output.
    escaped = escaped.replace("\x00ITAEND", "</i>").replace("\x00ITA", "<i>")
    escaped = escaped.replace(
        html.escape("[no case authority retrieved for this point]"),
        "<span style=\"color:#80868b;\">[no case authority retrieved for this point]</span>",
    )
    for citation_text, index, verdict in citation_spans:
        if not citation_text:
            continue
        # Build a regex that tolerates intervening <i> or </i> tags inserted
        # by the italics pass above (case names are typically wrapped in
        # *...* in the body but the citation_text comes without asterisks).
        escaped_cite = html.escape(citation_text)
        tolerant_pattern = r"(?:<i>|</i>)*".join(
            re.escape(ch) for ch in escaped_cite
        )
        color = _color_for_verdict(verdict)
        anchor = (
            f"<a href=\"citation:{index}\" "
            f"style=\"color:{color}; text-decoration:underline;\">"
            f"{escaped_cite}</a>"
        )
        escaped = re.sub(tolerant_pattern, anchor, escaped, count=0)
    return escaped


# ─── Verdict-aware citation rendering helpers (used by panel + dialog) ─────

_VERDICT_HEADER_COLORS = {
    "SUPPORTED": "#1e8e3e",
    "PARTIAL": "#f9ab00",
    "NOT_SUPPORTED": "#c5221f",
    "NOT_FOUND": "#c5221f",
    "UNVERIFIED": "#80868b",
}
_VERDICT_LABELS = {
    "SUPPORTED": "SUPPORTED",
    "PARTIAL": "PARTIAL",
    "NOT_SUPPORTED": "NOT SUPPORTED",
    "NOT_FOUND": "CITATION NOT FOUND",
    "UNVERIFIED": "UNVERIFIED",
}


def _citation_header_html(citation, verdict: str) -> str:
    color = _VERDICT_HEADER_COLORS.get(verdict, "#80868b")
    label = _VERDICT_LABELS.get(verdict, verdict or "UNVERIFIED")
    title = html.escape(citation.case_name or citation.citation_text or "Citation")
    return (
        f"<div style='border-left: 6px solid {color}; padding-left: 8px;'>"
        f"<div style='color:{color}; font-weight:bold; font-size:14pt;'>{html.escape(label)}</div>"
        f"<div style='font-size:11pt;'>{title}</div>"
        "</div>"
    )


def _citation_body_html(citation, verdict: str) -> str:
    parts: list[str] = []

    prop = (citation.proposition or "").strip()
    if prop:
        parts.append(
            f"<p><b>Brief's proposition:</b><br><i>{html.escape(prop)}</i></p>"
        )

    if verdict == "SUPPORTED":
        if citation.evidence:
            parts.append(
                f"<p><b>Verified holding:</b><br>{html.escape(citation.evidence)}</p>"
            )
        if citation.note:
            parts.append(f"<p><b>Verifier note:</b> {html.escape(citation.note)}</p>")
        if getattr(citation, "citation_count", None) is not None:
            year = (getattr(citation, "latest_citing_year", "") or "").strip()
            tail = f", most recently {html.escape(year)}" if year else ""
            parts.append(
                f"<p style='color:#80868b;'><b>Good-law hint:</b> cited by "
                f"{citation.citation_count} case(s){tail}. (Not a Shepard's check — "
                "confirm the case is still good law.)</p>"
            )

    elif verdict == "PARTIAL":
        if citation.evidence:
            parts.append(
                f"<p><b>Relevant passage:</b><br>{html.escape(citation.evidence)}</p>"
            )
        if citation.note:
            parts.append(
                f"<p><b>Why partial:</b> {html.escape(citation.note)}</p>"
            )
        if getattr(citation, "citation_count", None) is not None:
            year = (getattr(citation, "latest_citing_year", "") or "").strip()
            tail = f", most recently {html.escape(year)}" if year else ""
            parts.append(
                f"<p style='color:#80868b;'><b>Good-law hint:</b> cited by "
                f"{citation.citation_count} case(s){tail}. (Not a Shepard's check — "
                "confirm the case is still good law.)</p>"
            )

    elif verdict == "NOT_SUPPORTED":
        parts.append(
            "<p><b>⚠ The authority does NOT hold what the brief claims.</b></p>"
        )
        if citation.evidence:
            parts.append(
                f"<p><b>What it actually holds:</b><br>{html.escape(citation.evidence)}</p>"
            )
        if citation.note:
            parts.append(f"<p><b>Verifier note:</b> {html.escape(citation.note)}</p>")

    elif verdict == "NOT_FOUND":
        parts.append(
            "<p><b>⚠ This citation was not found in the authoritative source.</b></p>"
            "<p>Likely causes:</p>"
            "<ul>"
            "<li>The case or statute was invented by the LLM</li>"
            "<li>The citation is mis-typed</li>"
            "<li>The case is unpublished or pre-1900</li>"
            "</ul>"
        )
        if citation.note:
            parts.append(f"<p><b>Verifier note:</b> {html.escape(citation.note)}</p>")

    else:  # UNVERIFIED
        # The note (when present) explains the specific reason -- cluster
        # found but no opinion text, LLM call failed, prompt missing,
        # parse failure, etc. Show the note directly. Only fall back to
        # the generic "doesn't cover" message when no note is available
        # (which corresponds to genuinely out-of-scope citation kinds
        # like federal cases or local court rules).
        if citation.note:
            parts.append(
                f"<p><b>⚠ Could not verify:</b> {html.escape(citation.note)}</p>"
            )
        else:
            parts.append(
                "<p>The verifier doesn't cover this source in v1. Verify manually.</p>"
            )

    return "\n".join(parts)


def _run_find_replacement(parent_widget, citation):
    """Run the find-replacement LLM call and pop a result summary.

    Shared by the panel button and the (legacy) dialog button. Shows
    a QMessageBox with candidates or with the appropriate error. Returns
    nothing; intended as a fire-and-forget UI action.
    """
    from icharlotte_core.llm_config import call_llm
    from icharlotte_core.opposition.verifier import (
        build_opposition_verifier,
        find_replacement_candidates,
    )

    token = os.environ.get("COURTLISTENER_API_TOKEN", "").strip()
    if not token:
        QMessageBox.warning(
            parent_widget,
            "Missing API token",
            "COURTLISTENER_API_TOKEN is not set; replacement search is unavailable.",
        )
        return

    def llm(system_prompt, user_prompt):
        return call_llm(
            user_prompt,
            system_prompt,
            task_type="general",
            agent_id="agent_oppose_motion",
        ) or ""

    verifier = build_opposition_verifier(courtlistener_token=token, llm_callback=llm)
    candidates = find_replacement_candidates(
        failed_citation=citation,
        verifier=verifier,
        llm_callback=llm,
    )

    if not candidates:
        QMessageBox.information(
            parent_widget,
            "No replacements found",
            "No supported replacement candidates were found. Try editing the proposition or searching manually.",
        )
        return
    lines = [f"{c.citation_text} — {c.verdict}\n  {c.note}" for c in candidates]
    QMessageBox.information(parent_widget, "Replacement candidates", "\n\n".join(lines))


class CitationDetailPanel(QWidget):
    """Right-side panel showing verdict-aware citation detail.

    Replaces the popup dialog: anchor clicks in the draft body update this
    panel in-place. Exposes ``set_citation()`` and ``clear()``.
    """

    def __init__(self, parent: QWidget | None = None) -> None:
        super().__init__(parent)
        self.citation = None

        layout = QVBoxLayout(self)
        layout.setContentsMargins(8, 8, 8, 8)
        layout.setSpacing(8)

        self.header_label = QLabel("")
        self.header_label.setTextFormat(Qt.TextFormat.RichText)
        self.header_label.setWordWrap(True)
        layout.addWidget(self.header_label)

        self.body_browser = QTextBrowser()
        self.body_browser.setOpenExternalLinks(False)
        self.body_browser.setStyleSheet("QTextBrowser { background: transparent; }")
        layout.addWidget(self.body_browser, 1)

        button_row = QHBoxLayout()
        button_row.addStretch()
        self.find_btn = QPushButton("Find replacement case")
        self.find_btn.clicked.connect(self._on_find_replacement)
        self.find_btn.setVisible(False)
        button_row.addWidget(self.find_btn)
        self.open_btn = QPushButton("Open source")
        self.open_btn.clicked.connect(self._open_opinion_url)
        self.open_btn.setVisible(False)
        button_row.addWidget(self.open_btn)
        layout.addLayout(button_row)

        # Initial empty state.
        self.clear("Select a citation in the draft to see verification details.")

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    def set_citation(self, citation) -> None:
        self.citation = citation
        verdict = (citation.verdict or "UNVERIFIED").upper()
        self.header_label.setText(_citation_header_html(citation, verdict))
        self.body_browser.setHtml(_citation_body_html(citation, verdict))

        # Toggle action buttons based on what makes sense for this citation.
        if citation.opinion_url:
            label = "Open in CourtListener" if citation.kind == "case" else "Open in leginfo"
            self.open_btn.setText(label)
            self.open_btn.setVisible(True)
        else:
            self.open_btn.setVisible(False)

        self.find_btn.setVisible(verdict in {"NOT_SUPPORTED", "NOT_FOUND"})

    def clear(self, message: str = "") -> None:
        self.citation = None
        self.header_label.setText("")
        self.body_browser.setHtml(
            f"<p style='color:#80868b;'>{html.escape(message)}</p>" if message else ""
        )
        self.open_btn.setVisible(False)
        self.find_btn.setVisible(False)

    @property
    def body_html(self) -> str:
        """Current rendered body HTML (for tests)."""
        return self.body_browser.toHtml()

    # ------------------------------------------------------------------
    # Handlers
    # ------------------------------------------------------------------

    def _open_opinion_url(self) -> None:
        if self.citation and self.citation.opinion_url:
            QDesktopServices.openUrl(QUrl(self.citation.opinion_url))

    def _on_find_replacement(self) -> None:
        if self.citation is not None:
            _run_find_replacement(self, self.citation)


class CitationDetailDialog(QDialog):
    """Modal popup version (legacy). The output page no longer auto-opens
    this; it's retained as an opt-in surface for programmatic use and for
    backward compatibility with existing tests.
    """

    _VERDICT_HEADER_COLORS = _VERDICT_HEADER_COLORS
    _VERDICT_LABELS = _VERDICT_LABELS

    def __init__(self, citation, parent: QWidget | None = None):
        super().__init__(parent)
        self.citation = citation
        self.setWindowTitle(citation.case_name or citation.citation_text or "Citation")
        self.resize(720, 540)

        verdict = (citation.verdict or "UNVERIFIED").upper()

        layout = QVBoxLayout(self)
        layout.setContentsMargins(16, 16, 16, 16)
        layout.setSpacing(10)

        self.header = QLabel(_citation_header_html(citation, verdict))
        self.header.setTextFormat(Qt.TextFormat.RichText)
        self.header.setWordWrap(True)
        layout.addWidget(self.header)

        self.body_html = _citation_body_html(citation, verdict)
        self.body_label = QLabel(self.body_html)
        self.body_label.setTextFormat(Qt.TextFormat.RichText)
        self.body_label.setWordWrap(True)
        self.body_label.setOpenExternalLinks(False)
        layout.addWidget(self.body_label, 1)

        button_row = QHBoxLayout()
        button_row.addStretch()
        if citation.opinion_url:
            label = "Open in CourtListener" if citation.kind == "case" else "Open in leginfo"
            open_btn = QPushButton(label)
            open_btn.clicked.connect(self._open_opinion_url)
            button_row.addWidget(open_btn)
        close_btn = QPushButton("Close")
        close_btn.clicked.connect(self.accept)
        button_row.addWidget(close_btn)

        if verdict in {"NOT_SUPPORTED", "NOT_FOUND"}:
            self.find_btn = QPushButton("Find replacement case")
            self.find_btn.clicked.connect(self._on_find_replacement)
            button_row.insertWidget(button_row.count() - 1, self.find_btn)

        layout.addLayout(button_row)

    def _open_opinion_url(self) -> None:
        if self.citation.opinion_url:
            QDesktopServices.openUrl(QUrl(self.citation.opinion_url))

    def _on_find_replacement(self) -> None:
        _run_find_replacement(self, self.citation)


class CitationReviewOutputPage(QWidget):
    """Output page for tasks that draft a brief then verify its citations.

    Renders the draft body with color-coded clickable citation anchors, shows
    a verdict summary banner, and drives a right-side ``CitationDetailPanel``
    that updates when an anchor is clicked. Subclasses customize the title,
    the empty-citations message, and the bottom action buttons.
    """

    #: Fallback title when ``draft.title`` is empty; also used for the
    #: suggested save filename and the save dialog caption.
    default_title = "Memorandum"

    def __init__(self, parent: QWidget | None = None) -> None:
        super().__init__(parent)
        self.draft = DraftDocument()
        outer = QVBoxLayout(self)
        outer.setContentsMargins(24, 24, 24, 24)
        outer.setSpacing(12)

        self.summary_banner = QLabel("")
        self.summary_banner.setTextFormat(Qt.TextFormat.RichText)
        self.summary_banner.setWordWrap(True)
        self.summary_banner.setVisible(False)
        outer.addWidget(self.summary_banner)

        layout = QHBoxLayout()
        layout.setSpacing(12)
        self.editor = QTextBrowser()
        self.editor.setOpenLinks(False)
        self.editor.setOpenExternalLinks(False)
        self.editor.anchorClicked.connect(self._on_anchor_clicked)
        layout.addWidget(self.editor, 2)

        self.detail_panel = CitationDetailPanel()
        layout.addWidget(self.detail_panel, 1)
        outer.addLayout(layout, 1)

        row = QHBoxLayout()
        row.addStretch()
        self._build_action_buttons(row)
        outer.addLayout(row)

    # -- overridable seams ------------------------------------------------

    def empty_citations_message(self) -> str:
        return (
            "No citations were detected. If California case-law research "
            "returned no results, the draft was written without case "
            "citations. Review for any factual or statutory support that may "
            "need strengthening."
        )

    def _build_action_buttons(self, row: QHBoxLayout) -> None:
        """Populate the bottom button row. Base adds only a Save button;
        subclasses override and call ``_add_save_button`` where they want it."""
        self._add_save_button(row)

    def _add_save_button(self, row: QHBoxLayout) -> None:
        self.save_btn = QPushButton("Save")
        self.save_btn.clicked.connect(self.save_as)
        row.addWidget(self.save_btn)

    # -- rendering --------------------------------------------------------

    def show_result(self, draft: DraftDocument) -> None:
        self.draft = draft
        self.editor.setHtml(_render_draft_html(draft))
        self._refresh_summary_banner()
        if draft.citations:
            self.show_citation(0)
        else:
            self.detail_panel.clear(self.empty_citations_message())

    def show_citation(self, index: int) -> None:
        if index < 0 or index >= len(self.draft.citations):
            return
        self.detail_panel.set_citation(self.draft.citations[index])

    def _refresh_summary_banner(self) -> None:
        if not self.draft.citations:
            self.summary_banner.setVisible(False)
            return
        counts: dict[str, int] = {}
        for cv in self.draft.citations:
            verdict = (cv.verdict or "UNVERIFIED").upper()
            counts[verdict] = counts.get(verdict, 0) + 1
        total = sum(counts.values())
        red = counts.get("NOT_SUPPORTED", 0) + counts.get("NOT_FOUND", 0)
        parts = [
            f"<b>Verification:</b> {total} citation(s) checked &mdash; ",
            f"\U0001F7E2 {counts.get('SUPPORTED', 0)} supported, ",
            f"\U0001F7E1 {counts.get('PARTIAL', 0)} partial, ",
            f"\U0001F534 {red} flagged, ",
            f"⚪ {counts.get('UNVERIFIED', 0)} unverified.",
        ]
        warning = ""
        if red > 0:
            warning = (
                f"<br><span style='color:#c5221f;'>⚠ {red} citation(s) don't "
                "support what the brief claims. Review the red-flagged cites "
                "before filing.</span>"
            )
        self.summary_banner.setText("".join(parts) + warning)
        self.summary_banner.setVisible(True)

    def _on_anchor_clicked(self, url: QUrl) -> None:
        scheme = url.scheme()
        if scheme == "citation":
            try:
                index = int(url.path().lstrip("/") or url.host() or "0")
            except (TypeError, ValueError):
                return
            self.show_citation(index)
            return
        QDesktopServices.openUrl(url)

    @property
    def output_path(self) -> str:
        return self.draft.preview_path

    def load_output(self, output_path: str) -> None:
        body_text = ""
        if output_path and output_path.lower().endswith(".docx") and os.path.isfile(output_path):
            try:
                from docx import Document

                doc = Document(output_path)
                body_text = "\n".join(p.text for p in doc.paragraphs if p.text)
            except Exception:
                body_text = f"(Could not render generated document: {output_path})"
        self.show_result(
            DraftDocument(
                title=os.path.splitext(os.path.basename(output_path or ""))[0] or self.default_title,
                body_text=body_text,
                preview_path=output_path or "",
            )
        )

    @staticmethod
    def default_save_dir(preview_path: str) -> str:
        marker = os.path.join(".icharlotte", "wizard_previews")
        before_marker, _, _after_marker = preview_path.partition(marker)
        if before_marker:
            return os.path.dirname(os.path.normpath(before_marker))
        return os.path.dirname(preview_path)

    def save_as(self) -> None:
        if not self.draft.preview_path:
            QMessageBox.warning(
                self, "No preview", "No generated preview is available."
            )
            return

        red = sum(
            1 for cv in self.draft.citations
            if (cv.verdict or "").upper() in {"NOT_SUPPORTED", "NOT_FOUND"}
        )
        if red > 0:
            choice = QMessageBox.question(
                self,
                "Citations flagged",
                (
                    f"This document has {red} citation(s) flagged as "
                    "NOT_SUPPORTED or NOT_FOUND.\n\nSave anyway?"
                ),
                QMessageBox.StandardButton.Save | QMessageBox.StandardButton.Cancel,
                QMessageBox.StandardButton.Cancel,
            )
            if choice != QMessageBox.StandardButton.Save:
                return

        suggested = os.path.join(
            self.default_save_dir(self.draft.preview_path),
            f"{self.draft.title or self.default_title}.docx",
        )
        target, _ = QFileDialog.getSaveFileName(
            self,
            f"Save {self.default_title}",
            suggested,
            "Word Documents (*.docx);;All files (*.*)",
        )
        if not target:
            return
        if not target.lower().endswith(".docx"):
            target += ".docx"
        if os.path.abspath(target) == os.path.abspath(self.draft.preview_path):
            QMessageBox.warning(
                self,
                "Choose another location",
                "Select a location outside the internal preview file.",
            )
            return
        try:
            shutil.copyfile(self.draft.preview_path, target)
        except Exception as exc:
            QMessageBox.critical(self, "Save failed", f"Could not save file:\n{exc}")
            return
        QMessageBox.information(self, "Saved", f"Saved:\n{target}")
