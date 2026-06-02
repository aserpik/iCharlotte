"""Custom Wizard task page for generating motions from scratch.

Sibling of oppose_motion_page.py. Instead of analyzing a motion being opposed,
the user picks a motion type and target documents; an LLM proposes the grounds /
relief and a section outline (both editable), then the motion is drafted,
citation-verified, and assembled — reusing the opposition research/verify spine
and the motion_generation drafter/assembler.
"""
from __future__ import annotations

import os

from PySide6.QtCore import QThread, QUrl, Qt, Signal
from PySide6.QtGui import QDesktopServices
from PySide6.QtWidgets import (
    QComboBox,
    QHBoxLayout,
    QLabel,
    QLineEdit,
    QMessageBox,
    QPlainTextEdit,
    QPushButton,
    QStackedWidget,
    QTextBrowser,
    QTreeWidget,
    QTreeWidgetItem,
    QVBoxLayout,
    QWidget,
)

from icharlotte_core.discovery.assembler import DiscoveryAssembler
from icharlotte_core.opposition.argument_research import research_arguments
from icharlotte_core.opposition.citation_parser import extract_citations
from icharlotte_core.opposition.extraction import (
    extract_context_bundle,
    is_supported_context_file,
)
from icharlotte_core.opposition.models import (
    DraftDocument,
    MotionMetadata,
    OutlineNode,
)
from icharlotte_core.opposition.outline import selected_section_plan
from icharlotte_core.opposition.verifier import (
    build_local_opposition_verifier,
    build_opposition_verifier,
    enrich_with_pool_signals,
    pool_membership_check,
)
from icharlotte_core.ui.wizard.pages.oppose_motion_page import _make_local_corpus
from icharlotte_core.ui.wizard.pages.status_page import StatusPage
from icharlotte_core.ui.wizard.task_scaffold import WizardTaskContainer
from icharlotte_core.ui.context_files_dialog import ContextFilesDialog

from icharlotte_core.motion_generation.analyzer import analyze_target, outline_from_config
from icharlotte_core.motion_generation.assembler import assemble_motion_preview
from icharlotte_core.motion_generation.config import (
    MOTION_TYPE_CONFIGS,
    get_motion_config,
)
from icharlotte_core.motion_generation.drafter import draft_motion


SETTINGS_PAGE_REVIEW = 0
SETTINGS_PAGE_OUTLINE = 1
TASK_PAGE_SETTINGS = 0
TASK_PAGE_STATUS = 1
TASK_PAGE_OUTPUT = 2

# Order of the type dropdown (configured types first, generic last).
_TYPE_ORDER = ["compel", "demurrer", "strike", "generic"]


class GenerateMotionSettingsPage(QStackedWidget):
    """Review (grounds/relief) + editable outline screens for motion generation."""

    run_requested = Signal(dict)

    def __init__(
        self,
        case_root: str,
        file_number: str,
        motion_type_id: str,
        target_files: list[str],
        parent: QWidget | None = None,
    ):
        super().__init__(parent)
        self.case_root = case_root
        self.file_number = file_number
        self.motion_type_id = motion_type_id or "generic"
        self.target_files = list(target_files)
        self.metadata = MotionMetadata(motion_type=get_motion_config(self.motion_type_id).display_name)
        self.outline: list[OutlineNode] = []
        self._build_review_page()
        self._build_outline_page()

    # ---- Build ---------------------------------------------------------- #

    def _build_review_page(self) -> None:
        page = QWidget()
        layout = QVBoxLayout(page)
        layout.setContentsMargins(24, 24, 24, 24)
        layout.setSpacing(10)

        self.files_label = QLabel()
        layout.addWidget(self.files_label)
        self._refresh_files_label()

        layout.addWidget(QLabel("Motion type"))
        self.type_combo = QComboBox()
        for type_id in _TYPE_ORDER:
            self.type_combo.addItem(MOTION_TYPE_CONFIGS[type_id].display_name, type_id)
        idx = self.type_combo.findData(self.motion_type_id)
        if idx >= 0:
            self.type_combo.setCurrentIndex(idx)
        layout.addWidget(self.type_combo)

        layout.addWidget(QLabel("Relief requested (you can edit)"))
        self.relief_edit = QLineEdit()
        self.relief_edit.setPlaceholderText("Relief requested")
        layout.addWidget(self.relief_edit)

        layout.addWidget(QLabel("Grounds (one per line; add your own as needed)"))
        self.arguments_edit = QPlainTextEdit()
        self.arguments_edit.setPlaceholderText("Proposed grounds will appear here")
        layout.addWidget(self.arguments_edit, 1)

        row = QHBoxLayout()
        row.addStretch()
        self.continue_btn = QPushButton("Review Outline")
        self.continue_btn.clicked.connect(self._on_continue_to_outline)
        row.addWidget(self.continue_btn)
        layout.addLayout(row)
        self.addWidget(page)

    def _build_outline_page(self) -> None:
        page = QWidget()
        layout = QVBoxLayout(page)
        layout.setContentsMargins(24, 24, 24, 24)
        layout.setSpacing(10)

        layout.addWidget(QLabel("Motion Outline"))
        self.outline_tree = QTreeWidget()
        self.outline_tree.setHeaderLabels(["Include / Heading"])
        layout.addWidget(self.outline_tree, 1)

        row = QHBoxLayout()
        self.add_heading_btn = QPushButton("Add Heading")
        self.add_heading_btn.clicked.connect(self._add_heading)
        row.addWidget(self.add_heading_btn)
        row.addStretch()
        self.generate_btn = QPushButton("Generate Motion")
        self.generate_btn.clicked.connect(self._emit_run_requested)
        row.addWidget(self.generate_btn)
        layout.addLayout(row)
        self.addWidget(page)

    # ---- State ---------------------------------------------------------- #

    def current_motion_type_id(self) -> str:
        data = self.type_combo.currentData()
        return data or "generic"

    def set_metadata(self, metadata: MotionMetadata) -> None:
        self.metadata = metadata
        self.relief_edit.setText(metadata.relief_requested)
        self.arguments_edit.setPlainText("\n".join(metadata.principal_arguments))

    def current_metadata(self) -> MotionMetadata:
        metadata = MotionMetadata.from_dict(self.metadata.to_dict())
        metadata.motion_type = get_motion_config(self.current_motion_type_id()).display_name
        metadata.relief_requested = self.relief_edit.text().strip()
        metadata.principal_arguments = [
            line.strip()
            for line in self.arguments_edit.toPlainText().splitlines()
            if line.strip()
        ]
        return metadata

    def can_continue_to_outline(self) -> bool:
        return self.current_metadata().required_missing() == []

    def set_outline(self, outline: list[OutlineNode]) -> None:
        self.outline = list(outline)
        self.outline_tree.clear()
        for node in self.outline:
            self.outline_tree.addTopLevelItem(self._item_from_node(node))
        self.outline_tree.expandAll()

    def to_dict(self) -> dict:
        return {
            "motion_type_id": self.current_motion_type_id(),
            "target_files": list(self.target_files),
            "metadata": self.current_metadata().to_dict(),
            "outline": [node.to_dict() for node in self._outline_from_tree()],
        }

    def from_dict(self, data: dict | None) -> None:
        if not isinstance(data, dict):
            data = {}
        self.target_files = list(data.get("target_files", self.target_files))
        self.motion_type_id = data.get("motion_type_id", self.motion_type_id)
        self._refresh_files_label()
        idx = self.type_combo.findData(self.motion_type_id)
        if idx >= 0:
            self.type_combo.setCurrentIndex(idx)
        self.set_metadata(MotionMetadata.from_dict(data.get("metadata")))
        self.set_outline(
            [
                OutlineNode.from_dict(item)
                for item in data.get("outline", [])
                if isinstance(item, dict)
            ]
        )

    # ---- Handlers ------------------------------------------------------- #

    def _on_continue_to_outline(self) -> None:
        if not self.can_continue_to_outline():
            QMessageBox.warning(
                self,
                "Missing required fields",
                "Relief requested and at least one ground are required.",
            )
            return
        self.setCurrentIndex(SETTINGS_PAGE_OUTLINE)

    def _refresh_files_label(self) -> None:
        if not self.target_files:
            self.files_label.setText("Target documents: (none selected)")
            return
        names = ", ".join(os.path.basename(f) for f in self.target_files)
        self.files_label.setText(f"Target documents: {names}")

    def _emit_run_requested(self) -> None:
        self.run_requested.emit(self.to_dict())

    def _add_heading(self) -> None:
        item = QTreeWidgetItem(["New heading"])
        item.setFlags(item.flags() | Qt.ItemFlag.ItemIsEditable | Qt.ItemFlag.ItemIsUserCheckable)
        item.setCheckState(0, Qt.CheckState.Checked)
        item.setData(0, Qt.ItemDataRole.UserRole, "")
        self.outline_tree.addTopLevelItem(item)
        self.outline_tree.editItem(item, 0)

    def _item_from_node(self, node: OutlineNode) -> QTreeWidgetItem:
        item = QTreeWidgetItem([node.text])
        item.setData(0, Qt.ItemDataRole.UserRole, node.id)
        item.setFlags(item.flags() | Qt.ItemFlag.ItemIsEditable | Qt.ItemFlag.ItemIsUserCheckable)
        item.setCheckState(0, Qt.CheckState.Checked if node.selected else Qt.CheckState.Unchecked)
        for child in node.children:
            item.addChild(self._item_from_node(child))
        return item

    def _outline_from_tree(self) -> list[OutlineNode]:
        return [
            self._node_from_item(self.outline_tree.topLevelItem(index))
            for index in range(self.outline_tree.topLevelItemCount())
        ]

    def _node_from_item(self, item: QTreeWidgetItem) -> OutlineNode:
        return OutlineNode(
            id=item.data(0, Qt.ItemDataRole.UserRole) or "",
            text=item.text(0).strip(),
            selected=item.checkState(0) == Qt.CheckState.Checked,
            children=[
                self._node_from_item(item.child(index))
                for index in range(item.childCount())
            ],
        )


def _make_llms():
    """Return (analysis_llm, draft_llm, query_llm_factory) reusing the
    oppose_motion agent config so model selection is shared."""
    from icharlotte_core.llm_config import call_llm

    def analysis_llm(system_prompt, user_prompt):
        return call_llm(
            user_prompt, system_prompt, task_type="general", agent_id="agent_chat"
        ) or ""

    def draft_llm(system_prompt, user_prompt):
        return call_llm(
            user_prompt, system_prompt, task_type="general",
            agent_id="agent_oppose_motion",
        ) or ""

    def make_pass_llm(pass_name):
        def _llm(system_prompt, user_prompt):
            return call_llm(
                user_prompt, system_prompt, task_type="general",
                agent_id="agent_oppose_motion", pass_name=pass_name,
                pass_agent_id="agent_oppose_motion",
            ) or ""
        return _llm

    return analysis_llm, draft_llm, make_pass_llm


class GenerateMotionAnalysisWorker(QThread):
    progress = Signal(str)
    finished_analysis = Signal(bool, object)

    def __init__(self, settings: dict, parent=None):
        super().__init__(parent)
        self.settings = dict(settings or {})

    def run(self) -> None:
        try:
            config = get_motion_config(self.settings.get("motion_type_id"))
            self.progress.emit("Extracting target documents...")
            target_text, warnings = extract_context_bundle(
                self.settings.get("target_files", [])
            )
            for warning in warnings:
                self.progress.emit(f"WARNING: {warning}")
            if not target_text.strip():
                self.finished_analysis.emit(
                    False, "Could not read any text from the target documents."
                )
                return

            analysis_llm, _, _ = _make_llms()
            self.progress.emit("Proposing grounds and relief...")
            metadata = analyze_target(config, target_text, llm_callback=analysis_llm)
            outline = outline_from_config(config)
            self.finished_analysis.emit(
                True,
                {"metadata": metadata, "outline": outline, "target_text": target_text},
            )
        except Exception as exc:  # noqa: BLE001
            self.finished_analysis.emit(False, str(exc))


class GenerateMotionWorker(QThread):
    progress = Signal(str)
    finished_result = Signal(bool, object)

    def __init__(self, case_path: str, file_number: str, settings: dict, parent=None):
        super().__init__(parent)
        self.case_path = case_path
        self.file_number = file_number
        self.settings = dict(settings or {})

    def run(self) -> None:
        try:
            config = get_motion_config(self.settings.get("motion_type_id"))
            self.progress.emit("Extracting target documents...")
            target_text, warnings = extract_context_bundle(
                self.settings.get("target_files", [])
            )
            for warning in warnings:
                self.progress.emit(f"WARNING: {warning}")

            metadata = MotionMetadata.from_dict(self.settings.get("metadata"))
            outline = [
                OutlineNode.from_dict(item)
                for item in self.settings.get("outline", [])
                if isinstance(item, dict)
            ]
            plan = selected_section_plan(outline)

            _, draft_llm, make_pass_llm = _make_llms()

            # Retrieval-first grounding (prefer local corpus; fall back to CL API).
            corpus = _make_local_corpus()
            token = os.environ.get("COURTLISTENER_API_TOKEN", "").strip()
            retrieved = []
            if metadata.principal_arguments and (corpus is not None or token):
                client = corpus
                if client is None:
                    from icharlotte_core.legal_research.sources.courtlistener import (
                        CourtListenerClient,
                    )
                    client = CourtListenerClient(token)
                    self.progress.emit("Local corpus not built; using CourtListener API...")
                else:
                    self.progress.emit("Researching authorities locally...")
                retrieved = research_arguments(
                    metadata.principal_arguments,
                    cl_client=client,
                    query_llm=make_pass_llm("research_queries"),
                    rerank_llm=make_pass_llm("rerank_select"),
                    max_workers=4,
                    on_progress=self.progress.emit,
                )
                self.progress.emit(f"Retrieved {len(retrieved)} grounded authorities.")
            else:
                self.progress.emit(
                    "WARNING: no grounded research available; drafting from statutes only."
                )

            self.progress.emit("Drafting motion memorandum...")
            draft = draft_motion(
                config,
                metadata,
                plan,
                target_text,
                "",
                style_exemplars=[],
                retrieved_authorities=retrieved,
                llm_callback=draft_llm,
            )
            if not draft.body_text.strip():
                reason = (draft.rejection_reason or "unknown reason").strip()
                self.finished_result.emit(False, f"Drafting failed: {reason}")
                return

            citations = extract_citations(draft.body_text)
            if citations:
                to_verify, off_pool = pool_membership_check(citations, retrieved)
                self.progress.emit(f"Verifying citations ({len(to_verify)} found)...")
                if corpus is not None:
                    verifier = build_local_opposition_verifier(
                        corpus=corpus, llm_callback=draft_llm, max_workers=4
                    )
                    verified = verifier.verify_all(to_verify, on_progress=self.progress.emit)
                elif token:
                    verifier = build_opposition_verifier(
                        courtlistener_token=token, llm_callback=draft_llm, max_workers=4
                    )
                    verified = verifier.verify_all(to_verify, on_progress=self.progress.emit)
                else:
                    verified = []
                draft.citations = sorted(
                    list(verified) + list(off_pool),
                    key=lambda cv: cv.body_offset if cv.body_offset is not None else 0,
                )
                enrich_with_pool_signals(draft.citations, retrieved)
            else:
                self.progress.emit("WARNING: No citations detected in the draft.")
                draft.citations = []

            preview_dir = os.path.join(
                self.case_path, "NOTES", "AI OUTPUT", ".icharlotte",
                "wizard_previews", "generate_motion",
            )
            preview_path = os.path.join(preview_dir, "Motion Preview.docx")
            caption_path = DiscoveryAssembler.find_caption_page(self.case_path) or ""
            assemble_motion_preview(
                draft=draft, output_path=preview_path, config=config,
                caption_path=caption_path,
            )
            draft.preview_path = preview_path
            self.finished_result.emit(True, draft)
        except Exception as exc:  # noqa: BLE001
            self.finished_result.emit(False, str(exc))


class GenerateMotionOutputPage(QWidget):
    def __init__(self, parent: QWidget | None = None):
        super().__init__(parent)
        layout = QVBoxLayout(self)
        layout.setContentsMargins(24, 24, 24, 24)
        layout.setSpacing(10)

        self.heading = QLabel("Generated Motion")
        layout.addWidget(self.heading)

        self.body = QTextBrowser()
        layout.addWidget(self.body, 1)

        row = QHBoxLayout()
        self.open_btn = QPushButton("Open in Word")
        self.open_btn.clicked.connect(self._open_preview)
        self.open_btn.setEnabled(False)
        row.addWidget(self.open_btn)
        row.addStretch()
        layout.addLayout(row)

        self._preview_path = ""

    @property
    def output_path(self) -> str:
        """Path to the generated preview (consumed by the snapshot logic)."""
        return self._preview_path

    def show_result(self, draft: DraftDocument) -> None:
        self._preview_path = getattr(draft, "preview_path", "") or ""
        self.heading.setText(draft.title or "Generated Motion")
        self.body.setPlainText(draft.body_text or "")
        self.open_btn.setEnabled(bool(self._preview_path) and os.path.isfile(self._preview_path))

    def load_output(self, path: str) -> None:
        """Restore a previously-generated preview from disk (best-effort body text)."""
        self._preview_path = path or ""
        self.open_btn.setEnabled(bool(self._preview_path) and os.path.isfile(self._preview_path))
        try:
            from docx import Document

            text = "\n".join(p.text for p in Document(path).paragraphs)
            self.body.setPlainText(text)
        except Exception:  # noqa: BLE001
            pass

    def _open_preview(self) -> None:
        if self._preview_path and os.path.isfile(self._preview_path):
            QDesktopServices.openUrl(QUrl.fromLocalFile(self._preview_path))


class GenerateMotionTaskTab(WizardTaskContainer):
    task_completed = Signal(dict)

    def __init__(
        self,
        spec,
        case_path: str,
        file_number: str,
        motion_type_id: str,
        target_files: list[str],
        auto_analyze: bool = False,
        parent: QWidget | None = None,
    ):
        super().__init__(spec, parent=parent)
        self._case_path = case_path
        self._file_number = file_number
        self._files = list(target_files)
        self._worker = None
        self._last_settings: dict = {}
        self._finishing_worker = None
        self._analysis_worker = None
        self._finishing_analysis_worker = None

        self.settings_page = GenerateMotionSettingsPage(
            case_path, file_number, motion_type_id, target_files
        )
        self.status_page = StatusPage()
        self.output_page = GenerateMotionOutputPage()

        self.addWidget(self.settings_page)
        self.addWidget(self.status_page)
        self.addWidget(self.output_page)
        self.settings_page.run_requested.connect(self._on_run)
        if auto_analyze:
            self._start_analysis()

    @property
    def spec(self):
        return self._spec

    @property
    def files(self) -> list[str]:
        return list(self._files)

    def _start_analysis(self) -> None:
        if self._analysis_worker is not None or self._finishing_analysis_worker is not None:
            return
        self.status_page.reset()
        self.status_page.on_status("Analyzing target documents...")
        self.status_page.progress_bar.setRange(0, 0)
        self.setCurrentIndex(TASK_PAGE_STATUS)
        worker = GenerateMotionAnalysisWorker(
            settings={
                "motion_type_id": self.settings_page.current_motion_type_id(),
                "target_files": list(self.settings_page.target_files),
            },
            parent=None,
        )
        worker.progress.connect(self.status_page.on_status)
        worker.finished_analysis.connect(self._on_analysis_finished)
        worker.finished.connect(lambda w=worker: self._on_analysis_thread_finished(w))
        worker.finished.connect(worker.deleteLater)
        self._analysis_worker = worker
        worker.start()

    def _on_analysis_finished(self, success: bool, payload: object) -> None:
        if self.sender() is not None and self.sender() is not self._analysis_worker:
            return
        if self.sender() is self._analysis_worker:
            self._finishing_analysis_worker = self._analysis_worker
            self._analysis_worker = None
        elif self._analysis_worker is not None:
            self._analysis_worker = None

        if not success:
            self.status_page.on_status(f"FAILED: {payload}")
            return

        data = payload if isinstance(payload, dict) else {}
        metadata = data.get("metadata")
        if not isinstance(metadata, MotionMetadata):
            metadata = MotionMetadata.from_dict(metadata if isinstance(metadata, dict) else {})
        outline = [
            node if isinstance(node, OutlineNode) else OutlineNode.from_dict(node)
            for node in data.get("outline", []) or []
            if isinstance(node, (OutlineNode, dict))
        ]
        self.settings_page.set_metadata(metadata)
        self.settings_page.set_outline(outline)
        self.settings_page.setCurrentIndex(SETTINGS_PAGE_REVIEW)
        self.setCurrentIndex(TASK_PAGE_SETTINGS)

    def _on_analysis_thread_finished(self, worker) -> None:
        if self._analysis_worker is worker:
            self._analysis_worker = None
        if self._finishing_analysis_worker is worker:
            self._finishing_analysis_worker = None

    def _on_run(self, settings: dict) -> None:
        if self._analysis_worker is not None or self._finishing_analysis_worker is not None:
            self.status_page.on_status("Analysis is still running.")
            self.setCurrentIndex(TASK_PAGE_STATUS)
            return
        if self._worker is not None or self._finishing_worker is not None:
            self.status_page.on_status("Motion draft is already running.")
            self.setCurrentIndex(TASK_PAGE_STATUS)
            return
        self.status_page.reset()
        self.status_page.on_status("Drafting motion...")
        self.status_page.progress_bar.setRange(0, 0)
        self.setCurrentIndex(TASK_PAGE_STATUS)
        self._last_settings = dict(settings or {})
        worker = GenerateMotionWorker(
            case_path=self._case_path,
            file_number=self._file_number,
            settings=self._last_settings,
            parent=None,
        )
        worker.progress.connect(self.status_page.on_status)
        worker.finished_result.connect(self._on_worker_finished)
        worker.finished.connect(lambda w=worker: self._on_worker_thread_finished(w))
        worker.finished.connect(worker.deleteLater)
        self._worker = worker
        worker.start()

    def _on_worker_finished(self, success: bool, payload: object) -> None:
        from datetime import datetime

        if self.sender() is not None and self.sender() is not self._worker:
            return
        if self.sender() is self._worker:
            self._finishing_worker = self._worker
            self._worker = None
        elif self._worker is not None:
            self._worker = None
        if not success:
            self.status_page.on_status(f"FAILED: {payload}")
            return

        draft = payload if isinstance(payload, DraftDocument) else DraftDocument()
        self.output_page.show_result(draft)
        self.setCurrentIndex(TASK_PAGE_OUTPUT)
        self.task_completed.emit(
            {
                "task_id": self._spec.task_id,
                "title": self._spec.title,
                "files": list(self._files),
                "settings": self.settings_page.to_dict(),
                "output_path": draft.preview_path,
                "completed_at": datetime.now().isoformat(timespec="seconds"),
            }
        )

    def _on_worker_thread_finished(self, worker) -> None:
        if self._worker is worker:
            self._worker = None
        if self._finishing_worker is worker:
            self._finishing_worker = None

    def closeEvent(self, event) -> None:
        for worker in (
            self._analysis_worker,
            self._finishing_analysis_worker,
            self._worker,
            self._finishing_worker,
        ):
            if worker is not None and worker.isRunning():
                QMessageBox.information(
                    self,
                    "Task running",
                    "The motion draft is still running. Wait for it to finish before closing this tab.",
                )
                event.ignore()
                return
        super().closeEvent(event)


def _pick_motion_type(parent) -> str | None:
    """Small dialog to choose the motion type. Returns a type_id or None."""
    from PySide6.QtWidgets import QInputDialog

    labels = [MOTION_TYPE_CONFIGS[t].display_name for t in _TYPE_ORDER]
    choice, ok = QInputDialog.getItem(
        parent, "Motion type", "Which motion do you want to generate?", labels, 0, False
    )
    if not ok or not choice:
        return None
    for type_id in _TYPE_ORDER:
        if MOTION_TYPE_CONFIGS[type_id].display_name == choice:
            return type_id
    return "generic"


def build_generate_motion_tab(spec, case_path: str, file_number: str, parent: QWidget | None):
    motion_type_id = _pick_motion_type(parent)
    if motion_type_id is None:
        return None

    target_files = ContextFilesDialog.get_files(
        parent,
        title="Select target document(s) for the motion",
        start_dir=case_path,
        file_filter="Documents (*.pdf *.docx *.txt *.msg);;All files (*.*)",
    )
    target_files = [p for p in (target_files or []) if is_supported_context_file(p)]
    if not target_files:
        QMessageBox.warning(parent, "No documents", "Select at least one target document.")
        return None

    return GenerateMotionTaskTab(
        spec=spec,
        case_path=case_path,
        file_number=file_number,
        motion_type_id=motion_type_id,
        target_files=list(target_files),
        auto_analyze=True,
        parent=parent,
    )
