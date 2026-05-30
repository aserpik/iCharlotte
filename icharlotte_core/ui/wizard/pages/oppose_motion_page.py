"""Custom Wizard task page for drafting oppositions to motions."""

from __future__ import annotations

import html
import os
import re
import shutil

from PySide6.QtCore import QThread, QUrl, Qt, Signal
from PySide6.QtGui import QDesktopServices
from PySide6.QtWidgets import (
    QDialog,
    QFileDialog,
    QHBoxLayout,
    QLabel,
    QLineEdit,
    QMessageBox,
    QPlainTextEdit,
    QPushButton,
    QStackedWidget,
    QTextBrowser,
    QTreeWidget,
    QTreeWidgetItem,
    QVBoxLayout,
    QWidget,
)

from icharlotte_core.discovery.assembler import DiscoveryAssembler
from icharlotte_core.opposition.argument_research import research_arguments
from icharlotte_core.opposition.assembler import assemble_opposition_preview
from icharlotte_core.opposition.citation_parser import extract_citations
from icharlotte_core.opposition.drafter import draft_memorandum
from icharlotte_core.opposition.extraction import (
    extract_context_bundle,
    extract_document_text,
    is_supported_context_file,
    is_supported_motion_file,
)
from icharlotte_core.opposition.models import DraftDocument, MotionMetadata, OutlineNode
from icharlotte_core.opposition.motion_analyzer import analyze_motion, generate_outline
from icharlotte_core.opposition.outline import selected_section_plan
from icharlotte_core.opposition.style_examples import (
    StyleExampleRegistry,
    extract_exemplar_text,
)
from icharlotte_core.opposition.verifier import (
    build_local_opposition_verifier,
    build_opposition_verifier,
    enrich_with_pool_signals,
    pool_membership_check,
)
import os as _os_corpus
from icharlotte_core.config import CASELAW_DATA_DIR


def _corpus_paths() -> tuple[str, str]:
    return (_os_corpus.path.join(CASELAW_DATA_DIR, "corpus.db"),
            _os_corpus.path.join(CASELAW_DATA_DIR, "vectors.f16"))


def _corpus_available() -> bool:
    db, _vec = _corpus_paths()
    return _os_corpus.path.exists(db)


def _corpus_embedder():
    from icharlotte_core.legal_research.local_corpus.embedder import OnnxEmbedder
    return OnnxEmbedder()


def _make_local_corpus():
    if not _corpus_available():
        return None
    from icharlotte_core.legal_research.local_corpus.corpus import LocalCaseCorpus
    db, vec = _corpus_paths()
    return LocalCaseCorpus(db_path=db, vectors_path=vec, embedder=_corpus_embedder())
from icharlotte_core.ui.wizard.pages.status_page import StatusPage
from icharlotte_core.ui.context_files_dialog import ContextFilesDialog
from icharlotte_core.word_validator import validate_opposition_docx


SETTINGS_PAGE_CONFIRM = 0
SETTINGS_PAGE_OUTLINE = 1
TASK_PAGE_SETTINGS = 0
TASK_PAGE_STATUS = 1
TASK_PAGE_OUTPUT = 2


class OpposeMotionSettingsPage(QStackedWidget):
    """Confirmation and editable outline screens for the opposition workflow."""

    run_requested = Signal(dict)

    def __init__(
        self,
        case_root: str,
        file_number: str,
        motion_file: str,
        context_files: list[str],
        parent: QWidget | None = None,
    ):
        super().__init__(parent)
        self.case_root = case_root
        self.file_number = file_number
        self.motion_file = motion_file
        self.context_files = list(context_files)
        self.metadata = MotionMetadata()
        self.outline: list[OutlineNode] = []
        self._build_confirm_page()
        self._build_outline_page()

    def _build_confirm_page(self) -> None:
        page = QWidget()
        layout = QVBoxLayout(page)
        layout.setContentsMargins(24, 24, 24, 24)
        layout.setSpacing(10)

        self.motion_label = QLabel()
        layout.addWidget(self.motion_label)
        self._refresh_motion_label()

        self.motion_type_edit = QLineEdit()
        self.motion_type_edit.setPlaceholderText("Motion type")
        layout.addWidget(self.motion_type_edit)

        self.relief_edit = QLineEdit()
        self.relief_edit.setPlaceholderText("Relief requested")
        layout.addWidget(self.relief_edit)

        self.arguments_edit = QPlainTextEdit()
        self.arguments_edit.setPlaceholderText("Principal arguments, one per line")
        layout.addWidget(self.arguments_edit, 1)

        row = QHBoxLayout()
        row.addStretch()
        self.continue_btn = QPushButton("Review Outline")
        self.continue_btn.clicked.connect(self._on_continue_to_outline)
        row.addWidget(self.continue_btn)
        layout.addLayout(row)
        self.addWidget(page)

    def _build_outline_page(self) -> None:
        page = QWidget()
        layout = QVBoxLayout(page)
        layout.setContentsMargins(24, 24, 24, 24)
        layout.setSpacing(10)

        layout.addWidget(QLabel("Opposition Outline"))
        self.outline_tree = QTreeWidget()
        self.outline_tree.setHeaderLabels(["Include / Heading"])
        layout.addWidget(self.outline_tree, 1)

        row = QHBoxLayout()
        self.add_heading_btn = QPushButton("Add Heading")
        self.add_heading_btn.clicked.connect(self._add_heading)
        row.addWidget(self.add_heading_btn)
        row.addStretch()
        self.generate_btn = QPushButton("Generate Draft")
        self.generate_btn.clicked.connect(self._emit_run_requested)
        row.addWidget(self.generate_btn)
        layout.addLayout(row)
        self.addWidget(page)

    def set_metadata(self, metadata: MotionMetadata) -> None:
        self.metadata = metadata
        self.motion_type_edit.setText(metadata.motion_type)
        self.relief_edit.setText(metadata.relief_requested)
        self.arguments_edit.setPlainText("\n".join(metadata.principal_arguments))

    def current_metadata(self) -> MotionMetadata:
        metadata = MotionMetadata.from_dict(self.metadata.to_dict())
        metadata.motion_type = self.motion_type_edit.text().strip()
        metadata.relief_requested = self.relief_edit.text().strip()
        metadata.principal_arguments = [
            line.strip()
            for line in self.arguments_edit.toPlainText().splitlines()
            if line.strip()
        ]
        return metadata

    def can_continue_to_outline(self) -> bool:
        return self.current_metadata().required_missing() == []

    def set_outline(self, outline: list[OutlineNode]) -> None:
        self.outline = list(outline)
        self.outline_tree.clear()
        for node in self.outline:
            self.outline_tree.addTopLevelItem(self._item_from_node(node))
        self.outline_tree.expandAll()

    def to_dict(self) -> dict:
        return {
            "motion_file": self.motion_file,
            "context_files": list(self.context_files),
            "metadata": self.current_metadata().to_dict(),
            "outline": [node.to_dict() for node in self._outline_from_tree()],
        }

    def from_dict(self, data: dict | None) -> None:
        if not isinstance(data, dict):
            data = {}
        self.motion_file = data.get("motion_file", self.motion_file)
        self.context_files = list(data.get("context_files", self.context_files))
        self._refresh_motion_label()
        self.set_metadata(MotionMetadata.from_dict(data.get("metadata")))
        self.set_outline(
            [
                OutlineNode.from_dict(item)
                for item in data.get("outline", [])
                if isinstance(item, dict)
            ]
        )

    def _on_continue_to_outline(self) -> None:
        if not self.can_continue_to_outline():
            QMessageBox.warning(
                self,
                "Missing required fields",
                "Motion type, relief requested, and principal arguments are required.",
            )
            return
        self.setCurrentIndex(SETTINGS_PAGE_OUTLINE)

    def _refresh_motion_label(self) -> None:
        motion_name = os.path.basename(self.motion_file) or "(no motion selected)"
        self.motion_label.setText(f"Motion: {motion_name}")

    def _emit_run_requested(self) -> None:
        self.run_requested.emit(self.to_dict())

    def _add_heading(self) -> None:
        item = QTreeWidgetItem(["New heading"])
        item.setFlags(
            item.flags()
            | Qt.ItemFlag.ItemIsEditable
            | Qt.ItemFlag.ItemIsUserCheckable
        )
        item.setCheckState(0, Qt.CheckState.Checked)
        item.setData(0, Qt.ItemDataRole.UserRole, "")
        self.outline_tree.addTopLevelItem(item)
        self.outline_tree.editItem(item, 0)

    def _item_from_node(self, node: OutlineNode) -> QTreeWidgetItem:
        item = QTreeWidgetItem([node.text])
        item.setData(0, Qt.ItemDataRole.UserRole, node.id)
        item.setFlags(
            item.flags()
            | Qt.ItemFlag.ItemIsEditable
            | Qt.ItemFlag.ItemIsUserCheckable
        )
        item.setCheckState(
            0,
            Qt.CheckState.Checked if node.selected else Qt.CheckState.Unchecked,
        )
        for child in node.children:
            item.addChild(self._item_from_node(child))
        return item

    def _outline_from_tree(self) -> list[OutlineNode]:
        return [
            self._node_from_item(self.outline_tree.topLevelItem(index))
            for index in range(self.outline_tree.topLevelItemCount())
        ]

    def _node_from_item(self, item: QTreeWidgetItem) -> OutlineNode:
        return OutlineNode(
            id=item.data(0, Qt.ItemDataRole.UserRole) or "",
            text=item.text(0).strip(),
            selected=item.checkState(0) == Qt.CheckState.Checked,
            children=[
                self._node_from_item(item.child(index))
                for index in range(item.childCount())
            ],
        )


# ── DEV-ONLY: lightweight worker for the temporary Re-verify button. ──
# Re-runs just the citation extraction + verification step on a body
# that's already in the output page (e.g. after a restart, when the
# previous .docx has been reloaded). Avoids re-running the full draft
# pipeline. Remove this class along with the Re-verify button when no
# longer needed.
class _ReverifyWorker(QThread):
    finished_result = Signal(bool, object)  # (success, list[CitationVerification] | error str)

    def __init__(self, body_text: str, parent=None):
        super().__init__(parent)
        self.body_text = body_text

    def run(self) -> None:
        try:
            from icharlotte_core.llm_config import call_llm

            citations = extract_citations(self.body_text)
            if not citations:
                self.finished_result.emit(True, [])
                return

            token = os.environ.get("COURTLISTENER_API_TOKEN", "").strip()

            def llm(system_prompt, user_prompt):
                return call_llm(
                    user_prompt,
                    system_prompt,
                    task_type="general",
                    agent_id="agent_oppose_motion",
                ) or ""

            verifier = build_opposition_verifier(
                courtlistener_token=token,
                llm_callback=llm,
                max_workers=4,
            )
            results = verifier.verify_all(citations)
            self.finished_result.emit(True, results)
        except Exception as exc:
            self.finished_result.emit(False, str(exc))
# ── END DEV-ONLY ──


class OpposeMotionOutputPage(QWidget):
    rerun_requested = Signal()
    edit_settings_requested = Signal()

    def __init__(self, parent: QWidget | None = None):
        super().__init__(parent)
        self.draft = DraftDocument()
        outer = QVBoxLayout(self)
        outer.setContentsMargins(24, 24, 24, 24)
        outer.setSpacing(12)

        self.summary_banner = QLabel("")
        self.summary_banner.setTextFormat(Qt.TextFormat.RichText)
        self.summary_banner.setWordWrap(True)
        self.summary_banner.setVisible(False)
        outer.addWidget(self.summary_banner)

        layout = QHBoxLayout()
        layout.setSpacing(12)
        self.editor = QTextBrowser()
        self.editor.setOpenLinks(False)
        self.editor.setOpenExternalLinks(False)
        self.editor.anchorClicked.connect(self._on_anchor_clicked)
        layout.addWidget(self.editor, 2)

        self.detail_panel = CitationDetailPanel()
        layout.addWidget(self.detail_panel, 1)
        outer.addLayout(layout, 1)

        row = QHBoxLayout()
        row.addStretch()
        # ── DEV-ONLY: re-verify button. Remove when no longer needed. ──
        self.reverify_btn = QPushButton("Re-verify Citations (DEV)")
        self.reverify_btn.setToolTip(
            "Temporary developer button. Re-extracts citations from the "
            "current body and re-runs the verifier so you can test "
            "parser / verifier changes without re-running the full "
            "draft pipeline."
        )
        self.reverify_btn.setStyleSheet("QPushButton { color: #b06000; }")
        self.reverify_btn.clicked.connect(self._on_reverify_clicked)
        row.addWidget(self.reverify_btn)
        # ── END DEV-ONLY ──
        self.save_btn = QPushButton("Save")
        self.save_btn.clicked.connect(self.save_as)
        row.addWidget(self.save_btn)
        outer.addLayout(row)

        # Worker handle for the dev-only re-verify path.
        self._reverify_worker: _ReverifyWorker | None = None

    def show_result(self, draft: DraftDocument) -> None:
        self.draft = draft
        self.editor.setHtml(_render_draft_html(draft))
        self._refresh_summary_banner()
        if draft.citations:
            self.show_citation(0)
        else:
            self.detail_panel.clear(
                "No citations were detected in this opposition. If California "
                "case-law research returned no results, the draft was written "
                "without case citations. Review the brief for any factual or "
                "statutory support that may need strengthening."
            )

    def _refresh_summary_banner(self) -> None:
        if not self.draft.citations:
            self.summary_banner.setVisible(False)
            return
        counts: dict[str, int] = {}
        for cv in self.draft.citations:
            verdict = (cv.verdict or "UNVERIFIED").upper()
            counts[verdict] = counts.get(verdict, 0) + 1
        total = sum(counts.values())
        red = counts.get("NOT_SUPPORTED", 0) + counts.get("NOT_FOUND", 0)
        parts = [
            f"<b>Verification:</b> {total} citation(s) checked &mdash; ",
            f"\U0001F7E2 {counts.get('SUPPORTED', 0)} supported, ",
            f"\U0001F7E1 {counts.get('PARTIAL', 0)} partial, ",
            f"\U0001F534 {red} flagged, ",
            f"⚪ {counts.get('UNVERIFIED', 0)} unverified.",
        ]
        warning = ""
        if red > 0:
            warning = (
                f"<br><span style='color:#c5221f;'>⚠ {red} citation(s) don't "
                "support what the brief claims. Review the red-flagged cites "
                "before filing.</span>"
            )
        self.summary_banner.setText("".join(parts) + warning)
        self.summary_banner.setVisible(True)

    def _on_anchor_clicked(self, url: QUrl) -> None:
        scheme = url.scheme()
        if scheme == "citation":
            try:
                index = int(url.path().lstrip("/") or url.host() or "0")
            except (TypeError, ValueError):
                return
            # Click updates the right-side detail panel in place. No popup.
            self.show_citation(index)
            return
        QDesktopServices.openUrl(url)

    @property
    def output_path(self) -> str:
        return self.draft.preview_path

    def load_output(self, output_path: str) -> None:
        body_text = ""
        if output_path and output_path.lower().endswith(".docx") and os.path.isfile(output_path):
            try:
                from docx import Document

                doc = Document(output_path)
                body_text = "\n".join(p.text for p in doc.paragraphs if p.text)
            except Exception:
                body_text = f"(Could not render generated document: {output_path})"
        self.show_result(
            DraftDocument(
                title=os.path.splitext(os.path.basename(output_path or ""))[0] or "Opposition",
                body_text=body_text,
                preview_path=output_path or "",
            )
        )

    @staticmethod
    def default_save_dir(preview_path: str) -> str:
        marker = os.path.join(".icharlotte", "wizard_previews")
        before_marker, _, _after_marker = preview_path.partition(marker)
        if before_marker:
            return os.path.dirname(os.path.normpath(before_marker))
        return os.path.dirname(preview_path)

    def show_citation(self, index: int) -> None:
        if index < 0 or index >= len(self.draft.citations):
            return
        self.detail_panel.set_citation(self.draft.citations[index])

    # ── DEV-ONLY: re-verify handlers. Remove when no longer needed. ──

    def _on_reverify_clicked(self) -> None:
        if self._reverify_worker is not None:
            return  # already running
        body_text = (self.draft.body_text or "").strip()
        if not body_text:
            QMessageBox.warning(
                self,
                "Nothing to verify",
                "No draft body is loaded. Re-open or re-run the task first.",
            )
            return
        self.reverify_btn.setEnabled(False)
        self.reverify_btn.setText("Re-verifying… (DEV)")
        self.summary_banner.setText(
            "<b>Re-verifying citations…</b> (this runs the verifier on the "
            "current body without re-drafting)"
        )
        self.summary_banner.setVisible(True)
        worker = _ReverifyWorker(body_text=body_text, parent=None)
        worker.finished_result.connect(self._on_reverify_finished)
        worker.finished.connect(worker.deleteLater)
        self._reverify_worker = worker
        worker.start()

    def _on_reverify_finished(self, success: bool, payload: object) -> None:
        self._reverify_worker = None
        self.reverify_btn.setEnabled(True)
        self.reverify_btn.setText("Re-verify Citations (DEV)")
        if not success:
            QMessageBox.critical(
                self,
                "Re-verify failed",
                f"Re-verification failed:\n\n{payload}",
            )
            self._refresh_summary_banner()
            return
        citations = payload if isinstance(payload, list) else []
        self.draft.citations = citations
        # Re-render with updated colors + summary + panel.
        self.editor.setHtml(_render_draft_html(self.draft))
        self._refresh_summary_banner()
        if citations:
            self.show_citation(0)
        else:
            self.detail_panel.clear(
                "Re-verify found no citations in the current body text."
            )

    # ── END DEV-ONLY ──

    def save_as(self) -> None:
        if not self.draft.preview_path:
            QMessageBox.warning(
                self,
                "No preview",
                "No generated opposition preview is available.",
            )
            return

        red = sum(
            1 for cv in self.draft.citations
            if (cv.verdict or "").upper() in {"NOT_SUPPORTED", "NOT_FOUND"}
        )
        if red > 0:
            choice = QMessageBox.question(
                self,
                "Citations flagged",
                (
                    f"This opposition has {red} citation(s) flagged as "
                    "NOT_SUPPORTED or NOT_FOUND.\n\nSave anyway?"
                ),
                QMessageBox.StandardButton.Save | QMessageBox.StandardButton.Cancel,
                QMessageBox.StandardButton.Cancel,
            )
            if choice != QMessageBox.StandardButton.Save:
                return

        suggested = os.path.join(
            self.default_save_dir(self.draft.preview_path),
            f"{self.draft.title or 'Opposition Memorandum'}.docx",
        )
        target, _ = QFileDialog.getSaveFileName(
            self,
            "Save Opposition Memorandum",
            suggested,
            "Word Documents (*.docx);;All files (*.*)",
        )
        if not target:
            return
        if not target.lower().endswith(".docx"):
            target += ".docx"
        if os.path.abspath(target) == os.path.abspath(self.draft.preview_path):
            QMessageBox.warning(
                self,
                "Choose another location",
                "Select a location outside the internal preview file.",
            )
            return
        try:
            shutil.copyfile(self.draft.preview_path, target)
        except Exception as exc:
            QMessageBox.critical(self, "Save failed", f"Could not save file:\n{exc}")
            return
        QMessageBox.information(self, "Saved", f"Saved:\n{target}")


_HORIZONTAL_RULE_RE = re.compile(r"^[\*\-_]{3,}\s*$")
_MD_HEADING_RE = re.compile(r"^(#{1,6})\s+(.+?)\s*$")
_MD_ITALIC_RE = re.compile(r"\*{1,2}([^\*\n]+?)\*{1,2}")


_VERDICT_COLORS = {
    "SUPPORTED": "#1e8e3e",       # green
    "PARTIAL": "#f9ab00",          # yellow
    "NOT_SUPPORTED": "#c5221f",    # red
    "NOT_FOUND": "#c5221f",        # red (same as NOT_SUPPORTED)
    "UNVERIFIED": "#80868b",       # gray
}


def _color_for_verdict(verdict: str) -> str:
    return _VERDICT_COLORS.get((verdict or "").upper(), "#1a5dbf")  # default blue


def _render_draft_html(draft: DraftDocument) -> str:
    """Render the draft body into HTML with clickable citation anchors."""
    citation_spans = _build_citation_index(draft)
    body_html_lines: list[str] = []

    for raw_line in (draft.body_text or "").splitlines():
        stripped = raw_line.strip()
        if not stripped:
            body_html_lines.append("<p>&nbsp;</p>")
            continue
        if _HORIZONTAL_RULE_RE.match(stripped):
            continue
        heading_match = _MD_HEADING_RE.match(stripped)
        if heading_match:
            level = min(max(len(heading_match.group(1)), 2), 4)
            content = _format_inline_html(heading_match.group(2), citation_spans)
            body_html_lines.append(f"<h{level}>{content}</h{level}>")
            continue
        content = _format_inline_html(stripped, citation_spans)
        body_html_lines.append(f"<p>{content}</p>")

    title = html.escape(draft.title or "Opposition Memorandum")
    body = "\n".join(body_html_lines)
    return (
        "<html><body style=\"font-family:'Times New Roman',serif; font-size:13pt;\">"
        f"<h1 style=\"text-align:center;\">{title}</h1>{body}</body></html>"
    )


def _build_citation_index(draft: DraftDocument) -> list[tuple[str, int, str]]:
    """Return [(citation_text, draft_citation_index, verdict), ...] sorted by length desc.

    Sorting by length avoids partial-match collisions (e.g. "62 Cal. 4th 1081"
    must be wrapped before a shorter substring matches inside it).
    """
    spans: list[tuple[str, int, str]] = []
    for index, citation in enumerate(draft.citations or []):
        text = (citation.citation_text or "").strip()
        if text:
            spans.append((text, index, citation.verdict or ""))
    spans.sort(key=lambda triple: len(triple[0]), reverse=True)
    return spans


def _format_inline_html(line: str, citation_spans: list[tuple[str, int, str]]) -> str:
    """Escape, italicize *case names*, and wrap citation texts as clickable anchors."""
    # Mark italic spans with placeholders WITHOUT pre-escaping the inner
    # content. Pre-escaping plus the outer html.escape() double-escapes
    # entities (e.g. "&" → "&amp;" → "&amp;amp;"), which surfaces as
    # literal "&amp;" in the rendered HTML.
    italicized = _MD_ITALIC_RE.sub(
        lambda match: f"\x00ITA{match.group(1)}\x00ITAEND",
        line,
    )
    escaped = html.escape(italicized)
    # Replace closing placeholder before opening — otherwise "\x00ITA" matches
    # the prefix of "\x00ITAEND" and leaves "END" stranded in the output.
    escaped = escaped.replace("\x00ITAEND", "</i>").replace("\x00ITA", "<i>")
    escaped = escaped.replace(
        html.escape("[no case authority retrieved for this point]"),
        "<span style=\"color:#80868b;\">[no case authority retrieved for this point]</span>",
    )
    for citation_text, index, verdict in citation_spans:
        if not citation_text:
            continue
        # Build a regex that tolerates intervening <i> or </i> tags inserted
        # by the italics pass above (case names are typically wrapped in
        # *...* in the body but the citation_text comes without asterisks).
        escaped_cite = html.escape(citation_text)
        tolerant_pattern = r"(?:<i>|</i>)*".join(
            re.escape(ch) for ch in escaped_cite
        )
        color = _color_for_verdict(verdict)
        anchor = (
            f"<a href=\"citation:{index}\" "
            f"style=\"color:{color}; text-decoration:underline;\">"
            f"{escaped_cite}</a>"
        )
        escaped = re.sub(tolerant_pattern, anchor, escaped, count=0)
    return escaped


# ─── Verdict-aware citation rendering helpers (used by panel + dialog) ─────

_VERDICT_HEADER_COLORS = {
    "SUPPORTED": "#1e8e3e",
    "PARTIAL": "#f9ab00",
    "NOT_SUPPORTED": "#c5221f",
    "NOT_FOUND": "#c5221f",
    "UNVERIFIED": "#80868b",
}
_VERDICT_LABELS = {
    "SUPPORTED": "SUPPORTED",
    "PARTIAL": "PARTIAL",
    "NOT_SUPPORTED": "NOT SUPPORTED",
    "NOT_FOUND": "CITATION NOT FOUND",
    "UNVERIFIED": "UNVERIFIED",
}


def _citation_header_html(citation, verdict: str) -> str:
    color = _VERDICT_HEADER_COLORS.get(verdict, "#80868b")
    label = _VERDICT_LABELS.get(verdict, verdict or "UNVERIFIED")
    title = html.escape(citation.case_name or citation.citation_text or "Citation")
    return (
        f"<div style='border-left: 6px solid {color}; padding-left: 8px;'>"
        f"<div style='color:{color}; font-weight:bold; font-size:14pt;'>{html.escape(label)}</div>"
        f"<div style='font-size:11pt;'>{title}</div>"
        "</div>"
    )


def _citation_body_html(citation, verdict: str) -> str:
    parts: list[str] = []

    prop = (citation.proposition or "").strip()
    if prop:
        parts.append(
            f"<p><b>Brief's proposition:</b><br><i>{html.escape(prop)}</i></p>"
        )

    if verdict == "SUPPORTED":
        if citation.evidence:
            parts.append(
                f"<p><b>Verified holding:</b><br>{html.escape(citation.evidence)}</p>"
            )
        if citation.note:
            parts.append(f"<p><b>Verifier note:</b> {html.escape(citation.note)}</p>")
        if getattr(citation, "citation_count", None) is not None:
            year = (getattr(citation, "latest_citing_year", "") or "").strip()
            tail = f", most recently {html.escape(year)}" if year else ""
            parts.append(
                f"<p style='color:#80868b;'><b>Good-law hint:</b> cited by "
                f"{citation.citation_count} case(s){tail}. (Not a Shepard's check — "
                "confirm the case is still good law.)</p>"
            )

    elif verdict == "PARTIAL":
        if citation.evidence:
            parts.append(
                f"<p><b>Relevant passage:</b><br>{html.escape(citation.evidence)}</p>"
            )
        if citation.note:
            parts.append(
                f"<p><b>Why partial:</b> {html.escape(citation.note)}</p>"
            )
        if getattr(citation, "citation_count", None) is not None:
            year = (getattr(citation, "latest_citing_year", "") or "").strip()
            tail = f", most recently {html.escape(year)}" if year else ""
            parts.append(
                f"<p style='color:#80868b;'><b>Good-law hint:</b> cited by "
                f"{citation.citation_count} case(s){tail}. (Not a Shepard's check — "
                "confirm the case is still good law.)</p>"
            )

    elif verdict == "NOT_SUPPORTED":
        parts.append(
            "<p><b>⚠ The authority does NOT hold what the brief claims.</b></p>"
        )
        if citation.evidence:
            parts.append(
                f"<p><b>What it actually holds:</b><br>{html.escape(citation.evidence)}</p>"
            )
        if citation.note:
            parts.append(f"<p><b>Verifier note:</b> {html.escape(citation.note)}</p>")

    elif verdict == "NOT_FOUND":
        parts.append(
            "<p><b>⚠ This citation was not found in the authoritative source.</b></p>"
            "<p>Likely causes:</p>"
            "<ul>"
            "<li>The case or statute was invented by the LLM</li>"
            "<li>The citation is mis-typed</li>"
            "<li>The case is unpublished or pre-1900</li>"
            "</ul>"
        )
        if citation.note:
            parts.append(f"<p><b>Verifier note:</b> {html.escape(citation.note)}</p>")

    else:  # UNVERIFIED
        # The note (when present) explains the specific reason -- cluster
        # found but no opinion text, LLM call failed, prompt missing,
        # parse failure, etc. Show the note directly. Only fall back to
        # the generic "doesn't cover" message when no note is available
        # (which corresponds to genuinely out-of-scope citation kinds
        # like federal cases or local court rules).
        if citation.note:
            parts.append(
                f"<p><b>⚠ Could not verify:</b> {html.escape(citation.note)}</p>"
            )
        else:
            parts.append(
                "<p>The verifier doesn't cover this source in v1. Verify manually.</p>"
            )

    return "\n".join(parts)


def _run_find_replacement(parent_widget, citation):
    """Run the find-replacement LLM call and pop a result summary.

    Shared by the panel button and the (legacy) dialog button. Shows
    a QMessageBox with candidates or with the appropriate error. Returns
    nothing; intended as a fire-and-forget UI action.
    """
    from icharlotte_core.llm_config import call_llm
    from icharlotte_core.opposition.verifier import (
        build_opposition_verifier,
        find_replacement_candidates,
    )

    token = os.environ.get("COURTLISTENER_API_TOKEN", "").strip()
    if not token:
        QMessageBox.warning(
            parent_widget,
            "Missing API token",
            "COURTLISTENER_API_TOKEN is not set; replacement search is unavailable.",
        )
        return

    def llm(system_prompt, user_prompt):
        return call_llm(
            user_prompt,
            system_prompt,
            task_type="general",
            agent_id="agent_oppose_motion",
        ) or ""

    verifier = build_opposition_verifier(courtlistener_token=token, llm_callback=llm)
    candidates = find_replacement_candidates(
        failed_citation=citation,
        verifier=verifier,
        llm_callback=llm,
    )

    if not candidates:
        QMessageBox.information(
            parent_widget,
            "No replacements found",
            "No supported replacement candidates were found. Try editing the proposition or searching manually.",
        )
        return
    lines = [f"{c.citation_text} — {c.verdict}\n  {c.note}" for c in candidates]
    QMessageBox.information(parent_widget, "Replacement candidates", "\n\n".join(lines))


class CitationDetailPanel(QWidget):
    """Right-side panel showing verdict-aware citation detail.

    Replaces the popup dialog: anchor clicks in the draft body update this
    panel in-place. Exposes ``set_citation()`` and ``clear()``.
    """

    def __init__(self, parent: QWidget | None = None) -> None:
        super().__init__(parent)
        self.citation = None

        layout = QVBoxLayout(self)
        layout.setContentsMargins(8, 8, 8, 8)
        layout.setSpacing(8)

        self.header_label = QLabel("")
        self.header_label.setTextFormat(Qt.TextFormat.RichText)
        self.header_label.setWordWrap(True)
        layout.addWidget(self.header_label)

        self.body_browser = QTextBrowser()
        self.body_browser.setOpenExternalLinks(False)
        self.body_browser.setStyleSheet("QTextBrowser { background: transparent; }")
        layout.addWidget(self.body_browser, 1)

        button_row = QHBoxLayout()
        button_row.addStretch()
        self.find_btn = QPushButton("Find replacement case")
        self.find_btn.clicked.connect(self._on_find_replacement)
        self.find_btn.setVisible(False)
        button_row.addWidget(self.find_btn)
        self.open_btn = QPushButton("Open source")
        self.open_btn.clicked.connect(self._open_opinion_url)
        self.open_btn.setVisible(False)
        button_row.addWidget(self.open_btn)
        layout.addLayout(button_row)

        # Initial empty state.
        self.clear("Select a citation in the draft to see verification details.")

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    def set_citation(self, citation) -> None:
        self.citation = citation
        verdict = (citation.verdict or "UNVERIFIED").upper()
        self.header_label.setText(_citation_header_html(citation, verdict))
        self.body_browser.setHtml(_citation_body_html(citation, verdict))

        # Toggle action buttons based on what makes sense for this citation.
        if citation.opinion_url:
            label = "Open in CourtListener" if citation.kind == "case" else "Open in leginfo"
            self.open_btn.setText(label)
            self.open_btn.setVisible(True)
        else:
            self.open_btn.setVisible(False)

        self.find_btn.setVisible(verdict in {"NOT_SUPPORTED", "NOT_FOUND"})

    def clear(self, message: str = "") -> None:
        self.citation = None
        self.header_label.setText("")
        self.body_browser.setHtml(
            f"<p style='color:#80868b;'>{html.escape(message)}</p>" if message else ""
        )
        self.open_btn.setVisible(False)
        self.find_btn.setVisible(False)

    @property
    def body_html(self) -> str:
        """Current rendered body HTML (for tests)."""
        return self.body_browser.toHtml()

    # ------------------------------------------------------------------
    # Handlers
    # ------------------------------------------------------------------

    def _open_opinion_url(self) -> None:
        if self.citation and self.citation.opinion_url:
            QDesktopServices.openUrl(QUrl(self.citation.opinion_url))

    def _on_find_replacement(self) -> None:
        if self.citation is not None:
            _run_find_replacement(self, self.citation)


class CitationDetailDialog(QDialog):
    """Modal popup version (legacy). The output page no longer auto-opens
    this; it's retained as an opt-in surface for programmatic use and for
    backward compatibility with existing tests.
    """

    _VERDICT_HEADER_COLORS = _VERDICT_HEADER_COLORS
    _VERDICT_LABELS = _VERDICT_LABELS

    def __init__(self, citation, parent: QWidget | None = None):
        super().__init__(parent)
        self.citation = citation
        self.setWindowTitle(citation.case_name or citation.citation_text or "Citation")
        self.resize(720, 540)

        verdict = (citation.verdict or "UNVERIFIED").upper()

        layout = QVBoxLayout(self)
        layout.setContentsMargins(16, 16, 16, 16)
        layout.setSpacing(10)

        self.header = QLabel(_citation_header_html(citation, verdict))
        self.header.setTextFormat(Qt.TextFormat.RichText)
        self.header.setWordWrap(True)
        layout.addWidget(self.header)

        self.body_html = _citation_body_html(citation, verdict)
        self.body_label = QLabel(self.body_html)
        self.body_label.setTextFormat(Qt.TextFormat.RichText)
        self.body_label.setWordWrap(True)
        self.body_label.setOpenExternalLinks(False)
        layout.addWidget(self.body_label, 1)

        button_row = QHBoxLayout()
        button_row.addStretch()
        if citation.opinion_url:
            label = "Open in CourtListener" if citation.kind == "case" else "Open in leginfo"
            open_btn = QPushButton(label)
            open_btn.clicked.connect(self._open_opinion_url)
            button_row.addWidget(open_btn)
        close_btn = QPushButton("Close")
        close_btn.clicked.connect(self.accept)
        button_row.addWidget(close_btn)

        if verdict in {"NOT_SUPPORTED", "NOT_FOUND"}:
            self.find_btn = QPushButton("Find replacement case")
            self.find_btn.clicked.connect(self._on_find_replacement)
            button_row.insertWidget(button_row.count() - 1, self.find_btn)

        layout.addLayout(button_row)

    def _open_opinion_url(self) -> None:
        if self.citation.opinion_url:
            QDesktopServices.openUrl(QUrl(self.citation.opinion_url))

    def _on_find_replacement(self) -> None:
        _run_find_replacement(self, self.citation)


class OpposeMotionAnalysisWorker(QThread):
    progress = Signal(str)
    finished_analysis = Signal(bool, object)

    def __init__(self, settings: dict, parent=None):
        super().__init__(parent)
        self.settings = dict(settings or {})

    def run(self) -> None:
        try:
            from icharlotte_core.llm_config import call_llm

            self.progress.emit("Extracting motion text...")
            motion_result = extract_document_text(self.settings.get("motion_file", ""))
            if not motion_result.success:
                message = motion_result.error or "Could not read motion."
                self.finished_analysis.emit(False, message)
                return

            self.progress.emit("Extracting context documents...")
            context_text, warnings = extract_context_bundle(
                self.settings.get("context_files", [])
            )
            for warning in warnings:
                self.progress.emit(f"WARNING: {warning}")

            def llm(system_prompt, user_prompt):
                return call_llm(
                    user_prompt,
                    system_prompt,
                    task_type="general",
                    agent_id="agent_chat",
                ) or ""

            self.progress.emit("Analyzing motion...")
            metadata = analyze_motion(motion_result.text, llm_callback=llm)
            missing = metadata.required_missing()
            if missing:
                self.finished_analysis.emit(
                    False,
                    "Could not automatically identify required motion fields: "
                    + ", ".join(missing)
                    + ". Confirm the motion has extractable text.",
                )
                return

            self.progress.emit("Generating opposition outline...")
            outline = generate_outline(
                metadata,
                context_text,
                llm_callback=llm,
            )
            if not outline:
                self.finished_analysis.emit(
                    False,
                    "Could not automatically generate an opposition outline.",
                )
                return

            self.finished_analysis.emit(
                True,
                {
                    "metadata": metadata,
                    "outline": outline,
                    "motion_text": motion_result.text,
                    "context_text": context_text,
                },
            )
        except Exception as exc:
            self.finished_analysis.emit(False, str(exc))


class OpposeMotionWorker(QThread):
    progress = Signal(str)
    finished_result = Signal(bool, object)

    def __init__(self, case_path: str, file_number: str, settings: dict, parent=None):
        super().__init__(parent)
        self.case_path = case_path
        self.file_number = file_number
        self.settings = dict(settings or {})

    def run(self) -> None:
        try:
            from icharlotte_core.llm_config import call_llm

            self.progress.emit("Extracting motion text...")
            motion_result = extract_document_text(self.settings.get("motion_file", ""))
            if not motion_result.success:
                message = motion_result.error or "Could not read motion."
                self.finished_result.emit(False, message)
                return

            self.progress.emit("Extracting context documents...")
            context_text, warnings = extract_context_bundle(
                self.settings.get("context_files", [])
            )
            for warning in warnings:
                self.progress.emit(f"WARNING: {warning}")

            metadata = MotionMetadata.from_dict(self.settings.get("metadata"))
            outline = [
                OutlineNode.from_dict(item)
                for item in self.settings.get("outline", [])
                if isinstance(item, dict)
            ]
            plan = selected_section_plan(outline)

            def make_llm(pass_name):
                def _llm(system_prompt, user_prompt):
                    return call_llm(
                        user_prompt, system_prompt, task_type="general",
                        agent_id="agent_oppose_motion", pass_name=pass_name,
                        pass_agent_id="agent_oppose_motion",
                    ) or ""
                return _llm

            def llm(system_prompt, user_prompt):
                return call_llm(
                    user_prompt,
                    system_prompt,
                    task_type="general",
                    agent_id="agent_oppose_motion",
                ) or ""

            # Load style exemplars matching this motion type. The registry lives
            # under the repo's Scripts/prompts/oppose_motion/ directory; the
            # registry file is created on first save by the workbench so it may
            # be missing here — StyleExampleRegistry.load() handles that.
            self.progress.emit("Loading matching style exemplars...")
            repo_root = os.path.dirname(
                os.path.dirname(
                    os.path.dirname(
                        os.path.dirname(os.path.dirname(__file__))
                    )
                )
            )
            registry_path = os.path.join(
                repo_root,
                "Scripts",
                "prompts",
                "oppose_motion",
                "style_examples.json",
            )
            registry = StyleExampleRegistry.load(registry_path)
            matches = registry.matches_for_motion_type(metadata.motion_type)
            cache_dir = os.path.join(
                os.path.dirname(registry_path), ".cache", "style_examples"
            )
            exemplar_texts: list[str] = []
            for m in matches:
                text = extract_exemplar_text(m.path, cache_dir=cache_dir)
                if text.strip():
                    exemplar_texts.append(text)
            if matches:
                self.progress.emit(f"  Using {len(exemplar_texts)} style exemplar(s).")
            else:
                self.progress.emit("  No matching style exemplars; using default voice.")

            # Retrieval-first grounding: research real California authority for
            # each principal argument before drafting, so the drafter cites only
            # from a verified pool.
            # Prefer the local CA corpus (offline, unlimited). Fall back to the
            # live CourtListener API only if the corpus has not been built yet.
            corpus = _make_local_corpus()
            token = os.environ.get("COURTLISTENER_API_TOKEN", "").strip()
            retrieved = []
            opinion_cache = os.path.join(
                os.path.dirname(registry_path), ".cache", "opinions"
            )
            if corpus is not None and metadata.principal_arguments:
                self.progress.emit(
                    f"Researching authorities locally ({len(metadata.principal_arguments)} arguments)..."
                )
                retrieved = research_arguments(
                    metadata.principal_arguments,
                    cl_client=corpus,
                    query_llm=make_llm("research_queries"),
                    rerank_llm=make_llm("rerank_select"),
                    max_workers=4,
                    on_progress=self.progress.emit,
                    cache_dir=opinion_cache,
                )
                self.progress.emit(f"Retrieved {len(retrieved)} grounded authorities.")
            elif corpus is None and token and metadata.principal_arguments:
                from icharlotte_core.legal_research.sources.courtlistener import CourtListenerClient
                self.progress.emit(
                    "Local corpus not built; falling back to CourtListener API "
                    f"({len(metadata.principal_arguments)} arguments)..."
                )
                retrieved = research_arguments(
                    metadata.principal_arguments,
                    cl_client=CourtListenerClient(token),
                    query_llm=make_llm("research_queries"),
                    rerank_llm=make_llm("rerank_select"),
                    max_workers=4,
                    on_progress=self.progress.emit,
                    cache_dir=opinion_cache,
                )
                self.progress.emit(f"Retrieved {len(retrieved)} grounded authorities.")
            else:
                self.progress.emit(
                    "WARNING: no local corpus and no COURTLISTENER_API_TOKEN; "
                    "drafting without grounded research."
                )

            self.progress.emit("Drafting opposition memorandum...")
            draft = draft_memorandum(
                metadata=metadata,
                section_plan=plan,
                motion_text=motion_result.text,
                context_text=context_text,
                style_exemplars=exemplar_texts,
                retrieved_authorities=retrieved,
                llm_callback=llm,
            )
            if not draft.body_text.strip():
                reason = (draft.rejection_reason or "unknown reason").strip()
                self.finished_result.emit(False, f"Drafting failed: {reason}")
                return

            # Parse citations from the drafted body.
            citations = extract_citations(draft.body_text)
            if not citations:
                self.progress.emit(
                    "WARNING: No citations detected in the drafted opposition."
                )
                draft.citations = []
            else:
                to_verify, off_pool = pool_membership_check(citations, retrieved)
                if off_pool:
                    self.progress.emit(
                        f"{len(off_pool)} citation(s) were not in the researched pool "
                        "(flagged NOT_FOUND)."
                    )
                if corpus is None and not token:
                    self.progress.emit(
                        "WARNING: no local corpus and no COURTLISTENER_API_TOKEN; "
                        "case citations cannot be verified."
                    )
                self.progress.emit(f"Verifying citations ({len(to_verify)} found)...")
                if corpus is not None:
                    verifier = build_local_opposition_verifier(
                        corpus=corpus,
                        llm_callback=llm,
                        max_workers=4,
                    )
                else:
                    verifier = build_opposition_verifier(
                        courtlistener_token=token,
                        llm_callback=llm,
                        max_workers=4,
                    )
                verified = verifier.verify_all(to_verify, on_progress=self.progress.emit)
                # Merge verified + off-pool, restored to body order.
                draft.citations = sorted(
                    list(verified) + list(off_pool),
                    key=lambda cv: cv.body_offset if cv.body_offset is not None else 0,
                )
                enrich_with_pool_signals(draft.citations, retrieved)
                verdict_counts: dict[str, int] = {}
                for cv in draft.citations:
                    verdict_counts[cv.verdict] = verdict_counts.get(cv.verdict, 0) + 1
                summary = ", ".join(
                    f"{v.lower()}: {n}" for v, n in sorted(verdict_counts.items())
                )
                self.progress.emit(f"Verification complete ({summary}).")

            preview_dir = os.path.join(
                self.case_path,
                "NOTES",
                "AI OUTPUT",
                ".icharlotte",
                "wizard_previews",
                "oppose_motion",
            )
            preview_path = os.path.join(preview_dir, "Opposition Preview.docx")
            caption_path = DiscoveryAssembler.find_caption_page(self.case_path) or ""
            assemble_opposition_preview(
                draft=draft,
                output_path=preview_path,
                caption_path=caption_path,
            )
            validation = validate_opposition_docx(preview_path)
            if validation.has_errors:
                self.finished_result.emit(
                    False,
                    "Word validation failed for opposition preview.",
                )
                return
            draft.preview_path = preview_path
            self.finished_result.emit(True, draft)
        except Exception as exc:
            self.finished_result.emit(False, str(exc))


class OpposeMotionTaskTab(QStackedWidget):
    task_completed = Signal(dict)

    def __init__(
        self,
        spec,
        case_path: str,
        file_number: str,
        motion_file: str,
        context_files: list[str],
        auto_analyze: bool = False,
        parent: QWidget | None = None,
    ):
        super().__init__(parent)
        self._spec = spec
        self._case_path = case_path
        self._file_number = file_number
        self._files = [motion_file] + list(context_files)
        self._worker = None
        self._last_settings: dict = {}
        self._finishing_worker = None
        self._analysis_worker = None
        self._finishing_analysis_worker = None

        self.settings_page = OpposeMotionSettingsPage(
            case_path,
            file_number,
            motion_file,
            context_files,
        )
        self.status_page = StatusPage()
        self.output_page = OpposeMotionOutputPage()

        self.addWidget(self.settings_page)
        self.addWidget(self.status_page)
        self.addWidget(self.output_page)
        self.settings_page.run_requested.connect(self._on_run)
        if auto_analyze:
            self._start_analysis()

    @property
    def spec(self):
        return self._spec

    @property
    def files(self) -> list[str]:
        return list(self._files)

    def _start_analysis(self) -> None:
        if self._analysis_worker is not None or self._finishing_analysis_worker is not None:
            return
        self.status_page.reset()
        self.status_page.on_status("Analyzing selected motion and context...")
        self.status_page.progress_bar.setRange(0, 0)
        self.setCurrentIndex(TASK_PAGE_STATUS)
        worker = OpposeMotionAnalysisWorker(
            settings={
                "motion_file": self.settings_page.motion_file,
                "context_files": list(self.settings_page.context_files),
            },
            parent=None,
        )
        worker.progress.connect(self.status_page.on_status)
        worker.finished_analysis.connect(self._on_analysis_finished)
        if hasattr(worker, "finished"):
            worker.finished.connect(lambda w=worker: self._on_analysis_thread_finished(w))
        if hasattr(worker, "deleteLater"):
            worker.finished.connect(worker.deleteLater)
        self._analysis_worker = worker
        worker.start()

    def _on_analysis_finished(self, success: bool, payload: object) -> None:
        if self.sender() is not None and self.sender() is not self._analysis_worker:
            return
        if self.sender() is self._analysis_worker:
            self._finishing_analysis_worker = self._analysis_worker
            self._analysis_worker = None
        elif self._analysis_worker is not None:
            self._analysis_worker = None

        if not success:
            self.status_page.on_status(f"FAILED: {payload}")
            return

        data = payload if isinstance(payload, dict) else {}
        metadata = data.get("metadata")
        if not isinstance(metadata, MotionMetadata):
            metadata = MotionMetadata.from_dict(metadata if isinstance(metadata, dict) else {})
        outline = [
            node if isinstance(node, OutlineNode) else OutlineNode.from_dict(node)
            for node in data.get("outline", []) or []
            if isinstance(node, (OutlineNode, dict))
        ]
        self.settings_page.set_metadata(metadata)
        self.settings_page.set_outline(outline)
        # Skip the metadata-confirmation screen and go straight to the outline.
        # Metadata is still set on the page (line above) so it flows through to
        # the drafter via current_metadata() when the user clicks Generate Draft.
        self.settings_page.setCurrentIndex(SETTINGS_PAGE_OUTLINE)
        self.setCurrentIndex(TASK_PAGE_SETTINGS)

    def _on_analysis_thread_finished(self, worker) -> None:
        if self._analysis_worker is worker:
            self._analysis_worker = None
        if self._finishing_analysis_worker is worker:
            self._finishing_analysis_worker = None

    def _on_run(self, settings: dict) -> None:
        if self._analysis_worker is not None or self._finishing_analysis_worker is not None:
            self.status_page.on_status("Motion analysis is still running.")
            self.setCurrentIndex(TASK_PAGE_STATUS)
            return
        if self._worker is not None or self._finishing_worker is not None:
            self.status_page.on_status("Opposition draft is already running.")
            self.setCurrentIndex(TASK_PAGE_STATUS)
            return
        self.status_page.reset()
        self.status_page.on_status("Drafting opposition memorandum...")
        self.status_page.progress_bar.setRange(0, 0)
        self.setCurrentIndex(TASK_PAGE_STATUS)
        self._last_settings = dict(settings or {})
        worker = OpposeMotionWorker(
            case_path=self._case_path,
            file_number=self._file_number,
            settings=self._last_settings,
            parent=None,
        )
        worker.progress.connect(self.status_page.on_status)
        worker.finished_result.connect(self._on_worker_finished)
        if hasattr(worker, "finished"):
            worker.finished.connect(lambda w=worker: self._on_worker_thread_finished(w))
        if hasattr(worker, "deleteLater"):
            worker.finished.connect(worker.deleteLater)
        self._worker = worker
        worker.start()

    def _on_worker_finished(self, success: bool, payload: object) -> None:
        from datetime import datetime

        if self.sender() is not None and self.sender() is not self._worker:
            return
        if self.sender() is self._worker:
            self._finishing_worker = self._worker
            self._worker = None
        elif self._worker is not None:
            self._worker = None
        if not success:
            self.status_page.on_status(f"FAILED: {payload}")
            return

        draft = payload if isinstance(payload, DraftDocument) else DraftDocument()
        self.output_page.show_result(draft)
        self.setCurrentIndex(TASK_PAGE_OUTPUT)
        self.task_completed.emit(
            {
                "task_id": self._spec.task_id,
                "title": self._spec.title,
                "files": list(self._files),
                "settings": self.settings_page.to_dict(),
                "output_path": draft.preview_path,
                "completed_at": datetime.now().isoformat(timespec="seconds"),
            }
        )

    def _on_worker_thread_finished(self, worker) -> None:
        if self._worker is worker:
            self._worker = None
        if self._finishing_worker is worker:
            self._finishing_worker = None

    def closeEvent(self, event) -> None:
        for worker in (
            self._analysis_worker,
            self._finishing_analysis_worker,
            self._worker,
            self._finishing_worker,
        ):
            if worker is not None and worker.isRunning():
                QMessageBox.information(
                    self,
                    "Task running",
                    "The opposition draft is still running. Wait for it to finish before closing this tab.",
                )
                event.ignore()
                return
        super().closeEvent(event)


def build_oppose_motion_tab(
    spec,
    case_path: str,
    file_number: str,
    parent: QWidget | None,
):
    motion_file, _ = QFileDialog.getOpenFileName(
        parent,
        "Select motion to oppose",
        case_path,
        "Motion files (*.pdf *.docx)",
    )
    if not motion_file:
        return None
    if not is_supported_motion_file(motion_file):
        QMessageBox.warning(
            parent,
            "Unsupported motion file",
            "Select a PDF or DOCX motion.",
        )
        return None

    context_files = ContextFilesDialog.get_files(
        parent,
        title="Select context document(s)",
        start_dir=os.path.dirname(motion_file) or case_path,
        file_filter="Context files (*.pdf *.docx *.txt *.msg);;All files (*.*)",
    )
    context_files = [
        path for path in (context_files or []) if is_supported_context_file(path)
    ]
    return OpposeMotionTaskTab(
        spec=spec,
        case_path=case_path,
        file_number=file_number,
        motion_file=motion_file,
        context_files=list(context_files),
        auto_analyze=True,
        parent=parent,
    )
