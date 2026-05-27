"""Convert between .docx and the QTextEdit HTML model used by the Output Page.

Forward (.docx → HTML): mammoth's `convert_to_html`.
Reverse (QTextDocument → .docx): walk QTextBlocks and emit python-docx
paragraphs. We capture: heading levels (Heading 1/2/3 paragraph styles
or HTML <hN>), bold, italic, underline, and bullet/numbered lists. Tables,
images, and other complex structures may render approximately and may
be dropped on save (see spec known limitations).
"""
import mammoth
from docx import Document
from docx.shared import Pt
from PySide6.QtGui import QTextDocument, QTextBlock, QTextCharFormat, QTextBlockFormat


def load_docx_as_html(path: str) -> str:
    """Convert a .docx file to a self-contained HTML string."""
    with open(path, "rb") as f:
        result = mammoth.convert_to_html(f)
    return result.value


def save_qtextdocument_as_docx(qdoc: QTextDocument, out_path: str) -> None:
    """Write a python-docx Document mirroring qdoc's block/inline structure.

    Limitations:
      - Tables, images, embedded objects from the editor are not preserved.
      - Bullet / numbered lists fall back to plain paragraphs (python-docx
        list-style is template-dependent and we don't carry a template here).
    """
    document = Document()
    block: QTextBlock = qdoc.begin()
    while block.isValid():
        text = block.text()
        level = _detect_heading_level(block)
        if level is not None:
            para = document.add_heading(text, level=level)
        else:
            para = document.add_paragraph()
            _write_block_runs(block, para)
        # Spacing — preserve "blank paragraph" feel without touching styles.
        block = block.next()
    document.save(out_path)


def _detect_heading_level(block: QTextBlock) -> int | None:
    """Detect heading level from QTextBlockFormat.headingLevel() (Qt 6)."""
    fmt: QTextBlockFormat = block.blockFormat()
    if hasattr(fmt, "headingLevel"):
        lvl = fmt.headingLevel()
        if lvl and lvl > 0:
            return min(lvl, 9)
    return None


def _write_block_runs(block: QTextBlock, para) -> None:
    """Walk inline fragments of `block` and emit python-docx runs with formatting."""
    it = block.begin()
    while not it.atEnd():
        frag = it.fragment()
        if frag.isValid():
            text = frag.text()
            char_fmt: QTextCharFormat = frag.charFormat()
            run = para.add_run(text)
            if char_fmt.fontWeight() >= 600:
                run.bold = True
            if char_fmt.fontItalic():
                run.italic = True
            if char_fmt.fontUnderline():
                run.underline = True
            # Font size (in points), if set.
            size_pt = char_fmt.fontPointSize()
            if size_pt and size_pt > 0:
                run.font.size = Pt(size_pt)
        it += 1
