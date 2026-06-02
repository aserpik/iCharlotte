"""Configuration form for the multi-analysis Med-Cron task.

Reads the Phase 1 session JSON, presents:
- a checkbox per curated analysis (Rewrite Chronology pre-checked)
- a "Custom analyses" panel with add/remove rows
- a "narrative missing" warning banner when applicable

On commit_user_config(), validates selection and writes user_config back
to the session, flipping phase to ``ready_to_run``.

Custom analyses persist globally: rows that the user fills in are saved
to ``custom_analyses_store`` on every commit and pre-populated when the
form is next opened. Each row has an "Include in this run" checkbox so
the user can skip a saved analysis once without deleting it from the
global store.
"""

import os
from pathlib import Path
from typing import Optional

from PySide6.QtCore import QSettings, Qt, Signal
from PySide6.QtWidgets import (
    QCheckBox, QFileDialog, QFrame, QHBoxLayout, QLabel, QLineEdit,
    QPlainTextEdit, QPushButton, QScrollArea, QSplitter, QVBoxLayout, QWidget,
)

from icharlotte_core.med_chron import custom_analyses_store, session_manager


# ---- Context-document helpers ----

SUPPORTED_CONTEXT_EXTS = (".pdf", ".docx", ".txt")
MAX_CONTEXT_CHARS = 120_000  # mirrors the Phase 2 cap; only used by callers that want it


def sniff_text_layer(path: str) -> tuple[bool, str]:
    """Return (has_text, reason).

    Cheap, attach-time check: does this file appear to have extractable
    text without needing OCR? Used by the UI to warn the user.

    - .txt: any non-whitespace in the first 4 KB.
    - .docx: any non-empty paragraph OR table cell in the first ~50 paragraphs
      / ~5 tables. (python-docx's `doc.paragraphs` silently skips tables —
      see CLAUDE.md gotcha — so we also sample table cells.)
    - .pdf: at least 200 chars of text extractable from pages 0..2.
    - any error / unknown extension: (False, "...").
    """
    ext = os.path.splitext(path)[1].lower()
    try:
        if ext == ".txt":
            with open(path, "r", encoding="utf-8", errors="ignore") as f:
                sample = f.read(4096)
            if sample.strip():
                return (True, "")
            return (False, "no text content")
        if ext == ".docx":
            from docx import Document
            doc = Document(path)
            text = "".join(p.text for p in doc.paragraphs[:50])
            if not text.strip():
                # python-docx skips tables — sample a few cells too so we
                # don't falsely flag table-based docs as "no text layer".
                for table in doc.tables[:5]:
                    for row in table.rows[:20]:
                        for cell in row.cells:
                            text += cell.text
                            if text.strip():
                                break
                        if text.strip():
                            break
                    if text.strip():
                        break
            if text.strip():
                return (True, "")
            return (False, "no extractable text in paragraphs or tables")
        if ext == ".pdf":
            from pypdf import PdfReader
            reader = PdfReader(path)
            n = min(3, len(reader.pages))
            text = ""
            for i in range(n):
                try:
                    text += reader.pages[i].extract_text() or ""
                except Exception:
                    pass
            if len(text.strip()) > 200:
                return (True, "")
            return (False, "too little extractable text — OCR may be required")
        return (False, f"unsupported extension {ext}")
    except FileNotFoundError:
        return (False, "file not found")
    except Exception as e:
        return (False, f"could not read file: {e}")


_ANALYSES_SPLITTER_KEY = "med_chron_form/analyses_splitter_sizes"


class ContextDropTextEdit(QPlainTextEdit):
    """QPlainTextEdit that accepts drops of .pdf / .docx / .txt file URLs.

    Emits ``files_dropped(list[str])`` when one or more supported files are
    dropped on the widget. Non-file drops fall through to the base class so
    text editing still works.
    """

    files_dropped = Signal(list)

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setAcceptDrops(True)

    def _supported_paths(self, mime) -> list[str]:
        """Return paths for the drop if EVERY URL is a supported local file.

        All-or-nothing: a single non-local URL or unsupported extension
        rejects the whole drop. Mixed drops (e.g., a PDF + a PNG dragged
        together) are ambiguous, so the safer default is to reject and let
        the user retry with a clean selection.
        """
        if not mime.hasUrls():
            return []
        paths: list[str] = []
        for url in mime.urls():
            if not url.isLocalFile():
                return []
            p = os.path.normpath(url.toLocalFile())
            ext = os.path.splitext(p)[1].lower()
            if ext not in SUPPORTED_CONTEXT_EXTS:
                return []
            paths.append(p)
        return paths

    def dragEnterEvent(self, event):
        if self._supported_paths(event.mimeData()):
            event.acceptProposedAction()
        else:
            event.ignore()

    def dragMoveEvent(self, event):
        if self._supported_paths(event.mimeData()):
            event.acceptProposedAction()
        else:
            super().dragMoveEvent(event)

    def dropEvent(self, event):
        paths = self._supported_paths(event.mimeData())
        if paths:
            self.files_dropped.emit(paths)
            event.acceptProposedAction()
            return
        super().dropEvent(event)


class CustomAnalysisRow(QWidget):
    """One row in the custom-analyses list: include-checkbox + label + instruction + remove btn + context-doc chip strip."""

    def __init__(self, parent: QWidget, on_remove, included: bool = True):
        super().__init__(parent)
        self._on_remove = on_remove
        self._context_files: list[str] = []
        self._add_ctx_btn = None  # built lazily in _render_chip_strip

        layout = QVBoxLayout(self)
        layout.setContentsMargins(4, 4, 4, 4)
        layout.setSpacing(4)

        top = QHBoxLayout()
        self.include_cb = QCheckBox()
        self.include_cb.setChecked(included)
        self.include_cb.setToolTip("Include this analysis in the current run")
        top.addWidget(self.include_cb)
        top.addWidget(QLabel("Label:"))
        self.label_edit = QLineEdit()
        self.label_edit.setPlaceholderText("Short name (e.g. 'Left-knee mentions')")
        top.addWidget(self.label_edit, 1)
        self.remove_btn = QPushButton("−")
        self.remove_btn.setFixedSize(24, 24)
        self.remove_btn.setStyleSheet("QPushButton { color: #c62828; font-weight: bold; }")
        self.remove_btn.setToolTip("Remove this analysis (also deletes from global save)")
        self.remove_btn.clicked.connect(self._handle_remove)
        top.addWidget(self.remove_btn)
        layout.addLayout(top)

        layout.addWidget(QLabel("Request:"))
        self.instruction_edit = ContextDropTextEdit()
        self.instruction_edit.setPlaceholderText("Describe the analysis…")
        self.instruction_edit.setFixedHeight(60)
        self.instruction_edit.files_dropped.connect(self.add_context_files)
        layout.addWidget(self.instruction_edit)

        # Chip strip for attached context documents.
        self._chip_strip_container = QWidget()
        self._chip_strip_layout = QHBoxLayout(self._chip_strip_container)
        self._chip_strip_layout.setContentsMargins(0, 0, 0, 0)
        self._chip_strip_layout.setSpacing(4)
        layout.addWidget(self._chip_strip_container)

        # Warning label for files that look like they need OCR.
        self._context_warning_label = QLabel("")
        self._context_warning_label.setStyleSheet(
            "color: #856404; background-color: #FFF3CD; "
            "padding: 4px; border-radius: 3px; font-size: 11px;"
        )
        self._context_warning_label.setWordWrap(True)
        self._context_warning_label.setVisible(False)
        layout.addWidget(self._context_warning_label)

        self._render_chip_strip()

        self.setStyleSheet(
            "CustomAnalysisRow { border: 1px solid #ddd; border-radius: 4px; }"
        )

    def _handle_remove(self):
        self._on_remove(self)

    def label(self) -> str:
        return self.label_edit.text().strip()

    def instruction(self) -> str:
        return self.instruction_edit.toPlainText().strip()

    def is_included(self) -> bool:
        return self.include_cb.isChecked()

    def is_empty(self) -> bool:
        return not self.label() and not self.instruction()

    def context_files(self) -> list[str]:
        return list(self._context_files)

    def add_context_files(self, paths: list[str]) -> None:
        changed = False
        for p in paths:
            if p in self._context_files:
                continue
            self._context_files.append(p)
            changed = True
        if changed:
            self._render_chip_strip()
            self._refresh_context_warning()

    def _on_add_context_clicked(self) -> None:
        paths, _ = QFileDialog.getOpenFileNames(
            self,
            "Select context document(s)",
            "",
            "Context documents (*.pdf *.docx *.txt)",
        )
        if paths:
            self.add_context_files(list(paths))

    def _remove_context_file(self, path: str) -> None:
        if path in self._context_files:
            self._context_files.remove(path)
            self._render_chip_strip()
            self._refresh_context_warning()

    def _render_chip_strip(self) -> None:
        # Remove all widgets currently in the strip.
        while self._chip_strip_layout.count() > 0:
            item = self._chip_strip_layout.takeAt(0)
            w = item.widget()
            if w is not None and w is not self._add_ctx_btn:
                w.setParent(None)
                w.deleteLater()
        # One chip per attached file.
        for p in self._context_files:
            chip = self._build_chip(p)
            self._chip_strip_layout.addWidget(chip)
        # Trailing "+ Add context" button.
        if self._add_ctx_btn is None:
            self._add_ctx_btn = QPushButton("+ Add context")
            self._add_ctx_btn.setStyleSheet(
                "QPushButton { font-size: 11px; padding: 2px 8px; }"
            )
            self._add_ctx_btn.clicked.connect(self._on_add_context_clicked)
        # Re-parent the button each render so re-adding it works after takeAt cleared the layout.
        self._add_ctx_btn.setParent(self._chip_strip_container)
        self._chip_strip_layout.addWidget(self._add_ctx_btn)
        self._chip_strip_layout.addStretch()

    def _build_chip(self, path: str) -> QWidget:
        chip = QFrame()
        chip.setObjectName("ctx_chip")
        chip.setStyleSheet(
            "QFrame#ctx_chip { background: #E3F2FD; border: 1px solid #90CAF9; "
            "border-radius: 10px; padding: 2px 6px; }"
        )
        lay = QHBoxLayout(chip)
        lay.setContentsMargins(4, 0, 4, 0)
        lay.setSpacing(4)
        icon = QLabel("\U0001f4ce")
        icon.setStyleSheet("font-size: 11px;")
        lay.addWidget(icon)
        name = QLabel(os.path.basename(path))
        name.setObjectName("chip_filename")
        name.setStyleSheet("font-size: 11px; color: #0D47A1;")
        name.setToolTip(path)
        lay.addWidget(name)
        x_btn = QPushButton("✕")
        x_btn.setFixedSize(16, 16)
        x_btn.setStyleSheet(
            "QPushButton { font-size: 10px; color: #555; border: none; }"
            "QPushButton:hover { color: #c62828; }"
        )
        x_btn.setToolTip(f"Remove {os.path.basename(path)}")
        x_btn.clicked.connect(lambda _=False, p=path: self._remove_context_file(p))
        lay.addWidget(x_btn)
        return chip

    def _refresh_context_warning(self) -> None:
        bad: list[str] = []
        for p in self._context_files:
            has_text, _ = sniff_text_layer(p)
            if not has_text:
                bad.append(os.path.basename(p))
        if bad:
            self._context_warning_label.setText(
                "⚠ Likely needs OCR at run time: " + ", ".join(bad)
            )
            self._context_warning_label.setVisible(True)
        else:
            self._context_warning_label.setVisible(False)


class MedChronConfigForm(QWidget):
    """Pickable list of catalog analyses + custom analysis rows."""

    def __init__(self, session_path, parent: Optional[QWidget] = None):
        super().__init__(parent)
        self.session_path = Path(session_path)
        self._session = session_manager.read_session(self.session_path)
        self._settings = QSettings("iCharlotte", "iCharlotte")

        root = QVBoxLayout(self)
        root.setContentsMargins(8, 8, 8, 8)
        root.setSpacing(8)

        header = QLabel(
            f"<b>Analyses to run on {self._session.get('provider_name', '')}</b>"
        )
        root.addWidget(header)

        # Narrative-missing warning banner
        self.narrative_missing_banner = QLabel(
            "⚠ Narrative text not found in this document — "
            "Rewrite Chronology will be skipped."
        )
        self.narrative_missing_banner.setStyleSheet(
            "background-color: #FFF3CD; color: #856404; "
            "padding: 6px; border-radius: 4px;"
        )
        self.narrative_missing_banner.setVisible(
            bool(self._session.get("narrative_missing"))
        )
        root.addWidget(self.narrative_missing_banner)

        # Vertical splitter: Curated (top) ↔ Custom (bottom).
        self._analyses_splitter = QSplitter(Qt.Orientation.Vertical)
        self._analyses_splitter.setChildrenCollapsible(False)
        self._analyses_splitter.setHandleWidth(6)
        self._analyses_splitter.setStyleSheet(
            "QSplitter::handle { background: #e0e0e0; }"
            " QSplitter::handle:hover { background: #b0b0b0; }"
            " QSplitter::handle:pressed { background: #909090; }"
        )

        # --- Curated catalog pane ---
        curated_section = QWidget()
        curated_layout = QVBoxLayout(curated_section)
        curated_layout.setContentsMargins(0, 0, 0, 0)
        curated_layout.setSpacing(4)
        curated_layout.addWidget(QLabel("<b>Curated analyses:</b>"))

        curated_inner = QWidget()
        curated_inner_layout = QVBoxLayout(curated_inner)
        curated_inner_layout.setContentsMargins(0, 0, 0, 0)
        curated_inner_layout.setSpacing(0)

        self.catalog_checkboxes: dict[str, QCheckBox] = {}
        for entry in self._session.get("catalog", []):
            row = QWidget()
            row_layout = QVBoxLayout(row)
            row_layout.setContentsMargins(16, 2, 0, 2)
            row_layout.setSpacing(0)
            cb = QCheckBox(entry.get("title", entry.get("id", "")))
            cb.setChecked(bool(entry.get("default_selected")))
            self.catalog_checkboxes[entry["id"]] = cb
            row_layout.addWidget(cb)
            desc = QLabel(entry.get("description", ""))
            desc.setStyleSheet("color: #666; font-size: 11px; padding-left: 22px;")
            desc.setWordWrap(True)
            row_layout.addWidget(desc)
            curated_inner_layout.addWidget(row)
        curated_inner_layout.addStretch()

        curated_scroll = QScrollArea()
        curated_scroll.setWidget(curated_inner)
        curated_scroll.setWidgetResizable(True)
        curated_scroll.setFrameShape(QScrollArea.Shape.NoFrame)
        curated_layout.addWidget(curated_scroll, 1)
        self._analyses_splitter.addWidget(curated_section)

        # --- Custom analyses pane ---
        custom_section = QWidget()
        custom_layout = QVBoxLayout(custom_section)
        custom_layout.setContentsMargins(0, 0, 0, 0)
        custom_layout.setSpacing(4)

        custom_header = QHBoxLayout()
        custom_header.addWidget(QLabel("<b>Custom analyses:</b>"))
        custom_header.addStretch()
        custom_layout.addLayout(custom_header)

        # Scrollable container for custom rows.
        self._custom_container = QWidget()
        self._custom_container_layout = QVBoxLayout(self._custom_container)
        self._custom_container_layout.setContentsMargins(0, 0, 0, 0)
        self._custom_container_layout.setSpacing(4)
        self._custom_container_layout.addStretch()  # pin rows to top

        custom_scroll = QScrollArea()
        custom_scroll.setWidget(self._custom_container)
        custom_scroll.setWidgetResizable(True)
        custom_scroll.setFrameShape(QScrollArea.Shape.NoFrame)
        custom_layout.addWidget(custom_scroll, 1)

        self.custom_rows: list[CustomAnalysisRow] = []

        # Pre-populate previously-saved custom analyses so the user does
        # not have to re-type recurring requests. These load UNCHECKED — the
        # user must opt each one into this run, so saved customs never run
        # silently just because they exist in the global store.
        for saved in custom_analyses_store.load():
            row = self.add_custom_row(included=False)
            row.label_edit.setText(saved.get("label", ""))
            row.instruction_edit.setPlainText(saved.get("instruction", ""))

        add_btn = QPushButton("+ Add custom analysis")
        add_btn.clicked.connect(self.add_custom_row)
        custom_layout.addWidget(add_btn)

        self._analyses_splitter.addWidget(custom_section)
        self._analyses_splitter.setStretchFactor(0, 1)
        self._analyses_splitter.setStretchFactor(1, 1)
        self._restore_analyses_splitter_sizes()
        self._analyses_splitter.splitterMoved.connect(
            self._save_analyses_splitter_sizes
        )
        root.addWidget(self._analyses_splitter, 1)

        # Inline validation error label.
        self._error_label = QLabel("")
        self._error_label.setStyleSheet("color: #c62828; font-style: italic;")
        self._error_label.setVisible(False)
        root.addWidget(self._error_label)

    def _restore_analyses_splitter_sizes(self) -> None:
        saved = self._settings.value(_ANALYSES_SPLITTER_KEY)
        if saved:
            try:
                sizes = [int(x) for x in saved]
                if len(sizes) == 2 and all(s > 0 for s in sizes):
                    self._analyses_splitter.setSizes(sizes)
                    return
            except (TypeError, ValueError):
                pass
        self._analyses_splitter.setSizes([240, 200])

    def _save_analyses_splitter_sizes(self, *_args) -> None:
        self._settings.setValue(
            _ANALYSES_SPLITTER_KEY, self._analyses_splitter.sizes()
        )

    def add_custom_row(self, included: bool = True) -> "CustomAnalysisRow":
        row = CustomAnalysisRow(
            self._custom_container, self._remove_custom_row, included=included
        )
        # Insert above the stretch at the end.
        idx = self._custom_container_layout.count() - 1
        self._custom_container_layout.insertWidget(idx, row)
        self.custom_rows.append(row)
        return row

    def _remove_custom_row(self, row: "CustomAnalysisRow") -> None:
        if row in self.custom_rows:
            self.custom_rows.remove(row)
        self._custom_container_layout.removeWidget(row)
        row.setParent(None)
        row.deleteLater()

    def _selected_catalog_ids(self) -> list[str]:
        return [cid for cid, cb in self.catalog_checkboxes.items() if cb.isChecked()]

    def _validated_custom_rows(self) -> tuple[list[dict], list[dict], str]:
        """Return ``(persisted_rows, run_rows, error_msg)``.

        - ``persisted_rows`` — ``{label, instruction}`` only, for the global
          custom_analyses_store. Context files are intentionally excluded so
          they do not leak between sessions / cases.
        - ``run_rows`` — ``{label, instruction, context_files}`` for the
          session JSON; only rows whose include checkbox is checked.
        - ``error_msg`` — non-empty if a row is partially filled.
        """
        persisted: list[dict] = []
        run_rows: list[dict] = []
        for r in self.custom_rows:
            if r.is_empty():
                continue
            lbl, instr = r.label(), r.instruction()
            if not lbl or not instr:
                return [], [], (
                    "Custom analyses need both a label and an instruction. "
                    "Fill in (or remove) the partially-completed row."
                )
            persisted.append({"label": lbl, "instruction": instr})
            if r.is_included():
                run_rows.append({
                    "label": lbl,
                    "instruction": instr,
                    "context_files": r.context_files(),
                })
        return persisted, run_rows, ""

    def commit_user_config(self) -> bool:
        """Validate, persist saved analyses globally, and write user_config.

        Returns True on success.
        """
        self._error_label.setVisible(False)

        selected = self._selected_catalog_ids()
        persisted_custom, run_custom, err = self._validated_custom_rows()
        if err:
            self._error_label.setText(err)
            self._error_label.setVisible(True)
            return False
        if not selected and not run_custom:
            self._error_label.setText(
                "Select at least one analysis, or add a custom analysis."
            )
            self._error_label.setVisible(True)
            return False

        # Persist label + instruction only to the global store. Context
        # files are intentionally NOT persisted — they belong to this
        # session / case only.
        try:
            custom_analyses_store.save(persisted_custom)
        except OSError:
            # Persisting globally is best-effort; the run can still proceed.
            pass

        session_manager.update_user_config(
            self.session_path,
            {
                "selected_catalog_ids": selected,
                "custom_analyses": run_custom,
            },
        )
        return True
