"""
Process-safe DOCX Writer for iCharlotte

Provides file locking to prevent concurrent write conflicts when multiple
agents try to write to the same output file (e.g., AI_OUTPUT.docx).

Uses file-based locking that works across separate Python processes.
"""

import os
import time
import random
import datetime
import logging
from typing import Optional, Callable
from contextlib import contextmanager

try:
    from docx import Document
    from docx.shared import Pt, Inches
except ImportError:
    Document = None
    Pt = None
    Inches = None

# Lock timeout and retry settings
LOCK_TIMEOUT = 120  # Maximum seconds to wait for lock
LOCK_RETRY_INTERVAL = 0.5  # Base interval between retries
MAX_WRITE_RETRIES = 10  # Maximum file write attempts


class DocxWriteError(Exception):
    """Raised when docx write fails after all retries."""
    pass


class LockTimeoutError(Exception):
    """Raised when unable to acquire lock within timeout."""
    pass


def get_docx_lock(output_path: str, timeout: float = LOCK_TIMEOUT):
    """
    Get a file lock context manager for a docx file.

    Use this to wrap custom save logic that needs process-level locking.

    Example:
        with get_docx_lock(output_path):
            doc = Document(output_path)
            # ... modify doc ...
            doc.save(output_path)

    Args:
        output_path: Path to the docx file to lock
        timeout: Maximum seconds to wait for lock

    Returns:
        Context manager that acquires/releases the lock
    """
    lock_path = output_path + ".lock"
    return file_lock(lock_path, timeout)


@contextmanager
def file_lock(lock_path: str, timeout: float = LOCK_TIMEOUT):
    """
    Cross-process file lock using a .lock file.

    Uses exclusive file creation to ensure only one process can hold the lock.
    Falls back to checking lock file age if stale (process crashed).

    Args:
        lock_path: Path to the lock file (e.g., "AI_OUTPUT.docx.lock")
        timeout: Maximum seconds to wait for lock

    Yields:
        None when lock is acquired

    Raises:
        LockTimeoutError: If unable to acquire lock within timeout
    """
    start_time = time.time()
    lock_acquired = False
    lock_fd = None

    # Stale lock threshold (if lock file older than this, assume dead process)
    STALE_THRESHOLD = 300  # 5 minutes

    while time.time() - start_time < timeout:
        try:
            # Try to create lock file exclusively
            # os.O_CREAT | os.O_EXCL ensures atomic create-if-not-exists
            lock_fd = os.open(lock_path, os.O_CREAT | os.O_EXCL | os.O_WRONLY)

            # Write our PID to the lock file for debugging
            pid_bytes = f"{os.getpid()}\n{time.time()}".encode('utf-8')
            os.write(lock_fd, pid_bytes)

            lock_acquired = True
            break

        except FileExistsError:
            # Lock file exists - check if it's stale
            try:
                lock_age = time.time() - os.path.getmtime(lock_path)
                # Read holder PID for diagnostics
                try:
                    with open(lock_path, 'r') as lf:
                        holder_info = lf.read().strip()
                except Exception:
                    holder_info = "unknown"
                wait_elapsed = time.time() - start_time
                logging.info(f"Lock contention on {os.path.basename(lock_path)}: held by [{holder_info}], age={lock_age:.1f}s, waiting={wait_elapsed:.1f}s, my_pid={os.getpid()}")

                if lock_age > STALE_THRESHOLD:
                    # Lock is stale, remove it and retry
                    logging.warning(f"Removing stale lock {lock_path} (age={lock_age:.1f}s)")
                    try:
                        os.remove(lock_path)
                        continue
                    except OSError:
                        pass  # Another process might have removed it
            except OSError:
                pass  # File might have been removed between checks

            # Wait with jitter before retrying
            jitter = random.uniform(0, LOCK_RETRY_INTERVAL)
            time.sleep(LOCK_RETRY_INTERVAL + jitter)

        except OSError as e:
            # Other OS errors (permission, etc.)
            logging.warning(f"Lock file error: {e}")
            time.sleep(LOCK_RETRY_INTERVAL)

    if not lock_acquired:
        raise LockTimeoutError(f"Could not acquire lock on {lock_path} within {timeout}s")

    try:
        yield
    finally:
        # Release lock
        if lock_fd is not None:
            try:
                os.close(lock_fd)
            except OSError:
                pass

        # Remove lock file
        try:
            os.remove(lock_path)
        except OSError:
            pass


def safe_append_to_docx(
    output_path: str,
    content: str,
    title: str,
    format_func: Optional[Callable] = None,
    logger: Optional[Callable] = None
) -> str:
    """
    Safely append content to a DOCX file with process-level locking.

    If the file is locked by another process, waits until available.
    If the file can't be opened (e.g., open in Word), tries versioned names.

    Args:
        output_path: Target DOCX file path
        content: Text content to add (can include markdown)
        title: Title/header for this section
        format_func: Optional function(doc, content) to format content
        logger: Optional logging function

    Returns:
        Actual path where content was saved (may differ if versioning used)

    Raises:
        DocxWriteError: If write fails after all retries
    """
    if Document is None:
        raise DocxWriteError("python-docx not installed")

    def log(msg, level="info"):
        if logger:
            if callable(logger):
                logger(msg)
            elif hasattr(logger, level):
                getattr(logger, level)(msg)
        else:
            try:
                print(msg, flush=True)
            except OSError:
                pass  # stdout pipe broken

    # Ensure output directory exists
    output_dir = os.path.dirname(output_path)
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir, exist_ok=True)

    # Lock file path
    lock_path = output_path + ".lock"

    base_name, ext = os.path.splitext(output_path)
    version = 0
    current_path = output_path

    try:
        # Acquire process-level lock
        with file_lock(lock_path):
            log(f"Acquired write lock for: {os.path.basename(output_path)}")

            # Now try to write, with versioning fallback for open files
            for attempt in range(MAX_WRITE_RETRIES):
                try:
                    # Open or create document
                    if os.path.exists(current_path):
                        try:
                            doc = Document(current_path)
                            doc.add_page_break()
                            log(f"Appending to existing file: {current_path}")
                        except Exception as e:
                            # File might be corrupted or locked by another app
                            raise PermissionError(f"Cannot open file: {e}")
                    else:
                        doc = Document()
                        log(f"Creating new file: {current_path}")

                    # Apply default styles
                    _apply_default_styles(doc)

                    # Add title
                    p = doc.add_paragraph()
                    run = p.add_run(title)
                    run.bold = True
                    run.underline = True
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)

                    # Add timestamp
                    doc.add_paragraph(
                        f"Generated on: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
                    )

                    # Format and add content
                    if format_func:
                        format_func(doc, content)
                    else:
                        _add_markdown_to_doc(doc, content)

                    # Save document
                    doc.save(current_path)
                    log(f"Saved to: {current_path}")

                    return current_path

                except PermissionError as e:
                    # File locked by another app (e.g., Word) - try versioned name
                    version += 1
                    current_path = f"{base_name} v.{version}{ext}"
                    log(f"File locked, trying: {current_path}")

                    if attempt >= MAX_WRITE_RETRIES - 1:
                        raise DocxWriteError(
                            f"Could not write to {output_path} or any version after "
                            f"{MAX_WRITE_RETRIES} attempts"
                        )

                except Exception as e:
                    # Other errors - retry with backoff
                    if attempt >= MAX_WRITE_RETRIES - 1:
                        raise DocxWriteError(f"Write failed: {e}")

                    wait_time = (2 ** attempt) * 0.5 + random.uniform(0, 0.5)
                    log(f"Write error: {e}. Retrying in {wait_time:.1f}s...")
                    time.sleep(wait_time)

    except LockTimeoutError as e:
        raise DocxWriteError(f"Lock timeout: {e}")

    return current_path


def _apply_default_styles(doc):
    """Apply default Times New Roman 12pt styles to document."""
    if Pt is None:
        return

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.0

    for i in range(1, 10):
        if f'Heading {i}' in doc.styles:
            h_style = doc.styles[f'Heading {i}']
            h_style.font.name = 'Times New Roman'
            h_style.font.size = Pt(12)
            h_style.paragraph_format.line_spacing = 1.0


def _add_markdown_to_doc(doc, content: str):
    """Parse basic markdown and add to document."""
    if Pt is None or Inches is None:
        # Fallback: just add as plain text
        doc.add_paragraph(content)
        return

    lines = content.split('\n')
    active_paragraph = None

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if not stripped:
            i += 1
            continue

        # Markdown table block — convert to a real Word table.
        consumed, header, rows = try_consume_markdown_table(lines, i)
        if consumed:
            add_markdown_table_to_doc(doc, header, rows)
            i += consumed
            active_paragraph = None
            continue

        # Headings — strip emphasis markers; the heading itself is bolded.
        if stripped.startswith('#'):
            text = stripped.lstrip('#').strip().replace('**', '').replace('__', '')
            if not text.endswith('.'):
                text += "."
            active_paragraph = doc.add_paragraph()
            run = active_paragraph.add_run(text + " ")
            run.bold = True
            i += 1
            continue

        # List items
        if stripped.startswith('* ') or stripped.startswith('- '):
            text = stripped[2:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.first_line_indent = Inches(-0.25)

            tab_run = p.add_run("\t")
            tab_run.font.name = 'Times New Roman'
            tab_run.font.size = Pt(12)

            render_inline_markdown(p, text)

            active_paragraph = None
            i += 1
            continue

        # Normal text
        p = active_paragraph if active_paragraph else doc.add_paragraph()
        render_inline_markdown(p, stripped)
        active_paragraph = None
        i += 1


# =============================================================================
# Inline markdown helpers (shared by all Scripts/* add_markdown_to_doc impls)
# =============================================================================

import re as _re

# Single source of truth for inline emphasis. Tried in order, so `**` wins over
# `*`. Each alternation has its own capture group so we know which form matched.
#
#   group 1: ``**bold**``      content must not contain ``**``
#   group 2: ``*italic*``      content must not contain ``*``; no whitespace
#                              immediately inside the asterisks (so list
#                              bullets like ``* item`` don't match), and the
#                              outer ``*`` cannot be next to a word char on
#                              the closing side (avoids ``2*3*4`` triggering)
#   group 3: `` `code` ``      monospace
_INLINE_TOKEN = _re.compile(
    r"\*\*((?:(?!\*\*).)+?)\*\*"
    r"|(?<![\w*])\*(?!\s)([^*]+?)(?<!\s)\*(?!\w)"
    r"|`([^`]+?)`"
)


def _add_inline_run(para, text, *, bold=False, italic=False, code=False,
                    font_name="Times New Roman", font_size_pt=12):
    """Append a styled run to ``para``. Skip empty text."""
    if not text:
        return
    run = para.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if code:
        # Monospace for inline code; the size match keeps the line stable.
        run.font.name = "Consolas"
    elif font_name:
        run.font.name = font_name
    if font_size_pt and Pt is not None:
        run.font.size = Pt(font_size_pt)


def render_inline_markdown(para, text, *, font_name="Times New Roman",
                           font_size_pt=12):
    """Parse inline markdown in ``text`` and append styled runs to ``para``.

    Handles ``**bold**``, ``*italic*``, and `` `code` ``. Block-level syntax
    (headings, lists, tables) must be handled by the caller — this helper is
    inline-only, so passing a heading line through it will not bold the whole
    line.

    Callers that previously did ``re.split(r'(\\*\\*.*?\\*\\*)', text)`` should
    swap in a single call to this helper. It fixes three pre-existing bugs in
    Scripts/*: the broken ``\\**`` zero-or-more regex (med_record, med_chron,
    exposure), the greedy ``\\*\\*.*\\*\\*`` regex that swallowed multiple
    bold spans (liability), and the total lack of italic/code support.
    """
    if not text:
        return

    pos = 0
    for m in _INLINE_TOKEN.finditer(text):
        # Plain text before the matched token.
        if m.start() > pos:
            _add_inline_run(para, text[pos:m.start()],
                            font_name=font_name, font_size_pt=font_size_pt)
        bold_grp, italic_grp, code_grp = m.group(1), m.group(2), m.group(3)
        if bold_grp is not None:
            _add_inline_run(para, bold_grp, bold=True,
                            font_name=font_name, font_size_pt=font_size_pt)
        elif italic_grp is not None:
            _add_inline_run(para, italic_grp, italic=True,
                            font_name=font_name, font_size_pt=font_size_pt)
        elif code_grp is not None:
            _add_inline_run(para, code_grp, code=True,
                            font_size_pt=font_size_pt)
        pos = m.end()
    # Trailing plain text after the last token.
    if pos < len(text):
        _add_inline_run(para, text[pos:],
                        font_name=font_name, font_size_pt=font_size_pt)


# =============================================================================
# Markdown table helpers (shared by all Scripts/* add_markdown_to_doc impls)
# =============================================================================

_TABLE_SEPARATOR_CELL = _re.compile(r"^\s*:?-{3,}:?\s*$")


def _split_table_row(line: str) -> list[str]:
    """Split a markdown table row into cells.

    Accepts both ``| a | b | c |`` and ``a | b | c``. Trailing/leading empty
    cells from outer pipes are stripped. Returns ``[]`` if the line has no
    pipe character.
    """
    stripped = line.strip()
    if "|" not in stripped:
        return []
    if stripped.startswith("|"):
        stripped = stripped[1:]
    if stripped.endswith("|"):
        stripped = stripped[:-1]
    return [cell.strip() for cell in stripped.split("|")]


def _is_table_separator_line(line: str) -> bool:
    """True if ``line`` looks like a markdown table separator (``| --- | --- |``)."""
    cells = _split_table_row(line)
    if not cells:
        return False
    return all(_TABLE_SEPARATOR_CELL.match(c) for c in cells)


def try_consume_markdown_table(lines: list[str], start_idx: int):
    """Try to recognise a markdown table starting at ``lines[start_idx]``.

    A valid table is a header row followed immediately by a separator row,
    optionally followed by data rows. Returns
    ``(consumed_count, header_cells, list_of_row_cells)`` on success, or
    ``(0, None, None)`` if no table is recognised at this position.
    """
    if start_idx + 1 >= len(lines):
        return 0, None, None
    header_line = lines[start_idx]
    sep_line = lines[start_idx + 1]
    if "|" not in header_line:
        return 0, None, None
    if not _is_table_separator_line(sep_line):
        return 0, None, None
    header = _split_table_row(header_line)
    if not header:
        return 0, None, None

    rows: list[list[str]] = []
    i = start_idx + 2
    while i < len(lines):
        row_line = lines[i]
        if not row_line.strip():
            break
        if "|" not in row_line:
            break
        cells = _split_table_row(row_line)
        if not cells:
            break
        # Pad / trim to header width so the Word table stays rectangular.
        if len(cells) < len(header):
            cells = cells + [""] * (len(header) - len(cells))
        elif len(cells) > len(header):
            cells = cells[: len(header)]
        rows.append(cells)
        i += 1

    consumed = i - start_idx
    return consumed, header, rows


def add_markdown_table_to_doc(doc, header: list[str], rows: list[list[str]]):
    """Append a real Word table to ``doc``. The header row is bolded.

    Uses ``Table Grid`` style if available (gives visible borders that survive
    a round trip through mammoth → HTML → QTextEdit). Falls back to the
    default style if the template doesn't define ``Table Grid``.
    """
    if Document is None:
        return
    cols = len(header)
    table = doc.add_table(rows=1 + len(rows), cols=cols)
    try:
        table.style = "Table Grid"
    except KeyError:
        pass

    # NOTE: do NOT use ``cell.text = ""`` to clear — python-docx's text setter
    # creates an empty default run, so our subsequent add_run lands at runs[1]
    # and the bold/font formatting on it gets shadowed by an unstyled runs[0].
    # ``add_table`` already gives each cell one empty paragraph with zero runs;
    # we just append our styled run to it.
    hdr_cells = table.rows[0].cells
    for c_idx, text in enumerate(header):
        para = hdr_cells[c_idx].paragraphs[0]
        run = para.add_run(text)
        run.bold = True
        if Pt is not None:
            run.font.size = Pt(12)
        run.font.name = "Times New Roman"

    for r_idx, row in enumerate(rows, start=1):
        row_cells = table.rows[r_idx].cells
        for c_idx, text in enumerate(row):
            para = row_cells[c_idx].paragraphs[0]
            run = para.add_run(text)
            if Pt is not None:
                run.font.size = Pt(12)
            run.font.name = "Times New Roman"
