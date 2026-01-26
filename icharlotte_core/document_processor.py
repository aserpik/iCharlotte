"""
Shared Document Processing Module for iCharlotte

Consolidates text extraction, OCR, and document processing logic
from summarize.py, summarize_discovery.py, and summarize_deposition.py.
"""

import os
import gc
import re
from dataclasses import dataclass, field
from typing import List, Optional, Tuple
from enum import Enum

# PDF/Document libraries
from pypdf import PdfReader
from docx import Document as DocxDocument

# OCR setup
OCR_AVAILABLE = False
POPPLER_PATH = None

try:
    import pytesseract
    from pdf2image import convert_from_path
    OCR_AVAILABLE = True

    # Windows-specific path configuration
    if os.name == 'nt':
        tesseract_path = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
        if os.path.exists(tesseract_path):
            pytesseract.pytesseract.tesseract_cmd = tesseract_path

        poppler_path = r"C:\Program Files\poppler\Library\bin"
        if os.path.exists(poppler_path):
            POPPLER_PATH = poppler_path
except ImportError:
    pass


class ExtractionMethod(Enum):
    """How the text was extracted from the document."""
    NATIVE = "native"      # Direct text extraction
    OCR = "ocr"            # Full OCR
    MIXED = "mixed"        # Some pages native, some OCR
    FAILED = "failed"      # Extraction failed


@dataclass
class OCRConfig:
    """Configuration for OCR processing."""
    base_threshold: int = 50          # Chars per page to trigger OCR
    adaptive: bool = True             # Use dynamic threshold based on document type
    max_dpi: int = 300                # Maximum DPI for OCR rendering
    min_threshold: int = 20           # Minimum threshold (for sparse forms)
    max_threshold: int = 200          # Maximum threshold (for dense documents)
    gc_interval: int = 10             # Run GC every N pages


@dataclass
class ExtractResult:
    """Result of text extraction from a document."""
    text: str
    page_count: int
    ocr_pages: List[int] = field(default_factory=list)  # Pages that required OCR
    extraction_method: ExtractionMethod = ExtractionMethod.NATIVE
    char_count: int = 0
    char_density: float = 0.0         # Average chars per page
    file_path: str = ""
    error: Optional[str] = None

    @property
    def success(self) -> bool:
        """Whether extraction was successful."""
        return self.extraction_method != ExtractionMethod.FAILED and bool(self.text.strip())

    @property
    def ocr_percentage(self) -> float:
        """Percentage of pages that required OCR."""
        if self.page_count == 0:
            return 0.0
        return len(self.ocr_pages) / self.page_count * 100


class DocumentProcessor:
    """
    Unified document processor with text extraction and OCR support.

    Usage:
        processor = DocumentProcessor()
        result = processor.extract_text("path/to/document.pdf")
        if result.success:
            print(result.text)
    """

    def __init__(self, ocr_config: Optional[OCRConfig] = None, logger=None):
        """
        Initialize the document processor.

        Args:
            ocr_config: OCR configuration. Uses defaults if not provided.
            logger: Optional logger instance for output. If None, uses print().
        """
        self.ocr_config = ocr_config or OCRConfig()
        self.logger = logger
        self._ocr_available = OCR_AVAILABLE
        self._poppler_path = POPPLER_PATH

    def _log(self, message: str, level: str = "info"):
        """Log a message using the configured logger or print."""
        if self.logger:
            if hasattr(self.logger, level):
                getattr(self.logger, level)(message)
            else:
                self.logger.info(message)
        else:
            print(message, flush=True)

    def extract_text(self, file_path: str) -> ExtractResult:
        """
        Extract text from a document (PDF, DOCX, or plain text).

        Args:
            file_path: Path to the document file.

        Returns:
            ExtractResult with extracted text and metadata.
        """
        self._log(f"Extracting text from: {file_path}")

        if not os.path.exists(file_path):
            return ExtractResult(
                text="",
                page_count=0,
                extraction_method=ExtractionMethod.FAILED,
                file_path=file_path,
                error=f"File not found: {file_path}"
            )

        ext = os.path.splitext(file_path)[1].lower()

        try:
            if ext == ".pdf":
                return self._extract_from_pdf(file_path)
            elif ext == ".docx":
                return self._extract_from_docx(file_path)
            else:
                return self._extract_from_text(file_path)
        except Exception as e:
            self._log(f"Error extracting text: {e}", "error")
            return ExtractResult(
                text="",
                page_count=0,
                extraction_method=ExtractionMethod.FAILED,
                file_path=file_path,
                error=str(e)
            )

    def extract_with_dynamic_ocr(self, file_path: str) -> ExtractResult:
        """
        Extract text using adaptive OCR threshold based on document characteristics.

        The threshold adapts based on:
        - Document type (detected from content)
        - Page density patterns
        - First-pass extraction results

        Args:
            file_path: Path to the document file.

        Returns:
            ExtractResult with extracted text and metadata.
        """
        # First, do a quick scan to determine document characteristics
        ext = os.path.splitext(file_path)[1].lower()

        if ext != ".pdf":
            # Non-PDF files don't need dynamic OCR threshold
            return self.extract_text(file_path)

        try:
            reader = PdfReader(file_path)
            total_pages = len(reader.pages)

            # Sample first few pages to determine characteristics
            sample_densities = []
            sample_size = min(5, total_pages)

            for i in range(sample_size):
                page_text = reader.pages[i].extract_text() or ""
                sample_densities.append(len(page_text.strip()))

            # Calculate adaptive threshold
            threshold = self._calculate_dynamic_threshold(sample_densities, file_path)
            self._log(f"Dynamic OCR threshold calculated: {threshold} chars/page")

            # Now do full extraction with calculated threshold
            original_threshold = self.ocr_config.base_threshold
            self.ocr_config.base_threshold = threshold

            try:
                result = self._extract_from_pdf(file_path)
            finally:
                self.ocr_config.base_threshold = original_threshold

            return result

        except Exception as e:
            self._log(f"Error in dynamic OCR extraction: {e}", "error")
            return self.extract_text(file_path)

    def _calculate_dynamic_threshold(self, sample_densities: List[int], file_path: str) -> int:
        """
        Calculate adaptive OCR threshold based on document characteristics.

        Args:
            sample_densities: Character counts from sample pages.
            file_path: Path to file (for filename-based heuristics).

        Returns:
            Calculated threshold value.
        """
        if not sample_densities:
            return self.ocr_config.base_threshold

        avg_density = sum(sample_densities) / len(sample_densities)
        filename = os.path.basename(file_path).lower()

        # Filename-based heuristics
        is_form = any(term in filename for term in ['form', 'frog', 'rog', 'rfa', 'rfp'])
        is_pleading = any(term in filename for term in ['complaint', 'motion', 'demurrer', 'opposition'])
        is_deposition = any(term in filename for term in ['depo', 'transcript'])

        # Base threshold adjustments
        if is_form:
            # Forms often have sparse text with checkboxes
            base = self.ocr_config.min_threshold
        elif is_deposition:
            # Depositions are usually text-heavy
            base = self.ocr_config.max_threshold
        elif is_pleading:
            # Pleadings are usually well-formatted with good text
            base = 100
        else:
            # Default: scale based on observed density
            if avg_density < 100:
                base = self.ocr_config.min_threshold
            elif avg_density < 500:
                base = self.ocr_config.base_threshold
            else:
                base = min(avg_density * 0.1, self.ocr_config.max_threshold)

        # Clamp to configured range
        return max(self.ocr_config.min_threshold,
                   min(int(base), self.ocr_config.max_threshold))

    def _extract_from_pdf(self, file_path: str) -> ExtractResult:
        """Extract text from a PDF file with OCR fallback."""
        reader = PdfReader(file_path)
        total_pages = len(reader.pages)
        self._log(f"PDF has {total_pages} pages.")

        text_parts = []
        ocr_pages = []
        threshold = self.ocr_config.base_threshold

        for i, page in enumerate(reader.pages):
            # Memory management
            if i % self.ocr_config.gc_interval == 0:
                gc.collect()

            page_text = page.extract_text() or ""

            # Check if OCR is needed
            if len(page_text.strip()) < threshold:
                self._log(f"Page {i+1} has insufficient text ({len(page_text.strip())} chars). Attempting OCR...", "warning")

                ocr_text = self._ocr_page(file_path, i)
                if ocr_text and len(ocr_text.strip()) > len(page_text.strip()):
                    self._log(f"OCR successful for page {i+1}. Extracted {len(ocr_text)} chars.")
                    text_parts.append(ocr_text)
                    ocr_pages.append(i)
                else:
                    self._log(f"OCR did not improve page {i+1}. Using original.", "warning")
                    text_parts.append(page_text)
            else:
                text_parts.append(page_text)

        full_text = "\n".join(text_parts)
        char_count = len(full_text)

        # Determine extraction method
        if len(ocr_pages) == 0:
            method = ExtractionMethod.NATIVE
        elif len(ocr_pages) == total_pages:
            method = ExtractionMethod.OCR
        else:
            method = ExtractionMethod.MIXED

        return ExtractResult(
            text=full_text,
            page_count=total_pages,
            ocr_pages=ocr_pages,
            extraction_method=method,
            char_count=char_count,
            char_density=char_count / total_pages if total_pages > 0 else 0,
            file_path=file_path
        )

    def _ocr_page(self, file_path: str, page_index: int) -> Optional[str]:
        """OCR a specific page of a PDF."""
        if not self._ocr_available:
            self._log("OCR not available.", "warning")
            return None

        try:
            images = convert_from_path(
                file_path,
                first_page=page_index + 1,
                last_page=page_index + 1,
                poppler_path=self._poppler_path,
                dpi=self.ocr_config.max_dpi
            )

            if images:
                return pytesseract.image_to_string(images[0])
            return None

        except Exception as e:
            self._log(f"Error OCR'ing page {page_index + 1}: {e}", "warning")
            return None

    def _extract_from_docx(self, file_path: str) -> ExtractResult:
        """Extract text from a DOCX file."""
        doc = DocxDocument(file_path)

        text_parts = []
        for paragraph in doc.paragraphs:
            text_parts.append(paragraph.text)

        full_text = "\n".join(text_parts)

        # Estimate page count (rough approximation)
        estimated_pages = max(1, len(full_text) // 3000)

        return ExtractResult(
            text=full_text,
            page_count=estimated_pages,
            extraction_method=ExtractionMethod.NATIVE,
            char_count=len(full_text),
            char_density=len(full_text) / estimated_pages,
            file_path=file_path
        )

    def _extract_from_text(self, file_path: str) -> ExtractResult:
        """Extract text from a plain text file."""
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()

        # Estimate page count
        estimated_pages = max(1, len(text) // 3000)

        return ExtractResult(
            text=text,
            page_count=estimated_pages,
            extraction_method=ExtractionMethod.NATIVE,
            char_count=len(text),
            char_density=len(text) / estimated_pages,
            file_path=file_path
        )

    @staticmethod
    def get_page_count(file_path: str) -> int:
        """Get page count without full extraction."""
        ext = os.path.splitext(file_path)[1].lower()

        try:
            if ext == ".pdf":
                reader = PdfReader(file_path)
                return len(reader.pages)
            elif ext == ".docx":
                doc = DocxDocument(file_path)
                # Rough estimate based on paragraphs
                return max(1, len(doc.paragraphs) // 30)
            else:
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    text = f.read()
                return max(1, len(text) // 3000)
        except Exception:
            return 0


class VerificationDetector:
    """
    Detects verification pages in legal documents.

    Verification pages are typically:
    - Short documents (< 5 pages)
    - Contain "VERIFICATION" title
    - Contain perjury language
    - Lack actual response content
    """

    # Patterns indicating a verification page
    VERIFICATION_PATTERNS = [
        r'^\s*VERIFICATION\s*$',
        r'UNDER PENALTY OF PERJURY',
        r'TRUE AND CORRECT',
        r'EXECUTED ON',
        r'I DECLARE UNDER PENALTY',
    ]

    # Patterns indicating actual response content
    RESPONSE_PATTERNS = [
        r'RESPONSE\s+TO\s+',
        r'RESPONSES\s+TO\s+',
        r'ANSWER\s+TO\s+',
        r'INTERROGATORY\s+NO\.',
        r'REQUEST\s+NO\.',
    ]

    @classmethod
    def is_verification_page(cls, text: str) -> bool:
        """
        Check if text content indicates a verification-only page.

        Args:
            text: Text content to analyze.

        Returns:
            True if this appears to be a verification page.
        """
        if not text:
            return False

        text_upper = text.upper()
        lines = [l.strip() for l in text_upper.split('\n') if l.strip()]

        # Check for verification title in first 20 lines
        has_verification_title = False
        for i in range(min(20, len(lines))):
            if lines[i] == "VERIFICATION":
                has_verification_title = True
                break

        # Check for perjury language
        has_perjury_lang = (
            "UNDER PENALTY OF PERJURY" in text_upper and
            ("TRUE AND CORRECT" in text_upper or "EXECUTED ON" in text_upper)
        )

        # If we have the verification title, it's likely a verification
        if has_verification_title:
            return True

        # Check if it lacks response content but has perjury language
        has_response_content = any(
            re.search(pattern, text_upper) for pattern in cls.RESPONSE_PATTERNS
        )

        if has_perjury_lang and not has_response_content:
            return True

        return False

    @classmethod
    def is_verification_document(cls, file_path: str, max_pages: int = 5) -> Tuple[bool, str]:
        """
        Check if an entire document is just a verification.

        Improved detection that checks multiple pages and document length.

        Args:
            file_path: Path to the document.
            max_pages: Maximum page count for a verification-only document.

        Returns:
            Tuple of (is_verification, reason)
        """
        processor = DocumentProcessor()

        # Get page count first
        page_count = processor.get_page_count(file_path)

        # Verification documents are typically short
        if page_count > max_pages:
            return False, f"Document has {page_count} pages (> {max_pages})"

        # Extract text from first few pages
        ext = os.path.splitext(file_path)[1].lower()

        try:
            if ext == ".pdf":
                reader = PdfReader(file_path)
                # Check first 3 pages
                pages_to_check = min(3, len(reader.pages))
                first_pages_text = ""

                for i in range(pages_to_check):
                    page_text = reader.pages[i].extract_text() or ""
                    first_pages_text += page_text + "\n"

                    # If any page has response content, not a verification
                    if any(re.search(p, page_text.upper()) for p in cls.RESPONSE_PATTERNS):
                        return False, f"Page {i+1} contains response content"

                if cls.is_verification_page(first_pages_text):
                    return True, "First pages match verification pattern"

            elif ext == ".docx":
                doc = DocxDocument(file_path)
                # Get first 30 paragraphs (roughly first page)
                first_page_text = "\n".join(
                    p.text for p in doc.paragraphs[:30]
                )

                if cls.is_verification_page(first_page_text):
                    # Also check if there's substantial content after
                    full_text = "\n".join(p.text for p in doc.paragraphs)
                    if len(full_text) < 2000:  # Short document
                        return True, "Short DOCX matching verification pattern"

            else:
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    text = f.read()

                if len(text) < 2000 and cls.is_verification_page(text[:2000]):
                    return True, "Short text file matching verification pattern"

        except Exception as e:
            return False, f"Error checking document: {e}"

        return False, "No verification indicators found"


class DocumentTypeClassifier:
    """
    Classifies legal document types based on content.
    """

    DOCUMENT_TYPES = {
        'form_interrogatories': [
            r'FORM INTERROGATOR',
            r'FROG',
            r'FORM INT',
        ],
        'special_interrogatories': [
            r'SPECIAL INTERROGATOR',
            r'SROG',
            r'SPECIAL INT',
        ],
        'requests_for_production': [
            r'REQUEST\s+FOR\s+PRODUCTION',
            r'DEMAND\s+FOR\s+PRODUCTION',
            r'RFP',
            r'RFPD',
        ],
        'requests_for_admission': [
            r'REQUEST\s+FOR\s+ADMISSION',
            r'RFA',
        ],
        'deposition': [
            r'DEPOSITION\s+OF',
            r'DEPOSITION\s+TRANSCRIPT',
            r'ORAL\s+DEPOSITION',
        ],
        'motion': [
            r'MOTION\s+TO',
            r'MOTION\s+FOR',
            r'NOTICE\s+OF\s+MOTION',
        ],
        'opposition': [
            r'OPPOSITION\s+TO',
            r'MEMORANDUM\s+IN\s+OPPOSITION',
        ],
        'reply': [
            r'REPLY\s+TO',
            r'REPLY\s+IN\s+SUPPORT',
            r'REPLY\s+BRIEF',
        ],
        'complaint': [
            r'COMPLAINT\s+FOR',
            r'FIRST\s+AMENDED\s+COMPLAINT',
            r'SECOND\s+AMENDED\s+COMPLAINT',
        ],
        'answer': [
            r'ANSWER\s+TO\s+COMPLAINT',
            r'ANSWER\s+TO.*AMENDED\s+COMPLAINT',
        ],
    }

    @classmethod
    def classify(cls, text: str, filename: str = "") -> Tuple[str, float]:
        """
        Classify document type based on content and filename.

        Args:
            text: Document text content.
            filename: Optional filename for additional hints.

        Returns:
            Tuple of (document_type, confidence)
        """
        text_upper = text.upper()
        filename_upper = filename.upper()

        scores = {}

        for doc_type, patterns in cls.DOCUMENT_TYPES.items():
            score = 0
            for pattern in patterns:
                # Check content
                matches = len(re.findall(pattern, text_upper))
                score += matches * 2

                # Check filename (higher weight)
                if re.search(pattern, filename_upper):
                    score += 5

            if score > 0:
                scores[doc_type] = score

        if not scores:
            return 'unknown', 0.0

        # Get highest scoring type
        best_type = max(scores, key=scores.get)
        max_score = scores[best_type]

        # Calculate confidence (normalize to 0-1)
        confidence = min(1.0, max_score / 20.0)

        return best_type, confidence

    @classmethod
    def get_discovery_type(cls, text: str, filename: str = "") -> Optional[str]:
        """
        Get specific discovery type for discovery responses.

        Returns:
            Discovery type string or None if not a discovery document.
        """
        doc_type, confidence = cls.classify(text, filename)

        discovery_types = {
            'form_interrogatories': 'Form Interrogatories',
            'special_interrogatories': 'Special Interrogatories',
            'requests_for_production': 'Requests for Production',
            'requests_for_admission': 'Requests for Admission',
        }

        return discovery_types.get(doc_type)


class PartyExtractor:
    """
    Extracts party names from legal documents.
    """

    # Patterns for extracting party names
    PARTY_PATTERNS = [
        # "Plaintiff JOHN DOE's Responses"
        r"(Plaintiff|Defendant)\s+([A-Z][A-Za-z]+(?:\s+[A-Z][A-Za-z]+)*)'?s?\s+Response",
        # "Responses of Defendant JOHN DOE"
        r"Response[s]?\s+of\s+(Plaintiff|Defendant)\s+([A-Z][A-Za-z]+(?:\s+[A-Z][A-Za-z]+)*)",
        # "JOHN DOE, Plaintiff"
        r"([A-Z][A-Za-z]+(?:\s+[A-Z][A-Za-z]+)*),?\s+(Plaintiff|Defendant)",
        # Caption-style: "JOHN DOE, an individual"
        r"([A-Z][A-Z\s]+),?\s+an?\s+individual",
    ]

    @classmethod
    def extract_party_name(cls, text: str, party_type: str = None) -> Optional[str]:
        """
        Extract party name from document text.

        Args:
            text: Document text content.
            party_type: Optional filter for "Plaintiff" or "Defendant".

        Returns:
            Extracted party name or None.
        """
        text_sample = text[:5000]  # Check first portion

        for pattern in cls.PARTY_PATTERNS:
            matches = re.finditer(pattern, text_sample, re.IGNORECASE)
            for match in matches:
                groups = match.groups()

                # Determine which group has the party type and which has the name
                if len(groups) >= 2:
                    if groups[0].lower() in ('plaintiff', 'defendant'):
                        found_type = groups[0].title()
                        name = groups[1]
                    else:
                        name = groups[0]
                        found_type = groups[1].title() if len(groups) > 1 else None

                    # Filter by party type if specified
                    if party_type and found_type and party_type.lower() != found_type.lower():
                        continue

                    # Clean up the name
                    name = cls._clean_party_name(name)
                    if name:
                        return name

        return None

    @classmethod
    def _clean_party_name(cls, name: str) -> str:
        """Clean and normalize a party name."""
        if not name:
            return ""

        # Remove common suffixes
        name = re.sub(r',?\s*(an individual|a corporation|et al\.?|individually).*$', '', name, flags=re.IGNORECASE)

        # Remove extra whitespace
        name = ' '.join(name.split())

        # Title case if all caps
        if name.isupper():
            name = name.title()

        return name.strip()

    @classmethod
    def extract_deponent_name(cls, text: str) -> Optional[str]:
        """
        Extract deponent name from deposition transcript.

        Args:
            text: Deposition transcript text.

        Returns:
            Deponent name or None.
        """
        patterns = [
            r"DEPOSITION\s+OF\s+([A-Z][A-Za-z]+(?:\s+[A-Z][A-Za-z]+)*)",
            r"ORAL\s+DEPOSITION\s+OF\s+([A-Z][A-Za-z]+(?:\s+[A-Z][A-Za-z]+)*)",
            r"VIDEOTAPED\s+DEPOSITION\s+OF\s+([A-Z][A-Za-z]+(?:\s+[A-Z][A-Za-z]+)*)",
            r"WITNESS:\s*([A-Z][A-Za-z]+(?:\s+[A-Z][A-Za-z]+)*)",
        ]

        text_sample = text[:3000]  # Check header area

        for pattern in patterns:
            match = re.search(pattern, text_sample, re.IGNORECASE)
            if match:
                name = match.group(1)
                return cls._clean_party_name(name)

        return None

    @classmethod
    def extract_deposition_date(cls, text: str) -> Optional[str]:
        """
        Extract deposition date from transcript.

        Args:
            text: Deposition transcript text.

        Returns:
            Date string or None.
        """
        patterns = [
            r"(?:taken|held|commenced)\s+on\s+([A-Z][a-z]+\s+\d{1,2},?\s+\d{4})",
            r"(\d{1,2}/\d{1,2}/\d{4})",
            r"([A-Z][a-z]+\s+\d{1,2},?\s+\d{4})",
        ]

        text_sample = text[:2000]

        for pattern in patterns:
            match = re.search(pattern, text_sample, re.IGNORECASE)
            if match:
                return match.group(1)

        return None
