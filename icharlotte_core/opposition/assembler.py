"""Word assembly helpers for opposition memorandum previews."""

from __future__ import annotations

import os
import re

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from icharlotte_core.word_validator import validate_opposition_docx

from .models import DraftDocument


def assemble_opposition_preview(
    *,
    draft: DraftDocument,
    output_path: str,
    caption_path: str = "",
) -> str:
    """Render an opposition draft to a .docx preview and validate the file."""
    if caption_path and os.path.isfile(caption_path):
        if _same_path(caption_path, output_path):
            raise ValueError("Output path must be different from the caption template path")
        doc = Document(caption_path)
        if not _replace_caption_markers(doc, draft.title or "Opposition"):
            _add_title(doc, draft.title or "Opposition")
    else:
        doc = Document()
        _add_title(doc, draft.title or "Opposition")

    _add_paragraphs(doc, draft.body_text)

    output_dir = os.path.dirname(output_path)
    if output_dir:
        os.makedirs(output_dir, exist_ok=True)
    doc.save(output_path)

    validation = validate_opposition_docx(output_path)
    if validation.has_errors:
        messages = "; ".join(str(finding) for finding in validation.findings)
        raise ValueError(f"Generated opposition failed validation: {messages}")

    return output_path


def _add_title(doc, title: str) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(title)
    run.bold = True
    _format_run(run)


_HORIZONTAL_RULE_RE = re.compile(r"^[\*\-_]{3,}\s*$")
_MD_HEADING_RE = re.compile(r"^(#{1,6})\s+(.+?)\s*$")
_MD_ITALIC_RE = re.compile(r"\*([^\*\n]+?)\*")


def _add_paragraphs(doc, body_text: str) -> None:
    for raw_line in (body_text or "").splitlines():
        line = raw_line.rstrip()
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph("")
            continue

        # Skip markdown horizontal rules ("***", "---", "___") — they don't translate.
        if _HORIZONTAL_RULE_RE.match(stripped):
            continue

        # Markdown heading → bolded line. Higher levels (## / ###) are also bolded
        # but use a slightly smaller bold via paragraph formatting.
        heading_match = _MD_HEADING_RE.match(stripped)
        if heading_match:
            _add_heading_paragraph(
                doc,
                level=len(heading_match.group(1)),
                text=heading_match.group(2),
            )
            continue

        paragraph = doc.add_paragraph()
        _emit_inline(paragraph, stripped, base_bold=False)
        if _looks_like_legacy_heading(stripped):
            for run in paragraph.runs:
                run.bold = True


def _emit_inline(paragraph, text: str, *, base_bold: bool) -> None:
    """Emit text into a paragraph, converting *italic* segments to italic runs."""
    index = 0
    for match in _MD_ITALIC_RE.finditer(text):
        start, end = match.span()
        if start > index:
            run = paragraph.add_run(text[index:start])
            _format_run(run)
            run.bold = base_bold
        italic_run = paragraph.add_run(match.group(1))
        _format_run(italic_run)
        italic_run.italic = True
        italic_run.bold = base_bold
        index = end
    if index < len(text):
        run = paragraph.add_run(text[index:])
        _format_run(run)
        run.bold = base_bold


def _add_heading_paragraph(doc, *, level: int, text: str) -> None:
    paragraph = doc.add_paragraph()
    _emit_inline(paragraph, text, base_bold=True)
    if level >= 3:
        for run in paragraph.runs:
            run.italic = True


def _format_run(run) -> None:
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def _looks_like_legacy_heading(line: str) -> bool:
    return line.isupper() or line.startswith(
        ("I.", "II.", "III.", "IV.", "V.", "VI.", "VII.", "VIII.", "IX.", "X.")
    )


def _same_path(first: str, second: str) -> bool:
    if not first or not second:
        return False
    return os.path.normcase(os.path.abspath(first)) == os.path.normcase(
        os.path.abspath(second)
    )


def _replace_caption_markers(doc, title: str) -> bool:
    """Replace CAPTION PAGE markers in body, headers, and footers."""
    w_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    w_t = f"{{{w_ns}}}t"
    w_p = f"{{{w_ns}}}p"
    w_r = f"{{{w_ns}}}r"
    replaced = 0

    for root in _iter_story_roots(doc):
        for paragraph_element in list(root.iter(w_p)):
            paragraph_text = "".join(
                text_element.text or "" for text_element in paragraph_element.iter(w_t)
            )
            if "CAPTION PAGE" not in paragraph_text.upper():
                continue

            for run_element in list(paragraph_element.iter(w_r)):
                parent = run_element.getparent()
                if parent is not None:
                    parent.remove(run_element)

            paragraph_element.append(_make_title_run(title))
            replaced += 1

    return replaced > 0


def _iter_story_roots(doc):
    seen: set[int] = set()
    roots = [doc.element.body]
    for section in doc.sections:
        roots.extend(
            [
                section.header._element,
                section.first_page_header._element,
                section.even_page_header._element,
                section.footer._element,
                section.first_page_footer._element,
                section.even_page_footer._element,
            ]
        )
    for root in roots:
        root_id = id(root)
        if root_id in seen:
            continue
        seen.add(root_id)
        yield root


def _make_title_run(text: str):
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    properties.append(OxmlElement("w:b"))
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Times New Roman")
    fonts.set(qn("w:hAnsi"), "Times New Roman")
    properties.append(fonts)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "24")
    properties.append(size)
    complex_size = OxmlElement("w:szCs")
    complex_size.set(qn("w:val"), "24")
    properties.append(complex_size)
    run.append(properties)
    text_node = OxmlElement("w:t")
    text_node.text = text
    text_node.set(qn("xml:space"), "preserve")
    run.append(text_node)
    return run
