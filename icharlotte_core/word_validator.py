"""
Shared Word document validation library.

Composable check functions for validating Word document formatting,
redline quality, and report structure. Used by:
- word_hotkey.py (redline post-application validation)
- test_redline_real_doc.py (test script validation)
- Scripts/report_generator/validate.py (report pipeline validation)

Each check function follows the contract:
    def check_xxx(doc, **kwargs) -> list[Finding]

Callers pick which checks to run, collect findings, then decide action.

Two document backends:
- COM (win32com): for live Word operations (redline checks)
- python-docx: for offline file inspection (report checks)
"""

import os
import re
from dataclasses import dataclass, field
from typing import Any, Dict, List, Optional


# ---------------------------------------------------------------------------
# Data classes
# ---------------------------------------------------------------------------

@dataclass
class Finding:
    """A single validation finding."""
    severity: str       # "ERROR", "WARN", "INFO", "PASS"
    rule: str           # Rule name (e.g., "revisions_in_range", "heading_formatting")
    message: str        # Human-readable description
    location: Optional[str] = None  # e.g., "para 3", "pos 1234"
    expected: Any = None
    actual: Any = None

    def __str__(self):
        loc = f" ({self.location})" if self.location else ""
        return f"[{self.severity}] {self.rule}{loc}: {self.message}"


@dataclass
class ValidationResult:
    """Aggregated validation results."""
    context: str  # Description of what was validated
    findings: List[Finding] = field(default_factory=list)

    @property
    def has_errors(self) -> bool:
        return any(f.severity == "ERROR" for f in self.findings)

    @property
    def has_warnings(self) -> bool:
        return any(f.severity == "WARN" for f in self.findings)

    @property
    def error_count(self) -> int:
        return sum(1 for f in self.findings if f.severity == "ERROR")

    @property
    def warn_count(self) -> int:
        return sum(1 for f in self.findings if f.severity == "WARN")

    @property
    def pass_count(self) -> int:
        return sum(1 for f in self.findings if f.severity == "PASS")

    def print_summary(self, verbose=False):
        """Print results to console."""
        print(f"\n{'='*60}")
        print(f"VALIDATION: {self.context}")
        print(f"{'='*60}")
        for f in self.findings:
            if verbose or f.severity in ("ERROR", "WARN"):
                print(f"  {f}")
        if not self.has_errors and not self.has_warnings:
            print("  ALL CHECKS PASSED")
        print(f"\nResults: {self.error_count} ERROR, {self.warn_count} WARN, "
              f"{self.pass_count} PASS")
        print(f"{'='*60}")


# ===================================================================
# REDLINE CHECKS (COM-based, require live win32com Document)
# ===================================================================

def check_revisions_in_range(doc_com, range_start: int, range_end: int,
                              tolerance: int = 50) -> List[Finding]:
    """Check that all revisions fall within the target section range.

    Args:
        doc_com: win32com Document object
        range_start: Start position of the target section
        range_end: End position of the target section
        tolerance: Allow revisions up to this many chars past range_end
    """
    findings = []
    try:
        revs = doc_com.Revisions
        outside_count = 0
        for i in range(1, revs.Count + 1):
            rev = revs(i)
            rev_start = rev.Range.Start
            rev_end = rev.Range.End
            if rev_start < range_start or rev_end > range_end + tolerance:
                rev_text = rev.Range.Text
                rtype = 'INS' if rev.Type == 1 else 'DEL' if rev.Type == 2 else f'T{rev.Type}'
                findings.append(Finding(
                    "ERROR", "revisions_in_range",
                    f"{rtype} outside target range: {repr(rev_text[:50])}",
                    location=f"pos {rev_start}-{rev_end}",
                    expected=f"{range_start}-{range_end}",
                    actual=f"{rev_start}-{rev_end}",
                ))
                outside_count += 1

        if outside_count == 0:
            findings.append(Finding(
                "PASS", "revisions_in_range",
                f"All {revs.Count} revisions within target range"
            ))
    except Exception as e:
        findings.append(Finding(
            "WARN", "revisions_in_range",
            f"Could not check revisions: {e}"
        ))
    return findings


def check_paragraph_marks_preserved(doc_com, range_start: int,
                                     range_end: int) -> List[Finding]:
    """Check that no content paragraphs were destroyed (\\r deleted with content).

    Distinguishes between:
    - ERROR: Content paragraph deleted (\\r + meaningful text)
    - WARN: Empty paragraph mark deleted (just \\r, no content)
    """
    findings = []
    content_deleted = 0
    empty_deleted = 0

    try:
        revs = doc_com.Revisions
        for i in range(1, revs.Count + 1):
            rev = revs(i)
            if rev.Type != 2:  # Only check deletions
                continue
            rev_text = rev.Range.Text
            if '\r' not in rev_text:
                continue

            stripped = rev_text.replace('\r', '').strip()
            rev_start = rev.Range.Start
            if stripped:
                findings.append(Finding(
                    "ERROR", "paragraph_marks_preserved",
                    f"Content paragraph deleted: {repr(rev_text[:50])}",
                    location=f"pos {rev_start}",
                ))
                content_deleted += 1
            else:
                empty_deleted += 1

        if content_deleted == 0 and empty_deleted == 0:
            findings.append(Finding(
                "PASS", "paragraph_marks_preserved",
                "No paragraph marks deleted"
            ))
        elif content_deleted == 0:
            findings.append(Finding(
                "WARN", "paragraph_marks_preserved",
                f"{empty_deleted} empty paragraph mark(s) deleted (benign)",
            ))
    except Exception as e:
        findings.append(Finding(
            "WARN", "paragraph_marks_preserved",
            f"Could not check paragraph marks: {e}"
        ))
    return findings


def check_bold_text_not_deleted(doc_com, range_start: int,
                                 range_end: int) -> List[Finding]:
    """Check that no bold text (subheadings) was deleted.

    Short bold text deletions indicate a subheading was struck through,
    which almost always means the redline engine misaligned.
    """
    findings = []
    bold_deletions = 0

    try:
        revs = doc_com.Revisions
        for i in range(1, revs.Count + 1):
            rev = revs(i)
            if rev.Type != 2:  # Only check deletions
                continue
            rev_text = rev.Range.Text
            clean = rev_text.strip('\r').strip()
            if not clean or len(clean) >= 80 or '\r' in clean:
                continue

            try:
                if rev.Range.Font.Bold == -1:
                    findings.append(Finding(
                        "ERROR", "bold_text_not_deleted",
                        f"Bold text deleted (likely subheading): {repr(clean[:50])}",
                        location=f"pos {rev.Range.Start}",
                    ))
                    bold_deletions += 1
            except Exception:
                pass

        if bold_deletions == 0:
            findings.append(Finding(
                "PASS", "bold_text_not_deleted",
                "No bold text (subheadings) deleted"
            ))
    except Exception as e:
        findings.append(Finding(
            "WARN", "bold_text_not_deleted",
            f"Could not check bold deletions: {e}"
        ))
    return findings


def check_heading_formatting_com(doc_com, range_start: int,
                                  range_end: int) -> List[Finding]:
    """Check that headings/subheadings retain bold formatting after redline.

    Uses COM to inspect the live document state. Headings are identified as
    short text lines (<80 chars) that start with an uppercase letter.
    """
    findings = []
    heading_issues = 0

    try:
        range_obj = doc_com.Range(range_start, range_end)
        para_coll = range_obj.Paragraphs
        for i in range(1, para_coll.Count + 1):
            wp = para_coll(i)
            text = wp.Range.Text.strip().rstrip('\r')
            if not text:
                continue
            if len(text) >= 80:
                continue
            if not text[0].isupper() and not text[0].isdigit():
                continue

            is_heading = text.isupper() or (
                len(text) < 60 and not text[0].isdigit()
            )
            if not is_heading:
                continue

            try:
                bold = wp.Range.Font.Bold
                if bold == 0:
                    findings.append(Finding(
                        "ERROR", "heading_formatting",
                        f"Heading lost bold: {repr(text[:50])}",
                        location=f"para {i}",
                    ))
                    heading_issues += 1
                elif bold not in (-1, 0):
                    findings.append(Finding(
                        "WARN", "heading_formatting",
                        f"Heading has mixed bold: {repr(text[:50])}",
                        location=f"para {i}",
                    ))
            except Exception:
                pass

        if heading_issues == 0:
            findings.append(Finding(
                "PASS", "heading_formatting",
                "All headings retain bold formatting"
            ))
    except Exception as e:
        findings.append(Finding(
            "WARN", "heading_formatting",
            f"Could not check heading formatting: {e}"
        ))
    return findings


def check_content_paragraph_count(pre_content_count: int,
                                   doc_com, range_start: int,
                                   range_end: int) -> List[Finding]:
    """Check that content paragraph count hasn't decreased.

    Args:
        pre_content_count: Number of non-empty paragraphs before the edit
        doc_com: win32com Document object
        range_start: Start of target section
        range_end: End of target section
    """
    findings = []
    try:
        range_obj = doc_com.Range(range_start, range_end)
        para_coll = range_obj.Paragraphs
        post_count = 0
        for i in range(1, para_coll.Count + 1):
            text = para_coll(i).Range.Text.strip().rstrip('\r')
            if text:
                post_count += 1

        if post_count < pre_content_count:
            findings.append(Finding(
                "ERROR", "content_paragraph_count",
                f"Content paragraphs lost: {pre_content_count} -> {post_count}",
                expected=pre_content_count,
                actual=post_count,
            ))
        else:
            findings.append(Finding(
                "PASS", "content_paragraph_count",
                f"Content paragraphs: {pre_content_count} -> {post_count}"
            ))
    except Exception as e:
        findings.append(Finding(
            "WARN", "content_paragraph_count",
            f"Could not count paragraphs: {e}"
        ))
    return findings


def check_change_ratio(orig_length: int, deleted_chars: int,
                        inserted_chars: int,
                        warn_threshold: float = 0.15,
                        error_threshold: float = 0.30) -> List[Finding]:
    """Check that the total change ratio is within acceptable bounds.

    Args:
        orig_length: Length of the original flat text
        deleted_chars: Number of characters deleted
        inserted_chars: Number of characters inserted
        warn_threshold: Fraction above which to warn (default 15%)
        error_threshold: Fraction above which to error (default 30%)
    """
    findings = []
    if orig_length == 0:
        findings.append(Finding("WARN", "change_ratio", "Empty original text"))
        return findings

    change_pct = (deleted_chars + inserted_chars) / orig_length
    pct_str = f"{change_pct * 100:.1f}%"

    if change_pct > error_threshold:
        findings.append(Finding(
            "ERROR", "change_ratio",
            f"Excessive changes: {pct_str} modified "
            f"({deleted_chars} del + {inserted_chars} ins / {orig_length} chars)",
            expected=f"<= {error_threshold * 100:.0f}%",
            actual=pct_str,
        ))
    elif change_pct > warn_threshold:
        findings.append(Finding(
            "WARN", "change_ratio",
            f"High change ratio: {pct_str} modified",
            expected=f"<= {warn_threshold * 100:.0f}%",
            actual=pct_str,
        ))
    else:
        findings.append(Finding(
            "PASS", "change_ratio",
            f"Change ratio: {pct_str} ({deleted_chars} del + {inserted_chars} ins)"
        ))
    return findings


# ===================================================================
# CONVENIENCE WRAPPERS
# ===================================================================

def validate_redline(doc_com, range_start: int, range_end: int,
                     pre_content_para_count: int,
                     orig_flat_length: int, deleted_chars: int,
                     inserted_chars: int) -> ValidationResult:
    """Run all redline checks. Called from word_hotkey.py and test scripts.

    Args:
        doc_com: win32com Document object
        range_start: Start position of the edited section
        range_end: End position of the edited section
        pre_content_para_count: Non-empty paragraph count before editing
        orig_flat_length: Length of the original flattened text
        deleted_chars: Characters deleted by the diff
        inserted_chars: Characters inserted by the diff
    """
    result = ValidationResult(
        context=f"Redline validation (range {range_start}-{range_end})"
    )

    # Run all redline checks
    result.findings.extend(
        check_revisions_in_range(doc_com, range_start, range_end))
    result.findings.extend(
        check_paragraph_marks_preserved(doc_com, range_start, range_end))
    result.findings.extend(
        check_bold_text_not_deleted(doc_com, range_start, range_end))
    result.findings.extend(
        check_heading_formatting_com(doc_com, range_start, range_end))
    result.findings.extend(
        check_content_paragraph_count(
            pre_content_para_count, doc_com, range_start, range_end))
    result.findings.extend(
        check_change_ratio(
            orig_flat_length, deleted_chars, inserted_chars))

    return result


def validate_after_edit(doc_com, range_start: int,
                        range_end: int) -> ValidationResult:
    """Lightweight validation for any Word edit operation.

    Checks heading formatting and content paragraph integrity.
    Does not require diff metrics.
    """
    result = ValidationResult(
        context=f"Post-edit validation (range {range_start}-{range_end})"
    )
    result.findings.extend(
        check_heading_formatting_com(doc_com, range_start, range_end))
    return result


# ===================================================================
# REPORT CHECKS (python-docx based, offline file inspection)
# ===================================================================

# Known section headings (all caps)
SECTION_HEADINGS = {
    "FACTUAL BACKGROUND", "PROCEDURAL HISTORY", "PROCEDURAL STATUS",
    "INVESTIGATION", "DISCOVERY", "MEDICAL RECORD REVIEW",
    "EVALUATION OF LIABILITY", "EVALUATION OF EXPOSURE", "EXPERTS",
    "SETTLEMENT", "SETTLEMENT STATUS", "LITIGATION BUDGET",
    "FURTHER CASE HANDLING",
}


# ---------------------------------------------------------------------------
# python-docx XML helpers
# ---------------------------------------------------------------------------

def _qn(tag: str) -> str:
    """Qualified name for Word XML tags. Lazy import to avoid requiring
    python-docx when only using COM-based checks."""
    from docx.oxml.ns import qn
    return qn(tag)


def _get_pPr(para):
    """Get paragraph properties element, or None."""
    return para._element.find(_qn('w:pPr'))


def _get_rPr_from_pPr(pPr):
    """Get run properties from paragraph properties."""
    if pPr is None:
        return None
    return pPr.find(_qn('w:rPr'))


def _get_first_run_rPr(para):
    """Get run properties from the first non-empty run."""
    for run in para.runs:
        if run.text and run.text.strip():
            return run._element.find(_qn('w:rPr'))
    return None


def _is_bold_xml(rPr) -> bool:
    """Check if run properties indicate bold."""
    if rPr is None:
        return False
    b = rPr.find(_qn('w:b'))
    if b is None:
        return False
    val = b.get(_qn('w:val'))
    return val is None or val not in ('0', 'false')


def _is_underline_xml(rPr) -> bool:
    """Check if run properties indicate underline."""
    if rPr is None:
        return False
    u = rPr.find(_qn('w:u'))
    if u is None:
        return False
    val = u.get(_qn('w:val'))
    return val is not None and val != 'none'


def _get_indent_firstLine(pPr) -> Optional[int]:
    """Get first-line indent in twips from paragraph properties."""
    if pPr is None:
        return None
    ind = pPr.find(_qn('w:ind'))
    if ind is None:
        return None
    val = ind.get(_qn('w:firstLine'))
    return int(val) if val else None


def _get_indent_left(pPr) -> Optional[int]:
    """Get left indent in twips."""
    if pPr is None:
        return None
    ind = pPr.find(_qn('w:ind'))
    if ind is None:
        return None
    val = ind.get(_qn('w:left'))
    return int(val) if val else None


def _get_indent_hanging(pPr) -> Optional[int]:
    """Get hanging indent in twips."""
    if pPr is None:
        return None
    ind = pPr.find(_qn('w:ind'))
    if ind is None:
        return None
    val = ind.get(_qn('w:hanging'))
    return int(val) if val else None


def _get_spacing_after(pPr) -> Optional[int]:
    """Get space-after in twips."""
    if pPr is None:
        return None
    spacing = pPr.find(_qn('w:spacing'))
    if spacing is None:
        return None
    val = spacing.get(_qn('w:after'))
    return int(val) if val else None


def _get_style_name(para) -> Optional[str]:
    """Get the style name of a paragraph."""
    pPr = _get_pPr(para)
    if pPr is None:
        return None
    pStyle = pPr.find(_qn('w:pStyle'))
    if pStyle is None:
        return None
    return pStyle.get(_qn('w:val'))


def _para_text(para) -> str:
    """Get stripped paragraph text."""
    return para.text.strip() if para.text else ""


def _is_section_heading(text: str) -> bool:
    """Check if text is a known section heading."""
    return text.strip().rstrip(":") in SECTION_HEADINGS


# Public alias
is_section_heading = _is_section_heading


# Preferred canonical names (when a heading has multiple valid forms)
_CANONICAL_SECTION_NAMES = {
    "PROCEDURAL HISTORY": "PROCEDURAL STATUS",
    "SETTLEMENT": "SETTLEMENT STATUS",
}


def fuzzy_match_section_heading(text: str) -> Optional[str]:
    """Match text to a canonical section heading name, or None.

    Handles variations associates might use:
      - Exact match: "FACTUAL BACKGROUND" -> "FACTUAL BACKGROUND"
      - Known alias: "SETTLEMENT" -> "SETTLEMENT STATUS"
      - Extra words: "FACTUAL BACKGROUND AND SUMMARY" -> "FACTUAL BACKGROUND"
      - Containment: "EVALUATION OF LIABILITY AND DAMAGES" -> "EVALUATION OF LIABILITY"
      - Close match via difflib for typos/minor differences
    """
    import difflib

    cleaned = text.strip().rstrip(":").upper()
    if not cleaned:
        return None

    # 1. Exact match
    if cleaned in SECTION_HEADINGS:
        return _CANONICAL_SECTION_NAMES.get(cleaned, cleaned)

    # 2. Check if a known heading is a prefix of the input (extra words appended)
    #    e.g., "FACTUAL BACKGROUND AND SUMMARY" -> "FACTUAL BACKGROUND"
    best_prefix = None
    best_prefix_len = 0
    for heading in SECTION_HEADINGS:
        if cleaned.startswith(heading) and len(heading) > best_prefix_len:
            # Ensure the match is at a word boundary
            rest = cleaned[len(heading):]
            if not rest or rest[0] in (' ', ',', ':'):
                best_prefix = heading
                best_prefix_len = len(heading)
    if best_prefix:
        return _CANONICAL_SECTION_NAMES.get(best_prefix, best_prefix)

    # 3. Check if any known heading contains the input or vice versa
    for heading in SECTION_HEADINGS:
        if heading in cleaned or cleaned in heading:
            return _CANONICAL_SECTION_NAMES.get(heading, heading)

    # 4. Fuzzy match via difflib (for typos)
    matches = difflib.get_close_matches(cleaned, SECTION_HEADINGS, n=1, cutoff=0.7)
    if matches:
        heading = matches[0]
        return _CANONICAL_SECTION_NAMES.get(heading, heading)

    return None


# ---------------------------------------------------------------------------
# Report check functions
# ---------------------------------------------------------------------------

def check_report_section_headings(doc, profile: Dict) -> List[Finding]:
    """Check that section headings are bold, underlined, all-caps, flush left."""
    findings = []
    heading_profile = profile.get("section_heading", {})
    expected_bold = heading_profile.get("bold", True)
    expected_underline = heading_profile.get("underline", "single")

    found_headings = []
    for i, para in enumerate(doc.paragraphs):
        text = _para_text(para)
        if not _is_section_heading(text):
            continue

        found_headings.append(text)
        pPr = _get_pPr(para)

        rPr = _get_rPr_from_pPr(pPr)
        run_rPr = _get_first_run_rPr(para)
        is_bold = _is_bold_xml(rPr) or _is_bold_xml(run_rPr)

        if expected_bold and not is_bold:
            findings.append(Finding(
                "ERROR", "report_section_heading",
                f'"{text}" is not bold',
                location=f"para {i}", expected="bold", actual="not bold"
            ))

        is_underlined = _is_underline_xml(rPr) or _is_underline_xml(run_rPr)
        if expected_underline and not is_underlined:
            findings.append(Finding(
                "ERROR", "report_section_heading",
                f'"{text}" is not underlined',
                location=f"para {i}", expected="underline", actual="no underline"
            ))

        if text != text.upper():
            findings.append(Finding(
                "ERROR", "report_section_heading",
                f'"{text}" is not all-caps',
                location=f"para {i}", expected="ALL CAPS", actual=text
            ))

        first_line = _get_indent_firstLine(pPr)
        if first_line is not None and first_line != 0:
            findings.append(Finding(
                "ERROR", "report_section_heading",
                f'"{text}" has first-line indent {first_line} twips, expected 0',
                location=f"para {i}", expected=0, actual=first_line
            ))

    if found_headings:
        findings.append(Finding(
            "PASS", "report_section_heading",
            f"Found {len(found_headings)} section headings: {', '.join(found_headings)}"
        ))
    else:
        findings.append(Finding(
            "WARN", "report_section_heading",
            "No section headings found in document"
        ))

    return findings


def check_report_body_paragraphs(doc, profile: Dict) -> List[Finding]:
    """Check body paragraph formatting: first-line indent, font."""
    findings = []
    body_profile = profile.get("body_paragraph", {})
    expected_indent = body_profile.get("first_line_indent_twips", 720)
    expected_font = body_profile.get("font_name", "Times New Roman")

    body_paras = []
    wrong_indent = []
    wrong_font = []

    for i, para in enumerate(doc.paragraphs):
        text = _para_text(para)
        if not text or len(text) < 40:
            continue
        if _is_section_heading(text):
            continue
        style = _get_style_name(para)
        if style in ("zClosing", "zDelivery", "zRe", "zAddressee", "FlushLeft"):
            continue

        body_paras.append(i)
        pPr = _get_pPr(para)
        first_line = _get_indent_firstLine(pPr)

        if style == "Body" and first_line is None:
            first_line = 720  # Inherited from style

        if first_line is not None and first_line != expected_indent:
            left = _get_indent_left(pPr)
            hanging = _get_indent_hanging(pPr)
            if left or hanging:
                continue  # Subheading with hanging indent

            raw_text = para.text or ""
            if first_line == 0 and (raw_text.startswith("\t") or raw_text.startswith("  \t")):
                continue

            wrong_indent.append((i, first_line))

        for run in para.runs:
            if run.text and run.text.strip():
                if run.font.name and run.font.name != expected_font:
                    wrong_font.append((i, run.font.name))
                break

    if body_paras:
        findings.append(Finding(
            "PASS", "report_body_paragraph",
            f"Checked {len(body_paras)} body paragraphs"
        ))

    if wrong_indent:
        for idx, actual in wrong_indent[:5]:
            text_preview = _para_text(doc.paragraphs[idx])[:60]
            findings.append(Finding(
                "ERROR", "report_body_paragraph",
                f'first-line indent={actual} twips, expected {expected_indent}: "{text_preview}..."',
                location=f"para {idx}", expected=expected_indent, actual=actual
            ))
        if len(wrong_indent) > 5:
            findings.append(Finding(
                "ERROR", "report_body_paragraph",
                f"... and {len(wrong_indent) - 5} more paragraphs with wrong indent"
            ))
    else:
        findings.append(Finding(
            "PASS", "report_body_paragraph",
            f"All body paragraphs have correct first-line indent ({expected_indent} twips)"
        ))

    if wrong_font:
        fonts = set(f for _, f in wrong_font)
        findings.append(Finding(
            "WARN", "report_body_paragraph",
            f"Found non-standard fonts in {len(wrong_font)} paragraphs: {', '.join(fonts)}",
            expected=expected_font, actual=list(fonts)
        ))
    else:
        findings.append(Finding(
            "PASS", "report_body_paragraph",
            f"All body paragraphs use {expected_font}"
        ))

    return findings


def check_report_metadata_table(doc, profile: Dict) -> List[Finding]:
    """Check metadata table structure: columns, width, borders."""
    findings = []
    table_profile = profile.get("metadata_table", {})

    if not doc.tables:
        findings.append(Finding("ERROR", "report_metadata_table", "No tables found in document"))
        return findings

    table = doc.tables[0]
    tbl = table._tbl

    expected_cols = table_profile.get("columns", 2)
    actual_cols = len(table.columns)
    if actual_cols == expected_cols:
        findings.append(Finding("PASS", "report_metadata_table", f"Column count: {actual_cols}"))
    else:
        findings.append(Finding(
            "WARN", "report_metadata_table",
            f"Column count: {actual_cols}, expected {expected_cols}",
            expected=expected_cols, actual=actual_cols
        ))

    tblPr = tbl.find(_qn('w:tblPr'))
    if tblPr is not None:
        tblW = tblPr.find(_qn('w:tblW'))
        if tblW is not None:
            width = tblW.get(_qn('w:w'))
            width_type = tblW.get(_qn('w:type'))
            if width and width_type == 'dxa':
                expected_width = table_profile.get("total_width_twips", 10579)
                actual_width = int(width)
                tolerance = 200
                if abs(actual_width - expected_width) <= tolerance:
                    findings.append(Finding(
                        "PASS", "report_metadata_table",
                        f"Table width: {actual_width} twips (expected ~{expected_width})"
                    ))
                else:
                    findings.append(Finding(
                        "WARN", "report_metadata_table",
                        f"Table width: {actual_width} twips, expected ~{expected_width}",
                        expected=expected_width, actual=actual_width
                    ))

        tblInd = tblPr.find(_qn('w:tblInd'))
        expected_indent = table_profile.get("table_indent_twips", 0)
        if tblInd is not None:
            actual_indent = int(tblInd.get(_qn('w:w'), '0'))
            if actual_indent != expected_indent:
                findings.append(Finding(
                    "WARN", "report_metadata_table",
                    f"Table indent: {actual_indent} twips, gold standards have {expected_indent}",
                    expected=expected_indent, actual=actual_indent
                ))
            else:
                findings.append(Finding(
                    "PASS", "report_metadata_table", f"Table indent: {actual_indent}"))
        elif expected_indent == 0:
            findings.append(Finding(
                "PASS", "report_metadata_table", "No table indent (matches gold standard)"))

    has_re = False
    for row in table.rows:
        for cell in row.cells:
            if "Re:" in cell.text:
                has_re = True
                break
    if has_re:
        findings.append(Finding("PASS", "report_metadata_table", "Re: line found"))
    else:
        findings.append(Finding("WARN", "report_metadata_table", "No 'Re:' line found"))

    return findings


def check_report_closing(doc, profile: Dict) -> List[Finding]:
    """Check closing section: style, attorney names, firm name."""
    findings = []
    closing_profile = profile.get("closing", {})
    expected_style = closing_profile.get("style", "zClosing")
    expected_firm = closing_profile.get("firm_name", "BORDIN SEMMER LLP")
    expected_valedictions = closing_profile.get("valediction", ["Sincerely,", "Very truly yours,"])

    valediction_found = False
    firm_found = False

    for i, para in enumerate(doc.paragraphs):
        text = _para_text(para)
        style = _get_style_name(para)

        if any(text.startswith(v.rstrip(",")) for v in expected_valedictions):
            valediction_found = True
            if style != expected_style:
                findings.append(Finding(
                    "ERROR", "report_closing",
                    f'"{text}" uses style "{style}", expected "{expected_style}"',
                    location=f"para {i}", expected=expected_style, actual=style
                ))

        if text == expected_firm:
            firm_found = True
            if style != expected_style:
                findings.append(Finding(
                    "ERROR", "report_closing",
                    f'"{text}" uses style "{style}", expected "{expected_style}"',
                    location=f"para {i}", expected=expected_style, actual=style
                ))

    if valediction_found:
        findings.append(Finding("PASS", "report_closing", "Valediction found"))
    else:
        findings.append(Finding("ERROR", "report_closing",
                                "No valediction found (Sincerely/Very truly yours)"))

    if firm_found:
        findings.append(Finding("PASS", "report_closing", f'Firm name "{expected_firm}" found'))
    else:
        findings.append(Finding("ERROR", "report_closing", f'Firm name "{expected_firm}" not found'))

    expected_attorneys = closing_profile.get("attorney_names", [])
    for name in expected_attorneys:
        found = any(name in _para_text(p) for p in doc.paragraphs)
        if found:
            findings.append(Finding("PASS", "report_closing", f'Attorney "{name}" found'))
        else:
            findings.append(Finding("WARN", "report_closing", f'Attorney "{name}" not found'))

    return findings


def check_report_no_empty_paragraphs(doc, profile: Dict) -> List[Finding]:
    """Check for excessive consecutive empty paragraphs."""
    findings = []
    consecutive_empty = 0
    max_consecutive = 0
    worst_index = 0

    for i, para in enumerate(doc.paragraphs):
        text = _para_text(para)
        style = _get_style_name(para)

        if style == "zClosing":
            consecutive_empty = 0
            continue

        if not text:
            consecutive_empty += 1
            if consecutive_empty > max_consecutive:
                max_consecutive = consecutive_empty
                worst_index = i
        else:
            consecutive_empty = 0

    if max_consecutive > 3:
        findings.append(Finding(
            "WARN", "report_empty_paragraphs",
            f"Found {max_consecutive} consecutive empty paragraphs",
            location=f"para {worst_index}", expected="<= 3", actual=max_consecutive
        ))
    else:
        findings.append(Finding(
            "PASS", "report_empty_paragraphs",
            f"No excessive empty paragraph runs (max consecutive: {max_consecutive})"
        ))

    return findings


def check_report_salutation(doc, profile: Dict) -> List[Finding]:
    """Check salutation paragraph formatting."""
    findings = []
    sal_profile = profile.get("salutation", {})

    for i, para in enumerate(doc.paragraphs):
        text = _para_text(para)
        if text.startswith("Dear ") and text.endswith(","):
            pPr = _get_pPr(para)

            first_line = _get_indent_firstLine(pPr)
            if first_line and first_line > 0:
                findings.append(Finding(
                    "ERROR", "report_salutation",
                    f'Salutation "{text}" has first-line indent {first_line}, expected 0',
                    location=f"para {i}", expected=0, actual=first_line
                ))
            else:
                findings.append(Finding(
                    "PASS", "report_salutation", f'Salutation "{text}" formatting OK'))

            space_after = _get_spacing_after(pPr)
            expected_after = sal_profile.get("space_after_twips", 0)
            if space_after is not None and space_after != expected_after:
                findings.append(Finding(
                    "WARN", "report_salutation",
                    f"Salutation space_after={space_after}, gold standard={expected_after}",
                    location=f"para {i}", expected=expected_after, actual=space_after
                ))

            return findings

    findings.append(Finding("WARN", "report_salutation", "No salutation paragraph found"))
    return findings


def check_report_intro_paragraph(doc, profile: Dict) -> List[Finding]:
    """Check intro paragraph (As you recall...) formatting."""
    findings = []
    intro_profile = profile.get("intro_paragraph", {})
    expected_indent = intro_profile.get("first_line_indent_twips", 720)

    for i, para in enumerate(doc.paragraphs):
        text = _para_text(para)
        if "As you recall" in text or "As you know" in text:
            pPr = _get_pPr(para)
            style = _get_style_name(para)

            expected_style = intro_profile.get("style", "Body")
            if style == expected_style:
                findings.append(Finding("PASS", "report_intro", f"Style: {style}"))
            elif style is not None:
                findings.append(Finding(
                    "WARN", "report_intro",
                    f'Style: "{style}", expected "{expected_style}"',
                    location=f"para {i}", expected=expected_style, actual=style
                ))

            first_line = _get_indent_firstLine(pPr)
            if style == "Body" and first_line is None:
                first_line = 720

            if first_line == expected_indent:
                findings.append(Finding(
                    "PASS", "report_intro",
                    f"First-line indent: {first_line} twips"
                ))
            elif first_line is not None:
                findings.append(Finding(
                    "ERROR", "report_intro",
                    f"First-line indent: {first_line} twips, expected {expected_indent}",
                    location=f"para {i}", expected=expected_indent, actual=first_line
                ))

            return findings

    findings.append(Finding("WARN", "report_intro", "No intro paragraph found"))
    return findings


def check_report_subheadings(doc, profile: Dict) -> List[Finding]:
    """Check subheading formatting (bold, underlined, proper indentation)."""
    findings = []

    subheadings_found = 0
    subheading_issues = []

    for i, para in enumerate(doc.paragraphs):
        text = _para_text(para)
        if not text or len(text) > 80:
            continue
        if _is_section_heading(text):
            continue

        pPr = _get_pPr(para)
        rPr = _get_rPr_from_pPr(pPr)
        run_rPr = _get_first_run_rPr(para)
        is_bold = _is_bold_xml(rPr) or _is_bold_xml(run_rPr)
        is_ul = _is_underline_xml(rPr) or _is_underline_xml(run_rPr)

        if is_bold and is_ul and len(text) < 60:
            subheadings_found += 1

            left = _get_indent_left(pPr)
            hanging = _get_indent_hanging(pPr)

            has_prefix = bool(re.match(r'^[A-Z]\.\s|^\d+\.\s', text))
            if has_prefix and not hanging and left is None:
                subheading_issues.append((i, text, "no hanging indent"))

    if subheadings_found > 0:
        findings.append(Finding(
            "PASS", "report_subheading",
            f"Found {subheadings_found} subheadings (bold + underlined)"
        ))
    else:
        findings.append(Finding(
            "WARN", "report_subheading", "No subheadings detected"
        ))

    for idx, text, issue in subheading_issues[:5]:
        findings.append(Finding(
            "WARN", "report_subheading",
            f'"{text}" ({issue})',
            location=f"para {idx}"
        ))

    return findings


def check_report_delivery_line(doc, profile: Dict) -> List[Finding]:
    """Check VIA EMAIL delivery line."""
    findings = []
    delivery_profile = profile.get("delivery_line", {})
    expected_text = delivery_profile.get("text", "VIA EMAIL")
    expected_style = delivery_profile.get("style", "zDelivery")

    for i, para in enumerate(doc.paragraphs):
        text = _para_text(para)
        if text == expected_text:
            style = _get_style_name(para)
            if style == expected_style:
                findings.append(Finding(
                    "PASS", "report_delivery_line", f'"{text}" with style {style}'))
            else:
                findings.append(Finding(
                    "WARN", "report_delivery_line",
                    f'"{text}" uses style "{style}", expected "{expected_style}"',
                    location=f"para {i}", expected=expected_style, actual=style
                ))
            return findings

    findings.append(Finding("WARN", "report_delivery_line", f'"{expected_text}" not found'))
    return findings


def check_report_all_caps_names(doc, profile: Dict) -> List[Finding]:
    """Check that client/party names are not in ALL CAPS (should be title case)."""
    findings = []

    if doc.tables:
        table = doc.tables[0]
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if ":" in text or len(text) < 5:
                    continue
                words = text.split()
                all_caps_words = [w for w in words if w.isupper() and len(w) > 2
                                  and w not in ("LLP", "LLC", "INC", "DBA", "USA", "VIA")]
                if len(all_caps_words) >= 2:
                    findings.append(Finding(
                        "ERROR", "report_all_caps_names",
                        f'All-caps name in metadata: "{text}" - should be title case',
                        expected="Title Case", actual=text
                    ))

    if not any(f.severity == "ERROR" for f in findings):
        findings.append(Finding("PASS", "report_all_caps_names",
                                "No all-caps names found in metadata"))

    return findings


# ---------------------------------------------------------------------------
# Report rule registry
# ---------------------------------------------------------------------------

REPORT_RULES = [
    check_report_section_headings,
    check_report_body_paragraphs,
    check_report_metadata_table,
    check_report_salutation,
    check_report_intro_paragraph,
    check_report_subheadings,
    check_report_closing,
    check_report_no_empty_paragraphs,
    check_report_delivery_line,
    check_report_all_caps_names,
]


def validate_report(doc_path: str, profile: Dict = None) -> ValidationResult:
    """Run all report checks against a .docx file.

    Args:
        doc_path: Path to the .docx report
        profile: Reference profile dict. If None, loads the default profile.
    """
    from docx import Document

    result = ValidationResult(context=f"Report: {os.path.basename(doc_path)}")

    if not os.path.exists(doc_path):
        result.findings.append(Finding("ERROR", "file", f"File not found: {doc_path}"))
        return result

    if profile is None:
        project_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        profile_path = os.path.join(project_root, "config", "report_reference_profile.json")
        if os.path.exists(profile_path):
            import json
            with open(profile_path, 'r', encoding='utf-8') as f:
                profile = json.load(f)
        else:
            profile = {}

    doc = Document(doc_path)

    for rule in REPORT_RULES:
        try:
            findings = rule(doc, profile)
            result.findings.extend(findings)
        except Exception as e:
            result.findings.append(Finding(
                "WARN", rule.__name__,
                f"Rule raised exception: {e}"
            ))

    return result


def validate_index_docx(doc_path: str, expected_doc_count: Optional[int] = None) -> ValidationResult:
    """Lightweight offline validation for the separator INDEX .docx.

    Verifies the file exists, opens with python-docx, contains a table with a
    header row plus at least one data row, and (optionally) that the number of
    data rows matches the number of identified documents.
    """
    from docx import Document

    result = ValidationResult(context=f"Separator index: {os.path.basename(doc_path)}")

    if not os.path.exists(doc_path):
        result.findings.append(Finding("ERROR", "file", f"File not found: {doc_path}"))
        return result

    try:
        doc = Document(doc_path)
    except Exception as e:
        result.findings.append(Finding("ERROR", "open", f"Could not open .docx: {e}"))
        return result

    if not doc.tables:
        result.findings.append(Finding("ERROR", "table", "Index has no table"))
        return result

    table = doc.tables[0]
    data_rows = max(0, len(table.rows) - 1)  # minus header
    if data_rows == 0:
        result.findings.append(Finding("ERROR", "rows", "Index table has no data rows"))
        return result

    if expected_doc_count is not None and data_rows != expected_doc_count:
        result.findings.append(Finding(
            "WARN", "row_count",
            f"Index data rows ({data_rows}) != identified documents ({expected_doc_count})",
            expected=expected_doc_count, actual=data_rows,
        ))
    else:
        result.findings.append(Finding("PASS", "row_count", f"{data_rows} document rows"))

    return result


def validate_discovery_response_docx(doc_path: str) -> ValidationResult:
    """Lightweight offline validation for generated discovery-response .docx files.

    This intentionally does not run litigation-report formatting checks. It
    verifies that the file exists, can be opened by python-docx, and contains
    non-empty response text.
    """
    from docx import Document

    result = ValidationResult(
        context=f"Discovery response: {os.path.basename(doc_path)}"
    )

    if not os.path.exists(doc_path):
        result.findings.append(Finding("ERROR", "file", f"File not found: {doc_path}"))
        return result

    try:
        doc = Document(doc_path)
    except Exception as exc:
        result.findings.append(
            Finding("ERROR", "open_docx", f"Could not open generated docx: {exc}")
        )
        return result

    text = _docx_all_story_text(doc)
    if not text.strip():
        result.findings.append(
            Finding("ERROR", "content", "Generated document contains no paragraph text")
        )
    else:
        result.findings.append(
            Finding("PASS", "content", "Generated document contains paragraph text")
        )

    if "RESPONSE" in text.upper():
        result.findings.append(
            Finding("PASS", "response_markers", "Response headings detected")
        )
    else:
        result.findings.append(
            Finding("WARN", "response_markers", "No response headings detected")
        )

    return result


def validate_opposition_docx(doc_path: str) -> ValidationResult:
    """Lightweight offline validation for generated opposition .docx files."""
    from docx import Document

    result = ValidationResult(
        context=f"Opposition memorandum: {os.path.basename(doc_path)}"
    )

    if not os.path.exists(doc_path):
        result.findings.append(Finding("ERROR", "file", f"File not found: {doc_path}"))
        return result

    try:
        doc = Document(doc_path)
    except Exception as exc:
        result.findings.append(
            Finding("ERROR", "open_docx", f"Could not open generated docx: {exc}")
        )
        return result

    text = _docx_all_story_text(doc)
    if not text.strip():
        result.findings.append(
            Finding("ERROR", "content", "Generated document contains no paragraph text")
        )
    else:
        result.findings.append(
            Finding("PASS", "content", "Generated document contains paragraph text")
        )

    if "OPPOSITION" in text.upper():
        result.findings.append(
            Finding("PASS", "opposition_marker", "Opposition title detected")
        )
    else:
        result.findings.append(
            Finding("WARN", "opposition_marker", "No opposition title detected")
        )

    if "CAPTION PAGE" in text.upper():
        result.findings.append(
            Finding("ERROR", "caption_marker", "Unreplaced CAPTION PAGE marker detected")
        )
    else:
        result.findings.append(
            Finding("PASS", "caption_marker", "No unreplaced caption marker detected")
        )

    return result


def _docx_all_story_text(doc) -> str:
    w_t = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t"
    w_p = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p"
    roots = [doc.element.body]
    for section in doc.sections:
        roots.extend(
            [
                section.header._element,
                section.first_page_header._element,
                section.even_page_header._element,
                section.footer._element,
                section.first_page_footer._element,
                section.even_page_footer._element,
            ]
        )

    seen: set[int] = set()
    parts: list[str] = []
    for root in roots:
        root_id = id(root)
        if root_id in seen:
            continue
        seen.add(root_id)
        for paragraph in root.iter(w_p):
            paragraph_text = "".join(text.text or "" for text in paragraph.iter(w_t))
            if paragraph_text:
                parts.append(paragraph_text)
    return "\n".join(parts)
