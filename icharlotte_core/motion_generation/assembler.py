"""Word assembly for Generate Motion previews.

Produces a Notice of Motion + Memorandum of Points & Authorities, with per-type
attachments emitted as clearly-labeled placeholders. Reuses the opposition
assembler's docx rendering helpers and the project Word validator.
"""
from __future__ import annotations

import os

from docx import Document

from icharlotte_core.opposition.assembler import (
    _add_heading_paragraph,
    _add_paragraphs,
    _add_title,
    _replace_caption_markers,
    _same_path,
)
from icharlotte_core.opposition.models import DraftDocument
from icharlotte_core.word_validator import validate_opposition_docx

from .config import MotionTypeConfig


def _add_placeholder(doc, label: str) -> None:
    _add_heading_paragraph(
        doc, level=1, text=f"[ATTACHMENT — TO BE COMPLETED] {label}"
    )
    para = doc.add_paragraph()
    para.add_run(
        f"[This {label} must be completed before filing.]"
    ).italic = True


def assemble_motion_preview(
    *,
    draft: DraftDocument,
    output_path: str,
    config: MotionTypeConfig,
    caption_path: str = "",
) -> str:
    """Render a generated motion to a .docx preview and validate the file."""
    title = draft.title or config.display_name or "Motion"

    if caption_path and os.path.isfile(caption_path):
        if _same_path(caption_path, output_path):
            raise ValueError("Output path must differ from the caption template path")
        doc = Document(caption_path)
        if not _replace_caption_markers(doc, title):
            _add_title(doc, title)
    else:
        doc = Document()
        _add_title(doc, title)

    _add_heading_paragraph(doc, level=1, text="NOTICE OF MOTION AND MOTION")
    _add_paragraphs(doc, draft.body_text)

    for label in config.placeholder_attachments:
        _add_placeholder(doc, label)

    output_dir = os.path.dirname(output_path)
    if output_dir:
        os.makedirs(output_dir, exist_ok=True)
    doc.save(output_path)

    validation = validate_opposition_docx(output_path)
    if validation.has_errors:
        messages = "; ".join(str(finding) for finding in validation.findings)
        raise ValueError(f"Generated motion failed validation: {messages}")

    return output_path
