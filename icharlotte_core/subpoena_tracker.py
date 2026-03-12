"""
Subpoena Tracker — In-process worker for tracking subpoenas, received records,
and medical chronology coverage. Generates a hyperlinked .docx report.
"""

import os
import re
import datetime
import logging
from difflib import SequenceMatcher

import fitz  # PyMuPDF

from PySide6.QtCore import QThread, Signal

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import qn
from docx.oxml import OxmlElement
import docx.opc.constants as opc_constants

log = logging.getLogger(__name__)

# --- Regex patterns ---
ID_PATTERN = re.compile(r'(\d{3,7})[.\-](\d{1,4})')
DATE_PATTERN = re.compile(r'(\d{4})[.\-](\d{2})[.\-](\d{2})')
SEPARATOR_STRIP = re.compile(r'^[\s,_\-]+')
CNR_PATTERN = re.compile(r'\bCNR\b', re.IGNORECASE)
REPLY_OBJECTION_PATTERN = re.compile(r'(reply|objection)', re.IGNORECASE)

# Facility names that are too generic — trigger first-page PDF extraction
UNINFORMATIVE_FACILITY = {
    "pos", "file copy", "filecopy", "file_copy", "records", "unknown",
    "copy", "copies", "meds", "medical", "med", "billing", "bills",
    "radiology", "imaging", "films", "subpoena", "",
}

# Regex to find facility name after TO: / TO THE CUSTODIAN OF RECORDS on subpoena first page
_TO_CUSTODIAN_PATTERN = re.compile(
    r'TO(?:\s+THE)?\s*(?::|\s)\s*'
    r'(?:CUSTODIAN\s+OF\s+RECORDS\s*(?:OF|FOR|,|:|\n)\s*)?'
    r'([A-Z][A-Za-z0-9\s,.\-&\'()]+)',
    re.MULTILINE
)


def _extract_facility_from_pdf(pdf_path):
    """Try to extract the facility/custodian name from the first page of a subpoena PDF.

    Looks for patterns like:
      - "TO THE CUSTODIAN OF RECORDS OF <facility>"
      - "TO: <facility>"

    Returns:
        Facility name string, or None if extraction fails.
    """
    try:
        doc = fitz.open(pdf_path)
        if len(doc) == 0:
            doc.close()
            return None
        text = doc[0].get_text()
        doc.close()
    except Exception:
        return None

    if not text or len(text.strip()) < 20:
        return None

    match = _TO_CUSTODIAN_PATTERN.search(text)
    if not match:
        return None

    raw = match.group(1).strip()
    # Clean up: take lines until we hit a line that looks like an address or instruction
    lines = raw.split('\n')
    facility_lines = []
    for line in lines:
        line = line.strip()
        if not line:
            break
        # Stop at address-like lines (starts with digit) or legal boilerplate
        if re.match(r'^\d', line):
            break
        if re.search(r'\b(YOU ARE|HEREBY|COMMANDED|ORDERED|PLEASE|ATTACHED)\b', line, re.IGNORECASE):
            break
        facility_lines.append(line)
        # Facility names are typically 1-2 lines
        if len(facility_lines) >= 2:
            break

    if not facility_lines:
        return None

    facility = ' '.join(facility_lines).strip().rstrip(',').rstrip(':').strip()

    # Reject if still too short or too generic
    if len(facility) < 3 or facility.lower() in UNINFORMATIVE_FACILITY:
        return None

    return facility


def parse_subpoena_id(filename):
    """Parse a vendor-subpoena ID and remaining text from a filename.

    Returns:
        (normalized_id, remainder) — e.g. ("62122-0001", "Riverside County Fire Department")
        (None, None) if no ID found.
    """
    basename = os.path.splitext(filename)[0]

    # Find the first match that isn't a date (e.g., 2024-05 or 2025.12)
    match = None
    for m in ID_PATTERN.finditer(basename):
        vendor_candidate = m.group(1)
        sub_candidate = m.group(2)
        # Skip if it looks like a date: 4-digit year (2000-2099) + 1-2 digit month (1-12)
        if (len(vendor_candidate) == 4 and
                2000 <= int(vendor_candidate) <= 2099 and
                len(sub_candidate) <= 2 and
                1 <= int(sub_candidate) <= 12):
            continue
        match = m
        break

    if not match:
        return None, None

    vendor = match.group(1)
    subpoena = match.group(2).zfill(4)
    normalized_id = f"{vendor}-{subpoena}"

    # Extract remainder: everything after the matched ID
    remainder = basename[match.end():]
    remainder = SEPARATOR_STRIP.sub('', remainder).strip()

    return normalized_id, remainder


def _find_folder_ci(base_path, *folder_names):
    """Walk directory tree one level at a time, matching each folder name case-insensitively.

    Args:
        base_path: Starting directory.
        folder_names: Sequence of subfolder names to match (e.g., "DISCOVERY", "Subpoenas").

    Returns:
        Resolved path string, or None if not found.
    """
    current = base_path
    for target_name in folder_names:
        target_lower = target_name.lower()
        found = None
        try:
            for entry in os.scandir(current):
                if entry.is_dir() and entry.name.lower() == target_lower:
                    found = entry.path
                    break
        except OSError:
            return None
        if found is None:
            return None
        current = found
    return current


def add_hyperlink(paragraph, url, text, color="0000FF", underline=True):
    """Add a hyperlink to a paragraph pointing to a local file path."""
    part = paragraph.part
    r_id = part.relate_to(url, opc_constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    if color:
        c = OxmlElement('w:color')
        c.set(qn('w:val'), color)
        rPr.append(c)

    if underline:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)

    f = OxmlElement('w:rFonts')
    f.set(qn('w:ascii'), 'Times New Roman')
    f.set(qn('w:hAnsi'), 'Times New Roman')
    rPr.append(f)

    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '24')  # 12pt
    rPr.append(sz)

    new_run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)
    return hyperlink


class SubpoenaTrackerWorker(QThread):
    """Background worker that scans case folders and generates a subpoena tracking report."""

    progress = Signal(str)
    finished_result = Signal(bool, str)  # (success, output_path_or_error_message)
    warning = Signal(str)

    def __init__(self, case_path, file_number="", parent=None):
        super().__init__(parent)
        self.case_path = case_path
        self.file_number = file_number

    def run(self):
        try:
            self._run_phases()
        except Exception as e:
            log.exception("Subpoena tracker failed")
            self.finished_result.emit(False, f"Unexpected error: {e}")

    def _run_phases(self):
        # --- Phase 1: Scan issued subpoenas ---
        self.progress.emit("Scanning issued subpoenas...")
        subpoenas, phase1_ok = self._scan_issued_subpoenas()
        if not phase1_ok:
            return

        # --- Phase 2: Scan received records + build file index ---
        self.progress.emit("Scanning received records...")
        received = self._scan_received_records()
        file_index = self._build_file_index()

        # --- Phase 3: Scan medical chronologies ---
        self.progress.emit("Reading medical chronologies...")
        chronologies, pages_map, non_subpoena_records = self._scan_medical_chronologies()

        # --- Phase 4: Generate report ---
        self.progress.emit("Generating report...")
        self._generate_report(subpoenas, received, chronologies, pages_map,
                              non_subpoena_records, file_index)

    # ---------------------------------------------------------------
    # Phase 1
    # ---------------------------------------------------------------
    def _scan_issued_subpoenas(self):
        """Scan DISCOVERY/Subpoenas/ for issued subpoena PDFs.

        Returns:
            (subpoenas_dict, success_bool)
        """
        folder = _find_folder_ci(self.case_path, "DISCOVERY", "Subpoenas")
        if folder is None:
            self.finished_result.emit(False, "No DISCOVERY/Subpoenas folder found.")
            return {}, False

        subpoenas = {}
        skipped = []

        for root, _dirs, files in os.walk(folder):
            for fname in files:
                if not fname.lower().endswith('.pdf'):
                    continue
                sid, remainder = parse_subpoena_id(fname)
                if sid is None:
                    skipped.append(fname)
                    continue
                facility = remainder if remainder else "Unknown"
                full_path = os.path.join(root, fname)

                # If the facility name from the filename is uninformative,
                # try to extract it from the first page of the PDF
                if facility.lower().strip() in UNINFORMATIVE_FACILITY:
                    pdf_facility = _extract_facility_from_pdf(full_path)
                    if pdf_facility:
                        self.progress.emit(f"Extracted facility from PDF: {pdf_facility}")
                        facility = pdf_facility

                if sid not in subpoenas:
                    subpoenas[sid] = {
                        "facility": facility,
                        "subpoena_path": full_path,
                    }
                else:
                    # Prefer the entry with a real facility name over uninformative ones
                    existing = subpoenas[sid]["facility"]
                    if (existing.lower().strip() in UNINFORMATIVE_FACILITY
                            and facility.lower().strip() not in UNINFORMATIVE_FACILITY):
                        subpoenas[sid]["facility"] = facility
                        subpoenas[sid]["subpoena_path"] = full_path

        for s in skipped:
            self.warning.emit(f"Could not parse subpoena ID from: {s}")

        if not subpoenas:
            self.finished_result.emit(False, "No subpoenas found in DISCOVERY/Subpoenas folder.")
            return {}, False

        return subpoenas, True

    # ---------------------------------------------------------------
    # Phase 2
    # ---------------------------------------------------------------
    def _scan_received_records(self):
        """Scan RECORDS/Subpoenaed/ for received record files and subfolders.

        Returns:
            dict mapping normalized_id -> {"status": str, "path": str}
        """
        folder = _find_folder_ci(self.case_path, "RECORDS", "Subpoenaed")
        if folder is None:
            return {}

        received = {}
        skipped = []

        # Scan direct children and one level of subfolders (vendor folders like
        # JJ Photocopy, CNR, Gemini, COMPEX contain record files with subpoena IDs).
        try:
            for entry in os.scandir(folder):
                if entry.is_file():
                    self._classify_received_item(
                        entry.name, entry.path, received, skipped, is_file=True
                    )
                elif entry.is_dir():
                    # Scan inside vendor subfolders
                    try:
                        for sub_entry in os.scandir(entry.path):
                            self._classify_received_item(
                                sub_entry.name, sub_entry.path, received, skipped,
                                is_file=sub_entry.is_file()
                            )
                    except OSError:
                        pass
        except OSError as e:
            self.warning.emit(f"Error scanning Subpoenaed folder: {e}")

        for s in skipped:
            self.warning.emit(f"Could not parse ID from received item: {s}")

        return received

    def _classify_received_item(self, name, full_path, received, skipped, is_file):
        """Classify a single received record file or folder."""
        sid, remainder = parse_subpoena_id(name)
        if sid is None:
            skipped.append(name)
            return

        # Classify based on remainder text
        if CNR_PATTERN.search(remainder) or CNR_PATTERN.search(name):
            status = "CNR"
        elif REPLY_OBJECTION_PATTERN.search(remainder) or REPLY_OBJECTION_PATTERN.search(name):
            status = "Other"
        else:
            status = "Yes"

        if sid in received:
            existing = received[sid]
            # Prefer PDF files over folders
            if is_file and not os.path.isfile(existing["path"]):
                received[sid] = {"status": status, "path": full_path}
            # Prefer actual records over CNR/Reply
            elif status == "Yes" and existing["status"] != "Yes":
                received[sid] = {"status": status, "path": full_path}
            # Otherwise keep existing
        else:
            received[sid] = {"status": status, "path": full_path}

    def _build_file_index(self):
        """Walk RECORDS/ and DISCOVERY/RESPONSES/ once to build a filename-to-path lookup.

        Used to locate non-subpoena record files for hyperlinking.
        Matches are case-insensitive and handle filenames with or without .pdf extension.

        Returns:
            dict mapping lowercase filename (without extension) -> full file path
        """
        file_index = {}

        search_folders = [
            _find_folder_ci(self.case_path, "RECORDS"),
            _find_folder_ci(self.case_path, "DISCOVERY", "RESPONSES"),
        ]

        for folder in search_folders:
            if folder is None:
                continue
            for root, _dirs, files in os.walk(folder):
                for fname in files:
                    # Index by lowercase name without extension
                    key = os.path.splitext(fname)[0].lower().strip()
                    if key and key not in file_index:
                        file_index[key] = os.path.join(root, fname)

        return file_index

    @staticmethod
    def _lookup_file(file_index, filename):
        """Look up a filename in the file index, with fuzzy fallback for near-matches.

        First tries an exact match (case-insensitive, extension-stripped).
        If that fails, uses SequenceMatcher to find the closest match above 0.85 ratio.
        This handles minor spelling differences between med chron entries and actual filenames
        (e.g., "Orthodontics" vs "Orthodonics").

        Returns:
            Full file path, or None if no match found.
        """
        lookup_key = os.path.splitext(filename)[0].lower().strip()
        # Exact match
        found = file_index.get(lookup_key)
        if found:
            return found
        # Fuzzy fallback — find best match above threshold
        best_ratio = 0.0
        best_path = None
        for indexed_key, indexed_path in file_index.items():
            ratio = SequenceMatcher(None, lookup_key, indexed_key).ratio()
            if ratio > best_ratio:
                best_ratio = ratio
                best_path = indexed_path
        if best_ratio >= 0.85:
            return best_path
        return None

    # ---------------------------------------------------------------
    # Phase 3
    # ---------------------------------------------------------------
    def _scan_medical_chronologies(self):
        """Scan RECORDS/Medical Summary – DO NOT PRODUCE/ for chronology .docx files.

        Returns:
            (chronologies, pages_map, non_subpoena_records) where:
            - chronologies: {sid: [{"date": str, "path": str}, ...]}
            - pages_map: {sid_or_filename: str} — page counts from RECORDS SUMMARIZED
            - non_subpoena_records: {filename: [{"date": str, "path": str}, ...]}
        """
        # Try both dash types (regular hyphen and en-dash)
        folder = _find_folder_ci(self.case_path, "RECORDS", "Medical Summary – DO NOT PRODUCE")
        if folder is None:
            folder = _find_folder_ci(self.case_path, "RECORDS", "Medical Summary - DO NOT PRODUCE")
        if folder is None:
            return {}, {}, {}

        chronologies = {}
        pages_map = {}
        non_subpoena_records = {}

        for root, _dirs, files in os.walk(folder):
            for fname in files:
                if not fname.lower().endswith('.docx'):
                    continue
                # Skip Word temp files
                if fname.startswith('~$'):
                    continue

                full_path = os.path.join(root, fname)

                # Extract date from filename
                date_match = DATE_PATTERN.search(fname)
                if not date_match:
                    self.warning.emit(f"No date found in chronology filename: {fname}")
                    continue

                chron_date = f"{date_match.group(1)}-{date_match.group(2)}-{date_match.group(3)}"

                # Read the RECORDS SUMMARIZED table
                entries = self._read_records_summarized(full_path, fname)

                for entry in entries:
                    sid = entry["sid"]
                    pages = entry["pages"]
                    filename = entry["filename"]

                    if sid:
                        # Subpoena-matched record
                        if sid not in chronologies:
                            chronologies[sid] = []
                        if not any(c["date"] == chron_date and c["path"] == full_path
                                  for c in chronologies[sid]):
                            chronologies[sid].append({"date": chron_date, "path": full_path})
                        if pages:
                            pages_map[sid] = pages
                    else:
                        # Non-subpoena record — use filename as key
                        if filename not in non_subpoena_records:
                            non_subpoena_records[filename] = []
                        if not any(c["date"] == chron_date and c["path"] == full_path
                                  for c in non_subpoena_records[filename]):
                            non_subpoena_records[filename].append(
                                {"date": chron_date, "path": full_path}
                            )
                        if pages:
                            pages_map[filename] = pages

        return chronologies, pages_map, non_subpoena_records

    def _read_records_summarized(self, docx_path, fname):
        """Read the RECORDS SUMMARIZED table from a medical chronology .docx.

        Iterates tables in reverse, finds the first 2-column table.

        Returns:
            List of dicts: [{"sid": str|None, "filename": str, "pages": str}, ...]
        """
        try:
            doc = Document(docx_path)
        except Exception as e:
            self.warning.emit(f"Could not open chronology {fname}: {e}")
            return []

        # Iterate tables in reverse — first 2-column table is RECORDS SUMMARIZED
        for table in reversed(doc.tables):
            if not table.rows:
                continue
            # Check column count from first row
            if len(table.rows[0].cells) != 2:
                continue

            entries = []
            for row in table.rows:
                cell_text = row.cells[0].text.strip().upper()
                if cell_text in ("RECORDS SUMMARIZED:", "FILENAME", ""):
                    continue
                if cell_text == "TOTAL":
                    break

                raw_filename = row.cells[0].text.strip()
                pages = row.cells[1].text.strip()

                # Skip section headers and non-file entries (no pages, no extension)
                if not pages and not raw_filename.lower().endswith(('.pdf', '.doc', '.docx', '.xlsx', '.msg')):
                    continue

                # Try to parse subpoena ID
                sid, _ = parse_subpoena_id(raw_filename)
                entries.append({
                    "sid": sid,
                    "filename": raw_filename,
                    "pages": pages,
                })

            return entries

        self.warning.emit(f"No 2-column RECORDS SUMMARIZED table found in: {fname}")
        return []

    # ---------------------------------------------------------------
    # Phase 4
    # ---------------------------------------------------------------
    def _generate_report(self, subpoenas, received, chronologies, pages_map,
                         non_subpoena_records, file_index):
        """Generate the Tracked_Subpoenas .docx report."""
        doc = Document()

        # Set default style
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        font.color.rgb = RGBColor(0, 0, 0)
        pf = style.paragraph_format
        pf.line_spacing = 1.0
        pf.space_after = Pt(0)
        pf.space_before = Pt(0)

        # Title
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_text = 'Subpoena Tracking Report'
        if self.file_number:
            title_text += f' - {self.file_number}'
        run = title_p.add_run(title_text)
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = 'Times New Roman'

        doc.add_paragraph("")
        today = datetime.date.today()
        doc.add_paragraph(f"Last Updated: {today.strftime('%B %d, %Y')}")
        doc.add_paragraph("")

        # Create table with fixed column widths — 5 columns now
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        table.autofit = False

        # Column widths: narrow for Subpoena#, Received, Summarized, Pages; Facility gets the rest
        # Page width ~6.5" (letter with 1" margins)
        col_widths = [
            Inches(1.2),   # Subpoena Number
            None,          # Facility (auto — takes remaining space)
            Inches(0.9),   # Records Received
            Inches(1.2),   # Records Summarized
            Inches(0.6),   # Pages
        ]
        facility_width = Inches(6.5 - 1.2 - 0.9 - 1.2 - 0.6)  # ~2.6"
        col_widths[1] = facility_width

        for col_idx, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[col_idx].width = width

        # Header row
        headers = ['Subpoena Number', 'Facility', 'Records Received', 'Records Summarized', 'Pages']
        for i, txt in enumerate(headers):
            cell = table.rows[0].cells[i]
            p = cell.paragraphs[0]
            run = p.add_run(txt)
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

        # Data rows — sorted by subpoena number
        for sid in sorted(subpoenas.keys()):
            data = subpoenas[sid]
            row_cells = table.add_row().cells

            # Apply column widths to new row
            for col_idx, width in enumerate(col_widths):
                row_cells[col_idx].width = width

            # Column 1: Subpoena Number
            run = row_cells[0].paragraphs[0].add_run(sid)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

            # Column 2: Facility
            run = row_cells[1].paragraphs[0].add_run(data["facility"])
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

            # Column 3: Records Received
            rec = received.get(sid)
            rec_p = row_cells[2].paragraphs[0]
            if rec is None or rec["status"] == "No":
                run = rec_p.add_run("No")
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
            else:
                label = rec["status"]  # "Yes", "CNR", or "Other"
                add_hyperlink(rec_p, rec["path"], label)

            # Column 4: Records Summarized — dates in MM-DD-YYYY, one per line
            chron_list = chronologies.get(sid, [])
            self._write_chron_dates(row_cells[3], chron_list)

            # Column 5: Pages — from med chron table, blank if not summarized
            if chron_list and sid in pages_map:
                run = row_cells[4].paragraphs[0].add_run(pages_map[sid])
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)

        # Append non-subpoena records from medical chronologies
        if non_subpoena_records:
            for filename in sorted(non_subpoena_records.keys()):
                chron_list = non_subpoena_records[filename]
                row_cells = table.add_row().cells

                for col_idx, width in enumerate(col_widths):
                    row_cells[col_idx].width = width

                # Column 1: Subpoena Number — blank
                # Column 2: Facility — plain text filename
                run = row_cells[1].paragraphs[0].add_run(filename)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)

                # Column 3: Records Received — "Yes" hyperlinked to file if found
                rec_p = row_cells[2].paragraphs[0]
                found_path = self._lookup_file(file_index, filename)
                if found_path:
                    add_hyperlink(rec_p, found_path, "Yes")
                else:
                    run = rec_p.add_run("Yes")
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                # Column 4: Records Summarized
                self._write_chron_dates(row_cells[3], chron_list)

                # Column 5: Pages
                if filename in pages_map:
                    run = row_cells[4].paragraphs[0].add_run(pages_map[filename])
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)

        # Append unmatched received records (different vendor IDs from copy services)
        unmatched_received = {sid: rec for sid, rec in received.items()
                              if sid not in subpoenas}
        # Build a set of non-subpoena record basenames for dedup
        non_sub_basenames = {os.path.splitext(k)[0].lower() for k in non_subpoena_records}
        if unmatched_received:
            for sid in sorted(unmatched_received.keys()):
                rec = unmatched_received[sid]
                # Skip if already listed as a non-subpoena record from med chron
                rec_basename = os.path.splitext(os.path.basename(rec["path"]))[0].lower()
                if rec_basename in non_sub_basenames:
                    continue
                row_cells = table.add_row().cells
                for col_idx, width in enumerate(col_widths):
                    row_cells[col_idx].width = width

                # Column 1: Reference number from copy service
                run = row_cells[0].paragraphs[0].add_run(sid)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(128, 128, 128)

                # Column 2: Facility — extract from filename after client name
                facility = self._extract_facility_from_received(
                    os.path.basename(rec["path"]))
                run = row_cells[1].paragraphs[0].add_run(facility)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)

                # Column 3: Records Received — hyperlinked
                rec_p = row_cells[2].paragraphs[0]
                label = rec["status"]
                add_hyperlink(rec_p, rec["path"], label)

                # Column 4-5: Check med chron by filename
                basename = os.path.splitext(os.path.basename(rec["path"]))[0]
                chron_list = chronologies.get(sid, [])
                self._write_chron_dates(row_cells[3], chron_list)
                if chron_list and sid in pages_map:
                    run = row_cells[4].paragraphs[0].add_run(pages_map[sid])
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                elif basename in pages_map:
                    run = row_cells[4].paragraphs[0].add_run(pages_map[basename])
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)

        # Save
        output_dir = os.path.join(self.case_path, "NOTES", "AI OUTPUT")
        os.makedirs(output_dir, exist_ok=True)
        filename = f"Tracked_Subpoenas {today.strftime('%Y-%m-%d')}.docx"
        output_path = os.path.join(output_dir, filename)
        doc.save(output_path)

        self.finished_result.emit(True, output_path)

    @staticmethod
    def _extract_facility_from_received(filename):
        """Extract facility name from received record filename.

        Handles patterns like:
          6900848-01_Deshawn Stampley_Healthpointe Medical Group, Inc Meds.pdf
          6900848-13 Bear Valley Community Hospital Meds CNR.pdf
        """
        basename = os.path.splitext(filename)[0]
        # Pattern 1: ID_Client Name_Facility
        parts = basename.split('_')
        if len(parts) >= 3:
            return '_'.join(parts[2:]).strip()
        # Pattern 2: ID Facility (after the first space-separated ID)
        _, remainder = parse_subpoena_id(basename)
        if remainder:
            # Strip client name if present (e.g., "Deshawn Stampley_...")
            return remainder.strip()
        return basename

    def _write_chron_dates(self, cell, chron_list):
        """Write chronology dates as hyperlinked entries, one per line."""
        if not chron_list:
            return
        chron_list_sorted = sorted(chron_list, key=lambda c: c["date"])
        for idx, chron in enumerate(chron_list_sorted):
            try:
                dt = datetime.datetime.strptime(chron["date"], "%Y-%m-%d")
                display_date = dt.strftime("%m-%d-%Y")
            except ValueError:
                display_date = chron["date"]

            if idx == 0:
                p = cell.paragraphs[0]
            else:
                p = cell.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0
            add_hyperlink(p, chron["path"], display_date)
