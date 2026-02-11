"""
Testimony Formatter — Word document output and PDF highlighting.

Produces a Word document with verbatim Q/A pairs formatted per legal standards,
and optionally creates a highlighted PDF copy of the transcript.
"""

import os
import re
import logging
from typing import List, Optional, Tuple

from docx import Document
from docx.shared import Inches, Pt, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH

import fitz

from .models import TranscriptIndex, QAExchange, ExtractionResult

logger = logging.getLogger(__name__)

# Formatting constants
LEFT_INDENT = Inches(0.5)       # 0.5" indent for all Q/A text
HANGING_INDENT = Inches(0.5)    # 0.5" hanging indent (Q./A. marker hangs left)
FONT_NAME = "Times New Roman"
FONT_SIZE = Pt(12)
LINE_SPACING = 1.0


class TestimonyFormatter:
    """
    Formats extracted testimony into Word documents and optionally highlights PDFs.

    Usage:
        formatter = TestimonyFormatter()
        formatter.add_extraction(index, result)
        formatter.add_extraction(index, result2)
        output_path = formatter.save_word(output_dir)
    """

    def __init__(self):
        self._extractions: List[Tuple[TranscriptIndex, ExtractionResult]] = []

    def add_extraction(self, index: TranscriptIndex, result: ExtractionResult):
        """Add an extraction result to be included in the output document."""
        self._extractions.append((index, result))

    def clear(self):
        """Clear all accumulated extractions."""
        self._extractions.clear()

    # =========================================================================
    # Word Document Output
    # =========================================================================

    def save_word(self, output_dir: str = None, filename: str = None) -> str:
        """
        Save all extractions to a Word document.

        Args:
            output_dir: Directory for output. Defaults to transcript's directory.
            filename: Custom filename. Defaults to "[Extracted] [Deponent] Depo Trns.docx".

        Returns:
            Path to the saved document.
        """
        if not self._extractions:
            raise ValueError("No extractions to save")

        doc = Document()

        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = FONT_NAME
        font.size = FONT_SIZE

        for ext_idx, (index, result) in enumerate(self._extractions):
            if ext_idx > 0:
                # Add separator between different extraction prompts
                doc.add_paragraph()

            # Add prompt as header
            prompt_para = doc.add_paragraph()
            prompt_run = prompt_para.add_run(f"Extraction: {result.prompt}")
            prompt_run.bold = True
            prompt_run.font.size = FONT_SIZE
            prompt_run.font.name = FONT_NAME
            prompt_para.space_after = Pt(6)

            # Add each group of consecutive exchanges
            for group_idx, group_ids in enumerate(result.groups):
                if group_idx > 0:
                    # Small gap between citation blocks
                    doc.add_paragraph().space_after = Pt(6)

                # Get exchanges for this group
                exchange_map = {ex.id: ex for ex in index.exchanges}
                group_exchanges = [exchange_map[eid] for eid in group_ids if eid in exchange_map]

                if not group_exchanges:
                    continue

                # Add each Q/A pair
                for ex in group_exchanges:
                    self._add_qa_paragraph(doc, "Q.", ex.question)
                    self._add_qa_paragraph(doc, "A.", ex.answer)

                # Add citation after the group
                citation = self._format_citation(group_exchanges, index.deponent.last_name)
                cite_para = doc.add_paragraph()
                cite_para.paragraph_format.left_indent = LEFT_INDENT
                cite_para.space_before = Pt(6)
                cite_para.space_after = Pt(12)
                cite_run = cite_para.add_run(citation)
                cite_run.font.size = FONT_SIZE
                cite_run.font.name = FONT_NAME

        # Determine output path
        first_index = self._extractions[0][0]
        if output_dir is None:
            raise ValueError("output_dir is required — never save to transcript folder")

        if filename is None:
            deponent = first_index.deponent.last_name or "Unknown"
            filename = f"[Extracted] {deponent} Depo Trns.docx"

        output_path = os.path.join(output_dir, filename)
        doc.save(output_path)
        logger.info(f"Saved extraction document: {output_path}")
        return output_path

    def _add_qa_paragraph(self, doc: Document, marker: str, text: str):
        """Add a Q. or A. paragraph with proper formatting."""
        para = doc.add_paragraph()
        pf = para.paragraph_format

        # Left indent 0.5" with hanging indent for marker
        pf.left_indent = Inches(1.0)       # Text body at 1.0"
        pf.first_line_indent = Inches(-0.5)  # Marker hangs back to 0.5"
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)

        # Add marker + tab + text
        marker_run = para.add_run(f"{marker}\t")
        marker_run.font.name = FONT_NAME
        marker_run.font.size = FONT_SIZE

        text_run = para.add_run(text)
        text_run.font.name = FONT_NAME
        text_run.font.size = FONT_SIZE

        # Set tab stop at 1.0" from margin (= 1440 twips)
        self._set_tab_stop(para, 1440)

    def _set_tab_stop(self, paragraph, position_twips: int):
        """Set a tab stop at the given position (in twips)."""
        from docx.oxml.ns import qn
        from lxml import etree

        pPr = paragraph._p.get_or_add_pPr()
        tabs = pPr.find(qn('w:tabs'))
        if tabs is None:
            tabs = etree.SubElement(pPr, qn('w:tabs'))

        tab = etree.SubElement(tabs, qn('w:tab'))
        tab.set(qn('w:val'), 'left')
        tab.set(qn('w:pos'), str(position_twips))

    @staticmethod
    def _merge_citation_range(exchanges: List[QAExchange]) -> str:
        """
        Merge a group of consecutive exchanges into a single page:line range.

        Instead of "18:3-11; 18:12-20:2; 20:3-10", produces "18:3-20:10".
        """
        if not exchanges:
            return ""
        first = exchanges[0]
        last = exchanges[-1]
        if first.page_start == last.page_end:
            return f"{first.page_start}:{first.line_start}-{last.line_end}"
        return f"{first.page_start}:{first.line_start}-{last.page_end}:{last.line_end}"

    def _format_citation(self, exchanges: List[QAExchange], deponent_last_name: str) -> str:
        """
        Format a citation block for a group of consecutive exchanges.

        Format: (Exh. __, ([LastName] Depo. Trns.) at p. [range].)
        """
        range_str = self._merge_citation_range(exchanges)
        raw_name = deponent_last_name or "___"
        last_name = raw_name.title() if raw_name.isupper() else raw_name
        return f"(Exh. __ ({last_name} Depo. Trns.) at p. {range_str}.)"

    # =========================================================================
    # PDF Highlighting
    # =========================================================================

    def highlight_pdf(
        self, index: TranscriptIndex, result: ExtractionResult,
        output_dir: str = None, on_progress=None, color: str = None
    ) -> Optional[str]:
        """
        Create or update a highlighted PDF copy of the transcript.

        The highlighted copy is named "[H.AI] original_filename.pdf" and
        accumulates highlights across multiple extraction calls.

        Args:
            index: The transcript index.
            result: The extraction result with selected IDs.
            output_dir: Directory for output. Defaults to transcript's directory.

        Returns:
            Path to the highlighted PDF, or None on failure.
        """
        source_path = index.source_pdf
        source_name = os.path.basename(source_path)

        # Determine highlighted copy path
        if output_dir is None:
            raise ValueError("output_dir is required — never save to transcript folder")
        os.makedirs(output_dir, exist_ok=True)
        highlight_name = f"[H.AI] {source_name}"
        highlight_path = os.path.join(output_dir, highlight_name)

        # Open existing highlighted copy or create from original
        if os.path.exists(highlight_path):
            doc = fitz.open(highlight_path)
            logger.info(f"Updating existing highlighted copy: {highlight_path}")
        else:
            doc = fitz.open(source_path)
            logger.info(f"Creating new highlighted copy from: {source_path}")

        # Get selected exchanges
        exchange_map = {ex.id: ex for ex in index.exchanges}
        selected = [exchange_map[eid] for eid in result.selected_ids if eid in exchange_map]

        # Parse highlight color
        if color:
            r = int(color[1:3], 16) / 255
            g = int(color[3:5], 16) / 255
            b = int(color[5:7], 16) / 255
            stroke_color = (r, g, b)
        else:
            stroke_color = (1.0, 0.8, 0.0)

        # Build a mapping from transcript page → PDF page(s)
        page_map = self._build_page_map(doc, index.is_condensed)
        highlights_added = 0

        total_selected = len(selected)
        for ex_idx, ex in enumerate(selected):
            if on_progress:
                on_progress(ex_idx + 1, total_selected)

            pdf_pages = self._get_pdf_pages_for_transcript(
                ex.page_start, ex.page_end, page_map
            )

            for pdf_page_num in pdf_pages:
                if pdf_page_num >= len(doc):
                    continue

                page = doc[pdf_page_num]
                rects = self._get_line_range_rects(
                    page, ex.page_start, ex.page_end,
                    ex.line_start, ex.line_end, page_map
                )
                for rect in rects:
                    annot = page.add_highlight_annot(rect)
                    annot.set_colors(stroke=stroke_color)
                    annot.set_info(content=f"ex:{ex.id}")
                    annot.update()
                    highlights_added += 1

        # Save — must use temp file when overwriting the same file we opened
        import tempfile
        if os.path.exists(highlight_path) and os.path.samefile(
            doc.name, highlight_path
        ):
            # Opened from highlight_path → save incremental or via temp
            tmp_fd, tmp_path = tempfile.mkstemp(suffix=".pdf", dir=output_dir)
            os.close(tmp_fd)
            doc.save(tmp_path, garbage=3)
            doc.close()
            os.replace(tmp_path, highlight_path)
        else:
            doc.save(highlight_path)
            doc.close()

        logger.info(f"Added {highlights_added} highlights to {highlight_path}")
        return highlight_path

    def _build_page_map(self, doc, is_condensed: bool) -> dict:
        """
        Build a mapping from transcript page numbers to PDF page indices.

        For full-size: transcript page N ≈ PDF page N-1 (approximately)
        For condensed: transcript pages are packed 4-per-PDF-page
        """
        page_map = {}  # transcript_page -> [pdf_page_indices]

        for pdf_idx in range(len(doc)):
            page = doc[pdf_idx]
            text = page.get_text("text")

            # Skip word index pages that appear after the transcript.
            # These have dense "word line:col" entries (e.g., "accident 27:8")
            # but no Q./A. testimony markers. Detect by structure, not by
            # court reporter name (Veritext, US Legal, etc. all vary).
            has_qa = '\n       Q.' in text or '\n       A.' in text
            if not has_qa:
                # Count word index reference patterns (e.g., "27:8")
                # Exclude timestamp patterns like "10:32:13AM" by
                # stripping all timestamps first, then counting.
                cleaned = re.sub(r'\d{1,2}:\d{2}:\d{2}\s*[AP]M', '', text)
                word_refs = re.findall(r'\b\d{1,3}:\d{1,2}\b', cleaned)
                if len(word_refs) >= 10:
                    continue

            # Find all "Page N" markers on this PDF page
            for match in re.finditer(r'Page\s+(\d+)', text):
                tp_num = int(match.group(1))
                if tp_num not in page_map:
                    page_map[tp_num] = []
                if pdf_idx not in page_map[tp_num]:
                    page_map[tp_num].append(pdf_idx)

        return page_map

    def _get_pdf_pages_for_transcript(
        self, page_start: int, page_end: int, page_map: dict
    ) -> List[int]:
        """Get PDF page indices that contain the given transcript page range."""
        pdf_pages = set()
        for tp in range(page_start, page_end + 1):
            if tp in page_map:
                pdf_pages.update(page_map[tp])
            else:
                # Approximate: transcript page N is roughly PDF page N-1
                # (pages 1-4 are cover/index, so offset by ~4)
                approx = tp - 1
                if 0 <= approx:
                    pdf_pages.add(approx)
        return sorted(pdf_pages)

    @staticmethod
    def _find_line_info(page) -> dict:
        """Find transcript line positions and text extents on a PDF page.

        Returns dict mapping line_number (1-25) to:
            {"y0": float, "y1": float, "text_x0": float, "text_x1": float}
        where y0/y1 are the vertical bounds and text_x0/text_x1 are the
        horizontal extent of all text content on that line (excluding the
        line number itself).
        """
        # First pass: find line number positions and their y-coordinates
        line_num_info = {}  # line_num -> {"y0": float, "num_x1": float}
        text_dict = page.get_text("dict")
        for block in text_dict.get("blocks", []):
            if block.get("type") != 0:
                continue
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    text = span["text"].strip()
                    bbox = span["bbox"]
                    if bbox[0] < 100 and text.isdigit():
                        num = int(text)
                        if 1 <= num <= 25:
                            line_num_info[num] = {
                                "y0": bbox[1],
                                "y1": bbox[3],
                                "num_x1": bbox[2],
                            }

        if not line_num_info:
            return {}

        # Second pass: find text extents for each line by matching y-coordinates
        # A text span belongs to a line if its vertical center is within that line's y range
        line_info = {}
        for num, info in line_num_info.items():
            line_info[num] = {
                "y0": info["y0"],
                "y1": info["y1"],
                "text_x0": None,
                "text_x1": None,
            }

        for block in text_dict.get("blocks", []):
            if block.get("type") != 0:
                continue
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    bbox = span["bbox"]
                    span_y_center = (bbox[1] + bbox[3]) / 2
                    # Match this span to the closest line number by y
                    for num, info in line_num_info.items():
                        if abs(span_y_center - (info["y0"] + info["y1"]) / 2) < 8:
                            # Skip the line number span itself
                            if bbox[0] < 100 and span["text"].strip().isdigit():
                                continue
                            li = line_info[num]
                            if li["text_x0"] is None or bbox[0] < li["text_x0"]:
                                li["text_x0"] = bbox[0]
                            if li["text_x1"] is None or bbox[2] > li["text_x1"]:
                                li["text_x1"] = bbox[2]
                            break

        # Fill in defaults for lines with no text content
        for num, li in line_info.items():
            if li["text_x0"] is None:
                li["text_x0"] = line_num_info[num]["num_x1"] + 10
            if li["text_x1"] is None:
                li["text_x1"] = page.rect.width - 20

        return line_info

    @staticmethod
    def _get_line_range_rects(
        page, page_start: int, page_end: int,
        line_start: int, line_end: int, page_map: dict
    ) -> list:
        """
        Get per-line highlight rectangles for a transcript line range.

        Returns one fitz.Rect per transcript line, matching the style of
        manual highlights in Adobe Acrobat (individual line highlights
        rather than one giant block).

        For multi-page exchanges, determines whether this PDF page is the
        first, last, or a middle page and adjusts the line range:
        - First page: line_start to line 25
        - Middle page: line 1 to line 25
        - Last page: line 1 to line_end
        - Single page: line_start to line_end
        """
        line_info = TestimonyFormatter._find_line_info(page)
        if not line_info:
            return []

        # Determine which transcript page this PDF page contains
        pdf_idx = page.number
        tp_on_this_page = None
        for tp, pdf_pages in page_map.items():
            if pdf_idx in pdf_pages:
                tp_on_this_page = tp
                break

        # Determine the line range to highlight on this page
        is_first_page = (tp_on_this_page is not None and tp_on_this_page == page_start)
        is_last_page = (tp_on_this_page is not None and tp_on_this_page == page_end)
        single_page = (page_start == page_end)

        all_lines = sorted(line_info.keys())
        if not all_lines:
            return []

        if single_page or (is_first_page and is_last_page):
            start, end = line_start, line_end
        elif is_first_page:
            start, end = line_start, all_lines[-1]
        elif is_last_page:
            start, end = all_lines[0], line_end
        else:
            start, end = all_lines[0], all_lines[-1]

        # Build one rect per line
        rects = []
        for line_num in range(start, end + 1):
            if line_num not in line_info:
                continue
            li = line_info[line_num]
            rect = fitz.Rect(
                li["text_x0"] - 2,
                li["y0"] - 1,
                li["text_x1"] + 2,
                li["y1"] + 1,
            )
            rects.append(rect)

        return rects

